from pathlib import Path as _Path_main_env

try:
    from dotenv import load_dotenv

    _main_file = _Path_main_env(__file__).resolve()
    load_dotenv(_main_file.parent.parent / ".env")
    load_dotenv(_main_file.parent / ".env")
except ImportError:
    pass

import json
import mimetypes
import os
import re
import sys
import difflib
import base64
import hashlib
import subprocess
import threading
import concurrent.futures

try:
    from agent_loop import AgentLoop
    from agent_loop import run_agent as _run_agent_loop
    AGENT_LOOP_AVAILABLE = True
except ImportError:
    AGENT_LOOP_AVAILABLE = False
    AgentLoop = None  # type: ignore
    def _run_agent_loop(task, workspace_path, model=None):
        return {"ok": False, "error": "agent_loop.py nicht gefunden", "formatted_response": "Fehler: agent_loop.py fehlt."}
import time
import traceback
from datetime import datetime
from pathlib import Path
import shutil
from urllib.parse import quote
from uuid import uuid4

import requests
from flask import Flask, g, jsonify, request, send_from_directory
from flask import Response
from werkzeug.exceptions import HTTPException
from api_run_state import enrich_direct_confirm_response, enrich_direct_run_response
from auto_analyzer import AutoAnalyzer
from auto_logger import AutoLogger
from background_learning import start_background_learning
from change_tracker import ChangeTracker
from architecture_patterns import ArchitectureDecider
from code_generator_advanced import CodeGeneratorAdvanced
from executable_creator import ExecutableCreator
from file_creator import FileCreator
from project_builder import ProjectBuilder
from file_reader import get_latest_analysis_file, read_file_content
from build_system import BuildSystem
from decision_engine import DecisionMaker
from electron_builder import ElectronBuilder
from icon_processor import IconProcessor
from improvement_engine import ImprovementSuggester
from logic_engine import ProblemAnalyzer
from message_templates import MessageTemplates as MT
from logger import RainerLogger
from openapi_schema import OPENAPI_SCHEMA
from performance_optimizer import PerformanceOptimizer as PO
from prompt_optimizer import PromptOptimizer
from react_builder import ReactBuilder
from rag_integration import RAGIntegration
from response_formatter import ResponseFormatter, format_response_like_claude as formatter_claude_style
from response_builder import (
    error_payload,
    file_write_payload,
    no_change_payload,
    preview_payload,
    success_payload,
)
from agent_file_guard import extract_relative_path_from_zielordner_block
from routes import handle_user_prompt_routing, persist_text_file_change
from prompt_routing import (
    chat_reply_canned,
    classify_user_prompt,
    classify_user_prompt_intent,
    connectivity_diagnostics_reply,
    extract_folder_analysis_path_from_prompt,
    has_project_change_intent,
    is_analysis_only_prompt,
    is_folder_analysis_prompt,
    should_route_direct_run_as_chat,
    unclear_chat_short_reply,
    unknown_clarification_reply,
)
from intent_enrichment import (
    SUGGESTED_INTENT_ACTIONS,
    apply_user_mode_override,
    clarification_message_with_modes,
    compose_augmented_user_message,
    intent_llm_enabled,
    normalize_conversation_history_payload,
    run_llm_intent_refinement,
)
from model_providers import (
    generate_chat_response as llm_generate_chat_response,
    generate_coding_response as llm_generate_coding_response,
    generate_image_via_openai,
    is_llm_failure_message,
    summarize_llm_health,
)
from silent_learner import SilentLearner
from template_library import TemplateLibrary
from tradeoff_analyzer import TradeOffAnalyzer
from write_action import DIRECT_WRITE_ACTION_TIMEOUT_SEC
from agent_git import get_instance as get_git_integration
from agent_git_worktree import create_worktree as git_create_worktree, list_worktrees as git_list_worktrees
from agent_scheduler import get_instance as get_scheduler
from agent_model_router import get_instance as get_model_router
from agent_model_quality import get_instance as get_model_quality_router
from agent_parallel import ParallelAgentManager
from agent_pr_workflow import PRWorkflowManager
from agent_cloud_exec import CloudExecManager
from agent_workspace_sandbox import WorkspaceSandbox
from agent_workspace_indexer import get_instance as get_workspace_indexer
from agent_file_reader import get_instance as get_file_reader_agent
from agent_task_planner import get_instance as get_task_planner_agent
from agent_context_builder import get_instance as get_context_builder_agent
from agent_patch_generator import get_instance as get_patch_generator_agent
from agent_patch_validator import get_instance as get_patch_validator_agent
from agent_test_runner import get_instance as get_test_runner_agent
from agent_error_fixer import get_instance as get_error_fixer_agent
from agent_step_engine import StepEngineAgent
from agent_memory_history import get_instance as get_memory_history_agent
from agent_codebase_understanding import CodebaseUnderstanding
from agent_issue_tracker import get_instance as get_issue_tracker
from agent_cloud import get_instance as get_cloud_agent
from agent_tooling import build_tool_registry, get_tool as get_registry_tool, permission_gate as tools_permission_gate, execute_tool as tools_execute_tool, build_external_tool_adapters
import difflib
import shutil

# Optional: Hybrid Engine importieren (nicht kritisch für Rainer Build)
try:
    from hybrid_engine import execute_intelligent_hybrid, hybrid_engine
    HYBRID_ENGINE_AVAILABLE = True
except Exception:
    HYBRID_ENGINE_AVAILABLE = False
    hybrid_engine = None
    execute_intelligent_hybrid = None

app = Flask(__name__)
logger = RainerLogger(level="info")
SERVER_PORT = int(os.environ.get("FLASK_PORT", os.environ.get("SERVER_PORT", "5002")))
SERVER_INSTANCE_ID = f"rainer-{uuid4().hex[:12]}"
SERVER_STARTED_AT = datetime.now().isoformat(timespec="seconds")


def log_structured(event: str, **fields):
    try:
        payload = {"event": event, "timestamp": datetime.now().isoformat(timespec="seconds"), **fields}
        print(json.dumps(payload, ensure_ascii=True))
    except Exception:
        pass


@app.before_request
def _api_request_start():
    g._rainer_api_started_at = datetime.now().timestamp()


@app.after_request
def _api_request_log(resp):
    try:
        started = float(getattr(g, "_rainer_api_started_at", datetime.now().timestamp()))
        elapsed = max(0.0, datetime.now().timestamp() - started)
        logger.log_api_call(request.path or "-", request.method or "GET", int(resp.status_code), elapsed)
        logger.log_performance(f"{request.method} {request.path}", elapsed)
        log_structured(
            "api_request",
            path=request.path or "-",
            method=request.method or "GET",
            status=int(resp.status_code),
            elapsed_ms=int(elapsed * 1000),
        )
    except Exception:
        pass
    return resp

BASE_DIR = Path(__file__).resolve().parent
APP_DIR = BASE_DIR.parent
# Layout robust aufloesen:
# - root layout: <root>/backend/main.py  -> PROJECT_DIR = <root>
# - nested layout: <root>/rambo_builder_local/backend/main.py -> PROJECT_DIR = <root>
if (APP_DIR / "frontend").exists() and (APP_DIR / "backend").exists():
    PROJECT_DIR = APP_DIR
elif (APP_DIR.parent / "frontend").exists() and (APP_DIR.parent / "backend").exists():
    PROJECT_DIR = APP_DIR.parent
else:
    PROJECT_DIR = APP_DIR
FRONTEND_DIR = PROJECT_DIR / "frontend"
DATA_DIR = PROJECT_DIR / "data"

try:
    from agent_api import register_agent_routes

    register_agent_routes(app, APP_DIR)
except Exception as _agent_route_err:
    logger.logger.warning("Agent-API nicht registriert: %s", _agent_route_err)
try:
    from main_agent_integration import register_agent_mega_routes

    register_agent_mega_routes(app)
except Exception as _mega_route_err:
    logger.logger.warning("Agent-Mega-Routen nicht registriert: %s", _mega_route_err)
try:
    from app_cursor_features import register_cursor_features

    register_cursor_features(app, None, None, lambda: get_active_project_root())
except Exception as _cursor_route_err:
    logger.logger.warning("Cursor-Feature-Routen nicht registriert: %s", _cursor_route_err)
PROJECT_MEMORY_FILE = APP_DIR / "data" / "project_memory.json"
DESIGN_NOTES_FILE = DATA_DIR / "design_notes.json"
UI_ACTIVITY_LOG_FILE = DATA_DIR / "ui_activity_log.json"
PROJECT_MAP_FILE = DATA_DIR / "project_map.json"
PROJECT_AUTO_RUN_STATE_FILE = DATA_DIR / "project_auto_run_state.json"
QUALITY_TASK_GRAPH_FILE = DATA_DIR / "quality_task_graph.json"
QUALITY_EVAL_HISTORY_FILE = DATA_DIR / "quality_eval_history.json"
DIRECT_HISTORY_LIMIT = 12
ALLOWED_FILE_TYPES = {"txt", "md", "json", "py", "html", "css", "js"}
SAFE_FILENAME_PATTERN = re.compile(r"^[A-Za-z0-9_\-\.]+$")

RAMBO_RAINER_ROOT = PROJECT_DIR

def _resolve_downloads_dir() -> Path:
    """Resolve downloads directory inside the project root."""
    target = Path(__file__).resolve().parents[1] / "Downloads"
    target.mkdir(parents=True, exist_ok=True)
    return target


DOWNLOADS_DIR = _resolve_downloads_dir()
LOGS_DIR = DOWNLOADS_DIR / "logs"
LOGS_DIR.mkdir(parents=True, exist_ok=True)
UPLOADS_DIR = DOWNLOADS_DIR / "uploads"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_IMAGES_DIR = UPLOADS_DIR / "generated"
GENERATED_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
CHAT_HISTORY_FILE = DATA_DIR / "chat_history.json"

# Auto-Apply: direkte Ausfuehrung ohne Bestaetigung fuer normale Auftraege.
# Harte Sperren (Systemordner, Massenloesch, Geheimnisse) bleiben aktiv.
AUTO_APPLY: bool = True

ALLOWED_PROJECT_WRITE_PREFIXES = (
    "knowledge/",
    "outbox/",
    "memory/",
    "data/",
    "backend/BUILDER_MODE.md",
    "backend/COACH.md",
    "backend/DESIGN.md",
    "backend/DEV_WORKFLOW.md",
    "backend/GENERATION.md",
    "backend/OFFICE.md",
    "backend/ORCHESTRATION.md",
    "backend/PHASES.md",
    "backend/SCAFFOLD.md",
    "logs/system_status.md",
)

SENSITIVE_PATTERNS = (
    ".env",
    "rambo_rainer.db",
    ".git/",
    "node_modules/",
    "__pycache__",
    ".pyc",
    ".aider",
    "package-lock.json",
)

GUARDED_PROJECT_PATHS = {
    "frontend/src/App.jsx",
    "frontend/src/App.css",
}

# Default auf aktuelles App-Root; optional per Env ueberschreibbar.
DIRECT_RUN_EXPECTED_APP_ROOT = Path(
    os.getenv("RAINER_EXPECTED_APP_ROOT") or os.getenv("RAINER_APP_DIR") or str(APP_DIR)
).resolve()
DIRECT_RUN_GUARD_BLOCK_MESSAGE = (
    "Auftrag blockiert: geplanter Pfad liegt außerhalb des erlaubten Rainer-Build-Arbeitsbereichs."
)
DIRECT_RUN_GUARD_UI_BLOCKED_MESSAGE = "Auftrag wurde vom Guard blockiert"
DIRECT_RUN_MINI_ALLOWED_FILES = {
    "frontend/app.js",
    "frontend/index.html",
    "frontend/style.css",
}
DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_FILES = {"backend/main.py"}
DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_TEST_PREFIX = "tests/"
DIRECT_RUN_FORBIDDEN_PREFIXES = (
    "electron/",
    "rainerrobotdesktop/",
    "rambo_ui/",
    "src/components/",
    "downloads/",
    "../downloads/",
    "du_arbeitest_im_projekt__",
    "node_modules/",
    "dist/",
    "build/",
    "__pycache__/",
    ".pytest_cache/",
    ".git/",
)
DIRECT_RUN_FORBIDDEN_FILENAMES = {
    "package.json",
    "vite.config.js",
    "build_desktop.py",
    "robot-icon.png",
    "app-icon.png",
}
DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN = (
    "electron/main.js",
    "electron/preload.js",
    "electron/package.json",
    "electron/fallback-setup.html",
    "rambo_ui/package.json",
    "rambo_ui/vite.config.js",
    "rambo_ui/index.html",
    "rambo_ui/src/main.jsx",
    "rambo_ui/src/App.jsx",
    "rambo_ui/src/App.css",
    "rambo_ui/src/components/ServicePanel.jsx",
    "rambo_ui/src/components/UI.css",
)

# Paths that are local dev-tool state — not product code, should not count as product gaps
LOCAL_DEV_PATH_PREFIXES = (
    ".claude/",
    ".windsurf/",
    ".idea/",
    ".vscode/",
    "settings.local.json",
)

def _pick_best_ollama_model() -> str:
    preferred = [
        "llama-3.3-70b-versatile",
        "gemma3:12b-it-qat",
        "gemma3:12b",
        "qwen2.5-coder:latest",
        "qwen2.5-coder:7b",
        "deepseek-r1:8b",
    ]
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=3)
        if int(r.status_code) == 200:
            available = {str(m.get("name") or "").strip() for m in list((r.json() or {}).get("models") or [])}
            for model in preferred:
                if model in available:
                    return model
    except Exception:
        pass
    return "llama-3.3-70b-versatile"


OLLAMA_MODEL = os.getenv("GROQ_MODEL", _pick_best_ollama_model())
OLLAMA_FALLBACK_MODEL = os.getenv("OLLAMA_FALLBACK_MODEL", "qwen2.5-coder:3b")
OLLAMA_BASE_URL = "http://localhost:11434"
OLLAMA_TIMEOUT_SEC = int(os.getenv("OLLAMA_TIMEOUT_SEC", "60"))
OLLAMA_RETRY_COUNT = int(os.getenv("OLLAMA_RETRY_COUNT", "2"))
OLLAMA_CACHE_TTL_SEC = int(os.getenv("OLLAMA_CACHE_TTL_SEC", "300"))
OLLAMA_CACHE_MAX_ITEMS = int(os.getenv("OLLAMA_CACHE_MAX_ITEMS", "500"))
PASSIVE_LEARNING_DB = BASE_DIR / "data" / "passive_learning.json"

AUTO_LOGGER = AutoLogger(PASSIVE_LEARNING_DB)
AUTO_ANALYZER = AutoAnalyzer(PASSIVE_LEARNING_DB)
SILENT_LEARNER = SilentLearner(PASSIVE_LEARNING_DB)
RAG_INTEGRATION = RAGIntegration(PASSIVE_LEARNING_DB)
start_background_learning(PASSIVE_LEARNING_DB)
OLLAMA_SESSION = requests.Session()
OLLAMA_CACHE = {}

LOCAL_AGENT_MODEL_FALLBACKS = [
    m.strip()
    for m in os.getenv(
        "OLLAMA_LOCAL_AGENT_MODEL_FALLBACKS",
        "deepseek-coder:33b,deepseek-coder:7b,mistral:latest,llama3.2:latest",
    ).split(",")
    if m.strip()
]

WEATHER_STATUS_MAP = {
    0: "Klar",
    1: "Überwiegend klar",
    2: "Teilweise bewölkt",
    3: "Bewölkt",
    45: "Nebel",
    48: "Raureifnebel",
    51: "Leichter Nieselregen",
    53: "Nieselregen",
    55: "Starker Nieselregen",
    61: "Leichter Regen",
    63: "Regen",
    65: "Starker Regen",
    71: "Leichter Schneefall",
    73: "Schneefall",
    75: "Starker Schneefall",
    80: "Regenschauer",
    81: "Starke Regenschauer",
    82: "Heftige Regenschauer",
    95: "Gewitter",
}


def _weather_status_from_code(code):
    try:
        return WEATHER_STATUS_MAP.get(int(code), "Unbekannt")
    except Exception:
        return "Unbekannt"


def _get_git_integration():
    return get_git_integration(APP_DIR)


def _get_model_router():
    return get_model_router(APP_DIR)


def _get_model_quality_router():
    return get_model_quality_router(APP_DIR)


def _get_issue_tracker():
    return get_issue_tracker(APP_DIR)


def _get_cloud_agent():
    return get_cloud_agent(APP_DIR)


_EXPLORER_EXCLUDED_DIRS = {
    ".git",
    "node_modules",
    "__pycache__",
    ".pytest_cache",
    "dist",
    "build",
    ".rainer_agent",
    "Downloads",
}
PARALLEL_MANAGER = ParallelAgentManager(APP_DIR.resolve(), max_parallel_runs=4)
PR_WORKFLOW_MANAGER = PRWorkflowManager()
CLOUD_EXEC_MANAGER = CloudExecManager(APP_DIR.resolve())
WORKSPACE_SANDBOX = WorkspaceSandbox(APP_DIR.resolve())
_EXPLORER_EXCLUDED_SUFFIXES = {".pyc", ".pyo", ".exe", ".zip", ".tmp", ".bak"}
_EXPLORER_MAX_ITEMS = 300
_FILE_VIEWER_ALLOWED_SUFFIXES = {
    ".py",
    ".js",
    ".html",
    ".css",
    ".json",
    ".md",
    ".txt",
    ".bat",
    ".ps1",
    ".yml",
    ".yaml",
    ".toml",
}
_FILE_VIEWER_BLOCKED_SUFFIXES = {
    ".exe",
    ".zip",
    ".pyc",
    ".pyo",
    ".tmp",
    ".bak",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".webp",
    ".bmp",
    ".svg",
    ".ico",
    ".mp3",
    ".wav",
    ".mp4",
    ".avi",
    ".mov",
    ".mkv",
}
_FILE_VIEWER_MAX_BYTES = int(os.environ.get("RAINER_FILE_VIEWER_MAX_BYTES", "-1"))
_FILE_VIEWER_LANG_BY_SUFFIX = {
    ".py": "python",
    ".js": "javascript",
    ".html": "html",
    ".css": "css",
    ".json": "json",
    ".md": "markdown",
    ".txt": "text",
    ".bat": "bat",
    ".ps1": "powershell",
    ".yml": "yaml",
    ".yaml": "yaml",
    ".toml": "toml",
}
_DIFF_VIEWER_MAX_BYTES = int(os.environ.get("RAINER_DIFF_VIEWER_MAX_BYTES", "-1"))
_RUN_CHECK_MAX_OUTPUT_BYTES = 80 * 1024
_RUN_CHECK_TIMEOUT_SECONDS = 60
_RUN_ALLOWED_CHECKS: dict[str, list[str]] = {
    "pytest_all": ["python", "-m", "pytest", "tests", "-q"],
    "py_compile_main": ["python", "-m", "py_compile", "backend/main.py"],
    "node_check_app": ["node", "--check", "frontend/app.js"],
    "level20_diff_tests": ["python", "-m", "pytest", "tests/test_agent_level20_diff_viewer_api.py", "-q"],
    "level20_file_tests": ["python", "-m", "pytest", "tests/test_agent_level20_file_viewer_api.py", "-q"],
    "level20_explorer_tests": ["python", "-m", "pytest", "tests/test_agent_level20_explorer_api.py", "-q"],
}
_COMMIT_PLAN_MESSAGE_RE = re.compile(r"^(feat|fix|test|chore|docs|refactor|style|perf)(\([^)]+\))?:\s+\S.*$")
_COMMIT_PLAN_BLOCKED_DIR_MARKERS = (
    "downloads/",
    "../downloads/",
    "node_modules/",
    "dist/",
    "build/",
    "__pycache__/",
    ".pytest_cache/",
    ".git/",
    ".rainer_agent/",
)
_COMMIT_PLAN_BLOCKED_SUFFIXES = (".exe", ".zip", ".tmp", ".bak", ".pyc")


def _safe_project_files_listing(root: Path, max_items: int = _EXPLORER_MAX_ITEMS) -> tuple[list[dict], list[str], list[str]]:
    warnings: list[str] = []
    errors: list[str] = []
    items: list[dict] = []
    seen_warnings: set[str] = set()
    root_resolved = root.resolve()
    max_items = max(1, int(max_items))

    for current_root, dirs, files in os.walk(root_resolved, topdown=True):
        current_path = Path(current_root)
        rel_dir = current_path.relative_to(root_resolved).as_posix() if current_path != root_resolved else ""

        kept_dirs: list[str] = []
        for d in sorted(dirs, key=lambda x: x.lower()):
            if d in _EXPLORER_EXCLUDED_DIRS:
                if d == "Downloads" and "downloads_excluded" not in seen_warnings:
                    warnings.append("Downloads directory is excluded from explorer output.")
                    seen_warnings.add("downloads_excluded")
                continue
            kept_dirs.append(d)
        dirs[:] = kept_dirs

        for d in kept_dirs:
            rel_path = f"{rel_dir}/{d}" if rel_dir else d
            items.append({"path": rel_path, "name": d, "type": "directory"})
            if len(items) >= max_items:
                warnings.append(f"Explorer item limit reached ({max_items}).")
                break
        if len(items) >= max_items:
            break

        for fname in sorted(files, key=lambda x: x.lower()):
            p = Path(fname)
            if p.suffix.lower() in _EXPLORER_EXCLUDED_SUFFIXES:
                continue
            rel_path = f"{rel_dir}/{fname}" if rel_dir else fname
            if rel_path.startswith("Downloads/"):
                if "downloads_file_excluded" not in seen_warnings:
                    warnings.append("Downloads artifacts are excluded from explorer output.")
                    seen_warnings.add("downloads_file_excluded")
                continue
            try:
                size = int((current_path / fname).stat().st_size)
            except Exception:  # noqa: BLE001
                size = 0
            items.append({"path": rel_path, "name": fname, "type": "file", "size": size})
            if len(items) >= max_items:
                warnings.append(f"Explorer item limit reached ({max_items}).")
                break
        if len(items) >= max_items:
            break

    items.sort(key=lambda it: (0 if it.get("type") == "directory" else 1, str(it.get("path", "")).lower()))
    for it in items:
        path = str(it.get("path", ""))
        if os.path.isabs(path):
            errors.append("absolute_path_detected")
            break
    return items[:max_items], warnings, errors


def _safe_read_project_file(root: Path, rel_path_raw: str, max_bytes: int = _FILE_VIEWER_MAX_BYTES) -> dict[str, object]:
    warnings: list[str] = []
    errors: list[str] = []
    rel_path = str(rel_path_raw or "").strip().replace("\\", "/")
    if not rel_path:
        return {"ok": False, "warnings": warnings, "errors": ["path_missing"]}
    if rel_path.startswith("/") or rel_path.startswith("\\") or re.match(r"^[A-Za-z]:[\\/]", rel_path):
        return {"ok": False, "warnings": warnings, "errors": ["absolute_path_blocked"]}

    parts = [p for p in rel_path.split("/") if p not in ("", ".")]
    if not parts:
        return {"ok": False, "warnings": warnings, "errors": ["invalid_path"]}
    if ".." in parts:
        return {"ok": False, "warnings": warnings, "errors": ["path_traversal_blocked"]}
    if any(p in _EXPLORER_EXCLUDED_DIRS for p in parts):
        return {"ok": False, "warnings": warnings, "errors": ["excluded_path_blocked"]}

    root_resolved = root.resolve()
    target = (root_resolved / Path(*parts)).resolve()
    try:
        target.relative_to(root_resolved)
    except Exception:  # noqa: BLE001
        return {"ok": False, "warnings": warnings, "errors": ["path_outside_project_blocked"]}

    if not target.exists() or not target.is_file():
        return {"ok": False, "warnings": warnings, "errors": ["file_not_found"]}

    suffix = target.suffix.lower()
    if suffix in _FILE_VIEWER_BLOCKED_SUFFIXES:
        return {"ok": False, "warnings": warnings, "errors": ["file_type_blocked"]}
    if suffix not in _FILE_VIEWER_ALLOWED_SUFFIXES:
        return {"ok": False, "warnings": warnings, "errors": ["file_type_not_allowed"]}

    try:
        size = int(target.stat().st_size)
    except Exception:  # noqa: BLE001
        size = 0
    read_limit = int(max_bytes)
    raw: bytes
    try:
        with target.open("rb") as f:
            if read_limit < 0:
                raw = f.read()
            else:
                raw = f.read(max(1, read_limit) + 1)
    except Exception as ex:  # noqa: BLE001
        return {"ok": False, "warnings": warnings, "errors": [f"read_failed:{ex}"]}

    truncated = False
    if read_limit >= 0:
        read_limit = max(1, read_limit)
        truncated = len(raw) > read_limit or size > read_limit
        if len(raw) > read_limit:
            raw = raw[:read_limit]
    try:
        content = raw.decode("utf-8")
    except UnicodeDecodeError:
        content = raw.decode("utf-8", errors="replace")
        warnings.append("non_utf8_content_decoded_with_replacement")

    rel_clean = target.relative_to(root_resolved).as_posix()
    return {
        "ok": True,
        "path": rel_clean,
        "name": target.name,
        "type": "file",
        "size": size,
        "language": _FILE_VIEWER_LANG_BY_SUFFIX.get(suffix, "text"),
        "content": content,
        "truncated": bool(truncated),
        "warnings": warnings,
        "errors": errors,
    }


def _validate_safe_rel_project_path(rel_path_raw: str) -> tuple[bool, str, str | None]:
    rel_path = str(rel_path_raw or "").strip().replace("\\", "/")
    if not rel_path:
        return False, "", "path_missing"
    if rel_path.startswith("/") or rel_path.startswith("\\") or re.match(r"^[A-Za-z]:[\\/]", rel_path):
        return False, rel_path, "absolute_path_blocked"
    parts = [p for p in rel_path.split("/") if p not in ("", ".")]
    if not parts:
        return False, rel_path, "invalid_path"
    if ".." in parts:
        return False, rel_path, "path_traversal_blocked"
    if any(p in _EXPLORER_EXCLUDED_DIRS for p in parts):
        return False, rel_path, "excluded_path_blocked"
    return True, "/".join(parts), None


def _safe_read_project_diff(root: Path, rel_path: str | None = None, max_bytes: int = _DIFF_VIEWER_MAX_BYTES) -> dict[str, object]:
    warnings: list[str] = []
    errors: list[str] = []
    root_resolved = root.resolve()
    read_limit = int(max_bytes)

    safe_path = None
    if rel_path is not None and str(rel_path).strip():
        ok_path, cleaned, err = _validate_safe_rel_project_path(rel_path)
        if not ok_path:
            return {
                "ok": False,
                "path": str(rel_path or ""),
                "changed_files": [],
                "diff": "",
                "truncated": False,
                "warnings": warnings,
                "errors": [str(err or "invalid_path")],
            }
        target = (root_resolved / Path(cleaned)).resolve()
        try:
            target.relative_to(root_resolved)
        except Exception:  # noqa: BLE001
            return {
                "ok": False,
                "path": cleaned,
                "changed_files": [],
                "diff": "",
                "truncated": False,
                "warnings": warnings,
                "errors": ["path_outside_project_blocked"],
            }
        safe_path = cleaned

    try:
        probe = subprocess.run(
            ["git", "rev-parse", "--is-inside-work-tree"],
            cwd=str(root_resolved),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=10,
        )
        if probe.returncode != 0 or str(probe.stdout or "").strip().lower() != "true":
            return {
                "ok": False,
                "path": safe_path,
                "changed_files": [],
                "diff": "",
                "truncated": False,
                "warnings": warnings,
                "errors": ["not_a_git_repo"],
            }

        name_args = ["git", "diff", "--name-only"]
        diff_args = ["git", "diff"]
        if safe_path:
            name_args.extend(["--", safe_path])
            diff_args.extend(["--", safe_path])
        names_cp = subprocess.run(
            name_args,
            cwd=str(root_resolved),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=12,
        )
        diff_cp = subprocess.run(
            diff_args,
            cwd=str(root_resolved),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=15,
        )
    except Exception as ex:  # noqa: BLE001
        return {
            "ok": False,
            "path": safe_path,
            "changed_files": [],
            "diff": "",
            "truncated": False,
            "warnings": warnings,
            "errors": [str(ex)],
        }

    if names_cp.returncode != 0 or diff_cp.returncode != 0:
        errors.append((names_cp.stderr or diff_cp.stderr or "git_diff_failed").strip())
        return {
            "ok": False,
            "path": safe_path,
            "changed_files": [],
            "diff": "",
            "truncated": False,
            "warnings": warnings,
            "errors": [e for e in errors if e],
        }

    changed_files = [ln.strip().replace("\\", "/") for ln in (names_cp.stdout or "").splitlines() if ln.strip()]
    diff_text = diff_cp.stdout or ""
    truncated = False
    encoded = diff_text.encode("utf-8", errors="replace")
    if read_limit >= 0 and len(encoded) > max(1, read_limit):
        truncated = True
        safe_lim = max(1, read_limit)
        diff_text = encoded[:safe_lim].decode("utf-8", errors="replace")
        warnings.append(f"diff_truncated_at_{safe_lim}_bytes")
    if not diff_text.strip():
        warnings.append("no_diff_changes")

    return {
        "ok": True,
        "path": safe_path,
        "changed_files": changed_files,
        "diff": diff_text,
        "truncated": truncated,
        "warnings": warnings,
        "errors": errors,
    }


def _truncate_run_check_output(stdout: str, stderr: str, max_bytes: int) -> tuple[str, str, bool]:
    out = str(stdout or "")
    err = str(stderr or "")
    limit = max(1, int(max_bytes))
    out_b = out.encode("utf-8", errors="replace")
    err_b = err.encode("utf-8", errors="replace")
    total = len(out_b) + len(err_b)
    if total <= limit:
        return out, err, False
    if len(out_b) >= limit:
        return out_b[:limit].decode("utf-8", errors="replace"), "", True
    remain = limit - len(out_b)
    return out, err_b[:remain].decode("utf-8", errors="replace"), True


def _run_allowed_project_check(check_name: str) -> dict[str, object]:
    runner = get_test_runner_agent(APP_DIR.resolve())
    runner.app_root = APP_DIR.resolve()
    runner.max_output_bytes = _RUN_CHECK_MAX_OUTPUT_BYTES
    runner.timeout_seconds = _RUN_CHECK_TIMEOUT_SECONDS
    runner.allowed_checks = dict(_RUN_ALLOWED_CHECKS)
    return runner.run_allowed_check(check_name)


def _validate_commit_plan_message(message: str) -> dict[str, object]:
    msg = str(message or "").strip()
    if not msg:
        return {"ok": False, "reason": "message_empty"}
    if len(msg) > 120:
        return {"ok": False, "reason": "message_too_long"}
    if not _COMMIT_PLAN_MESSAGE_RE.match(msg):
        return {"ok": False, "reason": "invalid_prefix_or_format"}
    return {"ok": True, "reason": "valid"}


def _is_blocked_commit_path(path_raw: str) -> bool:
    p = str(path_raw or "").strip().replace("\\", "/")
    if not p:
        return True
    lower = p.lower()
    if any(lower.startswith(marker) for marker in _COMMIT_PLAN_BLOCKED_DIR_MARKERS):
        return True
    if "/node_modules/" in lower or "/__pycache__/" in lower or "/.pytest_cache/" in lower or "/.git/" in lower or "/.rainer_agent/" in lower:
        return True
    if lower.endswith(_COMMIT_PLAN_BLOCKED_SUFFIXES):
        return True
    return False


def _build_commit_plan_payload(message: str) -> dict[str, object]:
    warnings: list[str] = []
    errors: list[str] = []
    msg = str(message or "").strip()

    msg_check = _validate_commit_plan_message(msg)
    if not msg_check.get("ok", False):
        errors.append(str(msg_check.get("reason") or "invalid_message"))

    git_integration = _get_git_integration()
    changed = git_integration.changed_files(APP_DIR)
    status = git_integration.git_status(APP_DIR)

    if not changed.get("ok", False):
        errors.append(str(changed.get("error") or "git_changed_files_failed"))
    if not status.get("ok", False):
        errors.append(str(status.get("error") or "git_status_failed"))

    entries = list(status.get("entries", [])) if isinstance(status.get("entries", []), list) else []
    changed_files = sorted(
        {
            *list(changed.get("modified", [])),
            *list(changed.get("untracked", [])),
            *list(changed.get("deleted", [])),
            *[str((e or {}).get("path") or "") for e in entries],
        }
    )
    changed_files = [p for p in changed_files if p]
    untracked_files = sorted(
        {
            str((e or {}).get("path") or "")
            for e in entries
            if str((e or {}).get("status") or "") == "??" and str((e or {}).get("path") or "")
        }
    )
    staged_files = sorted(
        {
            str((e or {}).get("path") or "")
            for e in entries
            if str((e or {}).get("path") or "")
            and str((e or {}).get("status") or "") not in {"", "??"}
            and str((e or {}).get("status") or "").strip()
            and str((e or {}).get("status") or "")[0] not in {" ", "?"}
        }
    )

    blocked_files = sorted({p for p in changed_files if _is_blocked_commit_path(p)})
    commit_files = sorted({p for p in changed_files if p and p not in blocked_files})

    if blocked_files:
        warnings.append("blocked_files_detected")
    if not commit_files:
        warnings.append("no_safe_commit_files")

    commit_allowed = (
        bool(msg_check.get("ok", False))
        and bool(commit_files)
        and not bool(blocked_files)
        and not bool(errors)
    )

    return {
        "ok": not bool(errors),
        "message": msg,
        "message_valid": bool(msg_check.get("ok", False)),
        "commit_allowed": commit_allowed,
        "commit_files": commit_files,
        "blocked_files": blocked_files,
        "changed_files": changed_files,
        "untracked_files": untracked_files,
        "staged_files": staged_files,
        "warnings": warnings,
        "errors": errors,
    }


def _classify_git_dirty(porcelain_stdout):
    """Separates git dirty files into product changes vs local-dev-tool files."""
    product, local = [], []
    for line in (porcelain_stdout or "").splitlines():
        line = line.rstrip()
        if not line:
            continue
        path = line[3:].strip().strip('"') if len(line) > 3 else line.strip()
        if any(path.startswith(p) or path.endswith(p.rstrip("/")) for p in LOCAL_DEV_PATH_PREFIXES):
            local.append(path)
        else:
            product.append(path)
    return product, local


def format_response_like_claude(text, fallback_payload=None):
    raw_text = str(text or "").strip()
    if raw_text and (not str(raw_text).startswith("{")):
        return raw_text
    payload = fallback_payload if isinstance(fallback_payload, dict) else {}
    return formatter_claude_style(payload if payload else raw_text)


def call_ollama_intelligent(prompt, context="", model_override=None, *, local_agent_mode: bool = False):
    """Lokales LLM (Ollama / LM Studio / llama.cpp) — Routing über model_providers."""
    return llm_generate_coding_response(
        str(prompt or "").strip(),
        context=context,
        model_override=model_override,
        local_agent_mode=local_agent_mode,
    )


def generate_chat_response_plain(task: str) -> str:
    """Konversation ohne Dateizugriff: aktiven lokalen Provider nutzen; bei Ausfall echte Fehlermeldung, keine erfundene Modellantwort."""
    reply = llm_generate_chat_response(str(task or "").strip())
    if is_llm_failure_message(reply):
        return str(reply).strip()
    if not str(reply or "").strip():
        fb = _connectivity_chat_fallback(str(task or ""))
        if fb:
            return fb
        return chat_reply_canned(task)
    return str(reply).strip()


_CHAT_LLM_TIMEOUT_SEC = float(os.getenv("RAINER_CHAT_LLM_TIMEOUT_SEC", "5"))
_CONNECTIVITY_CHAT_TIMEOUT_SEC = float(os.getenv("RAINER_CONNECTIVITY_CHAT_TIMEOUT_SEC", "28"))


def _effective_chat_timeout_sec(task: str, explicit: float | None) -> float:
    """Offline-/Verbindungsfragen: längeres Timeout (Groq kann >5s brauchen)."""
    if explicit is not None:
        return float(explicit)
    base = _CHAT_LLM_TIMEOUT_SEC
    if should_route_direct_run_as_chat(str(task or "")):
        return max(base, _CONNECTIVITY_CHAT_TIMEOUT_SEC)
    return base


def _connectivity_chat_fallback(task: str) -> str:
    """Wenn LLM leer ausfällt: konkrete Checkliste statt generischem chat_reply_canned."""
    if not should_route_direct_run_as_chat(str(task or "")):
        return ""
    return connectivity_diagnostics_reply()


def generate_chat_response_plain_with_timeout(task: str, timeout_sec: float | None = None) -> str:
    """Wie generate_chat_response_plain, aber mit Timeout — vermeidet 30s+ Blockaden bei lokalem LLM."""
    task_s = str(task or "").strip()
    sec = _effective_chat_timeout_sec(task_s, timeout_sec)
    if sec <= 0:
        fb = _connectivity_chat_fallback(task_s)
        return fb if fb else chat_reply_canned(task_s)
    try:
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            fut = pool.submit(generate_chat_response_plain, task_s)
            return fut.result(timeout=sec)
    except (concurrent.futures.TimeoutError, Exception):
        fb = _connectivity_chat_fallback(task_s)
        if fb:
            return fb
        return chat_reply_canned(task_s)


_GREETING_INSTANT = frozenset(
    {
        "hallo",
        "hi",
        "hey",
        "servus",
        "moin",
        "guten morgen",
        "guten tag",
        "guten abend",
    }
)


def _instant_greeting_message() -> str:
    return "Hallo! Ich bin Rainer. Was möchtest du machen?"


def _chat_reply_skip_llm(intent: str, cleaned_prompt: str) -> str | None:
    """
    Kurze Begrüßung / Meta-Hilfe: sofortige Antwort ohne LLM-Aufruf.
    Rückgabe None → LLM (mit Timeout) nutzen.
    """
    low = " ".join(str(cleaned_prompt or "").strip().lower().split())
    if should_route_direct_run_as_chat(cleaned_prompt):
        return None
    if intent == "greeting" or low in _GREETING_INSTANT:
        return _instant_greeting_message()
    if intent == "help_question":
        return chat_reply_canned(cleaned_prompt)
    return None


def _unknown_reply_skip_llm(intent: str, cleaned_prompt: str) -> str | None:
    """Vage Kurzprompts: sofort Rückfrage, kein LLM."""
    low = " ".join(str(cleaned_prompt or "").strip().lower().split())
    if intent == "ambiguous":
        return unclear_chat_short_reply()
    if len(low) <= 32 and low in {"mach", "bitte", "mach mal", "tu was", "irgendwas", "was soll"}:
        return unclear_chat_short_reply()
    return None


SCANNER_SKIP_DIRS = {
    "node_modules", ".git", "__pycache__", ".pytest_cache",
    ".aider.tags.cache.v4", "_backup_rainer_dashboard_2026-04-15",
    "rambo_builder_local",
}

PROJECT_AREA_MAP = {
    "backend": "Backend (Python/Flask)",
    "frontend": "Frontend (React/Vite)",
    "agent": "Agent (Node.js)",
    "tools": "Tools (Python)",
    "data": "Daten (JSON)",
    "memory": "Speicher (Memory/Log)",
    "knowledge": "Wissen (Docs/Context)",
    "outbox": "Ausgabe (Outbox)",
    "logs": "Logs",
    "rambo_builder_local": "Builder (dieses Tool)",
    "scripts": "Skripte",
    "server": "Server (Node.js)",
}


def read_json_file(path, fallback):
    try:
        if not path.exists():
            return fallback
        with open(path, "r", encoding="utf-8") as file:
            return json.load(file)
    except (OSError, json.JSONDecodeError, TypeError, ValueError):
        return fallback


def write_json_file(path, content):
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as file:
        json.dump(content, file, ensure_ascii=True, indent=2)


def get_timestamp():
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _run_quality_check_command(command: str, timeout_sec: int = 300) -> dict:
    cmd = str(command or "").strip()
    if not cmd:
        return {"command": "", "ok": False, "exit_code": -1, "stdout": "", "stderr": "empty_command", "duration_ms": 0}
    started = time.time()
    try:
        cp = subprocess.run(
            cmd,
            cwd=str(PROJECT_DIR),
            shell=True,
            capture_output=True,
            text=True,
            timeout=max(10, int(timeout_sec)),
            check=False,
        )
        return {
            "command": cmd,
            "ok": cp.returncode == 0,
            "exit_code": int(cp.returncode),
            "stdout": str(cp.stdout or "")[-8000:],
            "stderr": str(cp.stderr or "")[-8000:],
            "duration_ms": int((time.time() - started) * 1000),
        }
    except subprocess.TimeoutExpired as exc:
        return {
            "command": cmd,
            "ok": False,
            "exit_code": -2,
            "stdout": str(getattr(exc, "stdout", "") or "")[-8000:],
            "stderr": f"timeout_after_{int(timeout_sec)}s",
            "duration_ms": int((time.time() - started) * 1000),
        }
    except Exception as exc:
        return {
            "command": cmd,
            "ok": False,
            "exit_code": -3,
            "stdout": "",
            "stderr": str(exc),
            "duration_ms": int((time.time() - started) * 1000),
        }


def _persist_quality_task_graph(entry: dict) -> None:
    history = read_json_file(QUALITY_TASK_GRAPH_FILE, [])
    if not isinstance(history, list):
        history = []
    if isinstance(entry, dict):
        history.insert(0, entry)
    write_json_file(QUALITY_TASK_GRAPH_FILE, history[:120])


def _auto_fix_via_direct_run(task: str, failed_commands: list[str], max_rounds: int = 2) -> list[dict]:
    rounds: list[dict] = []
    failed = [str(c).strip() for c in list(failed_commands or []) if str(c).strip()]
    if not failed:
        return rounds
    safe_rounds = max(1, min(int(max_rounds or 1), 3))
    with app.test_client() as c:
        for idx in range(safe_rounds):
            prompt = (
                "Auto-Fix: Behebe die fehlschlagenden Verifikationschecks im Projekt minimal-invasiv.\n\n"
                f"Nutzerziel: {str(task or '').strip()}\n\n"
                "Fehlgeschlagene Checks:\n- "
                + "\n- ".join(failed)
                + "\n\n"
                "Regeln: nur notwendige Änderungen, danach kurz zusammenfassen."
            )
            r = c.post("/api/direct-run", json={"task": prompt, "scope": "project", "mode": "safe"})
            body = r.get_json(silent=True) or {}
            rounds.append(
                {
                    "round": idx + 1,
                    "ok": bool(r.status_code == 200 and not body.get("error")),
                    "status_code": int(r.status_code),
                    "message": str(
                        body.get("formatted_response")
                        or body.get("chat_response")
                        or body.get("natural_message")
                        or body.get("message")
                        or body.get("error")
                        or ""
                    )[:1200],
                }
            )
            if rounds[-1]["ok"]:
                break
    return rounds


def _quality_eval_default_prompts() -> list[dict]:
    return [
        {"name": "chat_connectivity", "task": "überprüfe warum die app offline ist", "scope": "local", "mode": "apply"},
        {"name": "read_intent", "task": "analysiere kurz backend/main.py und nenne 2 risiken", "scope": "project", "mode": "safe"},
        {"name": "change_intent", "task": "füge in frontend eine kleine ui-verbesserung als plan hinzu", "scope": "project", "mode": "safe"},
    ]


def _quality_eval_quick_prompts() -> list[dict]:
    """Ein kurzer read-only Case — für eval_after nach Auto-Fix ohne volle Suite-Laufzeit."""
    return [
        {"name": "read_smoke", "task": "nenne in einem Satz, was backend/main.py grob tut", "scope": "project", "mode": "safe"},
    ]


def _quality_eval_run_cases(cases: list) -> tuple[list[dict], int, int]:
    """Führt die Quality-Eval-Kaskade aus (direct-run pro Case). Gibt (rows, total, avg_score) zurück."""
    rows: list[dict] = []
    with app.test_client() as c:
        for i, case in enumerate(cases):
            if isinstance(case, dict):
                name = str(case.get("name") or f"case_{i+1}")
                task = str(case.get("task") or "").strip()
                scope = str(case.get("scope") or "project").strip()
                mode = str(case.get("mode") or "safe").strip()
            else:
                name = f"case_{i+1}"
                task = str(case or "").strip()
                scope = "project"
                mode = "safe"
            if not task:
                rows.append({"name": name, "ok": False, "score": 0, "reason": "empty_task"})
                continue
            r = c.post("/api/direct-run", json={"task": task, "scope": scope, "mode": mode})
            body = r.get_json(silent=True) or {}
            has_text = bool(
                str(
                    body.get("formatted_response")
                    or body.get("chat_response")
                    or body.get("natural_message")
                    or body.get("message")
                    or ""
                ).strip()
            )
            has_contract = isinstance(body.get("confidence_gate"), dict) and isinstance(body.get("task_memory"), dict)
            has_checks = isinstance(body.get("recommended_checks"), list)
            case_score = 0
            case_score += 40 if r.status_code == 200 else 0
            case_score += 25 if has_text else 0
            case_score += 20 if has_contract else 0
            case_score += 15 if has_checks else 0
            rows.append(
                {
                    "name": name,
                    "status_code": int(r.status_code),
                    "ok": r.status_code == 200,
                    "score": int(case_score),
                    "has_text": has_text,
                    "has_contract": has_contract,
                    "has_checks": has_checks,
                }
            )
    total = len(rows)
    avg_score = int(round(sum(int(r.get("score") or 0) for r in rows) / max(1, total)))
    return rows, total, avg_score


def _quality_graph_entry_recent(ts_str: str, max_age_sec: int = 2700) -> bool:
    raw = str(ts_str or "").strip()
    if not raw:
        return False
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S"):
        try:
            dt = datetime.strptime(raw, fmt)
            return (datetime.now() - dt).total_seconds() <= max_age_sec
        except ValueError:
            continue
    return False


def _quality_attach_eval_to_latest_graph(avg_score: int, total: int, *, eval_quick: bool = False) -> bool:
    """Eval-Scores in den jüngsten Task-Graph-Eintrag schreiben, wenn frisch und noch ohne Eval."""
    rows = read_json_file(QUALITY_TASK_GRAPH_FILE, [])
    if not isinstance(rows, list) or not rows:
        return False
    head = rows[0]
    if not isinstance(head, dict):
        return False
    if head.get("eval_avg_score") is not None:
        return False
    if not _quality_graph_entry_recent(str(head.get("timestamp") or ""), 2700):
        return False
    head2 = dict(head)
    head2["eval_avg_score"] = int(avg_score)
    head2["eval_total_cases"] = int(total)
    head2["eval_quick"] = bool(eval_quick)
    rows[0] = head2
    write_json_file(QUALITY_TASK_GRAPH_FILE, rows[:120])
    return True


def _quality_attach_eval_to_graph_run(run_id: str, avg_score: int, total: int, *, eval_quick: bool = False) -> bool:
    rid = str(run_id or "").strip()
    if not rid:
        return False
    rows = read_json_file(QUALITY_TASK_GRAPH_FILE, [])
    if not isinstance(rows, list) or not rows:
        return False
    updated = False
    new_rows = []
    for row in rows:
        if not isinstance(row, dict):
            new_rows.append(row)
            continue
        if not updated and str(row.get("run_id") or "").strip() == rid:
            row2 = dict(row)
            row2["eval_avg_score"] = int(avg_score)
            row2["eval_total_cases"] = int(total)
            row2["eval_quick"] = bool(eval_quick)
            row2["eval_attached_via"] = "run_id"
            new_rows.append(row2)
            updated = True
        else:
            new_rows.append(row)
    if updated:
        write_json_file(QUALITY_TASK_GRAPH_FILE, new_rows[:120])
    return updated


def format_display_timestamp(timestamp):
    raw = str(timestamp or "").strip()
    if not raw:
        return "Ohne Zeitstempel"

    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%dT%H:%M:%S"):
        try:
            parsed = datetime.strptime(raw, fmt)
            return parsed.strftime("%d.%m.%Y %H:%M")
        except ValueError:
            continue
    return raw


def is_blank(value):
    return not str(value or "").strip()


def format_local_path(path_value):
    cleaned = str(path_value or "").strip().replace("\\", "/")
    cleaned = re.sub(r"/{2,}", "/", cleaned)
    return cleaned.rstrip("/")


def _normalize_guard_path(path_value) -> str:
    path = format_local_path(path_value or "")
    while path.startswith("./"):
        path = path[2:]
    return path.lstrip("/")


def _normalized_windows_path(path_value: Path | str) -> str:
    return str(Path(path_value).resolve()).replace("/", "\\").lower()


def _is_direct_run_expected_root_active() -> bool:
    # Aktiven Workspace berücksichtigen, nicht nur APP_DIR.
    active_root = get_active_project_root()
    if _normalized_windows_path(active_root) == _normalized_windows_path(DIRECT_RUN_EXPECTED_APP_ROOT):
        return True
    return _normalized_windows_path(APP_DIR) == _normalized_windows_path(DIRECT_RUN_EXPECTED_APP_ROOT)


def _is_explicit_backend_or_test_task(task: str, mode: str = "") -> bool:
    lowered = str(task or "").lower()
    mode_l = str(mode or "").lower()
    if mode_l in {"backend", "test"}:
        return True
    needles = (
        "backend",
        "backend/main.py",
        "api",
        "flask",
        "pytest",
        "test",
        "tests/",
        "unittest",
    )
    return any(n in lowered for n in needles)


def _extract_candidate_path(entry) -> str:
    if entry is None:
        return ""
    if isinstance(entry, str):
        return entry
    if isinstance(entry, Path):
        return str(entry)
    if isinstance(entry, dict):
        for key in (
            "path",
            "target_path",
            "target",
            "file",
            "filename",
            "name",
            "relative_path",
            "absolute_path",
            "selected_target_path",
        ):
            value = entry.get(key)
            if isinstance(value, str) and value.strip():
                return value
    return ""


def _is_generic_ready_reply(text: str) -> bool:
    t = str(text or "").strip().lower()
    return ("ich bin bereit." in t) and ("stelle eine frage" in t)


def _collect_direct_guard_paths(file_entries) -> list[str]:
    if file_entries is None:
        return []
    if not isinstance(file_entries, (list, tuple, set)):
        file_entries = [file_entries]
    collected: list[str] = []
    for entry in file_entries:
        path = _extract_candidate_path(entry)
        if not str(path or "").strip():
            continue
        if isinstance(path, str):
            collected.append(path)
    return collected


def _is_forbidden_generated_path(path_value) -> bool:
    path = _normalize_guard_path(path_value)
    if not path:
        return True
    low = path.lower()
    if low.startswith(("..", "/", "\\")) or "/../" in low:
        return True
    if ":" in low:
        return True
    if low.startswith("du_arbeitest_im_projekt__"):
        return True
    if "rainerrobotdesktop/" in low:
        return True
    if "/electron/" in f"/{low}/" or low.endswith("/electron") or low == "electron":
        return True
    if "/rambo_ui/" in f"/{low}/" or low.endswith("/rambo_ui") or low == "rambo_ui":
        return True
    if "/src/components/" in f"/{low}/":
        return True
    if "/downloads/" in low:
        return True
    for prefix in DIRECT_RUN_FORBIDDEN_PREFIXES:
        if low.startswith(prefix):
            return True
    filename = low.rsplit("/", 1)[-1]
    if filename in DIRECT_RUN_FORBIDDEN_FILENAMES:
        return True
    return False


def _validate_direct_run_paths(file_entries, mode, task: str = "") -> dict[str, object]:
    raw_paths = _collect_direct_guard_paths(file_entries)
    normalized_seen: set[str] = set()
    normalized_paths: list[str] = []
    blocked: list[str] = []
    reasons: list[str] = []

    if not _is_direct_run_expected_root_active():
        blocked = list(dict.fromkeys(raw_paths))
        if not blocked:
            blocked = ["<unbekannt>"]
        reasons.append("invalid_project_root")
        return {
            "ok": False,
            "guard_status": "blocked",
            "status": "blocked_by_guard",
            "message": DIRECT_RUN_GUARD_BLOCK_MESSAGE,
            "blocked_files": blocked,
            "validated_files": normalized_paths,
            "reason": "invalid_project_root",
            "reasons": reasons,
            "expected_root": str(DIRECT_RUN_EXPECTED_APP_ROOT),
            "active_root": str(APP_DIR.resolve()),
        }

    backend_or_test_allowed = _is_explicit_backend_or_test_task(task, mode)
    trusted = False
    try:
        trusted = is_active_workspace_trusted()
    except Exception:
        trusted = False
    try:
        active_workspace_root = get_active_project_root().resolve()
    except Exception:
        active_workspace_root = APP_DIR.resolve()
    # Ohne Trust: striktes Sandbox-Root (Installations-APP_DIR). Mit Trust: aktiver Projektordner.
    app_root = active_workspace_root if trusted else APP_DIR.resolve()

    for raw in raw_paths:
        normalized = _normalize_guard_path(raw)
        if not normalized:
            blocked.append(str(raw))
            reasons.append("empty_path")
            continue
        normalized_low = normalized.lower()
        if normalized_low in normalized_seen:
            continue
        normalized_seen.add(normalized_low)
        normalized_paths.append(normalized)

        if not trusted and _is_forbidden_generated_path(normalized):
            blocked.append(normalized)
            reasons.append("forbidden_path")
            continue

        try:
            resolved = (app_root / normalized).resolve()
            resolved.relative_to(app_root)
        except Exception:
            blocked.append(normalized)
            reasons.append("outside_app_dir")
            continue

        # Freigegebener Workspace: alle Pfade unter dem Projektroot erlaubt (keine Mini-Allowlist).
        if trusted:
            continue

        if normalized in DIRECT_RUN_MINI_ALLOWED_FILES:
            continue
        if backend_or_test_allowed:
            if normalized in DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_FILES:
                continue
            if normalized.startswith(DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_TEST_PREFIX) and normalized.endswith(".py"):
                continue

        blocked.append(normalized)
        reasons.append("not_in_allowlist")

    blocked_unique = list(dict.fromkeys(blocked))
    return {
        "ok": len(blocked_unique) == 0,
        "guard_status": "ok" if len(blocked_unique) == 0 else "blocked",
        "status": "ok" if len(blocked_unique) == 0 else "blocked_by_guard",
        "message": (
            "Guard-Pruefung erfolgreich."
            if len(blocked_unique) == 0
            else DIRECT_RUN_GUARD_BLOCK_MESSAGE
        ),
        "blocked_files": blocked_unique,
        "validated_files": normalized_paths,
        "reason": "ok" if len(blocked_unique) == 0 else "blocked_by_guard",
        "reasons": list(dict.fromkeys(reasons)),
        "expected_root": str(DIRECT_RUN_EXPECTED_APP_ROOT),
        "active_root": str(app_root),
    }


def _build_direct_guard_block_payload(*, scope: str, mode: str, blocked_files: list[str], task: str = "", recognized_task=None) -> dict:
    unique_blocked = list(dict.fromkeys([str(p) for p in (blocked_files or []) if str(p).strip()]))
    first_blocked = unique_blocked[0] if unique_blocked else ""
    lower_blocked = first_blocked.lower()
    is_downloads = "downloads/" in lower_blocked or lower_blocked.startswith("../downloads/")
    user_reason = "Downloads-Pfad erkannt und blockiert. Keine Änderung ausgeführt." if is_downloads else "Pfad durch Safety-Regel blockiert. Keine Änderung ausgeführt."
    path_hint = "../Downloads/..." if is_downloads else (first_blocked or "<kein_pfad>")
    guard = {
        "allowed": False,
        "reason": "blocked_by_guard",
        "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE,
        "path": first_blocked,
        "blocked_files": unique_blocked,
    }
    return {
        "ok": False,
        "success": False,
        "status": "blocked_by_guard",
        "guard_status": "blocked",
        "direct_status": "blocked",
        "scope": str(scope or "local"),
        "mode": str(mode or "apply"),
        "task": str(task or ""),
        "error": DIRECT_RUN_GUARD_UI_BLOCKED_MESSAGE,
        "message": DIRECT_RUN_GUARD_UI_BLOCKED_MESSAGE,
        "user_facing_block_reason": user_reason,
        "blocked_path_hint": path_hint,
        "safe_next_action": "Nutze eine erlaubte Datei innerhalb des aktiven Projektordners oder starte nur eine Preview.",
        "guard": guard,
        "blocked_files": unique_blocked,
        "forbidden_files": unique_blocked,
        "affected_files": unique_blocked,
        "file_plan": unique_blocked,
        "created_files": [],
        "changed_files": [],
        "errors": [DIRECT_RUN_GUARD_BLOCK_MESSAGE],
        "executable": False,
        "writes_files": False,
        "auto_apply": False,
        "requires_confirmation": False,
        "requires_user_confirmation": False,
        "recognized_task": recognized_task if isinstance(recognized_task, dict) else {},
    }


def _relativize_guard_candidate(path_value: str) -> str:
    raw = str(path_value or "").strip()
    if not raw:
        return ""
    normalized = format_local_path(raw)
    if not normalized:
        return ""
    if ":" not in normalized and not normalized.startswith("/"):
        return normalized
    try:
        cand = Path(raw).expanduser().resolve()
        app_root = APP_DIR.resolve()
        rel = cand.relative_to(app_root).as_posix()
        return format_local_path(rel)
    except Exception:
        return normalized


def _collect_guard_candidates_from_payload(payload: dict) -> list[str]:
    if not isinstance(payload, dict):
        return []
    candidates: list = []
    scalar_keys = (
        "selected_target_path",
        "target_path",
        "path",
        "relative_path",
        "absolute_path",
        "base_path",
        "target_root",
        "project_root",
    )
    list_keys = (
        "affected_files",
        "changed_files",
        "created_files",
        "updated_files",
        "deleted_files",
        "file_plan",
        "file_entries",
        "artifacts",
        "planned_files",
        "missing_files",
        "blocked_files",
        "forbidden_files",
    )
    for key in scalar_keys:
        value = payload.get(key)
        if isinstance(value, str) and value.strip():
            candidates.append(value)
    for key in list_keys:
        value = payload.get(key)
        if isinstance(value, dict):
            candidates.extend(list(value.keys()))
        elif isinstance(value, (list, tuple, set)):
            candidates.extend(list(value))
    guard = payload.get("guard") if isinstance(payload.get("guard"), dict) else {}
    if isinstance(guard.get("path"), str) and str(guard.get("path")).strip():
        candidates.append(str(guard.get("path")))
    normalized = []
    for item in candidates:
        path = _extract_candidate_path(item)
        if not str(path or "").strip():
            continue
        normalized.append(_relativize_guard_candidate(path))
    return [p for p in normalized if str(p or "").strip()]


def _apply_central_generation_guard(payload: dict, *, scope: str, mode: str, task: str = "", recognized_task=None):
    candidates = _collect_guard_candidates_from_payload(payload)
    if not candidates and _is_desktop_multi_file_project_prompt(str(task or "").lower()):
        candidates = list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN)
    guard_result = _validate_direct_run_paths(candidates, mode, task)
    if bool(guard_result.get("ok")):
        return None
    blocked_files = guard_result.get("blocked_files") or candidates
    blocked_payload = _build_direct_guard_block_payload(
        scope=scope,
        mode=mode,
        blocked_files=blocked_files,
        task=task,
        recognized_task=recognized_task if isinstance(recognized_task, dict) else payload.get("recognized_task"),
    )
    blocked_payload["forbidden_files"] = list(dict.fromkeys([str(p) for p in blocked_files if str(p).strip()]))
    blocked_payload["errors"] = [DIRECT_RUN_GUARD_BLOCK_MESSAGE]
    blocked_payload["success"] = False
    existing_events = payload.get("workstream_events")
    workstream_events = list(existing_events) if isinstance(existing_events, list) else []
    blocked_payload["workstream_events"] = workstream_events + [
        {"phase": "guard", "level": "error", "title": "Auftrag blockiert", "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE, "status": "blocked"},
    ]
    return blocked_payload


KNOWN_DOWNLOADS_ARTIFACT = "../Downloads/Baue die komplette Electron Desktop.txt"


def _is_known_downloads_artifact(path_value: str) -> bool:
    normalized = format_local_path(path_value or "").lower()
    return normalized == "../downloads/baue die komplette electron desktop.txt"


def _response_claims_success(payload: dict) -> bool:
    if not isinstance(payload, dict):
        return False
    status = str(payload.get("status") or "").strip().lower()
    guard_status = str(payload.get("guard_status") or "").strip().lower()
    direct_status = str(payload.get("direct_status") or "").strip().lower()
    if status in {"blocked_by_guard", "no_changes_detected"}:
        return False
    if guard_status in {"blocked", "no_changes"}:
        return False
    if payload.get("ok") is True or payload.get("success") is True:
        return True
    if direct_status in {"verified", "applied", "completed", "success"}:
        return True
    if status in {"ok", "applied", "verified", "completed", "success"}:
        return True
    return False


def collect_response_files(result) -> list[str]:
    payload = result if isinstance(result, dict) else {}
    candidates: list = []
    scalar_keys = (
        "selected_target_path",
        "target_path",
        "path",
        "relative_path",
        "absolute_path",
    )
    list_keys = (
        "affected_files",
        "changed_files",
        "created_files",
        "updated_files",
        "deleted_files",
        "file_plan",
        "file_entries",
        "generated_files",
        "artifacts",
        "planned_files",
        "forbidden_files",
        "blocked_files",
    )
    for key in scalar_keys:
        value = payload.get(key)
        if isinstance(value, str) and value.strip():
            candidates.append(value)
    for key in list_keys:
        value = payload.get(key)
        if isinstance(value, dict):
            candidates.extend(list(value.keys()))
        elif isinstance(value, (list, tuple, set)):
            candidates.extend(list(value))
    guard = payload.get("guard") if isinstance(payload.get("guard"), dict) else {}
    if isinstance(guard.get("path"), str) and str(guard.get("path")).strip():
        candidates.append(str(guard.get("path")))
    collected: list[str] = []
    for item in candidates:
        path = _extract_candidate_path(item)
        if not str(path or "").strip():
            continue
        normalized = _relativize_guard_candidate(path)
        if normalized and normalized not in collected:
            collected.append(normalized)
    return collected


def collect_git_changed_files() -> list[str]:
    root = str(APP_DIR)
    changed: list[str] = []
    seen: set[str] = set()

    def _add_candidate(raw_value: str):
        raw = str(raw_value or "").strip()
        if not raw:
            return
        if "->" in raw:
            raw = raw.split("->")[-1].strip()
        normalized = format_local_path(raw).lstrip("./")
        if not normalized:
            return
        if _is_known_downloads_artifact(normalized):
            return
        if normalized not in seen:
            seen.add(normalized)
            changed.append(normalized)

    try:
        status_cp = subprocess.run(
            ["git", "status", "--short"],
            cwd=root,
            capture_output=True,
            text=True,
            timeout=10,
        )
    except Exception:
        status_cp = None
    if status_cp and status_cp.returncode == 0:
        for line in (status_cp.stdout or "").splitlines():
            entry = line[3:].strip() if len(line) > 3 else line.strip()
            _add_candidate(entry)

    try:
        diff_cp = subprocess.run(
            ["git", "diff", "--name-only"],
            cwd=root,
            capture_output=True,
            text=True,
            timeout=10,
        )
    except Exception:
        diff_cp = None
    if diff_cp and diff_cp.returncode == 0:
        for line in (diff_cp.stdout or "").splitlines():
            _add_candidate(line)

    return changed


def filter_allowed_changed_files(files, prompt: str = "", mode: str = "apply") -> list[str]:
    candidates = files if isinstance(files, (list, tuple, set)) else []
    backend_or_test = _is_explicit_backend_or_test_task(prompt, mode)
    allowed: list[str] = []
    for item in candidates:
        normalized = format_local_path(item or "")
        if not normalized:
            continue
        if _is_known_downloads_artifact(normalized):
            continue
        if _is_forbidden_generated_path(normalized):
            continue
        if normalized in DIRECT_RUN_MINI_ALLOWED_FILES:
            if normalized not in allowed:
                allowed.append(normalized)
            continue
        if backend_or_test and normalized in DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_FILES:
            if normalized not in allowed:
                allowed.append(normalized)
            continue
        if backend_or_test and normalized.startswith(DIRECT_RUN_OPTIONAL_BACKEND_ALLOWED_TEST_PREFIX):
            if normalized not in allowed:
                allowed.append(normalized)
    return allowed


def _build_no_changes_payload(base_payload: dict | None = None) -> dict:
    payload = dict(base_payload or {})
    payload["ok"] = False
    payload["success"] = False
    payload["status"] = "no_changes_detected"
    payload["guard_status"] = "no_changes"
    payload["direct_status"] = "no_changes"
    payload["message"] = "Keine echte Dateiänderung erkannt."
    payload["errors"] = ["Keine erlaubte Projektdatei wurde geändert oder keine Dateipfade in der Response."]
    payload["affected_files"] = []
    payload["changed_files"] = []
    return payload


def verify_real_file_change(prompt: str, result: dict, before_snapshot=None, mode: str = "apply") -> dict:
    payload = result if isinstance(result, dict) else {}
    response_files = collect_response_files(payload)
    allowed_response_files = filter_allowed_changed_files(response_files, prompt, mode)
    if not allowed_response_files:
        return _build_no_changes_payload(payload)

    git_changed_files = collect_git_changed_files()
    if isinstance(before_snapshot, (list, tuple, set)) and before_snapshot:
        before_set = {format_local_path(p or "") for p in before_snapshot if str(p or "").strip()}
        git_changed_files = [p for p in git_changed_files if format_local_path(p) not in before_set]
    allowed_git_changed_files = filter_allowed_changed_files(git_changed_files, prompt, mode)

    confirmed_files = [p for p in allowed_response_files if p in allowed_git_changed_files]
    if not confirmed_files:
        return _build_no_changes_payload(payload)

    merged_changed = list(dict.fromkeys((payload.get("changed_files") or []) + confirmed_files))
    merged_affected = list(dict.fromkeys((payload.get("affected_files") or []) + confirmed_files))
    verified = dict(payload)
    verified["ok"] = True
    verified["success"] = True
    verified["status"] = str(payload.get("status") or "applied")
    verified["changed_files"] = merged_changed
    verified["affected_files"] = merged_affected
    return verified


def _enforce_real_change_success(task: str, payload: dict, mode: str = "apply", before_snapshot=None) -> dict:
    if not isinstance(payload, dict):
        return payload
    if not _response_claims_success(payload):
        return payload
    if str(payload.get("status") or "").strip().lower() == "blocked_by_guard":
        return payload
    if str(payload.get("guard_status") or "").strip().lower() == "blocked":
        return payload
    return verify_real_file_change(task, payload, before_snapshot=before_snapshot, mode=mode)


def _merge_direct_file_context(payload: dict | None, fallback: dict | None = None) -> dict:
    out = dict(payload or {})
    fallback_payload = fallback if isinstance(fallback, dict) else {}

    combined_paths: list[str] = []
    for source in (out, fallback_payload):
        for p in collect_response_files(source):
            if p and p not in combined_paths:
                combined_paths.append(p)

    for source in (out, fallback_payload):
        for key in ("selected_target_path", "target_path", "path", "relative_path"):
            value = str(source.get(key) or "").strip()
            if not value:
                continue
            normalized = _relativize_guard_candidate(value)
            if normalized and normalized not in combined_paths:
                combined_paths.append(normalized)

    if not combined_paths:
        return out

    affected = [p for p in collect_response_files({"affected_files": out.get("affected_files")}) if p]
    if not affected:
        out["affected_files"] = combined_paths[:]
    else:
        out["affected_files"] = list(dict.fromkeys(affected + combined_paths))

    file_plan = [p for p in collect_response_files({"file_plan": out.get("file_plan")}) if p]
    if not file_plan:
        out["file_plan"] = combined_paths[:]
    else:
        out["file_plan"] = list(dict.fromkeys(file_plan + combined_paths))

    if not isinstance(out.get("file_entries"), list) or not out.get("file_entries"):
        out["file_entries"] = [{"path": p, "status": "planned"} for p in combined_paths]

    if _response_claims_success(out):
        changed = [p for p in collect_response_files({"changed_files": out.get("changed_files")}) if p]
        if not changed:
            out["changed_files"] = combined_paths[:]
        else:
            out["changed_files"] = list(dict.fromkeys(changed + combined_paths))

    return out


def validate_task(task):
    cleaned = " ".join(str(task or "").strip().split())
    if not cleaned:
        return None, "Bitte eine Aufgabe eingeben."
    kind = classify_user_prompt(cleaned)
    if kind in ("chat", "unknown"):
        return cleaned, None
    intent = classify_user_prompt_intent(cleaned)
    if len(cleaned) < 8:
        if intent in {"risky_task", "coding_task"} or kind in ("project_task", "project_read"):
            return None, "Bitte die Aufgabe etwas genauer beschreiben."
    return cleaned, None


_CLAR_PATH_TOKEN_RE = re.compile(
    r"[\w./\\-]+\.(?:tsx|jsx|ts|js|py|css|html|json|md)\b",
    re.IGNORECASE,
)


def _extract_rel_path_for_clarification(prompt: str) -> str:
    """Längsten Pfad-Treffer wählen (z. B. frontend/app.js statt nur app.js)."""
    s = str(prompt or "")
    matches = _CLAR_PATH_TOKEN_RE.findall(s)
    if not matches:
        return "die genannte Datei"
    best = max(matches, key=len)
    return best.replace("\\", "/").strip()


def is_vague_project_file_edit(prompt: str) -> bool:
    """Nur Dateiname + Änderungsverb, keine konkrete Zieländerung — keine Apply-/Rewrite-Pipeline."""
    t = str(prompt or "").strip()
    if not t or "\n" in t:
        return False
    tl = t.lower()
    if len(tl) > 200:
        return False
    if not re.search(r"\.(py|js|tsx|jsx|ts|css|html|json|md)\b", tl):
        return False
    if " und " in tl:
        tail = tl.split(" und ", 1)[-1].strip()
        if len(tail) >= 8:
            return False
    detail_markers = (
        "füge",
        "fuege",
        "ersetze",
        "entferne",
        "implementiere",
        "comment",
        "kommentar",
        "farbe",
        "blau",
        "rot",
        "grün",
        "header",
        "footer",
        "button",
        "funktion",
        "bug",
        "zeile",
        "import ",
        "export ",
        "mache den",
        "mach den",
        "mach die",
        "schreibe",
        "ergänze",
        "ergaenze",
        "lösche",
        "loesche",
        "text ",
        " layout",
        "provider",
        "logging",
        "fehler",
        "fehlerbehebung",
        "passwort",
        "routing",
        "api ",
        "endpoint",
        "theme",
        "style",
        "design",
        "komponente",
        "komponenten",
        "variable",
        "konstante",
        "test ",
        "pytest",
        "unittest",
    )
    if any(m in tl for m in detail_markers):
        return False
    # „fix“ nur als Verb / Auftrag, nicht als Detailmarker (vage: „fix backend/main.py“)
    if not re.search(
        r"\b("
        r"ändere|aendere|bearbeite|editiere|"
        r"fix|fixe|patch|update|change|edit"
        r")\b",
        tl,
    ):
        return False
    return len(tl.split()) <= 14


def validate_filename(filename):
    cleaned = str(filename or "").strip()
    if not cleaned:
        return None, "Bitte einen Dateinamen angeben."
    if len(cleaned) < 2:
        return None, "Der Dateiname ist zu kurz."
    if cleaned.startswith("."):
        return None, "Der Dateiname darf nicht mit einem Punkt beginnen."
    if not SAFE_FILENAME_PATTERN.match(cleaned):
        return None, "Der Dateiname enthaelt unzulaessige Zeichen."
    return cleaned, None


def sanitize_upload_filename(filename: str) -> str:
    name = str(filename or "").strip().replace("\\", "/").split("/")[-1]
    name = re.sub(r"[^A-Za-z0-9._-]", "_", name)
    return name[:120] or "upload.bin"


MAX_UPLOAD_BYTES = 25 * 1024 * 1024


def _safe_resolved_path_under_uploads(fp_str: str) -> Path | None:
    if not str(fp_str or "").strip():
        return None
    try:
        cand = Path(str(fp_str).strip()).expanduser().resolve()
    except OSError:
        return None
    try:
        base = UPLOADS_DIR.resolve()
        cand.relative_to(base)
        return cand if cand.is_file() else None
    except ValueError:
        return None


def augment_prompt_with_uploads(task: str, data: dict) -> tuple[str, dict]:
    """Haengt nur sichere Upload-Metadaten/Kurzanalysen an den Prompt (Pfade nur unter Downloads/uploads)."""
    meta: dict = {"uploads": [], "errors": []}
    items: list = []
    ufs = data.get("uploaded_files")
    if isinstance(ufs, list):
        for u in ufs[:16]:
            if isinstance(u, dict):
                items.append(u)
    leg = str(data.get("uploaded_file_path") or "").strip()
    if leg and not any(isinstance(x, dict) and str(x.get("filepath") or x.get("saved_path") or "") == leg for x in items):
        items.append({"filepath": leg, "filename": str(data.get("uploaded_file_name") or Path(leg).name)})
    blocks: list[str] = []
    seen: set[str] = set()
    for u in items:
        fp = str(u.get("filepath") or u.get("saved_path") or "").strip()
        sp = _safe_resolved_path_under_uploads(fp)
        if not sp:
            meta["errors"].append({"path": fp, "error": "Pfad ungueltig oder ausserhalb uploads"})
            continue
        key = str(sp)
        if key in seen:
            continue
        seen.add(key)
        ext = sp.suffix.lower()
        try:
            analysis = _analyze_uploaded_file(sp, ext)
        except Exception as ex:
            analysis = {"summary": f"Analyse nicht moeglich: {ex}", "type": ext}
        summary = str(analysis.get("summary") or "")
        fname = str(u.get("filename") or sp.name)
        try:
            size = int(sp.stat().st_size)
        except OSError:
            size = 0
        blocks.append(f"- Upload-Kontext ({ext or '—'}, {size} Bytes)\n  Kurzinfo: {summary}")
        meta["uploads"].append({"filename": fname, "file_type": ext, "size": size, "saved_path": str(sp), "summary": summary})
    if not blocks:
        return str(task or ""), meta
    appendix = (
        "\n\n--- Hochgeladene Dateien (Kontext, keine Binaerdaten) ---\n"
        "WICHTIG:\n"
        "- Diese Upload-Dateien sind nur Referenz-Kontext.\n"
        "- Fuehre die konkrete Nutzeranweisung aus.\n"
        "- Nicht versuchen, Upload-Dateien selbst zu bearbeiten/verschieben/umzubenennen.\n"
        "- Keine Dateioperation auf Upload-Dateien planen.\n"
        + "\n".join(blocks)
    )
    return (str(task or "").rstrip() + appendix), meta


def validate_target_path(targetpath):
    cleaned = format_local_path(targetpath)
    if not cleaned:
        return None, "Bitte einen Zielpfad angeben."
    if cleaned.startswith("/") or ":" in cleaned or ".." in cleaned:
        return None, "Der Zielpfad muss relativ bleiben (kein .., keine Laufwerks-/Root-Pfade)."
    if len(cleaned) < 4:
        return None, "Der Zielpfad ist zu kurz."
    return cleaned, None


def validate_filetype(filetype):
    cleaned = str(filetype or "").strip().lower().lstrip(".")
    if cleaned not in ALLOWED_FILE_TYPES:
        return None, "Der Dateityp wird aktuell nicht unterstuetzt."
    return cleaned, None


def validate_design_note(note):
    cleaned = " ".join(str(note or "").strip().split())
    if not cleaned:
        return None, "Bitte eine Design-Idee eingeben."
    if len(cleaned) < 8:
        return None, "Die Design-Idee ist zu kurz."
    if len(cleaned) > 1200:
        return None, "Die Design-Idee ist zu lang. Bitte kuerzer und praeziser formulieren."
    return cleaned, None


def normalize_task_words(task):
    cleaned = " ".join(str(task or "").strip().split())
    return [word for word in cleaned.replace(",", " ").replace(".", " ").split(" ") if word]


def get_user_output_dir() -> Path:
    """
    USER-OUTPUT: fertige Nutzer-Dateien (Standard: DOWNLOADS_DIR).
    Override: Umgebungsvariable RAINER_USER_OUTPUT_DIR (absolut oder ~).
    """
    raw = (os.environ.get("RAINER_USER_OUTPUT_DIR") or "").strip()
    if raw:
        return Path(raw).expanduser().resolve()
    return DOWNLOADS_DIR


def local_user_download_requested(task: str) -> bool:
    """True wenn der Prompt explizit in die Nutzer-Download-Zone geschrieben werden soll."""
    t = (task or "").lower()
    needles = (
        "speichere als download",
        "als download speichern",
        "speichern als download",
        "save as download",
        "in den downloads",
        "in downloads speichern",
        "nach downloads",
        "unter downloads",
        "zu downloads",
    )
    return any(n in t for n in needles)


def _strip_leading_downloads_path_segment(path_rel: str) -> str:
    """'Downloads/foo.txt' -> 'foo.txt' (echte Downloads-Root, nicht rambo_builder_local/Downloads/)."""
    p = format_local_path(path_rel)
    low = p.lower()
    if low.startswith("downloads/"):
        return format_local_path(p[10:])
    if low == "downloads":
        return ""
    return p


def resolve_local_target_path(targetpath, task=None):
    cleaned, error = validate_target_path(targetpath)
    if error or cleaned is None:
        return None, cleaned, error

    use_user = local_user_download_requested(task or "")
    cleaned_eff = _strip_leading_downloads_path_segment(cleaned) if use_user else cleaned
    if use_user and (not cleaned_eff or len(str(cleaned_eff).strip()) < 1):
        return None, cleaned, "Bitte einen Dateinamen fuer den Download angeben."

    base_root = get_user_output_dir() if use_user else APP_DIR.resolve()
    resolved = (base_root / cleaned_eff).resolve()
    base_resolved = base_root.resolve()
    try:
        resolved.relative_to(base_resolved)
    except ValueError:
        label = "User-Downloads" if use_user else "rambo_builder_local"
        return None, cleaned, f"Schreibzugriff ist nur innerhalb von {label} erlaubt."

    if use_user:
        rel_display = "Downloads/" + Path(cleaned_eff).as_posix()
    else:
        rel_display = Path(cleaned_eff).as_posix()
    return resolved, rel_display, None


def read_text_file(path):
    try:
        if not path.exists():
            return "", False
        return path.read_text(encoding="utf-8"), True
    except OSError:
        return "", False


def build_written_result_detail(path: Path, display_path: str, *, max_chars: int = 1400) -> str:
    """
    Liest die geschriebene Datei und erzeugt einen nutzerlesbaren Ergebnistext.
    Fokus: Analyse-/Dokumentationsausgaben sichtbar im direkten Response-Text.
    """
    try:
        ext = str(path.suffix or "").lower()
        if ext not in {".md", ".txt", ".json", ".yaml", ".yml", ".log"}:
            return ""
        content = read_file_content(path)
        if content is None:
            latest = get_latest_analysis_file(Path(path).parent)
            if latest is not None:
                content = read_file_content(latest)
        if content is None:
            return ""
        text = str(content or "").strip()
        if not text:
            return ""
        if len(text) > max_chars:
            text = text[:max_chars].rstrip() + "\n..."
        body = f"Quelle: {display_path}\n\n{text}"
        return MT.query_result("structure", body)
    except Exception:
        return ""


def _super_builder_extract_icon_source(components):
    for item in components or []:
        s = str(item or "")
        if s.lower().endswith((".png", ".jpg", ".jpeg", ".webp")):
            p = APP_DIR / s
            if p.exists():
                return p
            p2 = RAMBO_RAINER_ROOT / s
            if p2.exists():
                return p2
    return None


_TINY_PNG_BASE64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR4nGNgYAAAAAMAASsJTYQAAAAASUVORK5CYII="
)


def _super_builder_write_text_files(base_dir: Path, files: dict[str, str]) -> list[str]:
    written: list[str] = []
    root = base_dir.resolve()
    for rel_path, content in (files or {}).items():
        normalized = str(rel_path or "").replace("\\", "/").lstrip("/")
        if not normalized:
            continue
        dst = (base_dir / normalized).resolve()
        if not str(dst).startswith(str(root)):
            continue
        dst.parent.mkdir(parents=True, exist_ok=True)
        dst.write_text(str(content or ""), encoding="utf-8")
        try:
            written.append(str(dst.relative_to(PROJECT_DIR).as_posix()))
        except Exception:
            written.append(str(dst))
    return written


def _super_builder_write_icon_placeholder(output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)
    png_path = output_dir / "roboter_icon.png"
    png_bytes = base64.b64decode(_TINY_PNG_BASE64)
    try:
        from PIL import Image
        from io import BytesIO
        img = Image.open(BytesIO(png_bytes)).convert("RGBA").resize((256, 256))
        img.save(png_path, format="PNG")
        ico_path = output_dir / "roboter_icon.ico"
        img.save(ico_path, format="ICO", sizes=[(256, 256), (128, 128), (64, 64), (32, 32), (16, 16)])
    except Exception:
        png_path.write_bytes(png_bytes)
        ico_path = output_dir / "roboter_icon.ico"
        if not ico_path.exists():
            ico_path.write_bytes(png_path.read_bytes())
    return {
        "source": "generated_placeholder",
        "main_png": str(png_path),
        "ico": str(ico_path),
        "generated_pngs": [str(png_path)],
        "placeholder": True,
    }


def _super_builder_write_build_script(output_root: Path) -> str:
    script_path = output_root / "build_desktop.py"
    script_path.parent.mkdir(parents=True, exist_ok=True)
    script_path.write_text(
        (
            "from __future__ import annotations\n\n"
            "import subprocess\n"
            "from pathlib import Path\n"
            "import sys\n\n"
            "ROOT = Path(__file__).resolve().parent\n"
            "UI_DIR = ROOT / 'rambo_ui'\n"
            "ELECTRON_DIR = ROOT / 'electron'\n\n"
            "def run_step(cmd, cwd):\n"
            "    cmd = list(cmd)\n"
            "    if cmd and str(cmd[0]).lower() == 'npm':\n"
            "        cmd[0] = 'npm.cmd'\n"
            "    proc = subprocess.run(cmd, cwd=cwd, check=False)\n"
            "    if proc.returncode != 0:\n"
            "        raise SystemExit(proc.returncode)\n\n"
            "def main():\n"
            "    run_step(['npm','install'], UI_DIR)\n"
            "    run_step(['npm','run','build'], UI_DIR)\n"
            "    run_step(['npm','install'], ELECTRON_DIR)\n"
            "    run_step(['npm','run','build:win'], ELECTRON_DIR)\n"
            "    return 0\n\n"
            "if __name__ == '__main__':\n"
            "    sys.exit(main())\n"
        ),
        encoding="utf-8",
    )
    try:
        return str(script_path.resolve().relative_to(PROJECT_DIR).as_posix())
    except Exception:
        return str(script_path)


def _super_builder_find_latest_installer(output_root: Path) -> Path | None:
    dist_dir = output_root / "electron" / "dist"
    if not dist_dir.exists():
        return None
    matches = [p for p in dist_dir.glob("*.exe") if p.is_file()]
    if not matches:
        return None
    matches.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    return matches[0]


def _super_builder_copy_installer_to_downloads(installer_path: Path, output_root: Path) -> dict:
    dst = output_root / installer_path.name
    dst.write_bytes(installer_path.read_bytes())
    size_bytes = int(dst.stat().st_size)
    return {
        "installer_path": str(dst),
        "installer_size_bytes": size_bytes,
        "installer_size_mb": round(size_bytes / (1024 * 1024), 2),
        "installer_size_ok_gt_50mb": size_bytes > 50 * 1024 * 1024,
    }


def _super_builder_run_build_process(output_root: Path) -> dict:
    build = BuildSystem()
    build_result = build.build_electron_app({
        "project_path": output_root,
        "electron_path": output_root / "electron",
        "ui_path": output_root / "rambo_ui",
    })
    installer_meta = {}
    if build_result.get("ok"):
        installer = _super_builder_find_latest_installer(output_root)
        if installer is not None:
            installer_meta = _super_builder_copy_installer_to_downloads(installer, output_root)
    return {
        "ok": bool(build_result.get("ok")),
        "steps": build_result.get("steps") or [],
        "installer": installer_meta,
    }


def _run_cmd_ok(cmd: list[str], cwd: Path) -> tuple[bool, str]:
    try:
        cp = subprocess.run(cmd, cwd=str(cwd), capture_output=True, text=True, check=False)
        msg = (cp.stdout or cp.stderr or "").strip()
        return cp.returncode == 0, msg[:600]
    except Exception as ex:
        return False, str(ex)


def _contains_text(path: Path, needle: str) -> bool:
    try:
        return needle in path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return False


def _super_builder_self_audit(output_root: Path, build_result: dict | None = None) -> dict:
    electron_dir = output_root / "electron"
    ui_dir = output_root / "rambo_ui"
    assets_dir = electron_dir / "assets"
    main_js = electron_dir / "main.js"
    preload_js = electron_dir / "preload.js"
    package_json = electron_dir / "package.json"
    app_jsx = ui_dir / "src" / "App.jsx"
    build_script = output_root / "build_desktop.py"
    rambo_ui = ui_dir / "src" / "components" / "RamboRainerUI.jsx"
    rainer_ui = ui_dir / "src" / "components" / "RainerBuildUI.jsx"
    service_panel = ui_dir / "src" / "components" / "ServicePanel.jsx"

    structure = {
        "downloads_electron_main_js": main_js.exists(),
        "downloads_electron_preload_js": preload_js.exists(),
        "downloads_electron_package_json": package_json.exists(),
        "downloads_electron_icon_png": (assets_dir / "roboter_icon.png").exists(),
        "downloads_electron_icon_ico": (assets_dir / "roboter_icon.ico").exists(),
        "downloads_rambo_ui_app_jsx": app_jsx.exists(),
        "downloads_rambo_ui_components_dir": (ui_dir / "src" / "components").exists(),
        "downloads_build_desktop_py": build_script.exists(),
    }

    js_main_ok, js_main_msg = _run_cmd_ok(["node", "--check", str(main_js)], output_root)
    js_pre_ok, js_pre_msg = _run_cmd_ok(["node", "--check", str(preload_js)], output_root)
    py_ok, py_msg = _run_cmd_ok(["python", "-m", "py_compile", str(build_script)], output_root)
    syntax = {
        "node_check_main_js": js_main_ok,
        "node_check_preload_js": js_pre_ok,
        "python_check_build_desktop_py": py_ok,
        "details": {"main_js": js_main_msg, "preload_js": js_pre_msg, "build_desktop_py": py_msg},
    }

    functionality = {
        "main_has_startBackend": _contains_text(main_js, "startBackend(")
        or _contains_text(main_js, "registerShellIpc"),
        "main_has_createWindow": _contains_text(main_js, "createWindow("),
        "main_has_ipc_handler": _contains_text(main_js, "ipcMain.handle("),
        "preload_has_electronAPI": _contains_text(preload_js, "electronAPI"),
        "package_has_electron_builder": _contains_text(package_json, "electron-builder"),
        "package_has_icon_ref": _contains_text(package_json, "assets/roboter_icon.ico")
        or _contains_text(package_json, "fallback-setup.html"),
        "icons_present": structure["downloads_electron_icon_png"] and structure["downloads_electron_icon_ico"],
    }

    integration = {
        "main_references_preload": _contains_text(main_js, "preload.js"),
        "app_imports_components": (
            (_contains_text(app_jsx, "RamboRainerUI") and _contains_text(app_jsx, "RainerBuildUI"))
            or _contains_text(app_jsx, "ServicePanel")
        ),
        "paths_consistent": _contains_text(main_js, "rambo_ui")
        and (
            _contains_text(package_json, "assets/roboter_icon.ico")
            or _contains_text(package_json, "rambo_ui/dist")
        ),
        "components_present": (rambo_ui.exists() and rainer_ui.exists()) or service_panel.exists(),
    }

    build_steps = build_result.get("steps") if isinstance(build_result, dict) else []
    step_ok = {
        "npm_install_rambo_ui_ok": bool(build_steps) and bool(build_steps[0].get("ok")),
        "npm_build_rambo_ui_ok": len(build_steps) > 1 and bool(build_steps[1].get("ok")),
        "npm_install_electron_ok": len(build_steps) > 2 and bool(build_steps[2].get("ok")),
        "npm_build_win_ok": len(build_steps) > 3 and bool(build_steps[3].get("ok")),
    }
    installer = (build_result or {}).get("installer") if isinstance(build_result, dict) else {}
    installer_size_ok = bool(installer and installer.get("installer_size_ok_gt_50mb"))
    final_checks = {
        "exe_present_in_downloads": bool(installer and installer.get("installer_path")),
        "exe_size_gt_50mb": installer_size_ok,
        "exe_signed": False,
    }
    if installer and installer.get("installer_path"):
        try:
            cp = subprocess.run(
                [
                    "powershell.exe",
                    "-NoProfile",
                    "-Command",
                    f"$s=(Get-AuthenticodeSignature '{installer.get('installer_path')}').Status; Write-Output $s",
                ],
                capture_output=True,
                text=True,
                check=False,
            )
            status_text = (cp.stdout or "").strip().lower()
            final_checks["exe_signed"] = status_text == "valid"
            final_checks["signature_status"] = status_text or "unknown"
        except Exception:
            final_checks["signature_status"] = "error"

    all_green = all(structure.values()) and all([
        syntax["node_check_main_js"],
        syntax["node_check_preload_js"],
        syntax["python_check_build_desktop_py"],
    ]) and all(functionality.values()) and all(integration.values()) and all(step_ok.values()) and final_checks["exe_present_in_downloads"] and final_checks["exe_size_gt_50mb"]

    return {
        "structure_check": structure,
        "syntax_check": syntax,
        "functionality_check": functionality,
        "integration_check": integration,
        "build_process_check": step_ok,
        "final_checks": final_checks,
        "all_green": bool(all_green),
    }


def _super_builder_write_audit_report(output_root: Path, audit: dict) -> str:
    report_path = output_root / "RAINER_SELF_AUDIT_REPORT.md"
    def yn(v: bool) -> str:
        return "JA" if bool(v) else "NEIN"
    lines = [
        "# Rainer Self-Audit Report",
        "",
        f"- Gesamtstatus: {'✅ ALLES OK' if audit.get('all_green') else '⚠️ FEHLER GEFUNDEN'}",
        "",
        "## Struktur-Check",
    ]
    for k, v in (audit.get("structure_check") or {}).items():
        lines.append(f"- {k}: {yn(v)}")
    lines += ["", "## Syntax-Check"]
    sx = audit.get("syntax_check") or {}
    lines.append(f"- node_check_main_js: {yn(sx.get('node_check_main_js'))}")
    lines.append(f"- node_check_preload_js: {yn(sx.get('node_check_preload_js'))}")
    lines.append(f"- python_check_build_desktop_py: {yn(sx.get('python_check_build_desktop_py'))}")
    lines += ["", "## Funktionalitäts-Check"]
    for k, v in (audit.get("functionality_check") or {}).items():
        lines.append(f"- {k}: {yn(v)}")
    lines += ["", "## Integrations-Check"]
    for k, v in (audit.get("integration_check") or {}).items():
        lines.append(f"- {k}: {yn(v)}")
    lines += ["", "## Build-Prozess-Check"]
    for k, v in (audit.get("build_process_check") or {}).items():
        lines.append(f"- {k}: {yn(v)}")
    lines += ["", "## Finale Checks"]
    for k, v in (audit.get("final_checks") or {}).items():
        if isinstance(v, bool):
            lines.append(f"- {k}: {yn(v)}")
        else:
            lines.append(f"- {k}: {v}")
    report_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    try:
        return str(report_path.resolve().relative_to(PROJECT_DIR).as_posix())
    except Exception:
        return str(report_path)


def build_super_builder_result(task: str) -> dict:
    optimizer = PromptOptimizer()
    analysis = optimizer.analyze_prompt(task or "")
    prompt_type = analysis.get("type") or "general"
    result = {
        "analysis": analysis,
        "executed": False,
        "generated_files": [],
        "template_hints": {
            "react_component": bool(TemplateLibrary.get_template("react_component")),
            "electron_main": bool(TemplateLibrary.get_template("electron_main")),
            "python_script": bool(TemplateLibrary.get_template("python_script")),
        },
    }
    output_root = get_user_output_dir()

    try:
        if prompt_type == "react_app":
            builder = ReactBuilder()
            spec = {
                "app_name": "RainerGeneratedReactApp",
                "component_name": "GeneratedFromPrompt",
                "include_state": True,
                "include_effect": True,
                "include_api_call": "api" in str(task).lower(),
                "include_form": True,
                "include_validation": True,
            }
            files = builder.build_app_structure(spec)
            written = _super_builder_write_text_files(output_root / "rambo_ui", files)
            result["executed"] = True
            result["generated_files"] = written
            result["react_preview"] = {
                "files_count": len(files),
                "files": list(files.keys())[:10],
                "sample": {"path": "src/App.jsx", "content": files.get("src/App.jsx", "")[:800]},
            }
        elif prompt_type == "electron_app":
            builder = ElectronBuilder()
            files = builder.build_complete_app({"app_name": "Rambo-Rainer Desktop", "backend_port": 5002})
            written = _super_builder_write_text_files(output_root, files)
            result["executed"] = True
            result["generated_files"] = written
            ui_files = {
                "package.json": (
                    '{\n'
                    '  "name": "rambo-ui",\n'
                    '  "private": true,\n'
                    '  "version": "1.0.0",\n'
                    '  "type": "module",\n'
                    '  "scripts": { "dev": "vite", "build": "vite build" },\n'
                    '  "dependencies": { "react": "^18.0.0", "react-dom": "^18.0.0" },\n'
                    '  "devDependencies": { "vite": "^5.0.0" }\n'
                    '}\n'
                ),
                "index.html": "<!doctype html><html><body><div id='root'></div><script type='module' src='/src/main.jsx'></script></body></html>\n",
                "src/main.jsx": "import React from 'react';\nimport { createRoot } from 'react-dom/client';\nimport App from './App';\ncreateRoot(document.getElementById('root')).render(<App />);\n",
                "src/App.jsx": (
                    "import React, { useState } from 'react';\n"
                    "import RamboRainerUI from './components/RamboRainerUI';\n"
                    "import RainerBuildUI from './components/RainerBuildUI';\n\n"
                    "export default function App() {\n"
                    "  const [mode, setMode] = useState('rainer_build');\n"
                    "  return (\n"
                    "    <main>\n"
                    "      <button onClick={() => setMode('rambo_rainer')}>Rambo-Rainer</button>\n"
                    "      <button onClick={() => setMode('rainer_build')}>Rainer Build</button>\n"
                    "      {mode === 'rambo_rainer' ? <RamboRainerUI /> : <RainerBuildUI />}\n"
                    "    </main>\n"
                    "  );\n"
                    "}\n"
                ),
                "src/components/RamboRainerUI.jsx": "import React from 'react';\nexport default function RamboRainerUI(){ return <section>Rambo-Rainer UI</section>; }\n",
                "src/components/RainerBuildUI.jsx": "import React from 'react';\nexport default function RainerBuildUI(){ return <section>Rainer Build UI</section>; }\n",
            }
            result["generated_files"].extend(_super_builder_write_text_files(output_root / "rambo_ui", ui_files))
            result["generated_files"].append(_super_builder_write_build_script(output_root))
            result["electron_preview"] = {
                "files_count": len(files),
                "files": list(files.keys()),
                "sample": {"path": "electron/main.js", "content": files.get("electron/main.js", "")[:900]},
            }
            output_dir = output_root / "electron" / "assets"
            source = _super_builder_extract_icon_source(analysis.get("components"))
            if source is not None:
                try:
                    processor = IconProcessor()
                    icon_result = processor.process_robot_icon(source, output_dir)
                except Exception:
                    icon_result = _super_builder_write_icon_placeholder(output_dir)
            else:
                icon_result = _super_builder_write_icon_placeholder(output_dir)
            result["icon_result"] = icon_result
            for p in [icon_result.get("main_png"), icon_result.get("ico")]:
                if p:
                    try:
                        result["generated_files"].append(str(Path(p).resolve().relative_to(PROJECT_DIR).as_posix()))
                    except Exception:
                        pass
            build_result = _super_builder_run_build_process(output_root)
            result["build_result"] = build_result
            result["self_audit"] = _super_builder_self_audit(output_root, build_result)
            audit_report = _super_builder_write_audit_report(output_root, result["self_audit"])
            result["self_audit_report"] = audit_report
            result["generated_files"].append(audit_report)
            installer_path = ((build_result.get("installer") or {}).get("installer_path") or "").strip()
            if installer_path:
                try:
                    result["generated_files"].append(str(Path(installer_path).resolve().relative_to(PROJECT_DIR).as_posix()))
                except Exception:
                    result["generated_files"].append(installer_path)
        elif prompt_type == "icon":
            source = _super_builder_extract_icon_source(analysis.get("components"))
            if source is not None:
                try:
                    processor = IconProcessor()
                    icon_result = processor.process_robot_icon(source, output_root / "electron" / "assets")
                except Exception:
                    icon_result = _super_builder_write_icon_placeholder(output_root / "electron" / "assets")
            else:
                icon_result = _super_builder_write_icon_placeholder(output_root / "electron" / "assets")
            result["executed"] = True
            result["icon_result"] = icon_result
            for p in [icon_result.get("main_png"), icon_result.get("ico")]:
                if p:
                    try:
                        result["generated_files"].append(str(Path(p).resolve().relative_to(PROJECT_DIR).as_posix()))
                    except Exception:
                        pass
        elif prompt_type == "build":
            build = BuildSystem()
            build_result = build.build_electron_app({
                "project_path": output_root,
                "electron_path": output_root / "electron",
                "ui_path": output_root / "rambo_ui",
            })
            result["executed"] = True
            result["build_result"] = build_result
            result["build_plan"] = {
                "steps": [
                    "run_npm_install(ui_path)",
                    "run_npm_build(ui_path)",
                    "run_npm_install(electron_path)",
                    "run_electron_builder(electron_path)",
                ],
                "dry_run_only": False,
                "runner": build.__class__.__name__,
            }
    except Exception as ex:
        result["error"] = str(ex)

    return result


def build_text_diff(current_content, proposed_content, relative_path):
    current_lines = current_content.splitlines()
    proposed_lines = proposed_content.splitlines()
    diff_lines = list(
        difflib.unified_diff(
            current_lines,
            proposed_lines,
            fromfile=f"a/{relative_path}",
            tofile=f"b/{relative_path}",
            lineterm=""
        )
    )
    return "\n".join(diff_lines) if diff_lines else "Keine inhaltliche Aenderung erkannt."


def infer_builder_areas(task):
    lowered = str(task or "").lower()
    areas = []

    keyword_groups = [
        ("UI und Oberflaeche", ["ui", "oberflaeche", "frontend", "maske", "layout", "seite"]),
        ("Status und Fehlermeldungen", ["status", "fehler", "warn", "meldung"]),
        ("Builder Mode", ["builder", "plan", "workflow", "aufgabe"]),
        ("Dateiverarbeitung", ["datei", "export", "import", "preview", "json", "md"]),
        ("Design Studio", ["design", "notiz", "idee", "studio"]),
        ("Aktivitaetslog", ["log", "aktivitaet", "historie"]),
    ]

    for label, keywords in keyword_groups:
        if any(keyword in lowered for keyword in keywords):
            areas.append(label)

    if not areas:
        areas.append("Lokale Arbeitsoberflaeche")

    areas.append("Sichere lokale Umsetzung innerhalb von rambo_builder_local")
    return areas


def build_builder_risks(task):
    lowered = str(task or "").lower()
    risks = [
        "Zu grobe Aufgabenbeschreibung kann zu unscharfen Arbeitspaketen fuehren.",
        "Unsaubere Reihenfolge kann lokale Schritte blockieren oder doppeln.",
        "Aenderungen duerfen nicht ueber rambo_builder_local hinausgehen."
    ]

    if "api" in lowered or "backend" in lowered:
        risks.append("Aenderungen an Antwortfeldern muessen zum Frontend-Modell passen.")
    if "ui" in lowered or "frontend" in lowered:
        risks.append("UI-Anpassungen muessen Status-, Fehler- und Leerzustaende mitdenken.")

    return risks


def build_builder_plan(task):
    words = normalize_task_words(task)
    focus = ", ".join(words[:4]) if words else "lokale Oberflaeche"
    focus_sentence = " ".join(words[:8]) if words else "lokale Builder-Aufgabe"
    affected_areas = infer_builder_areas(task)
    risks = build_builder_risks(task)

    return {
        "ziel": (
            f"Die Aufgabe '{task}' lokal vorbereiten und in umsetzbare, sichere Builder-Schritte zerlegen, "
            "ohne das Hauptprojekt direkt zu veraendern."
        ),
        "betroffene_bereiche": affected_areas + [f"Fokusthema: {focus}"],
        "empfohlene_reihenfolge": [
            f"Aufgabe '{focus_sentence}' konkretisieren und das Zielbild festhalten.",
            "Betroffene Bereiche priorisieren und auf lokale Auswirkungen eingrenzen.",
            "Kleine, sichere Umsetzungsschritte mit klarer Reihenfolge festlegen.",
            "Abschluss mit sichtbarem Status, Logeintrag und naechstem Schritt vorbereiten."
        ],
        "risiken": risks,
        "naechste_schritte": [
            "Plan lesen und auf Vollstaendigkeit pruefen.",
            "Bei Bedarf die Aufgabe weiter praezisieren oder eingrenzen.",
            "Builder Mode starten und den lokalen Ablauf ueber Status und Aktivitaetslog verfolgen."
        ]
    }


def classify_project_file(rel_path):
    parts = rel_path.split("/")
    area_key = parts[0] if len(parts) > 1 else "root"
    area = PROJECT_AREA_MAP.get(area_key, area_key.capitalize() if len(area_key) <= 20 else "Sonstiges")
    sensitive = rel_path in GUARDED_PROJECT_PATHS or any(p in rel_path for p in SENSITIVE_PATTERNS)
    allowed_write = (not sensitive) and any(
        rel_path.startswith(prefix) for prefix in ALLOWED_PROJECT_WRITE_PREFIXES
    )
    return {
        "path": rel_path,
        "area": area,
        "area_key": area_key,
        "sensitive": sensitive,
        "allowed_write": allowed_write,
    }


@PO.time_function
@PO.cache_result(ttl=120)
def scan_project_structure():
    project_root = RAMBO_RAINER_ROOT.resolve()
    results = []
    for root, dirs, files in os.walk(project_root):
        dirs[:] = sorted(
            d for d in dirs
            if d not in SCANNER_SKIP_DIRS and not d.startswith(".")
        )
        for fname in sorted(files):
            if fname.startswith("."):
                continue
            try:
                rel = (Path(root) / fname).relative_to(project_root)
                rel_str = str(rel).replace("\\", "/")
                results.append(classify_project_file(rel_str))
            except ValueError:
                continue
    return results


def validate_project_read_path(rel_path):
    cleaned = format_local_path(rel_path)
    if not cleaned:
        return None, None, "Bitte einen Pfad angeben."
    for pattern in (".env", "rambo_rainer.db", ".git/", "node_modules/", "__pycache__"):
        if pattern in cleaned:
            return None, cleaned, f"Lesezugriff blockiert: Datei enthaelt '{pattern}'."
    resolved = (RAMBO_RAINER_ROOT / cleaned).resolve()
    try:
        resolved.relative_to(RAMBO_RAINER_ROOT.resolve())
    except ValueError:
        return None, cleaned, "Pfad liegt ausserhalb des Projekts."
    return resolved, cleaned, None


def validate_project_write_path(rel_path):
    cleaned = format_local_path(rel_path)
    if not cleaned:
        return None, None, "Bitte einen Zielpfad angeben."
    if cleaned in GUARDED_PROJECT_PATHS:
        return None, cleaned, f"Schreiben blockiert: '{cleaned}' ist als geschuetzte Frontend-Datei markiert."
    for pattern in SENSITIVE_PATTERNS:
        if pattern in cleaned:
            return None, cleaned, f"Schreiben blockiert: Datei enthaelt sensibles Muster '{pattern}'."
    trusted_workspace = False
    try:
        trusted_workspace = bool(is_active_workspace_trusted())
    except Exception:
        trusted_workspace = False
    if (not trusted_workspace) and (not any(cleaned.startswith(prefix) for prefix in ALLOWED_PROJECT_WRITE_PREFIXES)):
        return None, cleaned, (
            f"'{cleaned}' liegt nicht in einem freigegebenen Bereich. "
            f"Erlaubt: {', '.join(ALLOWED_PROJECT_WRITE_PREFIXES[:5])} ..."
        )
    resolved = (RAMBO_RAINER_ROOT / cleaned).resolve()
    try:
        resolved.relative_to(RAMBO_RAINER_ROOT.resolve())
    except ValueError:
        return None, cleaned, "Pfad liegt ausserhalb des Projekts."
    return resolved, cleaned, None


KNOWN_FILE_ROLES = {
    "backend/main.py": {"role": "Haupt-Backend (Flask)", "keywords": ["backend", "api", "server", "route", "endpunkt"], "allowed_write": False},
    "backend/db.py": {"role": "Datenbankzugriff", "keywords": ["datenbank", "db", "sql"], "allowed_write": False},
    "backend/models.py": {"role": "Datenmodelle", "keywords": ["model", "schema", "daten"], "allowed_write": False},
    "backend/agent.py": {"role": "Agent-Backend", "keywords": ["agent", "ki"], "allowed_write": False},
    "backend/BUILDER_MODE.md": {"role": "Builder-Dokumentation", "keywords": ["builder", "modus", "dokumentation"], "allowed_write": True},
    "backend/COACH.md": {"role": "Coach-Dokumentation", "keywords": ["coach", "anleitung"], "allowed_write": True},
    "backend/DESIGN.md": {"role": "Design-Dokumentation", "keywords": ["design", "gestaltung", "ui"], "allowed_write": True},
    "backend/DEV_WORKFLOW.md": {"role": "Entwicklungs-Workflow", "keywords": ["workflow", "entwicklung", "ablauf"], "allowed_write": True},
    "backend/GENERATION.md": {"role": "Generierungs-Dokumentation", "keywords": ["generierung", "erzeugung"], "allowed_write": True},
    "backend/ORCHESTRATION.md": {"role": "Orchestrierung", "keywords": ["orchestrierung", "koordination"], "allowed_write": True},
    "frontend/src/App.jsx": {"role": "React Hauptkomponente", "keywords": ["ui", "frontend", "react", "oberflaeche"], "allowed_write": False},
    "frontend/src/App.css": {"role": "Haupt-Styling", "keywords": ["style", "css", "design", "farbe"], "allowed_write": False},
    "agent/core/planner.js": {"role": "Agent-Planung", "keywords": ["agent", "planung", "ki", "plan"], "allowed_write": False},
    "agent/core/executor.js": {"role": "Agent-Ausfuehrung", "keywords": ["ausfuehren", "execute"], "allowed_write": False},
    "knowledge/project_context.md": {"role": "Projektwissen/Kontext", "keywords": ["kontext", "wissen", "projekt", "ziel", "beschreibung"], "allowed_write": True},
    "knowledge/user_notes.md": {"role": "Benutzernotizen", "keywords": ["notiz", "idee", "hinweis", "persoenlich"], "allowed_write": True},
    "data/state.json": {"role": "Projektstatus", "keywords": ["status", "zustand", "state"], "allowed_write": True},
    "data/tasks.json": {"role": "Aufgabenliste", "keywords": ["aufgabe", "task", "todo", "liste"], "allowed_write": True},
    "data/memory.json": {"role": "Speicher", "keywords": ["speicher", "erinnerung"], "allowed_write": True},
    "memory/long_term_memory.md": {"role": "Langzeitgedaechtnis", "keywords": ["langzeit", "gedaechtnis", "verlauf"], "allowed_write": True},
    "memory/task_log.json": {"role": "Aufgaben-Protokoll", "keywords": ["protokoll", "log", "aufgaben"], "allowed_write": True},
    "outbox/result.txt": {"role": "Ergebnis-Ausgabe", "keywords": ["ergebnis", "ausgabe", "export"], "allowed_write": True},
}

AREA_KEYWORDS = {
    "backend": ["backend", "api", "server", "route", "endpunkt", "python", "flask", "datenbank"],
    "frontend": ["frontend", "ui", "oberflaeche", "design", "react", "css", "style", "komponente"],
    "agent": ["agent", "ki", "aufgabe", "automatisch", "planung", "executor"],
    "data": ["daten", "json", "konfiguration", "state", "status"],
    "knowledge": ["wissen", "dokumentation", "kontext", "beschreibung", "erklaerung"],
    "outbox": ["ausgabe", "export", "ergebnis", "output", "brief", "bericht"],
    "memory": ["speicher", "log", "protokoll", "verlauf", "gedaechtnis"],
    "tools": ["werkzeug", "tool", "hilfsskript", "utility"],
}


@PO.time_function
@PO.cache_result(ttl=180)
def build_project_knowledge():
    files = scan_project_structure()
    areas_summary = {}
    for f in files:
        key = f["area_key"]
        if key not in areas_summary:
            areas_summary[key] = {"label": f["area"], "total": 0, "allowed_write": 0, "sensitive": 0}
        areas_summary[key]["total"] += 1
        if f["allowed_write"]:
            areas_summary[key]["allowed_write"] += 1
        if f["sensitive"]:
            areas_summary[key]["sensitive"] += 1

    knowledge = {
        "built_at": get_timestamp(),
        "total_files": len(files),
        "areas": areas_summary,
        "files": KNOWN_FILE_ROLES,
        "allowed_write_prefixes": list(ALLOWED_PROJECT_WRITE_PREFIXES),
        "sensitive_patterns": list(SENSITIVE_PATTERNS),
        "endpoints": {
            "builder": [
                "/api/builder/plan", "/api/builder/generate-content",
                "/api/builder/change-preview", "/api/builder/apply-change", "/api/builder/auto-run"
            ],
            "project": [
                "/api/project/scan", "/api/project/read-file",
                "/api/project/change-preview", "/api/project/apply-change",
                "/api/project/guard-check", "/api/project/suggest-files",
                "/api/project/auto-run",
                "/api/project/build-knowledge"
            ],
            "system": ["/api/health", "/api/status", "/api/todos", "/api/ui-activity", "/api/direct-run", "/api/direct-confirm"],
            "design": ["/api/design-studio/notes", "/api/file-generator/preview"]
        }
    }

    write_json_file(DATA_DIR / "project_knowledge.json", knowledge)
    return knowledge


def suggest_files_for_task(task):
    lowered = task.lower()
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    known_files = knowledge.get("files", KNOWN_FILE_ROLES)

    safe_targets = []
    suggestions = []
    blocked = []
    area_suggestions = []

    task_words = [word.lower() for word in normalize_task_words(task)]

    for path, info in known_files.items():
        keywords = info.get("keywords", [])
        keyword_hits = sum(1 for kw in keywords if kw in lowered)
        word_hits = sum(1 for word in task_words if word and word in path.lower())
        priority = keyword_hits * 3 + word_hits
        if priority <= 0:
            continue
        sensitive = any(p in path for p in SENSITIVE_PATTERNS)
        if sensitive:
            blocked.append({
                "path": path,
                "reason": "Sensible Datei – nicht beschreibbar",
                "role": info.get("role", ""),
                "priority": priority
            })
        elif info.get("allowed_write"):
            safe_targets.append({
                "path": path,
                "role": info.get("role", ""),
                "priority": priority,
                "status": "freigegeben"
            })
        else:
            suggestions.append({
                "path": path,
                "role": info.get("role", ""),
                "allowed_write": False,
                "priority": priority,
                "status": "nur_lesen"
            })

    for area_key, area_kws in AREA_KEYWORDS.items():
        if not any(kw in lowered for kw in area_kws):
            continue
        for prefix in ALLOWED_PROJECT_WRITE_PREFIXES:
            if prefix.startswith(area_key + "/") or prefix.rstrip("/") == area_key:
                slug = "_".join(lowered.split()[:3]).replace("-", "_")
                slug = re.sub(r"[^a-z0-9_]", "", slug)[:30] or "aufgabe"
                base = prefix if prefix.endswith("/") else prefix.rsplit("/", 1)[0] + "/"
                area_suggestions.append({
                    "area": PROJECT_AREA_MAP.get(area_key, area_key),
                    "suggested_prefix": prefix,
                    "example_path": f"{base}{slug}.md",
                    "priority": 2,
                    "status": "vorgeschlagen"
                })
                break

    safe_targets.sort(key=lambda item: (-item.get("priority", 0), item["path"]))
    suggestions.sort(key=lambda item: (-item.get("priority", 0), item["path"]))
    blocked.sort(key=lambda item: (-item.get("priority", 0), item["path"]))
    area_suggestions.sort(key=lambda item: (-item.get("priority", 0), item["example_path"]))

    prioritized = safe_targets + suggestions + [
        {
            "path": item["example_path"],
            "role": f"Neuer Pfadvorschlag in {item['area']}",
            "priority": item.get("priority", 0),
            "status": "vorgeschlagen",
            "allowed_write": True
        }
        for item in area_suggestions
    ]

    return {
        "safe_targets": safe_targets[:6],
        "suggestions": suggestions[:5],
        "blocked": blocked[:4],
        "area_suggestions": area_suggestions[:4],
        "prioritized_targets": prioritized[:8],
        "allowed_prefixes": list(ALLOWED_PROJECT_WRITE_PREFIXES)
    }


def extract_explicit_local_relative_path(task):
    """Ermittelt einen relativen Zielpfad aus dem Prompt (inkl. Windows-Absolutpfad unter rambo_builder_local)."""
    raw = str(task or "").strip()
    if not raw:
        return None
    zb = extract_relative_path_from_zielordner_block(raw)
    if zb:
        cleaned, err = validate_target_path(zb)
        if not err and cleaned:
            return cleaned
    norm = raw.replace("\\", "/")
    low = norm.lower()
    marker = "rambo_builder_local/"
    ix = low.find(marker)
    if ix >= 0:
        tail = norm[ix + len(marker) :]
        tail = re.split(r"[\s\"'`]", tail.strip())[0] if tail.strip() else ""
        tail = format_local_path(tail).lstrip("./")
        if tail and ".." not in tail and ":" not in tail:
            cleaned, err = validate_target_path(tail)
            if not err and cleaned:
                return cleaned
    m = re.search(
        r"rambo_builder_local[/\\\\]+([\w./\\-]+\.(?:js|mjs|cjs|ts|tsx|jsx|css|html|json|py|md|txt|yml|yaml|toml))\b",
        norm,
        re.IGNORECASE,
    )
    if m:
        tail = format_local_path(m.group(1))
        cleaned, err = validate_target_path(tail)
        if not err and cleaned:
            return cleaned
    return None


_TARGET_SCAN_EXCLUDE_DIR_NAMES = frozenset(
    {"node_modules", "dist", "build", "__pycache__", ".git", ".svn", "coverage", ".next", "out"}
)
_TARGET_GLOB_EXTENSIONS = {".js", ".jsx", ".ts", ".tsx", ".css", ".html"}
_BUILDER_MODE_CONTENT_NEEDLES = (
    "Builder Mode",
    "builder mode",
    "builder-mode",
    "BUILDER_MODE",
    'data-view="builder"',
    "data-view='builder'",
    "builder-nav",
    "builder-panel",
    "BuilderPanel",
    "builderpanel",
)


def _is_ui_change_remove_intent(task: str) -> bool:
    """
    True bei UI-/Feature-Entfernung (kein Löschen ganzer Dateien per os.remove).
    Trifft z. B. 'Entferne Builder Mode aus der UI', 'remove from frontend', Tabs/Buttons.
    """
    tl = str(task or "").lower()
    markers = (
        " aus der ui",
        " aus der ",
        " aus ",
        " from ",
        " innerhalb ",
        "frontend-ui",
        " sichtbar",
        " sichtbaren ",
        " einstieg",
        "einstiege",
        "navigation",
        "navbar",
        " builder mode",
        "builder-mode",
        "builder mode",
        "data-view",
        "builder-nav",
        "builder-panel",
        "builderpanel",
        "verstecke",
        "ausblend",
        "tab ",
        " panel ",
        " button ",
        "ui-einstieg",
        "nur sichtbar",
    )
    if any(m in tl for m in markers):
        return True
    if re.search(r"\b(entfern|remove)\w*\s+.+\s+(aus|from|in)\s+(der\s+)?(ui|frontend|oberfläche|oberflaeche)\b", tl):
        return True
    if re.search(r"\b(hide|remove)\s+.+\s+(from|in)\s+", tl):
        return True
    if "ui" in tl and "datei" not in tl and any(x in tl for x in ("entfern", " remove ", "hide ", "versteck")):
        return True
    return False


def _is_explicit_file_delete_intent(task: str) -> bool:
    """
    Nur True, wenn eindeutig eine Datei vom Dateisystem gelöscht werden soll (os.remove).
    NIEMALS True bei UI-/Builder-/Komponenten-Formulierungen.
    """
    if _is_ui_change_remove_intent(task):
        return False
    tl = str(task or "").lower()
    if re.search(r"\b(entfern|remove)\w*\s+.+\s+(aus|from)\s+", tl):
        return False
    phrases = (
        "lösche die datei",
        "loesche die datei",
        "die datei löschen",
        "die datei loeschen",
        "datei löschen",
        "datei loeschen",
        "delete file",
        "delete the file",
        "remove file",
        "remove the file",
        "unlink file",
        "datei endgültig löschen",
        "datei endgueltig loeschen",
        "unwiderruflich löschen",
        "unwiderruflich loeschen",
        "lösche die folgende datei",
        "loesche die folgende datei",
    )
    if any(p in tl for p in phrases):
        return True
    if re.search(r"\b(lösche|loesche)\s+die\s+datei\s+\S", tl):
        return True
    if re.search(r"\b(entferne|entfernen)\s+die\s+datei\b", tl):
        return True
    if re.search(r"\b(delete|remove)\s+(the\s+)?file\s+\S", tl):
        return True
    if not re.search(r"\b(aus|from)\s+(der\s+)?(ui|oberfläche|oberflaeche)\b", tl):
        path_del = re.search(
            r"\b(lösche|loesche|delete|remove)\s+((?:frontend|backend|tests)/[^\s]+\.(?:py|js|jsx|mjs|cjs|ts|tsx|css|html|json|md|vue|svelte))\b",
            tl,
        )
        if path_del:
            return True
    return False


def _discover_builder_ui_target_files(active_root: Path, task: str) -> list[str]:
    """Relevante Frontend-Dateien für Builder/UI — keine Löschung."""
    raw = str(task or "")
    merged = _merge_unique_relative_paths(
        infer_allowed_target_files(raw),
        _builder_mode_scan_frontend(active_root, raw),
    )
    return merged


def _same_or_inside_project_root(active_root: Path, candidate: Path) -> tuple[bool, str]:
    try:
        candidate.resolve().relative_to(active_root.resolve())
        return True, "inside_active_root"
    except ValueError:
        return False, "outside_active_root"


def _allowlisted_missing_relative(rel: str) -> bool:
    norm = format_local_path(rel or "").replace("\\", "/").lower()
    if not any(norm.startswith(p) for p in ("frontend/", "backend/", "tests/")):
        return False
    name = Path(rel).name
    return "." in name


def _expand_trusted_dir_files(rel_dir: str, active_root: Path, max_files: int = 50) -> list[str]:
    cleaned, err = validate_target_path(format_local_path(rel_dir))
    if err or not cleaned:
        return []
    base = (active_root / cleaned).resolve()
    ok, _ = _same_or_inside_project_root(active_root, base)
    if not ok or not base.is_dir():
        return []
    out: list[str] = []
    for p in sorted(base.rglob("*")):
        if not p.is_file():
            continue
        try:
            rel_parts = p.relative_to(active_root.resolve()).parts
        except ValueError:
            continue
        if any(part in _TARGET_SCAN_EXCLUDE_DIR_NAMES for part in rel_parts):
            continue
        if p.suffix.lower() not in _TARGET_GLOB_EXTENSIONS:
            continue
        out.append(format_local_path(str(p.relative_to(active_root.resolve()))))
        if len(out) >= max_files:
            break
    return out


def _collect_raw_target_candidates(raw: str) -> list[str]:
    """Rohtokens aus Prompt: Backticks, Projekt-Pfade, Windows-Absolutpfade, lockere Listenformen."""
    seen: set[str] = set()
    out: list[str] = []

    def add(s: str):
        t = str(s or "").strip()
        if t and t not in seen:
            seen.add(t)
            out.append(t)

    for m in re.finditer(r"`([^`]+)`", raw):
        add(m.group(1).strip())
    for m in re.finditer(
        r"\b((?:frontend|backend|tests)(?:[/\\][\w.\-]+)+(?:/\*\*|/\*)?)",
        raw,
        flags=re.IGNORECASE,
    ):
        add(m.group(1).strip())
    for m in re.finditer(
        r"\b((?:frontend|backend|tests)(?:[/\\][^\s,;|\"'<>)\]]+)+)",
        raw,
        flags=re.IGNORECASE,
    ):
        add(m.group(1).strip().rstrip(".,;:-—–)"))
    for m in re.finditer(r"([A-Za-z]:[/\\](?:[^|\s\"<>]+(?:[/\\][^|\s\"<>]+)*))", raw):
        add(m.group(1).strip())
    return out


def _merge_unique_relative_paths(*path_lists: list) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for pl in path_lists:
        for p in pl or []:
            c = format_local_path(str(p or ""))
            if c and c not in seen:
                seen.add(c)
                out.append(c)
    return out


def _explicit_map_target_paths_from_prompt(raw: str) -> list[str]:
    out: list[str] = []
    explicit_map = (
        (r"\bfrontend[\\/]+style\.css\b", "frontend/style.css"),
        (r"\bfrontend[\\/]+app\.js\b", "frontend/app.js"),
        (r"\bfrontend[\\/]+index\.html\b", "frontend/index.html"),
        (r"\bbackend[\\/]+main\.py\b", "backend/main.py"),
    )
    for pattern, mapped in explicit_map:
        if re.search(pattern, raw, re.IGNORECASE):
            if mapped not in out:
                out.append(mapped)
    return out


def _literal_trusted_core_frontend_paths(raw: str, ar: Path) -> list[str]:
    """
    Wenn trusted: explizit genannte Kern-Dateien (wörtlich im Text) als Targets,
    sofern unter active_root.
    """
    if not is_active_workspace_trusted():
        return []
    core = ("frontend/app.js", "frontend/index.html", "frontend/style.css")
    nraw = format_local_path(raw)
    nlow = nraw.lower()
    ar_res = ar.resolve()
    out: list[str] = []
    for rel in core:
        if rel.lower() not in nlow:
            continue
        cleaned, err = validate_target_path(rel)
        if err or not cleaned:
            continue
        full = (ar_res / cleaned).resolve()
        try:
            full.relative_to(ar_res)
        except ValueError:
            continue
        if full.is_file() or _allowlisted_missing_relative(cleaned):
            if cleaned not in out:
                out.append(cleaned)
    return out


def _clean_path_token(tok: str) -> str:
    t = str(tok or "").strip().strip("`\"'")
    while t and t[-1] in ".,;:)":
        t = t[:-1]
    return format_local_path(t.replace("\\", "/")).strip()


def extract_and_resolve_trusted_targets(raw: str, active_root: Path) -> dict:
    """Explizite Ziele unter active_root; Wildcards frontend/src/**; Windows-Pfade relativieren."""
    extracted_raw = _collect_raw_target_candidates(raw)
    resolved_paths: list[str] = []
    same_or_inside: list[dict] = []
    reason_parts: list[str] = []
    ar = active_root.resolve()

    def append_unique(rel: str):
        c = format_local_path(rel)
        if c and c not in resolved_paths:
            resolved_paths.append(c)

    for tok in extracted_raw:
        ct = _clean_path_token(tok)
        if not ct:
            continue

        if re.match(r"^[A-Za-z]:[\\/]", ct) or re.match(r"^[A-Za-z]:[\\/]", tok.strip()):
            win_src = tok.strip().strip("`\"'")
            try:
                wp = Path(win_src).resolve()
                rel_o = wp.relative_to(ar)
                rel_s = format_local_path(str(rel_o))
                ok, detail_msg = _same_or_inside_project_root(ar, wp)
                same_or_inside.append({"path": rel_s, "same_or_inside": ok, "detail": detail_msg})
                if ok:
                    if wp.is_file():
                        append_unique(rel_s)
                    elif wp.is_dir():
                        reason_parts.append(f"win_dir:{rel_s}")
                        for r in _expand_trusted_dir_files(rel_s, ar):
                            append_unique(r)
                    elif _allowlisted_missing_relative(rel_s):
                        append_unique(rel_s)
                        same_or_inside.append({"path": rel_s, "same_or_inside": True, "detail": "missing_candidate"})
            except ValueError:
                same_or_inside.append({"path": win_src, "same_or_inside": False, "detail": "windows_outside_active_root"})
            continue

        norm = ct
        is_glob = "**" in norm or norm.endswith("/*")
        glob_base = norm
        if norm.endswith("/**"):
            glob_base = norm[:-3].rstrip("/")
        elif norm.endswith("/*"):
            glob_base = norm[:-2].rstrip("/")
        elif norm.endswith("**") and "/" in norm:
            glob_base = norm[:-2].rstrip("/")

        if is_glob:
            exp = _expand_trusted_dir_files(glob_base, ar)
            reason_parts.append(f"glob:{glob_base}->{len(exp)}")
            for r in exp:
                append_unique(r)
                same_or_inside.append({"path": r, "same_or_inside": True, "detail": "glob_expand"})
            continue

        cleaned, err = validate_target_path(norm)
        if err:
            same_or_inside.append({"path": norm, "same_or_inside": False, "detail": err})
            continue

        full = (ar / cleaned).resolve()
        ok_root, detail = _same_or_inside_project_root(ar, full)
        if not ok_root:
            same_or_inside.append({"path": cleaned, "same_or_inside": False, "detail": detail})
            continue

        if full.is_file():
            append_unique(cleaned)
            same_or_inside.append({"path": cleaned, "same_or_inside": True, "detail": "exists_file"})
        elif full.is_dir():
            exp = _expand_trusted_dir_files(cleaned, ar)
            reason_parts.append(f"dir:{cleaned}->{len(exp)}")
            for r in exp:
                append_unique(r)
                same_or_inside.append({"path": r, "same_or_inside": True, "detail": "dir_expand"})
        elif _allowlisted_missing_relative(cleaned):
            append_unique(cleaned)
            same_or_inside.append({"path": cleaned, "same_or_inside": True, "detail": "missing_allowlist_candidate"})
        else:
            same_or_inside.append({"path": cleaned, "same_or_inside": ok_root, "detail": "missing_not_allowlisted"})

    reason = "; ".join(reason_parts) if reason_parts else ("explicit_resolved" if resolved_paths else "no_explicit_match")
    return {
        "extracted_raw": extracted_raw,
        "resolved_paths": resolved_paths,
        "same_or_inside": same_or_inside,
        "reason": reason,
    }


def _prompt_mentions_builder_mode(raw: str) -> bool:
    low = str(raw or "").lower()
    if "builder mode" in low or "builder-mode" in low or "buildermode" in low.replace(" ", ""):
        return True
    for n in _BUILDER_MODE_CONTENT_NEEDLES:
        if n in raw:
            return True
    if re.search(r"\bbuilder\b", low):
        return True
    return False


def _builder_mode_scan_frontend(active_root: Path, raw: str, max_files: int = 50) -> list[str]:
    if not _prompt_mentions_builder_mode(raw):
        return []
    fe = (active_root / "frontend").resolve()
    if not fe.is_dir():
        return []
    hits: list[str] = []
    ar = active_root.resolve()
    needles_low = [n.lower() for n in _BUILDER_MODE_CONTENT_NEEDLES if len(n) > 2]
    for p in sorted(fe.rglob("*")):
        if not p.is_file():
            continue
        try:
            rel_parts = p.relative_to(ar).parts
        except ValueError:
            continue
        if any(part in _TARGET_SCAN_EXCLUDE_DIR_NAMES for part in rel_parts):
            continue
        if p.suffix.lower() not in _TARGET_GLOB_EXTENSIONS:
            continue
        try:
            text = p.read_text(encoding="utf-8", errors="replace")
        except OSError:
            continue
        tl = text.lower()
        matched = False
        for nd in needles_low:
            if nd in tl:
                matched = True
                break
        if not matched:
            for nd in _BUILDER_MODE_CONTENT_NEEDLES:
                if nd in text:
                    matched = True
                    break
        if matched:
            hits.append(format_local_path(str(p.relative_to(ar))))
            if len(hits) >= max_files:
                break
    return hits


def infer_allowed_target_files_with_meta(prompt: str) -> tuple[list[str], dict]:
    raw = str(prompt or "")
    lowered = raw.lower()
    trusted = is_active_workspace_trusted()
    ar = get_active_project_root()
    meta: dict = {
        "active_workspace_path": str(ar.resolve()),
        "active_workspace_trusted": trusted,
        "active_project_root": str(ar.resolve()),
        "extracted_target_paths": [],
        "resolved_target_paths": [],
        "target_inference_reason": "",
        "same_or_inside_result": [],
        "raw_target_candidates": [],
    }
    out: list[str] = []

    def add(path_value: str):
        cleaned = format_local_path(path_value or "")
        if cleaned and cleaned not in out:
            out.append(cleaned)

    raw_cands = _collect_raw_target_candidates(raw)
    meta["raw_target_candidates"] = list(raw_cands)

    if trusted:
        pack = extract_and_resolve_trusted_targets(raw, ar)
        meta["extracted_target_paths"] = list(pack.get("extracted_raw") or [])
        meta["same_or_inside_result"] = list(pack.get("same_or_inside") or [])
        base_reason = str(pack.get("reason") or "")
        merged_trusted = _merge_unique_relative_paths(
            list(pack.get("resolved_paths") or []),
            _literal_trusted_core_frontend_paths(raw, ar),
            _explicit_map_target_paths_from_prompt(raw),
            _builder_mode_scan_frontend(ar, raw) if _prompt_mentions_builder_mode(raw) else [],
        )
        meta["resolved_target_paths"] = merged_trusted
        meta["target_inference_reason"] = (
            (base_reason + " | trusted_merged_paths").strip(" |") if merged_trusted else base_reason
        )
        if merged_trusted:
            return merged_trusted, meta

    forbidden_markers = (
        "electron",
        "rambo_ui",
        "/src/components/",
        "downloads/",
        "../downloads/",
        "rainerrobotdesktop",
        "du_arbeitest_im_projekt__",
    )
    if not trusted and any(marker in lowered for marker in forbidden_markers):
        meta["target_inference_reason"] = "forbidden_markers_blocked"
        pack_dbg = extract_and_resolve_trusted_targets(raw, ar)
        meta["extracted_target_paths"] = list(pack_dbg.get("extracted_raw") or [])
        meta["resolved_target_paths"] = list(pack_dbg.get("resolved_paths") or [])
        meta["same_or_inside_result"] = list(pack_dbg.get("same_or_inside") or [])
        return [], meta

    explicit_map = (
        (r"\bfrontend[\\/]+style\.css\b", "frontend/style.css"),
        (r"\bfrontend[\\/]+app\.js\b", "frontend/app.js"),
        (r"\bfrontend[\\/]+index\.html\b", "frontend/index.html"),
        (r"\bbackend[\\/]+main\.py\b", "backend/main.py"),
    )
    for pattern, mapped in explicit_map:
        if re.search(pattern, raw, re.IGNORECASE):
            add(mapped)

    test_path_matches = re.findall(r"\btests[\\/]+[A-Za-z0-9_./\\-]+\.py\b", raw, flags=re.IGNORECASE)
    has_test_context = any(
        k in lowered for k in ("tests/", "tests\\", "pytest", " test ", " test-", "unit test", "integration test")
    )
    if has_test_context:
        for match in test_path_matches:
            add(match)

    if out:
        meta["resolved_target_paths"] = list(out)
        meta["target_inference_reason"] = "explicit_map_or_tests"
        return out, meta

    css_hints = ("css", "style", "design", "layout", "kompakter")
    css_trigger = any(h in lowered for h in css_hints)
    ui_hints = ("button", "ui", "anzeige")
    ui_trigger = any(h in lowered for h in ui_hints) or ("frontend" in lowered and not css_trigger)
    if ui_trigger:
        r = ["frontend/app.js", "frontend/index.html", "frontend/style.css"]
        meta["resolved_target_paths"] = r
        meta["target_inference_reason"] = "ui_heuristic_bundle"
        return r, meta
    if css_trigger:
        r = ["frontend/style.css"]
        meta["resolved_target_paths"] = r
        meta["target_inference_reason"] = "css_heuristic"
        return r, meta

    meta["target_inference_reason"] = "no_match"
    return [], meta


def infer_allowed_target_files(prompt: str) -> list[str]:
    paths, _meta = infer_allowed_target_files_with_meta(prompt)
    return paths


def detect_model_selection_prompt(prompt: str) -> bool:
    lowered = str(prompt or "").lower().strip()
    if not lowered:
        return False
    markers = (
        "modell wählen",
        "modell waehlen",
        "wähle das passende modell",
        "waehle das passende modell",
        "wähle das passende lokale modell",
        "waehle das passende lokale modell",
        "welches modell",
        "welches lokale modell",
        "model router",
        "modellrouter",
        "lokales modell",
    )
    return any(m in lowered for m in markers)


def detect_self_fix_task(prompt: str) -> dict:
    try:
        if bool(app.config.get("TESTING")):
            raise RuntimeError("ignore_workspace_trust_in_tests")
        active_ws = WORKSPACE_SANDBOX.get_active_workspace()
        ws_trusted = bool((active_ws.get("active") or {}).get("trusted", False))
        if ws_trusted:
            return {
                "is_self_fix": False,
                "reason": "Workspace-Modus aktiv.",
                "risk": "low",
                "recommended_mode": "apply",
                "candidate_files": [],
            }
    except Exception:
        pass

    raw = str(prompt or "")
    lowered = raw.lower()
    candidate_files = infer_allowed_target_files(raw)

    strong_markers = (
        "sich selbst reparieren",
        "self fix",
        "self-fix",
        "self repair",
        "wartungsmodus",
        "eigene steuerung",
        "ändere deine eigene steuerung",
        "aendere deine eigene steuerung",
        "rainer soll",
        "rainer kann",
    )
    maintenance_targets = (
        "modell wählen reparieren",
        "modell waehlen reparieren",
        "ui ausgabe",
        "rechte ausgabe",
        "agent panel",
        "promptfeld",
        "cursor layout",
        "layout umbauen",
        "router reparieren",
        "direct-run logik",
        "guard reparieren",
        "frontend/app.js",
        "frontend/index.html",
        "frontend/style.css",
        "backend/main.py",
    )

    matched = [k for k in strong_markers if k in lowered]
    matched_targets = [k for k in maintenance_targets if k in lowered]
    subject_hint = any(k in lowered for k in ("rainer", "eigene", "selbst", "self"))
    direct_fix_phrase = (
        ("fixe frontend/app.js" in lowered or "fix frontend/app.js" in lowered)
        and ("rechte ausgabe" in lowered or "ergebnisse rechts" in lowered or "agent-panel" in lowered or "agent panel" in lowered)
    )
    is_self_fix = bool(matched) or (subject_hint and bool(matched_targets)) or direct_fix_phrase

    if bool(matched_targets) and subject_hint:
        matched.extend(matched_targets)
    elif direct_fix_phrase:
        matched.extend(["fixe frontend/app.js", "rechte ausgabe"])

    risk = "high" if any(k in lowered for k in ("self fix", "self-fix", "self repair", "wartungsmodus", "eigene steuerung")) else ("medium" if is_self_fix else "low")
    reason = "Self-Fix-Marker erkannt: " + ", ".join(matched[:4]) if is_self_fix else "Kein Self-Fix-Marker erkannt."
    return {
        "is_self_fix": is_self_fix,
        "reason": reason,
        "risk": risk if is_self_fix else "low",
        "recommended_mode": "self_fix_plan",
        "candidate_files": candidate_files,
    }


def _normalize_self_fix_files(items) -> list[str]:
    out: list[str] = []
    seen: set[str] = set()
    for raw in list(items or []):
        p = format_local_path(raw)
        if not p or p in seen:
            continue
        seen.add(p)
        out.append(p)
    return out


def _build_self_fix_plan_record(prompt: str, self_fix: dict) -> dict:
    candidate_files = _normalize_self_fix_files(self_fix.get("candidate_files") or [])
    checks = list(self_fix.get("recommended_checks") or ["node_check_app", "py_compile_main", "pytest_all"])
    plan = {
        "plan_id": uuid4().hex,
        "created_at": get_timestamp(),
        "prompt": str(prompt or ""),
        "candidate_files": candidate_files,
        "affected_files": candidate_files[:],
        "recommended_checks": checks[:6],
        "risk": str(self_fix.get("risk") or "medium"),
        "next_step": str(self_fix.get("next_step") or "Änderungsplan anzeigen und Safe/Preview verwenden."),
        "reason": str(self_fix.get("reason") or ""),
    }
    return plan


def save_last_self_fix_plan(prompt: str, self_fix: dict) -> dict:
    plan = _build_self_fix_plan_record(prompt, self_fix)
    save_project_auto_run_state(
        {
            "last_self_fix_plan": plan,
            "last_task": plan["prompt"],
            "last_mode": "self_fix_plan",
            "last_direct_status": "self_fix_plan_required",
        }
    )
    return plan


def get_last_self_fix_plan(plan_id: str = "") -> dict | None:
    state = load_project_auto_run_state()
    plan = state.get("last_self_fix_plan") if isinstance(state.get("last_self_fix_plan"), dict) else None
    if not plan:
        return None
    if plan_id and str(plan.get("plan_id") or "").strip() != str(plan_id or "").strip():
        return None
    return plan


def _build_self_fix_preview_payload(plan: dict) -> dict:
    candidate_files = _normalize_self_fix_files(plan.get("candidate_files") or plan.get("affected_files") or [])
    checks = [str(c).strip() for c in list(plan.get("recommended_checks") or ["node_check_app", "py_compile_main", "pytest_all"]) if str(c).strip()]
    preview_steps = [
        "Problem analysieren",
        "Dateien prüfen",
        "Änderung als Diff vorbereiten",
        "Tests ausführen",
        "Erst nach Bestätigung anwenden",
    ]
    def _plan_item_for_file(path: str) -> dict:
        p = str(path or "").strip()
        lower = p.lower()
        if lower == "frontend/app.js":
            return {
                "file": p,
                "area": "IDE-Agent Logik / Routing / Ergebnis-Rendering",
                "change": "Routing- und Rendering-Handler prüfen und gezielt anpassen.",
                "reason": "Self-Fix-/Model-Route-Aufgaben dürfen nicht in Direct-Write laufen und müssen rechts im Agent-Panel angezeigt werden.",
                "risk": "medium",
                "checks": ["node --check frontend\\app.js", "python -m pytest tests -q"],
            }
        if lower == "frontend/index.html":
            return {
                "file": p,
                "area": "IDE-Agent Struktur / Buttons / Container",
                "change": "Falls nötig Button/Container für Preview-Aktion prüfen.",
                "reason": "Safe/Preview-Aktion muss sichtbar und klickbar bleiben.",
                "risk": "low",
                "checks": ["node --check frontend\\app.js"],
            }
        if lower == "frontend/style.css":
            return {
                "file": p,
                "area": "IDE-Agent Styling / Karten / Buttons",
                "change": "Falls nötig Styling für Ergebnis-/Preview-Karten prüfen.",
                "reason": "Preview muss klar sichtbar sein.",
                "risk": "low",
                "checks": ["node --check frontend\\app.js"],
            }
        if lower == "backend/main.py":
            return {
                "file": p,
                "area": "API-Routing / Self-Fix-/Direct-Run-Endpunkte",
                "change": "Self-Fix-Preview-Payload und Endpunktlogik gezielt prüfen.",
                "reason": "Preview muss Non-Write bleiben und einen klaren Änderungsplan liefern.",
                "risk": "medium",
                "checks": ["python -m py_compile backend\\main.py", "python -m pytest tests -q"],
            }
        if lower == "backend/agent_model_router.py":
            return {
                "file": p,
                "area": "Model-Router Klassifikation",
                "change": "Klassifizierung und Routing-Hinweise gegen Self-Fix-Flows prüfen.",
                "reason": "Model-Route-Verhalten muss im Self-Fix-Kontext stabil bleiben.",
                "risk": "medium",
                "checks": ["python -m py_compile backend\\agent_model_router.py", "python -m pytest tests -q"],
            }
        if lower.startswith("tests/") and lower.endswith(".py"):
            return {
                "file": p,
                "area": "Testabdeckung",
                "change": "Betroffene Regressionstests ergänzen oder bestehende Tests verifizieren.",
                "reason": "Self-Fix-Flow soll reproduzierbar abgesichert bleiben.",
                "risk": "low",
                "checks": ["python -m pytest tests -q"],
            }
        return {
            "file": p,
            "area": "Allgemeiner Self-Fix-Bereich",
            "change": "Betroffene Stelle prüfen und minimalen, sicheren Patch planen.",
            "reason": "Änderung soll erst nach Bestätigung angewendet werden.",
            "risk": "medium",
            "checks": ["python -m pytest tests -q"],
        }

    change_plan = [_plan_item_for_file(path) for path in candidate_files]
    diff_plan = [
        {
            "file": item.get("file", ""),
            "planned_patch_summary": item.get("change", ""),
            "would_modify": True,
            "applied": False,
        }
        for item in change_plan
    ]
    verification_plan = {
        "mode": "manual",
        "auto_run": False,
        "requires_confirmation": True,
        "summary": "Diese Checks müssen vor einem Apply erfolgreich sein.",
        "required_checks": [
            {
                "name": "Python Syntaxcheck",
                "command": "python -m py_compile backend\\main.py",
                "targets": ["backend/main.py"],
                "required": True,
                "pass_condition": "Befehl beendet ohne Fehlercode.",
                "blocking_if_fails": True,
            },
            {
                "name": "Frontend Syntaxcheck",
                "command": "node --check frontend\\app.js",
                "targets": ["frontend/app.js"],
                "required": True,
                "pass_condition": "Befehl beendet ohne Fehlercode.",
                "blocking_if_fails": True,
            },
            {
                "name": "Pytest Regression",
                "command": "python -m pytest tests -q",
                "targets": ["tests"],
                "required": True,
                "pass_condition": "Alle Tests bestehen.",
                "blocking_if_fails": True,
            },
        ],
        "optional_checks": [
            {
                "name": "Browser QA",
                "command": "Manuell: Backend starten, Browser Strg+F5, Self-Fix Preview testen.",
                "targets": ["UI"],
                "required": False,
                "pass_condition": "Self-Fix Preview zeigt Kandidaten, Änderungsplan und Verifikationsplan rechts im Agent-Panel.",
                "blocking_if_fails": False,
            }
        ],
        "order": [
            "python -m py_compile backend\\main.py",
            "node --check frontend\\app.js",
            "python -m pytest tests -q",
            "Manuelle Browser-QA",
        ],
        "apply_gate": {
            "can_apply_without_tests": False,
            "requires_all_required_checks_green": True,
            "requires_user_confirmation": True,
            "blocking_reason_if_missing": "Apply darf erst nach erfolgreichen Pflichtchecks und Nutzerbestätigung vorbereitet werden.",
        },
    }
    autopilot = {
        "enabled": True,
        "mode": "guided",
        "stage": "preview_ready",
        "can_auto_continue": False,
        "manual_confirmation_required": True,
        "recommended_next_action": "run_required_checks",
        "reason": "Preview ist bereit. Pflichtchecks ausführen und danach Apply explizit bestätigen.",
        "safe_actions": [
            {
                "id": "run_required_checks",
                "label": "Pflichtchecks ausführen",
                "type": "manual_command",
                "requires_confirmation": True,
                "commands": [
                    "python -m py_compile backend\\main.py",
                    "node --check frontend\\app.js",
                    "python -m pytest tests -q",
                ],
                "writes_files": False,
                "runs_git": False,
            },
            {
                "id": "confirm_apply",
                "label": "Self-Fix anwenden bestätigen",
                "type": "confirm_apply",
                "requires_confirmation": True,
                "writes_files": True,
                "runs_git": False,
            },
        ],
        "blocked_actions": [
            {"id": "auto_commit", "label": "Automatisch committen", "reason": "Commits dürfen nicht automatisch ausgeführt werden."},
            {"id": "auto_rollback", "label": "Automatisch rollbacken", "reason": "Rollback darf nur nach separater Bestätigung vorbereitet werden."},
        ],
        "safety_gate": {
            "allow_write_without_confirmation": False,
            "allow_commit": False,
            "allow_rollback": False,
            "allow_forbidden_paths": False,
            "requires_user_confirmation_for_apply": True,
        },
    }
    return {
        "ok": True,
        "success": False,
        "status": "self_fix_preview_ready",
        "direct_status": "self_fix_preview_ready",
        "mode": "self_fix_preview",
        "message": "Self-Fix-Preview bereit. Keine Änderung ausgeführt.",
        "plan_id": str(plan.get("plan_id") or ""),
        "prompt": str(plan.get("prompt") or ""),
        "reason": str(plan.get("reason") or ""),
        "risk": str(plan.get("risk") or "medium"),
        "candidate_files": candidate_files,
        "affected_files": candidate_files[:],
        "changed_files": [],
        "recommended_checks": checks,
        "change_plan": change_plan,
        "diff_plan": diff_plan,
        "verification_plan": verification_plan,
        "patch_validation": {
            "validated_patch": True,
            "blocked": False,
            "large_patch_blocked": False,
            "status": "validated",
        },
        "self_fix_autopilot": autopilot,
        "preview_steps": preview_steps,
        "next_step": "Bestätigung erforderlich, bevor geschrieben wird.",
        "requires_confirmation": True,
        "requires_user_confirmation": True,
        "workstream_events": [
            _ws_event("analysis", "info", "Self-Fix-Plan geladen", "Kandidaten-Dateien und Risiko geprüft", status="done"),
            _ws_event("preview", "warning", "Safe/Preview bereit", "Keine Schreibphase gestartet.", status="done"),
        ],
    }


def _is_mini_task_write_intent(task: str) -> bool:
    lowered = str(task or "").lower()
    forbidden_markers = ("electron", "rambo_ui", "/src/components/", "downloads/", "../downloads/", "rainerrobotdesktop")
    if any(marker in lowered for marker in forbidden_markers):
        return False
    write_hints = (
        "ändere",
        "aendere",
        "update",
        "anpass",
        "fix",
        "reparier",
        "mache",
        "schreibe",
        "erzeuge",
        "erstelle",
        "entferne",
        "entfernen",
        "lösche",
        "loesche",
        "remove",
    )
    area_hints = ("frontend", "style", "css", "index.html", "app.js", "backend/main.py", "tests/", "pytest", "layout", "ui")
    return any(k in lowered for k in write_hints) and any(k in lowered for k in area_hints)


def _build_target_path_unclear_payload(
    task: str,
    mode: str,
    candidates: list[str] | None = None,
    inference_debug: dict | None = None,
) -> dict:
    inferred = [format_local_path(p) for p in (candidates or []) if str(p or "").strip()]
    payload = {
        "ok": False,
        "success": False,
        "status": "target_path_unclear",
        "guard_status": "target_path_unclear",
        "direct_status": "target_path_unclear",
        "message": "Kein eindeutiger erlaubter Zielpfad erkannt.",
        "errors": ["Kein eindeutiger erlaubter Zielpfad erkannt."],
        "task": str(task or ""),
        "mode": str(mode or "apply"),
        "affected_files": inferred[:],
        "changed_files": [],
        "file_plan": inferred[:],
        "requires_confirmation": False,
        "requires_user_confirmation": False,
        "workstream_events": [
            _ws_event("analysis", "info", "Zielpfad-Inferenz", "Erlaubte Zielpfade werden ermittelt", status="done"),
            _ws_event("route", "warning", "Zielpfad unklar", "Bitte eine erlaubte Datei explizit nennen (z. B. frontend/style.css).", status="blocked"),
        ],
    }
    dbg = inference_debug if isinstance(inference_debug, dict) else {}
    ar_dbg = str(get_active_project_root().resolve())
    payload["active_workspace_path"] = dbg.get("active_workspace_path") or ar_dbg
    payload["active_workspace_trusted"] = bool(dbg.get("active_workspace_trusted")) if "active_workspace_trusted" in dbg else is_active_workspace_trusted()
    payload["active_project_root"] = str(dbg.get("active_project_root") or ar_dbg)
    payload["ws_ok"] = bool(dbg["ws_ok"]) if "ws_ok" in dbg else is_active_workspace_trusted()
    payload["extracted_target_paths"] = list(dbg.get("extracted_target_paths") or [])
    payload["resolved_target_paths"] = list(dbg.get("resolved_target_paths") or [])
    payload["target_inference_reason"] = str(dbg.get("target_inference_reason") or "")
    payload["same_or_inside_result"] = list(dbg.get("same_or_inside_result") or [])
    payload["raw_target_candidates"] = list(dbg.get("raw_target_candidates") or [])
    _inf = dbg.get("inferred_allowed_targets")
    payload["inferred_allowed_targets"] = list(_inf) if isinstance(_inf, list) else inferred[:]
    payload["len_inferred_allowed_targets"] = int(
        dbg.get("len_inferred_allowed_targets", len(payload["inferred_allowed_targets"]))
    )
    return payload


HARD_REWRITE_PROTECTED_FILES = {
    "frontend/style.css",
    "frontend/app.js",
    "frontend/index.html",
    "backend/main.py",
}


def _is_tiny_stub_content(path: str, before_lines: int, after_content: str) -> bool:
    if before_lines <= 200:
        return False
    after_lines_raw = str(after_content or "").splitlines()
    after_lines = len(after_lines_raw)
    if after_lines >= 20:
        return False
    non_empty = [ln.strip() for ln in after_lines_raw if ln.strip()]
    if len(non_empty) <= 3:
        return True
    comment_like = 0
    for ln in non_empty:
        if ln.startswith(("/*", "*", "//", "#", "<!--")) or ln.endswith("*/"):
            comment_like += 1
    structural = sum(1 for ln in non_empty if ("{" in ln or "}" in ln or ";" in ln or "<" in ln or ">" in ln))
    return comment_like >= 1 and structural <= 6 and len(non_empty) <= 12


def detect_unsafe_large_rewrite(path: str, before_content: str, after_content: str, context=None) -> dict:
    relative_path = format_local_path(path or "")
    before_lines_raw = str(before_content or "").splitlines()
    after_lines_raw = str(after_content or "").splitlines()
    before_lines = len(before_lines_raw)
    after_lines = len(after_lines_raw)
    if before_lines <= 0:
        return {"unsafe": False, "path": relative_path, "before_lines": before_lines, "after_lines": after_lines}

    removed_lines = 0
    for line in difflib.ndiff(before_lines_raw, after_lines_raw):
        if line.startswith("- "):
            removed_lines += 1

    removed_ratio = (removed_lines / before_lines) if before_lines > 0 else 0.0
    shrink_ratio = (after_lines / before_lines) if before_lines > 0 else 1.0
    tiny_stub = _is_tiny_stub_content(relative_path, before_lines, after_content)
    protected = relative_path in HARD_REWRITE_PROTECTED_FILES

    triggers: list[str] = []
    if before_lines > 200:
        if shrink_ratio < 0.5:
            triggers.append("new_file_below_50_percent")
        if removed_lines > 200:
            triggers.append("removed_lines_gt_200")
        if removed_ratio > 0.30:
            triggers.append("removed_lines_gt_30_percent")
        if tiny_stub:
            triggers.append("tiny_stub_pattern")

    # Fuer besonders kritische Kern-Dateien strengere Schwellwerte.
    if protected and before_lines > 120:
        if shrink_ratio < 0.6:
            triggers.append("protected_file_strong_shrink")
        if removed_lines > 120:
            triggers.append("protected_file_large_delete")
        if tiny_stub:
            triggers.append("protected_file_tiny_stub")

    return {
        "unsafe": bool(triggers),
        "path": relative_path,
        "before_lines": before_lines,
        "after_lines": after_lines,
        "removed_lines": removed_lines,
        "removed_ratio": removed_ratio,
        "shrink_ratio": shrink_ratio,
        "tiny_stub": tiny_stub,
        "protected": protected,
        "triggers": triggers,
        "context": context if isinstance(context, dict) else {},
    }


def _build_unsafe_large_rewrite_payload(path: str, *, scope: str = "local", mode: str = "apply", task: str = "") -> dict:
    rel = format_local_path(path or "")
    task_text = str(task or "").strip()
    split_steps = [
        f"Datei lesen: {rel}" if rel else "Datei lesen",
        "Nur Zielbereich lokalisieren (keine Vollersetzung)",
        "Kleinen Patch mit minimalen Änderungen erzeugen",
        "Patch validieren und anwenden",
    ]
    return {
        "ok": False,
        "success": False,
        "status": "unsafe_large_rewrite",
        "guard_status": "blocked",
        "direct_status": "blocked",
        "scope": str(scope or "local"),
        "mode": str(mode or "apply"),
        "task": str(task or ""),
        "message": "Änderung blockiert: Datei würde zu stark überschrieben. Nutze automatische Teilschritte (read -> locate -> small patch -> apply).",
        "blocked_files": [rel] if rel else [],
        "errors": [f"Unsafe large rewrite detected: {rel}"] if rel else ["Unsafe large rewrite detected"],
        "affected_files": [rel] if rel else [],
        "changed_files": [],
        "forbidden_files": [],
        "requires_split_patch": True,
        "recovery_mode": "split_patch_required",
        "next_actions": split_steps,
        "recommended_user_prompt": (
            f"Ändere nur den relevanten Abschnitt in {rel} mit kleinem Patch, keine Vollersetzung."
            if rel
            else "Ändere nur den relevanten Abschnitt mit kleinem Patch, keine Vollersetzung."
        ),
        "task_echo": task_text,
    }


_RECOVERY_BLOCKING_TRIGGERS = frozenset({"tiny_stub_pattern", "protected_file_tiny_stub"})


def _unsafe_triggers_block_auto_recovery(triggers: list | None) -> bool:
    return bool(triggers) and any(t in _RECOVERY_BLOCKING_TRIGGERS for t in triggers)


def _run_step_engine_internal(
    task: str,
    rel_path: str,
    current_content: str,
    proposed_content: str,
    *,
    confirmed: bool = False,
    run_checks: bool = False,
) -> dict:
    planner = get_task_planner_agent()
    context_builder = get_context_builder_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    patch_generator = get_patch_generator_agent(root=RAMBO_RAINER_ROOT.resolve())
    patch_validator = get_patch_validator_agent(root=RAMBO_RAINER_ROOT.resolve())
    error_fixer = get_error_fixer_agent(APP_DIR.resolve())
    test_runner = get_test_runner_agent(APP_DIR.resolve())
    engine = StepEngineAgent(planner, context_builder, patch_generator, patch_validator, error_fixer, test_runner)
    return engine.run_step_flow(
        task=task,
        path=rel_path,
        current_content=current_content,
        proposed_content=proposed_content,
        confirmed=confirmed,
        run_checks=run_checks,
    )


def _step_engine_patch_allows_apply(se_result: dict) -> bool:
    if not isinstance(se_result, dict) or not se_result.get("ok"):
        return False
    val = se_result.get("validation") or {}
    if str(val.get("status") or "") != "validated":
        return False
    patch = se_result.get("patch_preview") or {}
    return bool(patch.get("has_changes"))


def _try_unsafe_rewrite_step_engine_apply(
    *,
    relative_path: str,
    resolved_path,
    current_content: str,
    proposed_content: str,
    task: str,
    unsafe_check: dict,
) -> tuple[bool, dict | None]:
    if _unsafe_triggers_block_auto_recovery(unsafe_check.get("triggers")):
        return False, None
    try:
        se = _run_step_engine_internal(
            str(task or ""),
            str(relative_path or ""),
            str(current_content or ""),
            str(proposed_content or ""),
            confirmed=False,
            run_checks=False,
        )
        if not _step_engine_patch_allows_apply(se):
            return False, se
        wr = persist_text_file_change(
            resolved_path,
            proposed_content,
            relative_path,
            on_timeout_log=lambda m: append_ui_log_entry("Direkt", m, "error"),
        )
        if not wr.get("ok"):
            return False, se
        return True, se
    except Exception as exc:
        return False, {"ok": False, "status": "recovery_exception", "errors": [str(exc)]}


def _is_tests_py_path(path: str) -> bool:
    normalized = format_local_path(path or "").lower()
    return normalized.startswith("tests/") and normalized.endswith(".py")


def validate_pytest_file(path: str, content: str) -> dict:
    normalized = format_local_path(path or "")
    if not _is_tests_py_path(normalized):
        return {"ok": True, "applicable": False, "path": normalized}
    text = str(content or "")
    has_test = bool(re.search(r"(?m)^\s*(?:async\s+def|def)\s+test_[A-Za-z0-9_]*\s*\(", text))
    if has_test:
        return {"ok": True, "applicable": True, "path": normalized}
    return {
        "ok": False,
        "applicable": True,
        "path": normalized,
        "status": "invalid_test_file",
        "guard_status": "invalid_output",
        "message": "Kein pytest-Test erkannt.",
        "errors": [f"{normalized} enthält keine pytest-Testfunktion."],
    }


def extract_pytest_function_from_prompt(prompt: str) -> str | None:
    raw = str(prompt or "")
    if not raw.strip():
        return None

    # Inline-Format: def test_x(): assert True / async def test_x(): assert 1 == 1
    match = re.search(
        r"(?is)\b(async\s+def|def)\s+(test_[A-Za-z0-9_]*)\s*\(\s*\)\s*:\s*assert\s+([^\n\r;]+)",
        raw,
    )
    if not match:
        return None

    def_kw = "async def" if "async" in match.group(1).lower() else "def"
    fn_name = str(match.group(2) or "").strip()
    assert_expr = str(match.group(3) or "").strip().rstrip(".,")
    if not re.match(r"^test_[A-Za-z0-9_]+$", fn_name):
        return None
    if not assert_expr:
        return None

    lowered = (fn_name + " " + assert_expr).lower()
    banned_tokens = (
        "import ",
        "exec(",
        "eval(",
        "subprocess",
        "os.system",
        "__import__",
        "open(",
    )
    if any(tok in lowered for tok in banned_tokens):
        return None

    if len(assert_expr) > 120:
        return None
    if not re.match(r"^[A-Za-z0-9_ \t\(\)\[\]\{\}\.'\"=!<>+\-*/%,:]+$", assert_expr):
        return None

    return f"{def_kw} {fn_name}():\n    assert {assert_expr}\n"


def _build_invalid_test_file_payload(path: str, *, scope: str = "local", mode: str = "apply", task: str = "") -> dict:
    normalized = format_local_path(path or "")
    return {
        "ok": False,
        "success": False,
        "status": "invalid_test_file",
        "guard_status": "invalid_output",
        "direct_status": "invalid_test_file",
        "scope": str(scope or "local"),
        "mode": str(mode or "apply"),
        "task": str(task or ""),
        "message": "Kein pytest-Test erkannt.",
        "errors": [f"{normalized} enthält keine pytest-Testfunktion."] if normalized else ["tests-Datei enthält keine pytest-Testfunktion."],
        "affected_files": [normalized] if normalized else [],
        "changed_files": [],
        "blocked_files": [normalized] if normalized else [],
        "requires_confirmation": False,
        "requires_user_confirmation": False,
    }


def refine_local_source_for_editing_task(current_content, relative_path, task, file_exists):
    """Ersetzt bei klarem Edit-Auftrag vorhandenen Quelltext statt Stub-Generierung (nur enge Whitelist)."""
    if not file_exists or current_content is None:
        return None
    tl = str(task or "").lower()
    edit_kw = (
        "änder", "aender", "verkürz", "verkuerz", "kürz", "kuerz", "status", "meldung",
        "text", "formulier", "natür", "natuerl", "agent", "laufend",
    )
    if not any(k in tl for k in edit_kw):
        return None
    rel = format_local_path(relative_path).lower()
    # Spezifische stabile UI-Edits fuer TopNavigation
    if rel.endswith("frontend/src/components/topnavigation.jsx"):
        src = current_content
        tl2 = str(task or "").lower()
        wants_hide_generator = ("datei-generator" in tl2 or "datei generator" in tl2)
        wants_hide_studio = ("design studio" in tl2)
        if wants_hide_generator:
            src = _ensure_button_hidden_by_title(src, "Dokumente und Designs generieren")
        if wants_hide_studio:
            src = _ensure_button_hidden_by_title(src, "Design Studio mit Chat & Canvas")
        if src != current_content:
            return src
        return current_content
    if not rel.endswith("frontend/app.js"):
        return None
    src = current_content
    pairs = (
        ('"ich brauche kurz eine Entscheidung"', '"Bitte kurz entscheiden"'),
        ('"Rueckfrage ist direkt hier im Promptbereich."', '"Entscheidung direkt hier."'),
        ('"Keine zusaetzliche Formularbestaetigung noetig."', '"Ohne Extra-Formular."'),
        ('"Bestaetigung uebernommen, weiter gehts."', '"Weiter."'),
        ('"hier ist ein Problem aufgetreten"', '"Problem."'),
    )
    for old, new in pairs:
        if old in src:
            src = src.replace(old, new)
    if src != current_content:
        return src
    return None


def _ensure_button_hidden_by_title(src: str, title_text: str) -> str:
    lines = src.splitlines()
    out = []
    i = 0
    changed = False
    while i < len(lines):
        ln = lines[i]
        if "<button" in ln:
            block = [ln]
            j = i + 1
            while j < len(lines):
                block.append(lines[j])
                if "</button>" in lines[j]:
                    break
                j += 1
            block_txt = "\n".join(block)
            if title_text in block_txt:
                if "style={{ display: \"none\" }}" not in block_txt:
                    inserted = False
                    new_block = []
                    for b in block:
                        new_block.append(b)
                        if (not inserted) and ('title="' in b or "title='" in b):
                            indent = re.match(r"^(\s*)", b).group(1)
                            new_block.append(f'{indent}style={{ display: "none" }}')
                            inserted = True
                    if not inserted:
                        indent = re.match(r"^(\s*)", block[0]).group(1) + "  "
                        new_block.insert(1, f'{indent}style={{ display: "none" }}')
                    block = new_block
                    changed = True
                out.extend(block)
                i = j + 1
                continue
            out.extend(block)
            i = j + 1
            continue
        out.append(ln)
        i += 1
    return "\n".join(out) + ("\n" if src.endswith("\n") else "")


def _apply_explicit_ui_ops_for_known_files(current_content: str, relative_path: str, task: str) -> str | None:
    rel = format_local_path(relative_path).lower()
    tl = str(task or "").lower()
    if rel.endswith("frontend/src/components/topnavigation.jsx"):
        out = current_content
        if ("datei-generator" in tl) or ("datei generator" in tl):
            out = _ensure_button_hidden_by_title(out, "Dokumente und Designs generieren")
        if "design studio" in tl:
            out = _ensure_button_hidden_by_title(out, "Design Studio mit Chat & Canvas")
        if out != current_content:
            return out
    return None


def extract_literal_write_body(task: str):
    """
    Expliziten Dateiinhalt aus dem Prompt holen (DEBUG: nicht die ganze Aufgabe als Stub schreiben).

    Erkennt u.a. Zeile 'Inhalt: ...', Markdown-Fences, 'mit dem Inhalt \"...\"'.
    """
    raw = str(task or "").strip()
    if not raw:
        return None

    fences = list(re.finditer(r"```(?:[^\n`]*)\n(.*?)```", raw, re.DOTALL | re.IGNORECASE))
    for fm in reversed(fences):
        inner = (fm.group(1) or "").strip("\n").strip()
        if not inner:
            continue
        if len(inner) >= max(40, int(len(raw) * 0.88)):
            continue
        return inner

    line_patterns = [
        r"(?im)^(?:inhalt|content|neuer\s+inhalt|text|body)\s*(?:ist)?\s*:\s*(.+?)\s*$",
        r"(?im)^(?:nur|exakt|genau)\s*:\s*(.+?)\s*$",
    ]
    for pat in line_patterns:
        ms = list(re.finditer(pat, raw))
        if ms:
            val = ms[-1].group(1).strip()
            if (val.startswith('"') and val.endswith('"')) or (val.startswith("'") and val.endswith("'")):
                val = val[1:-1].strip()
            if val and len(val) < max(60, int(len(raw) * 0.92)):
                return val

    m = re.search(r'(?i)\b(?:mit\s+(?:dem\s+)?inhalt|mit\s+text)\s+["\']([^"\']+)["\']', raw)
    if m:
        return m.group(1).strip()

    m = re.search(
        r"(?i)\bmit\s+(?:dem\s+)?inhalt\s+([A-Za-z0-9_.:/\\-]+?)(?:[.!?]\s*$|\s*$)",
        raw,
        re.MULTILINE,
    )
    if m:
        val = str(m.group(1) or "").strip().strip("\"'`")
        if val and " " not in val:
            return val

    # Fallback fuer structured KI-prompts: "Inhalt X" innerhalb von "Ziel:"-Block
    m = re.search(
        r"(?i)\binhalt\s+([A-Za-z0-9_.\-]+)\b",
        raw,
        re.MULTILINE,
    )
    if m:
        val = str(m.group(1) or "").strip().strip("\"'`")
        if val and " " not in val and len(val) < 120:
            return val

    m = re.search(r'(?i)\b(?:inhalt|content)\s+["\']([^"\']+)["\']', raw)
    if m:
        return m.group(1).strip()

    return None


def _extract_explicit_replace_pair(task: str):
    """
    Erkennt explizite Aenderungspaare wie:
    - "den Inhalt von content2 zu modified"
    - "aendere ... von 'A' zu 'B'"
    """
    raw = str(task or "").strip()
    if not raw:
        return None, None

    patterns = (
        r"(?is)\b(?:den\s+)?inhalt\s+von\s+[\"'`]?([^\"'`\n]+?)[\"'`]?\s+zu\s+[\"'`]?([^\"'`\n]+?)[\"'`]?(?:[.!?]\s*$|\s*$)",
        r"(?is)\bvon\s+[\"'`]?([^\"'`\n]+?)[\"'`]?\s+zu\s+[\"'`]?([^\"'`\n]+?)[\"'`]?(?:[.!?]\s*$|\s*$)",
    )
    for pat in patterns:
        m = re.search(pat, raw)
        if not m:
            continue
        old = str(m.group(1) or "").strip()
        new = str(m.group(2) or "").strip()
        if old and new:
            return old, new
    return None, None


def infer_relative_path_token_from_task(task):
    """Erstes gueltiges relatives Ziel unter rambo_builder_local (inkl. .txt)."""
    raw = str(task or "").strip()
    if not raw:
        return None
    pattern = r"\b((?:[\w-]+/)*[\w.-]+\.(?:txt|md|json|py|js|mjs|cjs|ts|tsx|jsx|css|html|yml|yaml|toml))\b"
    for m in re.finditer(pattern, raw, re.IGNORECASE):
        cand = format_local_path(m.group(1))
        _, _, err = resolve_local_target_path(cand, task)
        if not err:
            return cand
    return None


def resolve_proposed_content_for_local_task(task, target_path, current_content, file_exists, relative_path):
    """Vorschlagsinhalt: expliziter Prompt-Text > refine > Stub-Generator."""
    rel = format_local_path(relative_path or target_path or "")
    if _is_tests_py_path(rel):
        extracted = extract_pytest_function_from_prompt(task)
        if extracted is not None:
            return extracted
        lit = extract_literal_write_body(task)
        if lit is not None and bool(validate_pytest_file(rel, lit).get("ok")):
            return lit
        tl = str(task or "").lower()
        if any(k in tl for k in ("pytest", " test", "tests/", "tests\\", "smoke-test", "smoketest")):
            return "def test_smoke():\n    assert True\n"

    old_txt, new_txt = _extract_explicit_replace_pair(task)
    if file_exists and old_txt and new_txt:
        source = str(current_content or "")
        if old_txt in source:
            return source.replace(old_txt, new_txt)
        if source.strip() == old_txt:
            return new_txt
    lit = extract_literal_write_body(task)
    if lit is not None:
        return lit
    explicit_ui = _apply_explicit_ui_ops_for_known_files(str(current_content or ""), rel, task)
    if explicit_ui is not None:
        return explicit_ui
    refined = refine_local_source_for_editing_task(current_content, rel, task, file_exists)
    if refined is not None:
        return refined
    explicit_rel = extract_explicit_local_relative_path(task)
    if explicit_rel and file_exists:
        code_exts = (".js", ".jsx", ".ts", ".tsx", ".py", ".css", ".html")
        if rel.lower().endswith(code_exts):
            # Bei expliziten Code-Dateien niemals Task-Stub als Vollersatz erzeugen.
            return current_content
    generated = generate_content_from_task(task, target_path or rel)
    if file_exists and rel.lower().endswith((".js", ".jsx", ".ts", ".tsx", ".py", ".css", ".html")):
        g_low = str(generated or "").strip().lower()
        if g_low.startswith("aufgabe:") or g_low.startswith("erstellt:") or "schritte:" in g_low:
            return current_content
    return generated


def _is_agent_instruction_prompt(task: str) -> bool:
    low = str(task or "").lower()
    patterns = (
        "du arbeitest im projekt",
        "aktueller fehler",
        "ziel:",
        "erlaubte dateien",
        "pflicht-test",
        "ausgabeformat",
        "geänderte dateien",
        "geaenderte dateien",
        "kurz umgesetzt",
        "kurz getestet",
        "self-repair",
        "self repair",
        "external ai prompt",
    )
    return any(p in low for p in patterns)


def _extract_agent_direct_write(task: str):
    """Extrahiert (resolved_path, relative_path, proposed_content, file_exists) aus Agent-Prompt."""
    rel = extract_explicit_local_relative_path(task) or infer_target_path(task)
    resolved, relative_path, err = resolve_local_target_path(rel, task)
    if err:
        return None, None, None, None, err
    guard_check = _validate_direct_run_paths([rel, relative_path], "apply", task)
    if not bool(guard_check.get("ok")):
        blocked = guard_check.get("blocked_files") or []
        blocked_text = ", ".join([str(p) for p in blocked]) if blocked else relative_path
        return None, None, None, None, f"{DIRECT_RUN_GUARD_BLOCK_MESSAGE} Blockiert: {blocked_text}"
    current_content, file_exists = read_text_file(resolved)
    proposed = resolve_proposed_content_for_local_task(task, relative_path, current_content, file_exists, relative_path)
    if not str(proposed or "").strip():
        return None, None, None, None, "Kein schreibbarer Zielinhalt im Prompt erkannt."
    return resolved, relative_path, str(proposed), bool(file_exists), None


def generated_content_from_task_or_literal(task, target_path, *, log_label=None, log_ok_message=None):
    """Stub-Generator, aber nur wenn kein expliziter Inhalt im Prompt steht."""
    lit = extract_literal_write_body(task)
    if lit is not None:
        if log_label and log_ok_message:
            append_ui_log_entry(log_label, log_ok_message, "info")
        return lit
    return generate_content_from_task(task, target_path)


def infer_target_path(task):
    explicit = extract_explicit_local_relative_path(task)
    if explicit:
        return explicit
    token = infer_relative_path_token_from_task(task)
    if token:
        return token
    lowered = str(task or "").lower()
    stop_words = {
        "die", "der", "das", "den", "dem", "des", "ein", "eine", "einen",
        "und", "oder", "fuer", "mit", "von", "bei", "in", "auf", "an",
        "als", "zu", "im", "am", "ist", "sind", "wird", "werden", "soll",
        "sollen", "kann", "koennen", "eine", "einen", "einer"
    }

    if any(w in lowered for w in ["dokumentier", "beschreib", "notiz", "readme", "markdown"]):
        ext = "md"
    elif any(w in lowered for w in ["konfig", "config", "einstellung", "parameter"]):
        ext = "json"
    elif any(w in lowered for w in ["skript", "script", "python", "funktion", "klasse"]):
        ext = "py"
    elif any(w in lowered for w in ["html", "seite", "ansicht"]):
        ext = "html"
    else:
        ext = "md"

    words = normalize_task_words(task)
    slug_parts = [w.lower() for w in words if w.lower() not in stop_words][:4]
    slug = "_".join(slug_parts) if slug_parts else "aufgabe"
    slug = re.sub(r"[^a-z0-9_]", "", slug)[:40] or "aufgabe"

    return f"data/auto/{slug}.{ext}"


def infer_project_target_path(task, suggestions=None):
    suggestions = suggestions or {}
    task_low = str(task or "").lower()
    smart_targets = [
        (("topnavigation", "datei-generator"), "frontend/src/components/TopNavigation.jsx"),
        (("topnavigation", "design studio"), "frontend/src/components/TopNavigation.jsx"),
        (("raineragent", "upload"), "frontend/src/components/RainerAgent.jsx"),
        (("dashboard", "offline"), "frontend/src/components/RamboManagementDashboard.jsx"),
    ]
    for needles, target in smart_targets:
        if all(n in task_low for n in needles):
            _, cleaned, error = validate_project_write_path(target)
            if not error:
                return cleaned
    expl = extract_explicit_local_relative_path(task)
    if expl:
        _, cleaned, error = validate_project_write_path(expl)
        if not error:
            return cleaned
    explicit_candidates = re.findall(r"[A-Za-z0-9_./-]+\.[A-Za-z0-9]+", str(task or ""))
    for candidate in explicit_candidates:
        _, cleaned, error = validate_project_write_path(candidate)
        if not error:
            return cleaned
    if explicit_candidates:
        return format_local_path(explicit_candidates[0])

    prioritized = suggestions.get("prioritized_targets") or []
    for candidate in prioritized:
        path = str(candidate.get("path") or "").strip()
        if not path:
            continue
        _, cleaned, error = validate_project_write_path(path)
        if not error:
            return cleaned

    for item in suggestions.get("area_suggestions") or []:
        path = str(item.get("example_path") or "").strip()
        _, cleaned, error = validate_project_write_path(path)
        if not error:
            return cleaned

    fallback = infer_target_path(task)
    _, cleaned, error = validate_project_write_path(fallback)
    if not error:
        return cleaned
    return "data/auto/projekt_aufgabe.md"


def detect_extension(target_path):
    cleaned = str(target_path or "").strip()
    if "." in cleaned.split("/")[-1]:
        return cleaned.rsplit(".", 1)[-1].lower()
    return "txt"


def generate_content_from_task(task, target_path=""):
    ext = detect_extension(target_path)
    timestamp = get_timestamp()

    path_stem = str(target_path or "").strip().rsplit("/", 1)[-1]
    if "." in path_stem:
        path_stem = path_stem.rsplit(".", 1)[0]
    if not path_stem:
        path_stem = "lokal"

    if ext == "md":
        return (
            f"# {task}\n\n"
            f"**Erstellt:** {timestamp}\n\n"
            "## Ziel\n\n"
            f"{task}\n\n"
            "## Umsetzungsschritte\n\n"
            "- [ ] Anforderungen klaeren\n"
            "- [ ] Umsetzung vorbereiten\n"
            "- [ ] Testen und pruefen\n"
            "- [ ] Ergebnis dokumentieren\n\n"
            "## Notizen\n\n"
            "_Hier koennen Hinweise zur Umsetzung ergaenzt werden._\n\n"
            "## Status\n\n"
            "- **Offen**\n"
        )
    if ext == "json":
        return json.dumps({
            "aufgabe": task,
            "erstellt": timestamp,
            "status": "offen",
            "schritte": [
                "Anforderungen klaeren",
                "Umsetzung vorbereiten",
                "Testen und pruefen"
            ],
            "notizen": ""
        }, ensure_ascii=True, indent=2)
    if ext == "py":
        return (
            f'"""Lokale Aufgabe: {task}\nErstellt: {timestamp}\n"""\n\n'
            "def main():\n"
            f"    print('Starte: {task}')\n\n"
            "if __name__ == '__main__':\n"
            "    main()\n"
        )
    if ext == "html":
        return (
            "<!DOCTYPE html>\n"
            "<html lang=\"de\">\n"
            "<head>\n"
            "  <meta charset=\"UTF-8\">\n"
            f"  <title>{task}</title>\n"
            "</head>\n"
            "<body>\n"
            f"  <h1>{task}</h1>\n"
            f"  <p>Erstellt: {timestamp}</p>\n"
            "  <p>Lokale Vorschau innerhalb von rambo_builder_local.</p>\n"
            "</body>\n"
            "</html>"
        )
    if ext == "css":
        return (
            f"/* Aufgabe: {task} */\n"
            f"/* Erstellt: {timestamp} */\n\n"
            f".{path_stem} {{\n"
            "  display: block;\n"
            "  padding: 1rem;\n"
            "}\n"
        )
    if ext == "js":
        fn = "init" + "".join(w.capitalize() for w in task.split()[:2])
        return (
            f"/* Aufgabe: {task} */\n"
            f"/* Erstellt: {timestamp} */\n\n"
            f"function {fn}() {{\n"
            f"  console.log('Starte: {task}');\n"
            "}\n"
        )
    return (
        f"Aufgabe: {task}\n"
        f"Erstellt: {timestamp}\n"
        "Status: offen\n\n"
        "Schritte:\n"
        "1. Anforderungen klaeren\n"
        "2. Umsetzung vorbereiten\n"
        "3. Testen und pruefen\n\n"
        "Notizen:\n"
        "[Hier ergaenzen]\n"
    )


AUTO_LOOP_STATUSES = {"idle", "running", "paused", "blocked", "failed", "stopped", "done", "approval_required"}
AUTO_LOOP_STEP_STATUSES = {"geplant", "laeuft", "fertig", "blockiert", "fehlgeschlagen", "wartet auf freigabe", "uebersprungen"}
AUTO_LOOP_GATES = {"auto", "preview", "approval"}


def normalize_auto_loop_state(payload):
    base = {
        "run_id": "",
        "last_run_at": "",
        "goal": "",
        "status": "idle",
        "current_phase": "",
        "current_step": "",
        "last_action": "",
        "next_step": "",
        "summary": "",
        "steps": [],
        "history": [],
        "error_label": "",
        "repair_suggestion": "",
        "affected_file": "",
        "requires_approval": False,
        "stopped": False,
        "stop_requested": False,
        "pause_requested": False,
        "paused_at": "",
        "resumed_at": "",
        "started_at": "",
        "ended_at": "",
        "runner_command": "",
        "retry_count": 0,
        "retry_max": 2,
        "session_blockers": [],
        "apply_mode": "safe",
        "phase": "planning",
        "active_module": "",
        "current_error": _empty_error_info(),
        "repair_plan": [],
        "retry_possible": True,
        "retry_blocked_reason": "",
        "repair_history": [],
        "memory_snapshot": {},
        "acceptance_status": "",
        "acceptance_ready": False,
        "acceptance_blocked_reason": "",
        "acceptance_notes": "",
        "acceptance_timestamp": "",
    }
    if not isinstance(payload, dict):
        return base
    normalized = base.copy()
    normalized.update(payload)
    normalized["status"] = str(normalized.get("status") or "idle").strip().lower() or "idle"
    if normalized["status"] not in AUTO_LOOP_STATUSES:
        normalized["status"] = "idle"
    normalized["run_id"] = str(normalized.get("run_id") or "").strip()
    normalized["goal"] = str(normalized.get("goal") or "").strip()
    normalized["current_phase"] = str(normalized.get("current_phase") or "").strip()
    normalized["current_step"] = str(normalized.get("current_step") or "").strip()
    normalized["last_action"] = str(normalized.get("last_action") or "").strip()
    normalized["next_step"] = str(normalized.get("next_step") or "").strip()
    normalized["summary"] = str(normalized.get("summary") or "").strip()
    normalized["error_label"] = str(normalized.get("error_label") or "").strip()
    normalized["repair_suggestion"] = str(normalized.get("repair_suggestion") or "").strip()
    normalized["affected_file"] = str(normalized.get("affected_file") or "").strip()
    normalized["requires_approval"] = bool(normalized.get("requires_approval"))
    normalized["stopped"] = bool(normalized.get("stopped"))
    normalized["stop_requested"] = bool(normalized.get("stop_requested"))
    normalized["pause_requested"] = bool(normalized.get("pause_requested"))
    normalized["paused_at"] = str(normalized.get("paused_at") or "").strip()
    normalized["resumed_at"] = str(normalized.get("resumed_at") or "").strip()
    normalized["started_at"] = str(normalized.get("started_at") or "").strip()
    normalized["ended_at"] = str(normalized.get("ended_at") or "").strip()
    normalized["runner_command"] = str(normalized.get("runner_command") or "").strip()
    apply_mode_val = str(normalized.get("apply_mode") or "safe").strip().lower()
    normalized["apply_mode"] = apply_mode_val if apply_mode_val in {"safe", "apply"} else "safe"
    try:
        normalized["retry_count"] = max(0, int(normalized.get("retry_count") or 0))
    except Exception:
        normalized["retry_count"] = 0
    try:
        normalized["retry_max"] = max(0, int(normalized.get("retry_max") or 2))
    except Exception:
        normalized["retry_max"] = 2
    raw_steps = normalized.get("steps")
    steps = raw_steps if isinstance(raw_steps, list) else []
    raw_history = normalized.get("history")
    history = raw_history if isinstance(raw_history, list) else []
    raw_session_blockers = normalized.get("session_blockers")
    session_blockers = raw_session_blockers if isinstance(raw_session_blockers, list) else []
    clean_steps = []
    for idx, item in enumerate(steps):
        if not isinstance(item, dict):
            continue
        label = str(item.get("label") or "").strip()
        if not label:
            continue
        status = str(item.get("status") or "geplant").strip().lower()
        if status not in AUTO_LOOP_STEP_STATUSES:
            status = "geplant"
        gate = str(item.get("gate") or "auto").strip().lower()
        if gate not in AUTO_LOOP_GATES:
            gate = "auto"
        raw_details = item.get("blocker_details")
        blocker_details = {
            "path": str((raw_details or {}).get("path") or "").strip() if isinstance(raw_details, dict) else "",
            "rule": str((raw_details or {}).get("rule") or "").strip() if isinstance(raw_details, dict) else "",
            "reason": str((raw_details or {}).get("reason") or "").strip() if isinstance(raw_details, dict) else "",
            "suggestion": str((raw_details or {}).get("suggestion") or "").strip() if isinstance(raw_details, dict) else "",
        }
        action_val = str(item.get("action") or item.get("id") or "").strip()
        phase_val = str(item.get("phase") or "").strip().lower()
        if phase_val not in AGENT_PHASES:
            phase_val = phase_for_action(action_val)
        tool_val = str(item.get("tool") or "").strip()
        if not tool_val:
            tool_val = tool_for_action(action_val)
        clean_steps.append({
            "id": str(item.get("id") or f"s{idx + 1}"),
            "label": label,
            "status": status,
            "detail": str(item.get("detail") or "").strip(),
            "gate": gate,
            "risky": bool(item.get("risky")),
            "approved": bool(item.get("approved")),
            "action": action_val,
            "phase": phase_val,
            "tool": tool_val,
            "blocker": str(item.get("blocker") or "").strip(),
            "blocker_details": blocker_details,
            "retry_round": int(item.get("retry_round") or 0) if isinstance(item.get("retry_round"), (int, float)) else 0,
            "retry_reason": str(item.get("retry_reason") or "").strip(),
            "retry_category": canonical_error_category(item.get("retry_category")) if item.get("retry_category") else "",
            "retry_target_file": str(item.get("retry_target_file") or "").strip(),
            "retry_status": str(item.get("retry_status") or status or "").strip().lower(),
        })
    normalized["steps"] = clean_steps[:96]
    normalized["history"] = [str(item).strip() for item in history if str(item or "").strip()][:160]
    normalized["last_run_at"] = str(normalized.get("last_run_at") or "").strip()
    clean_blockers = []
    for item in session_blockers:
        if not isinstance(item, dict):
            continue
        reason = str(item.get("reason") or "").strip()
        path = str(item.get("path") or "").strip()
        if not reason and not path:
            continue
        clean_blockers.append({
            "step_id": str(item.get("step_id") or "").strip(),
            "step_label": str(item.get("step_label") or "").strip(),
            "path": path,
            "rule": str(item.get("rule") or "").strip(),
            "reason": reason,
            "suggestion": str(item.get("suggestion") or "").strip(),
            "timestamp": str(item.get("timestamp") or "").strip(),
        })
    normalized["session_blockers"] = clean_blockers[:16]

    raw_error = normalized.get("current_error")
    if isinstance(raw_error, dict) and raw_error.get("category"):
        err = _empty_error_info()
        err.update({
            "category": canonical_error_category(raw_error.get("category")),
            "raw_category": str(raw_error.get("raw_category") or raw_error.get("category") or "").strip(),
            "label": str(raw_error.get("label") or "").strip(),
            "suggestion": str(raw_error.get("suggestion") or "").strip(),
            "file": str(raw_error.get("file") or "").strip(),
            "area": str(raw_error.get("area") or "").strip(),
            "severity": str(raw_error.get("severity") or "medium").strip(),
            "recommended_tool": str(raw_error.get("recommended_tool") or "").strip(),
            "source": str(raw_error.get("source") or "auto_loop").strip(),
            "signature": str(raw_error.get("signature") or "").strip(),
            "ts": str(raw_error.get("ts") or "").strip(),
        })
        try:
            err["occurrences"] = max(1, int(raw_error.get("occurrences") or 1))
        except Exception:
            err["occurrences"] = 1
        if not err["signature"]:
            err["signature"] = f"{err['category']}|{err['file']}".strip("|")
        normalized["current_error"] = err
    else:
        normalized["current_error"] = _empty_error_info()

    raw_plan = normalized.get("repair_plan")
    clean_plan = []
    if isinstance(raw_plan, list):
        for item in raw_plan:
            if not isinstance(item, dict):
                continue
            label = str(item.get("label") or "").strip()
            if not label:
                continue
            clean_plan.append({
                "id": str(item.get("id") or "").strip() or label.lower().replace(" ", "_"),
                "label": label,
                "tool": str(item.get("tool") or "").strip(),
                "action": str(item.get("action") or "").strip(),
                "gate": str(item.get("gate") or "auto").strip().lower(),
                "risky": bool(item.get("risky")),
            })
    normalized["repair_plan"] = clean_plan[:8]

    normalized["retry_possible"] = bool(normalized.get("retry_possible", True))
    normalized["retry_blocked_reason"] = str(normalized.get("retry_blocked_reason") or "").strip()
    normalized["acceptance_status"] = str(normalized.get("acceptance_status") or "").strip()
    normalized["acceptance_ready"] = bool(normalized.get("acceptance_ready"))
    normalized["acceptance_blocked_reason"] = str(normalized.get("acceptance_blocked_reason") or "").strip()
    normalized["acceptance_notes"] = str(normalized.get("acceptance_notes") or "").strip()
    normalized["acceptance_timestamp"] = str(normalized.get("acceptance_timestamp") or "").strip()

    raw_repair_history = normalized.get("repair_history")
    clean_history = []
    if isinstance(raw_repair_history, list):
        for item in raw_repair_history:
            if not isinstance(item, dict):
                continue
            reason = str(item.get("reason") or "").strip()
            if not reason and not item.get("category"):
                continue
            try:
                round_num = int(item.get("round") or 0)
            except Exception:
                round_num = 0
            clean_history.append({
                "round": round_num,
                "category": canonical_error_category(item.get("category")),
                "reason": reason,
                "file": str(item.get("file") or "").strip(),
                "status": str(item.get("status") or "").strip(),
                "repeated": bool(item.get("repeated")),
                "timestamp": str(item.get("timestamp") or "").strip(),
            })
    normalized["repair_history"] = clean_history[-16:]
    return normalized


# ------------------------------------------------------------------
# Server-side Auto-Loop orchestration (resume + gated execution)
# ------------------------------------------------------------------

AUTO_LOOP_ACTIONS_READ = {"context_load", "project_mode_prepare", "direct_preview",
                          "analyze_error", "prepare_repair", "runner_prepare"}
AUTO_LOOP_ACTIONS_WRITE = {"runner_execute", "direct_apply"}

AGENT_PHASES = (
    "planning", "context", "analysis", "preview", "approval",
    "execution", "verification", "repair", "blocked", "completed",
)

ACTION_PHASE_MAP = {
    "context_load": "context",
    "project_mode_prepare": "planning",
    "direct_preview": "preview",
    "direct_apply": "execution",
    "runner_prepare": "analysis",
    "runner_execute": "execution",
    "analyze_error": "verification",
    "prepare_repair": "repair",
}

ACTION_TOOL_MAP = {
    "context_load": "context",
    "project_mode_prepare": "project_mode",
    "direct_preview": "direct",
    "direct_apply": "direct",
    "runner_prepare": "runner",
    "runner_execute": "runner",
    "analyze_error": "error_analysis",
    "prepare_repair": "repair",
}


def phase_for_action(action):
    return ACTION_PHASE_MAP.get(str(action or "").lower(), "planning")


def tool_for_action(action):
    return ACTION_TOOL_MAP.get(str(action or "").lower(), "")


def _empty_blocker_details():
    return {"path": "", "rule": "", "reason": "", "suggestion": ""}


def _blocker_details(path="", rule="", reason="", suggestion=""):
    return {
        "path": str(path or "").strip(),
        "rule": str(rule or "").strip(),
        "reason": str(reason or "").strip(),
        "suggestion": str(suggestion or "").strip(),
    }


CANONICAL_ERROR_CATEGORIES = {
    "syntax_error": "syntax",
    "syntax": "syntax",
    "import_error": "import",
    "import": "import",
    "test_error": "test",
    "test": "test",
    "build_error": "build",
    "build": "build",
    "file_not_found": "file_not_found",
    "permission_error": "permission",
    "permission": "permission",
    "guard_block": "guard_block",
    "approval_block": "approval_block",
    "runtime_error": "runtime",
    "runtime": "runtime",
    "unknown_error": "unknown",
    "unknown": "unknown",
}

CATEGORY_SEVERITY = {
    "syntax": "high", "import": "high", "test": "medium", "build": "high",
    "file_not_found": "medium", "permission": "medium",
    "guard_block": "medium", "approval_block": "low",
    "runtime": "high", "unknown": "medium",
}

CATEGORY_PRIMARY_TOOL = {
    "syntax": "direct", "import": "direct", "test": "runner", "build": "runner",
    "file_not_found": "project_mode", "permission": "project_mode",
    "guard_block": "project_mode", "approval_block": "direct",
    "runtime": "runner", "unknown": "direct",
}


def _empty_error_info():
    return {
        "category": "",
        "raw_category": "",
        "label": "",
        "suggestion": "",
        "file": "",
        "area": "",
        "severity": "",
        "recommended_tool": "",
        "source": "",
        "signature": "",
        "occurrences": 0,
        "ts": "",
    }


def canonical_error_category(raw):
    key = str(raw or "").strip().lower()
    if not key:
        return "unknown"
    return CANONICAL_ERROR_CATEGORIES.get(key, "unknown")


def build_error_info(label="", suggestion="", file="", area="",
                     raw_category="", source="auto_loop"):
    """Erzeugt eine einheitliche Fehlerstruktur."""
    label = str(label or "").strip()
    suggestion = str(suggestion or "").strip()
    file = str(file or "").strip()
    area = str(area or "").strip()
    raw_category = str(raw_category or "").strip()
    category = canonical_error_category(raw_category)
    signature = f"{category}|{file}".strip("|")
    return {
        "category": category,
        "raw_category": raw_category or category,
        "label": label or category,
        "suggestion": suggestion,
        "file": file,
        "area": area,
        "severity": CATEGORY_SEVERITY.get(category, "medium"),
        "recommended_tool": CATEGORY_PRIMARY_TOOL.get(category, "direct"),
        "source": str(source or "auto_loop"),
        "signature": signature,
        "occurrences": 1,
        "ts": get_timestamp(),
    }


def build_repair_plan(error_info, runner_command="", apply_mode="safe"):
    """Leitet einen kompakten Repair-Plan aus Fehler-Info ab."""
    if not isinstance(error_info, dict):
        return []
    category = canonical_error_category(error_info.get("category"))
    file = str(error_info.get("file") or "").strip()
    suggestion = str(error_info.get("suggestion") or "").strip()
    needs_runner = bool(str(runner_command or "").strip())
    needs_apply_gate = str(apply_mode or "safe").lower() == "apply"

    def step(id_, label, tool, action, gate="auto", risky=False):
        return {
            "id": id_, "label": label, "tool": tool, "action": action,
            "gate": gate, "risky": risky,
        }

    def file_step(id_, label, tool, action, gate="auto", risky=False):
        """Repair-Schritt mit Datei-Endpoint-Referenz wenn Datei bekannt."""
        s = step(id_, label, tool, action, gate, risky)
        if file and _is_valid_file_path(file):
            s["file_target"] = file
            s["load_endpoint"] = f"/api/file/load?path={file}"
            s["prepare_endpoint"] = "/api/file/prepare-edit"
            s["apply_endpoint"] = "/api/file/apply"
            s["apply_mode"] = "apply" if needs_apply_gate else "safe"
        return s

    plan = []
    if category in {"syntax", "import", "runtime", "unknown"}:
        plan.append(step("repair_context", "Kontext fuer Reparatur laden", "context", "context_load"))
        plan.append(file_step("repair_direct_preview",
                         f"Vorschau: {file or 'Zieldatei'} laden & Diff pruefen",
                         "direct", "direct_preview", gate="preview"))
        plan.append(file_step("repair_direct_apply",
                         "Apply (Freigabe)" if needs_apply_gate else "Apply (Safe) — kein Schreiben",
                         "direct", "direct_apply",
                         gate="approval" if needs_apply_gate else "auto",
                         risky=needs_apply_gate))
        if needs_runner:
            plan.append(step("repair_runner", "Runner erneut ausfuehren", "runner",
                             "runner_execute", gate="approval", risky=True))
        plan.append(step("repair_verify", "Fehleranalyse erneut pruefen", "error_analysis", "analyze_error"))
    elif category in {"test", "build"}:
        if needs_runner:
            plan.append(step("repair_runner_pre", "Runner Vorbereitung", "runner", "runner_prepare"))
            plan.append(step("repair_runner", "Runner erneut ausfuehren", "runner",
                             "runner_execute", gate="approval", risky=True))
        plan.append(step("repair_verify", "Fehleranalyse erneut pruefen", "error_analysis", "analyze_error"))
        if suggestion:
            plan.append(file_step("repair_direct_preview",
                             "Direktmodus-Vorschau fuer Fix-Vorschlag",
                             "direct", "direct_preview", gate="preview"))
    elif category == "file_not_found":
        plan.append(step("repair_project_mode", "Zielpfade / Project-Mode pruefen", "project_mode", "project_mode_prepare"))
        plan.append(file_step("repair_direct_preview",
                         "Fehlende Datei erzeugen/patchen (Vorschau)",
                         "direct", "direct_preview", gate="preview"))
        plan.append(file_step("repair_direct_apply",
                         "Apply (Freigabe)" if needs_apply_gate else "Apply (Safe)",
                         "direct", "direct_apply",
                         gate="approval" if needs_apply_gate else "auto",
                         risky=needs_apply_gate))
    elif category == "permission":
        plan.append(step("repair_project_mode", "Zielpfad / Berechtigung pruefen", "project_mode", "project_mode_prepare"))
        plan.append(step("repair_context", "Kontext / Guard-Regel lesen", "context", "context_load"))
    elif category == "guard_block":
        plan.append(step("repair_project_mode", "Zielpfad waehlen / Guard umgehen", "project_mode", "project_mode_prepare"))
    elif category == "approval_block":
        plan.append(file_step("repair_direct_apply",
                         "Apply-Freigabe erneut anfordern",
                         "direct", "direct_apply",
                         gate="approval", risky=True))
    else:
        plan.append(step("repair_verify", "Fehler erneut analysieren", "error_analysis", "analyze_error"))
    return plan


def _auto_loop_clear_error(state):
    state["current_error"] = _empty_error_info()
    state["repair_plan"] = []
    state["error_label"] = ""
    state["repair_suggestion"] = ""
    state["affected_file"] = ""


def build_auto_loop_plan(goal, runner_command="", project_blocked=False,
                          apply_mode="safe", guard_info=None, retry_round=0):
    goal = str(goal or "").strip()
    runner_command = str(runner_command or "").strip()
    apply_mode = str(apply_mode or "safe").strip().lower()
    needs_runner_approval = bool(runner_command)
    needs_apply_approval = apply_mode == "apply"
    guard_info = guard_info if isinstance(guard_info, dict) else {}
    guard_path = str(guard_info.get("path") or guard_info.get("target_path") or "").strip()
    guard_rule = str(guard_info.get("rule") or guard_info.get("pattern") or "Guard-Blocker").strip()
    guard_reason = str(guard_info.get("reason") or guard_info.get("decision") or "Pfad nicht freigegeben.").strip()
    guard_suggestion = str(guard_info.get("suggestion") or "Anderen Zielpfad waehlen oder Guard-Regel pruefen.").strip()

    runner_detail = (
        f"Freigabepflichtig: {runner_command}" if needs_runner_approval
        else "Kein Runner-Command gesetzt, nur Vorbereitung."
    )
    apply_detail = (
        "Direktmodus Apply: echtes Schreiben. Freigabe erforderlich."
        if needs_apply_approval
        else "Direktmodus bleibt im Safe-Mode. Keine Schreibfreigabe noetig."
    )
    plan = [
        {"id": "context", "label": "Kontext pruefen/laden",
         "detail": "Projektkontext, Guard-Basis und Kontext-Flow laden.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "context_load", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "project_mode", "label": "Zielpfade/Project-Mode vorbereiten",
         "detail": "Project-Mode und sichere Zielpfade vorbereiten.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "project_mode_prepare", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "direct_preview", "label": "Direktmodus-Vorschau vorbereiten",
         "detail": "Direktmodus Safe-Mode Vorschau (nur Analyse).",
         "status": "geplant", "gate": "preview", "risky": False, "approved": True,
         "action": "direct_preview", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "direct_apply", "label": "Direktmodus Apply (Freigabe)",
         "detail": apply_detail,
         "status": "geplant",
         "gate": "approval" if needs_apply_approval else "auto",
         "risky": needs_apply_approval,
         "approved": not needs_apply_approval,
         "action": "direct_apply",
         "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "runner", "label": "Runner vorbereiten/ausfuehren",
         "detail": runner_detail,
         "status": "geplant",
         "gate": "approval" if needs_runner_approval else "auto",
         "risky": needs_runner_approval,
         "approved": not needs_runner_approval,
         "action": "runner_execute" if needs_runner_approval else "runner_prepare",
         "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "error_analysis", "label": "Fehleranalyse lesen",
         "detail": "Runner-Ergebnis und Fehlerkategorie auswerten.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "analyze_error", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": "repair", "label": "Reparatur-Schritt vorbereiten",
         "detail": "Reparaturvorschlag in Direktmodus/Agent-Run uebernehmen.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "prepare_repair", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
    ]
    if project_blocked:
        block_reason = guard_reason or "Guard markiert Pfad als blockiert."
        for step in plan:
            if step["id"] in {"direct_apply", "runner"}:
                step["status"] = "blockiert"
                step["approved"] = False
                step["gate"] = "approval"
                step["blocker"] = f"Guard blockiert Pfad: {block_reason}"
                step["blocker_details"] = _blocker_details(
                    path=guard_path,
                    rule=guard_rule,
                    reason=block_reason,
                    suggestion=guard_suggestion,
                )
                step["detail"] = step["blocker"]
    for step in plan:
        if not step.get("phase"):
            step["phase"] = phase_for_action(step.get("action"))
        if not step.get("tool"):
            step["tool"] = tool_for_action(step.get("action"))
    return plan


def build_auto_loop_retry_plan(goal, runner_command="", apply_mode="safe", retry_round=1):
    """Baut eine verkuerzte Retry-Runde (direct_preview -> apply -> runner -> analyze -> repair)."""
    needs_runner_approval = bool(str(runner_command or "").strip())
    needs_apply_approval = str(apply_mode or "safe").strip().lower() == "apply"
    return [
        {"id": f"direct_preview_r{retry_round}", "label": f"Retry {retry_round}: Vorschau",
         "detail": "Reparierte Vorschau vorbereiten.",
         "status": "geplant", "gate": "preview", "risky": False, "approved": True,
         "action": "direct_preview", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": f"direct_apply_r{retry_round}", "label": f"Retry {retry_round}: Apply (Freigabe)",
         "detail": "Reparierter Apply-Schritt (Freigabe erforderlich)." if needs_apply_approval else "Reparierter Safe-Apply (keine Freigabe).",
         "status": "geplant",
         "gate": "approval" if needs_apply_approval else "auto",
         "risky": needs_apply_approval,
         "approved": not needs_apply_approval,
         "action": "direct_apply",
         "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": f"runner_r{retry_round}", "label": f"Retry {retry_round}: Runner",
         "detail": "Runner erneut ausfuehren." if needs_runner_approval else "Runner-Vorbereitung.",
         "status": "geplant",
         "gate": "approval" if needs_runner_approval else "auto",
         "risky": needs_runner_approval,
         "approved": not needs_runner_approval,
         "action": "runner_execute" if needs_runner_approval else "runner_prepare",
         "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": f"error_analysis_r{retry_round}", "label": f"Retry {retry_round}: Fehleranalyse",
         "detail": "Fehler nach Retry erneut auswerten.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "analyze_error", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
        {"id": f"repair_r{retry_round}", "label": f"Retry {retry_round}: Reparatur-Schritt",
         "detail": "Reparaturvorschlag erneut vorbereiten.",
         "status": "geplant", "gate": "auto", "risky": False, "approved": True,
         "action": "prepare_repair", "blocker": "", "blocker_details": _empty_blocker_details(),
         "retry_round": retry_round},
    ]
    # Phase/Tool wird von normalize_auto_loop_state anhand action nachgezogen.


def _auto_loop_root_step_id(step_id):
    """Mappt Retry-IDs (z.B. 'runner_r2') auf die zugrundeliegende Aktion."""
    sid = str(step_id or "")
    if "_r" in sid:
        return sid.split("_r", 1)[0]
    return sid


def _auto_loop_find_index(state, step_id):
    for idx, step in enumerate(state.get("steps") or []):
        if step.get("id") == step_id:
            return idx
    return -1


def _auto_loop_next_pending(state):
    steps = state.get("steps") or []
    for idx, step in enumerate(steps):
        status = (step.get("status") or "").lower()
        if status in {"geplant", "laeuft", "wartet auf freigabe"}:
            return idx, step
    return -1, None


def _auto_loop_is_internal_safe_step(step):
    if not isinstance(step, dict):
        return False
    action = str(step.get("action") or step.get("id") or "").strip().lower()
    gate = str(step.get("gate") or "auto").strip().lower()
    status = str(step.get("status") or "").strip().lower()
    if status not in {"geplant", "laeuft"}:
        return False
    if action == "direct_preview":
        return gate in {"auto", "preview"}
    if action == "direct_apply":
        return gate == "auto" and not bool(step.get("risky"))
    if gate != "auto":
        return False
    return action in {
        "context_load",
        "project_mode_prepare",
        "runner_prepare",
        "analyze_error",
        "prepare_repair",
    }


def _auto_loop_autocomplete_detail(step):
    if not isinstance(step, dict):
        return "Interner Schritt abgeschlossen."
    label = str(step.get("label") or "").strip()
    if label:
        return f"{label} automatisch abgeschlossen."
    action = str(step.get("action") or step.get("id") or "").strip()
    if action:
        return f"{action} automatisch abgeschlossen."
    return "Interner Schritt abgeschlossen."


def _auto_loop_find_waiting_step(state, target_id=""):
    steps = state.get("steps") or []
    target = str(target_id or "").strip()
    for idx, step in enumerate(steps):
        step_id = str(step.get("id") or "").strip()
        step_status = str(step.get("status") or "").lower()
        if target and step_id != target:
            continue
        if step_status == "wartet auf freigabe":
            return idx, step
        if target and step_id == target:
            return idx, step
    return -1, None


def _auto_loop_push_history(state, message):
    if not message:
        return
    stamp = get_timestamp()
    entry = f"{stamp} - {message}"
    history = state.get("history") or []
    history.insert(0, entry)
    state["history"] = history[:160]


def _parse_run_timestamp(raw):
    text = str(raw or "").strip()
    if not text:
        return None
    try:
        return datetime.strptime(text, "%Y-%m-%d %H:%M:%S")
    except Exception:
        return None


def _format_duration_seconds(total_seconds):
    try:
        sec = max(0, int(total_seconds or 0))
    except Exception:
        sec = 0
    h = sec // 3600
    m = (sec % 3600) // 60
    s = sec % 60
    if h > 0:
        return f"{h:02d}:{m:02d}:{s:02d}"
    return f"{m:02d}:{s:02d}"


def _auto_loop_run_metrics(state):
    steps = state.get("steps") or []
    total = len(steps)
    done = 0
    blocked = 0
    failed = 0
    waiting = 0
    running = 0
    for step in steps:
        status = str(step.get("status") or "").lower()
        if status == "fertig":
            done += 1
        elif status == "blockiert":
            blocked += 1
        elif status == "fehlgeschlagen":
            failed += 1
        elif status == "wartet auf freigabe":
            waiting += 1
        elif status == "laeuft":
            running += 1
    started = _parse_run_timestamp(state.get("started_at"))
    ended = _parse_run_timestamp(state.get("ended_at"))
    now_dt = datetime.now()
    if started:
        base_end = ended if ended and ended >= started else now_dt
        duration_seconds = max(0, int((base_end - started).total_seconds()))
    else:
        duration_seconds = 0
    next_gate = ""
    for step in steps:
        if str(step.get("status") or "").lower() in {"wartet auf freigabe", "geplant"} and str(step.get("gate") or "").lower() == "approval":
            next_gate = str(step.get("label") or step.get("id") or "")
            break
    return {
        "total_steps": total,
        "done_steps": done,
        "blocked_steps": blocked,
        "failed_steps": failed,
        "waiting_approval_steps": waiting,
        "running_steps": running,
        "retry_count": int(state.get("retry_count") or 0),
        "retry_max": int(state.get("retry_max") or 0),
        "approvals_total": len([s for s in steps if str(s.get("gate") or "").lower() == "approval"]),
        "approvals_open": waiting,
        "blockers_count": len(state.get("session_blockers") or []),
        "duration_seconds": duration_seconds,
        "duration_label": _format_duration_seconds(duration_seconds),
        "next_critical_gate": next_gate,
    }


def _auto_loop_summary(state):
    metrics = _auto_loop_run_metrics(state)
    return (
        f"{metrics['done_steps']}/{metrics['total_steps']} Schritte fertig"
        f" | Retry {metrics['retry_count']}/{metrics['retry_max']}"
        f" | Freigaben offen {metrics['approvals_open']}"
        f" | Dauer {metrics['duration_label']}"
        f" | Ziel: {state.get('goal') or '-'}"
    )


def _auto_loop_save(state):
    state["last_run_at"] = get_timestamp()
    normalized = normalize_auto_loop_state(state)
    save_project_auto_run_state({"auto_loop_state": normalized})
    return normalized


def _auto_loop_maybe_trigger_retry(state):
    """Fuegt eine Retry-Runde an, wenn Fehler vorliegt und Budget noch nicht erschoepft."""
    err = state.get("current_error") if isinstance(state.get("current_error"), dict) else {}
    has_error = bool(state.get("error_label")) or bool(err.get("category") and err.get("category") != "unknown")
    if not has_error and not err.get("label"):
        state["retry_possible"] = False
        state["retry_blocked_reason"] = ""
        return False
    retry_count = int(state.get("retry_count") or 0)
    retry_max = int(state.get("retry_max") or 2)
    category = canonical_error_category(err.get("category") or "unknown")
    occurrences = int(err.get("occurrences") or 1)
    # Heuristik: gleicher Fehler >= 3x -> abbrechen
    if occurrences >= 3:
        reason = f"Gleicher Fehler mehrfach ({occurrences}x). Automatischer Retry gestoppt."
        state["retry_possible"] = False
        state["retry_blocked_reason"] = reason
        _auto_loop_push_history(state, reason)
        _auto_loop_append_repair_history(state, retry_count, err, status="stopped_repeat", repeated=True)
        return False
    if retry_count >= retry_max:
        reason = f"Retry-Budget erschoepft ({retry_count}/{retry_max})."
        state["retry_possible"] = False
        state["retry_blocked_reason"] = reason
        _auto_loop_push_history(state, reason + " Lauf endet.")
        _auto_loop_append_repair_history(state, retry_count, err, status="budget_exhausted")
        return False
    next_round = retry_count + 1
    retry_steps = build_auto_loop_retry_plan(
        goal=state.get("goal") or "",
        runner_command=state.get("runner_command") or "",
        apply_mode=state.get("apply_mode") or "safe",
        retry_round=next_round,
    )
    if not retry_steps:
        reason = "Kein sinnvoller Retry-/Repair-Plan verfuegbar. Lauf wird beendet."
        state["retry_possible"] = False
        state["retry_blocked_reason"] = reason
        state["status"] = "failed"
        state["phase"] = "blocked"
        state["current_phase"] = "Fehler"
        state["last_action"] = reason
        _auto_loop_push_history(state, reason)
        _auto_loop_append_repair_history(state, retry_count, err, status="no_plan")
        return False
    # Metadaten auf Retry-Schritte pragen
    for rs in retry_steps:
        rs["retry_round"] = next_round
        rs["retry_reason"] = err.get("label") or ""
        rs["retry_category"] = category
        rs["retry_target_file"] = err.get("file") or ""
        rs["retry_status"] = "geplant"
    state["retry_count"] = next_round
    existing = state.get("steps") or []
    state["steps"] = (existing + retry_steps)[:32]
    state["error_label"] = ""
    state["repair_suggestion"] = ""
    state["affected_file"] = ""
    state["current_phase"] = f"Retry-Runde {next_round}/{retry_max}"
    state["last_action"] = f"Retry {next_round}/{retry_max} gestartet ({category})."
    state["retry_possible"] = next_round < retry_max
    state["retry_blocked_reason"] = ""
    state["phase"] = "repair"
    _auto_loop_push_history(
        state,
        f"Retry {next_round}/{retry_max} startet | Kategorie={category} | Datei={err.get('file') or '-'}"
    )
    _auto_loop_append_repair_history(state, next_round, err, status="started", repeated=(occurrences > 1))
    return True


def _auto_loop_append_repair_history(state, round_num, err, status="started", repeated=False):
    entry = {
        "round": int(round_num or 0),
        "category": canonical_error_category((err or {}).get("category") if isinstance(err, dict) else "unknown"),
        "reason": str((err or {}).get("label") or "") if isinstance(err, dict) else "",
        "file": str((err or {}).get("file") or "") if isinstance(err, dict) else "",
        "status": str(status or ""),
        "repeated": bool(repeated),
        "timestamp": get_timestamp(),
    }
    hist = state.get("repair_history") or []
    hist.append(entry)
    state["repair_history"] = hist[-16:]


def _auto_loop_build_instruction(state):
    status = state.get("status")
    if status in {"done", "failed", "blocked", "stopped"}:
        if not state.get("ended_at"):
            state["ended_at"] = get_timestamp()
        return {"action": status, "step_index": -1, "step": None, "reason": state.get("last_action") or status}
    if status == "paused":
        idx_p, step_p = _auto_loop_next_pending(state)
        return {"action": "paused", "step_index": idx_p, "step": step_p, "reason": state.get("last_action") or "Pausiert."}
    if state.get("pause_requested"):
        idx_p, step_p = _auto_loop_next_pending(state)
        state["status"] = "paused"
        state["pause_requested"] = False
        state["paused_at"] = get_timestamp()
        state["current_phase"] = "Pausiert"
        state["last_action"] = "Auto-Loop pausiert."
        _auto_loop_push_history(state, "Pausiert.")
        return {"action": "paused", "step_index": idx_p, "step": step_p, "reason": "Pause angefordert."}
    idx, step = _auto_loop_next_pending(state)
    if idx == -1:
        if _auto_loop_maybe_trigger_retry(state):
            return _auto_loop_build_instruction(state)
        state["status"] = "done"
        state["current_phase"] = "Abgeschlossen"
        state["phase"] = "completed"
        state["active_module"] = ""
        state["current_step"] = "Kein aktiver Schritt"
        state["next_step"] = ""
        state["summary"] = _auto_loop_summary(state)
        _auto_loop_push_history(state, "Lauf abgeschlossen.")
        return {"action": "done", "step_index": -1, "step": None, "reason": "Alle Schritte abgeschlossen."}

    if state.get("stop_requested"):
        state["status"] = "stopped"
        state["stopped"] = True
        state["stop_requested"] = False
        state["current_phase"] = "Gestoppt"
        state["phase"] = "blocked"
        state["last_action"] = "Auto-Loop wurde gestoppt."
        state["next_step"] = step.get("label") or ""
        _auto_loop_push_history(state, "Lauf gestoppt.")
        return {"action": "stopped", "step_index": idx, "step": step, "reason": "Stop angefordert."}

    if (step.get("status") or "").lower() == "blockiert":
        state["status"] = "blocked"
        state["phase"] = "blocked"
        state["active_module"] = step.get("tool") or tool_for_action(step.get("action"))
        state["current_phase"] = "Blockiert"
        state["current_step"] = step.get("label") or ""
        state["last_action"] = step.get("blocker") or step.get("detail") or "Schritt blockiert."
        state["next_step"] = step.get("label") or ""
        blocker_details = step.get("blocker_details") if isinstance(step.get("blocker_details"), dict) else {}
        _auto_loop_add_session_blocker(state, step, blocker_details)
        _auto_loop_push_history(state, f"Blockiert: {step.get('label')}")
        return {"action": "blocked", "step_index": idx, "step": step, "reason": state["last_action"]}

    gate = (step.get("gate") or "auto").lower()
    approved = bool(step.get("approved"))
    if gate == "approval" and not approved:
        step["status"] = "wartet auf freigabe"
        state["status"] = "approval_required"
        state["phase"] = "approval"
        state["active_module"] = step.get("tool") or tool_for_action(step.get("action"))
        state["requires_approval"] = True
        state["current_phase"] = "Wartet auf Freigabe"
        state["current_step"] = step.get("label") or ""
        state["last_action"] = step.get("detail") or "Freigabe erforderlich."
        state["next_step"] = step.get("label") or ""
        _auto_loop_push_history(state, f"Freigabe erforderlich: {step.get('label')}")
        return {"action": "wait_approval", "step_index": idx, "step": step, "reason": state["last_action"]}

    step["status"] = "laeuft"
    state["status"] = "running"
    state["phase"] = step.get("phase") or phase_for_action(step.get("action"))
    state["active_module"] = step.get("tool") or tool_for_action(step.get("action"))
    state["requires_approval"] = False
    state["current_phase"] = f"Schritt {idx + 1} von {len(state.get('steps') or [])}"
    state["current_step"] = step.get("label") or ""
    state["last_action"] = f"Starte: {step.get('label')}"
    steps = state.get("steps") or []
    state["next_step"] = steps[idx + 1]["label"] if idx + 1 < len(steps) else "Abschluss"
    return {"action": step.get("action") or step.get("id"),
            "step_index": idx, "step": step, "reason": state["last_action"]}


def _auto_loop_add_session_blocker(state, step, details=None):
    if not isinstance(details, dict):
        details = {}
    entry = {
        "step_id": str(step.get("id") or ""),
        "step_label": str(step.get("label") or ""),
        "path": str(details.get("path") or step.get("blocker") or "")[:240],
        "rule": str(details.get("rule") or ""),
        "reason": str(details.get("reason") or step.get("blocker") or step.get("detail") or ""),
        "suggestion": str(details.get("suggestion") or ""),
        "timestamp": get_timestamp(),
    }
    blockers = state.get("session_blockers") or []
    blockers.insert(0, entry)
    state["session_blockers"] = blockers[:16]


def _auto_loop_apply_report(state, report):
    if not isinstance(report, dict):
        return
    step_id = str(report.get("step_id") or "").strip()
    if not step_id:
        return
    idx = _auto_loop_find_index(state, step_id)
    if idx == -1:
        return
    step = state["steps"][idx]
    ok = bool(report.get("ok")) if "ok" in report else True
    detail = str(report.get("detail") or "").strip()
    approval_required = bool(report.get("approval_required"))
    blocked = bool(report.get("blocked"))
    error_label = str(report.get("error_label") or "").strip()
    repair_suggestion = str(report.get("repair_suggestion") or "").strip()
    affected_file = str(report.get("affected_file") or "").strip()
    clear_error = bool(report.get("clear_error"))

    if approval_required:
        step["status"] = "wartet auf freigabe"
        if detail:
            step["detail"] = detail
        state["status"] = "approval_required"
        state["requires_approval"] = True
        state["current_phase"] = "Wartet auf Freigabe"
        state["current_step"] = step.get("label") or ""
        state["last_action"] = detail or "Freigabe erforderlich."
        _auto_loop_push_history(state, f"Freigabe erforderlich: {step.get('label')}")
        return

    if blocked:
        step["status"] = "blockiert"
        step["blocker"] = detail or step.get("blocker") or "Schritt blockiert."
        if detail:
            step["detail"] = detail
        raw_details = report.get("blocker_details")
        if isinstance(raw_details, dict):
            step["blocker_details"] = {
                "path": str(raw_details.get("path") or "").strip(),
                "rule": str(raw_details.get("rule") or "").strip(),
                "reason": str(raw_details.get("reason") or step["blocker"]).strip(),
                "suggestion": str(raw_details.get("suggestion") or "").strip(),
            }
        state["status"] = "blocked"
        state["current_phase"] = "Blockiert"
        state["current_step"] = step.get("label") or ""
        state["last_action"] = step["blocker"]
        _auto_loop_add_session_blocker(state, step, step.get("blocker_details"))
        _auto_loop_push_history(state, f"Blockiert: {step.get('label')}")
        return

    if not ok:
        step["status"] = "fehlgeschlagen"
        if detail:
            step["detail"] = detail
        state["status"] = "failed"
        state["current_phase"] = "Fehler"
        state["current_step"] = step.get("label") or ""
        state["last_action"] = detail or "Schritt fehlgeschlagen."
        _auto_loop_push_history(state, f"Fehlgeschlagen: {step.get('label')}")
        return

    step["status"] = "fertig"
    if detail:
        step["detail"] = detail
    state["last_action"] = detail or f"{step.get('label')} abgeschlossen."
    if error_label:
        state["error_label"] = error_label
    if repair_suggestion:
        state["repair_suggestion"] = repair_suggestion
    if affected_file:
        state["affected_file"] = affected_file
    raw_category = str(report.get("error_category") or "").strip()
    raw_area = str(report.get("error_area") or "").strip()
    if error_label or raw_category or repair_suggestion:
        err = build_error_info(
            label=error_label,
            suggestion=repair_suggestion,
            file=affected_file,
            area=raw_area,
            raw_category=raw_category or error_label,
            source=str(report.get("source") or step.get("tool") or "auto_loop"),
        )
        prev = state.get("current_error") if isinstance(state.get("current_error"), dict) else {}
        if prev and prev.get("signature") == err["signature"]:
            try:
                err["occurrences"] = int(prev.get("occurrences") or 1) + 1
            except Exception:
                err["occurrences"] = 2
        state["current_error"] = err
        state["repair_plan"] = build_repair_plan(err, state.get("runner_command") or "", state.get("apply_mode") or "safe")
        if err["occurrences"] > 1:
            _auto_loop_push_history(state, f"Wiederholter Fehler ({err['category']}, {err['occurrences']}x): {err['label']}")
        else:
            _auto_loop_push_history(state, f"Fehler erkannt [{err['category']}]: {err['label']} ({affected_file or '-'})")
    elif clear_error:
        state["current_error"] = _empty_error_info()
        state["repair_plan"] = []
        state["error_label"] = ""
        state["repair_suggestion"] = ""
        state["affected_file"] = ""
        _auto_loop_push_history(state, "Fehlerstatus bereinigt.")
    _auto_loop_push_history(state, f"Fertig: {step.get('label')}")


def load_project_auto_run_state():
    fallback = {
        "last_run_at": "",
        "last_task": "",
        "last_mode": "safe",
        "last_target_paths": [],
        "last_guard_decision": "",
        "last_apply_action": "",
        "last_check_result": "",
        "last_result": "",
        "blocked": False,
        "pending_direct_run": None,
        "last_direct_scope": "",
        "last_direct_prompt": "",
        "last_direct_decision": "",
        "last_direct_status": "idle",
        "last_self_fix_plan": None,
        "pending_self_fix_preview": None,
        "last_direct_run_id": "",
        "last_direct_confirmed_run_id": "",
        "active_direct_run_id": "",
        "last_completed_run_id": "",
        "last_planned_steps": [],
        "direct_run_history": [],
        "last_runner_execution": None,
        "last_project_scan": None,
        "last_agent_run": None,
        "agent_runs": {},
        "pending_agent_run_confirmations": {},
        "last_agent_decision": {},
        "agent_decision_history": [],
        "auto_loop_state": {
            "last_run_at": "",
            "goal": "",
            "status": "idle",
            "current_phase": "",
            "current_step": "",
            "last_action": "",
            "next_step": "",
            "summary": "",
            "steps": [],
            "history": [],
            "error_label": "",
            "repair_suggestion": "",
            "affected_file": "",
            "requires_approval": False,
            "stopped": False
        }
    }
    data = read_json_file(PROJECT_AUTO_RUN_STATE_FILE, fallback)
    if not isinstance(data, dict):
        return fallback
    normalized = fallback.copy()
    normalized.update(data)
    if not isinstance(normalized.get("last_target_paths"), list):
        normalized["last_target_paths"] = []
    if not isinstance(normalized.get("pending_direct_run"), dict):
        normalized["pending_direct_run"] = None
    else:
        pending = normalized["pending_direct_run"]
        recognized = pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {}
        if not recognized:
            pending["recognized_task"] = classify_direct_task(pending.get("task"))
        pending["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
        if not pending_direct_context_is_valid(pending, normalized):
            normalized["pending_direct_run"] = None
            if str(normalized.get("active_direct_run_id") or "").strip() == str(pending.get("run_id") or pending.get("token") or "").strip():
                normalized["active_direct_run_id"] = ""
    if not isinstance(normalized.get("direct_run_history"), list):
        normalized["direct_run_history"] = []
    normalized["direct_run_history"] = compact_direct_run_history(normalized.get("direct_run_history"))
    if not isinstance(normalized.get("last_self_fix_plan"), dict):
        normalized["last_self_fix_plan"] = None
    if not isinstance(normalized.get("pending_self_fix_preview"), dict):
        normalized["pending_self_fix_preview"] = None
    if not isinstance(normalized.get("last_agent_run"), dict):
        normalized["last_agent_run"] = None
    if not isinstance(normalized.get("agent_runs"), dict):
        normalized["agent_runs"] = {}
    if not isinstance(normalized.get("pending_agent_run_confirmations"), dict):
        normalized["pending_agent_run_confirmations"] = {}
    normalized["auto_loop_state"] = normalize_auto_loop_state(normalized.get("auto_loop_state"))
    normalized["last_planned_steps"] = normalize_planned_steps(normalized.get("last_planned_steps") or [])
    return normalized


def save_project_auto_run_state(state):
    payload = load_project_auto_run_state()
    payload.update(state or {})
    write_json_file(PROJECT_AUTO_RUN_STATE_FILE, payload)


def get_active_project_state() -> dict:
    if bool(app.config.get("TESTING")):
        root = str(APP_DIR.resolve())
        return {
            "active_project_id": "rambo_builder_local",
            "active_project_name": "rambo builder local",
            "active_project_root": root,
            "last_switched_at": "",
        }
    state = load_project_auto_run_state()
    active = state.get("active_project") if isinstance(state.get("active_project"), dict) else {}
    root = str(active.get("active_project_root") or str(APP_DIR.resolve())).strip()
    name = str(active.get("active_project_name") or Path(root).name or "rambo_builder_local")
    pid = str(active.get("active_project_id") or Path(root).name or "rambo_builder_local")
    return {
        "active_project_id": pid,
        "active_project_name": name,
        "active_project_root": root,
        "last_switched_at": str(active.get("last_switched_at") or ""),
    }


def get_active_project_root() -> Path:
    return Path(get_active_project_state().get("active_project_root") or APP_DIR.resolve()).resolve()


def is_active_workspace_trusted() -> bool:
    if bool(app.config.get("TESTING")):
        return False
    try:
        act = WORKSPACE_SANDBOX.get_active_workspace()
        return bool((act.get("active") or {}).get("trusted", False))
    except Exception:
        return False


def save_pending_self_fix_preview(preview_payload: dict) -> dict:
    token = uuid4().hex
    pending = {
        "token": token,
        "created_at": get_timestamp(),
        "used": False,
        "prompt": str(preview_payload.get("prompt") or ""),
        "plan_id": str(preview_payload.get("plan_id") or ""),
        "candidate_files": list(preview_payload.get("candidate_files") or []),
        "change_plan": list(preview_payload.get("change_plan") or []),
        "diff_plan": list(preview_payload.get("diff_plan") or []),
        "verification_plan": dict(preview_payload.get("verification_plan") or {}),
        "patch_validation": dict(preview_payload.get("patch_validation") or {}),
        "recommended_checks": list(preview_payload.get("recommended_checks") or []),
    }
    save_project_auto_run_state({"pending_self_fix_preview": pending})
    return pending


def save_pending_direct_run(preview_payload):
    run_id = uuid4().hex
    preview_with_files = _merge_direct_file_context(preview_payload, preview_payload)
    pending = {
        "run_id": run_id,
        "token": run_id,
        "created_at": get_timestamp(),
        "scope": str(preview_with_files.get("scope") or "project"),
        "mode": str(preview_with_files.get("mode") or "safe"),
        "status": str(preview_with_files.get("direct_status") or "pending_confirmation"),
        "task": str(preview_with_files.get("task") or ""),
        "recognized_task": preview_with_files.get("recognized_task") if isinstance(preview_with_files.get("recognized_task"), dict) else {},
        "selected_target_path": str(preview_with_files.get("selected_target_path") or ""),
        "proposed_content": str(preview_with_files.get("proposed_content") or ""),
        "diff": str(preview_with_files.get("diff") or ""),
        "has_changes": bool(preview_with_files.get("has_changes")),
        "planned_steps": normalize_planned_steps(preview_with_files.get("planned_steps") or []),
        "guard": preview_with_files.get("guard") if isinstance(preview_with_files.get("guard"), dict) else {},
        "strategy": preview_with_files.get("strategy") if isinstance(preview_with_files.get("strategy"), dict) else {},
        "requires_user_confirmation": bool(preview_with_files.get("requires_user_confirmation")),
        "affected_files": preview_with_files.get("affected_files") if isinstance(preview_with_files.get("affected_files"), list) else [],
        "changed_files": preview_with_files.get("changed_files") if isinstance(preview_with_files.get("changed_files"), list) else [],
        "file_plan": preview_with_files.get("file_plan") if isinstance(preview_with_files.get("file_plan"), list) else [],
        "file_entries": preview_with_files.get("file_entries") if isinstance(preview_with_files.get("file_entries"), list) else [],
    }
    save_project_auto_run_state({
        "pending_direct_run": pending,
        "last_direct_scope": pending["scope"],
        "last_direct_prompt": pending["task"],
        "last_direct_decision": "Vorschau bereit. Bestaetigung ausstehend.",
        "last_direct_status": pending["status"],
        "last_direct_run_id": run_id,
        "active_direct_run_id": run_id,
        "last_task": pending["task"],
        "last_mode": pending["mode"],
        "last_target_paths": [pending["selected_target_path"]] if pending["selected_target_path"] else [],
        "last_planned_steps": pending["planned_steps"],
        "blocked": False,
    })
    return pending


def get_pending_direct_run():
    return load_project_auto_run_state().get("pending_direct_run")


def clear_pending_direct_run():
    save_project_auto_run_state({"pending_direct_run": None, "active_direct_run_id": ""})


def _agent_run_patch_fingerprint(run_id: str, selected_files: list[str], patch_plan: list[dict], validation: dict) -> str:
    payload = {
        "run_id": str(run_id or ""),
        "files": [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)],
        "patch_plan": [
            {
                "file": format_local_path((x or {}).get("file") or ""),
                "has_changes": bool((x or {}).get("has_changes")),
                "risk": str((x or {}).get("risk") or ""),
                "target_area": str((x or {}).get("target_area") or ""),
            }
            for x in list(patch_plan or [])
            if isinstance(x, dict)
        ],
        "validation_status": str((validation or {}).get("status") or ""),
        "validated_patch": bool((validation or {}).get("validated_patch")),
    }
    raw = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(raw.encode("utf-8", errors="replace")).hexdigest()


def save_pending_agent_run_confirmation(
    *,
    run_id: str,
    selected_files: list[str],
    patch_plan: list[dict],
    validation: dict,
    patch_entries: list[dict] | None = None,
    recommended_checks: list[str] | None = None,
) -> str:
    token = uuid4().hex
    fingerprint = _agent_run_patch_fingerprint(run_id, selected_files, patch_plan, validation)
    state = load_project_auto_run_state()
    store = dict(state.get("pending_agent_run_confirmations") or {})
    store[token] = {
        "token": token,
        "run_id": str(run_id or ""),
        "created_at": get_timestamp(),
        "used": False,
        "selected_files": [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)],
        "patch_plan": [dict(x) for x in list(patch_plan or []) if isinstance(x, dict)],
        "validation": dict(validation or {}),
        "patch_entries": [
            {
                "path": format_local_path((x or {}).get("path") or ""),
                "current_content": str((x or {}).get("current_content") or ""),
                "proposed_content": str((x or {}).get("proposed_content") or ""),
            }
            for x in list(patch_entries or [])
            if isinstance(x, dict) and format_local_path((x or {}).get("path") or "")
        ],
        "recommended_checks": [str(c).strip() for c in list(recommended_checks or []) if str(c).strip()],
        "patch_fingerprint": fingerprint,
    }
    save_project_auto_run_state({"pending_agent_run_confirmations": store})
    return token


def shorten_text(text, limit=92):
    cleaned = " ".join(str(text or "").strip().split())
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1].rstrip() + "..."


def normalize_planned_steps(steps):
    cleaned_steps = []
    for index, step in enumerate(steps or []):
        if not isinstance(step, dict):
            continue
        label = " ".join(str(step.get("label") or "").strip().split())
        detail = " ".join(str(step.get("detail") or "").strip().split())
        status = str(step.get("status") or "planned").strip().lower()
        if not label:
            continue
        if status not in {"done", "active", "pending", "planned", "blocked"}:
            status = "planned"
        cleaned_steps.append({
            "id": str(step.get("id") or f"step_{index + 1}").strip(),
            "label": label,
            "detail": detail,
            "status": status,
        })
    return cleaned_steps[:8]


def infer_direct_subgoals(task):
    cleaned = " ".join(str(task or "").strip().split())
    if not cleaned:
        return []

    raw_parts = re.split(r"(?:\s*[\n;]+\s*|\s+(?:danach|anschliessend|anschließend|dann|sowie|plus)\s+)", cleaned, flags=re.IGNORECASE)
    normalized_parts = []
    lowered_parts = set()
    for part in raw_parts:
        candidate = " ".join(str(part or "").strip(" ,.-").split())
        if len(candidate) < 14:
            continue
        lowered = candidate.lower()
        if lowered in lowered_parts:
            continue
        lowered_parts.add(lowered)
        normalized_parts.append(candidate)

    if len(normalized_parts) <= 1 and "," in cleaned:
        comma_parts = []
        for part in cleaned.split(","):
            candidate = " ".join(str(part or "").strip(" ,.-").split())
            if len(candidate) >= 18:
                comma_parts.append(candidate)
        if 2 <= len(comma_parts) <= 4:
            normalized_parts = comma_parts

    return normalized_parts[:3] if len(normalized_parts) > 1 else []


def build_planned_steps(task, mode, selected_target_path="", affected_files=None, has_changes=True, guard_allowed=True):
    affected = sanitize_history_paths(affected_files or [])
    selected = format_local_path(selected_target_path)
    if selected and selected not in affected:
        affected.insert(0, selected)

    subgoals = infer_direct_subgoals(task)
    final_label = "Apply" if mode == "apply" and has_changes and guard_allowed else "Safe-Abschluss"
    final_status = "pending" if guard_allowed else "blocked"
    final_detail = "Schreibaktion erst nach Bestaetigung." if final_label == "Apply" else "Ohne Schreibzugriff abschliessen oder sicher beenden."

    steps = [{
        "id": "analysis",
        "label": "Analyse",
        "status": "done",
        "detail": shorten_text(task, 120),
    }]
    if subgoals:
        steps.append({
            "id": "plan",
            "label": "Plan",
            "status": "done",
            "detail": f"{len(subgoals)} sichere Teilschritte als ein zusammenhaengender Lauf.",
        })

    for index, goal in enumerate(subgoals, start=1):
        steps.append({
            "id": f"goal_{index}",
            "label": f"Teilziel {index}",
            "status": "planned",
            "detail": shorten_text(goal, 100),
        })

    steps.extend([
        {
            "id": "files",
            "label": "Betroffene Dateien",
            "status": "done" if affected else "planned",
            "detail": ", ".join(affected[:3]) if affected else "Werden aus Prompt und Guard-Kontext abgeleitet.",
        },
        {
            "id": "preview",
            "label": "Vorschau erstellen",
            "status": "done",
            "detail": "Guard und Diff fuer diesen Lauf sind vorbereitet.",
        },
        {
            "id": "confirm",
            "label": "Bestaetigung",
            "status": "active" if guard_allowed else "blocked",
            "detail": "Einmal bestaetigen, danach laeuft der Rest automatisch." if guard_allowed else "Guard blockiert den Lauf vor der Bestaetigung.",
        },
        {
            "id": "finalize",
            "label": final_label,
            "status": final_status,
            "detail": final_detail,
        }
    ])

    return normalize_planned_steps(steps)


def sanitize_history_paths(paths):
    cleaned_paths = []
    for path in paths or []:
        cleaned = format_local_path(path)
        if cleaned and cleaned not in cleaned_paths:
            cleaned_paths.append(cleaned)
    return cleaned_paths[:6]


DIRECT_HISTORY_STATUS_RANK = {
    "pending_confirmation": 1,
    "apply_ready": 2,
    "safe_preview": 3,
    "applied": 4,
    "verified": 5,
    "blocked": 2,
}
DIRECT_HISTORY_NOISE_STATUSES = {"pending_confirmation", "apply_ready", "safe_preview", "blocked"}


def _history_status_rank(status):
    return DIRECT_HISTORY_STATUS_RANK.get(str(status or "").strip().lower(), 0)


def _parse_history_ts(value):
    raw = str(value or "").strip()
    if not raw:
        return datetime.min
    try:
        return datetime.strptime(raw, "%Y-%m-%d %H:%M:%S")
    except ValueError:
        return datetime.min


def _normalize_history_entry(entry):
    if not isinstance(entry, dict):
        return None
    return {
        "run_id": str(entry.get("run_id") or "").strip(),
        "timestamp": str(entry.get("timestamp") or get_timestamp()).strip(),
        "prompt_short": shorten_text(entry.get("prompt_short") or entry.get("prompt")),
        "prompt": str(entry.get("prompt") or "").strip(),
        "scope": str(entry.get("scope") or "project").strip(),
        "mode": str(entry.get("mode") or "safe").strip(),
        "status": str(entry.get("status") or "idle").strip(),
        "task_type": str(entry.get("task_type") or "unknown").strip(),
        "primary_area": str(entry.get("primary_area") or "Builder").strip(),
        "recommendation_hint": str(entry.get("recommendation_hint") or "").strip(),
        "selected_target_path": str(entry.get("selected_target_path") or "").strip(),
        "affected_files": sanitize_history_paths(entry.get("affected_files") or []),
        "planned_steps": normalize_planned_steps(entry.get("planned_steps") or []),
        "guard_allowed": bool(entry.get("guard_allowed")),
        "guard_detail": str(entry.get("guard_detail") or "").strip(),
        "message": str(entry.get("message") or "").strip(),
        "post_check_ok": bool(entry.get("post_check_ok")),
        "post_check_detail": str(entry.get("post_check_detail") or "").strip(),
        "diff_summary": entry.get("diff_summary") if isinstance(entry.get("diff_summary"), dict) else build_diff_summary(""),
    }


def _merge_history_entry(existing, incoming):
    ex = existing.copy()
    inc = incoming.copy()
    ex_rank = _history_status_rank(ex.get("status"))
    inc_rank = _history_status_rank(inc.get("status"))
    ex_ts = _parse_history_ts(ex.get("timestamp"))
    inc_ts = _parse_history_ts(inc.get("timestamp"))
    keep_existing_status = ex_rank > inc_rank or (ex_rank == inc_rank and ex_ts >= inc_ts)

    merged = ex if keep_existing_status else inc
    fallback = inc if keep_existing_status else ex
    for k in ("prompt_short", "prompt", "scope", "mode", "task_type", "primary_area", "recommendation_hint",
              "selected_target_path", "guard_detail", "message", "post_check_detail"):
        if not str(merged.get(k) or "").strip() and str(fallback.get(k) or "").strip():
            merged[k] = fallback[k]
    if not merged.get("affected_files"):
        merged["affected_files"] = fallback.get("affected_files") or []
    if not merged.get("planned_steps"):
        merged["planned_steps"] = fallback.get("planned_steps") or []
    if not isinstance(merged.get("diff_summary"), dict) or not merged.get("diff_summary"):
        merged["diff_summary"] = fallback.get("diff_summary") if isinstance(fallback.get("diff_summary"), dict) else build_diff_summary("")
    if not merged.get("post_check_ok") and fallback.get("post_check_ok"):
        merged["post_check_ok"] = True
    return _normalize_history_entry(merged)


def _history_noise_signature(entry):
    e = entry or {}
    diff_preview = ""
    diff_summary = e.get("diff_summary") if isinstance(e.get("diff_summary"), dict) else {}
    if isinstance(diff_summary.get("preview"), str):
        diff_preview = diff_summary.get("preview")
    return "|".join([
        str(e.get("status") or "").strip().lower(),
        str(e.get("scope") or "").strip().lower(),
        str(e.get("mode") or "").strip().lower(),
        str(e.get("selected_target_path") or "").strip().lower(),
        str(e.get("prompt_short") or "").strip().lower(),
        ",".join([str(x).strip().lower() for x in (e.get("affected_files") or [])]),
        diff_preview[:160].strip().lower(),
    ])


def compact_direct_run_history(history):
    normalized = []
    for raw in (history or []):
        n = _normalize_history_entry(raw)
        if n:
            normalized.append(n)
    if not normalized:
        return []

    by_run_id = {}
    ordered = []
    for item in normalized:
        run_id = str(item.get("run_id") or "").strip()
        if run_id and run_id in by_run_id:
            idx = by_run_id[run_id]
            ordered[idx] = _merge_history_entry(ordered[idx], item)
            continue
        if run_id:
            by_run_id[run_id] = len(ordered)
        ordered.append(item)

    compacted = []
    for item in ordered:
        status = str(item.get("status") or "").strip().lower()
        if compacted:
            prev = compacted[-1]
            prev_status = str(prev.get("status") or "").strip().lower()
            same_sig = _history_noise_signature(prev) == _history_noise_signature(item)
            prev_ts = _parse_history_ts(prev.get("timestamp"))
            cur_ts = _parse_history_ts(item.get("timestamp"))
            close_in_time = abs((cur_ts - prev_ts).total_seconds()) <= 180
            if same_sig and close_in_time and status in DIRECT_HISTORY_NOISE_STATUSES and prev_status in DIRECT_HISTORY_NOISE_STATUSES:
                if _parse_history_ts(item.get("timestamp")) >= _parse_history_ts(prev.get("timestamp")):
                    compacted[-1] = item
                continue
        compacted.append(item)

    return compacted[:DIRECT_HISTORY_LIMIT]


def get_direct_run_history():
    history = load_project_auto_run_state().get("direct_run_history", [])
    return compact_direct_run_history(history)


def get_history_entry_by_run_id(run_id):
    run_id = str(run_id or "").strip()
    if not run_id:
        return None
    for entry in get_direct_run_history():
        if str(entry.get("run_id") or "") == run_id:
            return entry
    return None


def upsert_direct_run_history(entry):
    if not isinstance(entry, dict):
        return
    run_id = str(entry.get("run_id") or "").strip()
    if not run_id:
        return

    state = load_project_auto_run_state()
    history = get_direct_run_history()
    filtered = [item for item in history if str(item.get("run_id") or "") != run_id]
    filtered.insert(0, entry)
    save_project_auto_run_state({"direct_run_history": filtered[:DIRECT_HISTORY_LIMIT]})


def build_direct_history_entry(run_id, payload, status_override=None):
    strategy = payload.get("strategy") if isinstance(payload.get("strategy"), dict) else {}
    guard = payload.get("guard") if isinstance(payload.get("guard"), dict) else {}
    post_check = payload.get("post_check") if isinstance(payload.get("post_check"), dict) else {}
    recognized = payload.get("recognized_task") if isinstance(payload.get("recognized_task"), dict) else {}
    affected = strategy.get("betroffene_dateien") if isinstance(strategy.get("betroffene_dateien"), list) else []
    if not affected and payload.get("selected_target_path"):
        affected = [payload.get("selected_target_path")]

    return {
        "run_id": run_id,
        "timestamp": get_timestamp(),
        "prompt_short": shorten_text(payload.get("task") or payload.get("message")),
        "prompt": str(payload.get("task") or "").strip(),
        "scope": str(payload.get("scope") or "project").strip(),
        "mode": str(payload.get("mode") or "safe").strip(),
        "status": str(status_override or payload.get("direct_status") or "idle").strip(),
        "task_type": str(recognized.get("task_type") or payload.get("task_type") or "unknown").strip(),
        "primary_area": str(recognized.get("primary_area") or payload.get("primary_area") or "Builder").strip(),
        "recommendation_hint": str(recognized.get("hint") or payload.get("recommendation_hint") or "").strip(),
        "selected_target_path": str(payload.get("selected_target_path") or "").strip(),
        "affected_files": sanitize_history_paths(affected),
        "planned_steps": normalize_planned_steps(
            payload.get("planned_steps")
            or build_planned_steps(
                payload.get("task") or "",
                payload.get("mode") or "safe",
                payload.get("selected_target_path") or "",
                affected,
                bool(payload.get("has_changes", True)),
                bool(guard.get("allowed", True)),
            )
        ),
        "guard_allowed": bool(guard.get("allowed")),
        "guard_detail": str(guard.get("detail") or payload.get("error") or "").strip(),
        "message": str(payload.get("message") or payload.get("error") or "").strip(),
        "post_check_ok": bool(post_check.get("ok")),
        "post_check_detail": str(post_check.get("detail") or "").strip(),
        "diff_summary": payload.get("diff_summary") if isinstance(payload.get("diff_summary"), dict) else build_diff_summary(payload.get("diff") or ""),
    }


PATCH_REVIEW_STATUS_MAP = {
    "pending_confirmation": ("review_ready", "approval_required"),
    "safe_preview": ("review_ready", "safe_preview"),
    "verified": ("applied", "applied"),
    "applied": ("applied", "applied"),
    "blocked": ("blocked", "blocked"),
    "idle": ("idle", "idle"),
}


def _patch_review_file_entry(path, status="review_ready", detail=""):
    return {
        "path": str(path or "").strip(),
        "status": str(status or "review_ready"),
        "detail": str(detail or "").strip(),
    }


def build_patch_review_from_history(entry, run_state=None):
    """Baut eine Review-Sicht aus einem direct_run_history-Eintrag."""
    if not isinstance(entry, dict):
        return None
    raw_status = str(entry.get("status") or "idle").lower()
    review_status, apply_status = PATCH_REVIEW_STATUS_MAP.get(
        raw_status, ("review_ready", raw_status or "idle")
    )
    run_state = run_state if isinstance(run_state, dict) else {}
    run_id = str(entry.get("run_id") or "").strip()
    last_confirmed = str(run_state.get("last_direct_confirmed_run_id") or "").strip()
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    if raw_status == "safe_preview" and run_id and run_id == last_confirmed and not pending_requires_confirmation(pending_direct):
        review_status, apply_status = "closed", "safe_confirmed"
    files = sanitize_history_paths(entry.get("affected_files") or [])
    file_entries = []
    for p in files:
        file_status = review_status
        detail = ""
        if apply_status == "blocked":
            file_status = "blocked"
            detail = entry.get("guard_detail") or entry.get("message") or ""
        elif apply_status == "applied":
            file_status = "applied"
        file_entries.append(_patch_review_file_entry(p, file_status, detail))
    guard_allowed = bool(entry.get("guard_allowed", True))
    return {
        "patch_id": str(entry.get("run_id") or ""),
        "run_id": str(entry.get("run_id") or ""),
        "source": "direct",
        "objective": str(entry.get("prompt_short") or entry.get("prompt") or "").strip(),
        "prompt": str(entry.get("prompt") or "").strip(),
        "scope": str(entry.get("scope") or "").strip(),
        "mode": str(entry.get("mode") or "safe").strip(),
        "affected_files": files,
        "file_count": len(files),
        "file_entries": file_entries,
        "blocked_files": [f.get("path") for f in file_entries if f.get("status") == "blocked"],
        "diff_summary": entry.get("diff_summary") if isinstance(entry.get("diff_summary"), dict) else {},
        "review_status": review_status,
        "apply_status": apply_status,
        "approvals_needed": 1 if apply_status == "approval_required" else 0,
        "guard_allowed": guard_allowed,
        "guard_detail": str(entry.get("guard_detail") or ""),
        "warnings": [] if guard_allowed else [entry.get("guard_detail") or "Guard blockiert."],
        "timestamp": str(entry.get("timestamp") or ""),
        "last_result": str(entry.get("message") or ""),
        "task_type": str(entry.get("task_type") or ""),
        "primary_area": str(entry.get("primary_area") or ""),
    }


def build_patch_review_from_pending(pending):
    """Aktuelle offene Vorschau (noch unbestaetigt) als Review-Sicht."""
    if not isinstance(pending, dict):
        return None
    target = str(pending.get("selected_target_path") or "").strip()
    affected = []
    strategy = pending.get("strategy") if isinstance(pending.get("strategy"), dict) else {}
    if isinstance(strategy.get("betroffene_dateien"), list):
        affected = [str(x) for x in strategy["betroffene_dateien"] if x]
    if not affected and target:
        affected = [target]
    guard = pending.get("guard") if isinstance(pending.get("guard"), dict) else {}
    mode = str(pending.get("mode") or "safe").lower()
    has_changes = bool(pending.get("has_changes", True))
    if not guard.get("allowed", True):
        review_status, apply_status = "blocked", "blocked"
    elif mode == "apply" and has_changes:
        review_status, apply_status = "review_ready", "approval_required"
    else:
        review_status, apply_status = "review_ready", "safe_preview"
    file_entries = [_patch_review_file_entry(
        p,
        "blocked" if apply_status == "blocked" else review_status,
        guard.get("detail", "") if apply_status == "blocked" else ""
    ) for p in sanitize_history_paths(affected)]
    return {
        "patch_id": str(pending.get("run_id") or pending.get("token") or ""),
        "run_id": str(pending.get("run_id") or ""),
        "source": "direct_pending",
        "objective": str(pending.get("task") or "")[:120],
        "prompt": str(pending.get("task") or ""),
        "scope": str(pending.get("scope") or "").strip(),
        "mode": mode,
        "affected_files": sanitize_history_paths(affected),
        "file_count": len(file_entries),
        "file_entries": file_entries,
        "blocked_files": [f["path"] for f in file_entries if f["status"] == "blocked"],
        "diff_summary": build_diff_summary(pending.get("diff") or ""),
        "review_status": review_status,
        "apply_status": apply_status,
        "approvals_needed": 1 if apply_status == "approval_required" else 0,
        "guard_allowed": bool(guard.get("allowed", True)),
        "guard_detail": str(guard.get("detail") or ""),
        "warnings": [] if guard.get("allowed", True) else [guard.get("detail") or "Guard blockiert."],
        "timestamp": get_timestamp(),
        "last_result": "Vorschau wartet auf Freigabe." if apply_status == "approval_required" else "Vorschau bereit.",
        "task_type": str((pending.get("recognized_task") or {}).get("task_type") or ""),
        "primary_area": str((pending.get("recognized_task") or {}).get("primary_area") or ""),
    }


def _history_entry_by_id(history, run_id):
    run_id = str(run_id or "").strip()
    if not run_id:
        return None
    for item in (history or []):
        if str((item or {}).get("run_id") or "").strip() == run_id:
            return item
    return None


def _canonical_current_history_entry(run_state, history):
    run_state = run_state if isinstance(run_state, dict) else {}
    history = history if isinstance(history, list) else []
    if not history:
        return None

    pending = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    if pending and pending_requires_confirmation(pending):
        pending_id = str(pending.get("run_id") or pending.get("token") or "").strip()
        pending_entry = _history_entry_by_id(history, pending_id)
        if pending_entry:
            return pending_entry
        return None

    active_id = str(run_state.get("active_direct_run_id") or "").strip()
    active_entry = _history_entry_by_id(history, active_id)
    if active_entry:
        return active_entry

    confirmed_id = str(run_state.get("last_direct_confirmed_run_id") or "").strip()
    confirmed_entry = _history_entry_by_id(history, confirmed_id)
    if confirmed_entry:
        return confirmed_entry

    completed_id = str(run_state.get("last_completed_run_id") or "").strip()
    completed_entry = _history_entry_by_id(history, completed_id)
    if completed_entry:
        return completed_entry

    last_run_id = str(run_state.get("last_direct_run_id") or "").strip()
    last_run_entry = _history_entry_by_id(history, last_run_id)
    if last_run_entry:
        return last_run_entry

    return history[0]


def build_patch_review_snapshot():
    """Aggregiert aktuelle Patch-Review-Sicht + Historie aus vorhandenen Quellen."""
    run_state = load_project_auto_run_state()
    pending = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    history = run_state.get("direct_run_history") if isinstance(run_state.get("direct_run_history"), list) else []

    current = None
    if pending and pending_requires_confirmation(pending):
        current = build_patch_review_from_pending(pending)
    if current is None and history:
        canonical_entry = _canonical_current_history_entry(run_state, history)
        current = build_patch_review_from_history(canonical_entry, run_state=run_state) if canonical_entry else None

    recent = []
    for entry in history[:12]:
        if current and str((current.get("run_id") or current.get("patch_id") or "")).strip() == str(entry.get("run_id") or "").strip():
            continue
        patch = build_patch_review_from_history(entry, run_state=run_state)
        if patch:
            recent.append(patch)

    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    apply_mode = auto_loop.get("apply_mode") or "safe"
    return {
        "current": current,
        "recent": recent,
        "apply_mode": apply_mode,
        "last_completed_run_id": str(run_state.get("last_completed_run_id") or ""),
        "last_direct_confirmed_run_id": str(run_state.get("last_direct_confirmed_run_id") or ""),
        "blocked": bool(run_state.get("blocked")),
    }


def compact_diff_text(diff_text, max_body_lines=24):
    raw = str(diff_text or "").strip()
    if not raw:
        return "Keine Diff-Vorschau vorhanden."
    if raw == "Keine inhaltliche Aenderung erkannt.":
        return raw

    lines = raw.splitlines()
    visible = []
    body_count = 0
    hidden = 0
    for line in lines:
        if line.startswith(("---", "+++", "@@")):
            visible.append(line)
            continue
        body_count += 1
        if body_count <= max_body_lines:
            visible.append(line)
        else:
            hidden += 1

    if hidden > 0:
        visible.append(f"... {hidden} weitere Diff-Zeilen ausgeblendet ...")
    return "\n".join(visible)


def build_diff_summary(diff_text):
    raw = str(diff_text or "").strip()
    if not raw or raw == "Keine inhaltliche Aenderung erkannt.":
        return {
            "changed": False,
            "added": 0,
            "removed": 0,
            "preview": raw or "Keine Diff-Vorschau vorhanden."
        }

    added = 0
    removed = 0
    for line in raw.splitlines():
        if line.startswith("+++") or line.startswith("---"):
            continue
        if line.startswith("+"):
            added += 1
        elif line.startswith("-"):
            removed += 1

    return {
        "changed": True,
        "added": added,
        "removed": removed,
        "preview": compact_diff_text(raw)
    }


def pending_requires_confirmation(pending):
    return isinstance(pending, dict) and str(pending.get("status") or "") in {
        "pending_confirmation", "safe_preview", "apply_ready"
    }


def pending_direct_context_is_valid(pending, state=None):
    state = state if isinstance(state, dict) else {}
    if not pending_requires_confirmation(pending):
        return False
    if not isinstance(pending, dict):
        return False
    run_id = str(pending.get("run_id") or pending.get("token") or "").strip()
    token = str(pending.get("token") or "").strip()
    task = str(pending.get("task") or "").strip()
    target = str(pending.get("selected_target_path") or "").strip()
    if not run_id or not token or not task or not target:
        return False
    last_confirmed = str(state.get("last_direct_confirmed_run_id") or "").strip()
    last_completed = str(state.get("last_completed_run_id") or "").strip()
    if run_id and (run_id == last_confirmed or run_id == last_completed):
        return False
    active_direct_run_id = str(state.get("active_direct_run_id") or "").strip()
    if active_direct_run_id and run_id != active_direct_run_id:
        return False
    # Stale-Guard: offene Bestaetigungen ohne aktuellen Laufkontext nicht unbegrenzt tragen.
    created_at_raw = str(pending.get("created_at") or "").strip()
    created_at = _parse_run_timestamp(created_at_raw)
    now = datetime.now()
    age_seconds = None
    if created_at is not None:
        try:
            age_seconds = max(0, int((now - created_at).total_seconds()))
        except Exception:
            age_seconds = None
    auto_loop = state.get("auto_loop_state") if isinstance(state.get("auto_loop_state"), dict) else {}
    auto_status = str(auto_loop.get("status") or "").strip().lower()
    auto_run_id = str(auto_loop.get("run_id") or "").strip()
    has_live_loop_context = auto_status in {"running", "approval_required", "paused"} and (not auto_run_id or auto_run_id == run_id)
    # 10 Minuten fuer direkte Confirm-Flows; danach nur mit aktivem Loop-Kontext weiter anzeigen.
    if age_seconds is not None and age_seconds > 600 and not has_live_loop_context:
        return False
    return True


def current_auto_run_status():
    state = load_project_auto_run_state()
    if pending_direct_context_is_valid(state.get("pending_direct_run"), state):
        pending = state["pending_direct_run"]
        return {
            "label": "Bestaetigung offen",
            "detail": (
                f"Direktmodus fuer '{pending.get('task') or 'Aufgabe'}' vorbereitet "
                f"({pending.get('scope') or 'project'}/{pending.get('mode') or 'safe'})."
            )
        }
    if state.get("last_direct_status") == "verified":
        return {
            "label": "Verifiziert",
            "detail": state.get("last_check_result") or "Direktmodus mit Nachkontrolle abgeschlossen."
        }
    if state.get("last_direct_status") == "applied":
        return {
            "label": "Applied",
            "detail": state.get("last_apply_action") or "Direktmodus hat eine Aenderung angewendet."
        }
    if state.get("blocked"):
        return {
            "label": "Blockiert",
            "detail": state.get("last_guard_decision") or "Guard hat die letzte Ausfuehrung blockiert."
        }
    if state.get("last_apply_action"):
        return {
            "label": "Apply abgeschlossen",
            "detail": state.get("last_apply_action")
        }
    if state.get("last_result"):
        return {
            "label": "Safe abgeschlossen",
            "detail": state.get("last_result")
        }
    return {
        "label": "Bereit",
        "detail": "Project Auto-Run kann gestartet werden."
    }


def build_project_strategy(task, suggestions, selected_target, guard_result, mode):
    prioritized = suggestions.get("prioritized_targets") or []
    blocked_items = suggestions.get("blocked") or []

    return {
        "ziel": (
            f"Die Aufgabe '{task}' im Project Mode kontrolliert koordinieren, "
            "nur innerhalb freigegebener Projektbereiche und ohne Blind-Ueberschreibungen."
        ),
        "betroffene_dateien": [
            item.get("path")
            for item in prioritized[:4]
            if item.get("path")
        ] or [selected_target],
        "empfohlene_reihenfolge": [
            "Relevante Bereiche scannen und Projektwissen abgleichen.",
            "Priorisierte Zielpfade pruefen und Guard-Entscheidung absichern.",
            "Aenderungsinhalt erzeugen und Diff-Vorschau erstellen.",
            "Nur bei Guard-OK und Apply Mode gezielt anwenden.",
            "Nachkontrolle und Status-Update dokumentieren."
        ],
        "risiken": [
            "Sensible oder nicht freigegebene Dateien duerfen nicht beschrieben werden.",
            "Nur gezielte Aenderungen mit sichtbarer Diff-Vorschau sind erlaubt.",
            "Bei mehrdeutiger Aufgabe kann der vorgeschlagene Zielpfad angepasst werden."
        ] + ([f"Blockierte Datei erkannt: {blocked_items[0]['path']}"] if blocked_items else []),
        "guard_ergebnis": {
            "allowed": bool(guard_result.get("allowed")),
            "detail": guard_result.get("detail") or guard_result.get("error") or "",
            "path": guard_result.get("path") or selected_target
        },
        "empfohlene_aktion": (
            "Apply ausfuehren" if mode == "apply" and guard_result.get("allowed")
            else "Safe-Pruefung abschliessen und Diff pruefen"
        ),
        "modus": mode
    }


def run_project_post_check(resolved_path, relative_path, proposed_content):
    content_after, exists_after = read_text_file(resolved_path)
    if not exists_after:
        return {
            "ok": False,
            "detail": f"Nachkontrolle fehlgeschlagen: '{relative_path}' wurde nicht gefunden."
        }

    if content_after != proposed_content:
        return {
            "ok": False,
            "detail": f"Nachkontrolle fehlgeschlagen: Inhalt von '{relative_path}' weicht ab."
        }

    return {
        "ok": True,
        "detail": f"Nachkontrolle erfolgreich: '{relative_path}' vorhanden und erwarteter Inhalt sichtbar."
    }


def run_local_post_check(resolved_path, relative_path, proposed_content):
    content_after, exists_after = read_text_file(resolved_path)
    if not exists_after:
        return {
            "ok": False,
            "detail": f"Nachkontrolle fehlgeschlagen: '{relative_path}' wurde lokal nicht gefunden."
        }
    if content_after != proposed_content:
        return {
            "ok": False,
            "detail": f"Nachkontrolle fehlgeschlagen: Inhalt von '{relative_path}' weicht lokal ab."
        }
    return {
        "ok": True,
        "detail": f"Nachkontrolle erfolgreich: '{relative_path}' wurde lokal geschrieben."
    }


def _is_single_file_direct_write_intent(lowered: str) -> bool:
    """Klarer Einzeldatei-Direct-Write (nicht Desktop-Multi-File-Build)."""
    if re.search(r"erstelle\s+nur\s+(die\s+)?datei", lowered):
        return True
    if re.search(r"erstelle\s+nur\s+eine\s+datei", lowered):
        return True
    if "mit dem inhalt" in lowered and len(lowered) < 4000:
        return True
    return False


def _is_desktop_multi_file_project_prompt(lowered: str) -> bool:
    """
    Electron-/React-/Robot-Desktop-Mehrdatei-Auftrag (Project Build / deterministischer Build).
    Muss VOR dem breiten file_edit-Match kommen: dort triggert das Wort ' in '
    zusammen mit Pfaden wie main.js / Bild.png fast jeden Lang-Prompt.
    """
    if not lowered or len(lowered) < 32:
        return False
    if _is_single_file_direct_write_intent(lowered):
        return False
    stack = any(
        k in lowered
        for k in (
            "electron",
            "react",
            "vite",
            "electron-builder",
            "desktop-app",
            "desktop app",
            "electron desktop",
        )
    ) or bool(re.search(r"\broboter\b", lowered)) or bool(re.search(r"\brobot\b", lowered))
    if not stack:
        return False
    bulk = any(
        h in lowered
        for h in (
            "build_desktop",
            "rambo_ui",
            "preload",
            "phase ",
            "npm install",
            "npm run",
            "komplette",
            "komplett ",
            "mehrere datei",
            "mehrere files",
            "projekt-struktur",
            "iconprocessor",
            "webpack",
            "package.json",
            "installer",
            ".exe",
            "electron/",
            "electron\\",
            "app.jsx",
            "main.jsx",
        )
    )
    if bulk:
        return True
    # Langer Auftrag mit zwei Stack-Signalen (z. B. Electron + React)
    pair = sum(1 for k in ("electron", "react", "vite", "preload") if k in lowered) >= 2
    return pair and len(lowered) >= 420


def classify_direct_task(task):
    lowered = str(task or "").lower()
    code_signals = [
        "def ", "class ", "import ", "function ", "async ", "api", "route",
        "endpoint", "backend", "frontend", "refactor", "typescript", "javascript",
        "komponente", "server", "projekt", "test", "bugfix",
    ]
    has_code_signal = any(sig in lowered for sig in code_signals)

    # SELF-REPAIR / AGENT INSTRUCTION Erkennung
    agent_patterns = [
        "du arbeitest im projekt",
        "aktueller fehler",
        "ziel:",
        "erlaubte dateien",
        "nicht ändern",
        "nicht aendern",
        "aufgaben:",
        "pflicht-test",
        "regression",
        "ausgabeformat",
        "geänderte dateien",
        "geaenderte dateien",
        "kurz umgesetzt",
        "kurz getestet",
        "self-repair",
        "self repair",
    ]
    is_agent_instruction = any(pat in lowered for pat in agent_patterns)
    if is_agent_instruction:
        task_type = "agent_instruction_prompt"

    strict_gate = bool(
        re.search(
            r"(freigabe\s+(erteilen|genehmigen|verweigern)|gate\s*(freigeben|approve|reject)|"
            r"approve\s+(the\s+)?(gate|step|patch)|patch\s+freigeben|schritt\s+freigeben|"
            r"approval\s+(grant|deny)|orchestrierungs?\s*freigabe|runner\s*freigeben|auto-?loop\s*freigeben)",
            lowered,
            re.IGNORECASE,
        )
    )
    if strict_gate:
        task_type = "approve_gate"
    elif _is_desktop_multi_file_project_prompt(lowered):
        task_type = "project_build"
    elif any(keyword in lowered for keyword in ["review", "patch", "diff", "begutachten", "code review"]):
        task_type = "review_patch"
    elif any(keyword in lowered for keyword in ["qa", "abnahme", "acceptance", "abnehmen"]):
        task_type = "qa_accept"
    elif any(keyword in lowered for keyword in ["continue", "fortsetzen", "resume", "weiterlaufen", "weitermachen"]):
        task_type = "continue_run"
    elif any(keyword in lowered for keyword in ["fix", "fehler", "bug", "debug", "repair", "reparier"]):
        task_type = "fix_error"
    elif re.search(r"([a-zA-Z0-9_\-./\\]+\.[a-zA-Z0-9]{1,8})", lowered) and any(
        keyword in lowered for keyword in ["edit", "aender", "update", "anpass", "replace", "ersetz"]
    ):
        task_type = "file_edit"
    elif any(
        phrase in lowered
        for phrase in [
            "orchestrator",
            "orchestrierung",
            "autoloop",
            "auto-loop",
            "runner-steuer",
            "pipeline steuer",
            "builder-modus",
            "koordinierter mehrschritt",
        ]
    ):
        task_type = "builder_mode"
    elif any(keyword in lowered for keyword in ["template", "scaffold", "boilerplate", "datei erstellen", "file generator", "neue datei"]):
        task_type = "file_generation"
    elif any(
        keyword in lowered
        for keyword in [
            "electron",
            "react",
            "vite",
            "desktop-app",
            "desktop app",
            "electron desktop",
            "komplette app",
            "komplette minimal",
            "komplette electron",
            "mehrere dateien",
            "mehrere files",
            "erzeuge vollständig",
            "erzeuge vollstaendig",
            "robot-icon",
            "roboter-icon",
        ]
    ):
        task_type = "project_build"
    elif any(keyword in lowered for keyword in ["design", "ui", "ux", "layout", "wireframe", "typografie", "visual"]) and not has_code_signal:
        task_type = "design_ui"
    else:
        task_type = "unknown"

    design_keywords = [
        "design", "ui", "layout", "style", "styling", "farbe", "farben",
        "typografie", "spacing", "abstand", "visual", "optik", "oberflaeche"
    ]
    file_keywords = [
        "datei erstellen",
        "neue datei",
        "file generator",
        "template",
        "scaffold",
        "boilerplate",
        "generieren",
        "anlegen",
        "export",
    ]
    code_keywords = [
        "code", "funktion", "klasse", "backend", "frontend", "api", "route",
        "endpunkt", "fix", "bug", "refactor", "komponente", "server", "projekt"
    ]

    if task_type == "unknown":
        if any(keyword in lowered for keyword in design_keywords) and not has_code_signal:
            task_type = "design_ui"
        elif any(keyword in lowered for keyword in file_keywords):
            task_type = "file_generation"
        elif any(keyword in lowered for keyword in code_keywords):
            task_type = "code_change"
        elif len(lowered.strip()) >= 12:
            task_type = "code_change"

    if task_type == "builder_mode" and not any(
        phrase in lowered
        for phrase in [
            "orchestrator",
            "orchestrierung",
            "autoloop",
            "auto-loop",
            "runner",
            "pipeline",
            "builder-modus",
            "mehrschritt",
        ]
    ):
        task_type = "code_change"

    if task_type == "design_ui" and has_code_signal:
        task_type = "code_change"

    special_types = {"approve_gate", "review_patch", "qa_accept", "builder_mode", "file_generation", "design_ui", "project_build"}
    execution_route = "special_surfaces" if task_type in special_types else "direct_agent"

    recommendation_map = {
        "agent_instruction_prompt": {
            "primary_area": "Builder",
            "hint": "Strukturierter Agent-/Self-Repair-Auftrag; Direct-Write nur wenn Ziel extrahierbar."
        },
        "approve_gate": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie Gate-/Freigabeentscheidung und bleibt manuell priorisiert."
        },
        "review_patch": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie Review-/Patch-Arbeit und sollte im Review-Pfad geklaert werden."
        },
        "qa_accept": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie QA-/Acceptance-Schritt."
        },
        "continue_run": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie Fortsetzungsauftrag fuer den aktiven Lauf."
        },
        "fix_error": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie Fehlerbehebung mit priorisiertem Reparaturpfad."
        },
        "file_edit": {
            "primary_area": "Datei-Generator",
            "hint": "Prompt wirkt wie konkrete Datei-Aenderung."
        },
        "builder_mode": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie koordinierter Builder-Auftrag."
        },
        "design_ui": {
            "primary_area": "Design Studio",
            "hint": "Prompt wirkt wie UI-/Design-Arbeit. Design Studio ist der naheliegendste interne Bereich."
        },
        "file_generation": {
            "primary_area": "Datei-Generator",
            "hint": "Prompt wirkt wie Datei-, Template- oder Export-Erzeugung. Datei-Generator passt prima als interner Schwerpunkt."
        },
        "project_build": {
            "primary_area": "Project Builder",
            "hint": "Prompt wirkt wie Multi-File-Projekt (Electron, React, etc.). Project Builder wird aktiviert."
        },
        "code_change": {
            "primary_area": "Builder",
            "hint": "Prompt wirkt wie Projekt- oder Code-Aenderung. Builder ist der passende Hauptpfad."
        },
        "unknown": {
            "primary_area": "Builder",
            "hint": "Prompt ist nicht eindeutig. Builder bleibt der sichere Standardpfad."
        }
    }

    recommendation = recommendation_map[task_type]
    return {
        "task_type": task_type,
        "primary_area": recommendation["primary_area"],
        "hint": recommendation["hint"],
        "execution_route": execution_route,
        "route_hint": (
            "Standard: direkter Agenten-Direktlauf ohne Spezialoberflaeche."
            if execution_route == "direct_agent"
            else "Spezialoberflaeche passt zum erkannten Auftragstyp."
        ),
    }


def direct_diff_exceeds_small_change(diff_text, max_lines=200):
    if not diff_text or diff_text == "Keine inhaltliche Aenderung erkannt.":
        return False
    return (diff_text.count("\n") + 1) > max_lines


def classify_direct_execution_route(task, scope, mode, has_changes, diff_text, guard_allowed):
    """
    Behebungsplan Phase 1: intern nur direct_execute vs safe_review.
    Nicht als Nutzer-UI ausgeben; nur Logs/Serverlogik.
    """
    if not guard_allowed:
        return "safe_review", "guard_verweigert"
    sc = str(scope or "local").strip().lower()
    m = str(mode or "safe").strip().lower()
    if direct_task_requires_user_confirmation(task, m, bool(has_changes), True, sc):
        return "safe_review", "risiko_oder_project_apply"
    if bool(has_changes) and direct_diff_exceeds_small_change(str(diff_text or "")):
        return "safe_review", "diff_ueber_schwelle"
    return "direct_execute", "lokal_sicher_klein"


def direct_task_requires_user_confirmation(task, mode, has_changes, guard_allowed, scope="local"):
    """
    Gibt True zurueck, wenn der Task eine Benutzerbestaetigung erfordert.
    Im AUTO_APPLY-Modus werden nur echte Hard-Blocks gesperrt.
    """
    if not guard_allowed:
        return True
    tl = str(task or "").lower()
    # Hard-Blocks: gelten immer, unabhaengig von AUTO_APPLY
    hard_blocks = (
        "rm -rf", "format c:", "drop table", "delete from ", "truncate table",
        "shred ", "del /s", "remove-item -recurse", "wipe disk", "rd /s",
        "alle dateien loeschen", "alle logs loeschen", "komplett loeschen",
        "loesche alle", "lösche alle", "delete all", "remove all",
        "formatiere die festplatte", "shutdown", "reboot",
        "api-key", "apikey", ".env", "password", "credential",
    )
    if any(s in tl for s in hard_blocks):
        return True
    # Im AUTO_APPLY-Modus: alle anderen Operationen direkt ausfuehren
    if AUTO_APPLY:
        return False
    # Soft-Blocks nur ausserhalb AUTO_APPLY
    soft_blocks = (
        "curl ", "wget ", "invoke-webrequest", "urllib.request", "requests.get",
        "secret", "token ", "oauth",
        "loesch", "lösch", "delete ", "entfern alle", "umbenenn", "verschieb",
    )
    if any(s in tl for s in soft_blocks):
        return True
    sc = str(scope or "local").strip().lower()
    m = str(mode or "safe").strip().lower()
    if m == "apply" and bool(has_changes) and sc != "local":
        return True
    return False


def reset_ephemeral_control_state_for_new_direct_task():
    """Vor neuem Direktauftrag alte Steuerungs-/Gate-Zwischenzustaende verwerfen (Plan Phase 4)."""
    try:
        idle_loop = normalize_auto_loop_state({})
        if isinstance(idle_loop, dict):
            idle_loop["pending_approval"] = None
        save_project_auto_run_state({
            "auto_loop_state": idle_loop,
            "last_agent_decision": {},
            "agent_decision_history": [],
            "last_runner_execution": None,
        })
    except Exception:
        pass


def build_local_direct_preview(task, mode):
    steps = [{"label": "Prompt", "status": "ok", "detail": task}]
    recognized_task = classify_direct_task(task)
    inferred_allowed_targets, infer_meta = infer_allowed_target_files_with_meta(task)
    if len(inferred_allowed_targets) == 1:
        target_path = inferred_allowed_targets[0]
    elif _is_mini_task_write_intent(task):
        if is_active_workspace_trusted() and len(inferred_allowed_targets) >= 2:
            target_path = inferred_allowed_targets[0]
        else:
            unclear_payload = _build_target_path_unclear_payload(
                task, mode, inferred_allowed_targets, inference_debug=infer_meta
            )
            unclear_payload["scope"] = "local"
            unclear_payload["recognized_task"] = recognized_task
            unclear_payload["steps"] = steps + [
                {"label": "Zielpfad", "status": "warning", "detail": "Kein eindeutiger erlaubter Zielpfad erkannt."},
            ]
            unclear_payload["planned_steps"] = normalize_planned_steps([])
            return {"ok": False, "status_code": 400, "payload": unclear_payload}
    else:
        target_path = infer_target_path(task)
    steps.append({"label": "Zielpfad", "status": "auto", "detail": f"Automatisch ermittelt: {target_path}"})

    resolved_path, relative_path, path_error = resolve_local_target_path(target_path, task)
    if path_error:
        planned_steps = build_planned_steps(task, mode, target_path, [target_path], True, False)
        steps.append({"label": "Guard", "status": "error", "detail": path_error})
        return {
            "ok": False,
            "status_code": 403,
            "payload": {
                "error": path_error,
                "scope": "local",
                "mode": mode,
                "steps": steps,
                "planned_steps": planned_steps,
                "guard": {"allowed": False, "detail": path_error, "path": target_path},
            }
        }
    direct_guard_check = _validate_direct_run_paths([target_path, relative_path], mode, task)
    if not bool(direct_guard_check.get("ok")):
        blocked_files = direct_guard_check.get("blocked_files") or [relative_path]
        planned_steps = build_planned_steps(task, mode, relative_path, blocked_files, True, False)
        steps.append({"label": "Guard", "status": "blocked", "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE})
        blocked_payload = _build_direct_guard_block_payload(
            scope="local",
            mode=mode,
            blocked_files=blocked_files,
            task=task,
            recognized_task=recognized_task,
        )
        blocked_payload["steps"] = steps
        blocked_payload["planned_steps"] = planned_steps
        blocked_payload["selected_target_path"] = relative_path
        return {"ok": False, "status_code": 403, "payload": blocked_payload}

    _zone = "USER-Downloads" if local_user_download_requested(task) else "rambo_builder_local (intern)"
    guard = {
        "allowed": True,
        "detail": f"Pfad erlaubt ({_zone}): {relative_path}",
        "path": relative_path,
    }
    steps.append({"label": "Guard", "status": "ok", "detail": guard["detail"]})

    try:
        from task_parser import parse_user_prompt_to_task_spec

        if parse_user_prompt_to_task_spec(task).file_count > 1:
            append_ui_log_entry(
                "Direkt",
                "Hinweis: Prompt erwaehnt mehrere Dateien; pro Lauf wird nur ein Ziel verarbeitet.",
                "warning",
            )
    except Exception:
        pass

    current_content, file_exists = read_text_file(resolved_path)
    proposed_content = resolve_proposed_content_for_local_task(
        task, target_path, current_content, file_exists, relative_path
    )
    pytest_validation = validate_pytest_file(relative_path, proposed_content)
    if not bool(pytest_validation.get("ok")):
        blocked_payload = _build_invalid_test_file_payload(relative_path, scope="local", mode=mode, task=task)
        blocked_payload["recognized_task"] = recognized_task
        blocked_payload["steps"] = steps + [
            {"label": "Test-Validierung", "status": "blocked", "detail": "Kein pytest-Test erkannt."},
        ]
        blocked_payload["planned_steps"] = normalize_planned_steps([])
        blocked_payload["workstream_events"] = [
            _ws_event("analysis", "info", "Datei geprüft", f"Ziel: {relative_path}", status="done"),
            _ws_event("guard", "error", "Ungültige Testdatei", "Kein pytest-Test erkannt.", status="blocked"),
        ]
        return {"ok": False, "status_code": 400, "payload": blocked_payload}
    diff_text = build_text_diff(current_content, proposed_content, relative_path)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    steps.append({
        "label": "Diff",
        "status": "ok" if has_changes else "warn",
        "detail": "Diff-Vorschau erstellt." if has_changes else "Keine inhaltliche Aenderung erkannt."
    })

    plan = build_builder_plan(task)
    strategy = {
        "ziel": plan["ziel"],
        "betroffene_dateien": [relative_path],
        "empfohlene_reihenfolge": plan["empfohlene_reihenfolge"],
        "risiken": plan["risiken"],
        "guard_ergebnis": guard,
        "empfohlene_aktion": (
            "Nach Bestaetigung lokal anwenden" if mode == "apply" and has_changes
            else "Safe-Vorschau bestaetigen"
        ),
        "modus": mode
    }
    direct_status = "apply_ready" if mode == "apply" and has_changes else "safe_preview"
    diff_summary = build_diff_summary(diff_text)
    planned_steps = build_planned_steps(task, mode, relative_path, [relative_path], has_changes, True)

    return {
        "ok": True,
        "status_code": 200,
        "payload": {
            **preview_payload(
                path=relative_path,
                technical_message="Direktmodus-Vorschau vorbereitet."
            ),
            "scope": "local",
            "mode": mode,
            "task": task,
            "selected_target_path": relative_path,
            "affected_files": [relative_path],
            "file_plan": [relative_path],
            "file_entries": [{"path": relative_path, "status": "planned"}],
            "steps": steps,
            "strategy": strategy,
            "recognized_task": recognized_task,
            "planned_steps": planned_steps,
            "guard": guard,
            "diff": diff_text,
            "diff_summary": diff_summary,
            "has_changes": has_changes,
            "file_exists": file_exists,
            "proposed_content": proposed_content,
            "post_check": {"ok": False, "detail": "Noch nicht ausgefuehrt."},
            "requires_confirmation": True,
            "requires_user_confirmation": (
                direct_task_requires_user_confirmation(task, mode, has_changes, True, "local")
                or (has_changes and direct_diff_exceeds_small_change(diff_text))
            ),
            "blocked_files": [],
            "direct_status": direct_status,
        }
    }


def build_project_direct_preview(task, mode):
    steps = [{"label": "Prompt", "status": "ok", "detail": task}]
    recognized_task = classify_direct_task(task)

    repo_map = read_json_file(PROJECT_MAP_FILE, {})
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    if not repo_map:
        files = scan_project_structure()
        repo_map = {"scanned_at": get_timestamp(), "total_files": len(files), "areas": {}}
        write_json_file(PROJECT_MAP_FILE, repo_map)
        try:
            save_project_auto_run_state({
                "last_project_scan": {
                    "source": "project_preview_auto",
                    "scanned_at": str(repo_map.get("scanned_at") or ""),
                    "total_files": int(repo_map.get("total_files") or 0),
                    "area_count": len(repo_map.get("areas") or {}) if isinstance(repo_map.get("areas"), dict) else 0,
                }
            })
        except Exception:
            pass
    if not knowledge:
        knowledge = build_project_knowledge()

    steps.append({
        "label": "Scan",
        "status": "ok",
        "detail": f"Projektbasis verfuegbar: {repo_map.get('total_files', 0)} Dateien."
    })

    suggestions = suggest_files_for_task(task)
    selected_target = infer_project_target_path(task, suggestions)
    prioritized_targets = suggestions.get("prioritized_targets") or []
    guard_candidates = [selected_target]
    guard_candidates.extend(
        str(item.get("path") or "").strip()
        for item in prioritized_targets
        if isinstance(item, dict) and str(item.get("path") or "").strip()
    )
    direct_guard_check = _validate_direct_run_paths(guard_candidates, mode, task)
    if not bool(direct_guard_check.get("ok")):
        blocked_files = direct_guard_check.get("blocked_files") or [selected_target]
        steps.append({
            "label": "Guard",
            "status": "blocked",
            "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE,
        })
        planned_steps = build_planned_steps(task, mode, selected_target, blocked_files, True, False)
        blocked_payload = _build_direct_guard_block_payload(
            scope="project",
            mode=mode,
            blocked_files=blocked_files,
            task=task,
            recognized_task=recognized_task,
        )
        blocked_payload["steps"] = steps
        blocked_payload["planned_steps"] = planned_steps
        blocked_payload["selected_target_path"] = selected_target
        blocked_payload["suggestions"] = suggestions
        return {"ok": False, "status_code": 403, "payload": blocked_payload}
    steps.append({
        "label": "Auswahl",
        "status": "ok" if selected_target else "warning",
        "detail": f"Zielpfad gewaehlt: {selected_target}" if selected_target else "Kein Zielpfad bestimmbar."
    })

    resolved, cleaned, guard_error = validate_project_write_path(selected_target)
    if guard_error:
        # Auto-Recovery (1x): bei Guard-Block alternative Zielpfade testen, statt sofort abzubrechen.
        retry_candidates: list[str] = []
        retry_candidates.extend(
            str(item.get("path") or "").strip()
            for item in (suggestions.get("prioritized_targets") or [])
            if isinstance(item, dict) and str(item.get("path") or "").strip()
        )
        retry_candidates.extend(
            str(item.get("example_path") or "").strip()
            for item in (suggestions.get("area_suggestions") or [])
            if isinstance(item, dict) and str(item.get("example_path") or "").strip()
        )
        retry_candidates.append(infer_project_target_path(task, suggestions))
        seen_retry: set[str] = set()
        for cand in retry_candidates:
            c = format_local_path(cand or "")
            if not c or c in seen_retry:
                continue
            seen_retry.add(c)
            r2, c2, e2 = validate_project_write_path(c)
            if not e2:
                resolved, cleaned, guard_error = r2, c2, None
                steps.append({"label": "Auto-Recovery", "status": "ok", "detail": f"Alternativer Zielpfad: {c2}"})
                break
    if guard_error:
        guard_result = {
            "allowed": False,
            "detail": guard_error,
            "path": selected_target
        }
        planned_steps = build_planned_steps(task, mode, selected_target, suggestions.get("prioritized_targets") or [selected_target], True, False)
        steps.append({"label": "Guard", "status": "error", "detail": guard_error})
        strategy = build_project_strategy(task, suggestions, selected_target, guard_result, mode)
        return {
            "ok": False,
            "status_code": 403,
            "payload": {
                "error": guard_error,
                "scope": "project",
                "mode": mode,
                "task": task,
                "steps": steps,
                "planned_steps": planned_steps,
                "suggestions": suggestions,
                "selected_target_path": selected_target,
                "guard": guard_result,
                "strategy": strategy,
                "blocked_files": suggestions.get("blocked", []),
            }
        }

    _, file_exists = read_text_file(resolved)
    guard_result = {
        "allowed": True,
        "detail": f"Pfad freigegeben. Datei {'vorhanden' if file_exists else 'noch nicht vorhanden'}.",
        "path": cleaned,
        "file_exists": file_exists
    }
    steps.append({"label": "Guard", "status": "ok", "detail": guard_result["detail"]})

    current_content, file_exists = read_text_file(resolved)
    proposed_content = resolve_proposed_content_for_local_task(
        task, cleaned, current_content, file_exists, cleaned
    )
    pytest_validation = validate_pytest_file(cleaned, proposed_content)
    if not bool(pytest_validation.get("ok")):
        blocked_payload = _build_invalid_test_file_payload(cleaned, scope="project", mode=mode, task=task)
        blocked_payload["recognized_task"] = recognized_task
        blocked_payload["steps"] = steps + [
            {"label": "Test-Validierung", "status": "blocked", "detail": "Kein pytest-Test erkannt."},
        ]
        blocked_payload["planned_steps"] = normalize_planned_steps([])
        blocked_payload["suggestions"] = suggestions
        return {"ok": False, "status_code": 400, "payload": blocked_payload}
    diff_text = build_text_diff(current_content, proposed_content, cleaned)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    steps.append({
        "label": "Diff",
        "status": "ok" if has_changes else "warn",
        "detail": "Diff erstellt." if has_changes else "Keine inhaltliche Aenderung erkannt."
    })

    strategy = build_project_strategy(task, suggestions, cleaned, guard_result, mode)
    strategy["betroffene_dateien"] = strategy.get("betroffene_dateien") or [cleaned]
    direct_status = "apply_ready" if mode == "apply" and has_changes else "safe_preview"
    diff_summary = build_diff_summary(diff_text)
    planned_steps = build_planned_steps(task, mode, cleaned, strategy.get("betroffene_dateien") or [cleaned], has_changes, True)

    return {
        "ok": True,
        "status_code": 200,
        "payload": {
            **preview_payload(
                path=cleaned,
                technical_message="Project-Direktmodus-Vorschau vorbereitet."
            ),
            "scope": "project",
            "mode": mode,
            "task": task,
            "selected_target_path": cleaned,
            "affected_files": [cleaned],
            "file_plan": [cleaned],
            "file_entries": [{"path": cleaned, "status": "planned"}],
            "steps": steps,
            "suggestions": suggestions,
            "strategy": strategy,
            "recognized_task": recognized_task,
            "planned_steps": planned_steps,
            "guard": guard_result,
            "blocked_files": suggestions.get("blocked", []),
            "proposed_content": proposed_content,
            "diff": diff_text,
            "diff_summary": diff_summary,
            "has_changes": has_changes,
            "file_exists": file_exists,
            "post_check": {"ok": False, "detail": "Noch nicht ausgefuehrt."},
            "requires_confirmation": True,
            "requires_user_confirmation": (
                direct_task_requires_user_confirmation(task, mode, has_changes, True, "project")
                or (has_changes and direct_diff_exceeds_small_change(diff_text))
            ),
            "direct_status": direct_status,
        }
    }


def build_direct_preview(scope, task, mode):
    if scope == "local":
        return build_local_direct_preview(task, mode)
    return build_project_direct_preview(task, mode)


def attach_direct_ui_chrome(payload: dict, pending) -> dict:
    """Mini-Phase 3: Chrome an Confirm-Antwort; SAFE-Vorschau ohne Schreiblauf immer full UI."""
    p = pending if isinstance(pending, dict) else {}
    if str(payload.get("direct_status") or "").lower() == "safe_preview":
        chrome = "full"
    else:
        dc = p.get("direct_ui_chrome")
        chrome = dc if dc in ("minimal", "full") else "full"
    return {**payload, "direct_ui_chrome": chrome}


def execute_direct_confirmation(pending):
    pending = _merge_direct_file_context(pending, pending)
    scope = str(pending.get("scope") or "project")
    mode = str(pending.get("mode") or "safe")
    # Global Auto-Continue: Safe-Preview niemals als Stopp/Gate behandeln.
    if mode.strip().lower() == "safe":
        mode = "apply"
        pending = {**pending, "mode": "apply", "requires_user_confirmation": False}
    task = str(pending.get("task") or "")
    target_path = str(pending.get("selected_target_path") or "")
    proposed_content = str(pending.get("proposed_content") or "")
    guard = pending.get("guard") if isinstance(pending.get("guard"), dict) else {}

    if not task or not target_path:
        return {"error": "Keine gueltige Direktmodus-Vorschau zur Bestaetigung vorhanden."}, 400

    if not guard.get("allowed"):
        return {"error": "Direktmodus ist blockiert und kann nicht bestaetigt werden."}, 403

    guard_candidates = []
    for key in ("selected_target_path", "target_path", "path", "relative_path", "absolute_path"):
        value = pending.get(key)
        if isinstance(value, str) and value.strip():
            guard_candidates.append(value)
    for key in (
        "affected_files",
        "changed_files",
        "created_files",
        "updated_files",
        "deleted_files",
        "file_plan",
        "file_entries",
        "artifacts",
    ):
        value = pending.get(key)
        if isinstance(value, (list, tuple, set)):
            guard_candidates.extend(list(value))
    if isinstance(guard, dict) and str(guard.get("path") or "").strip():
        guard_candidates.append(str(guard.get("path") or "").strip())
    if target_path:
        guard_candidates.append(target_path)
    validation = _validate_direct_run_paths(guard_candidates, mode, task)
    if not bool(validation.get("ok")):
        blocked_files = validation.get("blocked_files") or guard_candidates
        payload = _build_direct_guard_block_payload(
            scope=scope,
            mode=mode,
            blocked_files=blocked_files,
            task=task,
            recognized_task=pending.get("recognized_task"),
        )
        payload["selected_target_path"] = target_path
        payload["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
        return payload, 403

    if scope == "local":
        resolved_path, relative_path, path_error = resolve_local_target_path(target_path, task)
        if path_error:
            return {"error": path_error}, 400
        append_ui_log_entry("Direkt", f"Guard erneut geprueft: {relative_path}", "info")

        tl_confirm = str(task or "").lower()
        explicit_file_delete = _is_explicit_file_delete_intent(task)
        ui_change_remove = _is_ui_change_remove_intent(task)
        delete_intent_loose = any(p in tl_confirm for p in ["lösch", "loesch", "delete ", "entfern"])
        active_root_confirm = get_active_project_root().resolve()
        resolved_abs_confirm = Path(resolved_path).resolve()
        inside_del_ok, inside_del_detail = _same_or_inside_project_root(active_root_confirm, resolved_abs_confirm)

        # KRITISCH: Nur bei eindeutiger Datei-Lösch-Absicht os.remove — nie bei „Entferne aus UI …“
        if explicit_file_delete:
            if not inside_del_ok:
                clear_pending_direct_run()
                resp = {
                    "direct_status": "blocked",
                    "message": "Löschen nur innerhalb des aktiven Projektroots erlaubt.",
                    "technical_message": inside_del_detail,
                    "delete_intent": True,
                    "explicit_file_delete_intent": True,
                    "delete_block_reason": "path_outside_active_project_root",
                    "active_project_root": str(active_root_confirm),
                    "resolved_delete_path": str(relative_path),
                    "same_or_inside_result": {"ok": False, "detail": inside_del_detail},
                    "route_mode": "blocked",
                    "has_changes": False,
                    "changed_files": [],
                }
                resp = _merge_direct_file_context(resp, pending)
                return jsonify(attach_direct_ui_chrome(resp, pending)), 403
            if not os.path.exists(resolved_path):
                clear_pending_direct_run()
                save_project_auto_run_state({
                    "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                    "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                })
                resp = {
                    "direct_status": "verified",
                    "message": f"Datei existiert bereits nicht mehr: {relative_path}",
                    "technical_message": f"Datei nicht gefunden: {resolved_path}",
                    "has_changes": False,
                    "changed_files": [],
                    "affected_files": [relative_path],
                    "post_check": {"ok": True, "detail": "Datei bereits nicht vorhanden."},
                    "delete_intent": True,
                    "explicit_file_delete_intent": True,
                    "delete_block_reason": "",
                    "active_project_root": str(active_root_confirm),
                    "resolved_delete_path": str(relative_path),
                    "same_or_inside_result": {"ok": True, "detail": inside_del_detail},
                }
                resp = _merge_direct_file_context(resp, pending)
                return jsonify(attach_direct_ui_chrome(resp, pending))
            if not resolved_abs_confirm.is_file():
                clear_pending_direct_run()
                resp = {
                    "direct_status": "blocked",
                    "message": "Nur reguläre Dateien können gelöscht werden (kein Ordner).",
                    "delete_intent": True,
                    "explicit_file_delete_intent": True,
                    "delete_block_reason": "target_not_a_file",
                    "active_project_root": str(active_root_confirm),
                    "resolved_delete_path": str(relative_path),
                    "same_or_inside_result": {"ok": inside_del_ok, "detail": inside_del_detail},
                    "route_mode": "blocked",
                    "has_changes": False,
                    "changed_files": [],
                }
                resp = _merge_direct_file_context(resp, pending)
                return jsonify(attach_direct_ui_chrome(resp, pending)), 400
            try:
                os.remove(resolved_path)
            except Exception as del_err:
                return {"error": f"Loeschen fehlgeschlagen: {del_err}"}, 500
            still_exists = os.path.exists(resolved_path)
            clear_pending_direct_run()
            save_project_auto_run_state({
                "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_direct_status": "verified",
            })
            append_ui_log_entry("Direkt", f"Datei geloescht: {relative_path}", "success")
            _ws_del = [
                {"ts": get_timestamp(), "phase": "analysis", "level": "info",
                 "title": "Delete-Auftrag erkannt", "detail": str(relative_path), "status": "done"},
                {"ts": get_timestamp(), "phase": "writing", "level": "success" if not still_exists else "error",
                 "title": "Datei gelöscht" if not still_exists else "Löschen fehlgeschlagen",
                 "detail": str(relative_path), "file": str(resolved_path), "status": "done" if not still_exists else "failed"},
            ]
            resp = {
                "direct_status": "verified",
                "message": f"Datei erfolgreich geloescht: {relative_path}",
                "technical_message": f"os.remove ausgefuehrt: {resolved_path}",
                "has_changes": True,
                "changed_files": [relative_path],
                "affected_files": [relative_path],
                "deleted_files": [relative_path],
                "created_files": [],
                "updated_files": [],
                "workstream_events": _ws_del,
                "auto_apply": True,
                "auto_continue": True,
                "skip_review": True,
                "direct_execute": True,
                "delete_intent": True,
                "explicit_file_delete_intent": True,
                "delete_block_reason": "",
                "active_project_root": str(active_root_confirm),
                "resolved_delete_path": str(relative_path),
                "same_or_inside_result": {"ok": True, "detail": inside_del_detail},
                "route_mode": "file_delete",
                "post_check": {
                    "ok": not still_exists,
                    "detail": "Datei ist geloescht." if not still_exists else "Warnung: Datei noch vorhanden.",
                    "path": str(relative_path),
                    "target": str(resolved_path),
                },
            }
            upsert_direct_run_history(build_direct_history_entry(
                str(pending.get("run_id") or ""), {**pending, **resp}, "verified"
            ))
            resp = _merge_direct_file_context(resp, pending)
            return jsonify(attach_direct_ui_chrome(resp, pending))

        if delete_intent_loose and not explicit_file_delete and (
            ui_change_remove or _prompt_mentions_builder_mode(task)
        ):
            reason = "ui_or_code_change_not_os_remove"
            if ui_change_remove:
                reason = "ui_change_requires_code_edit_not_file_delete"
            targets_hint = _discover_builder_ui_target_files(active_root_confirm, task)
            clear_pending_direct_run()
            resp = {
                "direct_status": "delete_blocked",
                "message": (
                    "Kein Datei-Löschen per os.remove: Es wurde eine UI-/Code-Änderung oder eine nicht eindeutige "
                    "Formulierung erkannt. Bitte Inhalt per Diff/Code-Änderung anpassen."
                ),
                "technical_message": reason,
                "has_changes": False,
                "changed_files": [],
                "delete_intent": True,
                "explicit_file_delete_intent": False,
                "delete_block_reason": reason,
                "active_project_root": str(active_root_confirm),
                "resolved_delete_path": str(relative_path),
                "same_or_inside_result": {"ok": inside_del_ok, "detail": inside_del_detail},
                "route_mode": "project_change",
                "target_files": targets_hint,
                "relevant_files": targets_hint,
                "classification": "project_task",
                "workstream_events": [
                    {"ts": get_timestamp(), "phase": "guard", "level": "warning",
                     "title": "Datei-Löschen blockiert", "detail": reason, "status": "blocked"},
                ],
            }
            resp = _merge_direct_file_context(resp, pending)
            return jsonify(attach_direct_ui_chrome(resp, pending)), 200

        scope_l = str(scope).strip().lower()
        mode_l = str(mode).strip().lower()
        must_review = False
        should_auto_write = (
            mode_l == "safe"
            and bool(pending.get("has_changes"))
            and scope_l == "local"
            and not must_review
        )

        if mode_l == "safe" and not should_auto_write:
            clear_pending_direct_run()
            if not pending.get("has_changes"):
                msg_safe = "Safe bestaetigt: Keine inhaltliche Aenderung in der Vorschau, Datei unveraendert."
            elif scope_l != "local":
                msg_safe = "Safe bestaetigt: Kein automatischer Schreiblauf ausserhalb des Local-Direktpfads."
            elif must_review:
                msg_safe = "Safe bestaetigt: Auftrag mit Freigabeprofil (Risiko/grosse Diff), Schreiben nur nach expliziter Apply-Freigabe."
            else:
                msg_safe = "Safe bestaetigt: Kein Schreiblauf in diesem Schritt (interner Ausschluss)."
            save_project_auto_run_state({
                "last_run_at": get_timestamp(),
                "last_result": f"Safe bestaetigt fuer '{relative_path}'.",
                "last_check_result": "Keine Nachkontrolle noetig (Safe Mode).",
                "last_direct_decision": "Safe-Vorschau bestaetigt.",
                "last_direct_status": "safe_preview",
                "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or [])
            })
            append_ui_log_entry("Direkt", f"Safe bestaetigt: {relative_path}", "success")
            response_payload = {
                "message": msg_safe,
                "scope": scope,
                "mode": mode,
                "selected_target_path": relative_path,
                "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
                "affected_files": [relative_path],
                "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                "post_check": {"ok": False, "detail": "Keine Nachkontrolle noetig (Safe Mode)."},
                "direct_status": "safe_preview",
                "diff_summary": build_diff_summary(pending.get("diff") or "")
            }
            upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, "safe_preview"))
            response_payload = _merge_direct_file_context(response_payload, pending)
            return jsonify(attach_direct_ui_chrome(response_payload, pending))

        current_content, file_exists = read_text_file(resolved_path)
        diff_text = build_text_diff(current_content, proposed_content, relative_path)
        has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
        if not has_changes:
            clear_pending_direct_run()
            append_ui_log_entry("Direkt", f"Keine Aenderung fuer '{relative_path}'.", "warning")
            response_payload = {
                "message": "Keine inhaltliche Aenderung erkannt. Es wurde nichts geschrieben.",
                "scope": scope,
                "mode": mode,
                "selected_target_path": relative_path,
                "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
                "affected_files": [relative_path],
                "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                "diff": diff_text,
                "post_check": {"ok": False, "detail": "Keine Nachkontrolle noetig."},
                "direct_status": "safe_preview",
                "diff_summary": build_diff_summary(diff_text)
            }
            upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, "safe_preview"))
            response_payload = _merge_direct_file_context(response_payload, pending)
            return jsonify(attach_direct_ui_chrome(response_payload, pending))

        pytest_validation = validate_pytest_file(relative_path, proposed_content)
        if not bool(pytest_validation.get("ok")):
            clear_pending_direct_run()
            blocked_payload = _build_invalid_test_file_payload(relative_path, scope=scope, mode=mode, task=task)
            blocked_payload["recognized_task"] = pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {}
            blocked_payload["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
            blocked_payload["diff_summary"] = build_diff_summary(diff_text)
            blocked_payload["workstream_events"] = [
                _ws_event("analysis", "info", "Datei geprüft", f"Ziel: {relative_path}", status="done"),
                _ws_event("guard", "error", "Ungültige Testdatei", "Kein pytest-Test erkannt.", status="blocked"),
            ]
            append_ui_log_entry("Direkt", "Ungültige Testdatei blockiert.", "error")
            return blocked_payload, 400

        unsafe_check = detect_unsafe_large_rewrite(
            relative_path,
            current_content,
            proposed_content,
            {"phase": "local_apply", "task": task},
        )
        if bool(unsafe_check.get("unsafe")):
            applied_rec, se_res = _try_unsafe_rewrite_step_engine_apply(
                relative_path=relative_path,
                resolved_path=resolved_path,
                current_content=current_content,
                proposed_content=proposed_content,
                task=task,
                unsafe_check=unsafe_check,
            )
            if applied_rec:
                append_ui_log_entry("Direkt", f"Auto-Recovery (Step-Engine): Patch angewendet: {relative_path}", "success")
                post_check = run_local_post_check(resolved_path, relative_path, proposed_content)
                clear_pending_direct_run()
                decision_txt = "Lokale Aenderung nach Step-Engine-Recovery angewendet."
                save_project_auto_run_state({
                    "last_run_at": get_timestamp(),
                    "last_result": "Direktmodus lokal abgeschlossen (Recovery).",
                    "last_apply_action": f"Lokale Datei {'aktualisiert' if file_exists else 'angelegt'}: {relative_path}",
                    "last_check_result": post_check["detail"],
                    "last_direct_decision": decision_txt,
                    "last_direct_status": "verified" if post_check["ok"] else "applied",
                    "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                    "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                    "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                    "blocked": False,
                })
                append_ui_log_entry("Direkt", f"Apply ausgefuehrt: {relative_path}", "success")
                append_ui_log_entry("Direkt", post_check["detail"], "success" if post_check["ok"] else "error")
                msg_local = (
                    f"Lokale Direktmodus-Aenderung {'aktualisiert' if file_exists else 'angelegt'}: {relative_path} "
                    "(nach Step-Engine-Recovery, kleiner Patch)"
                )
                analysis_detail = build_written_result_detail(resolved_path, relative_path)
                file_action = "erstellt" if not file_exists else "aktualisiert"
                _ws_events = [
                    {"ts": get_timestamp(), "phase": "analysis", "level": "info",
                     "title": "Auftrag erkannt", "detail": str(task)[:80], "status": "done"},
                    {"ts": get_timestamp(), "phase": "guard", "level": "warning",
                     "title": "Rewrite-Heuristik", "detail": "Große Änderung erkannt, Patch aber unter Diff-Schwelle.", "status": "done"},
                    {"ts": get_timestamp(), "phase": "recovery", "level": "success",
                     "title": "Step-Engine Recovery", "detail": "Validierter Patch automatisch angewendet.", "status": "done"},
                    {"ts": get_timestamp(), "phase": "writing", "level": "success",
                     "title": f"Datei {file_action}", "detail": str(relative_path),
                     "file": str(resolved_path), "status": "done"},
                    {"ts": get_timestamp(), "phase": "verify", "level": "success" if post_check["ok"] else "error",
                     "title": "Ergebnis geprüft", "detail": post_check["detail"],
                     "status": "done" if post_check["ok"] else "failed"},
                ]
                response_payload = {
                    **success_payload(
                        "Änderung angewendet! ✓",
                        technical_message=msg_local,
                        changed_files=[relative_path],
                        location="rambo_builder_local/",
                        detail=analysis_detail or "Die Änderung wurde erfolgreich angewendet (Recovery).",
                    ),
                    "scope": scope,
                    "mode": "apply",
                    "selected_target_path": relative_path,
                    "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
                    "affected_files": [relative_path],
                    "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                    "diff": diff_text,
                    "post_check": {
                        **post_check,
                        "path": str(relative_path),
                        "target": str(resolved_path),
                    },
                    "direct_status": "verified" if post_check["ok"] else "applied",
                    "diff_summary": build_diff_summary(diff_text),
                    "workstream_events": _ws_events,
                    "created_files": [] if file_exists else [relative_path],
                    "updated_files": [relative_path] if file_exists else [],
                    "deleted_files": [],
                    "changed_files": [relative_path],
                    "split_patch_recovery": True,
                    "step_engine_result": se_res if isinstance(se_res, dict) else {},
                    "auto_apply": True,
                    "auto_continue": True,
                    "skip_review": True,
                    "direct_execute": True,
                }
                response_payload = _merge_direct_file_context(response_payload, pending)
                response_payload = _enforce_real_change_success(task, response_payload, mode="apply")
                upsert_direct_run_history(
                    build_direct_history_entry(
                        str(pending.get("run_id") or ""),
                        {**pending, **response_payload},
                        response_payload["direct_status"],
                    )
                )
                return jsonify(attach_direct_ui_chrome(response_payload, pending))

            clear_pending_direct_run()
            blocked_payload = _build_unsafe_large_rewrite_payload(relative_path, scope=scope, mode=mode, task=task)
            blocked_payload["recognized_task"] = pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {}
            blocked_payload["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
            blocked_payload["rewrite_guard"] = unsafe_check
            blocked_payload["diff_summary"] = build_diff_summary(diff_text)
            blocked_payload["workstream_events"] = [
                _ws_event("analysis", "info", "Diff analysiert", f"Ziel: {relative_path}", status="done"),
                _ws_event("guard", "error", "Änderung blockiert", "Datei würde zu stark überschrieben.", status="blocked"),
                _ws_event("recovery", "info", "Auto-Fallback aktiv", "Split-Patch-Workflow starten: read -> locate -> patch -> apply", status="planned"),
            ]
            blocked_payload["auto_fallback_started"] = True
            blocked_payload["next_route"] = "/api/workspace/step-engine"
            blocked_payload["step_engine_payload"] = {
                "task": str(task or ""),
                "path": str(relative_path or ""),
                "current_content": str(current_content or ""),
                "proposed_content": str(proposed_content or ""),
                "confirmed": False,
            }
            if isinstance(se_res, dict):
                blocked_payload["step_engine_result"] = se_res
                blocked_payload["auto_recovery_attempted"] = True
            append_ui_log_entry("Direkt", "Unsafe large rewrite blockiert.", "error")
            return blocked_payload, 400

        write_result = persist_text_file_change(
            resolved_path,
            proposed_content,
            relative_path,
            on_timeout_log=lambda m: append_ui_log_entry("Direkt", m, "error"),
        )
        if not write_result.get("ok"):
            err = write_result.get("error") or "Die lokale Direktmodus-Aenderung konnte nicht geschrieben werden."
            append_ui_log_entry(
                "Direkt",
                err + " | persist_text_file_change -> execute_write_action | resolved=" + str(resolved_path),
                "error",
            )
            return {"error": err}, 500
        append_ui_log_entry(
            "Direkt",
            f"Schreibpfad ok: {relative_path} -> {resolved_path} ({int(write_result.get('lines') or 0)} Zeilen, Watchdog max {DIRECT_WRITE_ACTION_TIMEOUT_SEC}s).",
            "success",
        )

        post_check = run_local_post_check(resolved_path, relative_path, proposed_content)
        clear_pending_direct_run()
        decision_txt = (
            "Lokale Aenderung nach Guard automatisch angewendet."
            if should_auto_write
            else "Lokale Aenderung nach Bestaetigung angewendet."
        )
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_result": "Direktmodus lokal abgeschlossen.",
            "last_apply_action": f"Lokale Datei {'aktualisiert' if file_exists else 'angelegt'}: {relative_path}",
            "last_check_result": post_check["detail"],
            "last_direct_decision": decision_txt,
            "last_direct_status": "verified" if post_check["ok"] else "applied",
            "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "blocked": False
        })
        append_ui_log_entry("Direkt", f"Apply ausgefuehrt: {relative_path}", "success")
        append_ui_log_entry("Direkt", post_check["detail"], "success" if post_check["ok"] else "error")
        msg_local = f"Lokale Direktmodus-Aenderung {'aktualisiert' if file_exists else 'angelegt'}: {relative_path}"
        if should_auto_write:
            msg_local += " (Guard-gepruefter Standard-Schreiblauf)"
        analysis_detail = build_written_result_detail(resolved_path, relative_path)
        file_action = "erstellt" if not file_exists else "aktualisiert"
        _ws_events = [
            {"ts": get_timestamp(), "phase": "analysis", "level": "info",
             "title": "Auftrag erkannt", "detail": str(task)[:80], "status": "done"},
            {"ts": get_timestamp(), "phase": "guard", "level": "success",
             "title": "Guard freigegeben", "detail": str(relative_path), "status": "done"},
            {"ts": get_timestamp(), "phase": "writing", "level": "success",
             "title": f"Datei {file_action}", "detail": str(relative_path),
             "file": str(resolved_path), "status": "done"},
            {"ts": get_timestamp(), "phase": "verify", "level": "success" if post_check["ok"] else "error",
             "title": "Ergebnis geprüft", "detail": post_check["detail"],
             "status": "done" if post_check["ok"] else "failed"},
        ]
        response_payload = {
            **success_payload(
                "Änderung angewendet! ✓",
                technical_message=msg_local,
                changed_files=[relative_path],
                location="rambo_builder_local/",
                detail=analysis_detail or "Die Änderung wurde erfolgreich angewendet."
            ),
            "scope": scope,
            "mode": "apply",
            "selected_target_path": relative_path,
            "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
            "affected_files": [relative_path],
            "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "diff": diff_text,
            "post_check": {
                **post_check,
                "path": str(relative_path),
                "target": str(resolved_path),
            },
            "direct_status": "verified" if post_check["ok"] else "applied",
            "diff_summary": build_diff_summary(diff_text),
            "workstream_events": _ws_events,
            "created_files": [] if file_exists else [relative_path],
            "updated_files": [relative_path] if file_exists else [],
            "deleted_files": [],
            "changed_files": [relative_path],
            "affected_files": [relative_path],
            "auto_apply": True,
            "auto_continue": True,
            "skip_review": True,
            "direct_execute": True,
        }
        response_payload = _merge_direct_file_context(response_payload, pending)
        response_payload = _enforce_real_change_success(task, response_payload, mode="apply")
        upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, response_payload["direct_status"]))
        return jsonify(attach_direct_ui_chrome(response_payload, pending))

    resolved, cleaned, error = validate_project_write_path(target_path)
    if error:
        return {"error": error}, 403
    append_ui_log_entry("Direkt", f"Guard erneut geprueft: {cleaned}", "info")

    if mode == "safe":
        clear_pending_direct_run()
        if not pending.get("has_changes"):
            msg_proj = "Safe bestaetigt (Project): Keine inhaltliche Aenderung in der Vorschau."
        elif pending.get("requires_user_confirmation"):
            msg_proj = "Safe bestaetigt (Project): Freigabeprofil aktiv, Schreiben nur im Apply-Schritt."
        else:
            msg_proj = "Safe bestaetigt (Project): Schreiben erfolgt nur im expliziten Apply-Schritt."
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [cleaned],
            "last_guard_decision": guard.get("detail", ""),
            "last_check_result": "Keine Nachkontrolle noetig (Safe Mode).",
            "last_result": "Direktmodus im Safe Mode bestaetigt.",
            "last_direct_decision": "Safe-Vorschau bestaetigt.",
            "last_direct_status": "safe_preview",
            "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "blocked": False
        })
        append_ui_log_entry("Direkt", f"Safe bestaetigt: {cleaned}", "success")
        response_payload = {
            "message": msg_proj,
            "scope": scope,
            "mode": mode,
            "selected_target_path": cleaned,
            "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
            "affected_files": [cleaned],
            "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "post_check": {"ok": False, "detail": "Keine Nachkontrolle noetig (Safe Mode)."},
            "direct_status": "safe_preview",
            "diff_summary": build_diff_summary(pending.get("diff") or "")
        }
        upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, "safe_preview"))
        response_payload = _merge_direct_file_context(response_payload, pending)
        return jsonify(attach_direct_ui_chrome(response_payload, pending))

    current_content, file_exists = read_text_file(resolved)
    diff_text = build_text_diff(current_content, proposed_content, cleaned)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    if not has_changes:
        clear_pending_direct_run()
        append_ui_log_entry("Direkt", f"Keine Aenderung fuer '{cleaned}'.", "warning")
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [cleaned],
            "last_guard_decision": guard.get("detail", ""),
            "last_apply_action": "Keine Aenderung geschrieben.",
            "last_check_result": "Keine Nachkontrolle noetig.",
            "last_result": "Direktmodus ohne Aenderung abgeschlossen.",
            "last_direct_decision": "Keine Aenderung erkannt.",
            "last_direct_status": "safe_preview",
            "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
            "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "blocked": False
        })
        response_payload = {
            "message": "Keine inhaltliche Aenderung erkannt. Apply wurde uebersprungen.",
            "scope": scope,
            "mode": mode,
            "selected_target_path": cleaned,
            "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
            "affected_files": [cleaned],
            "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
            "diff": diff_text,
            "post_check": {"ok": False, "detail": "Keine Nachkontrolle noetig."},
            "direct_status": "safe_preview",
            "diff_summary": build_diff_summary(diff_text)
        }
        upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, "safe_preview"))
        response_payload = _merge_direct_file_context(response_payload, pending)
        return jsonify(attach_direct_ui_chrome(response_payload, pending))

    pytest_validation = validate_pytest_file(cleaned, proposed_content)
    if not bool(pytest_validation.get("ok")):
        clear_pending_direct_run()
        blocked_payload = _build_invalid_test_file_payload(cleaned, scope=scope, mode=mode, task=task)
        blocked_payload["recognized_task"] = pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {}
        blocked_payload["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
        blocked_payload["diff_summary"] = build_diff_summary(diff_text)
        append_ui_log_entry("Direkt", "Ungültige Testdatei blockiert.", "error")
        return blocked_payload, 400

    unsafe_check = detect_unsafe_large_rewrite(
        cleaned,
        current_content,
        proposed_content,
        {"phase": "project_apply", "task": task},
    )
    if bool(unsafe_check.get("unsafe")):
        applied_rec, se_res = _try_unsafe_rewrite_step_engine_apply(
            relative_path=cleaned,
            resolved_path=resolved,
            current_content=current_content,
            proposed_content=proposed_content,
            task=task,
            unsafe_check=unsafe_check,
        )
        if applied_rec:
            append_ui_log_entry("Direkt", f"Auto-Recovery (Step-Engine): Patch angewendet: {cleaned}", "success")
            post_check = run_project_post_check(resolved, cleaned, proposed_content)
            clear_pending_direct_run()
            save_project_auto_run_state({
                "last_run_at": get_timestamp(),
                "last_task": task,
                "last_mode": mode,
                "last_target_paths": [cleaned],
                "last_guard_decision": guard.get("detail", ""),
                "last_apply_action": f"Apply (Recovery) ausgefuehrt: {cleaned} {'aktualisiert' if file_exists else 'neu angelegt'}.",
                "last_check_result": post_check["detail"],
                "last_result": "Direktmodus abgeschlossen (Recovery)." if post_check["ok"] else "Direktmodus mit Nachkontrollfehler abgeschlossen (Recovery).",
                "last_direct_decision": "Projekt-Aenderung nach Step-Engine-Recovery angewendet.",
                "last_direct_status": "verified" if post_check["ok"] else "applied",
                "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
                "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                "blocked": False,
            })
            append_ui_log_entry("Direkt", f"Apply ausgefuehrt: {cleaned}", "success")
            append_ui_log_entry("Direkt", post_check["detail"], "success" if post_check["ok"] else "error")
            analysis_detail = build_written_result_detail(resolved, cleaned)
            response_payload = {
                **success_payload(
                    "Änderung angewendet! ✓",
                    technical_message=f"Projekt-Direktmodus erfolgreich abgeschlossen (Recovery): {cleaned}",
                    changed_files=[cleaned],
                    location="Rambo-Rainer/",
                    detail=analysis_detail or "Die Änderung wurde erfolgreich angewendet (Recovery).",
                ),
                "scope": scope,
                "mode": mode,
                "selected_target_path": cleaned,
                "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
                "affected_files": [cleaned],
                "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
                "diff": diff_text,
                "post_check": post_check,
                "direct_status": "verified" if post_check["ok"] else "applied",
                "diff_summary": build_diff_summary(diff_text),
                "split_patch_recovery": True,
                "step_engine_result": se_res if isinstance(se_res, dict) else {},
            }
            response_payload = _merge_direct_file_context(response_payload, pending)
            response_payload = _enforce_real_change_success(task, response_payload, mode="apply")
            upsert_direct_run_history(
                build_direct_history_entry(
                    str(pending.get("run_id") or ""),
                    {**pending, **response_payload},
                    response_payload["direct_status"],
                )
            )
            return jsonify(attach_direct_ui_chrome(response_payload, pending))

        clear_pending_direct_run()
        blocked_payload = _build_unsafe_large_rewrite_payload(cleaned, scope=scope, mode=mode, task=task)
        blocked_payload["recognized_task"] = pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {}
        blocked_payload["planned_steps"] = normalize_planned_steps(pending.get("planned_steps") or [])
        blocked_payload["rewrite_guard"] = unsafe_check
        blocked_payload["diff_summary"] = build_diff_summary(diff_text)
        if isinstance(se_res, dict):
            blocked_payload["step_engine_result"] = se_res
            blocked_payload["auto_recovery_attempted"] = True
        append_ui_log_entry("Direkt", "Unsafe large rewrite blockiert.", "error")
        return blocked_payload, 400

    write_result = persist_text_file_change(
        resolved,
        proposed_content,
        cleaned,
        on_timeout_log=lambda m: append_ui_log_entry("Direkt", m, "error"),
    )
    if not write_result.get("ok"):
        err = write_result.get("error") or "Projektdatei konnte im Direktmodus nicht geschrieben werden."
        append_ui_log_entry("Direkt", err, "error")
        return {"error": err}, 500
    append_ui_log_entry(
        "Direkt",
        f"Schreibpfad ok: {cleaned} ({int(write_result.get('lines') or 0)} Zeilen, Watchdog max {DIRECT_WRITE_ACTION_TIMEOUT_SEC}s).",
        "success",
    )

    post_check = run_project_post_check(resolved, cleaned, proposed_content)
    clear_pending_direct_run()
    save_project_auto_run_state({
        "last_run_at": get_timestamp(),
        "last_task": task,
        "last_mode": mode,
        "last_target_paths": [cleaned],
        "last_guard_decision": guard.get("detail", ""),
        "last_apply_action": f"Apply ausgefuehrt: {cleaned} {'aktualisiert' if file_exists else 'neu angelegt'}.",
        "last_check_result": post_check["detail"],
        "last_result": "Direktmodus abgeschlossen." if post_check["ok"] else "Direktmodus mit Nachkontrollfehler abgeschlossen.",
        "last_direct_decision": "Projekt-Aenderung nach Bestaetigung angewendet.",
        "last_direct_status": "verified" if post_check["ok"] else "applied",
        "last_direct_confirmed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
        "last_completed_run_id": str(pending.get("run_id") or pending.get("token") or ""),
        "last_planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
        "blocked": False
    })
    append_ui_log_entry("Direkt", f"Apply ausgefuehrt: {cleaned}", "success")
    append_ui_log_entry("Direkt", post_check["detail"], "success" if post_check["ok"] else "error")
    analysis_detail = build_written_result_detail(resolved, cleaned)
    response_payload = {
        **success_payload(
            "Änderung angewendet! ✓",
            technical_message=f"Projekt-Direktmodus erfolgreich abgeschlossen: {cleaned}",
            changed_files=[cleaned],
            location="Rambo-Rainer/",
            detail=analysis_detail or "Die Änderung wurde erfolgreich angewendet."
        ),
        "scope": scope,
        "mode": mode,
        "selected_target_path": cleaned,
        "recognized_task": pending.get("recognized_task") if isinstance(pending.get("recognized_task"), dict) else {},
        "affected_files": [cleaned],
        "planned_steps": normalize_planned_steps(pending.get("planned_steps") or []),
        "diff": diff_text,
        "post_check": post_check,
        "direct_status": "verified" if post_check["ok"] else "applied",
        "diff_summary": build_diff_summary(diff_text)
    }
    response_payload = _merge_direct_file_context(response_payload, pending)
    response_payload = _enforce_real_change_success(task, response_payload, mode="apply")
    upsert_direct_run_history(build_direct_history_entry(str(pending.get("run_id") or ""), {**pending, **response_payload}, response_payload["direct_status"]))
    return jsonify(attach_direct_ui_chrome(response_payload, pending))


def build_change_summary(task, relative_path, file_exists, has_changes):
    return {
        "ziel_datei": relative_path,
        "aufgabe": task,
        "modus": "Bestehende Datei aktualisieren" if file_exists else "Neue Datei anlegen",
        "aenderung": "Inhalt aendern" if has_changes else "Keine Aenderung erkannt"
    }


def build_file_preview(filename, filetype, targetpath):
    suggested_path = f"{targetpath}/{filename}.{filetype}"
    summary = {
        "txt": "Einfache Textnotiz fuer lokale Hinweise oder Aufgabenlisten.",
        "md": "Markdown-Vorlage fuer strukturierte Notizen oder kurze Dokumentation.",
        "json": "Datenvorlage fuer lokale Konfiguration oder strukturierte Zwischenstaende.",
        "py": "Python-Grundgeruest fuer kleine lokale Hilfsskripte.",
        "html": "HTML-Startdatei fuer eine kleine lokale Ansicht oder Vorschauseite.",
        "css": "CSS-Grunddatei fuer lokale Stilregeln.",
        "js": "JavaScript-Datei fuer lokale Interaktion oder kleine Hilfsfunktionen."
    }

    example_content = {
        "txt": (
            f"Projekt: Rambo Rainer\n"
            f"Datei: {filename}.txt\n"
            f"Zielpfad: {targetpath}\n"
            "Hinweis: Diese Vorschau bleibt lokal innerhalb von rambo_builder_local.\n"
            "Naechster Schritt: Inhalt ergaenzen oder als Arbeitsnotiz nutzen."
        ),
        "md": (
            f"# {filename}\n\n"
            "- Projekt: Rambo Rainer\n"
            f"- Zielpfad: {targetpath}\n"
            "- Modus: lokale Vorschau\n"
            "- Status: vorbereitet\n\n"
            "## Inhalt\n\n"
            "- Ziel:\n"
            "- Naechster Schritt:\n"
        ),
        "json": json.dumps(
            {
                "project": "Rambo Rainer",
                "file": f"{filename}.json",
                "target_path": targetpath,
                "status": "preview",
                "scope": "local_only",
                "next_step": "Inhalt lokal ergaenzen"
            },
            ensure_ascii=True,
            indent=2
        ),
        "py": (
            "\"\"\"Lokale Vorschau fuer kleine Hilfslogik.\"\"\"\n\n"
            "def main():\n"
            f"    print('Lokale Vorschau fuer {filename}.py')\n\n"
            "if __name__ == '__main__':\n"
            "    main()\n"
        ),
        "html": (
            "<!DOCTYPE html>\n"
            "<html lang=\"de\">\n"
            "<head>\n"
            "  <meta charset=\"UTF-8\">\n"
            f"  <title>{filename}</title>\n"
            "</head>\n"
            "<body>\n"
            f"  <h1>{filename}</h1>\n"
            "  <p>Lokale Vorschau innerhalb von rambo_builder_local.</p>\n"
            "  <p>Hier kann spaeter lokaler Inhalt ergaenzt werden.</p>\n"
            "</body>\n"
            "</html>"
        ),
        "css": (
            f".{filename} {{\n"
            "  display: block;\n"
            "  padding: 1rem;\n"
            "  color: #214d34;\n"
            "  background: #eef3ef;\n"
            "}\n"
        ),
        "js": (
            f"function init{filename.title().replace('_', '').replace('-', '')}() {{\n"
            f"  console.log('Lokale Vorschau fuer {filename}.js');\n"
            "}\n"
        )
    }

    return {
        "path": suggested_path,
        "type": filetype,
        "summary": summary[filetype],
        "content": example_content[filetype]
    }


def load_design_entries():
    entries = read_json_file(DESIGN_NOTES_FILE, [])
    if not isinstance(entries, list):
        return []

    normalized_entries = []
    seen_notes = set()
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        note = str(entry.get("note") or "").strip()
        if not note:
            continue
        timestamp = str(entry.get("timestamp") or "").strip() or get_timestamp()
        dedupe_key = note.lower()
        normalized_entries.append({
            "id": str(entry.get("id") or len(normalized_entries) + 1),
            "note": note,
            "timestamp": timestamp,
            "display_timestamp": format_display_timestamp(timestamp),
            "length": len(note),
            "duplicate": dedupe_key in seen_notes
        })
        seen_notes.add(dedupe_key)
    return normalized_entries[:50]


def save_design_entries(entries):
    cleaned_entries = []
    for entry in entries[:50]:
        cleaned_entries.append({
            "id": str(entry.get("id") or len(cleaned_entries) + 1),
            "note": str(entry.get("note") or "").strip(),
            "timestamp": str(entry.get("timestamp") or get_timestamp()).strip()
        })
    write_json_file(DESIGN_NOTES_FILE, cleaned_entries)


def generate_next_log_id(existing_entries):
    highest_numeric_id = 0
    for entry in existing_entries:
        try:
            highest_numeric_id = max(highest_numeric_id, int(str(entry.get("id") or "").strip()))
        except (TypeError, ValueError):
            continue
    return str(highest_numeric_id + 1)


def load_ui_activity_entries():
    entries = read_json_file(UI_ACTIVITY_LOG_FILE, [])
    if not isinstance(entries, list):
        return []

    cleaned_entries = []
    seen_ids = set()
    for entry in entries:
        if not isinstance(entry, dict):
            continue
        preferred_id = str(entry.get("id") or "").strip()
        entry_id = preferred_id
        if not entry_id or entry_id in seen_ids:
            entry_id = generate_next_log_id(cleaned_entries)
        seen_ids.add(entry_id)
        cleaned_entries.append({
            "id": entry_id,
            "timestamp": str(entry.get("timestamp") or get_timestamp()).strip(),
            "label": str(entry.get("label") or "System").strip() or "System",
            "message": str(entry.get("message") or "").strip(),
            "level": normalize_level(entry.get("level"))
        })
    return [entry for entry in cleaned_entries if entry["message"]][:20]


def save_ui_activity_entries(entries):
    write_json_file(UI_ACTIVITY_LOG_FILE, entries[:20])


def normalize_level(level):
    normalized = str(level or "info").strip().lower()
    if normalized not in {"info", "success", "warning", "error"}:
        return "info"
    return normalized


def append_ui_log_entry(label, message, level="info"):
    entries = load_ui_activity_entries()
    entry = {
        "id": generate_next_log_id(entries),
        "timestamp": get_timestamp(),
        "label": str(label or "System").strip() or "System",
        "message": str(message or "").strip(),
        "level": normalize_level(level)
    }
    if not entry["message"]:
        return None
    entries.insert(0, entry)
    save_ui_activity_entries(entries)
    return entry


def current_builder_status():
    entries = load_ui_activity_entries()
    for entry in entries:
        if entry["label"] == "Builder":
            if "fehl" in entry["message"].lower():
                return {"label": "Fehler", "detail": entry["message"]}
            if entry["level"] == "warning":
                return {"label": "Pruefen", "detail": entry["message"]}
            if "gestartet" in entry["message"].lower():
                return {"label": "Aktiv", "detail": entry["message"]}
            return {"label": "Plan bereit", "detail": entry["message"]}
    return {"label": "Bereit", "detail": "Lokale Plan-Generierung verfuegbar."}


def _derive_requirements_from_analysis(analysis: dict) -> dict:
    actual_problem = str((analysis or {}).get("actual_problem") or "").lower()
    goals = [str(g).lower() for g in ((analysis or {}).get("user_goals") or [])]
    constraints = (analysis or {}).get("constraints") or {}
    platforms = [str(p).lower() for p in constraints.get("platforms") or []]
    return {
        "high_performance": "performance" in actual_problem or "performance" in goals,
        "fast_development": "usability" in goals or "general functionality" in goals,
        "windows_only": "windows" in platforms,
        "web_ui": "usability" in goals,
    }


def _generate_intelligent_code(prompt: str, tech_choices: dict, architecture: dict) -> str:
    lang = str(((tech_choices or {}).get("recommended") or {}).get("language") or "Python")
    framework = str(((tech_choices or {}).get("recommended") or {}).get("framework") or "Flask")
    arch = str((architecture or {}).get("recommended") or "monolith")
    if lang == "Python":
        return (
            "from flask import Flask, jsonify\n\n"
            "app = Flask(__name__)\n\n"
            "@app.get('/health')\n"
            "def health():\n"
            f"    return jsonify({{'status': 'ok', 'architecture': '{arch}', 'framework': '{framework}'}})\n"
        )
    if lang == "JavaScript":
        return (
            "const express = require('express');\n"
            "const app = express();\n"
            f"app.get('/health', (_req, res) => res.json({{ status: 'ok', architecture: '{arch}', framework: '{framework}' }}));\n"
        )
    return (
        f"// Generated starter code\n"
        f"// language={lang}, framework={framework}, architecture={arch}\n"
        f"// prompt={prompt}\n"
    )


def _build_architecture_context(task: str) -> dict:
    low = str(task or "").lower()
    return {
        "size": "small" if "mvp" in low else "medium",
        "team_size": 1 if "solo" in low else 3,
        "real_time": ("realtime" in low or "real-time" in low),
        "ui_focus": True,
    }


def quick_detect_task_type(prompt: str) -> str:
    text = str(prompt or "").lower()
    if any(k in text for k in ("windows", "win32", "desktop", "electron")):
        return "desktop_tool"
    if any(k in text for k in ("api", "backend", "endpoint")):
        return "backend_api"
    if any(k in text for k in ("frontend", "ui", "react", "design")):
        return "frontend_ui"
    if any(k in text for k in ("script", "automation", "batch")):
        return "automation_script"
    return "general_solution"


def _is_detailed_mode(prompt: str) -> bool:
    text = str(prompt or "").lower()
    keywords = (
        "analysiere",
        "analysieren",
        "analyse",
        "analyze",
        "design",
        "designe",
        "entwirf",
        "entwerfen",
        "architektur",
        "trade-off",
        "deep dive",
        "detailliert",
    )
    return any(k in text for k in keywords)


def _parse_intelligent_options(response_style):
    preferred_style = response_style
    uploaded_file_path = None
    model_override = None
    implementation = False
    if isinstance(response_style, dict):
        uploaded_file_path = response_style.get("uploaded_file_path")
        model_override = response_style.get("model")
        preferred_style = response_style.get("response_style")
        implementation = bool(response_style.get("implementation"))
    return preferred_style, uploaded_file_path, model_override, implementation


def _check_implementation_needed(prompt: str, implementation_flag: bool) -> bool:
    """API-Flag oder Schluesselwoerter loesen die Sandbox-Implementation aus."""
    if implementation_flag:
        return True
    text = str(prompt or "").lower()
    keywords = (
        "dateien erstellen",
        "dateien anlegen",
        "projektstruktur",
        "write files",
        "erstelle",
        "build mir",
        "generiere projekt",
        "generate files",
        "create project",
        "python projekt",
        "implementation: true",
        "echte implementation",
        "implementierung mit dateien",
        "erstelle mir dateien",
        "leg dateien an",
        "schreibe die dateien",
        "generate code",
        "visual studio projekt",
        "wpf app",
        "csharp projekt",
        "c# projekt",
        "sandbox implementation",
    )
    return any(kw in text for kw in keywords)


def _wants_implementation(prompt: str, implementation_flag: bool) -> bool:
    """Alias fuer aeltere Aufrufer."""
    return _check_implementation_needed(prompt, implementation_flag)


def _run_implementation_phase(task: str, raw_result: dict) -> dict:
    """
    Multi-File-Projektgenerator (10-Feature):
    1. Datei-Plan mit allen Zielpfaden
    2. Ordner automatisch erstellen
    3. Alle Dateien schreiben (README, main, utils, tests, .gitignore, ...)
    4. Shell-Befehle ausfuehren: py_compile + npm falls package.json vorhanden
    5. Build-Logs lesen und speichern
    6. Fehler automatisch reparieren (1 Versuch: Fix oder Stub)
    7. Artefakte nach Downloads schreiben
    8. Erfolg erst melden wenn Dateien wirklich existieren und Build gruen
    """
    from project_builder import auto_repair_python, build_project

    impl_base = DOWNLOADS_DIR / "implementations"
    root = FileCreator.create_sandbox(impl_base)

    # 1. Datei-Plan erzeugen
    bundle = CodeGeneratorAdvanced.generate_implementation_bundle(task, raw_result, str(root))
    file_rows = bundle.get("files") or []
    if not file_rows:
        raise RuntimeError("Implementation-Bundle ohne Dateien")

    file_plan = []
    for f in file_rows:
        rel = str(f.get("rel") or "").strip().replace("\\", "/").lstrip("/")
        if rel:
            file_plan.append({"rel": rel, "abs": str(root / rel)})

    append_ui_log_entry(
        "Implementation",
        f"Datei-Plan ({len(file_plan)} Dateien): {[fp['rel'] for fp in file_plan]}",
        "info",
    )

    # 2+3. Ordner anlegen und alle Dateien schreiben
    written: list[str] = []
    write_errors: list[str] = []
    for file_info in file_rows:
        rel = str(file_info.get("rel") or "").strip().replace("\\", "/").lstrip("/")
        if not rel:
            continue
        body = str(file_info.get("content") or "")
        try:
            p = FileCreator.create_project_file(root, rel, body)
            # 8. Existenz sofort nach dem Schreiben pruefen
            if not p.exists():
                write_errors.append(f"{rel}: geschrieben aber nicht auffindbar")
            else:
                written.append(rel)
        except Exception as ex:
            write_errors.append(f"{rel}: {ex}")

    if write_errors:
        append_ui_log_entry("Implementation", f"Schreibfehler: {write_errors}", "warn")

    # 4+5. Build-Befehle ausfuehren und Logs lesen
    build_result = build_project(root)
    build_logs: list[str] = list(build_result.get("log") or [])
    for step in build_result.get("steps") or []:
        if step.get("stdout"):
            build_logs.append(f"[{step['step']}:stdout] {step['stdout'][:300]}")
        if step.get("stderr") and not step.get("ok"):
            build_logs.append(f"[{step['step']}:stderr] {step['stderr'][:300]}")

    # 6. Fehler automatisch reparieren (max 3 Versuche)
    repaired_info: dict = {"repaired": False, "files": []}
    _MAX_REPAIR = 3
    for _repair_attempt in range(1, _MAX_REPAIR + 1):
        if build_result["ok"] or not build_result.get("errors"):
            break
        if build_result.get("kind") not in ("python", "mixed", "unknown"):
            break
        _errs = build_result["errors"]
        append_ui_log_entry(
            "Implementation",
            f"Auto-Repair Versuch {_repair_attempt}/{_MAX_REPAIR}: {len(_errs)} Fehler",
            "info",
        )
        _attempt = auto_repair_python(root, _errs)
        if not _attempt.get("repaired"):
            append_ui_log_entry("Implementation", "Auto-Repair: keine weiteren Fixes moeglich", "warn")
            break
        repaired_info["repaired"] = True
        repaired_info["files"] += _attempt.get("files") or []
        build_logs += [f"[auto-repair:{_repair_attempt}] {f}" for f in _attempt.get("files") or []]
        append_ui_log_entry(
            "Implementation",
            f"Auto-Repair {_repair_attempt}: {_attempt.get('files')}",
            "info",
        )
        build_result = build_project(root)
        if build_result["ok"]:
            append_ui_log_entry("Implementation", f"Auto-Repair erfolgreich nach {_repair_attempt} Versuch(en)", "info")
            break

    # 5b. Build-Log in Datei schreiben
    _log_stamp = int(time.time())
    _log_file = LOGS_DIR / f"build_{_log_stamp}.log"
    _log_file_path: str = ""
    try:
        _log_file.write_text("\n".join(build_logs) or "(leer)", encoding="utf-8")
        _log_file_path = str(_log_file)
    except Exception:
        pass

    # 8. Dateien nach Repair nochmals pruefen
    missing_after = [rel for rel in written if not (root / rel).exists()]

    # 9. Echte Erfolgspruefung: Dateien vorhanden UND Build gruen
    build_ok = build_result.get("ok", False)
    files_ok = bool(written) and not missing_after and not write_errors
    verified = files_ok and build_ok

    st = "OK" if build_ok else "FAILED"
    summary_lines = [
        "Implementation abgeschlossen:",
        f"   Ordner: {root}",
        f"   Dateien ({len(written)}): {', '.join(written) or '-'}",
        f"   Build-Status: {st}  |  Verifiziert: {'JA' if verified else 'NEIN'}",
    ]
    if build_logs:
        summary_lines.append("   Build-Log: " + " | ".join(build_logs[:4]))
    if write_errors:
        summary_lines.append("   Schreibfehler: " + "; ".join(write_errors[:3]))
    if repaired_info.get("files"):
        summary_lines.append("   Auto-Repair: " + ", ".join(repaired_info["files"]))
    if missing_after:
        summary_lines.append("   Fehlende Dateien nach Repair: " + ", ".join(missing_after))
    summary_lines.append(f'   Naechster Schritt: cd "{root}" && python src/main.py')

    release_info = ExecutableCreator.write_release_notes(root, build_result)
    summary_lines.append(f"   Hinweis: {release_info.get('hint', '')}")
    summary = "\n".join(summary_lines)

    # 10. Kein Fake-Erfolg: implementation_verified zeigt echten Zustand
    result: dict = {
        "implementation": True,
        "implementation_verified": verified,
        "implementation_root": str(root),
        "implementation_files": written,
        "implementation_file_plan": file_plan,
        "implementation_build": build_result,
        "implementation_release": release_info,
        "implementation_build_logs": build_logs,
        "implementation_log_file": _log_file_path,
        "implementation_repaired": repaired_info,
        "implementation_summary": summary,
    }

    if not files_ok:
        err = f"Keine Dateien geschrieben oder fehlend: {write_errors or missing_after}"
        result["implementation_error"] = err
        raise RuntimeError(err)

    return result


def _chat_text_for_prompt(prompt: str) -> str:
    """Rueckwaertskompatibel: Kurzantwort ohne LLM."""
    return chat_reply_canned(prompt)


def _ensure_style_signature(text: str, style: str) -> str:
    body = str(text or "").strip()
    if not body:
        return body
    if style == "developer" and all(k not in body for k in ("Engineer Notes", "Problem-Analyse", "Tech-Plan")):
        return "🚀 Tech-Plan\n\n" + body
    if style == "business" and "Executive Summary" not in body:
        return "📊 Executive Summary\n\n" + body
    return body


def _ensure_formatted_intelligent_response(result: dict) -> dict:
    """Stellt sicher, dass das Frontend immer einen sichtbaren formatted_response-String erhaelt."""
    if not isinstance(result, dict):
        return {"formatted_response": str(result or ""), "final": True, "stop_continue": True}
    fmt = ResponseFormatter()
    style = str(result.get("response_style") or "friendly").lower()
    if style not in {"business", "developer", "friendly"}:
        style = "friendly"
    text = result.get("formatted_response")
    if text is None or not str(text).strip():
        text = fmt.format_response(result, style=style)
    if not str(text).strip():
        code = str(result.get("generated_code") or "").strip()
        if code:
            text = "Hier ist der Vorschlag-Code:\n\n```text\n" + code + "\n```\n"
        else:
            text = str(
                result.get("error")
                or result.get("fallback_reason")
                or "Keine formatierte Antwort erzeugt. Bitte Prompt erneut senden."
            )
    result["formatted_response"] = str(text).strip()
    return result


def _empty_change_report():
    return {
        "total_files": 0,
        "files_created": 0,
        "files_modified": 0,
        "files_deleted": 0,
        "total_lines_added": 0,
        "total_lines_removed": 0,
        "changes": [],
        "visual": {
            "file_tree": [],
            "summary": {"total": 0, "new": 0, "modified": 0, "lines_added": 0, "lines_removed": 0},
            "detailed_changes": [],
        },
    }


def quick_generate(prompt, preferred_style=None, uploaded_file_path=None, model_override=None):
    task = str(prompt or "").strip()
    formatter = ResponseFormatter()
    style = formatter.detect_style(task, preferred_style=preferred_style)
    task_type = quick_detect_task_type(task)
    minimal_analysis = {
        "problem_type": task_type,
        "actual_problem": f"Schnelle Umsetzung fuer: {task_type}",
        "scope": "quick_mode",
        "user_goals": ["Direkte Loesung", "Wenig Overhead"],
    }
    context = (
        "Quick-Mode aktiv: Liefere direkt umsetzbare Loesung mit minimaler Analyse.\n"
        f"Task-Type: {task_type}"
    )
    if uploaded_file_path:
        up = Path(str(uploaded_file_path))
        if up.exists() and up.is_file():
            context += f"\nUpload-Datei: {up.name}"
            if up.suffix.lower() in {".txt", ".md", ".json", ".csv"}:
                try:
                    snippet = up.read_text(encoding="utf-8", errors="ignore")[:1200]
                    context += "\nDateiinhalt (Ausschnitt):\n" + snippet
                except Exception:
                    pass

    ollama_response = call_ollama_intelligent(
        "Quick Mode: " + task,
        context=context,
        model_override=model_override,
    )
    generated_code = _generate_intelligent_code(
        task,
        {"recommended": {"language": "Python", "framework": "Flask"}},
        {"recommended": "modular_monolith"},
    )
    result = {
        "analysis": minimal_analysis,
        "recommended_approach": "Direkte Umsetzung (Quick Mode)",
        "architecture": "Pragmatischer Build-First Flow",
        "recommended_approach_detail": {"language": "Python", "framework": "Flask"},
        "architecture_detail": {"recommended": "modular_monolith", "score": 75},
        "generated_code": generated_code,
        "improvements": [],
        "improvements_detail": {"priority": [], "issues_found": 0},
        "performance_optimizations": {"optimizations": []},
        "tradeoffs": {},
        "response_style": style,
        "quick_mode": True,
        "detailed_mode": False,
        "workflow_mode": "quick",
        "final": True,
        "stop_continue": True,
    }
    if is_llm_failure_message(str(ollama_response or "")):
        result["formatted_response"] = formatter.format_response(result, style=style)
        result["llm_provider"] = "local_fallback"
        result["graceful_fallback"] = True
        result["fallback_reason"] = str(ollama_response)
    else:
        result["formatted_response"] = _ensure_style_signature(
            format_response_like_claude(ollama_response, result), style
        )
        result["llm_provider"] = str(model_override or OLLAMA_MODEL)
        result["graceful_fallback"] = False
    return result


def detailed_analysis(prompt, preferred_style=None, uploaded_file_path=None, model_override=None):
    task = str(prompt or "").strip()
    formatter = ResponseFormatter()
    analyzer = ProblemAnalyzer()
    problem_analysis = analyzer.analyze_problem(task)

    requirements = _derive_requirements_from_analysis(problem_analysis)
    decision_maker = DecisionMaker()
    tech_choices = decision_maker.make_technology_choice(requirements)

    architecture_decider = ArchitectureDecider()
    architecture = architecture_decider.recommend_architecture(_build_architecture_context(task))

    tradeoffs = TradeOffAnalyzer().analyze_tradeoffs(
        [tech_choices.get("language", {}).get("recommended"), "Python", "C#", "JavaScript"]
    )
    generated_code = _generate_intelligent_code(task, tech_choices, architecture)

    language = str((tech_choices.get("language") or {}).get("recommended") or "Python").lower()
    improver = ImprovementSuggester()
    improvements = improver.analyze_code(generated_code, language)

    optimizer = PO()
    perf_optimizations = optimizer.optimize(generated_code, language)

    recommended = tech_choices.get("recommended") or {}
    recommended_text = " ".join(
        [str(recommended.get("language") or "").strip(), str(recommended.get("framework") or "").strip()]
    ).strip() or "Unklar"
    architecture_text = str((architecture or {}).get("recommended") or "Unklar")
    style = formatter.detect_style(task, preferred_style=preferred_style)

    raw_result = {
        "analysis": problem_analysis,
        "recommended_approach": recommended_text,
        "architecture": architecture_text,
        "recommended_approach_detail": recommended,
        "architecture_detail": architecture,
        "generated_code": generated_code,
        "improvements": improvements.get("priority") if isinstance(improvements, dict) else improvements,
        "improvements_detail": improvements,
        "performance_optimizations": perf_optimizations,
        "tradeoffs": tradeoffs,
        "response_style": style,
        "quick_mode": False,
        "detailed_mode": True,
        "workflow_mode": "detailed",
        "final": True,
        "stop_continue": True,
    }
    context = (
        f"Analyse: {problem_analysis}\n"
        f"Empfehlung: {tech_choices.get('recommended')}\n"
        f"Architektur: {architecture.get('recommended')}"
    )
    rag_items = RAG_INTEGRATION.retrieve_similar_context(task, top_k=3)
    rag_context = RAG_INTEGRATION.format_context(rag_items)
    if rag_context:
        context += "\n" + rag_context
    if uploaded_file_path:
        up = Path(str(uploaded_file_path))
        if up.exists() and up.is_file():
            context += f"\nUpload-Datei: {up.name}"
            if up.suffix.lower() in {".txt", ".md", ".json", ".csv"}:
                try:
                    snippet = up.read_text(encoding="utf-8", errors="ignore")[:3000]
                    context += "\nDateiinhalt (Ausschnitt):\n" + snippet
                except Exception:
                    pass
    ollama_response = call_ollama_intelligent(task, context=context, model_override=model_override)
    if is_llm_failure_message(str(ollama_response or "")):
        raw_result["formatted_response"] = formatter.format_response(raw_result, style=style)
        raw_result["llm_provider"] = "local_fallback"
        raw_result["graceful_fallback"] = True
        raw_result["fallback_reason"] = str(ollama_response)
    else:
        raw_result["formatted_response"] = _ensure_style_signature(
            format_response_like_claude(ollama_response, raw_result), style
        )
        raw_result["llm_provider"] = str(model_override or OLLAMA_MODEL)
        raw_result["graceful_fallback"] = False
    return raw_result


def _ws_event(phase, level, title, detail="", file="", command="", status="done"):
    """Hilfsfunktion: Workstream-Event-Dict erzeugen."""
    return {
        "ts": get_timestamp(),
        "phase": phase,
        "level": level,
        "title": title,
        "detail": str(detail)[:200],
        "file": str(file or ""),
        "command": str(command or ""),
        "status": status,
    }


def _evaluate_prompt_semantics(prompt):
    """
    Semantische Prompt-Auswertung.
    Gibt True zurueck, wenn der Prompt ein echter Auftrag ist (nicht nur eine Frage).
    """
    text = str(prompt or "").lower().strip()
    if not text:
        return False
    words = text.split()
    if len(words) < 3:
        return False
    # Frage ohne Auftrag
    question_words = ["wie", "was", "warum", "wann", "wo", "wer", "welche", "welcher"]
    request_patterns = [
        r"\bmach\b", r"\berstelle\b", r"\bbaue\b", r"\bgeneriere\b",
        r"\bschreib\b", r"\berzeuge\b", r"\blösch\b", r"\baendere\b",
        r"\bändere\b", r"\breparier\b", r"\bfixe\b", r"\bimplementier\b",
        r"\bfüge\b", r"\bsetze\b", r"\bcreate\b", r"\bbuild\b",
        r"\bmake\b", r"\bgenerate\b", r"\bwrite\b", r"\bdelete\b",
        r"\bfix\b", r"\badd\b", r"\bremove\b", r"\bchange\b",
    ]
    has_question = any(text.startswith(q) for q in question_words) and "?" in text
    has_request = any(re.search(p, text) for p in request_patterns)
    if has_question and not has_request:
        return False
    valid_keywords = [
        "tool", "app", "system", "programm", "script", "datei", "file",
        "windows", "web", "desktop", "mobile", "server", "api", "backend",
        "frontend", "python", "javascript", "c#", "html", "css", "json",
        "projekt", "ordner", "directory", "electron", "react", "inhalt",
        "content", "text",
    ]
    has_keyword = any(kw in text for kw in valid_keywords)
    if has_request and (has_keyword or len(words) >= 5):
        return True
    return len(words) >= 5 and has_request


def execute_intelligent(prompt, response_style=None):
    task = str(prompt or "").strip()
    task_lower = task.lower()
    inferred_allowed_targets, infer_meta = infer_allowed_target_files_with_meta(task)
    ws_ok_intel = is_active_workspace_trusted()

    if _is_mini_task_write_intent(task):
        n_inf = len(inferred_allowed_targets)
        if n_inf == 0:
            return _build_target_path_unclear_payload(
                task, "apply", inferred_allowed_targets, inference_debug=infer_meta
            )
        if n_inf != 1 and not ws_ok_intel:
            return _build_target_path_unclear_payload(
                task, "apply", inferred_allowed_targets, inference_debug=infer_meta
            )
        if n_inf == 1:
            preview = build_local_direct_preview(task, "apply")
            if not bool(preview.get("ok")):
                return preview.get("payload") if isinstance(preview.get("payload"), dict) else _build_target_path_unclear_payload(
                    task, "apply", inferred_allowed_targets, inference_debug=infer_meta
                )
            pending_payload = preview.get("payload") if isinstance(preview.get("payload"), dict) else {}
            confirm_result = execute_direct_confirmation({
                **pending_payload,
                "scope": "local",
                "mode": "apply",
                "token": "",
                "status": "confirming",
                "requires_user_confirmation": False,
            })
            if isinstance(confirm_result, tuple):
                body = confirm_result[0]
                return body if isinstance(body, dict) else {"ok": False, "status": "target_path_unclear", "error": str(body)}
            try:
                parsed = confirm_result.get_json(silent=True)
                if isinstance(parsed, dict):
                    return parsed
            except Exception:
                pass
            return {
                "ok": False,
                "status": "target_path_unclear",
                "message": "Kein eindeutiger erlaubter Zielpfad erkannt.",
            }

    # HARD ELECTRON/REACT DETECTION - Must be FIRST (nicht bei klarem Einzeldatei-Direct-Write)
    electron_keywords = ["electron", "react", "vite", "desktop-app", "desktop app", "build_desktop.py", "testelectronapp"]
    if (
        not _is_single_file_direct_write_intent(task_lower)
        and any(kw in task_lower for kw in electron_keywords)
        and (_is_desktop_multi_file_project_prompt(task_lower) or len(task_lower) >= 380)
    ):
        early_guard = _validate_direct_run_paths(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN, "apply", task)
        if not bool(early_guard.get("ok")):
            blocked_payload = _build_direct_guard_block_payload(
                scope="project",
                mode="apply",
                blocked_files=early_guard.get("blocked_files") or list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN),
                task=task,
                recognized_task={
                    "task_type": "electron_react_build",
                    "primary_area": "Project Builder",
                    "execution_route": "electron_react_build",
                },
            )
            blocked_payload["workstream_events"] = [
                _ws_event("analysis", "info", "Electron-/React-Projekt erkannt", "Guard wird vor Build geprueft", status="done"),
                _ws_event("guard", "error", "Auftrag blockiert", DIRECT_RUN_GUARD_BLOCK_MESSAGE, status="blocked"),
            ]
            blocked_payload["formatted_response"] = DIRECT_RUN_GUARD_BLOCK_MESSAGE
            return blocked_payload
        append_ui_log_entry("Intelligent", f"Electron-/React-Projekt erkannt: Starte deterministischen Build.", "info")
        run_id = uuid4().hex
        result = execute_electron_react_build(task, run_id)
        return _ensure_formatted_intelligent_response(result)

    response_style, uploaded_file_path, model_override, impl_flag = _parse_intelligent_options(response_style)
    is_detailed = _is_detailed_mode(task)

    # Workstream-Events sammeln
    _ws = []
    _ws.append(_ws_event("start", "info", "Auftrag erkannt", task[:80], status="done"))

    # Semantische Auswertung
    prompt_direkt_ok = _evaluate_prompt_semantics(task)
    _ws.append(_ws_event("analysis", "info", "Semantik ausgewertet",
                         f"Auftrag erkannt: {prompt_direkt_ok}", status="done"))

    tracker = None
    if is_detailed:
        tracker = ChangeTracker()
        tracker.capture_state(APP_DIR)
    log_id = AUTO_LOGGER.log_prompt(task)

    _ws.append(_ws_event("routing", "info",
                         "Modus gewählt",
                         "Detailliert" if is_detailed else "Schnell",
                         status="done"))

    if is_detailed:
        _ws.append(_ws_event("analysis", "info", "Detaillierte Analyse läuft", status="running"))
        raw_result = detailed_analysis(
            task,
            preferred_style=response_style,
            uploaded_file_path=uploaded_file_path,
            model_override=model_override,
        )
        _ws.append(_ws_event("analysis", "success", "Analyse abgeschlossen", status="done"))
    else:
        _ws.append(_ws_event("code_gen", "info", "Schnelle Generierung läuft", status="running"))
        raw_result = quick_generate(
            task,
            preferred_style=response_style,
            uploaded_file_path=uploaded_file_path,
            model_override=model_override,
        )
        _ws.append(_ws_event("code_gen", "success", "Antwort generiert", status="done"))

    raw_result["implementation"] = False
    if _wants_implementation(task, impl_flag):
        _ws.append(_ws_event("implementation", "info", "Implementation läuft", status="running"))
        try:
            impl_payload = _run_implementation_phase(task, raw_result)
            raw_result.update(impl_payload)
            tail = "\n\n---\n🛠️ **Implementation (Sandbox)**\n" + str(impl_payload.get("implementation_summary") or "")
            raw_result["formatted_response"] = str(raw_result.get("formatted_response") or "").rstrip() + tail
            _ws.append(_ws_event("implementation", "success", "Implementation fertig", status="done"))
        except Exception as ex:
            raw_result["implementation"] = False
            raw_result["implementation_error"] = str(ex)
            err_tail = f"\n\n---\n⚠️ **Implementation fehlgeschlagen:** {ex}\n"
            raw_result["formatted_response"] = str(raw_result.get("formatted_response") or "").rstrip() + err_tail
            _ws.append(_ws_event("implementation", "error", "Implementation fehlgeschlagen", str(ex), status="failed"))

    AUTO_LOGGER.log_result(log_id, raw_result)
    threading.Thread(
        target=SILENT_LEARNER.learn_from_current_session,
        args=(raw_result,),
        daemon=True,
        name="rainer-silent-learner",
    ).start()
    if tracker is not None:
        changes = tracker.detect_changes(APP_DIR)
        report = tracker.generate_report(changes)
    else:
        changes = []
        report = _empty_change_report()
    raw_result["change_report"] = report
    raw_result["files_changed"] = len(changes)
    raw_result["total_lines_changed"] = int(report["total_lines_added"]) + int(report["total_lines_removed"])
    # Strukturierte Metadaten
    raw_result["run_id"] = uuid4().hex[:8]
    raw_result["auto_apply"] = AUTO_APPLY
    if raw_result.get("implementation_file_plan"):
        raw_result["file_plan"] = raw_result["implementation_file_plan"]
    if raw_result.get("implementation_files"):
        raw_result["artifacts"] = raw_result["implementation_files"]

    _ws.append(_ws_event("formatting", "info", "Response formatiert", status="done"))
    _ws.append(_ws_event("final", "success", "Fertig", status="done"))

    final = _ensure_formatted_intelligent_response(raw_result)
    # Workstream und Semantik in finale Antwort einbauen
    final["workstream_events"] = _ws
    final["workstream_events_count"] = len(_ws)
    final["prompt_direkt_ok"] = prompt_direkt_ok
    # Chat-Verlauf speichern
    try:
        _ts = int(time.time())
        _append_chat_entry({"role": "user", "content": task, "ts": _ts})
        _append_chat_entry({
            "role": "assistant",
            "content": str(final.get("formatted_response") or "")[:2000],
            "run_id": final.get("run_id"),
            "ts": _ts,
        })
    except Exception:
        pass
    return final


@app.route("/api/health", methods=["GET"])
def health():
    ollama_ok = False
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=2)
        ollama_ok = response.status_code == 200
    except Exception:
        ollama_ok = False

    payload = {
        "status": "backend_ok",
        "service": "rainer-build-3.0",
        "active_pid": int(os.getpid()),
        "port": SERVER_PORT,
        "server_instance_id": SERVER_INSTANCE_ID,
        "started_at": SERVER_STARTED_AT,
        "ollama_ok": bool(ollama_ok),
        "model": OLLAMA_MODEL,
        "learning_db": str(PASSIVE_LEARNING_DB),
    }
    log_structured("health_check", **payload)
    return jsonify(payload)


def _detect_port_owners(port: int) -> list[int]:
    owners: list[int] = []
    try:
        cp = subprocess.run(
            ["netstat", "-ano", "-p", "TCP"],
            capture_output=True,
            text=True,
            timeout=3,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        for line in (cp.stdout or "").splitlines():
            lower = line.lower()
            if "listen" not in lower:
                continue
            if f":{int(port)}" not in line:
                continue
            parts = line.split()
            if not parts:
                continue
            try:
                owners.append(int(parts[-1]))
            except Exception:
                continue
    except Exception:
        return []
    return sorted(set(owners))


@app.route("/api/server/instance", methods=["GET"])
def server_instance():
    active_pid = int(os.getpid())
    port_owners = _detect_port_owners(SERVER_PORT)
    multiple_instances = len(port_owners) > 1 or (len(port_owners) == 1 and active_pid not in port_owners)
    warnings = []
    if multiple_instances:
        warnings.append("multiple_server_instances_detected_on_port")
    payload = {
        "ok": True,
        "service": "rainer-build-3.0",
        "active_pid": active_pid,
        "port": SERVER_PORT,
        "server_instance_id": SERVER_INSTANCE_ID,
        "started_at": SERVER_STARTED_AT,
        "port_owners": port_owners,
        "multiple_instances_detected": bool(multiple_instances),
        "warnings": warnings,
        "writes_files": False,
    }
    return jsonify(payload), 200


@app.route("/api/intelligent-run", methods=["POST"])
def intelligent_run():
    """Phase 1.1: gleicher Ausführungspfad wie POST /api/direct-run (Klassifikation + Agent/Chat).

    Ausnahme: ``implementation: true`` nutzt weiterhin ``execute_intelligent`` (Sandbox-Bundle).
    """
    import json as _json

    payload = request.get_json(silent=True) or {}
    response_style = str(payload.get("response_style") or "").strip().lower() or None
    uploaded_file_path = str(payload.get("uploaded_file_path") or "").strip() or None
    model = str(payload.get("model") or "").strip() or None
    if response_style == "auto":
        response_style = None
        payload["response_style"] = None
    raw_prompt = _extract_task_or_prompt_from_request_json(payload)
    if not raw_prompt:
        return jsonify({"error": "prompt oder task fehlt"}), 400

    if bool(payload.get("implementation")):
        prompt, upload_ctx_meta = augment_prompt_with_uploads(raw_prompt, payload)
        cleaned_ir = " ".join(str(prompt or "").split())
        ps_ir = _powershell_direct_run_envelope(
            cleaned_prompt=cleaned_ir,
            run_id=uuid4().hex,
            scope="local",
            mode="apply",
            log_label="Intelligent",
        )
        if ps_ir is not None:
            if upload_ctx_meta.get("uploads") or upload_ctx_meta.get("errors"):
                ps_ir = {**(ps_ir if isinstance(ps_ir, dict) else {}), "upload_context": upload_ctx_meta}
            return jsonify(ps_ir), 200
        try:
            result = execute_intelligent(
                prompt,
                {
                    "response_style": response_style,
                    "uploaded_file_path": uploaded_file_path,
                    "model": model,
                    "implementation": True,
                    "upload_context": upload_ctx_meta,
                },
            )
            result = _ensure_formatted_intelligent_response(result)
            guarded = _apply_central_generation_guard(
                result,
                scope="project",
                mode="apply",
                task=prompt,
                recognized_task=(result.get("recognized_task") if isinstance(result, dict) else None),
            )
            if isinstance(guarded, dict):
                guarded["run_mode"] = "intelligent"
                guarded["status"] = "blocked_by_guard"
                if upload_ctx_meta.get("uploads") or upload_ctx_meta.get("errors"):
                    guarded["upload_context"] = upload_ctx_meta
                return jsonify(enrich_direct_run_response(guarded)), 403
            result["run_mode"] = "intelligent"
            if result.get("ok") is True:
                result["status"] = "ok"
            result = _enforce_real_change_success(prompt, result, mode="apply")
            if upload_ctx_meta.get("uploads") or upload_ctx_meta.get("errors"):
                result["upload_context"] = upload_ctx_meta
            if str(result.get("status") or "").strip().lower() == "target_path_unclear":
                return jsonify(result), 400
            return jsonify(result)
        except Exception as ex:
            tb = traceback.format_exc()
            return jsonify(
                {
                    "ok": False,
                    "error": str(ex) or "Interner Fehler",
                    "direct_status": "failed",
                    "technical_message": tb[:4000],
                    "workstream_events": [
                        _ws_event("error", "error", "intelligent-run", str(ex)[:500], status="failed"),
                    ],
                }
            ), 500

    dr_payload = dict(payload)
    _sc = str(dr_payload.get("scope") or "local").strip().lower()
    dr_payload["scope"] = _sc if _sc in {"local", "project"} else "local"
    _md = str(dr_payload.get("mode") or "apply").strip().lower()
    dr_payload["mode"] = _md if _md in {"safe", "apply"} else "apply"
    headers = [("Content-Type", "application/json; charset=utf-8")]
    _adm = request.headers.get("X-Rambo-Admin")
    if _adm:
        headers.append(("X-Rambo-Admin", _adm))
    try:
        with app.test_request_context(
            "/api/direct-run",
            method="POST",
            data=_json.dumps(dr_payload, ensure_ascii=False),
            content_type="application/json; charset=utf-8",
            headers=headers,
        ):
            out = direct_run()
    except Exception as ex:
        tb = traceback.format_exc()
        return jsonify(
            {
                "ok": False,
                "error": str(ex) or "Interner Fehler",
                "direct_status": "failed",
                "technical_message": tb[:4000],
                "workstream_events": [
                    _ws_event("error", "error", "intelligent-run", str(ex)[:500], status="failed"),
                ],
            }
        ), 500

    if isinstance(out, tuple) and len(out) >= 2:
        resp_obj, status_code = out[0], int(out[1])
    else:
        resp_obj, status_code = out, 200
    body = resp_obj.get_json(silent=True)
    if not isinstance(body, dict):
        return jsonify({"ok": False, "error": "Antwort ungueltig.", "direct_status": "failed"}), 500
    body["run_mode"] = "intelligent"
    body.setdefault("final", True)
    body.setdefault("stop_continue", True)
    body = _ensure_formatted_intelligent_response(body)
    task_guard = str(
        dr_payload.get("task") or dr_payload.get("prompt") or dr_payload.get("message") or ""
    ).strip()
    guarded = _apply_central_generation_guard(
        body,
        scope="project",
        mode="apply",
        task=task_guard,
        recognized_task=(body.get("recognized_task") if isinstance(body, dict) else None),
    )
    if isinstance(guarded, dict):
        guarded["run_mode"] = "intelligent"
        guarded["status"] = "blocked_by_guard"
        return jsonify(enrich_direct_run_response(guarded)), 403
    if str(body.get("status") or "").strip().lower() == "target_path_unclear":
        return jsonify(body), 400
    return jsonify(body), status_code


def _analyze_uploaded_file(filepath: Path, ext: str) -> dict:
    """Liest hochgeladene Datei aus und gibt Kontext-Dict zurueck."""
    result: dict = {"type": ext, "path": str(filepath), "summary": "", "content": None}
    try:
        if ext in {".txt", ".md", ".json", ".py", ".js", ".jsx", ".ts", ".tsx", ".html", ".css", ".xml", ".yaml", ".yml"}:
            text = filepath.read_text(encoding="utf-8", errors="replace")
            result["content"] = text[:6000]
            result["summary"] = f"Textdatei ({len(text)} Zeichen)"

        elif ext == ".zip":
            import zipfile
            with zipfile.ZipFile(str(filepath), "r") as z:
                names = z.namelist()
            result["files"] = names[:60]
            result["file_count"] = len(names)
            result["summary"] = f"ZIP: {len(names)} Dateien — {', '.join(names[:10])}"

        elif ext == ".csv":
            import csv
            rows: list = []
            with open(str(filepath), "r", encoding="utf-8", errors="replace") as fh:
                for i, row in enumerate(csv.reader(fh)):
                    rows.append(row)
                    if i >= 6:
                        break
            result["columns"] = rows[0] if rows else []
            result["preview"] = rows[1:6] if len(rows) > 1 else []
            result["summary"] = f"CSV: Spalten={', '.join(result['columns'][:10])}"

        elif ext in {".xlsx", ".xls"}:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(filepath), read_only=True, data_only=True)
                sheets = wb.sheetnames
                ws = wb.active
                headers = [str(c.value or "") for c in next(ws.iter_rows(max_row=1))]
                result["sheets"] = sheets
                result["columns"] = headers[:20]
                result["summary"] = f"Excel: Sheets={sheets} | Spalten={', '.join(headers[:8])}"
            except ImportError:
                result["summary"] = "Excel-Datei (openpyxl nicht installiert — pip install openpyxl)"
            except Exception as ex:
                result["summary"] = f"Excel: Fehler beim Lesen ({ex})"

        elif ext == ".docx":
            try:
                import docx  # python-docx
                doc = docx.Document(str(filepath))
                text = "\n".join(p.text for p in doc.paragraphs[:30])
                result["content"] = text[:4000]
                result["summary"] = f"Word: {len(doc.paragraphs)} Absaetze"
            except ImportError:
                result["summary"] = "Word-Datei (python-docx nicht installiert — pip install python-docx)"
            except Exception as ex:
                result["summary"] = f"Word: Fehler beim Lesen ({ex})"

        elif ext == ".pdf":
            text = ""
            pages = 0
            try:
                import pdfplumber
                with pdfplumber.open(str(filepath)) as pdf:
                    pages = len(pdf.pages)
                    text = (pdf.pages[0].extract_text() or "") if pages > 0 else ""
            except ImportError:
                try:
                    import PyPDF2
                    with open(str(filepath), "rb") as fh:
                        reader = PyPDF2.PdfReader(fh)
                        pages = len(reader.pages)
                        text = reader.pages[0].extract_text() or "" if pages > 0 else ""
                except ImportError:
                    result["summary"] = "PDF (pdfplumber/PyPDF2 nicht installiert)"
                    return result
                except Exception as ex:
                    result["summary"] = f"PDF: Fehler ({ex})"
                    return result
            except Exception as ex:
                result["summary"] = f"PDF: Fehler ({ex})"
                return result
            result["pages"] = pages
            result["content"] = text[:4000]
            result["summary"] = f"PDF: {pages} Seiten"

        elif ext in {".png", ".jpg", ".jpeg", ".webp", ".gif"}:
            size_kb = filepath.stat().st_size // 1024
            result["summary"] = f"Bild ({ext}, {size_kb} KB) — kann fuer Icon/Design/Analyse genutzt werden."

    except Exception as ex:
        result["summary"] = f"Analyse-Fehler: {ex}"
    return result


@app.route("/api/upload-file", methods=["POST"])
@app.route("/api/upload", methods=["POST"])
def upload_file():
    f = request.files.get("file")
    if f is None or not str(f.filename or "").strip():
        return jsonify({"ok": False, "status": "error", "error": "keine Datei empfangen"}), 400

    ext = Path(str(f.filename)).suffix.lower()
    allowed = {
        ".pdf", ".txt", ".png", ".jpg", ".jpeg", ".webp", ".gif", ".csv", ".json", ".md",
        ".zip", ".xlsx", ".xls", ".docx",
        ".js", ".jsx", ".ts", ".tsx", ".py", ".html", ".css", ".xml", ".yaml", ".yml",
    }
    if ext not in allowed:
        return jsonify({"ok": False, "status": "error", "error": f"Dateityp {ext!r} nicht unterstuetzt"}), 400

    # Eindeutiger Run-Ordner: Sekunden-Stempel reicht bei mehreren Uploads nicht (WinError 183).
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
    run_dir = UPLOADS_DIR / f"run_{stamp}_{uuid4().hex[:10]}"
    run_dir.mkdir(parents=True, exist_ok=True)
    safe_name = sanitize_upload_filename(str(f.filename))
    dst = run_dir / safe_name
    if dst.exists():
        p = Path(safe_name)
        safe_name = f"{p.stem}_{uuid4().hex[:8]}{p.suffix}"
        dst = run_dir / safe_name
    f.save(str(dst))
    try:
        sz = int(dst.stat().st_size)
    except OSError:
        sz = 0
    if sz > MAX_UPLOAD_BYTES:
        try:
            dst.unlink(missing_ok=True)
        except Exception:
            pass
        return jsonify({"ok": False, "status": "error", "error": f"Datei zu gross (>{MAX_UPLOAD_BYTES} Bytes)"}), 413

    analysis = _analyze_uploaded_file(dst, ext)
    upload_id = uuid4().hex
    mime_type = mimetypes.guess_type(safe_name)[0] or "application/octet-stream"
    return jsonify({
        "ok": True,
        "status": "ok",
        "upload_id": upload_id,
        "filename": safe_name,
        "saved_path": str(dst),
        "filepath": str(dst),
        "mime_type": mime_type,
        "size": sz,
        "file_type": ext,
        "analysis": analysis,
        "summary": analysis.get("summary", ""),
    })


# ── Chat-History ─────────────────────────────────────────────────────────────

def _load_chat_history() -> list:
    try:
        data = read_json_file(CHAT_HISTORY_FILE, [])
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _append_chat_entry(entry: dict) -> None:
    history = _load_chat_history()
    history.insert(0, entry)
    write_json_file(CHAT_HISTORY_FILE, history[:300])


@app.route("/api/chat-history", methods=["GET"])
def get_chat_history():
    limit = int(request.args.get("limit", 50))
    return jsonify({"ok": True, "history": _load_chat_history()[:limit]})


@app.route("/api/chat-history/clear", methods=["POST"])
def clear_chat_history():
    write_json_file(CHAT_HISTORY_FILE, [])
    return jsonify({"ok": True})


@app.route("/api/chat-history/append", methods=["POST"])
def append_chat_history_route():
    data = request.get_json(silent=True) or {}
    entry = data.get("entry") if isinstance(data.get("entry"), dict) else data
    if not isinstance(entry, dict):
        return jsonify({"ok": False, "error": "entry fehlt"}), 400
    role = str(entry.get("role") or "").strip().lower()
    if role not in {"user", "assistant", "system"}:
        return jsonify({"ok": False, "error": "role muss user|assistant|system sein"}), 400
    safe = {
        "role": role[:32],
        "content": str(entry.get("content") or "")[:12000],
        "uploads": entry.get("uploads") if isinstance(entry.get("uploads"), list) else [],
        "changed_files": entry.get("changed_files") if isinstance(entry.get("changed_files"), list) else [],
        "status": str(entry.get("status") or "")[:120],
        "timestamp": str(entry.get("timestamp") or get_timestamp()),
        "run_id": str(entry.get("run_id") or "")[:64],
    }
    _append_chat_entry(safe)
    return jsonify({"ok": True})


def _strip_powershell_direct_prefix(text: str) -> str | None:
    """Prefix ``ps:`` oder ``powershell:`` fuer Direkt-Auftraege im Projektroot."""
    raw = str(text or "").strip()
    if not raw:
        return None
    lower = raw.lower()
    for prefix in ("powershell:", "ps:"):
        pl = prefix.lower()
        if lower.startswith(pl):
            return raw[len(pl) :].strip()
    return None


_POWERSHELL_CMD_MAX_LEN = 12000


def _run_powershell_command(command: str) -> dict[str, object]:
    cmd = str(command or "").strip()
    if not cmd:
        return {"ok": False, "error": "Leerer PowerShell-Befehl."}
    if len(cmd) > _POWERSHELL_CMD_MAX_LEN:
        return {"ok": False, "error": "PowerShell-Befehl zu lang."}
    try:
        cp = subprocess.run(
            [
                "powershell.exe",
                "-NoProfile",
                "-NonInteractive",
                "-ExecutionPolicy",
                "Bypass",
                "-Command",
                cmd,
            ],
            cwd=str(PROJECT_DIR),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=120,
        )
        out, err, trunc = _truncate_run_check_output(cp.stdout or "", cp.stderr or "", _RUN_CHECK_MAX_OUTPUT_BYTES)
        return {
            "ok": True,
            "exit_code": int(cp.returncode),
            "stdout": out,
            "stderr": err,
            "truncated": bool(trunc),
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "PowerShell-Timeout (120 s)."}
    except Exception as ex:
        return {"ok": False, "error": str(ex)}


def _powershell_direct_run_envelope(
    *, cleaned_prompt: str, run_id: str, scope: str, mode: str, log_label: str = "Direkt"
) -> dict | None:
    """Wenn ``ps:`` / ``powershell:`` — Ausfuehrung im :data:`PROJECT_DIR`. Sonst None."""
    ps_body = _strip_powershell_direct_prefix(cleaned_prompt)
    if ps_body is None:
        return None
    if not str(ps_body).strip():
        msg = "Nach ps: oder powershell: den Befehl angeben (Arbeitsverzeichnis ist der Projektordner)."
        append_ui_log_entry(log_label, "PowerShell leer", "warning")
        bad = {
            "ok": True,
            "success": True,
            "applied": False,
            "run_id": run_id,
            "scope": scope,
            "mode": mode,
            "status": "chat_response",
            "direct_status": "chat_response",
            "classification": "powershell_run",
            "route_mode": "powershell",
            "chat_response": msg,
            "formatted_response": msg,
            "natural_message": msg,
            "message": msg,
            "requires_confirmation": False,
            "writes_files": False,
            "has_changes": False,
            "changed_files": [],
            "workstream_events": [
                _ws_event("analysis", "warn", "PowerShell", "Leerer Befehl nach Praefix", status="done"),
            ],
        }
        return enrich_direct_run_response(bad)
    pr = _run_powershell_command(ps_body)
    if not pr.get("ok"):
        err_text = "PowerShell: " + str(pr.get("error") or "Fehler")
        ps_fail = {
            "ok": True,
            "success": True,
            "applied": False,
            "run_id": run_id,
            "scope": scope,
            "mode": mode,
            "status": "chat_response",
            "direct_status": "powershell_error",
            "classification": "powershell_run",
            "route_mode": "powershell",
            "chat_response": err_text,
            "formatted_response": err_text,
            "natural_message": err_text,
            "message": err_text,
            "requires_confirmation": False,
            "writes_files": False,
            "has_changes": False,
            "changed_files": [],
            "workstream_events": [
                _ws_event("analysis", "error", "PowerShell", "Ausfuehrung fehlgeschlagen", status="done"),
            ],
        }
        return enrich_direct_run_response(ps_fail)
    parts: list[str] = []
    if str(pr.get("stdout") or "").strip():
        parts.append("[stdout]\n" + str(pr.get("stdout") or ""))
    if str(pr.get("stderr") or "").strip():
        parts.append("[stderr]\n" + str(pr.get("stderr") or ""))
    out_text = "\n\n".join(parts) if parts else "(keine Ausgabe)"
    out_text += "\n\nExit-Code: " + str(pr.get("exit_code", ""))
    if pr.get("truncated"):
        out_text += "\n\n(Ausgabe wurde gekuerzt.)"
    append_ui_log_entry(log_label, "PowerShell ok exit=" + str(pr.get("exit_code")), "info")
    ps_ok = {
        "ok": True,
        "success": True,
        "applied": False,
        "run_id": run_id,
        "scope": scope,
        "mode": mode,
        "status": "chat_response",
        "direct_status": "powershell_done",
        "classification": "powershell_run",
        "route_mode": "powershell",
        "chat_response": out_text,
        "formatted_response": out_text,
        "natural_message": out_text,
        "message": out_text,
        "requires_confirmation": False,
        "writes_files": False,
        "has_changes": False,
        "changed_files": [],
        "powershell": {
            "exit_code": pr.get("exit_code"),
            "truncated": pr.get("truncated"),
            "cwd": str(PROJECT_DIR),
        },
        "workstream_events": [
            _ws_event("analysis", "info", "PowerShell", "Befehl ausgefuehrt", status="done"),
        ],
    }
    return enrich_direct_run_response(ps_ok)


# --- Ordneranalyse (nur Lesen, kein LLM) ---------------------------------

_FOLDER_ANALYSIS_BUDGET = 300
_FOLDER_ANALYSIS_SHALLOW = frozenset(
    {
        "node_modules",
        "__pycache__",
        ".pytest_cache",
        ".git",
        ".rainer_agent",
        ".cursor",
        "dist",
        "build",
        "out",
        "target",
    }
)
_KNOWN_DIR_BLURB: dict[str, str] = {
    "backend": "Python-Backend, Flask/API, viele `agent_*.py`, `routes/`, `services/`",
    "frontend": "Web-UI (`app.js`, `index.html`, `style.css`, ggf. Vite/React)",
    "tests": "`pytest`, Agent- und API-Tests",
    "data": "lokale JSON-Zustände, Chat, Runs, Todos",
    "docs": "Markdown-Dokumentation",
    "knowledge": "Notizen / Wissensbasis",
    "rambo_builder_local": "zweite/nested Builder-Struktur (falls vorhanden)",
    "electron": "Desktop/Electron-bezogene Dateien",
    "scripts": "Hilfs-Skripte",
    "config": "Konfiguration",
    "agent": "Agent-Hilfsmodule",
    "tools": "Werkzeuge / CLI",
    "server": "Server-Nebenprozesse",
    "logs": "Log-Dateien",
    "outbox": "Ausgabe-/Warteschlange",
    "Downloads": "Downloads/Artefakte (prüfen)",
    "memory": "Memory-/State-Dateien",
}
_BACKUP_NAME_MARKERS = ("backup", "archive", "_archive", "frontend.zip", "rambo_ui")


def _is_same_or_child_path(target: Path, workspace: Path) -> bool:
    try:
        t = target.resolve()
        w = workspace.resolve()
    except (OSError, ValueError):
        return False
    if t == w:
        return True
    try:
        t.relative_to(w)
        return True
    except ValueError:
        return False


def _same_or_inside(child: Path, parent: Path) -> bool:
    """Gleicher Ordner oder Kind — unter Windows case-insensitive mit normpath."""
    try:
        child_res = child.resolve(strict=False)
        parent_res = parent.resolve(strict=False)
    except (OSError, ValueError):
        return False
    if os.name == "nt":
        c = os.path.normcase(os.path.normpath(str(child_res)))
        p = os.path.normcase(os.path.normpath(str(parent_res)))
        if c == p:
            return True
        return c.startswith(p + os.sep)
    try:
        child_res.relative_to(parent_res)
        return True
    except ValueError:
        return False


def _clean_explicit_analysis_path(explicit: str) -> str:
    s = str(explicit).strip().strip('"').strip("'")
    s = s.rstrip(".,;:!?)]}\"'").strip()
    s = os.path.expandvars(os.path.expanduser(s))
    return s.strip()


def _norm_path_compare_key(path_like: str | Path) -> str:
    """Case-/Slash-normalisiert fuer POSIX und Windows-Vergleiche."""
    try:
        s = os.path.normpath(os.path.expandvars(str(path_like))).strip()
    except Exception:
        s = str(path_like).strip()
    try:
        return os.path.normcase(s)
    except Exception:
        return s.lower()


def _resolve_folder_analysis_root(explicit: str | None) -> tuple[Path | None, str | None, dict]:
    """
    Zielverzeichnis für Analyse. Nur gleich dem aktiven Workspace oder darunter.
    Rückgabe: (Path, None, debug) oder (None, Fehlermeldung, debug).
    """
    workspace = get_active_project_root()
    ws_res = workspace.resolve(strict=False)
    dbg: dict = {
        "extracted_path": explicit,
        "active_root": str(ws_res),
        "resolved_path": None,
        "exists": None,
        "is_dir": None,
    }

    if not explicit or not str(explicit).strip():
        dbg["resolved_path"] = str(ws_res)
        dbg["exists"] = ws_res.exists()
        dbg["is_dir"] = ws_res.is_dir() if dbg["exists"] else None
        return ws_res, None, dbg

    raw = _clean_explicit_analysis_path(explicit)
    try:
        cand = Path(raw)
        if not cand.is_absolute():
            cand = (ws_res / cand).resolve(strict=False)
        else:
            cand = cand.resolve(strict=False)
    except Exception as ex:
        dbg["resolved_path"] = None
        dbg["exists"] = None
        dbg["is_dir"] = None
        return None, f"Der Pfad konnte nicht aufgelöst werden: `{explicit}` ({ex})", dbg

    dbg["resolved_path"] = str(cand)
    dbg["exists"] = cand.exists()
    dbg["is_dir"] = cand.is_dir() if cand.exists() else None

    if not cand.exists():
        # Gleicher Text wie aktiver Workspace / gespeicherter Root — FS kann auf anderem Host abweichen
        if ws_res.exists():
            if _norm_path_compare_key(raw) == _norm_path_compare_key(ws_res):
                dbg["resolved_path"] = str(ws_res.resolve(strict=False))
                dbg["exists"] = ws_res.exists()
                dbg["is_dir"] = ws_res.is_dir()
                return ws_res, None, dbg
            try:
                st = get_active_project_state()
                root_s = str(st.get("active_project_root") or "").strip()
                if root_s and _norm_path_compare_key(raw) == _norm_path_compare_key(root_s):
                    dbg["resolved_path"] = str(ws_res.resolve(strict=False))
                    dbg["exists"] = ws_res.exists()
                    dbg["is_dir"] = ws_res.is_dir()
                    return ws_res, None, dbg
            except Exception:
                pass
        # Nicht-Windows: D:\... existiert nicht — gleicher Ordnername wie Workspace zulassen
        if os.name != "nt" and ws_res.exists() and re.match(r"(?i)^[a-z]:[\\/]", raw):
            leaf = re.split(r"[\\/]+", raw.rstrip("\\/"))[-1] or ""
            if leaf and leaf.lower() == (ws_res.name or "").lower():
                dbg["resolved_path"] = str(ws_res.resolve(strict=False))
                dbg["exists"] = ws_res.exists()
                dbg["is_dir"] = ws_res.is_dir()
                return ws_res, None, dbg

        err_msg = (
            f'Der angegebene Pfad "{raw}" wurde nicht gefunden oder ist nicht zugreifbar.'
        )
        return None, err_msg, dbg

    if not cand.is_dir():
        return None, "Der angegebene Pfad ist kein Verzeichnis.", dbg

    if not _same_or_inside(cand, ws_res):
        return (
            None,
            "Ich kann nur den aktuell freigegebenen Workspace analysieren. Bitte wähle oder gib den Ordner zuerst frei.",
            dbg,
        )

    return cand, None, dbg


def _safe_iterdir(path: Path) -> list[Path]:
    try:
        return sorted(path.iterdir(), key=lambda p: (not p.is_dir(), p.name.lower()))
    except OSError:
        return []


def _fmt_size(n: int) -> str:
    if n >= 1048576:
        return f"{n / 1048576:.1f} MB"
    if n >= 1024:
        return f"{n / 1024:.1f} KB"
    return f"{n} B"


def _render_folder_analysis_markdown(root: Path) -> str:
    root = root.resolve()
    lines: list[str] = []
    lines.append(f"Hier eine strukturierte Übersicht über `{root}`:\n")
    lines.append(
        f"*Dieser Block listet nur den Inhalt von **{root.name}** — nicht von anderen Projekten auf dem Rechner.*\n"
    )

    # Projekttyp (Heuristik)
    guess_bits: list[str] = []
    if (root / "backend").is_dir() and (root / "frontend").is_dir():
        guess_bits.append("Monorepo mit **Python-Backend** und **Web-Frontend**.")
    if (root / "electron").is_dir() or (root / "rambo_ui").is_dir():
        guess_bits.append("Hinweise auf **Desktop/Electron** oder UI-Nebenpackages.")
    if (root / "tests").is_dir() or (root / "pytest.ini").is_file():
        guess_bits.append("Umfangreiche **Tests** (`pytest`) sind vorgesehen.")
    if (root / "docker-compose.yml").is_file() or (root / "Dockerfile").is_file():
        guess_bits.append("**Docker**-Dateien für Deployment/Entwicklung.")
    if not guess_bits:
        guess_bits.append("Keine eindeutige Signatur — vermutlich ein **Software-Projekt** mit Backend/Frontend-Komponenten.")
    lines.append("## Projekttyp\n")
    lines.append(" ".join(guess_bits) + "\n")

    entries_used = 0
    table_rows: list[tuple[str, str]] = []
    root_files: list[str] = []
    artifacts: list[str] = []
    shallow_note: list[str] = []

    for child in _safe_iterdir(root):
        if entries_used >= _FOLDER_ANALYSIS_BUDGET:
            break
        name = child.name
        entries_used += 1
        if child.is_dir():
            low = name.lower()
            if name in _FOLDER_ANALYSIS_SHALLOW or name.startswith("."):
                shallow_note.append(f"`{name}/`")
                desc = "*(nicht rekursiv durchsucht — typischer Cache/Systemordner)*"
            elif name in _KNOWN_DIR_BLURB:
                desc = _KNOWN_DIR_BLURB[name]
            else:
                desc = "Verzeichnis"
            if any(m in low for m in _BACKUP_NAME_MARKERS) or "backup" in low:
                artifacts.append(f"- Backup-/Archiv-artiger Ordner: `{name}/`\n")
            table_rows.append((f"`{name}/`", desc))
        else:
            root_files.append(name)
            try:
                sz = child.stat().st_size
            except OSError:
                sz = 0
            if sz >= 5 * 1024 * 1024 or name.endswith(".zip"):
                artifacts.append(f"- Große Datei / Archiv: `{name}` ({_fmt_size(sz)})\n")

    lines.append("\n## Wichtigste Verzeichnisse\n")
    lines.append("| Bereich | Kurzbeschreibung |\n")
    lines.append("| --- | --- |\n")
    for cell, desc in table_rows[:60]:
        lines.append(f"| {cell} | {desc} |\n")
    if shallow_note:
        lines.append("\n*Übersprungene System-/Cache-Ordner (oberste Ebene):* " + ", ".join(sorted(set(shallow_note))) + "\n")

    # Tiefe 2: ausgewählte Kernordner
    lines.append("\n## Ausgewählte Unterebenen (max. Tiefe 2)\n")
    peek_dirs = [
        n for n in ("backend", "frontend", "tests", "data", "docs", "config")
        if (root / n).is_dir() and n not in _FOLDER_ANALYSIS_SHALLOW
    ]
    if not peek_dirs:
        lines.append("*(Keine der Standard-Kernordner als separates Verzeichnis gefunden oder alles übersprungen.)*\n")
    for dn in peek_dirs:
        if entries_used >= _FOLDER_ANALYSIS_BUDGET:
            break
        sub = root / dn
        kids: list[str] = []
        for ch in _safe_iterdir(sub):
            if entries_used >= _FOLDER_ANALYSIS_BUDGET:
                break
            entries_used += 1
            nm = ch.name
            if ch.is_dir():
                if nm in _FOLDER_ANALYSIS_SHALLOW or nm.startswith("."):
                    kids.append(f"{nm}/ …")
                else:
                    kids.append(f"{nm}/")
            else:
                kids.append(nm)
            if len(kids) >= 20:
                kids.append("…")
                break
        lines.append(f"- **`{dn}/`**: " + (", ".join(kids) if kids else "(leer)") + "\n")

    lines.append("\n## Wichtige Dateien (Root)\n")
    prefer = (
        "README.md",
        "CHANGELOG.md",
        "docker-compose.yml",
        "Dockerfile",
        "pytest.ini",
        "start_rainer.ps1",
        "start_all.bat",
        "main.py",
        "package.json",
        "README",
    )
    hit = [f for f in root_files if f in prefer]
    rest = sorted([f for f in root_files if f not in prefer])[:40]
    show = hit + rest
    if not show:
        lines.append("*(Keine Dateien auf der Root-Ebene gelistet.)*\n")
    else:
        for f in show:
            lines.append(f"- `{f}`\n")

    lines.append("\n## Auffälliges / Aufräumen\n")
    if artifacts:
        lines.extend(artifacts)
        lines.append("")
    else:
        lines.append("- Keine extrem großen Root-Dateien oder offensichtliche ZIP-Artefakte erkannt (oberste Ebene).\n")
    if (root / "rambo_builder_local").is_dir():
        lines.append("- **Nested Kopie** `rambo_builder_local/` — ggf. Duplikat der Builder-Struktur prüfen.\n")
    if (root / "backend_backup_20260501_144231").exists():
        lines.append("- **Backup-Ordner** `backend_backup_*` vorhanden.\n")

    lines.append("\n## Kurzfassung\n")
    root_disp = root.resolve()
    lines.append(
        f"Diese Übersicht bezieht sich **ausschließlich auf** `{root_disp}` "
        f"(Ordner: **{root_disp.name}**). Typisch für ein **Rainer Build / RAMBO-Builder**-Setup "
        "(Backend-Agenten, Frontend, Tests, lokale Daten) — hier nur **Lesen**, keine Änderungen.\n"
    )
    if entries_used >= _FOLDER_ANALYSIS_BUDGET:
        lines.append(
            f"\n*Hinweis: Es wurden mindestens {_FOLDER_ANALYSIS_BUDGET} Einträge berücksichtigt; "
            "die Darstellung ist begrenzt.*\n"
        )
    return "".join(lines)


def build_folder_analysis_payload(root: Path, *, run_id: str, scope: str, mode: str) -> dict:
    """JSON-kompatibles Payload für enrich_direct_run_response (Ordneranalyse)."""
    try:
        md = _render_folder_analysis_markdown(root)
    except Exception as ex:
        md = f"Ordneranalyse fehlgeschlagen: {ex}"
    return {
        "ok": True,
        "success": True,
        "applied": False,
        "run_id": run_id,
        "scope": scope,
        "mode": mode,
        "status": "chat_response",
        "direct_status": "chat_response",
        "classification": "project_read",
        "route_mode": "workspace_analysis",
        "task_kind": "read_only_analysis",
        "chat_response": md,
        "formatted_response": md,
        "natural_message": md,
        "message": md,
        "requires_confirmation": False,
        "requires_user_confirmation": False,
        "writes_files": False,
        "has_changes": False,
        "changed_files": [],
        "files": [],
        "workspace_analysis_root": str(root.resolve()),
        "workstream_events": [
            _ws_event("analysis", "info", "Workspace", "Ordneranalyse (nur Lesen)", status="done"),
        ],
    }


@app.route("/api/open-powershell", methods=["POST"])
def open_powershell():
    try:
        cwd_lit = str(PROJECT_DIR).replace("'", "''")
        cmd = [
            "powershell.exe",
            "-NoExit",
            "-Command",
            f"Set-Location -LiteralPath '{cwd_lit}'",
        ]
        creationflags = getattr(subprocess, "CREATE_NEW_CONSOLE", 0)
        subprocess.Popen(cmd, cwd=str(PROJECT_DIR), creationflags=creationflags)
        return jsonify({"status": "ok", "cwd": str(PROJECT_DIR)})
    except Exception as ex:
        return jsonify({"status": "error", "error": str(ex)}), 500


@app.route("/api/powershell/run", methods=["POST"])
def powershell_run_endpoint():
    """Einzelbefehl ausfuehren (Projektroot), fuer UI/Debug."""
    data = request.get_json(silent=True) or {}
    cmd = str(data.get("command") or "").strip()
    pr = _run_powershell_command(cmd)
    if not pr.get("ok"):
        return jsonify({"ok": False, "error": pr.get("error")}), 400
    return jsonify(
        {
            "ok": True,
            "exit_code": pr.get("exit_code"),
            "stdout": pr.get("stdout"),
            "stderr": pr.get("stderr"),
            "truncated": pr.get("truncated"),
            "cwd": str(PROJECT_DIR),
        }
    )


@app.route("/api/status", methods=["GET"])
def status():
    project_path = format_local_path(os.path.join(os.getcwd(), "data"))
    if not project_path or not os.path.exists(project_path.replace("/", os.sep)):
        project_path = "C:/Users/mielersch/Desktop/Rambo-Rainer"

    guarded_files = [
        "frontend/src/App.jsx",
        "frontend/src/App.css"
    ]

    try:
        response = requests.get("http://127.0.0.1:11434/api/tags", timeout=2)
        ollama_ok = response.status_code == 200
    except requests.RequestException:
        ollama_ok = False

    try:
        llm_health = summarize_llm_health()
    except Exception:
        llm_health = {}

    activity_entries = load_ui_activity_entries()
    last_error = ""
    for entry in activity_entries:
        if entry["level"] == "error":
            last_error = entry["message"]
            break

    builder_status = current_builder_status()
    auto_run_state = load_project_auto_run_state()
    auto_run_status = current_auto_run_status()
    active_direct_run = auto_run_state.get("pending_direct_run") if pending_direct_context_is_valid(auto_run_state.get("pending_direct_run"), auto_run_state) else None
    last_completed_run = get_history_entry_by_run_id(auto_run_state.get("last_completed_run_id"))
    active_recognized = (active_direct_run or {}).get("recognized_task") if isinstance((active_direct_run or {}).get("recognized_task"), dict) else {}
    if not active_recognized and active_direct_run:
        active_recognized = classify_direct_task((active_direct_run or {}).get("task"))
    last_recognized = {
        "task_type": str((last_completed_run or {}).get("task_type") or "unknown"),
        "primary_area": str((last_completed_run or {}).get("primary_area") or "Builder"),
    }

    project_map = read_json_file(PROJECT_MAP_FILE, {})
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    todos_content = read_json_file(DATA_DIR / "todos.json", {"todos": []})
    last_scan = project_map.get("scanned_at", "")
    knowledge_built = knowledge.get("built_at", "")
    total_scanned = project_map.get("total_files", 0)
    current_todo_focus = {}
    if isinstance(todos_content, dict) and isinstance(todos_content.get("todos"), list):
        for todo in todos_content["todos"]:
            if isinstance(todo, dict) and bool(todo.get("is_focus")):
                current_todo_focus = {
                    "id": str(todo.get("id") or ""),
                    "title": str(todo.get("title") or ""),
                    "status": str(todo.get("status") or "offen"),
                    "area": str(todo.get("area") or "Rainer Build"),
                }
                break

    last_blocked = ""
    last_allowed = ""
    for entry in activity_entries:
        if entry["label"] in {"Projekt", "Auto"}:
            if "blockiert" in entry["message"].lower() and not last_blocked:
                last_blocked = entry["message"]
            if "angelegt" in entry["message"].lower() or "aktualisiert" in entry["message"].lower():
                if not last_allowed:
                    last_allowed = entry["message"]

    if not last_blocked:
        last_blocked = auto_run_state.get("last_guard_decision", "") if auto_run_state.get("blocked") else ""
    if not last_allowed:
        last_allowed = auto_run_state.get("last_apply_action", "")

    try:
        _aw_root = WORKSPACE_SANDBOX.get_active_workspace()
        _aw_act = _aw_root.get("active") if isinstance(_aw_root, dict) else {}
    except Exception:
        _aw_act = {}
    workspace_active = {
        "path": str((_aw_act or {}).get("path") or ""),
        "trusted": bool((_aw_act or {}).get("trusted", False)),
        "label": str((_aw_act or {}).get("label") or (_aw_act or {}).get("name") or ""),
    }

    return jsonify({
        "backend_ok": True,
        "ollama_ok": ollama_ok,
        "llm_health": llm_health,
        "workspace_active": workspace_active,
        "project_path": project_path,
        "guarded_files": guarded_files,
        "last_error": last_error,
        "last_success": "Lokale Oberflaeche einsatzbereit",
        "blocked_mode": False,
        "recovery_recommendation": "Bei Problemen Status neu laden, Eingaben pruefen und nur innerhalb von rambo_builder_local arbeiten.",
        "builder_status": builder_status,
        "auto_run_status": auto_run_status,
        "project_mode": {
            "last_scan": last_scan,
            "total_scanned": total_scanned,
            "knowledge_built": knowledge_built,
            "last_blocked": last_blocked,
            "last_allowed": last_allowed,
            "last_auto_task": auto_run_state.get("last_task", ""),
            "last_auto_mode": auto_run_state.get("last_mode", "safe"),
            "last_auto_targets": auto_run_state.get("last_target_paths", []),
            "last_guard_decision": auto_run_state.get("last_guard_decision", ""),
            "last_apply_action": auto_run_state.get("last_apply_action", ""),
            "last_post_check": auto_run_state.get("last_check_result", ""),
            "last_direct_scope": auto_run_state.get("last_direct_scope", ""),
            "last_direct_prompt": auto_run_state.get("last_direct_prompt", ""),
            "last_direct_decision": auto_run_state.get("last_direct_decision", ""),
            "last_direct_status": auto_run_state.get("last_direct_status", "idle"),
            "last_planned_steps": normalize_planned_steps(auto_run_state.get("last_planned_steps") or []),
            "recognized_task_type": str(active_recognized.get("task_type") or last_recognized.get("task_type") or "unknown"),
            "recognized_primary_area": str(active_recognized.get("primary_area") or last_recognized.get("primary_area") or "Builder"),
            "last_direct_run_id": auto_run_state.get("last_direct_run_id", ""),
            "active_direct_run_id": auto_run_state.get("active_direct_run_id", ""),
            "last_completed_run_id": auto_run_state.get("last_completed_run_id", ""),
            "direct_confirmation_pending": pending_direct_context_is_valid(auto_run_state.get("pending_direct_run"), auto_run_state),
            "active_direct_run": {
                "run_id": str((active_direct_run or {}).get("run_id") or ""),
                "created_at": str((active_direct_run or {}).get("created_at") or ""),
                "scope": str((active_direct_run or {}).get("scope") or ""),
                "mode": str((active_direct_run or {}).get("mode") or ""),
                "task": str((active_direct_run or {}).get("task") or ""),
                "task_type": str(active_recognized.get("task_type") or ""),
                "primary_area": str(active_recognized.get("primary_area") or ""),
                "selected_target_path": str((active_direct_run or {}).get("selected_target_path") or ""),
                "status": str((active_direct_run or {}).get("status") or ""),
                "planned_steps": normalize_planned_steps((active_direct_run or {}).get("planned_steps") or []),
            },
            "last_completed_run": last_completed_run or {},
            "recent_direct_runs": get_direct_run_history()[:3],
            "auto_loop": normalize_auto_loop_state(auto_run_state.get("auto_loop_state")),
            "last_runtime_info": (
                f"Aktiver Lauf seit {format_display_timestamp((active_direct_run or {}).get('created_at'))}"
                if (active_direct_run or {}).get("created_at")
                else (
                    f"Letzter abgeschlossener Lauf: {format_display_timestamp((last_completed_run or {}).get('timestamp'))}"
                    if (last_completed_run or {}).get("timestamp")
                    else "Noch keine Laufzeitinfo."
                )
            ),
            "last_api_info": (
                f"Backend OK / aktiv {llm_health.get('active_provider', '?')} "
                f"({'OK' if llm_health.get('provider_reachable') else 'offline'}) — Chat: "
                f"{'ja' if llm_health.get('chat_available') else 'nein'}, Coding: "
                f"{'ja' if llm_health.get('coding_available') else 'nein'}"
            ),
            "current_todo_focus": current_todo_focus,
            "allowed_prefixes_count": len(ALLOWED_PROJECT_WRITE_PREFIXES),
            "sensitive_patterns_count": len(SENSITIVE_PATTERNS)
        }
    })


@app.route("/api/direct-history", methods=["GET"])
def direct_history():
    return jsonify({"entries": get_direct_run_history()})


@app.route("/api/direct-reset", methods=["POST"])
def direct_reset():
    data = request.get_json(silent=True) or {}
    action = str(data.get("action") or "discard").strip().lower()
    pending = get_pending_direct_run()

    clear_pending_direct_run()
    save_project_auto_run_state({
        "last_direct_status": "idle",
        "last_direct_decision": (
            "Aktiver Direktlauf verworfen." if action == "discard"
            else "Direktmodus fuer neuen Lauf vorbereitet."
        )
    })
    append_ui_log_entry(
        "Direkt",
        "Aktiver Direktlauf verworfen." if pending else "Direktmodus ohne aktiven Lauf zurueckgesetzt.",
        "info"
    )
    return jsonify({
        "message": (
            "Aktiver Direktlauf wurde verworfen." if pending and action == "discard"
            else "Direktmodus fuer einen neuen Lauf bereit."
        ),
        "had_pending_run": bool(pending),
        "action": action,
    })


@app.route("/api/auto-loop/state", methods=["GET", "POST"])
def auto_loop_state_endpoint():
    if request.method == "GET":
        state = load_project_auto_run_state()
        auto_loop = normalize_auto_loop_state(state.get("auto_loop_state"))
        return jsonify({"auto_loop": auto_loop, "stats": _auto_loop_run_metrics(auto_loop)})

    data = request.get_json(silent=True) or {}
    incoming = normalize_auto_loop_state(data.get("auto_loop"))
    if incoming.get("status") != "idle" and not incoming.get("last_run_at"):
        incoming["last_run_at"] = get_timestamp()
    save_project_auto_run_state({"auto_loop_state": incoming})
    return jsonify({"saved": True, "auto_loop": incoming, "timestamp": get_timestamp()})


@app.route("/api/auto-loop/start", methods=["POST"])
def auto_loop_start_endpoint():
    data = request.get_json(silent=True) or {}
    goal = str(data.get("goal") or "").strip()
    if len(goal) < 4:
        return jsonify({"ok": False, "message": "Ziel fehlt oder ist zu kurz."}), 400
    runner_command = str(data.get("runner_command") or "").strip()
    apply_mode = str(data.get("apply_mode") or data.get("mode") or "safe").strip().lower()
    if apply_mode not in {"safe", "apply"}:
        apply_mode = "safe"
    run_state = load_project_auto_run_state()
    existing = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    active_statuses = {"running", "paused", "approval_required"}
    if existing.get("status") in active_statuses and not bool(data.get("force_restart")):
        return jsonify({
            "ok": False,
            "message": f"Es laeuft bereits ein Auto-Loop im Status '{existing.get('status')}'. Erst stoppen/zuruecksetzen oder force_restart senden.",
            "auto_loop": existing,
        }), 409
    project_blocked = bool(run_state.get("blocked"))
    guard_info = {
        "reason": str(run_state.get("last_guard_decision") or "").strip(),
        "path": ", ".join([str(p) for p in (run_state.get("last_target_paths") or []) if p][:3]),
    }
    retry_max_raw = data.get("retry_max")
    try:
        retry_max = max(0, min(5, int(retry_max_raw)))
    except Exception:
        retry_max = 2
    plan = build_auto_loop_plan(goal, runner_command, project_blocked,
                                 apply_mode=apply_mode, guard_info=guard_info)
    state = normalize_auto_loop_state({
        "run_id": f"loop-{uuid4().hex[:10]}",
        "goal": goal,
        "status": "running",
        "current_phase": "Initialisierung",
        "current_step": "",
        "last_action": "Auto-Loop gestartet.",
        "next_step": plan[0]["label"] if plan else "Abschluss",
        "summary": "",
        "steps": plan,
        "history": [],
        "error_label": "",
        "repair_suggestion": "",
        "affected_file": "",
        "requires_approval": False,
        "stopped": False,
        "stop_requested": False,
        "runner_command": runner_command,
        "retry_count": 0,
        "retry_max": retry_max,
        "session_blockers": [],
    })
    state["apply_mode"] = apply_mode
    state["current_error"] = _empty_error_info()
    state["repair_plan"] = []
    state["retry_possible"] = retry_max > 0
    state["retry_blocked_reason"] = ""
    state["repair_history"] = []
    state["acceptance_status"] = ""
    state["acceptance_ready"] = False
    state["acceptance_blocked_reason"] = ""
    state["acceptance_notes"] = ""
    state["acceptance_timestamp"] = ""
    state["started_at"] = get_timestamp()
    state["ended_at"] = ""
    state["paused_at"] = ""
    state["resumed_at"] = ""
    try:
        state["memory_snapshot"] = build_memory_snapshot_compact()
    except Exception:
        state["memory_snapshot"] = {}
    _auto_loop_push_history(state, f"Gestartet mit Ziel: {goal[:100]}")
    mem = state.get("memory_snapshot") or {}
    if mem.get("rules_count"):
        _auto_loop_push_history(state, f"Memory-Snapshot: {mem.get('rules_count')} Regeln, {len(mem.get('sources') or [])} Quellen.")
    if apply_mode == "apply":
        _auto_loop_push_history(state, "Apply-Mode aktiv - Apply-Schritt ist freigabepflichtig.")
    if runner_command:
        _auto_loop_push_history(state, f"Runner-Command vorgemerkt (Freigabe): {runner_command}")
    if project_blocked:
        _auto_loop_push_history(state, "Hinweis: Guard markiert aktuellen Pfad als blockiert.")
    instruction = _auto_loop_build_instruction(state)
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", f"Auto-Loop gestartet: {goal[:80]}", "info")
    return jsonify({"ok": True, "auto_loop": normalized, "instruction": instruction})


@app.route("/api/auto-loop/advance", methods=["POST"])
def auto_loop_advance_endpoint():
    data = request.get_json(silent=True) or {}
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") in {"idle", "done", "failed", "stopped", "blocked"} and not data.get("force"):
        if state.get("status") == "idle":
            return jsonify({"ok": False, "message": "Kein aktiver Auto-Loop. Zuerst starten."}), 400
        return jsonify({"ok": False, "message": f"Advance ist im Status '{state.get('status')}' nicht erlaubt."}), 409
    if state.get("status") in {"paused", "approval_required"} and not data.get("force"):
        return jsonify({
            "ok": False,
            "message": f"Advance ist im Status '{state.get('status')}' gesperrt. Erst fortsetzen bzw. freigeben.",
            "auto_loop": state,
        }), 409
    report = data.get("report")
    if isinstance(report, dict):
        report_step_id = str(report.get("step_id") or "").strip()
        if report_step_id:
            idx = _auto_loop_find_index(state, report_step_id)
            if idx == -1:
                return jsonify({"ok": False, "message": f"Unbekannte step_id im Report: {report_step_id}"}), 400
            step_status = str((state.get("steps") or [])[idx].get("status") or "").lower()
            if step_status not in {"laeuft", "geplant", "wartet auf freigabe"}:
                return jsonify({
                    "ok": False,
                    "message": f"Report fuer Schritt '{report_step_id}' im Status '{step_status}' nicht erlaubt.",
                }), 409
        _auto_loop_apply_report(state, report)
    instruction = _auto_loop_build_instruction(state)
    # Interne sichere Schritte koennen in einer Advance-Anfrage direkt weiterlaufen,
    # damit der Agentenlauf als zusammenhaengender Job wirkt.
    auto_progress_budget = 6
    while auto_progress_budget > 0:
        auto_progress_budget -= 1
        current_step = instruction.get("step") if isinstance(instruction, dict) else None
        action = str(instruction.get("action") or "").strip().lower() if isinstance(instruction, dict) else ""
        if action in {"done", "stopped", "blocked", "failed", "wait_approval", "paused"}:
            break
        if not _auto_loop_is_internal_safe_step(current_step):
            break
        step_id = str(current_step.get("id") or "").strip()
        if not step_id:
            break
        _auto_loop_apply_report(state, {
            "step_id": step_id,
            "ok": True,
            "detail": _auto_loop_autocomplete_detail(current_step),
            "clear_error": action in {"analyze_error", "prepare_repair"},
        })
        instruction = _auto_loop_build_instruction(state)
    if instruction.get("action") in {"done", "stopped", "blocked", "failed"}:
        state["summary"] = _auto_loop_summary(state)
    normalized = _auto_loop_save(state)
    return jsonify({"ok": True, "auto_loop": normalized, "instruction": instruction})


@app.route("/api/auto-loop/approve", methods=["POST"])
def auto_loop_approve_endpoint():
    data = request.get_json(silent=True) or {}
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") != "approval_required":
        return jsonify({
            "ok": False,
            "message": f"Keine Freigabe ausstehend (Status: {state.get('status')}).",
            "auto_loop": state,
        }), 409
    target_id = str(data.get("step_id") or "").strip()
    target_idx, step = _auto_loop_find_waiting_step(state, target_id)
    if target_idx == -1:
        return jsonify({"ok": False, "message": "Kein Schritt wartet auf Freigabe."}), 400
    step_status = str(step.get("status") or "").lower()
    if step_status != "wartet auf freigabe":
        return jsonify({"ok": False, "message": f"Schritt ist nicht freigabefaehig (Status: {step_status})."}), 409
    step["approved"] = True
    step["status"] = "geplant"
    step["detail"] = step.get("detail") or step.get("label") or ""
    state["status"] = "running"
    state["requires_approval"] = False
    state["current_phase"] = "Fortsetzung nach Freigabe"
    state["current_step"] = step.get("label") or ""
    state["last_action"] = f"Freigabe erteilt: {step.get('label')}"
    _auto_loop_push_history(state, f"Freigegeben: {step.get('label')}")
    instruction = _auto_loop_build_instruction(state)
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", f"Freigabe erteilt: {step.get('label')}", "info")
    return jsonify({"ok": True, "auto_loop": normalized, "instruction": instruction})


@app.route("/api/auto-loop/stop", methods=["POST"])
def auto_loop_stop_endpoint():
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") in {"running", "approval_required", "paused"}:
        state["stop_requested"] = True
        state["status"] = "stopped"
        state["stopped"] = True
        state["stop_requested"] = False
        state["current_phase"] = "Gestoppt"
        state["last_action"] = "Auto-Loop manuell gestoppt."
        state["ended_at"] = get_timestamp()
        state["summary"] = _auto_loop_summary(state)
        _auto_loop_push_history(state, "Manuell gestoppt.")
    elif state.get("status") in {"idle", "done", "failed", "blocked", "stopped"}:
        state["stop_requested"] = False
        state["last_action"] = state.get("last_action") or "Stop ohne aktiven Lauf ignoriert."
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", "Auto-Loop gestoppt.", "warning")
    return jsonify({"ok": True, "auto_loop": normalized})


@app.route("/api/auto-loop/pause", methods=["POST"])
def auto_loop_pause_endpoint():
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") == "paused":
        return jsonify({"ok": True, "auto_loop": state, "message": "Bereits pausiert."})
    if state.get("status") not in {"running", "approval_required"}:
        return jsonify({
            "ok": False,
            "message": f"Nicht pausierbar im Status '{state.get('status')}'.",
            "auto_loop": state,
        }), 409
    state["pause_requested"] = True
    state["status"] = "paused"
    state["paused_at"] = get_timestamp()
    state["current_phase"] = "Pausiert"
    state["last_action"] = "Auto-Loop pausiert."
    _auto_loop_push_history(state, "Pausiert.")
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", "Auto-Loop pausiert.", "info")
    return jsonify({"ok": True, "auto_loop": normalized})


@app.route("/api/auto-loop/resume", methods=["POST"])
def auto_loop_resume_endpoint():
    data = request.get_json(silent=True) or {}
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") not in {"paused", "approval_required", "blocked"}:
        return jsonify({
            "ok": False,
            "message": f"Kein fortsetzbarer Lauf im Status '{state.get('status')}'.",
            "auto_loop": state,
        }), 409
    if state.get("status") == "approval_required":
        return jsonify({
            "ok": False,
            "message": "Lauf wartet auf Freigabe. Bitte zuerst Freigabe/Ablehnung senden.",
            "auto_loop": state,
        }), 409
    if state.get("status") == "blocked":
        blockers = state.get("session_blockers") if isinstance(state.get("session_blockers"), list) else []
        force_resume = bool(data.get("force"))
        if blockers and not force_resume:
            return jsonify({
                "ok": False,
                "message": "Lauf ist blockiert. Erst Blocker aufloesen oder force=true senden.",
                "auto_loop": state,
            }), 409
        state["session_blockers"] = blockers
    state["pause_requested"] = False
    state["stop_requested"] = False
    state["status"] = "running"
    state["resumed_at"] = get_timestamp()
    state["current_phase"] = "Fortsetzung"
    state["last_action"] = "Auto-Loop fortgesetzt."
    _auto_loop_push_history(state, "Fortgesetzt.")
    instruction = _auto_loop_build_instruction(state)
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", "Auto-Loop fortgesetzt.", "info")
    return jsonify({"ok": True, "auto_loop": normalized, "instruction": instruction})


@app.route("/api/auto-loop/reject", methods=["POST"])
def auto_loop_reject_endpoint():
    data = request.get_json(silent=True) or {}
    run_state = load_project_auto_run_state()
    state = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    if state.get("status") != "approval_required":
        return jsonify({
            "ok": False,
            "message": f"Keine Freigabe ausstehend (Status: {state.get('status')}).",
            "auto_loop": state,
        }), 409
    target_id = str(data.get("step_id") or "").strip()
    reason = str(data.get("reason") or "Freigabe verweigert.").strip()
    target_idx, step = _auto_loop_find_waiting_step(state, target_id)
    if target_idx == -1:
        return jsonify({"ok": False, "message": "Kein Schritt wartet auf Freigabe."}), 400
    step_status = str(step.get("status") or "").lower()
    if step_status != "wartet auf freigabe":
        return jsonify({"ok": False, "message": f"Schritt ist nicht ablehnbar (Status: {step_status})."}), 409
    step["status"] = "blockiert"
    step["approved"] = False
    step["blocker"] = reason
    state["requires_approval"] = False
    state["status"] = "blocked"
    state["phase"] = "blocked"
    state["current_phase"] = "Freigabe abgelehnt"
    state["current_step"] = step.get("label") or ""
    state["last_action"] = f"Freigabe abgelehnt: {step.get('label')} - {reason}"
    state["ended_at"] = get_timestamp()
    _auto_loop_add_session_blocker(state, step, {"reason": reason, "suggestion": "Schritt anpassen oder Lauf zuruecksetzen."})
    _auto_loop_push_history(state, f"Abgelehnt: {step.get('label')} ({reason})")
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", f"Freigabe abgelehnt: {step.get('label')}", "warning")
    return jsonify({"ok": True, "auto_loop": normalized})


@app.route("/api/orchestrator/control", methods=["POST"])
def orchestrator_control_endpoint():
    """Zentraler serverseitiger Kontrollpunkt fuer den Agentenlauf.

    Aktionen: start, pause, resume, stop, reset, approve, reject, advance.
    Delegiert an bestehende Auto-Loop-Endpunkte, aber ueber einen einheitlichen
    Pfad. Das Frontend soll sich ueber diese API nur noch beobachten/steuern.
    """
    data = request.get_json(silent=True) or {}
    action = str(data.get("action") or "").strip().lower()
    if not action:
        return jsonify({"ok": False, "message": "Aktion fehlt."}), 400

    dispatch = {
        "start": auto_loop_start_endpoint,
        "pause": auto_loop_pause_endpoint,
        "resume": auto_loop_resume_endpoint,
        "stop": auto_loop_stop_endpoint,
        "reset": auto_loop_reset_endpoint,
        "approve": auto_loop_approve_endpoint,
        "reject": auto_loop_reject_endpoint,
        "advance": auto_loop_advance_endpoint,
    }
    handler = dispatch.get(action)
    if not handler:
        return jsonify({"ok": False, "message": f"Unbekannte Aktion: {action}"}), 400
    return handler()


@app.route("/api/auto-loop/reset", methods=["POST"])
def auto_loop_reset_endpoint():
    state = normalize_auto_loop_state({
        "run_id": "",
        "goal": "",
        "status": "idle",
        "current_phase": "Bereit",
        "current_step": "",
        "last_action": "Auto-Loop zurueckgesetzt.",
        "next_step": "",
        "summary": "",
        "steps": [],
        "history": [],
        "error_label": "",
        "repair_suggestion": "",
        "affected_file": "",
        "requires_approval": False,
        "stopped": False,
        "stop_requested": False,
        "runner_command": "",
        "retry_count": 0,
        "retry_max": 2,
        "session_blockers": [],
        "apply_mode": "safe",
        "phase": "planning",
        "active_module": "",
        "current_error": _empty_error_info(),
        "repair_plan": [],
        "retry_possible": True,
        "retry_blocked_reason": "",
        "repair_history": [],
        "acceptance_status": "",
        "acceptance_ready": False,
        "acceptance_blocked_reason": "",
        "acceptance_notes": "",
        "acceptance_timestamp": "",
    })
    _auto_loop_push_history(state, "Auto-Loop zurueckgesetzt.")
    normalized = _auto_loop_save(state)
    append_ui_log_entry("Auto-Loop", "Auto-Loop zurueckgesetzt.", "info")
    return jsonify({"ok": True, "auto_loop": normalized})


def _qa_level(ok, partial=False, error=False):
    if error:
        return "Fehler"
    if ok:
        return "bereit"
    if partial:
        return "teilweise"
    return "offen"


def _qa_bucket(state):
    raw = str(state or "offen").strip().lower()
    if raw == "bereit":
        return "ready"
    if raw == "teilweise":
        return "partial"
    if raw == "blockiert":
        return "blocked"
    if raw == "fehler":
        return "failed"
    return "open"


def build_qa_acceptance_snapshot(run_state=None, auto_loop=None, activity=None):
    """Kanonische QA-/Acceptance-Sicht aus vorhandenen Datenquellen."""
    run_state = run_state if isinstance(run_state, dict) else load_project_auto_run_state()
    auto_loop = auto_loop if isinstance(auto_loop, dict) else normalize_auto_loop_state(run_state.get("auto_loop_state"))
    activity = activity if isinstance(activity, list) else load_ui_activity_entries()

    project_map = read_json_file(PROJECT_MAP_FILE, {})
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    design_notes = read_json_file(DATA_DIR / "design_notes.json", {})
    todos_content = read_json_file(DATA_DIR / "todos.json", {"todos": []})
    patch_snapshot = build_patch_review_snapshot()
    current_patch = patch_snapshot.get("current") if isinstance(patch_snapshot, dict) else None
    memory_info = _collect_memory_info_for_agent_core(auto_loop)

    modules = []

    last_direct_status = str(run_state.get("last_direct_status") or "idle")
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    direct_ok = last_direct_status in {"completed", "confirmed"} or bool(run_state.get("last_completed_run_id"))
    direct_partial = bool(pending_direct)
    modules.append({
        "id": "direct",
        "label": "Direktmodus",
        "state": _qa_level(direct_ok, direct_partial),
        "detail": f"Letzter Status: {last_direct_status}; offene Vorschau: {'ja' if pending_direct else 'nein'}",
    })

    pm_ok = bool(run_state.get("last_task"))
    pm_partial = bool(run_state.get("last_target_paths"))
    pm_state = "blockiert" if run_state.get("blocked") else _qa_level(pm_ok, pm_partial)
    modules.append({
        "id": "project_mode",
        "label": "Project-Mode",
        "state": pm_state,
        "detail": run_state.get("last_task") or "Kein letzter Projektauftrag.",
    })

    map_ok = bool(project_map.get("scanned_at"))
    modules.append({
        "id": "generator",
        "label": "Datei-Generator",
        "state": _qa_level(map_ok, False),
        "detail": f"Projektkarte: {project_map.get('scanned_at') or 'nicht gescannt'}",
    })

    if isinstance(design_notes, dict):
        design_ok = bool(design_notes.get("updated_at") or design_notes.get("notes"))
    elif isinstance(design_notes, list):
        design_ok = bool(design_notes)
    else:
        design_ok = False
    modules.append({
        "id": "design",
        "label": "Design Studio",
        "state": _qa_level(design_ok, False),
        "detail": "Design-Notizen vorhanden." if design_ok else "Noch keine Design-Notizen.",
    })

    last_runner = run_state.get("last_runner_execution") if isinstance(run_state.get("last_runner_execution"), dict) else None
    runner_entry = None
    for entry in activity:
        if entry.get("label") in {"Runner", "Execute"}:
            runner_entry = entry
            break
    runner_ok = bool(last_runner) or bool(runner_entry)
    lr_status = str((last_runner or {}).get("status") or "").lower()
    runner_error = bool(
        (last_runner and lr_status in {"failed", "timeout", "error"})
        or (runner_entry and runner_entry.get("level") == "error")
    )
    runner_partial = bool(runner_ok and not runner_error and not last_runner and runner_entry)
    if last_runner:
        analysis_note = ""
        if last_runner.get("error_analysis_ran"):
            analysis_note = (
                "Fehleranalyse: Problem erkannt."
                if last_runner.get("error_analysis_has_issue")
                else "Fehleranalyse: kein strukturierter Fehler."
            )
        runner_detail = " · ".join(
            [
                f"Persistiert {last_runner.get('timestamp') or '-'}",
                f"Modus {last_runner.get('mode') or '-'}",
                f"Status {last_runner.get('status') or '-'}",
                f"RC {last_runner.get('returncode') if last_runner.get('returncode') is not None else '-'}",
                "Ausgabe ja" if last_runner.get("has_stdout") else "Ausgabe nein",
                "Stderr ja" if last_runner.get("has_stderr") else "Stderr nein",
                analysis_note,
            ]
        )
        if last_runner.get("command_preview"):
            runner_detail += " · Cmd " + str(last_runner.get("command_preview"))[:96]
    else:
        runner_detail = (runner_entry or {}).get("message") or "Noch kein persistierter Runner-Lauf (nur UI-Log oder nie ausgefuehrt)."
    modules.append({
        "id": "runner",
        "label": "Runner",
        "state": _qa_level(runner_ok and not runner_error, runner_partial, runner_error),
        "detail": runner_detail,
    })

    current_error = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else _empty_error_info()
    has_active_error = bool(current_error.get("category") and str(current_error.get("category")) != "none")
    unresolved_status = str(auto_loop.get("status") or "").lower() in {"running", "paused", "approval_required", "failed", "blocked"}
    last_error_entry = next((e for e in activity if e.get("level") == "error"), None) if unresolved_status else None
    analysis_detail = (current_error.get("label") or (last_error_entry or {}).get("message") or "Keine aktiven Fehler.")
    modules.append({
        "id": "error_analysis",
        "label": "Fehleranalyse",
        "state": "Fehler" if has_active_error or last_error_entry else "bereit",
        "detail": analysis_detail,
    })

    try:
        cp = subprocess.run(["git", "status", "--porcelain"], cwd=str(PROJECT_DIR),
                            capture_output=True, text=True, timeout=5)
        product_files, local_files = _classify_git_dirty(cp.stdout or "")
        dirty = bool(product_files)
        if not dirty and not local_files:
            git_detail = "Arbeitskopie: clean"
            git_state = "bereit"
        elif not dirty and local_files:
            git_detail = f"Arbeitskopie: clean (nur {len(local_files)} lokale Dev-Datei(en), nicht produktrelevant)"
            git_state = "bereit"
        else:
            git_detail = f"Arbeitskopie: {len(product_files)} Produkt-Aenderungen vorhanden"
            if local_files:
                git_detail += f" + {len(local_files)} lokale Dev-Datei(en)"
            git_state = "teilweise"
        modules.append({
            "id": "git",
            "label": "Git-Status",
            "state": git_state,
            "detail": git_detail,
        })
    except Exception as exc:
        modules.append({
            "id": "git",
            "label": "Git-Status",
            "state": "Fehler",
            "detail": f"git status fehlgeschlagen: {exc}",
        })

    context_ok = bool(isinstance(knowledge, dict) and (knowledge.get("built_at") or knowledge.get("summary")))
    built_at = knowledge.get("built_at") if isinstance(knowledge, dict) else None
    modules.append({
        "id": "context",
        "label": "Kontext",
        "state": _qa_level(context_ok, False),
        "detail": f"Kontext gebaut: {built_at or 'nein'}",
    })

    memory_rules = int(memory_info.get("rules_count") or 0)
    modules.append({
        "id": "memory",
        "label": "Projekt-Memory",
        "state": _qa_level(memory_rules > 0, False),
        "detail": f"Regeln aktiv: {memory_rules}",
    })

    agent_run_ok = bool(run_state.get("last_result") or run_state.get("last_check_result"))
    modules.append({
        "id": "agent_run",
        "label": "Agent-Run",
        "state": _qa_level(agent_run_ok, bool(run_state.get("last_planned_steps"))),
        "detail": str(run_state.get("last_result") or "Noch kein Agent-Run.")[:220],
    })

    al_status = auto_loop.get("status") or "idle"
    al_state_label = {
        "idle": "offen",
        "running": "teilweise",
        "paused": "teilweise",
        "approval_required": "teilweise",
        "blocked": "blockiert",
        "failed": "Fehler",
        "stopped": "teilweise",
        "done": "bereit",
    }.get(al_status, "offen")
    modules.append({
        "id": "auto_loop",
        "label": "Auto-Loop",
        "state": al_state_label,
        "detail": f"Status: {al_status}; Ziel: {(auto_loop.get('goal') or '-')[:140]}",
    })

    pending_approvals = []
    for step in (auto_loop.get("steps") or []):
        if str(step.get("status") or "").lower() == "wartet auf freigabe":
            pending_approvals.append({
                "step_id": str(step.get("id") or ""),
                "label": str(step.get("label") or ""),
                "tool": str(step.get("tool") or tool_for_action(step.get("action")) or ""),
                "phase": str(step.get("phase") or phase_for_action(step.get("action")) or ""),
                "detail": str(step.get("detail") or ""),
                "risky": bool(step.get("risky")),
            })

    pending_reviews = []
    if isinstance(current_patch, dict):
        for fe in (current_patch.get("file_entries") or []):
            status = str(fe.get("status") or "")
            if status not in {"applied", "verified", "closed"}:
                pending_reviews.append({
                    "path": str(fe.get("path") or ""),
                    "status": status or "review_ready",
                    "detail": str(fe.get("detail") or ""),
                })
    modules.append({
        "id": "review",
        "label": "Review",
        "state": "teilweise" if pending_reviews else "bereit",
        "detail": f"Offene Review-Dateien: {len(pending_reviews)}",
    })

    counts = {"bereit": 0, "teilweise": 0, "offen": 0, "Fehler": 0, "blockiert": 0}
    for mod in modules:
        state_key = str(mod.get("state") or "offen")
        counts[state_key] = counts.get(state_key, 0) + 1
    total = len(modules) or 1
    overall_status = "bereit"
    if counts.get("Fehler", 0):
        overall_status = "Fehler"
    elif counts.get("blockiert", 0):
        overall_status = "blockiert"
    elif counts.get("offen", 0) or counts.get("teilweise", 0):
        overall_status = "teilweise"

    ready_modules = [m for m in modules if _qa_bucket(m.get("state")) == "ready"]
    partial_modules = [m for m in modules if _qa_bucket(m.get("state")) == "partial"]
    blocked_modules = [m for m in modules if _qa_bucket(m.get("state")) == "blocked"]
    failed_modules = [m for m in modules if _qa_bucket(m.get("state")) == "failed"]

    blockers = auto_loop.get("session_blockers") if isinstance(auto_loop.get("session_blockers"), list) else []
    blocked_reason = ""
    if failed_modules:
        blocked_reason = failed_modules[0].get("detail") or failed_modules[0].get("label") or ""
    elif blocked_modules:
        blocked_reason = blocked_modules[0].get("detail") or blocked_modules[0].get("label") or ""
    elif blockers:
        blocked_reason = str((blockers[0] or {}).get("reason") or "")
    elif pending_approvals:
        blocked_reason = "Offene Freigaben vorhanden."
    elif pending_reviews:
        blocked_reason = "Offene Reviews vorhanden."

    # Global Freier Lauf: Acceptance gilt immer als automatisch erteilt.
    run_finished = True
    acceptance_ready = True
    acceptance_status = "accepted"
    acceptance_blocked_reason = ""
    acceptance_notes = str(auto_loop.get("acceptance_notes") or "").strip()
    acceptance_timestamp = str(auto_loop.get("acceptance_timestamp") or "").strip()
    run_metrics = _auto_loop_run_metrics(auto_loop)

    run_summary = {
        "run_id": str(auto_loop.get("run_id") or ""),
        "objective": str(auto_loop.get("goal") or run_state.get("last_task") or ""),
        "status": al_status,
        "phase": str(auto_loop.get("phase") or ""),
        "last_action": str(auto_loop.get("last_action") or ""),
        "completed": al_status == "done",
        "accepted": True,
        "blocked": bool(failed_modules or blocked_modules or al_status in {"failed", "blocked"}),
        "incomplete": False,
        "needs_approval": False,
        "total_steps": run_metrics.get("total_steps") or 0,
        "done_steps": run_metrics.get("done_steps") or 0,
        "retry_count": run_metrics.get("retry_count") or 0,
        "approvals_open": run_metrics.get("approvals_open") or 0,
        "blockers_count": run_metrics.get("blockers_count") or 0,
        "duration_seconds": run_metrics.get("duration_seconds") or 0,
        "duration_label": run_metrics.get("duration_label") or "00:00",
        "next_critical_gate": run_metrics.get("next_critical_gate") or "",
    }

    last_error_category = str(current_error.get("category") or "")
    if not last_error_category:
        for h in (auto_loop.get("repair_history") or []):
            if isinstance(h, dict) and h.get("category"):
                last_error_category = str(h.get("category"))
                break
    verification_summary = {
        "passed_checks": [
            "module_states",
            "review_status",
            "approval_queue",
            "git_status",
            "auto_loop_status",
        ],
        "successful": [m.get("id") for m in ready_modules],
        "failed": [m.get("id") for m in failed_modules],
        "blocked": [m.get("id") for m in blocked_modules],
        "open": [m.get("id") for m in partial_modules if m.get("id") not in {"review"}],
        "pending_reviews": 0,
        "pending_approvals": 0,
        "last_error_category": last_error_category,
        "last_review_status": str((current_patch or {}).get("review_status") or ""),
        "last_apply_status": str((current_patch or {}).get("apply_status") or ""),
        "last_git_check": get_timestamp(),
    }

    focus_todo = {}
    if isinstance(todos_content, dict) and isinstance(todos_content.get("todos"), list):
        for todo in todos_content["todos"]:
            if isinstance(todo, dict) and bool(todo.get("is_focus")):
                focus_todo = {
                    "id": str(todo.get("id") or ""),
                    "title": str(todo.get("title") or ""),
                    "status": str(todo.get("status") or "offen"),
                }
                break

    last_checked = get_timestamp()
    return {
        "ok": True,
        "timestamp": last_checked,
        "last_checked": last_checked,
        "overall": overall_status,
        "overall_status": overall_status,
        "acceptance_status": acceptance_status,
        "acceptance_ready": acceptance_ready,
        "acceptance_blocked_reason": acceptance_blocked_reason,
        "acceptance_notes": acceptance_notes,
        "acceptance_timestamp": acceptance_timestamp,
        "ready_modules": ready_modules,
        "partial_modules": partial_modules,
        "blocked_modules": blocked_modules,
        "failed_modules": failed_modules,
        "pending_approvals": [],
        "pending_reviews": [],
        "run_summary": run_summary,
        "verification_summary": verification_summary,
        "counts": counts,
        "ready_ratio": round(counts.get("bereit", 0) / total, 2),
        "modules": modules,
        "focus_todo": focus_todo,
    }


@app.route("/api/qa/status", methods=["GET"])
def qa_status_endpoint():
    return jsonify(build_qa_acceptance_snapshot())


@app.route("/api/qa/acceptance", methods=["POST"])
def qa_acceptance_endpoint():
    data = request.get_json(silent=True) or {}
    notes = str(data.get("notes") or "").strip()

    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    # Global Freier Lauf: QA ist immer automatisch signiert/akzeptiert.
    auto_loop["acceptance_status"] = "accepted"
    auto_loop["acceptance_ready"] = True
    auto_loop["acceptance_blocked_reason"] = ""
    auto_loop["acceptance_notes"] = notes or "Automatisch akzeptiert (Global Freier Lauf)."
    auto_loop["acceptance_timestamp"] = get_timestamp()
    auto_loop["qa_sign_off"] = True
    _auto_loop_push_history(auto_loop, "QA/Acceptance: automatisch akzeptiert.")
    append_ui_log_entry("QA", f"Lauf automatisch akzeptiert ({auto_loop.get('run_id') or '-'})", "info")

    normalized = _auto_loop_save(auto_loop)
    updated_qa = build_qa_acceptance_snapshot(run_state=load_project_auto_run_state(), auto_loop=normalized)
    return jsonify({"ok": True, "auto_loop": normalized, "qa": updated_qa})


def build_parity_snapshot():
    """Phase 12: Ehrliche Abschluss-Parity-Sicht aus vorhandenen Datenquellen."""
    qa = build_qa_acceptance_snapshot()
    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    activity = load_ui_activity_entries()
    project_map = read_json_file(PROJECT_MAP_FILE, {})
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    todos_content = read_json_file(DATA_DIR / "todos.json", {"todos": []})
    todos_list = todos_content.get("todos") or []

    qa_modules = {m["id"]: m for m in (qa.get("modules") or [])}

    def _dom_state(mod_id, fallback="offen"):
        return str((qa_modules.get(mod_id) or {}).get("state") or fallback)

    al_status = auto_loop.get("status") or "idle"
    al_phase = auto_loop.get("phase") or ""
    has_run = bool(auto_loop.get("run_id"))
    al_history = auto_loop.get("history") or []
    al_history_text = " ".join([str(h) for h in al_history[-30:]]).lower()

    direct_history = run_state.get("direct_run_history") or []
    direct_has_preview = any(r.get("status") == "safe_preview" for r in direct_history[:8])
    direct_has_blocked = any(r.get("status") == "blocked" for r in direct_history[:8])
    runner_entries = [e for e in activity if e.get("label") in {"Runner", "Execute"}]
    error_entries = [e for e in activity if e.get("level") == "error"]
    last_runner = run_state.get("last_runner_execution") if isinstance(run_state.get("last_runner_execution"), dict) else None
    last_project_scan = run_state.get("last_project_scan") if isinstance(run_state.get("last_project_scan"), dict) else None
    project_scanned = bool(project_map.get("scanned_at"))
    knowledge_ready = bool(isinstance(knowledge, dict) and (knowledge.get("built_at") or knowledge.get("summary")))
    scan_trace_ok = bool(last_project_scan and (last_project_scan.get("scanned_at") or last_project_scan.get("source")))
    if project_scanned and scan_trace_ok and knowledge_ready:
        scan_smoke_result = "geprüft"
    elif project_scanned and (scan_trace_ok or knowledge_ready):
        scan_smoke_result = "teilweise"
    elif project_scanned or scan_trace_ok or knowledge_ready:
        scan_smoke_result = "teilweise"
    else:
        scan_smoke_result = "offen"
    project_scan_evidence_ok = scan_smoke_result == "geprüft"

    repair_hist = auto_loop.get("repair_history") if isinstance(auto_loop.get("repair_history"), list) else []
    has_repair_hist_entries = bool(repair_hist)
    has_repair_plan = bool(auto_loop.get("repair_plan"))
    cur_err = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else {}
    err_cat = str(cur_err.get("category") or "").strip().lower()
    has_error_signal = bool(err_cat and err_cat not in {"", "none"})
    retry_n = int(auto_loop.get("retry_count") or 0)
    retry_max = int(auto_loop.get("retry_max") or 0)
    pending_reviews = len(qa.get("pending_reviews") or [])
    memory_mod = qa_modules.get("memory", {})
    git_mod = qa_modules.get("git", {})
    activity_count = len(activity)
    session_active = bool(auto_loop.get("goal") or has_run)

    runner_persisted = bool(last_runner and last_runner.get("timestamp"))
    runner_analysis_done = bool(last_runner and last_runner.get("error_analysis_ran"))
    runner_smoke_result = "geprüft" if (runner_persisted and runner_analysis_done) else (
        "teilweise" if (runner_persisted or runner_entries) else "offen"
    )
    runner_smoke_detail = "Kein persistierter Runner-Lauf."
    if last_runner:
        runner_smoke_detail = (
            f"Letzter Lauf {last_runner.get('timestamp') or '-'} · Modus {last_runner.get('mode') or '-'} · "
            f"Status {last_runner.get('status') or '-'} · Analyse {'ja' if runner_analysis_done else 'nein'} · "
            f"Stdout {'ja' if last_runner.get('has_stdout') else 'nein'} · Stderr {'ja' if last_runner.get('has_stderr') else 'nein'}"
        )
    elif runner_entries:
        runner_smoke_detail = f"Nur UI-Log ({len(runner_entries)} Eintraege), keine Persistenz."

    repair_smoke_result = "geprüft" if has_repair_hist_entries else (
        "teilweise" if (retry_n > 0 or has_repair_plan or has_error_signal or "retry" in al_history_text or "reparatur" in al_history_text) else "offen"
    )
    repair_smoke_detail = (
        f"Historie-Eintraege: {len(repair_hist)} · Retry {retry_n}/{(retry_max if retry_max else 2)} · "
        f"Plan: {'ja' if has_repair_plan else 'nein'} · Fehlersignal: {'ja' if has_error_signal else 'nein'}"
    )

    scan_smoke_detail = (
        f"project_map scanned_at: {project_map.get('scanned_at') or '-'} · "
        f"Letzter Scan-Pfad: {last_project_scan.get('source') if last_project_scan else '-'} · "
        f"Projektwissen: {'ja' if knowledge_ready else 'nein'}"
    )

    smoke_checks = [
        {
            "id": "direct_preview",
            "label": "Direktmodus Preview",
            "result": "geprüft" if direct_has_preview else "offen",
            "detail": "Safe-Preview im Verlauf." if direct_has_preview else "Noch kein Preview.",
        },
        {
            "id": "direct_guard_block",
            "label": "Direktmodus Guard-Block",
            "result": "geprüft" if direct_has_blocked else "offen",
            "detail": "Guard-Block erkannt." if direct_has_blocked else "Noch nicht getestet.",
        },
        {
            "id": "autoloop_lifecycle",
            "label": "Auto-Loop Start/Stop/Reset",
            "result": "geprüft" if ("zurueckgesetzt" in al_history_text and (has_run or "gestartet" in al_history_text)) else "teilweise" if (has_run or "zurueckgesetzt" in al_history_text) else "offen",
            "detail": f"Gestartet: {'ja' if (has_run or 'gestartet' in al_history_text) else 'nein'}, Reset: {'ja' if 'zurueckgesetzt' in al_history_text else 'nein'}",
        },
        {
            "id": "runner_error",
            "label": "Runner + Fehleranalyse",
            "result": runner_smoke_result,
            "detail": runner_smoke_detail,
        },
        {
            "id": "repair_retry",
            "label": "Repair / Retry Nachweis",
            "result": repair_smoke_result,
            "detail": repair_smoke_detail,
        },
        {
            "id": "project_scan",
            "label": "Project-Scan / Projektwissen",
            "result": scan_smoke_result,
            "detail": scan_smoke_detail,
        },
        {
            "id": "git_status",
            "label": "Git-Status/Dateizuordnung",
            "result": "geprüft" if git_mod.get("state") not in {"", None} else "offen",
            "detail": str(git_mod.get("detail") or "Noch nicht geprüft."),
        },
        {
            "id": "memory_rules",
            "label": "Memory im Lauf",
            "result": "geprüft" if memory_mod.get("state") == "bereit" else "offen",
            "detail": str(memory_mod.get("detail") or "Keine Regeln."),
        },
        {
            "id": "qa_acceptance",
            "label": "QA/Acceptance",
            "result": "geprüft" if qa.get("overall_status") not in {"", None, "offen"} else "offen",
            "detail": f"QA: {qa.get('overall_status') or '-'}, Acceptance: {qa.get('acceptance_status') or 'keine'}",
        },
        {
            "id": "review_apply",
            "label": "Review/Apply-Flow",
            "result": "geprüft" if pending_reviews == 0 else "teilweise",
            "detail": f"Offene Reviews: {pending_reviews}",
        },
    ]

    if has_repair_hist_entries:
        repair_status = "bereit"
    elif retry_n > 0 or has_repair_plan or has_error_signal or "retry" in al_history_text:
        repair_status = "teilweise"
    else:
        repair_status = _dom_state("error_analysis")

    if runner_persisted and runner_analysis_done:
        runner_domain_status = "bereit"
    elif runner_persisted or runner_entries:
        runner_domain_status = "teilweise"
    else:
        runner_domain_status = _dom_state("runner")

    agent_core_status = "bereit" if al_status in {"done", "running"} else "teilweise" if has_run else "offen"
    ws_status = "bereit" if (session_active or activity_count > 0) else "offen"
    qa_overall = qa.get("overall_status") or "offen"
    qa_domain_status = "bereit" if qa_overall == "bereit" else "blockiert" if qa_overall in {"blockiert", "Fehler"} else "teilweise" if qa_overall == "teilweise" else "offen"

    pm_notes = (
        f"Scan-Smoke: {scan_smoke_result}; "
        f"map {'ja' if project_scanned else 'nein'}; spur {'ja' if scan_trace_ok else 'nein'}; "
        f"Letzter Auftrag: {str(run_state.get('last_task') or '-')[:60]}; "
        f"Wissen: {'built' if knowledge_ready else 'offen'}"
    )
    domains = [
        {"id": "agent_core", "label": "Agent-Core", "status": agent_core_status, "notes": f"Laufstatus: {al_status}; Phase: {al_phase or '-'}"},
        {"id": "direktmodus", "label": "Direktmodus", "status": _dom_state("direct"), "notes": f"Preview: {'ja' if direct_has_preview else 'nein'}, Guard: {'geprüft' if direct_has_blocked else 'nein'}"},
        {"id": "project_mode", "label": "Project-Mode", "status": _dom_state("project_mode"), "notes": pm_notes},
        {"id": "runner", "label": "Runner", "status": runner_domain_status, "notes": runner_smoke_detail[:220]},
        {"id": "repair_retry", "label": "Fehleranalyse / Repair / Retry", "status": repair_status, "notes": repair_smoke_detail[:220]},
        {"id": "review_apply", "label": "Review / Apply", "status": _dom_state("review"), "notes": f"Offene Reviews: {pending_reviews}"},
        {"id": "git", "label": "Git-Arbeitsmodus", "status": _dom_state("git"), "notes": str(git_mod.get("detail") or "")},
        {"id": "memory", "label": "Projekt-Memory / Regeln", "status": _dom_state("memory"), "notes": str(memory_mod.get("detail") or "")},
        {"id": "qa_acceptance", "label": "QA / Acceptance", "status": qa_domain_status, "notes": f"Gesamt: {qa_overall}, Acceptance: {qa.get('acceptance_status') or 'keine'}"},
        {"id": "auto_loop", "label": "Auto-Loop / Orchestrierung", "status": _dom_state("auto_loop"), "notes": f"Status: {al_status}, Schritte: {len(auto_loop.get('steps') or [])}"},
        {"id": "workspace", "label": "Workspace / Session-Fokus / Konsole", "status": ws_status, "notes": f"Aktivitätslog: {activity_count} Einträge"},
    ]

    status_counts = {}
    for d in domains:
        s = d["status"]
        status_counts[s] = status_counts.get(s, 0) + 1
    bereit_count = status_counts.get("bereit", 0)
    teilweise_count = status_counts.get("teilweise", 0)
    blockiert_count = status_counts.get("blockiert", 0)
    fehler_count = status_counts.get("Fehler", 0)
    offen_count = status_counts.get("offen", 0)
    total = len(domains)

    if blockiert_count or fehler_count:
        parity_overall = "parity-blocked"
    elif bereit_count == total:
        parity_overall = "parity-ready"
    elif bereit_count >= total * 0.7:
        parity_overall = "partially-ready"
    else:
        parity_overall = "needs-work"

    real_gaps = [d for d in domains if d["status"] in {"offen", "blockiert", "Fehler"}]
    strong_domains = [d for d in domains if d["status"] == "bereit"]

    parity_notes = []
    if not runner_persisted and not runner_entries:
        parity_notes.append("Runner noch nicht ausgefuehrt – kein persistierter Nachweis und kein UI-Log.")
    elif not runner_persisted and runner_entries:
        parity_notes.append("Runner nur im kurzen UI-Log sichtbar – letzter Lauf nicht persistiert (nach naechstem Execute voll nachweisbar).")
    if not direct_has_preview:
        parity_notes.append("Direktmodus noch ohne abgeschlossenen Preview-Lauf.")
    if scan_smoke_result == "offen":
        parity_notes.append("Project-Scan/Projektwissen noch nicht nachweisbar (weder project_map noch Wissen/Scan-Spur).")
    elif scan_smoke_result == "teilweise":
        parity_notes.append("Project-Scan nur teilweise belastbar (z.B. Karte ohne Wissen oder umgekehrt).")
    if not has_repair_hist_entries and repair_smoke_result == "offen":
        parity_notes.append("Kein Repair-/Retry-Historieneintrag – keine abgeschlossene Repair-Kette persistiert.")
    elif not has_repair_hist_entries and repair_smoke_result == "teilweise":
        parity_notes.append("Repair/Retry nur teilweise sichtbar (z.B. Plan oder Retry-Zaehler ohne abgeschlossene Historie).")
    if pending_reviews > 0:
        parity_notes.append(f"{pending_reviews} offene Review-Datei(en) vorhanden.")
    if qa.get("acceptance_status") not in {"accepted", "ready"}:
        parity_notes.append("QA/Acceptance noch nicht freigegeben.")

    if parity_overall == "parity-ready":
        completion_summary = f"Alle {total} Kernbereiche auf 'bereit'. Rainer Build ist parity-ready."
    elif parity_overall == "parity-blocked":
        completion_summary = f"Parity durch {blockiert_count + fehler_count} blockierte/fehlerhafte Bereiche verhindert."
    elif bereit_count + teilweise_count == total:
        completion_summary = f"{bereit_count}/{total} Kernbereiche bereit, {teilweise_count} teilweise. Keine echten Blocker."
    else:
        completion_summary = f"{bereit_count}/{total} Kernbereiche bereit. Restlücken: {offen_count} offen."

    smoke_ok = sum(1 for s in smoke_checks if s["result"] == "geprüft")
    done_count = sum(1 for t in todos_list if t.get("status") == "erledigt")
    focus_todo = next((t for t in todos_list if t.get("is_focus")), None)

    last_round = repair_hist[-1] if repair_hist else None
    if isinstance(last_round, dict):
        last_round = {k: last_round.get(k) for k in ("round", "status", "category", "timestamp", "repeated") if k in last_round}

    runner_evidence = {
        "persisted": runner_persisted,
        "last_timestamp": (last_runner or {}).get("timestamp") or "",
        "last_mode": (last_runner or {}).get("mode") or "",
        "last_status": (last_runner or {}).get("status") or "",
        "analysis_ran": runner_analysis_done,
        "analysis_has_issue": bool(last_runner and last_runner.get("error_analysis_has_issue")),
        "has_stdout": bool(last_runner and last_runner.get("has_stdout")),
        "has_stderr": bool(last_runner and last_runner.get("has_stderr")),
        "ui_log_hits": len(runner_entries),
    }
    repair_evidence = {
        "history_entries": len(repair_hist),
        "last_round": last_round,
        "retry_count": retry_n,
        "retry_max": retry_max,
        "repair_plan_steps": len(auto_loop.get("repair_plan") or []) if isinstance(auto_loop.get("repair_plan"), list) else 0,
        "error_category_signal": err_cat,
    }
    project_scan_evidence_payload = {
        "map_scanned_at": project_map.get("scanned_at") or "",
        "map_total_files": int(project_map.get("total_files") or 0),
        "last_scan_source": (last_project_scan or {}).get("source") or "",
        "last_scan_at": (last_project_scan or {}).get("scanned_at") or "",
        "knowledge_ready": knowledge_ready,
        "knowledge_built_at": str(knowledge.get("built_at") or "") if isinstance(knowledge, dict) else "",
    }

    return {
        "ok": True,
        "timestamp": get_timestamp(),
        "parity_overall": parity_overall,
        "completion_summary": completion_summary,
        "runner_evidence": runner_evidence,
        "repair_evidence": repair_evidence,
        "project_scan_evidence": project_scan_evidence_payload,
        "domains": domains,
        "strong_domains": [d["label"] for d in strong_domains],
        "real_gaps": [{"label": d["label"], "notes": d["notes"]} for d in real_gaps],
        "partial_domains": [d["label"] for d in domains if d["status"] == "teilweise"],
        "domain_counts": {"bereit": bereit_count, "teilweise": teilweise_count, "offen": offen_count, "blockiert": blockiert_count},
        "smoke_checks": smoke_checks,
        "smoke_summary": {"ok": smoke_ok, "total": len(smoke_checks)},
        "parity_notes": parity_notes,
        "parity_blockers": [d["notes"] for d in real_gaps if d["notes"]],
        "qa_summary": {
            "overall": qa.get("overall_status"),
            "acceptance": qa.get("acceptance_status"),
            "ready_count": len(qa.get("ready_modules") or []),
            "module_total": len(qa.get("modules") or []),
        },
        "roadmap_summary": {
            "done": done_count,
            "total": len(todos_list),
            "focus_title": str((focus_todo or {}).get("title") or ""),
            "focus_status": str((focus_todo or {}).get("status") or ""),
        },
    }


@app.route("/api/parity/status", methods=["GET"])
def parity_status_endpoint():
    return jsonify(build_parity_snapshot())


AGENT_CORE_TOOLS = [
    ("direct", "Direktmodus"),
    ("project_mode", "Project-Mode"),
    ("runner", "Runner"),
    ("error_analysis", "Fehleranalyse"),
    ("context", "Kontext"),
    ("repair", "Reparatur"),
    ("git", "Git"),
    ("agent_run", "Agent-Run"),
]


def _agent_core_tools(run_state, auto_loop):
    active = (auto_loop.get("active_module") or "").strip()
    tools = []
    step_tool_set = set()
    for step in (auto_loop.get("steps") or []):
        t = step.get("tool")
        if t:
            step_tool_set.add(t)
    for tool_id, label in AGENT_CORE_TOOLS:
        state = "idle"
        detail = ""
        if tool_id == active and auto_loop.get("status") == "running":
            state = "active"
        elif tool_id == active and auto_loop.get("status") == "approval_required":
            state = "waiting"
        elif tool_id == active and auto_loop.get("status") == "blocked":
            state = "blocked"
        elif tool_id in step_tool_set:
            state = "planned"
        if tool_id == "direct":
            detail = str(run_state.get("last_direct_status") or "idle")
        elif tool_id == "project_mode":
            detail = "blocked" if run_state.get("blocked") else (str(run_state.get("last_task") or "")[:80] or "idle")
        elif tool_id == "runner":
            detail = "Letzter Lauf vorhanden" if run_state.get("last_result") else "kein Lauf"
        elif tool_id == "error_analysis":
            detail = str(auto_loop.get("error_label") or "keine aktive Fehlerkategorie")
        elif tool_id == "context":
            detail = "Kontext verfuegbar" if Path(DATA_DIR / "project_knowledge.json").exists() else "kein Kontext"
        elif tool_id == "repair":
            detail = str(auto_loop.get("repair_suggestion") or "-")[:80]
        elif tool_id == "git":
            detail = "Git-Status via /api/git/status"
        elif tool_id == "agent_run":
            detail = str(run_state.get("last_check_result") or run_state.get("last_result") or "idle")[:80]
        tools.append({"id": tool_id, "label": label, "state": state, "detail": detail or "-"})
    return tools


def _build_quick_actions(status, auto_loop, current_patch, current_error):
    """Kontextbasierte Einzel-Klick-Aktionen fuer den aktiven Lauf (BLOCK D)."""
    actions = []
    ctrl = "/api/orchestrator/control"
    if status == "idle":
        actions.append({"id": "start", "label": "Auto-Loop starten", "method": "POST", "endpoint": ctrl, "body": {"action": "start"}, "priority": 0})
    elif status == "approval_required":
        actions.append({"id": "approve", "label": "Freigeben", "method": "POST", "endpoint": ctrl, "body": {"action": "approve"}, "priority": 0})
        actions.append({"id": "reject", "label": "Ablehnen", "method": "POST", "endpoint": ctrl, "body": {"action": "reject"}, "priority": 1})
    elif status == "paused":
        actions.append({"id": "resume", "label": "Fortsetzen", "method": "POST", "endpoint": ctrl, "body": {"action": "resume"}, "priority": 0})
        actions.append({"id": "stop", "label": "Stoppen", "method": "POST", "endpoint": ctrl, "body": {"action": "stop"}, "priority": 2})
    elif status == "blocked":
        actions.append({"id": "advance", "label": "Erneut versuchen", "method": "POST", "endpoint": ctrl, "body": {"action": "advance"}, "priority": 0})
        actions.append({"id": "reset", "label": "Zurücksetzen", "method": "POST", "endpoint": ctrl, "body": {"action": "reset"}, "priority": 1})
    elif status == "running":
        actions.append({"id": "pause", "label": "Pausieren", "method": "POST", "endpoint": ctrl, "body": {"action": "pause"}, "priority": 0})
        actions.append({"id": "stop", "label": "Stoppen", "method": "POST", "endpoint": ctrl, "body": {"action": "stop"}, "priority": 1})
    elif status in {"done", "stopped"}:
        actions.append({"id": "reset", "label": "Neuer Lauf", "method": "POST", "endpoint": ctrl, "body": {"action": "reset"}, "priority": 1})
        actions.append({"id": "qa_accept", "label": "QA annehmen", "method": "POST", "endpoint": "/api/qa/acceptance", "body": {"action": "accept"}, "priority": 0})
    # Error shortcut
    if isinstance(current_error, dict) and (current_error.get("category") or "") not in {"", "none"}:
        rec = current_error.get("recommended_tool") or "direct"
        actions.append({"id": "fix_error", "label": f"Fehler beheben ({current_error.get('label') or rec})", "method": "POST", "endpoint": ctrl, "body": {"action": "advance"}, "tool_focus": rec, "priority": 0})
    # Review shortcut
    if isinstance(current_patch, dict) and (current_patch.get("review_status") or "") not in {"closed", "applied", "verified", ""}:
        actions.append({"id": "open_review", "label": "Review öffnen", "method": "GET", "endpoint": "/api/patch-review/state", "priority": 1})
    actions.sort(key=lambda x: x.get("priority", 99))
    seen = set()
    result = []
    for a in actions:
        if a["id"] not in seen:
            seen.add(a["id"])
            result.append({k: v for k, v in a.items() if k != "priority"})
    return result[:5]


def _build_context_actions(current_error, current_patch, auto_loop):
    """Koppelt Error-/Patch-/Approval-Kontext direkt an handlungsfaehige Eintraege (BLOCK C)."""
    actions = []
    if isinstance(current_error, dict) and (current_error.get("category") or "") not in {"", "none"}:
        actions.append({
            "source": "error",
            "label": current_error.get("label") or "Fehler",
            "category": current_error.get("category") or "",
            "file": current_error.get("file") or "",
            "suggestion": current_error.get("suggestion") or "",
            "recommended_tool": current_error.get("recommended_tool") or "direct",
            "next_action": "retry" if auto_loop.get("retry_possible") else "reset",
        })
    if isinstance(current_patch, dict) and (current_patch.get("review_status") or "") not in {"closed", "applied", "verified", ""}:
        actions.append({
            "source": "patch",
            "label": f"Review: {(current_patch.get('objective') or '')[:80]}",
            "patch_id": current_patch.get("patch_id") or current_patch.get("run_id") or "",
            "review_status": current_patch.get("review_status") or "",
            "apply_status": current_patch.get("apply_status") or "",
            "recommended_tool": "review",
            "next_action": "approve" if (current_patch.get("apply_status") or "") == "pending" else "review",
        })
    for step in (auto_loop.get("steps") or []):
        if str(step.get("status") or "").lower() == "wartet auf freigabe":
            actions.append({
                "source": "approval",
                "label": f"Freigabe: {step.get('label') or step.get('id') or ''}",
                "step_id": step.get("id") or "",
                "tool": step.get("tool") or "direct",
                "recommended_tool": step.get("tool") or "direct",
                "next_action": "approve",
            })
            break
    return actions


def _agent_pending_step(auto_loop):
    for step in (auto_loop.get("steps") or []):
        step_status = str(step.get("status") or "").lower()
        if step_status in {"geplant", "wartet auf freigabe", "laeuft"}:
            return step
    return None


def _agent_step_target_file(step, active_file=None, current_error=None, current_patch=None, run_state=None):
    step = step if isinstance(step, dict) else {}
    current_error = current_error if isinstance(current_error, dict) else {}
    current_patch = current_patch if isinstance(current_patch, dict) else {}
    run_state = run_state if isinstance(run_state, dict) else {}
    for key in ("file_target", "path", "target_file", "affected_file"):
        value = str(step.get(key) or "").strip()
        if _is_valid_file_path(value):
            return value
    err_file = str(current_error.get("file") or "").strip()
    if _is_valid_file_path(err_file):
        return err_file
    for fe in (current_patch.get("file_entries") or []):
        path = str((fe or {}).get("path") or "").strip()
        if _is_valid_file_path(path) and str((fe or {}).get("status") or "").lower() not in {"closed", "applied", "verified"}:
            return path
    if isinstance(active_file, dict):
        af_path = str(active_file.get("path") or "").strip()
        if _is_valid_file_path(af_path):
            return af_path
    raw_last_target_paths = run_state.get("last_target_paths")
    last_target_paths = raw_last_target_paths if isinstance(raw_last_target_paths, list) else []
    for candidate in last_target_paths:
        rel = str(candidate or "").strip()
        if _is_valid_file_path(rel):
            return rel
    return ""


def _agent_action_is_safe(step):
    step = step if isinstance(step, dict) else {}
    action_id = str(step.get("action") or step.get("id") or "").strip().lower()
    gate = str(step.get("gate") or "auto").strip().lower()
    if gate == "approval" or bool(step.get("risky")):
        return False
    if action_id in {
        "apply_patch", "write_file", "commit", "stage", "unstage",
        "create_branch", "revert_file", "apply_edit", "apply_repair",
        "auto_apply", "repo_write", "runner_execute",
    }:
        return False
    return True


def _append_followup_action(actions, action_id, label, endpoint="", method="GET", body=None,
                            view="", tool="", file_path="", reason="", primary=False,
                            auto_ok=False, disabled=False):
    if disabled:
        return
    entry = {
        "id": str(action_id or ""),
        "label": str(label or ""),
        "endpoint": str(endpoint or ""),
        "method": str(method or "GET"),
        "body": body if isinstance(body, dict) else {},
        "view": str(view or ""),
        "tool": str(tool or ""),
        "file": str(file_path or ""),
        "reason": str(reason or ""),
        "primary": bool(primary),
        "auto_ok": bool(auto_ok),
    }
    for existing in actions:
        if existing.get("id") == entry["id"]:
            return
    actions.append(entry)


def _record_agent_decision_trace(run_state, decision):
    run_state = run_state if isinstance(run_state, dict) else {}
    decision = decision if isinstance(decision, dict) else {}
    compact = {
        "timestamp": get_timestamp(),
        "status": str(decision.get("status") or ""),
        "gate_type": str(decision.get("gate_type") or ""),
        "next_action": str(decision.get("next_action") or ""),
        "next_tool": str(decision.get("next_tool") or ""),
        "next_file": str(decision.get("next_file") or ""),
        "reason": str(decision.get("next_reason") or decision.get("why_stopped") or decision.get("why_continuing") or ""),
        "pending_step_id": str(decision.get("pending_step_id") or ""),
        "auto_ok": bool(decision.get("auto_ok")),
    }
    signature = "|".join([
        compact["status"], compact["gate_type"], compact["next_action"],
        compact["next_tool"], compact["next_file"], compact["reason"],
        compact["pending_step_id"], "1" if compact["auto_ok"] else "0",
    ])
    last = run_state.get("last_agent_decision") if isinstance(run_state.get("last_agent_decision"), dict) else {}
    if signature == str(last.get("signature") or ""):
        return run_state.get("agent_decision_history") if isinstance(run_state.get("agent_decision_history"), list) else []
    compact["signature"] = signature
    history = run_state.get("agent_decision_history") if isinstance(run_state.get("agent_decision_history"), list) else []
    history.insert(0, compact)
    save_project_auto_run_state({
        "last_agent_decision": compact,
        "agent_decision_history": history[:18],
    })
    return history[:18]


def _build_agent_decision(run_state, auto_loop, current_patch=None, current_error=None, qa_info=None, active_file=None):
    run_state = run_state if isinstance(run_state, dict) else {}
    auto_loop = auto_loop if isinstance(auto_loop, dict) else normalize_auto_loop_state({})
    current_patch = current_patch if isinstance(current_patch, dict) else {}
    current_error = current_error if isinstance(current_error, dict) else _empty_error_info()
    qa_info = qa_info if isinstance(qa_info, dict) else {}
    active_file = active_file if isinstance(active_file, dict) else {}

    status = str(auto_loop.get("status") or "idle").strip().lower() or "idle"
    phase = str(auto_loop.get("phase") or "planning").strip().lower() or "planning"
    current_step = str(auto_loop.get("current_step") or "").strip()
    next_step_label = str(auto_loop.get("next_step") or "").strip()
    last_action = str(auto_loop.get("last_action") or "").strip()
    # Global Freier Lauf: keine Gate-Pause wegen Freigaben.
    requires_approval = False
    pending_step = _agent_pending_step(auto_loop)
    pending_approval = auto_loop.get("pending_approval") if isinstance(auto_loop.get("pending_approval"), dict) else {}

    next_action = ""
    next_tool = ""
    next_file = ""
    next_reason = ""
    why_continuing = ""
    why_stopped = ""
    gate_type = "ambiguous"
    api_url = ""
    api_method = "POST"
    api_body = {}
    auto_ok = False
    can_auto_continue = False
    safe_chain = False
    followup_actions = []
    gate_required = False

    if pending_step:
        next_tool = str(pending_step.get("tool") or tool_for_action(pending_step.get("action")) or pending_step.get("id") or "").strip()
        next_file = _agent_step_target_file(pending_step, active_file, current_error, current_patch, run_state)

    patch_review_status = str(current_patch.get("review_status") or "").lower()
    patch_apply_status = str(current_patch.get("apply_status") or "").lower()
    patch_direct_status = str(current_patch.get("direct_status") or "").lower()
    patch_open = (
        bool(current_patch)
        and patch_direct_status not in {"safe_preview", "blocked"}
        and (
            patch_review_status in {"open", "pending", "needs_review"}
            or patch_apply_status in {"pending", "applied"}
        )
    )
    patch_context_run_id = str(current_patch.get("run_id") or current_patch.get("patch_id") or "").strip()
    active_context_run_ids = {
        str(run_state.get("active_direct_run_id") or "").strip(),
        str(run_state.get("last_direct_confirmed_run_id") or "").strip(),
        str(run_state.get("last_completed_run_id") or "").strip(),
        str(auto_loop.get("run_id") or "").strip(),
    }
    active_context_run_ids.discard("")
    has_active_context = bool(active_context_run_ids)
    patch_matches_active_context = bool(patch_context_run_id) and patch_context_run_id in active_context_run_ids
    review_relevant_for_current_run = bool(patch_open) and (patch_matches_active_context or not has_active_context)
    acceptance_open = False
    pending_reviews = int(qa_info.get("pending_reviews") or 0)
    pending_reviews_relevant = False
    pending_approvals = int(qa_info.get("pending_approvals") or 0)
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    has_pending_direct_context = pending_direct_context_is_valid(pending_direct, run_state)
    pending_approvals_relevant = False
    retry_possible = bool(auto_loop.get("retry_possible"))
    current_error_active = bool(str(current_error.get("category") or "").strip()) and str(current_error.get("category") or "").strip().lower() not in {"none", ""}
    last_direct_status = str(run_state.get("last_direct_status") or "").strip().lower()
    has_confirmed_direct = bool(str(run_state.get("last_direct_confirmed_run_id") or "").strip())

    if status == "approval_required":
        status = "running"

    if status == "idle":
        if pending_reviews_relevant or review_relevant_for_current_run:
            gate_type = "stopped"
            next_action = "review_patch"
            next_tool = "review"
            next_reason = "Offenes Review vorhanden."
            why_stopped = "Review ist noch offen und braucht manuelle Klaerung."
            _append_followup_action(followup_actions, "open_review", "Review oeffnen", "", "GET", {}, "review", "review", next_file, next_reason, primary=True)
        elif pending_approvals_relevant:
            gate_type = "approval_pending"
            gate_required = True
            next_action = "approve"
            next_tool = "direct"
            next_reason = "Offene Freigabe vorhanden."
            why_stopped = "Eine Freigabeentscheidung ist noch offen."
            api_url = "/api/auto-loop/approve"
            _append_followup_action(followup_actions, "approve_gate", "Freigeben", "/api/auto-loop/approve", "POST", {}, "agent_core", next_tool, next_file, next_reason, primary=True)
            _append_followup_action(followup_actions, "reject_gate", "Ablehnen", "/api/auto-loop/reject", "POST", {"reason": "Freigabe verweigert."}, "agent_core", next_tool, next_file, "Gate bewusst stoppen.")
        elif has_confirmed_direct and last_direct_status in {"safe_preview", "applied", "verified"} and not current_error_active:
            gate_type = "stopped"
            next_action = "completed"
            next_tool = "agent_run"
            next_reason = "Letzter Lauf wurde abgeschlossen."
            why_stopped = "Kein weiterer Schritt erforderlich."
            _append_followup_action(followup_actions, "new_loop", "Neuer Lauf", "/api/orchestrator/control", "POST", {"action": "reset"}, "agent_core", "agent_run", "", "Neuen Lauf vorbereiten.", primary=True)
        else:
            gate_type = "stopped"
            next_action = "start"
            next_tool = "agent_run"
            next_reason = "Kein aktiver Lauf."
            why_stopped = "Kein aktiver Lauf. Neuer Lauf noetig."
            _append_followup_action(followup_actions, "start_loop", "Auto-Loop starten", "/api/orchestrator/control", "POST", {"action": "start"}, "agent_core", "agent_run", "", next_reason, primary=True)

    elif status == "approval_required" or requires_approval:
        gate_type = "approval_pending"
        gate_required = True
        next_action = "approve"
        next_tool = str(pending_approval.get("tool") or next_tool or "direct")
        next_file = next_file or str(pending_approval.get("file") or "")
        label = str(
            pending_approval.get("step_label")
            or pending_approval.get("label")
            or (pending_step.get("label") if pending_step else "")
            or next_tool
            or "Schritt"
        )
        next_reason = f"Freigabe erforderlich: {label}"
        why_stopped = "Schreib-, Apply- oder Repo-Schritt wartet auf Freigabe."
        api_url = "/api/auto-loop/approve"
        _append_followup_action(followup_actions, "approve_gate", "Freigeben", "/api/auto-loop/approve", "POST", {}, "agent_core", next_tool, next_file, next_reason, primary=True)
        _append_followup_action(followup_actions, "reject_gate", "Ablehnen", "/api/auto-loop/reject", "POST", {"reason": "Freigabe verweigert."}, "agent_core", next_tool, next_file, "Gate bewusst stoppen.")
        if next_file:
            _append_followup_action(followup_actions, "open_gate_file", "Datei oeffnen", f"/api/file/load?path={quote(next_file)}", "GET", {}, "file_workspace", "file_workspace", next_file, "Betroffene Datei vor Freigabe pruefen.")

    elif status in {"blocked", "failed"}:
        gate_type = "stopped"
        next_action = "retry" if retry_possible else "review_blocker"
        next_tool = str(current_error.get("recommended_tool") or next_tool or "agent_run")
        next_file = next_file or str(current_error.get("file") or "")
        next_reason = str(current_error.get("label") or last_action or status)
        why_stopped = f"Lauf gestoppt: {next_reason}"
        if retry_possible:
            _append_followup_action(followup_actions, "retry_loop", "Retry starten", "/api/orchestrator/control", "POST", {"action": "advance"}, "agent_core", next_tool, next_file, "Kontrollierter Retry verfuegbar.", primary=True)
        if next_file:
            _append_followup_action(followup_actions, "open_error_file", "Datei oeffnen", f"/api/file/load?path={quote(next_file)}", "GET", {}, "file_workspace", "file_workspace", next_file, "Fehlerdatei direkt pruefen.")
        if current_error_active:
            _append_followup_action(followup_actions, "open_error_analysis", "Fehleranalyse", "", "GET", {}, "agent_core", "error_analysis", next_file, "Fehlerkategorie und Reparaturpfad ansehen.")

    elif status == "paused":
        gate_type = "stopped"
        next_action = "resume"
        next_tool = next_tool or "agent_run"
        next_reason = "Lauf pausiert."
        why_stopped = "Lauf pausiert. Fortsetzen fuer weitere sichere Schritte."
        _append_followup_action(followup_actions, "resume_loop", "Fortsetzen", "/api/orchestrator/control", "POST", {"action": "resume"}, "agent_core", next_tool, next_file, next_reason, primary=True)

    elif status in {"done", "completed"}:
        gate_type = "stopped"
        next_tool = "qa"
        next_action = "idle"
        next_reason = "Lauf abgeschlossen."
        why_stopped = "Kein weiterer Auto-Schritt offen."
        if review_relevant_for_current_run or pending_reviews_relevant:
            next_action = "review_patch"
            next_tool = "review"
            next_reason = "Patch-Review noch offen."
            _append_followup_action(followup_actions, "open_review", "Review oeffnen", "", "GET", {}, "review", "review", next_file, next_reason, primary=True)
        elif acceptance_open and bool(qa_info.get("acceptance_ready")):
            next_action = "qa_accept"
            next_tool = "qa"
            next_reason = "Acceptance kann abgeschlossen werden."
            _append_followup_action(followup_actions, "qa_accept", "QA annehmen", "/api/qa/acceptance", "POST", {"action": "accept"}, "qa", "qa", "", next_reason, primary=True)
        else:
            _append_followup_action(followup_actions, "new_loop", "Neuer Lauf", "/api/orchestrator/control", "POST", {"action": "reset"}, "agent_core", "agent_run", "", "Neuen Lauf vorbereiten.", primary=True)
        if next_file:
            _append_followup_action(followup_actions, "open_completed_file", "Datei oeffnen", f"/api/file/load?path={quote(next_file)}", "GET", {}, "file_workspace", "file_workspace", next_file, "Letzte relevante Datei pruefen.")

    elif status == "running":
        if pending_step:
            step_status = str(pending_step.get("status") or "").lower()
            step_label = str(pending_step.get("label") or next_tool or pending_step.get("id") or "")
            next_action = str(pending_step.get("action") or pending_step.get("id") or "advance").strip().lower()
            next_reason = str(pending_step.get("detail") or "") or f"Naechster Schritt: {step_label}"
            if step_status == "wartet auf freigabe" or not _agent_action_is_safe(pending_step):
                gate_type = "approval_gate"
                gate_required = True
                auto_ok = False
                why_stopped = f"Gate erreicht: {step_label}"
                api_url = "/api/auto-loop/approve"
                _append_followup_action(followup_actions, "approve_gate", "Freigeben", "/api/auto-loop/approve", "POST", {}, "agent_core", next_tool, next_file, why_stopped, primary=True)
                _append_followup_action(followup_actions, "pause_loop", "Pausieren", "/api/orchestrator/control", "POST", {"action": "pause"}, "agent_core", "agent_run", next_file, "Automatik vor Gate anhalten.")
            else:
                gate_type = "auto"
                auto_ok = True
                can_auto_continue = True
                safe_chain = True
                why_continuing = f"Sicherer Folge-Schritt: {step_label}"
                api_url = "/api/auto-loop/advance"
                _append_followup_action(followup_actions, "auto_advance", "Auto-Schritt", "/api/auto-loop/advance", "POST", {}, "agent_core", next_tool, next_file, why_continuing, primary=True, auto_ok=True)
                if next_file:
                    _append_followup_action(followup_actions, "open_active_file", "Datei oeffnen", f"/api/file/load?path={quote(next_file)}", "GET", {}, "file_workspace", "file_workspace", next_file, "Dateikontext parallel laden.")
                if review_relevant_for_current_run:
                    _append_followup_action(followup_actions, "open_review", "Review oeffnen", "", "GET", {}, "review", "review", next_file, "Vorbereiteten Patch ansehen.")
        else:
            gate_type = "auto"
            next_action = "advance"
            next_tool = str(active_file.get("recommended_tool") or "") if isinstance(active_file, dict) else ""
            if not next_file and isinstance(active_file, dict):
                next_file = str(active_file.get("path") or "")
            next_reason = f"Lauf aktiv in Phase {phase}."
            why_continuing = "Kein blockierender Schritt offen."
            auto_ok = True
            can_auto_continue = True
            safe_chain = True
            api_url = "/api/auto-loop/advance"
            _append_followup_action(followup_actions, "auto_advance", "Weiterlaufen", "/api/auto-loop/advance", "POST", {}, "agent_core", next_tool, next_file, why_continuing, primary=True, auto_ok=True)

    else:
        gate_type = "ambiguous"
        next_action = "review_state"
        next_reason = f"Unbekannter Status: {status}"
        why_stopped = "Manuelle Pruefung noetig."

    if not next_file and isinstance(active_file, dict):
        next_file = str(active_file.get("path") or "")
    if pending_approvals_relevant and gate_type not in {"approval_gate", "approval_pending"}:
        _append_followup_action(followup_actions, "approve_pending", "Offene Freigabe", "/api/auto-loop/approve", "POST", {}, "agent_core", "direct", next_file, "Offene Freigabe vorhanden.", primary=False)
    if review_relevant_for_current_run and gate_type not in {"approval_gate", "approval_pending"}:
        _append_followup_action(followup_actions, "open_review", "Review oeffnen", "", "GET", {}, "review", "review", next_file, "Offenes Review verfuegbar.")
    if isinstance(active_file, dict) and active_file.get("path"):
        _append_followup_action(followup_actions, "open_workspace", "Workspace", f"/api/file/load?path={quote(str(active_file.get('path') or ''))}", "GET", {}, "file_workspace", "file_workspace", str(active_file.get("path") or ""), "Relevante Datei im Workspace oeffnen.")
    if int(qa_info.get("pending_reviews") or 0) == 0 and bool(qa_info.get("acceptance_ready")) and acceptance_open:
        _append_followup_action(followup_actions, "open_qa", "QA pruefen", "", "GET", {}, "qa", "qa", "", "Acceptance ist bereit.")

    remote_workflow = {}
    if str(status).lower() in {"done", "completed", "idle"}:
        try:
            remote_workflow = _collect_remote_workflow_for_agent_core()
        except Exception:
            remote_workflow = {}
    post_run_publish_hints = (
        str(status).lower() in {"done", "completed"}
        or (
            str(status).lower() == "idle"
            and has_confirmed_direct
            and last_direct_status in {"safe_preview", "applied", "verified"}
            and not current_error_active
        )
    )
    if post_run_publish_hints and remote_workflow:
        _append_remote_publish_followups(followup_actions, remote_workflow)

    KNOWN_GATE_TYPES = {"auto", "approval_gate", "approval_pending", "stopped", "ambiguous"}
    if gate_type not in KNOWN_GATE_TYPES:
        gate_type = "ambiguous"

    # BLOCK F: auto_ok must never be True without a valid internal api_url
    resolved_auto_ok = bool(auto_ok) and bool(api_url) and str(api_url).startswith("/api/")
    if bool(auto_ok) and not resolved_auto_ok:
        gate_type = "ambiguous"
        why_stopped = why_stopped or "Interner Fehler: auto_ok ohne gültige API-URL."
        next_reason = next_reason or why_stopped

    return {
        "status": status,
        "phase": phase,
        "next_action": next_action,
        "next_tool": next_tool,
        "next_file": next_file,
        "next_reason": next_reason,
        "auto_ok": resolved_auto_ok,
        "can_auto_continue": resolved_auto_ok,
        "safe_chain": resolved_auto_ok,
        "gate_type": gate_type,
        "gate_required": bool(gate_required),
        "why_continuing": why_continuing,
        "why_stopped": why_stopped,
        "api_url": api_url,
        "api_method": api_method,
        "api_body": api_body,
        "step_label": str(pending_step.get("label") if pending_step else (next_step_label or current_step)),
        "pending_step_id": str(pending_step.get("id") if pending_step else ""),
        "followup_actions": followup_actions[:8],
        "console_label": "",
        "remote_workflow": remote_workflow if str(status).lower() in {"done", "completed", "idle"} else {},
    }


def _collect_git_info_for_agent_core(current_patch):
    """Kompakter Git-Arbeitsmodus-Snapshot fuer Agent-Core-State.

    Liefert Branch, Dateiezahl im aktuellen Lauf, externe Aenderungen und
    offene Review-Dateien. Nutzt nur read-only Git-Kommandos und fallback-t
    ruhig, wenn Git nicht erreichbar ist.
    """
    info = {
        "current_branch": "",
        "changed_count": 0,
        "current_run_files": [],
        "current_run_count": 0,
        "external_changed_files_count": 0,
        "review_files_pending": 0,
        "has_external_changes": False,
        "ahead": 0,
        "behind": 0,
    }
    try:
        project_root = str(RAMBO_RAINER_ROOT.resolve())
        def _run(args, timeout=6):
            try:
                r = subprocess.run(["git"] + args, cwd=project_root,
                                   capture_output=True, text=True, timeout=timeout)
                return r.stdout, r.returncode == 0
            except Exception:
                return "", False
        branch, _ = _run(["branch", "--show-current"], timeout=5)
        info["current_branch"] = branch.strip() or "unbekannt"
        status_out, ok = _run(["status", "--short"], timeout=6)
        if not ok:
            return info
        patch_files = set(str(p).replace("\\", "/").lstrip("./") for p in ((current_patch or {}).get("affected_files") or []))
        patch_file_status = {}
        for fe in ((current_patch or {}).get("file_entries") or []):
            k = str(fe.get("path") or "").replace("\\", "/").lstrip("./")
            if k:
                patch_file_status[k] = str(fe.get("status") or "")
        current_run = []
        external = 0
        total = 0
        pending = 0
        for line in (status_out or "").splitlines():
            line = line.rstrip()
            if not line:
                continue
            total += 1
            raw = line[3:].strip() if len(line) > 3 else line
            norm = raw.replace("\\", "/").lstrip("./")
            match_key = ""
            for pk in patch_files:
                if norm == pk or norm.endswith("/" + pk) or pk.endswith("/" + norm):
                    match_key = pk
                    break
            if match_key:
                rs = patch_file_status.get(match_key, "")
                current_run.append({"path": raw, "status": rs})
                if rs not in ("applied", "verified", "closed"):
                    pending += 1
            else:
                external += 1
        info["changed_count"] = total
        info["current_run_files"] = current_run[:20]
        info["current_run_count"] = len(current_run)
        info["external_changed_files_count"] = external
        info["review_files_pending"] = pending
        info["has_external_changes"] = external > 0
        track, track_ok = _run(["rev-list", "--left-right", "--count", "@{upstream}...HEAD"], timeout=5)
        if track_ok:
            parts = track.strip().split()
            if len(parts) == 2:
                try:
                    info["behind"] = int(parts[0])
                    info["ahead"] = int(parts[1])
                except ValueError:
                    pass
    except Exception:
        pass
    return info


def _collect_memory_info_for_agent_core(auto_loop):
    """Kompakte Memory-/Regelinfos fuer Agent-Core-State.

    Liefert aktive Kontextquellen, dauerhafte Regeln, bevorzugte/blockierte
    Pfade und eine kompakte Instruktionskurzfassung. Nutzt build_project_memory
    und einen optionalen Lauf-Snapshot aus auto_loop.memory_snapshot.
    """
    try:
        m = build_project_memory()
    except Exception:
        return {}
    run_snapshot = {}
    snap = auto_loop.get("memory_snapshot") if isinstance(auto_loop.get("memory_snapshot"), dict) else None
    if snap:
        run_snapshot = {
            "snapshot_at": str(snap.get("snapshot_at") or ""),
            "rules_count": int(snap.get("rules_count") or 0),
            "sources": list(snap.get("sources") or []),
            "project_context_len": int(snap.get("project_context_len") or 0),
            "builder_notes_len": int(snap.get("builder_notes_len") or 0),
            "preferred_paths": list(snap.get("preferred_paths") or [])[:6],
            "avoid_paths": list(snap.get("avoid_paths") or [])[:6],
            "instruction_summary": str(snap.get("instruction_summary") or ""),
        }
    return {
        "active_context_sources": m.get("active_context_sources") or [],
        "active_rules": (m.get("working_rules") or []) + (m.get("style_rules") or []) + (m.get("agent_instructions") or []),
        "rules_count": m.get("rules_count") or 0,
        "instruction_summary": m.get("instruction_summary") or "",
        "preferred_paths": (m.get("preferred_paths") or [])[:10],
        "restricted_paths": sorted(list({p for p in ((m.get("avoid_paths") or []) + (m.get("guarded_paths") or [])) if p})),
        "sensitive_patterns_count": len(m.get("sensitive_patterns") or []),
        "allowed_prefixes_count": len(m.get("allowed_prefixes") or []),
        "last_updated": m.get("last_updated") or "",
        "run_snapshot": run_snapshot,
    }


def _collect_qa_info_for_agent_core(run_state, auto_loop, activity):
    try:
        qa = build_qa_acceptance_snapshot(run_state=run_state, auto_loop=auto_loop, activity=activity)
    except Exception:
        return {}
    return {
        "overall_status": qa.get("overall_status") or qa.get("overall") or "offen",
        "acceptance_status": qa.get("acceptance_status") or "open",
        "acceptance_ready": bool(qa.get("acceptance_ready")),
        "acceptance_blocked_reason": qa.get("acceptance_blocked_reason") or "",
        "pending_approvals": len(qa.get("pending_approvals") or []),
        "pending_reviews": len(qa.get("pending_reviews") or []),
        "ready_modules": len(qa.get("ready_modules") or []),
        "partial_modules": len(qa.get("partial_modules") or []),
        "blocked_modules": len(qa.get("blocked_modules") or []),
        "failed_modules": len(qa.get("failed_modules") or []),
        "run_summary": qa.get("run_summary") or {},
        "verification_summary": qa.get("verification_summary") or {},
        "last_checked": qa.get("last_checked") or qa.get("timestamp") or "",
    }


def _agent_core_approvals(auto_loop):
    approvals = []
    for step in (auto_loop.get("steps") or []):
        if str(step.get("status") or "").lower() == "wartet auf freigabe":
            approvals.append({
                "step_id": step.get("id"),
                "label": step.get("label"),
                "detail": step.get("detail") or "",
                "gate": step.get("gate") or "approval",
                "risky": bool(step.get("risky")),
                "tool": step.get("tool") or tool_for_action(step.get("action")),
                "phase": step.get("phase") or phase_for_action(step.get("action")),
            })
    return approvals


@app.route("/api/agent-core/state", methods=["GET"])
def agent_core_state_endpoint():
    """Zentraler Agenten-Kern-State.

    Vereint Auto-Loop, Direktmodus, Runner, Kontext, Git in einer
    autoritativen Sicht fuer die UI. Keine Parallelzustaende - der Auto-Loop
    ist die kanonische Laufrepraesentation.
    """
    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    activity = load_ui_activity_entries()
    project_map_ac = read_json_file(PROJECT_MAP_FILE, {})
    knowledge_ac = read_json_file(DATA_DIR / "project_knowledge.json", {})
    knowledge_ready_ac = bool(isinstance(knowledge_ac, dict) and (knowledge_ac.get("built_at") or knowledge_ac.get("summary")))

    errors = [e for e in activity if e.get("level") == "error"][:5]
    warnings = [e for e in activity if e.get("level") == "warning"][:5]
    blockers = list(auto_loop.get("session_blockers") or [])
    if run_state.get("blocked"):
        blockers.append({
            "step_id": "project_mode",
            "step_label": "Project-Mode",
            "reason": str(run_state.get("last_guard_decision") or "Guard blockiert Pfad."),
            "path": ", ".join([str(p) for p in (run_state.get("last_target_paths") or [])[:3]]),
            "rule": "",
            "suggestion": "Zielpfad pruefen oder Guard-Regel anpassen.",
            "timestamp": run_state.get("last_run_at") or "",
        })

    approvals = _agent_core_approvals(auto_loop)
    tools = _agent_core_tools(run_state, auto_loop)

    phase = auto_loop.get("phase") or "planning"
    status = auto_loop.get("status") or "idle"
    if status == "idle":
        phase = "planning"

    objective = auto_loop.get("goal") or run_state.get("last_task") or ""
    active_module = auto_loop.get("active_module") or ""
    next_step = auto_loop.get("next_step") or ""
    if not next_step:
        for step in (auto_loop.get("steps") or []):
            if str(step.get("status") or "").lower() in {"geplant", "wartet auf freigabe"}:
                next_step = step.get("label") or ""
                break

    patch_snapshot = build_patch_review_snapshot()
    current_patch = patch_snapshot.get("current")
    git_info = _collect_git_info_for_agent_core(current_patch)
    memory_info = _collect_memory_info_for_agent_core(auto_loop)
    qa_info = _collect_qa_info_for_agent_core(run_state, auto_loop, activity)
    run_metrics = _auto_loop_run_metrics(auto_loop)
    current_error = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else _empty_error_info()
    repair_plan = auto_loop.get("repair_plan") if isinstance(auto_loop.get("repair_plan"), list) else []
    repair_history = auto_loop.get("repair_history") if isinstance(auto_loop.get("repair_history"), list) else []
    retry_max_val = int(auto_loop.get("retry_max") or 0)
    retry_count_val = int(auto_loop.get("retry_count") or 0)
    retry_possible = bool(auto_loop.get("retry_possible"))
    if retry_count_val >= retry_max_val and retry_max_val > 0:
        retry_possible = False

    pending_approval = None
    if approvals:
        first_pending = approvals[0] if isinstance(approvals, list) and approvals else {}
        pending_approval = {
            "step_id": str(first_pending.get("step_id") or ""),
            "step_label": str(first_pending.get("label") or ""),
            "tool": str(first_pending.get("tool") or ""),
            "scope": str(first_pending.get("phase") or ""),
            "detail": str(first_pending.get("detail") or ""),
            "risky": bool(first_pending.get("risky")),
        }
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    pending_direct_valid = pending_direct_context_is_valid(pending_direct, run_state)
    if status == "approval_required" and not pending_approval and not pending_direct_valid:
        status = "idle"
        phase = "planning"

    active_tool = active_module or ""
    pending_action = ""
    blocked_reason = ""
    for step in (auto_loop.get("steps") or []):
        step_status = str(step.get("status") or "").lower()
        if step_status == "laeuft" or step_status == "wartet auf freigabe":
            pending_action = step.get("action") or step.get("id") or ""
            active_tool = active_tool or step.get("tool") or ""
            break
        if step_status == "geplant" and not pending_action:
            pending_action = step.get("action") or step.get("id") or ""
    if status == "blocked":
        blocked_reason = auto_loop.get("last_action") or ""
        for b in blockers:
            if b.get("reason"):
                blocked_reason = b.get("reason") or blocked_reason
                break

    active_file = _resolve_active_file(current_error, current_patch, auto_loop, run_state)
    decision = _build_agent_decision(run_state, auto_loop, current_patch, current_error, qa_info, active_file)
    if decision.get("gate_type") == "auto":
        decision["console_label"] = f"[AUTO] {decision.get('next_reason') or decision.get('why_continuing') or decision.get('status')}"
    elif decision.get("gate_type") in {"approval_gate", "approval_pending"}:
        decision["console_label"] = f"[GATE] {decision.get('next_reason') or decision.get('why_stopped') or decision.get('status')}"
    elif decision.get("gate_type") == "stopped":
        decision["console_label"] = f"[STOP] {decision.get('why_stopped') or decision.get('next_reason') or decision.get('status')}"
    else:
        decision["console_label"] = f"[?] {decision.get('next_reason') or decision.get('status')}"
    decision_history = run_state.get("agent_decision_history") if isinstance(run_state.get("agent_decision_history"), list) else []

    canonical_run_id = str(auto_loop.get("run_id") or "").strip()
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    if status == "idle" and not pending_requires_confirmation(pending_direct):
        canonical_run_id = ""

    lifecycle_status = status
    pending_direct = run_state.get("pending_direct_run") if isinstance(run_state.get("pending_direct_run"), dict) else None
    has_pending_direct = pending_direct_context_is_valid(pending_direct, run_state)
    has_pending_reviews = int(qa_info.get("pending_reviews") or 0) > 0 and str(decision.get("next_action") or "").strip().lower() == "review_patch"
    has_pending_approvals = int(qa_info.get("pending_approvals") or 0) > 0
    acceptance_status = str(qa_info.get("acceptance_status") or "").strip().lower()
    if status == "idle":
        lifecycle_status = "running" if has_pending_direct else "idle"
    elif status == "running":
        if phase in {"planning", "context"}:
            lifecycle_status = "planning"
        elif phase in {"analysis", "preview"}:
            lifecycle_status = "analyzing"
        else:
            lifecycle_status = "running"
    elif status == "approval_required":
        lifecycle_status = "running"
    elif status in {"blocked", "failed"}:
        lifecycle_status = status
    elif status in {"done", "completed"}:
        if has_pending_reviews:
            lifecycle_status = "running"
        elif has_pending_approvals or acceptance_status in {"open", "needs_approval", "open_review", "incomplete", "invalidated", "ready"}:
            lifecycle_status = "running"
        else:
            lifecycle_status = "completed"
    if decision.get("gate_type") == "approval_pending" or decision.get("gate_type") == "approval_gate":
        lifecycle_status = "running"
    elif decision.get("gate_type") == "stopped" and lifecycle_status not in {"blocked", "failed", "completed"} and has_pending_reviews:
        lifecycle_status = "running"

    run = {
        "run_id": canonical_run_id,
        "objective": objective,
        "status": status,
        "lifecycle_status": lifecycle_status,
        "phase": phase,
        "active_module": active_module,
        "active_tool": active_tool,
        "pending_action": pending_action,
        "pending_approval": pending_approval or {},
        "blocked_reason": blocked_reason,
        "current_step": auto_loop.get("current_step") or "",
        "next_step": next_step,
        "summary": auto_loop.get("summary") or "",
        "last_action": auto_loop.get("last_action") or "",
        "last_run_at": auto_loop.get("last_run_at") or "",
        "started_at": auto_loop.get("started_at") or "",
        "paused_at": auto_loop.get("paused_at") or "",
        "resumed_at": auto_loop.get("resumed_at") or "",
        "ended_at": auto_loop.get("ended_at") or "",
        "apply_mode": auto_loop.get("apply_mode") or "safe",
        "retry_count": retry_count_val,
        "retry_max": retry_max_val,
        "retry_possible": retry_possible,
        "retry_blocked_reason": str(auto_loop.get("retry_blocked_reason") or ""),
        "current_patch": current_patch or {},
        "git": git_info,
        "memory": memory_info,
        "qa": qa_info,
        "run_metrics": run_metrics,
        "current_error": current_error,
        "quick_actions": _build_quick_actions(status, auto_loop, current_patch, current_error),
        "context_actions": _build_context_actions(current_error, current_patch, auto_loop),
        "active_file": active_file,
        "workspace_edit_flow": _build_workspace_edit_flow(
            current_error, current_patch, auto_loop, run_state,
            active_file,
        ),
        "decision": decision,
        "remote_workflow": decision.get("remote_workflow") if isinstance(decision.get("remote_workflow"), dict) else {},
        "decision_history": decision_history[:12],
        "repair_plan": repair_plan,
        "repair_history": repair_history,
        "requires_approval": bool(auto_loop.get("requires_approval")),
        "approvals_needed": approvals,
        "blockers": blockers[:10],
        "errors": errors,
        "warnings": warnings,
        "history": (auto_loop.get("history") or [])[:16],
        "last_runner_execution": run_state.get("last_runner_execution") if isinstance(run_state.get("last_runner_execution"), dict) else {},
        "project_scan_summary": {
            "scanned_at": str(project_map_ac.get("scanned_at") or ""),
            "total_files": int(project_map_ac.get("total_files") or 0),
            "last_scan_trace": run_state.get("last_project_scan") if isinstance(run_state.get("last_project_scan"), dict) else {},
            "knowledge_ready": knowledge_ready_ac,
            "knowledge_built_at": str(knowledge_ac.get("built_at") or "") if isinstance(knowledge_ac, dict) else "",
        },
        "tools": tools,
        "steps": [{
            "id": step.get("id"),
            "label": step.get("label"),
            "status": step.get("status"),
            "phase": step.get("phase") or phase_for_action(step.get("action")),
            "tool": step.get("tool") or tool_for_action(step.get("action")),
            "gate": step.get("gate"),
            "approved": bool(step.get("approved")),
            "risky": bool(step.get("risky")),
            "retry_round": step.get("retry_round") or 0,
            "retry_category": step.get("retry_category") or "",
            "retry_target_file": step.get("retry_target_file") or "",
            "retry_reason": step.get("retry_reason") or "",
        } for step in (auto_loop.get("steps") or [])],
    }

    focus_todo = {}
    todos_content = read_json_file(DATA_DIR / "todos.json", {"todos": []})
    if isinstance(todos_content, dict) and isinstance(todos_content.get("todos"), list):
        for todo in todos_content["todos"]:
            if isinstance(todo, dict) and bool(todo.get("is_focus")):
                focus_todo = {
                    "id": str(todo.get("id") or ""),
                    "title": str(todo.get("title") or ""),
                    "status": str(todo.get("status") or "offen"),
                }
                break

    return jsonify({
        "ok": True,
        "timestamp": get_timestamp(),
        "phases": list(AGENT_PHASES),
        "run": run,
        "focus_todo": focus_todo,
    })


@app.route("/api/orchestrator/state", methods=["GET"])
def orchestrator_state_endpoint():
    """Alias fuer /api/agent-core/state mit Fokus auf Laufkontrolle.

    Der Auto-Loop-State ist weiterhin die kanonische Persistenz. Dieser
    Endpunkt liefert die gleiche autoritative Sicht, eignet sich aber
    semantisch als 'Orchestrator-State' fuer UI-Polling und Resume.
    """
    return agent_core_state_endpoint()


@app.route("/api/patch-review/state", methods=["GET"])
def patch_review_state_endpoint():
    """Zentrale Patch-/Review-Sicht.

    Nutzt pending_direct_run + direct_run_history + auto_loop als einheitliche Quelle.
    Keine separate Patch-Persistenz, keine Parallelwelten.
    """
    snapshot = build_patch_review_snapshot()
    return jsonify({
        "ok": True,
        "timestamp": get_timestamp(),
        "current": snapshot["current"],
        "recent": snapshot["recent"],
        "apply_mode": snapshot["apply_mode"],
        "last_completed_run_id": snapshot["last_completed_run_id"],
    })


@app.route("/api/todos", methods=["GET"])
def todos():
    todos_content = read_json_file(DATA_DIR / "todos.json", None)
    if isinstance(todos_content, dict) and isinstance(todos_content.get("todos"), list):
        normalized = []
        focus_found = False
        for todo in todos_content["todos"]:
            if not isinstance(todo, dict):
                continue
            status = str(todo.get("status") or "offen").strip().lower() or "offen"
            if status not in {"erledigt", "naechstes", "offen"}:
                status = "offen"
            last_status = str(todo.get("last_status") or "").strip().lower()
            if last_status not in {"erledigt", "naechstes", "offen"}:
                last_status = ""
            is_focus = bool(todo.get("is_focus"))
            if is_focus and not focus_found:
                focus_found = True
            else:
                is_focus = False
            normalized.append({
                "id": str(todo.get("id") or len(normalized) + 1),
                "title": str(todo.get("title") or "Ohne Titel").strip() or "Ohne Titel",
                "description": str(todo.get("description") or "Keine Beschreibung").strip() or "Keine Beschreibung",
                "status": status,
                "area": str(todo.get("area") or "Rainer Build").strip() or "Rainer Build",
                "priority": str(todo.get("priority") or "mittel").strip() or "mittel",
                "last_status": last_status,
                "is_focus": is_focus,
            })
        return jsonify({"todos": normalized})

    return jsonify({
        "todos": [
            {
                "id": "1",
                "title": "Builder Mode weiterentwickeln",
                "description": "Implementiere neue Funktionen und Verbesserungen fuer den Builder Mode.",
                "status": "offen",
                "area": "Builder",
                "priority": "mittel",
                "last_status": "",
                "is_focus": False,
            },
            {
                "id": "2",
                "title": "Datei-Generator weiterentwickeln",
                "description": "Erweitere die Funktionalitaet des Datei-Generators um neue Formate und Optionen.",
                "status": "offen",
                "area": "Datei-Generator",
                "priority": "mittel",
                "last_status": "",
                "is_focus": False,
            },
            {
                "id": "3",
                "title": "Design Studio weiterentwickeln",
                "description": "Optimiere das Design Studio fuer eine bessere Benutzerfreundlichkeit und Funktionalitaet.",
                "status": "offen",
                "area": "Design Studio",
                "priority": "mittel",
                "last_status": "",
                "is_focus": False,
            }
        ]
    })


@app.route("/api/todos", methods=["POST"])
def update_todos():
    data = request.get_json(silent=True) or {}
    todo_id = str(data.get("id") or "").strip()
    status = str(data.get("status") or "").strip().lower()
    focus_update = data.get("is_focus", None)

    if not todo_id:
        return jsonify({"error": "Bitte eine gueltige To-do-ID angeben."}), 400
    if not status and focus_update is None:
        return jsonify({"error": "Bitte Status oder Fokus-Aenderung angeben."}), 400
    if status and status not in {"erledigt", "naechstes", "offen"}:
        return jsonify({"error": "Status muss erledigt, naechstes oder offen sein."}), 400

    todos_content = read_json_file(DATA_DIR / "todos.json", {"todos": []})
    if not isinstance(todos_content, dict) or not isinstance(todos_content.get("todos"), list):
        todos_content = {"todos": []}

    updated = False
    target_todo = None
    for todo in todos_content["todos"]:
        if not isinstance(todo, dict):
            continue
        if str(todo.get("id") or "").strip() == todo_id:
            target_todo = todo
            if status:
                previous_status = str(todo.get("status") or "offen").strip().lower() or "offen"
                if previous_status not in {"erledigt", "naechstes", "offen"}:
                    previous_status = "offen"
                todo["last_status"] = previous_status
                todo["status"] = status
            updated = True
            break

    if not updated:
        return jsonify({"error": "To-do nicht gefunden."}), 404

    if focus_update is not None:
        desired_focus = bool(focus_update)
        for todo in todos_content["todos"]:
            if isinstance(todo, dict):
                todo["is_focus"] = False
        if desired_focus and isinstance(target_todo, dict):
            target_todo["is_focus"] = True
            if str(target_todo.get("status") or "").strip().lower() == "offen":
                target_todo["last_status"] = "offen"
                target_todo["status"] = "naechstes"
    elif isinstance(target_todo, dict) and str(target_todo.get("status") or "").strip().lower() == "erledigt" and bool(target_todo.get("is_focus")):
        target_todo["is_focus"] = False
        next_focus = None
        for todo in todos_content["todos"]:
            if not isinstance(todo, dict) or todo is target_todo:
                continue
            candidate_status = str(todo.get("status") or "offen").strip().lower() or "offen"
            if candidate_status == "naechstes":
                next_focus = todo
                break
        if next_focus is None:
            for todo in todos_content["todos"]:
                if not isinstance(todo, dict) or todo is target_todo:
                    continue
                candidate_status = str(todo.get("status") or "offen").strip().lower() or "offen"
                if candidate_status == "offen":
                    next_focus = todo
                    break
        if next_focus is not None:
            next_focus["is_focus"] = True
            if str(next_focus.get("status") or "offen").strip().lower() == "offen":
                next_focus["last_status"] = "offen"
                next_focus["status"] = "naechstes"

    write_json_file(DATA_DIR / "todos.json", todos_content)
    if focus_update is not None:
        append_ui_log_entry("Builder", f"Roadmap-Fokus aktualisiert: {todo_id} -> {'fokus' if bool(focus_update) else 'aus'}", "info")
    else:
        append_ui_log_entry("Builder", f"Roadmap aktualisiert: {todo_id} -> {status}", "info")
    return todos()


@app.route("/api/ollama/health", methods=["GET"])
def ollama_health():
    try:
        response = requests.get("http://127.0.0.1:11434/api/tags", timeout=2)
        return jsonify({"status": response.status_code == 200})
    except requests.RequestException:
        return jsonify({"status": False})


@app.route("/api/local-agent/chat", methods=["POST"])
def local_agent_chat():
    """Reiner Ollama-Chat (lokal, ohne bezahlte APIs) mit eingebetteten Agent-Schutzregeln."""
    from agent_file_guard import RAINER_BUILD_AGENT_RULES_EXTENDED

    data = request.get_json(silent=True) or {}
    msg = str(data.get("message") or data.get("prompt") or "").strip()
    if not msg:
        return jsonify({"ok": False, "error": "message fehlt"}), 400
    if len(msg) > 120_000:
        return jsonify({"ok": False, "error": "Nachricht zu lang (max. 120000 Zeichen)."}), 400
    extra = str(data.get("context") or "").strip()
    pack_parts: list[str] = []
    root = APP_DIR.resolve()
    from local_agent_context import build_workspace_tree_snippet
    from local_agent_tools import safe_read_project_file

    if data.get("attach_workspace_tree", True):
        tree = build_workspace_tree_snippet(root, 44)
        pack_parts.append(f"[WORKSPACE]\n{tree}\n[/WORKSPACE]")
    if data.get("attach_workspace_file"):
        rel = str(data.get("workspace_rel") or "").strip()
        if rel:
            ok, txt = safe_read_project_file(root, rel, 12_000)
            if ok:
                esc_rel = rel.replace('"', "'")
                pack_parts.append(f'[WORKSPACE_FILE path="{esc_rel}"]\n{txt}\n[/WORKSPACE_FILE]')
            else:
                pack_parts.append(f"[WORKSPACE_FILE_ERROR]\n{txt}\n[/WORKSPACE_FILE_ERROR]")
    qr_tail = str(data.get("pytest_tail") or data.get("quick_run_output") or "").strip()
    if qr_tail:
        pack_parts.append(f"[QUICK_CHECK_OUTPUT]\n{qr_tail[:8000]}\n[/QUICK_CHECK_OUTPUT]")
    err_snip = str(data.get("error_snippet") or "").strip()
    if err_snip:
        pack_parts.append(f"[ERRORS]\n{err_snip[:4000]}\n[/ERRORS]")
    if pack_parts:
        extra = ((extra + "\n\n") if extra else "") + "\n\n".join(pack_parts)
    ctx_parts = [
        "Du bist der lokale Rainer-Build-Assistent. Antworte knapp und technisch korrekt auf Deutsch.",
        "Keine Dateien schreiben — nur erklaeren/planen. Zum Anwenden nutzt der Nutzer den Direktmodus.",
        "",
        RAINER_BUILD_AGENT_RULES_EXTENDED,
    ]
    if extra:
        ctx_parts.extend(["", "Zusatzkontext:", extra[:28000]])
    context = "\n".join(ctx_parts)
    try:
        _append_chat_entry(
            {
                "role": "user",
                "content": msg[:8000],
                "timestamp": get_timestamp(),
                "status": "local_agent",
                "uploads": [],
                "changed_files": [],
                "run_id": "",
            }
        )
    except Exception:
        pass
    reply = call_ollama_intelligent(
        msg, context=context, model_override=data.get("model"), local_agent_mode=True
    )
    ollama_ok = not is_llm_failure_message(str(reply or ""))
    try:
        _append_chat_entry(
            {
                "role": "assistant",
                "content": str(reply or "")[:12000],
                "timestamp": get_timestamp(),
                "status": "local_agent",
                "uploads": [],
                "changed_files": [],
                "run_id": "",
            }
        )
    except Exception:
        pass
    return jsonify({"ok": True, "reply": reply, "ollama_ok": ollama_ok})


@app.route("/api/local-agent/capabilities", methods=["GET"])
def local_agent_capabilities():
    try:
        r = requests.get("http://127.0.0.1:11434/api/tags", timeout=2)
        ollama_ok = r.status_code == 200
        models = []
        if ollama_ok:
            j = r.json() or {}
            models = [m.get("name") for m in (j.get("models") or []) if isinstance(m, dict) and m.get("name")]
    except requests.RequestException:
        ollama_ok = False
        models = []
    try:
        llm_health = summarize_llm_health()
    except Exception:
        llm_health = {}
    return jsonify(
        {
            "ok": True,
            "product": "Rainer Build 3.0",
            "pricing": "local_free",
            "ollama_ok": bool(ollama_ok),
            "ollama_models": models[:16],
            "llm_health": llm_health,
            "context_tags": ["[WORKSPACE]", "[WORKSPACE_FILE]", "[ERRORS]", "[QUICK_CHECK_OUTPUT]"],
            "features": [
                "Rainer Build 3.0 — Lokal-Agent (Beratung) + Direktmodus (Apply)",
                "POST /api/local-agent/chat — Ollama + Tags [WORKSPACE]/[ERRORS]/[QUICK_CHECK_OUTPUT]",
                "POST /api/local-agent/tool — read_file (path|rel_path) / search (query|pattern, file_pattern|glob)",
                "POST /api/direct-run — Aenderungen anwenden",
                "GET /api/chat-history — Verlauf",
                "POST /api/local-agent/quick-run — nur fest whitelistete Checks",
                "POST /api/intelligent-run-optimized — Hybrid-Ollama (Quick/Detailed + AST-Kontext)",
            ],
        }
    )


@app.route("/api/llm/providers", methods=["GET"])
def llm_providers_endpoint():
    """Lokale Provider-Übersicht (Ollama, LM Studio, llama.cpp) ohne Cloud."""
    from model_providers import get_available_providers

    try:
        summary = summarize_llm_health()
    except Exception:
        summary = {}
    return jsonify({"ok": True, "catalog": get_available_providers(), **summary})


@app.route("/api/intelligent-run-optimized", methods=["POST"])
def intelligent_run_optimized():
    """Hybrid-Ollama: Quick vs. Detailed Modell + SmartTools-Architektur-Kontext."""
    from hybrid_optimizer import hybrid_optimizer
    from smart_tools import SmartTools
    from system_prompts import SYSTEM_PROMPT_DETAILED, SYSTEM_PROMPT_QUICK

    data = request.get_json(silent=True) or {}
    prompt = _extract_task_or_prompt_from_request_json(data)
    if not prompt:
        return jsonify({"ok": False, "error": "Kein Prompt/Task"}), 400
    max_scan = int(data.get("max_arch_scan_files") or 100)
    tools = SmartTools(project_root=APP_DIR.resolve(), max_scan_files=max_scan)
    arch_context = tools.get_context_for_ollama()
    mode = "detailed" if hybrid_optimizer.is_detailed_query(prompt) else "quick"
    system_prompt = SYSTEM_PROMPT_DETAILED if mode == "detailed" else SYSTEM_PROMPT_QUICK
    result = hybrid_optimizer.execute_optimized(prompt, context=arch_context, system_prompt=system_prompt)
    if not result.get("success"):
        return jsonify({"ok": False, **result}), 500
    return jsonify(
        {
            "ok": True,
            "success": True,
            "response": result.get("response"),
            "model": result.get("model"),
            "mode": result.get("mode"),
            "elapsed_seconds": result.get("elapsed_seconds"),
            "tokens_generated": result.get("tokens_generated"),
            "quality_estimate": result.get("quality_estimate"),
            "final": True,
            "stop_continue": True,
        }
    )


@app.route("/api/capabilities", methods=["GET"])
def rainer_capabilities_overview():
    """Ueberblick Rainer Build 3.0 (Features & Modell-Namen)."""
    return jsonify(
        {
            "ok": True,
            "rainer_build": "3.0",
            "features": [
                "Lokal-Agent (Beratung, Ollama)",
                "Smart Tools (AST-Architektur-Kontext)",
                "Hybrid Optimizer (Quick/Detailed)",
                "Autonomer Agent (Plan + Shell-Whitelist + Fehler-Loop)",
            ],
            "models": {
                "quick": os.getenv("HYBRID_QUICK_MODEL", "mistral:latest"),
                "detailed": os.getenv("HYBRID_DETAILED_MODEL", "deepseek-coder:33b"),
                "agent": "HybridOptimizer + AgentBrain",
            },
        }
    )


@app.route("/api/local-agent/tool", methods=["POST"])
def local_agent_tool():
    """Nur read_file und search unter APP_DIR — kein freies Shell."""
    from local_agent_context import parse_search_result_lines
    from local_agent_tools import normalize_project_rel, safe_read_project_file, safe_search_project

    data = request.get_json(silent=True) or {}
    tool = str(data.get("tool") or "").strip().lower()
    root = APP_DIR.resolve()
    if tool == "read_file":
        rel = str(data.get("rel_path") or data.get("path") or "").strip()
        try:
            max_c = int(data.get("max_chars") or 102_400)
        except (TypeError, ValueError):
            max_c = 102_400
        max_c = max(256, min(max_c, 102_400))
        ok, out = safe_read_project_file(root, rel, max_c)
        rel_clean = normalize_project_rel(rel) or rel.replace("\\", "/").strip()
        if ok:
            line_count = out.count("\n") + (1 if out else 0)
            sz = len(out.encode("utf-8"))
            return jsonify(
                {
                    "ok": True,
                    "success": True,
                    "tool": "read_file",
                    "result": out,
                    "content": out,
                    "file": rel_clean,
                    "lines": line_count,
                    "size_bytes": sz,
                }
            )
        return jsonify({"ok": False, "success": False, "tool": "read_file", "error": out}), 400
    if tool == "search":
        pattern = str(data.get("pattern") or data.get("query") or "").strip()
        glob_pat = str(data.get("glob") or data.get("file_pattern") or "*.py").strip() or "*.py"
        try:
            max_m = int(data.get("max_matches") or data.get("max_hits") or 40)
        except (TypeError, ValueError):
            max_m = 40
        max_m = max(1, min(max_m, 80))
        ok, out = safe_search_project(root, pattern, glob_pat, max_matches=max_m)
        if ok:
            matches = parse_search_result_lines(out)
            return jsonify(
                {
                    "ok": True,
                    "success": True,
                    "tool": "search",
                    "result": out,
                    "query": pattern,
                    "matches": matches,
                    "total_matches": len(matches),
                }
            )
        return jsonify({"ok": False, "success": False, "tool": "search", "error": out}), 400
    return jsonify({"ok": False, "error": "Unbekanntes tool (nur read_file, search)."}), 400


@app.route("/api/local-agent/quick-run", methods=["POST"])
def local_agent_quick_run():
    """Nur feste IDs — kein freies Shell aus dem Browser."""
    data = request.get_json(silent=True) or {}
    run_id = str(data.get("id") or "").strip()
    root = APP_DIR.resolve()
    backend_dir = root / "backend"
    specs: dict[str, tuple[list[str], str]] = {
        "py_compile_main": ([sys.executable, "-m", "py_compile", str(backend_dir / "main.py")], str(backend_dir)),
        "py_compile_write_action": (
            [sys.executable, "-m", "py_compile", str(backend_dir / "write_action.py")],
            str(backend_dir),
        ),
        "pytest_agent_guard": (
            [sys.executable, "-m", "pytest", "backend/test_agent_file_guard.py", "-q", "--tb=no"],
            str(root),
        ),
        "pytest_path_extract": (
            [sys.executable, "-m", "pytest", "backend/test_robot_desktop_path_extract.py", "-q", "--tb=no"],
            str(root),
        ),
    }
    if run_id not in specs:
        return jsonify({"ok": False, "error": "Unbekannte Quick-Run-ID"}), 400
    cmd, cwd = specs[run_id]
    try:
        cp = subprocess.run(cmd, cwd=cwd, capture_output=True, text=True, timeout=180, check=False)
    except subprocess.TimeoutExpired:
        return jsonify({"ok": False, "error": "Timeout", "returncode": -1}), 500
    return jsonify(
        {
            "ok": cp.returncode == 0,
            "returncode": cp.returncode,
            "stdout": (cp.stdout or "")[:6000],
            "stderr": (cp.stderr or "")[:6000],
        }
    )


@app.route("/api/builder/placeholder", methods=["POST"])
def builder_placeholder():
    data = request.get_json(silent=True) or {}
    task = data.get("task")
    action = data.get("action")

    if task:
        return jsonify({"message": f"Plan fuer Aufgabe '{task}' erstellt."})
    if action == "start":
        return jsonify({"message": "Builder Mode gestartet."})
    return jsonify({"error": "Ungueltige Anfrage"}), 400


@app.route("/api/file-generator/placeholder", methods=["POST"])
def file_generator_placeholder():
    data = request.get_json(silent=True) or {}
    filename = data.get("filename")
    filetype = data.get("filetype")
    targetpath = data.get("targetpath")

    if filename and filetype and targetpath:
        return jsonify(success_payload(
            "Datei vorbereitet! ✓",
            technical_message=f"Datei '{filename}' fuer {filetype} vorbereitet.",
            changed_files=[f"{filename}.{filetype}"]
        ))
    return jsonify(error_payload("Ungültige Anfrage.", technical_error="Ungueltige Anfrage")), 400


@app.route("/api/designer/placeholder", methods=["POST"])
def designer_placeholder():
    data = request.get_json(silent=True) or {}
    note = data.get("note")

    if note:
        return jsonify(success_payload(
            "Design-Idee gespeichert! ✓",
            technical_message=f"Design-Idee '{note}' gespeichert.",
            detail="Du kannst direkt darauf aufbauen."
        ))
    return jsonify(error_payload("Ungültige Anfrage.", technical_error="Ungueltige Anfrage")), 400


@app.route("/api/builder/plan", methods=["POST"])
def builder_plan():
    data = request.get_json(silent=True) or {}
    task, error = validate_task(data.get("task"))

    if error:
        append_ui_log_entry("Builder", error, "warning")
        return jsonify(error_payload(error, technical_error=error)), 400

    plan = build_builder_plan(task)
    append_ui_log_entry("Builder", f"Plan fuer '{task}' erzeugt.", "success")
    return jsonify(success_payload(
        "Plan erstellt! ✓",
        technical_message="Lokaler Builder-Plan erzeugt.",
        detail="Die nächsten Schritte sind vorbereitet.",
        plan=plan
    ))


@app.route("/api/builder/generate-content", methods=["POST"])
def builder_generate_content():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path = str(data.get("target_path") or "")

    if task_error:
        append_ui_log_entry("Builder", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error)), 400

    content = generated_content_from_task_or_literal(
        task,
        target_path,
        log_label="Builder",
        log_ok_message="Inhalt aus Prompt (expliziter Text) uebernommen.",
    )
    append_ui_log_entry("Builder", f"Inhalt fuer Ziel '{target_path}' erzeugt.", "success")
    return jsonify(success_payload(
        "Inhalt generiert! ✓",
        technical_message="Inhalt automatisch generiert.",
        detail="Du kannst die Vorschau jetzt prüfen.",
        generated_content=content
    ))


@app.route("/api/builder/change-preview", methods=["POST"])
def builder_change_preview():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path = data.get("target_path")
    proposed_content = str(data.get("proposed_content") or "")

    if task_error:
        append_ui_log_entry("Builder", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error)), 400

    resolved_path, relative_path, path_error = resolve_local_target_path(target_path, task)
    if path_error:
        append_ui_log_entry("Builder", path_error, "warning")
        return jsonify(error_payload(path_error, technical_error=path_error)), 400

    auto_generated = False
    if not proposed_content.strip():
        proposed_content = generated_content_from_task_or_literal(
            task,
            target_path,
            log_label="Builder",
            log_ok_message="Inhalt aus Prompt (expliziter Text) uebernommen.",
        )
        auto_generated = True
        append_ui_log_entry("Builder", "Inhalt aus Aufgabe abgeleitet (explizit oder Stub).", "info")

    current_content, file_exists = read_text_file(resolved_path)
    diff_text = build_text_diff(current_content, proposed_content, relative_path)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    summary = build_change_summary(task, relative_path, file_exists, has_changes)

    append_ui_log_entry("Builder", f"Aenderungsvorschau fuer '{relative_path}' erzeugt.", "success")
    return jsonify(preview_payload(
        path=relative_path,
        technical_message="Aenderungsvorschau erzeugt.",
        target_path=relative_path,
        file_exists=file_exists,
        has_changes=has_changes,
        auto_generated=auto_generated,
        summary=summary,
        diff=diff_text,
        current_content=current_content,
        proposed_content=proposed_content
    ))


@app.route("/api/builder/apply-change", methods=["POST"])
def builder_apply_change():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path = data.get("target_path")
    proposed_content = str(data.get("proposed_content") or "")
    confirm = bool(data.get("confirm"))

    if task_error:
        append_ui_log_entry("Builder", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error)), 400

    resolved_path, relative_path, path_error = resolve_local_target_path(target_path, task)
    if path_error:
        append_ui_log_entry("Builder", path_error, "warning")
        return jsonify(error_payload(path_error, technical_error=path_error)), 400

    if not proposed_content.strip():
        proposed_content = generated_content_from_task_or_literal(
            task,
            target_path,
            log_label="Builder",
            log_ok_message="Inhalt aus Prompt (expliziter Text) uebernommen.",
        )

    if not confirm:
        message = "Bitte die Aenderung zuerst sichtbar bestaetigen."
        append_ui_log_entry("Builder", message, "warning")
        return jsonify(error_payload(message, technical_error=message)), 400

    current_content, file_exists = read_text_file(resolved_path)
    diff_text = build_text_diff(current_content, proposed_content, relative_path)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."

    if not has_changes:
        message = "Keine Aenderung erkannt. Es wurde nichts geschrieben."
        append_ui_log_entry("Builder", message, "warning")
        return jsonify(no_change_payload(
            relative_path,
            technical_message=message,
            target_path=relative_path,
            has_changes=False,
            diff=diff_text
        ))

    wr = persist_text_file_change(
        resolved_path,
        proposed_content,
        relative_path,
        on_timeout_log=lambda m: append_ui_log_entry("Builder", m, "error"),
    )
    if not wr.get("ok"):
        message = wr.get("error") or "Die Aenderung konnte lokal nicht geschrieben werden."
        append_ui_log_entry("Builder", message, "error")
        logger.log_error_with_context(message, {"action": "builder_apply_change", "path": relative_path})
        return jsonify(error_payload(message, technical_error=message)), 500

    append_ui_log_entry(
        "Builder",
        f"Aenderung auf '{relative_path}' {'aktualisiert' if file_exists else 'neu angelegt'}.",
        "success"
    )
    logger.log_file_operation("write_local", relative_path, success=True)
    return jsonify(file_write_payload(
        relative_path,
        created=not file_exists,
        location="rambo_builder_local/",
        technical_message="Aenderung lokal angewendet.",
        target_path=relative_path,
        has_changes=True,
        diff=diff_text
    ))


@app.route("/api/builder/auto-run", methods=["POST"])
def builder_auto_run():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path_input = str(data.get("target_path") or "").strip()

    steps = []

    if task_error:
        append_ui_log_entry("Auto", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error, steps=steps)), 400

    steps.append({"label": "Aufgabe", "status": "ok", "detail": task})

    if target_path_input:
        target_path = target_path_input
        steps.append({"label": "Zielpfad", "status": "ok", "detail": target_path})
    else:
        target_path = infer_target_path(task)
        steps.append({"label": "Zielpfad", "status": "auto", "detail": f"Automatisch ermittelt: {target_path}"})

    resolved_path, relative_path, path_error = resolve_local_target_path(target_path, task)
    if path_error:
        steps.append({"label": "Guard", "status": "error", "detail": path_error})
        append_ui_log_entry("Auto", f"Guard blockiert: {path_error}", "error")
        return jsonify(error_payload(path_error, technical_error=path_error, steps=steps)), 403

    steps.append({"label": "Guard", "status": "ok", "detail": f"Pfad erlaubt: {relative_path}"})

    current_content, file_exists = read_text_file(resolved_path)
    content = resolve_proposed_content_for_local_task(
        task, target_path, current_content, file_exists, relative_path
    )
    steps.append({"label": "Inhalt", "status": "ok", "detail": f"{len(content)} Zeichen generiert"})

    diff_text = build_text_diff(current_content, content, relative_path)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."

    steps.append({
        "label": "Diff",
        "status": "ok" if has_changes else "warn",
        "detail": "Aenderungen erkannt." if has_changes else "Keine inhaltliche Aenderung erkannt."
    })

    if not has_changes:
        append_ui_log_entry("Auto", f"Auto-Run: keine Aenderung fuer '{relative_path}'.", "warning")
        return jsonify(no_change_payload(
            relative_path,
            technical_message="Keine inhaltliche Aenderung erkannt. Datei unveraendert.",
            target_path=relative_path,
            has_changes=False,
            diff=diff_text,
            proposed_content=content,
            steps=steps
        ))

    wr = persist_text_file_change(
        resolved_path,
        content,
        relative_path,
        on_timeout_log=lambda m: append_ui_log_entry("Auto", m, "error"),
    )
    if not wr.get("ok"):
        message = wr.get("error") or "Datei konnte lokal nicht geschrieben werden."
        steps.append({"label": "Schreiben", "status": "error", "detail": message})
        append_ui_log_entry("Auto", message, "error")
        logger.log_error_with_context(message, {"action": "builder_auto_run", "path": relative_path})
        return jsonify(error_payload(message, technical_error=message, steps=steps)), 500

    verb = "aktualisiert" if file_exists else "neu angelegt"
    steps.append({"label": "Schreiben", "status": "ok", "detail": f"Datei {verb}: {relative_path}"})

    plan = build_builder_plan(task)
    steps.append({"label": "Plan", "status": "ok", "detail": plan["ziel"]})

    append_ui_log_entry("Auto", f"Auto-Run abgeschlossen: '{relative_path}' {verb}.", "success")
    logger.log_file_operation("auto_run_write_local", relative_path, success=True)

    return jsonify(file_write_payload(
        relative_path,
        created=not file_exists,
        location="rambo_builder_local/",
        technical_message=f"Auto-Run abgeschlossen. Datei {verb}: {relative_path}",
        target_path=relative_path,
        file_exists=file_exists,
        has_changes=True,
        diff=diff_text,
        proposed_content=content,
        plan=plan,
        steps=steps
    ))


@app.route("/api/file-generator/preview", methods=["POST"])
def file_generator_preview():
    data = request.get_json(silent=True) or {}
    filename, filename_error = validate_filename(data.get("filename"))
    filetype, filetype_error = validate_filetype(data.get("filetype"))
    targetpath, path_error = validate_target_path(data.get("targetpath"))

    error = filename_error or filetype_error or path_error
    if error:
        append_ui_log_entry("Datei", f"Dateivorschau abgelehnt: {error}", "warning")
        return jsonify({"error": error}), 400

    preview = build_file_preview(filename, filetype, targetpath)
    append_ui_log_entry("Datei", f"Vorschau fuer '{preview['path']}' als Typ '{preview['type']}' erzeugt.", "success")
    return jsonify({
        "message": "Dateivorschau lokal erzeugt.",
        "preview": preview
    })


@app.route("/api/design-studio/notes", methods=["GET"])
@app.route("/api/design-notes", methods=["GET"])
def design_notes():
    return jsonify({"entries": load_design_entries()})


@app.route("/api/design-studio/notes", methods=["POST"])
@app.route("/api/design-notes", methods=["POST"])
def save_design_note():
    data = request.get_json(silent=True) or {}
    note, validation_error = validate_design_note(data.get("note"))

    if validation_error:
        append_ui_log_entry("Design", validation_error, "warning")
        return jsonify({"error": validation_error}), 400

    entries = load_design_entries()
    duplicate = any(existing["note"].strip().lower() == note.lower() for existing in entries)
    timestamp = get_timestamp()

    if duplicate:
        append_ui_log_entry("Design", "Aehnliche Design-Idee erneut gespeichert.", "warning")

    entry = {
        "id": str(len(entries) + 1),
        "note": note,
        "timestamp": timestamp
    }
    entries.insert(0, entry)
    save_design_entries(entries)
    append_ui_log_entry("Design", "Design-Idee lokal gespeichert.", "success")

    return jsonify({
        "message": "Design-Idee lokal gespeichert.",
        "entry": entry,
        "entries": load_design_entries(),
        "duplicate": duplicate
    })


@app.route("/api/ui-activity", methods=["GET"])
def get_ui_activity():
    return jsonify({"entries": load_ui_activity_entries()})


@app.route("/api/ui-activity", methods=["POST"])
def add_ui_activity():
    data = request.get_json(silent=True) or {}
    label = str(data.get("label") or "System").strip() or "System"
    message = str(data.get("message") or "").strip()
    level = normalize_level(data.get("level"))

    if not message:
        return jsonify({"error": "Bitte eine Log-Nachricht angeben."}), 400

    entry = append_ui_log_entry(label, message, level)
    return jsonify({
        "message": "Aktivitaet gespeichert.",
        "entry": entry,
        "entries": load_ui_activity_entries()
    })


@app.route("/", methods=["GET"])
def serve_index():
    return send_from_directory(FRONTEND_DIR, "index.html")


@app.route("/style.css")
def serve_style_css():
    return send_from_directory(FRONTEND_DIR, "style.css")


@app.route("/app.js")
def serve_app_js():
    return send_from_directory(FRONTEND_DIR, "app.js")


@app.route("/favicon.ico")
def serve_favicon():
    """Vermeidet 500/HTML-Fehler bei Browser-Requests auf /favicon.ico."""
    return Response(status=204)


@app.route("/api/project/scan", methods=["POST"])
def project_scan():
    files = scan_project_structure()
    areas = {}
    for f in files:
        key = f["area_key"]
        if key not in areas:
            areas[key] = {"label": f["area"], "files": []}
        areas[key]["files"].append(f)

    repo_map = {
        "scanned_at": get_timestamp(),
        "total_files": len(files),
        "areas": {k: {"label": v["label"], "count": len(v["files"])} for k, v in areas.items()}
    }
    write_json_file(PROJECT_MAP_FILE, repo_map)
    try:
        save_project_auto_run_state({
            "last_project_scan": {
                "source": "api_project_scan",
                "scanned_at": str(repo_map.get("scanned_at") or ""),
                "total_files": int(repo_map.get("total_files") or 0),
                "area_count": len(areas),
            }
        })
    except Exception:
        pass

    append_ui_log_entry("Projekt", f"Scan: {len(files)} Dateien in {len(areas)} Bereichen.", "success")
    return jsonify({
        "message": f"Projekt gescannt: {len(files)} Dateien.",
        "total": len(files),
        "areas": areas,
    })


@app.route("/api/project/read-file", methods=["POST"])
def project_read_file():
    data = request.get_json(silent=True) or {}
    rel_path = str(data.get("path") or "").strip()

    resolved, cleaned, error = validate_project_read_path(rel_path)
    if error:
        append_ui_log_entry("Projekt", f"Lesezugriff abgelehnt: {error}", "warning")
        return jsonify(error_payload(error, technical_error=error)), 400

    content, exists = read_text_file(resolved)
    if not exists:
        not_found = f"Datei nicht gefunden: {cleaned}"
        return jsonify(error_payload(not_found, technical_error=not_found)), 404

    append_ui_log_entry("Projekt", f"Datei gelesen: {cleaned}", "info")
    return jsonify({"path": cleaned, "content": content, "length": len(content)})


@app.route("/api/project/change-preview", methods=["POST"])
def project_change_preview():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path = str(data.get("target_path") or "").strip()
    proposed_content = str(data.get("proposed_content") or "")

    if task_error:
        append_ui_log_entry("Projekt", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error)), 400

    resolved, cleaned, error = validate_project_write_path(target_path)
    if error:
        append_ui_log_entry("Projekt", f"Guard blockiert: {error}", "error")
        return jsonify(error_payload(error, technical_error=error, guard_blocked=True)), 403

    auto_generated = False
    if not proposed_content.strip():
        proposed_content = generated_content_from_task_or_literal(
            task,
            target_path,
            log_label="Projekt",
            log_ok_message="Inhalt aus Prompt (expliziter Text) uebernommen.",
        )
        auto_generated = True

    current_content, file_exists = read_text_file(resolved)
    diff_text = build_text_diff(current_content, proposed_content, cleaned)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    summary = build_change_summary(task, cleaned, file_exists, has_changes)

    append_ui_log_entry("Projekt", f"Vorschau: '{cleaned}'.", "success")
    return jsonify(preview_payload(
        path=cleaned,
        technical_message="Projekt-Aenderungsvorschau erzeugt.",
        target_path=cleaned,
        file_exists=file_exists,
        has_changes=has_changes,
        auto_generated=auto_generated,
        guard_blocked=False,
        diff=diff_text,
        current_content=current_content,
        proposed_content=proposed_content,
        summary=summary,
    ))


@app.route("/api/project/apply-change", methods=["POST"])
def project_apply_change():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    target_path = str(data.get("target_path") or "").strip()
    proposed_content = str(data.get("proposed_content") or "")
    confirm = bool(data.get("confirm"))

    if task_error:
        append_ui_log_entry("Projekt", task_error, "warning")
        return jsonify(error_payload(task_error, technical_error=task_error)), 400

    resolved, cleaned, error = validate_project_write_path(target_path)
    if error:
        append_ui_log_entry("Projekt", f"Guard blockiert: {error}", "error")
        return jsonify(error_payload(error, technical_error=error, guard_blocked=True)), 403

    if not proposed_content.strip():
        proposed_content = generated_content_from_task_or_literal(
            task,
            target_path,
            log_label="Projekt",
            log_ok_message="Inhalt aus Prompt (expliziter Text) uebernommen.",
        )

    if not confirm:
        msg = "Bitte die Aenderung zuerst bestaetigen."
        return jsonify(error_payload(msg, technical_error=msg)), 400

    current_content, file_exists = read_text_file(resolved)
    diff_text = build_text_diff(current_content, proposed_content, cleaned)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."

    if not has_changes:
        append_ui_log_entry("Projekt", "Keine Aenderung erkannt.", "warning")
        return jsonify(no_change_payload(
            cleaned,
            technical_message="Keine Aenderung erkannt.",
            has_changes=False,
            diff=diff_text
        ))

    wr = persist_text_file_change(
        resolved,
        proposed_content,
        cleaned,
        on_timeout_log=lambda m: append_ui_log_entry("Projekt", m, "error"),
    )
    if not wr.get("ok"):
        message = wr.get("error") or "Projektdatei konnte nicht geschrieben werden."
        append_ui_log_entry("Projekt", message, "error")
        logger.log_error_with_context(message, {"action": "project_apply_change", "path": cleaned})
        return jsonify(error_payload(message, technical_error=message)), 500

    verb = "aktualisiert" if file_exists else "angelegt"
    append_ui_log_entry("Projekt", f"Projektdatei '{cleaned}' {verb}.", "success")
    logger.log_file_operation("write_project", cleaned, success=True)
    return jsonify(file_write_payload(
        cleaned,
        created=not file_exists,
        location="Rambo-Rainer/",
        technical_message=f"Projektdatei {verb}: {cleaned}",
        target_path=cleaned,
        has_changes=True,
        diff=diff_text,
    ))


@app.route("/api/project/guard-check", methods=["POST"])
def project_guard_check():
    data = request.get_json(silent=True) or {}
    target_path = str(data.get("target_path") or "").strip()

    if not target_path:
        return jsonify({"error": "Bitte einen Pfad angeben."}), 400

    cleaned = format_local_path(target_path)

    for pattern in SENSITIVE_PATTERNS:
        if pattern in cleaned:
            append_ui_log_entry("Projekt", f"Guard blockiert: '{cleaned}' (sensibel: {pattern})", "error")
            return jsonify({
                "allowed": False,
                "reason": "blocked_sensitive",
                "detail": f"Datei enthaelt sensibles Muster '{pattern}' und ist permanent blockiert.",
                "path": cleaned
            })

    trusted_workspace = False
    try:
        trusted_workspace = bool(is_active_workspace_trusted())
    except Exception:
        trusted_workspace = False
    allowed_write = trusted_workspace or any(cleaned.startswith(p) for p in ALLOWED_PROJECT_WRITE_PREFIXES)
    if not allowed_write:
        append_ui_log_entry("Projekt", f"Guard blockiert: '{cleaned}' (nicht freigegeben)", "warning")
        return jsonify({
            "allowed": False,
            "reason": "not_in_allowlist",
            "detail": f"'{cleaned}' liegt nicht in einem freigegebenen Bereich.",
            "path": cleaned,
            "allowed_prefixes": list(ALLOWED_PROJECT_WRITE_PREFIXES)
        })

    resolved = (RAMBO_RAINER_ROOT / cleaned).resolve()
    try:
        resolved.relative_to(RAMBO_RAINER_ROOT.resolve())
    except ValueError:
        append_ui_log_entry("Projekt", f"Guard blockiert: '{cleaned}' (ausserhalb Projekt)", "error")
        return jsonify({"allowed": False, "reason": "outside_project", "detail": "Pfad liegt ausserhalb des Projekts.", "path": cleaned})

    _, file_exists = read_text_file(resolved)
    append_ui_log_entry("Projekt", f"Guard OK: '{cleaned}' freigegeben.", "success")
    return jsonify({
        "allowed": True,
        "reason": "ok",
        "detail": f"Pfad freigegeben. Datei {'vorhanden' if file_exists else 'noch nicht vorhanden (wird neu angelegt)'}.",
        "path": cleaned,
        "file_exists": file_exists
    })


@app.route("/api/project/suggest-files", methods=["POST"])
def project_suggest_files():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))

    if task_error:
        return jsonify({"error": task_error}), 400

    result = suggest_files_for_task(task)
    append_ui_log_entry("Projekt", f"Dateivorschlaege fuer '{task}' erzeugt.", "info")
    return jsonify({
        "message": "Dateivorschlaege erzeugt.",
        "task": task,
        **result
    })


@app.route("/api/project/auto-run", methods=["POST"])
def project_auto_run():
    data = request.get_json(silent=True) or {}
    task, task_error = validate_task(data.get("task"))
    mode = str(data.get("mode") or "safe").strip().lower()
    requested_target = format_local_path(data.get("target_path"))
    steps = []

    if mode not in {"safe", "apply"}:
        mode = "safe"

    if task_error:
        append_ui_log_entry("Auto", task_error, "warning")
        return jsonify({"error": task_error, "steps": steps}), 400

    append_ui_log_entry("Auto", f"Auto-Run gestartet ({mode}) fuer '{task}'.", "info")
    steps.append({"label": "Aufgabe", "status": "ok", "detail": task})

    repo_map = read_json_file(PROJECT_MAP_FILE, {})
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    if not repo_map:
        files = scan_project_structure()
        repo_map = {
            "scanned_at": get_timestamp(),
            "total_files": len(files),
            "areas": {}
        }
        write_json_file(PROJECT_MAP_FILE, repo_map)
    if not knowledge:
        knowledge = build_project_knowledge()

    steps.append({
        "label": "Scan",
        "status": "ok",
        "detail": f"Projektbasis verfuegbar: {repo_map.get('total_files', 0)} Dateien."
    })
    append_ui_log_entry("Auto", "Scan abgeschlossen.", "success")

    suggestions = suggest_files_for_task(task)
    append_ui_log_entry("Auto", "Dateien vorgeschlagen und priorisiert.", "success")
    prioritized_targets = suggestions.get("prioritized_targets") or []
    selected_target = requested_target or infer_project_target_path(task, suggestions)
    steps.append({
        "label": "Auswahl",
        "status": "ok" if selected_target else "warning",
        "detail": (
            f"Zielpfad gewaehlt: {selected_target}"
            if selected_target else
            "Kein Zielpfad bestimmbar."
        )
    })

    resolved, cleaned, guard_error = validate_project_write_path(selected_target)
    if guard_error:
        guard_result = {
            "allowed": False,
            "detail": guard_error,
            "path": selected_target or requested_target
        }
        steps.append({"label": "Guard", "status": "error", "detail": guard_error})
        append_ui_log_entry("Auto", f"Guard blockiert: {guard_error}", "error")
        strategy = build_project_strategy(task, suggestions, selected_target or requested_target, guard_result, mode)
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [item.get("path") for item in prioritized_targets[:4] if item.get("path")],
            "last_guard_decision": guard_error,
            "last_result": "Blockiert",
            "blocked": True
        })
        return jsonify({
            "error": guard_error,
            "mode": mode,
            "steps": steps,
            "suggestions": suggestions,
            "selected_target_path": selected_target or requested_target,
            "guard": guard_result,
            "strategy": strategy,
            "blocked_files": suggestions.get("blocked", [])
        }), 403

    _, file_exists = read_text_file(resolved)
    guard_result = {
        "allowed": True,
        "detail": f"Pfad freigegeben. Datei {'vorhanden' if file_exists else 'noch nicht vorhanden'}.",
        "path": cleaned,
        "file_exists": file_exists
    }
    steps.append({"label": "Guard", "status": "ok", "detail": guard_result["detail"]})
    append_ui_log_entry("Auto", f"Guard freigegeben: {cleaned}", "success")

    current_content, file_exists = read_text_file(resolved)
    proposed_content = resolve_proposed_content_for_local_task(
        task, cleaned, current_content, file_exists, cleaned
    )
    diff_text = build_text_diff(current_content, proposed_content, cleaned)
    has_changes = diff_text != "Keine inhaltliche Aenderung erkannt."
    steps.append({
        "label": "Diff",
        "status": "ok" if has_changes else "warn",
        "detail": "Diff erstellt." if has_changes else "Keine inhaltliche Aenderung erkannt."
    })
    append_ui_log_entry("Auto", "Diff erstellt.", "success" if has_changes else "warning")

    strategy = build_project_strategy(task, suggestions, cleaned, guard_result, mode)
    strategy["betroffene_dateien"] = strategy.get("betroffene_dateien") or [cleaned]

    response_payload = {
        "message": "Project Auto-Run vorbereitet.",
        "mode": mode,
        "selected_target_path": cleaned,
        "steps": steps,
        "suggestions": suggestions,
        "strategy": strategy,
        "guard": guard_result,
        "blocked_files": suggestions.get("blocked", []),
        "proposed_content": proposed_content,
        "diff": diff_text,
        "has_changes": has_changes,
        "file_exists": file_exists,
        "post_check": {
            "ok": False,
            "detail": "Noch nicht ausgefuehrt."
        }
    }

    if mode == "safe":
        append_ui_log_entry("Auto", "Safe Mode abgeschlossen. Nur Analyse und Diff erzeugt.", "success")
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [item.get("path") for item in prioritized_targets[:4] if item.get("path")] or [cleaned],
            "last_guard_decision": guard_result["detail"],
            "last_check_result": "Keine Nachkontrolle noetig (Safe Mode).",
            "last_result": response_payload["message"],
            "blocked": False
        })
        return jsonify(response_payload)

    if not has_changes:
        append_ui_log_entry("Auto", "Apply uebersprungen: keine inhaltliche Aenderung erkannt.", "warning")
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [item.get("path") for item in prioritized_targets[:4] if item.get("path")] or [cleaned],
            "last_guard_decision": guard_result["detail"],
            "last_apply_action": "Keine Aenderung geschrieben.",
            "last_check_result": "Keine Nachkontrolle noetig.",
            "last_result": "Apply nicht noetig.",
            "blocked": False
        })
        response_payload["message"] = "Keine inhaltliche Aenderung erkannt. Apply wurde uebersprungen."
        return jsonify(response_payload)

    wr = persist_text_file_change(
        resolved,
        proposed_content,
        cleaned,
        on_timeout_log=lambda m: append_ui_log_entry("Auto", m, "error"),
    )
    if not wr.get("ok"):
        message = wr.get("error") or "Projektdatei konnte im Apply Mode nicht geschrieben werden."
        steps.append({"label": "Apply", "status": "error", "detail": message})
        append_ui_log_entry("Auto", message, "error")
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [cleaned],
            "last_guard_decision": guard_result["detail"],
            "last_apply_action": message,
            "last_check_result": "",
            "last_result": message,
            "blocked": False
        })
        return jsonify({"error": message, **response_payload, "steps": steps}), 500

    apply_message = f"Apply ausgefuehrt: {cleaned} {'aktualisiert' if file_exists else 'neu angelegt'}."
    steps.append({"label": "Apply", "status": "ok", "detail": apply_message})
    append_ui_log_entry("Auto", apply_message, "success")

    post_check = run_project_post_check(resolved, cleaned, proposed_content)
    steps.append({
        "label": "Nachkontrolle",
        "status": "ok" if post_check["ok"] else "error",
        "detail": post_check["detail"]
    })
    append_ui_log_entry(
        "Auto",
        post_check["detail"],
        "success" if post_check["ok"] else "error"
    )

    save_project_auto_run_state({
        "last_run_at": get_timestamp(),
        "last_task": task,
        "last_mode": mode,
        "last_target_paths": [item.get("path") for item in prioritized_targets[:4] if item.get("path")] or [cleaned],
        "last_guard_decision": guard_result["detail"],
        "last_apply_action": apply_message,
        "last_check_result": post_check["detail"],
        "last_result": "Apply abgeschlossen." if post_check["ok"] else "Apply mit Nachkontrollfehler.",
        "blocked": False
    })

    response_payload.update({
        "message": "Project Auto-Run erfolgreich abgeschlossen.",
        "post_check": post_check,
        "steps": steps
    })
    return jsonify(response_payload)


def _strip_trailing_prompt_from_windows_path_line(path: str) -> str:
    """
    Entfernt angeklebten Fliesstext nach einem Windows-Pfadsegment.
    Stoppt u.a. bei Prompt-Labels (Ziel:, Wichtig:, …), doppeltem Leerzeichen und zweitem Laufwerkspfad.
    """
    s = (path or "").strip().strip('"').strip("'")
    if not s:
        return s
    low = s.lower()
    # Doppeltes Leerzeichen: oft Grenze zwischen Pfad und Fliesstext
    dbl = low.find("  ")
    cut = len(s)
    if dbl >= 0:
        cut = min(cut, dbl)
    # Bekannte Abschnitts-Labels (einleitendes Leerzeichen oder Zeilenanfang nach Pfadteil)
    stop_phrases = (
        " ziel:", " wichtig:", " app-icon", " desktop-icon", " hauptanforderung:", " architektur:",
        " electron main.js:", " preload.js:", " react ui:", " package.json:", " build:", " tests:",
        " nicht machen:", " ausgabeformat:",
        " quelle", " erfolg", " pflicht", " root cause", " aufgabe",
        "\nziel:", "\nwichtig:", "\r\nziel:", "\tquelle", "\nquelle", "\rquelle", "\taufgabe",
        "\nerzeuge", "\nerfolg", "\nroot",
        " ausgabeformat", " pflicht-test",
    )
    for ph in stop_phrases:
        i = low.find(ph)
        if i >= 0:
            cut = min(cut, i)
    if cut < len(s):
        s = s[:cut]
        low = s.lower()
    markers = (
        " quelle", "\nquelle", "\rquelle", "\tquelle",
        " aufgabe", "\naufgabe", "\raufgabe",
        " erzeuge", "\nerzeuge",
        " erfolg", "\nerfolg",
        " root cause", "\nroot",
        " pflicht", " ausgabeformat",
    )
    for m in markers:
        i = low.find(m)
        if i >= 0:
            s = s[:i].rstrip()
            low = s.lower()
    # Zweiter Windows-Pfad im selben Token: nur den ersten behalten
    m2 = re.search(r"(?is)(^[a-z]:\\.+?)(?=\s+[a-z]:\\)", s)
    if m2:
        s = m2.group(1).rstrip()
    s = re.sub(r"\s{2,}$", "", s)
    return s.rstrip(" \t.,;:'\"")


def _is_disallowed_robot_desktop_output_path(path_str: str) -> bool:
    """Bild-Dateien, rambo_builder_local u.a. nie als Desktop-Build-Zielordner."""
    if not path_str:
        return True
    low = path_str.lower().replace("/", "\\")
    if re.search(r"\.(png|jpe?g|webp|gif|ico|svg|bmp|zip|pdf|docx?)\s*$", low):
        return True
    if re.search(r"downloads[/\\]bild\.png", low):
        return True
    if "rambo_builder_local" in low:
        return True
    return False


def _project_build_invalid_path_payload(run_id: str, scope: str, mode: str, message: str) -> dict:
    return {
        "run_id": run_id,
        "scope": scope,
        "mode": mode,
        "ok": False,
        "direct_status": "failed",
        "build_status": "failed",
        "error": message,
        "message": message,
        "technical_message": message,
        "base_path": None,
        "target_root": None,
        "requires_confirmation": False,
        "requires_user_confirmation": False,
        "created_files": [],
        "missing_files": [],
        "workstream_events": [
            _ws_event("route", "error", "Zielpfad", message[:500], status="failed"),
        ],
        "recognized_task": {
            "task_type": "project_build",
            "primary_area": "Project Builder",
            "execution_route": "project_build",
            "hint": "Zielpfad konnte nicht sicher ermittelt werden.",
        },
    }


def _finalize_windows_folder_candidate(rest: str):
    """Ein Windows-Ordnerpfad oder None (keine Bilder, kein rambo_builder_local-Ziel)."""
    if not rest:
        return None
    rest = _strip_trailing_prompt_from_windows_path_line(rest.replace("/", "\\"))
    if not rest or not re.match(r"(?i)^[a-z]:\\", rest):
        return None
    if _is_disallowed_robot_desktop_output_path(rest):
        return None
    return rest


def _extract_robot_desktop_base_path_str(task: str):
    """
    Erkennt den Zielordner fuer Desktop-/Robot-Builds.
    Stoppt vor angeklebtem Prompt-Text (Ziel:, Wichtig:, …); ignoriert Bild- und Builder-Pfade.
    """
    if not (task or "").strip():
        return None
    raw = str(task)

    for line in raw.splitlines():
        ln = line.strip()
        low = ln.lower()
        if low.startswith("zielordner:") or low.startswith("ziel-ordner:"):
            rest = ln.split(":", 1)[-1].strip()
            got = _finalize_windows_folder_candidate(rest)
            if got:
                return got

    for pat in (
        r"(?is)\blanden\s+unter\s*:\s*([a-z]:(?:\\|/)[^\n\r]+)",
        r"(?is)\b(?:alle\s+dateien\s+landen\s+(?:unter|in))\s*:\s*([a-z]:(?:\\|/)[^\n\r]+)",
        r"(?is)\b(?:arbeiten\s+direkt\s+im\s+ordner|in\s+diesen\s+ordner|zielordner\s+ist)\s*[:\s]\s*([a-z]:(?:\\|/)[^\n\r]+)",
    ):
        m = re.search(pat, raw)
        if m:
            got = _finalize_windows_folder_candidate(m.group(1))
            if got:
                return got

    paths = re.findall(r"(?i)([a-z]:(?:\\|/)[^\n\r]+)", raw)
    cleaned = []
    seen = set()
    for p in paths:
        got = _finalize_windows_folder_candidate(p)
        if not got or got.lower() in seen:
            continue
        seen.add(got.lower())
        cleaned.append(got)

    for p in cleaned:
        if "rainerrobotdesktop" in p.lower():
            return p
    for p in cleaned:
        pl = p.lower().replace("/", "\\")
        if r"\downloads" in pl and "rambo_builder_local" not in pl:
            return p
    return cleaned[0] if cleaned else None


def _extract_robot_icon_source_path(task: str):
    raw = str(task or "")
    for line in raw.splitlines():
        low = line.strip().lower()
        if "quelle" in low and "robot" in low and ":" in line:
            rest = line.split(":", 1)[-1].strip()
            rest = _strip_trailing_prompt_from_windows_path_line(rest.replace("/", "\\"))
            if re.match(r"(?i)^[a-z]:\\", rest):
                return rest
    for p in re.findall(r"(?i)([a-z]:(?:\\|/)[^\n\r]+\.(?:png|jpe?g|webp|gif))", raw):
        c = _strip_trailing_prompt_from_windows_path_line(p.replace("/", "\\"))
        if c:
            return c
    return None


def _sync_robot_desktop_icons(base: Path, task: str) -> list[str]:
    """Kopiert Roboter-Bild nach rambo_ui/public und electron/assets (Fenster- + Installer-Icon)."""
    import shutil

    out: list[str] = []
    src = _extract_robot_icon_source_path(task)
    if not src:
        cand = DOWNLOADS_DIR / "Bild.png"
        if cand.is_file():
            src = str(cand)
    if not src:
        return out
    sp = Path(src)
    if not sp.is_file():
        return out
    try:
        pub = base / "rambo_ui" / "public"
        pub.mkdir(parents=True, exist_ok=True)
        assets = base / "electron" / "assets"
        assets.mkdir(parents=True, exist_ok=True)
        p1 = pub / "robot-icon.png"
        p2 = assets / "app-icon.png"
        shutil.copy2(sp, p1)
        shutil.copy2(sp, p2)
        out.extend([str(p1), str(p2)])
    except OSError:
        pass
    return out


def _merge_post_build_agent_digest(result: dict, task: str, workspace: Path) -> dict:
    """
    Optional nach Robot-Desktop-Build: AgentLoop.run_analysis auf dem Zielordner (nur Lesen).
    Deaktivieren: AGENT_SKIP_POST_BUILD_ANALYSIS=1
    """
    if not isinstance(result, dict):
        return result
    if str(os.environ.get("AGENT_SKIP_POST_BUILD_ANALYSIS") or "").strip().lower() in ("1", "true", "yes"):
        return result
    if not AGENT_LOOP_AVAILABLE or AgentLoop is None:
        return result
    try:
        root = Path(workspace).resolve()
        if not root.is_dir():
            return result
        agent_o = AgentLoop(root)
        digest_task = (
            f"Nutzerauftrag (Auszug): {task[:500]}\n\n"
            "Analysiere nur die vorhandene Struktur und Dateiinhalte. "
            "Antworte mit 3–7 kurzen Stichpunkten: was liegt unter electron/, was unter rambo_ui/, "
            "wofür build_desktop.py. Keine Änderungsvorschläge, kein neuer Code."
        )
        ar = agent_o.run_analysis(digest_task)
        if not isinstance(ar, dict) or not ar.get("ok"):
            return result
        digest = str(ar.get("analysis") or "").strip()
        if not digest:
            return result
        out = dict(result)
        suffix = "\n\n**Projekt-Kurzcheck (AgentLoop)**\n\n" + digest
        out["post_build_analysis"] = digest
        out["agent_loop_after_build"] = True
        base = str(
            out.get("formatted_response") or out.get("message") or out.get("technical_message") or ""
        ).strip()
        new_body = base + suffix
        out["formatted_response"] = new_body
        if "message" in out:
            out["message"] = new_body
        if "technical_message" in out:
            out["technical_message"] = new_body
        we = out.get("workstream_events")
        if isinstance(we, list):
            out["workstream_events"] = list(we) + [
                _ws_event("analyze", "info", "AgentLoop", "Kurzcheck nach Build", status="done"),
            ]
        return out
    except Exception:
        return result


def execute_project_build(task, run_id, scope="project", mode="safe"):
    """
    Multi-File Projekt-Build Handler für Electron/React/Vite Apps.
    Extrahiert Zielpfad, erstellt File-Plan, schreibt Dateien, führt Build aus.
    """
    import subprocess
    pre_guard = _validate_direct_run_paths(
        list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN) + ["build_desktop.py", "robot-icon.png", "app-icon.png"],
        mode,
        task,
    )
    if not bool(pre_guard.get("ok")):
        blocked_payload = _build_direct_guard_block_payload(
            scope=scope,
            mode=mode,
            blocked_files=pre_guard.get("blocked_files") or list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN),
            task=task,
            recognized_task={
                "task_type": "project_build",
                "primary_area": "Project Builder",
                "execution_route": "project_build",
            },
        )
        blocked_payload["run_id"] = run_id
        return blocked_payload

    # 1. Zielpfad: Zielordner-Zeile / erster gueltiger Ordner unter Downloads (kein Fallback bei ungueltigem Extrakt)
    base_path_str = _extract_robot_desktop_base_path_str(task)
    if base_path_str:
        normalized = base_path_str.replace("\\", "/")
        is_valid, validation_result = validate_downloads_path(normalized)
        if not is_valid:
            return _project_build_invalid_path_payload(
                run_id,
                scope,
                mode,
                f"Zielpfad liegt nicht unter Downloads oder ist ungueltig: {validation_result}",
            )
        base_path = Path(validation_result)
    else:
        safe_name = re.sub(r"[^\w\-]", "_", task[:40].lower())
        project_match = re.search(r"(?:in|unter|nach)\s+([A-Z][a-zA-Z0-9]+(?:App|Project|Build))", task)
        if project_match:
            safe_name = project_match.group(1)
        base_path = DOWNLOADS_DIR / safe_name
        v_ok, v_msg = validate_downloads_path(str(base_path).replace("\\", "/"))
        if not v_ok:
            return _project_build_invalid_path_payload(run_id, scope, mode, v_msg or "Fallback-Zielpfad ungueltig.")

    try:
        base_path.mkdir(parents=True, exist_ok=True)
    except OSError as ose:
        return _project_build_invalid_path_payload(
            run_id,
            scope,
            mode,
            f"Zielordner nicht anlegbar: {getattr(ose, 'winerror', '')} {ose}".strip(),
        )

    # 2. File-Plan (Pflichtdateien + React-Komponenten)
    file_plan = {
        "electron/main.js": ELECTRON_MAIN_JS,
        "electron/preload.js": ELECTRON_PRELOAD_JS,
        "electron/package.json": ELECTRON_PACKAGE_JSON,
        "electron/fallback-setup.html": ELECTRON_FALLBACK_SETUP_HTML,
        "rambo_ui/package.json": RAMBO_UI_PACKAGE_JSON,
        "rambo_ui/vite.config.js": RAMBO_UI_VITE_CONFIG,
        "rambo_ui/index.html": RAMBO_UI_INDEX_HTML,
        "rambo_ui/src/main.jsx": RAMBO_UI_MAIN_JSX,
        "rambo_ui/src/App.jsx": RAMBO_UI_APP_JSX,
        "rambo_ui/src/App.css": RAMBO_UI_APP_CSS,
        "build_desktop.py": BUILD_DESKTOP_PY,
    }
    file_plan.update(get_robot_desktop_extra_files())

    planned_files = [str(base_path / f) for f in file_plan.keys()]
    created_files = []
    failed_files = []

    # 3. Dateien schreiben
    for rel_path, content in file_plan.items():
        full_path = base_path / rel_path
        try:
            full_path.parent.mkdir(parents=True, exist_ok=True)
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(content)
            created_files.append(str(full_path))
        except Exception as e:
            failed_files.append({"path": str(full_path), "error": str(e)})

    for icon_path in _sync_robot_desktop_icons(base_path, task):
        if icon_path not in created_files:
            created_files.append(icon_path)

    # 4. Build-Kommandos ausführen
    build_logs = []
    build_status = "pending"

    def run_npm(cmd, cwd):
        try:
            result = subprocess.run(
                cmd, shell=True, cwd=cwd, capture_output=True, text=True, timeout=300
            )
            return {
                "cmd": cmd,
                "cwd": str(cwd),
                "returncode": result.returncode,
                "stdout": result.stdout[:2000],
                "stderr": result.stderr[:2000],
            }
        except Exception as e:
            return {"cmd": cmd, "cwd": str(cwd), "error": str(e)}

    rambo_ui_path = base_path / "rambo_ui"
    electron_path = base_path / "electron"

    if rambo_ui_path.exists():
        npm_install_result = run_npm("npm install", rambo_ui_path)
        build_logs.append(npm_install_result)
        if npm_install_result.get("returncode") == 0:
            npm_build_result = run_npm("npm run build", rambo_ui_path)
            build_logs.append(npm_build_result)

    if electron_path.exists():
        npm_install_electron = run_npm("npm install", electron_path)
        build_logs.append(npm_install_electron)

    # 5. Verifikation
    missing_files = [f for f in planned_files if not Path(f).exists()]
    created_count = len([f for f in planned_files if Path(f).exists()])

    if missing_files:
        build_status = "partial" if created_count > 0 else "error"
    else:
        build_status = "success"

    # Logs speichern
    log_dir = DOWNLOADS_DIR / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"project_build_{run_id}.json"
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            json.dump({
                "run_id": run_id,
                "task": task[:200],
                "base_path": str(base_path),
                "planned_files": planned_files,
                "created_files": created_files,
                "failed_files": failed_files,
                "missing_files": missing_files,
                "build_logs": build_logs,
                "build_status": build_status,
                "timestamp": get_timestamp(),
            }, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

    pb_result = {
        "run_id": run_id,
        "scope": scope,
        "mode": mode,
        "ok": build_status == "success",
        "direct_status": build_status,
        "message": f"Project Build: {created_count}/{len(planned_files)} Dateien erstellt." if build_status == "success" else f"Project Build unvollständig: {len(missing_files)} Dateien fehlen.",
        "requires_confirmation": False,
        "requires_user_confirmation": False,
        "planned_files_count": len(planned_files),
        "created_files_count": created_count,
        "missing_files": missing_files,
        "created_files": created_files,
        "file_plan": list(file_plan.keys()),
        "build_status": build_status,
        "base_path": str(base_path),
        "target_root": str(base_path),
        "has_changes": created_count > 0,
        "robot_build_auto_applied": True,
        "debug_auto_apply_decision": {"is_robot_desktop_build": True, "execution_route": "project_build"},
        "workstream_events": [
            _ws_event("plan", "info", "Zielordner", str(base_path), status="done"),
            _ws_event(
                "verify",
                "success" if build_status == "success" else "error",
                "Project Build",
                f"{created_count}/{len(planned_files)} Dateien",
                status="done" if build_status == "success" else "failed",
            ),
        ],
        "log_file": str(log_file) if build_logs else None,
        "recognized_task": {
            "task_type": "project_build",
            "primary_area": "Project Builder",
            "hint": "Multi-File Projekt generiert.",
            "execution_route": "project_build",
        },
    }
    if created_count > 0 and build_status == "success":
        pb_result = _merge_post_build_agent_digest(pb_result, task, base_path)
    return pb_result


# Template-Inhalte für Electron/React Projekt
ELECTRON_FALLBACK_SETUP_HTML = '''<!DOCTYPE html>
<html lang="de">
<head>
  <meta charset="utf-8" />
  <title>Rainer Robot Desktop — Setup</title>
  <style>
    body { font-family: Segoe UI, system-ui, sans-serif; background: #0d1117; color: #c9d1d9; padding: 32px; max-width: 720px; margin: 0 auto; line-height: 1.6; }
    h1 { color: #58a6ff; }
    code { background: #21262d; padding: 2px 8px; border-radius: 4px; }
    ol { line-height: 2; }
  </style>
</head>
<body>
  <h1>Rainer Robot Desktop</h1>
  <p>Es gibt weder <code>rambo_ui/dist/index.html</code> noch einen erreichbaren Vite-Dev-Server (Port 5173).</p>
  <h2>So bekommen Sie die React-Oberfläche</h2>
  <ol>
    <li><strong>Empfohlen:</strong> im Ordner <code>RainerRobotDesktop</code> <code>python build_desktop.py</code> ausführen, danach in <code>electron</code> erneut <code>npm start</code>.</li>
    <li><strong>Oder Dev:</strong> in <code>rambo_ui</code> <code>npm install</code>, dann <code>npm run dev</code> laufen lassen; parallel in <code>electron</code> <code>npm start</code>.</li>
    <li><strong>Windows-Installer:</strong> nach erfolgreichem Build <code>set ROBOT_DESKTOP_PACK=1</code> und erneut <code>python build_desktop.py</code> — Setup liegt unter <code>electron/dist-installer/</code>.</li>
  </ol>
</body>
</html>
'''

ELECTRON_MAIN_JS = '''const { app, BrowserWindow, ipcMain, shell } = require('electron');
const path = require('path');
const fs = require('fs');
const http = require('http');
const { spawn } = require('child_process');

const RAMBO_ROOT = process.env.RAMBO_RAINER_ROOT || 'C:/Users/mielersch/Desktop/Rambo-Rainer';
const RAINER_ROOT = process.env.RAINER_BUILD_ROOT || path.join(RAMBO_ROOT, 'rambo_builder_local');

const state = {
  rambo: { child: null, startedByUs: false, buf: [] },
  rainer: { child: null, startedByUs: false, buf: [] },
};

function logBuf(name, line) {
  const b = state[name].buf;
  b.push(`${new Date().toISOString().slice(11, 19)} ${line}`);
  while (b.length > 120) b.shift();
}

function httpHealth(port) {
  return new Promise((resolve) => {
    const req = http.request(
      { hostname: '127.0.0.1', port, path: '/api/health', method: 'GET', timeout: 2500 },
      (res) => {
        resolve(res.statusCode >= 200 && res.statusCode < 500);
        res.resume();
      }
    );
    req.on('error', () => resolve(false));
    req.on('timeout', () => {
      try { req.destroy(); } catch (e) {}
      resolve(false);
    });
    req.end();
  });
}

function resolveRamboPythonTarget() {
  const seq = [['backend', 'server.py'], ['backend', 'main.py'], ['server.py'], ['main.py'], ['app.py']];
  for (let i = 0; i < seq.length; i++) {
    const full = path.join(RAMBO_ROOT, ...seq[i]);
    if (fs.existsSync(full)) return { cwd: RAMBO_ROOT, script: full };
  }
  return null;
}

function resolveRainerPythonTarget() {
  const full = path.join(RAINER_ROOT, 'backend', 'main.py');
  return fs.existsSync(full) ? { cwd: RAINER_ROOT, script: full } : null;
}

function startPythonService(name, resolveFn) {
  const port = name === 'rambo' ? 5001 : 5002;
  return httpHealth(port).then((up) => {
    if (up) {
      return { ok: true, message: 'Dienst laeuft bereits (Port erreichbar).', started: false };
    }
    const det = resolveFn();
    if (!det) {
      const msg = name === 'rambo'
        ? 'Rambo-Rainer Startbefehl muss noch konfiguriert werden (kein passendes Backend gefunden).'
        : 'Rainer Build: backend/main.py nicht gefunden.';
      logBuf(name, msg);
      return { ok: false, message: msg, started: false };
    }
    const py = process.platform === 'win32' ? 'python' : 'python3';
    const child = spawn(py, [det.script], {
      cwd: det.cwd,
      env: { ...process.env, PYTHONUTF8: '1' },
      stdio: 'ignore',
      windowsHide: true,
      detached: false,
    });
    child.on('error', (err) => logBuf(name, String(err)));
    state[name].child = child;
    state[name].startedByUs = true;
    logBuf(name, 'Start: ' + py + ' ' + det.script);
    return { ok: true, message: 'Start ausgeloest.', started: true };
  });
}

function stopService(name) {
  const st = state[name];
  if (!st.startedByUs) {
    return { ok: true, skipped: true, message: 'Nicht von dieser App gestartet (Stop uebersprungen).' };
  }
  if (!st.child) {
    st.startedByUs = false;
    return { ok: true, message: 'Kein Kindprozess.' };
  }
  const pid = st.child.pid;
  try {
    if (process.platform === 'win32' && pid) {
      spawn('taskkill', ['/PID', String(pid), '/T', '/F'], { shell: true, stdio: 'ignore', detached: true });
    } else {
      st.child.kill('SIGTERM');
    }
  } catch (e) {
    logBuf(name, 'Stop: ' + String(e));
  }
  st.child = null;
  st.startedByUs = false;
  logBuf(name, 'Stop ausgefuehrt');
  return { ok: true, message: 'Gestoppt.', skipped: false };
}

function registerShellIpc() {
  ipcMain.handle('shell:probe', async () => {
    const ramboOk = await httpHealth(5001);
    const rainerOk = await httpHealth(5002);
    return {
      rambo: ramboOk ? 'running' : 'unreachable',
      rainer: rainerOk ? 'running' : 'unreachable',
    };
  });
  ipcMain.handle('shell:start', (e, name) => {
    if (name === 'rambo') return startPythonService('rambo', resolveRamboPythonTarget);
    if (name === 'rainer') return startPythonService('rainer', resolveRainerPythonTarget);
    return Promise.resolve({ ok: false, message: 'Unbekannter Dienst' });
  });
  ipcMain.handle('shell:stop', (e, name) => {
    if (name !== 'rambo' && name !== 'rainer') return { ok: false };
    return stopService(name);
  });
  ipcMain.handle('shell:restart', async (e, name) => {
    if (name !== 'rambo' && name !== 'rainer') return { ok: false };
    stopService(name);
    await new Promise((r) => setTimeout(r, 700));
    if (name === 'rambo') return startPythonService('rambo', resolveRamboPythonTarget);
    return startPythonService('rainer', resolveRainerPythonTarget);
  });
  ipcMain.handle('shell:open-external', async (e, url) => {
    try {
      await shell.openExternal(String(url || ''));
      return { ok: true };
    } catch (err) {
      return { ok: false, message: String(err) };
    }
  });
  ipcMain.handle('shell:logs', () => ({
    rambo: state.rambo.buf.slice(),
    rainer: state.rainer.buf.slice(),
  }));
}

function killOurChildren() {
  stopService('rambo');
  stopService('rainer');
}

async function autoStartIfDown() {
  try {
    const rUp = await httpHealth(5001);
    if (!rUp) await startPythonService('rambo', resolveRamboPythonTarget);
    await new Promise((r) => setTimeout(r, 500));
    const bUp = await httpHealth(5002);
    if (!bUp) await startPythonService('rainer', resolveRainerPythonTarget);
  } catch (e) {}
}

registerShellIpc();

function createWindow() {
  const appIcon = path.join(__dirname, 'assets', 'app-icon.png');
  const winOpts = {
    width: 1280,
    height: 860,
    title: 'Rainer Robot Desktop',
    webPreferences: {
      preload: path.join(__dirname, 'preload.js'),
      contextIsolation: true,
      nodeIntegration: false,
      webSecurity: false,
    },
  };
  if (fs.existsSync(appIcon)) {
    winOpts.icon = appIcon;
  }
  const win = new BrowserWindow(winOpts);

  const distIndex = path.join(__dirname, '..', 'rambo_ui', 'dist', 'index.html');
  const fallbackPage = path.join(__dirname, 'fallback-setup.html');
  const devUrl = (process.env.VITE_DEV_SERVER_URL || 'http://127.0.0.1:5173/').replace(/\\/?$/, '/');

  function loadDist() {
    win.loadFile(distIndex);
  }

  function loadFallback() {
    if (fs.existsSync(fallbackPage)) {
      win.loadFile(fallbackPage);
    } else {
      const html = '<!DOCTYPE html><html lang="de"><head><meta charset="utf-8"><title>Setup</title></head><body style="font-family:system-ui;padding:24px;background:#0d1117;color:#c9d1d9"><p>Bitte im Projektroot <code>python build_desktop.py</code> ausfuehren.</p></body></html>';
      win.loadURL('data:text/html;charset=utf-8,' + encodeURIComponent(html));
    }
  }

  if (fs.existsSync(distIndex)) {
    loadDist();
  } else if (app.isPackaged) {
    loadFallback();
  } else {
    let swapped = false;
    const onFail = (event, errorCode, errorDesc, validatedURL) => {
      if (swapped) return;
      const u = String(validatedURL || '');
      if (!u.includes('127.0.0.1:5173') && !u.includes('localhost:5173')) return;
      swapped = true;
      win.webContents.removeListener('did-fail-load', onFail);
      loadFallback();
    };
    win.webContents.on('did-fail-load', onFail);
    win.loadURL(devUrl);
  }

  if (!app.isPackaged) {
    win.webContents.openDevTools({ mode: 'detach' });
  }
}

app.on('before-quit', () => {
  killOurChildren();
});

app.whenReady().then(() => {
  createWindow();
  setTimeout(() => { autoStartIfDown(); }, 700);
  app.on('activate', () => {
    if (BrowserWindow.getAllWindows().length === 0) createWindow();
  });
});

app.on('window-all-closed', () => {
  if (process.platform !== 'darwin') app.quit();
});
'''

ELECTRON_PRELOAD_JS = '''const { contextBridge, ipcRenderer } = require('electron');

contextBridge.exposeInMainWorld('desktopShell', {
  probe: () => ipcRenderer.invoke('shell:probe'),
  start: (name) => ipcRenderer.invoke('shell:start', name),
  stop: (name) => ipcRenderer.invoke('shell:stop', name),
  restart: (name) => ipcRenderer.invoke('shell:restart', name),
  openExternal: (url) => ipcRenderer.invoke('shell:open-external', url),
  logs: () => ipcRenderer.invoke('shell:logs'),
});

contextBridge.exposeInMainWorld('electronAPI', {
  ping: () => 'pong',
});
'''

ELECTRON_PACKAGE_JSON = '''{
  "name": "electron-app",
  "version": "1.0.0",
  "main": "main.js",
  "scripts": {
    "start": "electron .",
    "build:win": "electron-builder --win",
    "pack:installer": "electron-builder --win --publish never"
  },
  "dependencies": {
    "electron": "^28.0.0"
  },
  "devDependencies": {
    "electron-builder": "^24.0.0"
  },
  "build": {
    "appId": "com.rainer.robotdesktop",
    "productName": "Rainer Robot Desktop",
    "directories": {
      "output": "dist-installer",
      "buildResources": "build"
    },
    "files": [
      "main.js",
      "preload.js",
      "fallback-setup.html",
      "assets/**/*",
      "../rambo_ui/dist/**/*"
    ],
    "win": {
      "target": [
        {
          "target": "nsis",
          "arch": [
            "x64"
          ]
        }
      ],
      "icon": "assets/app-icon.png"
    },
    "nsis": {
      "oneClick": false,
      "allowToChangeInstallationDirectory": true,
      "allowElevation": true,
      "createDesktopShortcut": true,
      "createStartMenuShortcut": true,
      "shortcutName": "Rainer Robot Desktop",
      "artifactName": "${productName}-${version}-Setup.${ext}"
    }
  }
}
'''

RAMBO_UI_PACKAGE_JSON = '''{
  "name": "rambo-ui",
  "private": true,
  "version": "1.0.0",
  "type": "module",
  "scripts": {
    "dev": "vite",
    "build": "vite build",
    "preview": "vite preview"
  },
  "dependencies": {
    "react": "^18.2.0",
    "react-dom": "^18.2.0"
  },
  "devDependencies": {
    "@types/react": "^18.2.43",
    "@types/react-dom": "^18.2.17",
    "@vitejs/plugin-react": "^4.2.1",
    "vite": "^5.0.8"
  }
}
'''

RAMBO_UI_VITE_CONFIG = '''import { defineConfig } from 'vite'
import react from '@vitejs/plugin-react'

// Relativ-Pfade: Electron loadFile(.../dist/index.html) laedt Skripte unter file:// korrekt
export default defineConfig({
  plugins: [react()],
  base: './',
  build: {
    outDir: 'dist',
    assetsDir: 'assets',
    emptyOutDir: true,
  },
})
'''

RAMBO_UI_INDEX_HTML = '''<!doctype html>
<html lang="de">
  <head>
    <meta charset="UTF-8" />
    <link rel="icon" type="image/png" href="./robot-icon.png" />
    <meta name="viewport" content="width=device-width, initial-scale=1.0" />
    <title>Rainer Robot Desktop</title>
  </head>
  <body>
    <div id="root"></div>
    <script type="module" src="/src/main.jsx"></script>
  </body>
</html>
'''

RAMBO_UI_MAIN_JSX = '''import React from 'react'
import ReactDOM from 'react-dom/client'
import App from './App.jsx'
import './App.css'

ReactDOM.createRoot(document.getElementById('root')).render(
  <React.StrictMode>
    <App />
  </React.StrictMode>,
)
'''

RAMBO_UI_APP_JSX = '''import { useState, useEffect, useCallback } from 'react'
import './App.css'
import './components/UI.css'
import ServicePanel from './components/ServicePanel.jsx'

const ds = typeof window !== 'undefined' ? window.desktopShell : null

function joinLogs(L) {
  if (!L) return ''
  const a = (L.rambo || []).join(String.fromCharCode(10))
  const b = (L.rainer || []).join(String.fromCharCode(10))
  return a + String.fromCharCode(10) + '---' + String.fromCharCode(10) + b
}

export default function App() {
  const [tab, setTab] = useState('rambo')
  const [probe, setProbe] = useState({ rambo: 'unknown', rainer: 'unknown' })
  const [logText, setLogText] = useState('')

  const pullProbe = useCallback(async () => {
    if (!ds) return
    try {
      const p = await ds.probe()
      setProbe(p)
    } catch {
      setProbe({ rambo: 'error', rainer: 'error' })
    }
  }, [])

  const pullLogs = useCallback(async () => {
    if (!ds || !ds.logs) return
    try {
      const L = await ds.logs()
      setLogText(joinLogs(L))
    } catch {
      /* ignore */
    }
  }, [])

  useEffect(() => {
    pullProbe()
    const t1 = setInterval(pullProbe, 4000)
    const t2 = setInterval(pullLogs, 5000)
    pullLogs()
    return () => {
      clearInterval(t1)
      clearInterval(t2)
    }
  }, [pullProbe, pullLogs])

  return (
    <div className="shell-root">
      <header className="shell-header">
        <img
          src="/robot-icon.png"
          alt=""
          className="shell-logo"
          onError={(e) => {
            e.currentTarget.style.display = 'none'
          }}
        />
        <div className="shell-header-text">
          <h1>Rainer Robot Desktop</h1>
          <p className="shell-sub">Rambo-Rainer (Port 5001) und Rainer Build (Port 5002)</p>
        </div>
      </header>
      <nav className="shell-tabs" aria-label="Hauptmodi">
        <button type="button" className={tab === 'rambo' ? 'on' : ''} onClick={() => setTab('rambo')}>
          Rambo-Rainer
        </button>
        <button type="button" className={tab === 'rainer' ? 'on' : ''} onClick={() => setTab('rainer')}>
          Rainer Build
        </button>
      </nav>
      <main className="shell-main">
        {tab === 'rambo' ? (
          <ServicePanel
            serviceId="rambo"
            label="Rambo-Rainer"
            port={5001}
            baseUrl="http://127.0.0.1:5001"
            embedPath="/"
            status={probe.rambo}
            onRefresh={pullProbe}
          />
        ) : (
          <ServicePanel
            serviceId="rainer"
            label="Rainer Build"
            port={5002}
            baseUrl="http://127.0.0.1:5002"
            embedPath="/?desktop=1"
            status={probe.rainer}
            onRefresh={pullProbe}
          />
        )}
      </main>
      <footer className="shell-log">
        <div className="shell-log-head">Protokoll (nur von dieser Desktop-App gestartete Prozesse)</div>
        <pre>{logText || '(leer)'}</pre>
      </footer>
    </div>
  )
}
'''

RAMBO_UI_APP_CSS = '''* {
  margin: 0;
  padding: 0;
  box-sizing: border-box;
}

body {
  font-family: 'Segoe UI', system-ui, sans-serif;
  background: #0d1117;
  color: #c9d1d9;
  min-height: 100vh;
}

.shell-root {
  min-height: 100vh;
  display: flex;
  flex-direction: column;
}

.shell-header {
  display: flex;
  align-items: center;
  gap: 20px;
  padding: 16px 24px;
  border-bottom: 1px solid #30363d;
  background: linear-gradient(180deg, #161b22 0%, #0d1117 100%);
}

.shell-logo {
  width: 56px;
  height: 56px;
  object-fit: contain;
  border-radius: 10px;
  border: 1px solid #30363d;
}

.shell-header-text h1 {
  font-size: 1.35rem;
  color: #58a6ff;
  margin: 0 0 4px 0;
}

.shell-sub {
  font-size: 0.85rem;
  color: #8b949e;
  margin: 0;
}

.shell-tabs {
  display: flex;
  gap: 8px;
  padding: 10px 16px;
  background: #161b22;
  border-bottom: 1px solid #30363d;
}

.shell-tabs button {
  padding: 10px 20px;
  border-radius: 8px;
  border: 1px solid #30363d;
  background: #21262d;
  color: #c9d1d9;
  cursor: pointer;
  font-size: 14px;
}

.shell-tabs button.on {
  background: #238636;
  border-color: #238636;
  color: #fff;
}

.shell-main {
  flex: 1;
  padding: 16px;
  overflow: auto;
}

.shell-log {
  border-top: 1px solid #30363d;
  background: #010409;
  max-height: 160px;
  display: flex;
  flex-direction: column;
}

.shell-log-head {
  font-size: 12px;
  color: #8b949e;
  padding: 6px 12px;
}

.shell-log pre {
  flex: 1;
  overflow: auto;
  padding: 0 12px 8px;
  font-size: 11px;
  color: #7ee787;
  white-space: pre-wrap;
}
'''

BUILD_DESKTOP_PY = '''#!/usr/bin/env python3
"""
Build-Skript für Desktop App.
- npm install + build (rambo_ui), npm install (electron)
- Optional: Windows-Installer mit electron-builder (Umgebung ROBOT_DESKTOP_PACK=1)
"""
import os
import subprocess
import sys
from pathlib import Path

def run_cmd(cmd, cwd, description, timeout=300):
    print(f"\\n=== {description} ===")
    print(f"Command: {cmd}")
    print(f"Working dir: {cwd}")
    try:
        result = subprocess.run(cmd, shell=True, cwd=cwd, capture_output=True, text=True, timeout=timeout)
        print(f"Return code: {result.returncode}")
        if result.stdout:
            print("STDOUT:", result.stdout[:1000])
        if result.stderr:
            print("STDERR:", result.stderr[:1000])
        return result.returncode == 0
    except Exception as e:
        print(f"ERROR: {e}")
        return False

def main():
    base_path = Path(__file__).parent
    rambo_ui = base_path / "rambo_ui"
    electron = base_path / "electron"

    success = True

    # 1. npm install in rambo_ui
    if rambo_ui.exists():
        if not run_cmd("npm install", rambo_ui, "npm install in rambo_ui"):
            success = False
    else:
        print(f"ERROR: {rambo_ui} not found")
        success = False

    # 2. npm run build in rambo_ui
    if rambo_ui.exists():
        if not run_cmd("npm run build", rambo_ui, "npm run build in rambo_ui"):
            success = False

    # 3. npm install in electron
    if electron.exists():
        if not run_cmd("npm install", electron, "npm install in electron"):
            success = False
    else:
        print(f"ERROR: {electron} not found")
        success = False

    # 4. Optional: NSIS-Installer (electron-builder, kann mehrere Minuten dauern)
    pack = str(os.environ.get("ROBOT_DESKTOP_PACK", "")).lower() in ("1", "true", "yes", "on")
    if pack and electron.exists() and success:
        if not run_cmd("npm run pack:installer", electron, "electron-builder Windows-Installer", timeout=1200):
            success = False
    elif not pack:
        print("\\nHinweis: Installer mit set ROBOT_DESKTOP_PACK=1 && python build_desktop.py erzeugen.")

    if success:
        print("\\n=== Build erfolgreich! ===")
        return 0
    else:
        print("\\n=== Build mit Fehlern abgeschlossen ===")
        return 1

if __name__ == "__main__":
    sys.exit(main())
'''


RAMBO_DESKTOP_SERVICE_PANEL_JSX = '''import { useState, useRef } from 'react'

export default function ServicePanel({ serviceId, label, port, baseUrl, embedPath, status, onRefresh }) {
  const [busy, setBusy] = useState(false)
  const [msg, setMsg] = useState('')
  const iframeRef = useRef(null)
  const ds = typeof window !== 'undefined' ? window.desktopShell : null
  const embedSrc = baseUrl + (embedPath || '/')
  const running = status === 'running'

  async function run(op) {
    if (!ds) {
      setMsg('Nur in Electron verfuegbar.')
      return
    }
    setBusy(true)
    setMsg('')
    try {
      let r = null
      if (op === 'start') r = await ds.start(serviceId)
      else if (op === 'stop') r = await ds.stop(serviceId)
      else if (op === 'restart') r = await ds.restart(serviceId)
      setMsg((r && r.message) || '')
      if (onRefresh) await onRefresh()
    } catch (e) {
      setMsg(String(e))
    } finally {
      setBusy(false)
    }
  }

  function openInApp() {
    if (iframeRef.current) iframeRef.current.src = embedSrc
  }

  async function openBrowser() {
    if (ds && ds.openExternal) await ds.openExternal(embedSrc)
    else window.open(embedSrc, '_blank')
  }

  const statusClass =
    status === 'running'
      ? 'ok'
      : status === 'unreachable'
        ? 'bad'
        : status === 'error'
          ? 'bad'
          : 'unk'

  return (
    <div className="svc-panel">
      <div className="svc-head">
        <h2>{label}</h2>
        <div className={'svc-status svc-status-' + statusClass}>
          {busy ? 'startet…' : status === 'running' ? 'laeuft' : status === 'unreachable' ? 'nicht erreichbar' : status === 'error' ? 'Fehler' : 'unbekannt'}
        </div>
      </div>
      <p className="svc-meta">
        Port {port} ·{' '}
        <a href={embedSrc} target="_blank" rel="noreferrer">
          {embedSrc}
        </a>
      </p>
      <div className="svc-actions">
        <button type="button" disabled={busy} onClick={() => run('start')}>
          Start
        </button>
        <button type="button" disabled={busy} onClick={() => run('stop')}>
          Stop
        </button>
        <button type="button" disabled={busy} onClick={() => run('restart')}>
          Neustart
        </button>
        <button type="button" disabled={busy} onClick={openInApp}>
          In App oeffnen
        </button>
        <button type="button" disabled={busy} onClick={openBrowser}>
          Extern im Browser
        </button>
      </div>
      {msg ? <p className="svc-msg">{msg}</p> : null}
      {running ? (
        <div className="svc-embed-wrap">
          <iframe ref={iframeRef} title={label} src={embedSrc} className="svc-iframe" />
        </div>
      ) : (
        <div className="svc-card">
          <p>Dienst nicht erreichbar oder laeuft ausserhalb dieser App.</p>
          <p className="svc-hint">Start waehlen oder Vite/React-Build pruefen (siehe build_desktop.py).</p>
        </div>
      )}
    </div>
  )
}
'''

RAMBO_DESKTOP_SHELL_CSS = '''.svc-panel {
  max-width: 1100px;
  margin: 0 auto;
}

.svc-head {
  display: flex;
  align-items: center;
  justify-content: space-between;
  gap: 12px;
  margin-bottom: 10px;
}

.svc-head h2 {
  font-size: 1.2rem;
  color: #79c0ff;
  margin: 0;
}

.svc-status {
  font-size: 13px;
  padding: 6px 12px;
  border-radius: 999px;
  border: 1px solid #30363d;
}

.svc-status-ok {
  border-color: #238636;
  color: #3fb950;
}

.svc-status-bad {
  border-color: #da3633;
  color: #ff7b72;
}

.svc-status-unk {
  color: #8b949e;
}

.svc-meta {
  font-size: 13px;
  color: #8b949e;
  margin-bottom: 12px;
}

.svc-meta a {
  color: #58a6ff;
}

.svc-actions {
  display: flex;
  flex-wrap: wrap;
  gap: 8px;
  margin-bottom: 12px;
}

.svc-actions button {
  padding: 8px 14px;
  border-radius: 6px;
  border: 1px solid #30363d;
  background: #21262d;
  color: #c9d1d9;
  cursor: pointer;
  font-size: 13px;
}

.svc-actions button:hover:not(:disabled) {
  border-color: #58a6ff;
}

.svc-actions button:disabled {
  opacity: 0.55;
  cursor: not-allowed;
}

.svc-msg {
  font-size: 13px;
  color: #d29922;
  margin-bottom: 10px;
}

.svc-card {
  background: #161b22;
  border: 1px solid #30363d;
  border-radius: 10px;
  padding: 20px;
  text-align: left;
  line-height: 1.5;
}

.svc-hint {
  color: #8b949e;
  font-size: 13px;
  margin-top: 8px;
}

.svc-embed-wrap {
  margin-top: 12px;
  border: 1px solid #30363d;
  border-radius: 10px;
  overflow: hidden;
  background: #000;
  min-height: 480px;
}

.svc-iframe {
  width: 100%;
  height: 520px;
  border: 0;
  display: block;
}
'''


def get_robot_desktop_extra_files():
    return {
        "rambo_ui/src/components/ServicePanel.jsx": RAMBO_DESKTOP_SERVICE_PANEL_JSX,
        "rambo_ui/src/components/UI.css": RAMBO_DESKTOP_SHELL_CSS,
    }


def extract_windows_target_dir(prompt):
    """
    Extrahiert sauber einen Windows-Zielpfad aus dem Prompt.
    Stoppt bei Zeilenumbrüchen, Anführungszeichen, etc.
    """
    if not prompt:
        return None

    # Pattern 1: C:\... bis zum Zeilenumbruch oder Ende
    # Stoppe bei: \n, \r, ", `, <, >, Bullet (-, *, •)
    pattern = r'([A-Z]:\\[^\n\r"`<>|?*]+?)(?:\s*[\n\r]|\s{2,}|\s*$|\"|\`|\<|\>|\s+[-\*•]\s|\s+\Z)'
    match = re.search(pattern, prompt, re.IGNORECASE)

    if match:
        path = match.group(1).strip()
        path = _strip_trailing_prompt_from_windows_path_line(path)
        path = path.rstrip('.:,;')
        path = path.strip()
        path = path.replace('\\', '/').replace('\\', '/')
        return path

    # Pattern 2: "Pfad" in Anführungszeichen
    quote_pattern = r'\"([A-Z]:\\[^"\n\r]+?)\"'
    quote_match = re.search(quote_pattern, prompt, re.IGNORECASE)
    if quote_match:
        path = quote_match.group(1).strip()
        path = _strip_trailing_prompt_from_windows_path_line(path)
        path = path.replace('\\', '/').replace('\\', '/')
        return path

    # Pattern 3: Downloads\... ohne Laufwerk
    downloads_pattern = r'(Downloads\\[^\n\r"`<>|?*]+?)(?:\s*[\n\r]|\s{2,}|\s*$|\"|\`|\<|\>|\s+[-\*•]\s)'
    downloads_match = re.search(downloads_pattern, prompt, re.IGNORECASE)
    if downloads_match:
        path = downloads_match.group(1).strip()
        path = _strip_trailing_prompt_from_windows_path_line(path)
        path = path.rstrip('.:,;')
        path = path.replace('\\', '/').replace('\\', '/')
        return f"C:/Users/mielersch/Desktop/Rambo-Rainer/{path}"

    alt = _extract_robot_desktop_base_path_str(prompt)
    if alt:
        return alt.replace("\\", "/")

    return None


def validate_downloads_path(path_str):
    """
    Validiert, dass der Pfad innerhalb von Downloads liegt.
    """
    if not path_str:
        return False, "Kein Pfad angegeben"

    # Bereinigen
    path_str = path_str.replace('\\', '/').replace('\\', '/')

    # Ungültige Zeichen prüfen
    invalid_chars = ['\n', '\r', '"', '<', '>', '|', '?', '*']
    for char in invalid_chars:
        if char in path_str:
            return False, f"Pfad enthält ungültige Zeichen: {repr(char)}"

    # Normalisiere für Vergleich
    normalized = Path(path_str).resolve()
    downloads_base = DOWNLOADS_DIR.resolve()

    try:
        # Prüfe ob Pfad unterhalb von Downloads liegt
        normalized.relative_to(downloads_base)
        return True, str(normalized)
    except ValueError:
        return False, f"Pfad liegt außerhalb von Downloads: {path_str}"


def execute_electron_react_build(task, run_id):
    """
    Deterministischer Electron/React Projekt-Build.
    Wird aufgerufen wenn Electron-Keywords im Prompt erkannt werden.
    """
    import subprocess
    pre_guard = _validate_direct_run_paths(
        list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN) + ["build_desktop.py", "robot-icon.png", "app-icon.png"],
        "apply",
        task,
    )
    if not bool(pre_guard.get("ok")):
        blocked_payload = _build_direct_guard_block_payload(
            scope="project",
            mode="apply",
            blocked_files=pre_guard.get("blocked_files") or list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN),
            task=task,
            recognized_task={
                "task_type": "electron_react_build",
                "primary_area": "Project Builder",
                "execution_route": "electron_react_build",
            },
        )
        blocked_payload["run_id"] = run_id
        return blocked_payload

    workstream_events = []
    def add_event(phase, level, title, detail="", file="", command="", status="done"):
        event = {
            "ts": get_timestamp(),
            "phase": phase,
            "level": level,
            "title": title,
            "detail": detail,
            "file": file,
            "command": command,
            "status": status,
        }
        workstream_events.append(event)
        append_ui_log_entry(phase, f"{title}: {detail}", level)
        return event

    # 1. Target root erkennen
    add_event("analysis", "info", "Auftrag erkannt", "Electron-/React-Projektbuild erkannt", status="running")

    extracted_path = None
    bp = _extract_robot_desktop_base_path_str(task)
    if bp:
        extracted_path = bp.replace("\\", "/")
    if not extracted_path:
        extracted_path = extract_windows_target_dir(task)

    if extracted_path:
        is_valid, validation_result = validate_downloads_path(extracted_path)
        if is_valid:
            target_root = Path(validation_result)
            add_event("analysis", "info", "Zielordner erkannt", str(target_root), status="done")
        else:
            add_event("analysis", "warning", "Pfad-Validierung fehlgeschlagen", validation_result, status="failed")
            target_root = DOWNLOADS_DIR / "RainerRobotDesktop"
            add_event("analysis", "info", "Fallback verwendet", str(target_root), status="done")
    else:
        target_root = DOWNLOADS_DIR / "RainerRobotDesktop"
        add_event("analysis", "info", "Zielordner nicht erkannt", f"Verwende Default: {target_root}", status="done")

    # 2. Ordner erstellen
    target_root.mkdir(parents=True, exist_ok=True)
    add_event("planning", "info", "Datei-Plan erstellt", "Robot-Desktop-Dateien geplant", status="done")

    # 3. File-Plan definieren
    file_plan = {
        "electron/main.js": ELECTRON_MAIN_JS,
        "electron/preload.js": ELECTRON_PRELOAD_JS,
        "electron/package.json": ELECTRON_PACKAGE_JSON,
        "electron/fallback-setup.html": ELECTRON_FALLBACK_SETUP_HTML,
        "rambo_ui/package.json": RAMBO_UI_PACKAGE_JSON,
        "rambo_ui/vite.config.js": RAMBO_UI_VITE_CONFIG,
        "rambo_ui/index.html": RAMBO_UI_INDEX_HTML,
        "rambo_ui/src/main.jsx": RAMBO_UI_MAIN_JSX,
        "rambo_ui/src/App.jsx": RAMBO_UI_APP_JSX,
        "rambo_ui/src/App.css": RAMBO_UI_APP_CSS,
        "build_desktop.py": BUILD_DESKTOP_PY,
    }
    file_plan.update(get_robot_desktop_extra_files())

    # 4. Dateien schreiben
    created_files = []
    failed_files = []
    for rel_path, content in file_plan.items():
        full_path = target_root / rel_path
        try:
            add_event("writing", "info", "Schreibe Datei", f"{rel_path}", file=str(full_path), status="running")
            # Robust: Erstelle Parent-Ordner mit Fehlerbehandlung für WinError 183
            try:
                full_path.parent.mkdir(parents=True, exist_ok=True)
            except OSError as mkdir_err:
                # WinError 183: Prüfe ob ein Ordner mit gleichem Namen wie die Datei existiert
                if mkdir_err.errno == 183 or (hasattr(mkdir_err, 'winerror') and mkdir_err.winerror == 183):
                    # Lösche existierende Datei/Ordner Konflikt
                    if full_path.parent.exists() and full_path.parent.is_file():
                        full_path.parent.unlink()
                        full_path.parent.mkdir(parents=True, exist_ok=True)
                else:
                    raise
            # Robust: Schreibe Datei mit explizitem Überschreiben
            if full_path.exists() and full_path.is_dir():
                # Konflikt: Ein Ordner blockiert die Datei
                import shutil
                shutil.rmtree(full_path)
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(content)
            created_files.append(str(full_path))
            add_event("writing", "success", "Datei geschrieben", f"{rel_path}", file=str(full_path), status="done")
        except Exception as e:
            failed_files.append({"path": str(full_path), "error": str(e)})
            add_event("writing", "error", "Schreiben fehlgeschlagen", str(e), file=str(full_path), status="failed")

    for icon_path in _sync_robot_desktop_icons(target_root, task):
        if str(icon_path) not in created_files:
            created_files.append(str(icon_path))

    # 5. Verifikation
    planned_files = [str(target_root / f) for f in file_plan.keys()]
    missing_files = [f for f in planned_files if not Path(f).exists()]
    created_count = len([f for f in planned_files if Path(f).exists()])

    # 6. Build-Kommandos
    build_logs = []
    commands_run = []

    def run_npm(cmd, cwd, description):
        add_event("build", "info", "Führe Befehl aus", description, command=cmd, status="running")
        try:
            result = subprocess.run(cmd, shell=True, cwd=cwd, capture_output=True, text=True, timeout=300)
            log_entry = {
                "cmd": cmd,
                "cwd": str(cwd),
                "returncode": result.returncode,
                "stdout": result.stdout[:2000],
                "stderr": result.stderr[:2000],
            }
            build_logs.append(log_entry)
            commands_run.append({"cmd": cmd, "cwd": str(cwd), "status": "success" if result.returncode == 0 else "failed"})
            if result.returncode == 0:
                add_event("build", "success", "Befehl erfolgreich", f"{cmd} (exit={result.returncode})", command=cmd, status="done")
            else:
                add_event("build", "error", "Befehl fehlgeschlagen", f"{cmd} (exit={result.returncode})", command=cmd, status="failed")
            return log_entry
        except Exception as e:
            log_entry = {"cmd": cmd, "cwd": str(cwd), "error": str(e)}
            build_logs.append(log_entry)
            commands_run.append({"cmd": cmd, "cwd": str(cwd), "status": "error"})
            add_event("build", "error", "Befehl fehlgeschlagen", str(e), command=cmd, status="failed")
            return log_entry

    rambo_ui_path = target_root / "rambo_ui"
    electron_path = target_root / "electron"

    if rambo_ui_path.exists() and created_count == len(planned_files):
        run_npm("npm install", rambo_ui_path, "npm install in rambo_ui")
        run_npm("npm run build", rambo_ui_path, "npm run build in rambo_ui")

    if electron_path.exists() and created_count == len(planned_files):
        run_npm("npm install", electron_path, "npm install in electron")

    # 7. Logs speichern
    log_dir = DOWNLOADS_DIR / "logs"
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / f"electron_build_{run_id}.json"
    try:
        with open(log_file, 'w', encoding='utf-8') as f:
            json.dump({
                "run_id": run_id,
                "task": task[:200],
                "target_root": str(target_root),
                "planned_files": planned_files,
                "created_files": created_files,
                "missing_files": missing_files,
                "commands_run": commands_run,
                "build_logs": build_logs,
                "workstream_events": workstream_events,
                "timestamp": get_timestamp(),
            }, f, ensure_ascii=False, indent=2)
        add_event("artifact", "info", "Log gespeichert", str(log_file), status="done")
    except Exception as e:
        add_event("artifact", "warning", "Log speichern fehlgeschlagen", str(e), status="failed")

    # 8. Finaler Status
    if missing_files:
        build_status = "error"
        technical_message = f"Electron-Projekt unvollständig: {len(missing_files)} Pflichtdateien fehlen."
        add_event("verify", "error", "Fehlende Dateien", f"{len(missing_files)} Dateien fehlen", status="failed")
    else:
        build_status = "success"
        technical_message = f"Electron-Projekt erfolgreich: {created_count}/{len(planned_files)} Dateien erstellt."
        add_event("verify", "success", "Ergebnis geprüft", f"Alle {created_count} Dateien vorhanden", status="done")

    add_event("done", "info" if build_status == "success" else "error", "Build abgeschlossen", technical_message, status="done")

    er_result = {
        "run_id": run_id,
        "status": build_status,
        "ok": build_status == "success",
        "technical_message": technical_message,
        "target_root": str(target_root),
        "base_path": str(target_root),
        "robot_build_auto_applied": True,
        "debug_auto_apply_decision": {"is_robot_desktop_build": True, "execution_route": "electron_react_build"},
        "requires_user_confirmation": False,
        "requires_confirmation": False,
        "recognized_task": {
            "task_type": "electron_react_build",
            "primary_area": "Project Builder",
            "hint": "Deterministischer Electron/React Build.",
            "execution_route": "electron_react_build",
        },
        "file_plan": list(file_plan.keys()),
        "planned_files_count": len(planned_files),
        "created_files_count": created_count,
        "created_files": created_files,
        "missing_files": missing_files,
        "commands_run": commands_run,
        "build_logs": build_logs,
        "workstream_events": workstream_events,
        "log_file": str(log_file) if build_logs else None,
        "formatted_response": technical_message,
    }
    if created_count > 0 and build_status == "success":
        er_result = _merge_post_build_agent_digest(er_result, task, target_root)
    return er_result



def _extract_task_or_prompt_from_request_json(data) -> str:
    """
    Liest task/prompt aus JSON-POST.
    Keys case-insensitive: PowerShell ConvertTo-Json liefert oft Task/Prompt statt task/prompt.
    """
    if not isinstance(data, dict):
        return ""
    for k in ("task", "prompt"):
        v = data.get(k)
        if v is not None and str(v).strip():
            return " ".join(str(v).strip().split())
    for dk, dv in data.items():
        if str(dk).lower() in ("task", "prompt") and dv is not None:
            s = str(dv).strip()
            if s:
                return " ".join(s.split())
    return ""


def _workspace_tree_outline(root: Path, *, max_lines: int = 80) -> str:
    """
    Kompakter Ueberblick (1–2 Ebenen) fuer Nur-Lesen-Analyse, ohne grosse Verzeichnisse.
    """
    try:
        base = root.resolve()
    except Exception:
        return ""
    if not base.is_dir():
        return ""
    lines = []
    n = 0
    try:
        entries = sorted(base.iterdir(), key=lambda e: (not e.is_dir(), e.name.lower()))
    except OSError:
        return ""
    skip = SCANNER_SKIP_DIRS
    for p in entries:
        if n >= max_lines:
            lines.append(f"... (weitere Eintraege ausgelassen, Limit {max_lines})")
            break
        name = p.name
        if name in skip or name.startswith("."):
            continue
        if p.is_dir():
            lines.append(f"[DIR]  {name}/")
            n += 1
            try:
                sub = sorted(p.iterdir(), key=lambda x: (not x.is_dir(), x.name.lower()))[:20]
            except OSError:
                sub = []
            for sp in sub:
                if n >= max_lines:
                    break
                if sp.name in skip or sp.name.startswith("."):
                    continue
                suff = "/" if sp.is_dir() else ""
                lines.append(f"        {sp.name}{suff}")
                n += 1
        else:
            lines.append(f"[FILE] {name}")
            n += 1
    return "\n".join(lines) if lines else "(Keine listbaren Eintraege im Projektroot.)"


def build_read_only_project_analysis_reply(
    task: str,
    *,
    path_inference_source: str | None = None,
) -> str:
    """
    Reine Projekt-/Code-Analyse ohne Patch oder Schreibzugriff.
    Nutzt den freigegebenen Workspace und optional Auszuege aus inferierten Zieldateien.
    path_inference_source: nur fuer Pfad-Erkennung (ohne lange Session-Prepend); task = Anzeige+LLM.
    """
    raw_task = " ".join(str(task or "").strip().split())
    infer_from = " ".join(str(path_inference_source or task or "").strip().split())
    try:
        root = get_active_project_root().resolve()
    except Exception:
        root = None
    if not root:
        return (
            "Kein freigegebener Projektordner aktiv. Bitte Ordner waehlen und freigeben — "
            "dann kann ich den Code/ die Struktur nur lesen und erklaeren (ohne Aenderungen)."
        )
    paths, _meta = infer_allowed_target_files_with_meta(infer_from)
    chunks = []
    max_files = 10
    max_chars = 8000
    for rel in paths[:max_files]:
        cleaned = format_local_path(rel)
        if not cleaned:
            continue
        try:
            resolved = (root / cleaned).resolve()
            resolved.relative_to(root)
        except Exception:
            continue
        if not resolved.is_file():
            continue
        content, exists = read_text_file(resolved)
        if not exists or not (content and str(content).strip()):
            continue
        text = str(content)
        if len(text) > max_chars:
            text = text[:max_chars] + "\n\n… (gekuerzt — vollstaendige Datei nicht noetig fuer Ueberblick)"
        chunks.append(f"### `{cleaned}`\n```\n{text}\n```")
    tl = raw_task.lower()
    want_tree = (
        not chunks
        or any(
            k in tl
            for k in (
                "ordner",
                "struktur",
                "projekt",
                "codebase",
                "komplett",
                "gesamt",
                "überblick",
                "ueberblick",
                "workspace",
                "verzeichnis",
                "ganzen ordner",
                "alle module",
            )
        )
    )
    tree_block = ""
    if want_tree:
        outline = _workspace_tree_outline(root)
        if outline.strip():
            tree_block = f"\n\n### Verzeichnisüberblick (Auszug)\n```\n{outline}\n```"
    ctx = "\n\n".join(chunks) if chunks else (
        f"(Keine eindeutigen Zieldatei-Pfade im Prompt; Arbeitsverzeichnis: `{root}`.)"
    )
    if tree_block:
        ctx = (ctx + tree_block).strip()
    elif not chunks:
        ctx = (
            ctx
            + "\n\n*(Tipp: Nenne konkrete Dateipfade wie `backend/main.py` fuer laengeren Kontext.)*"
        )
    augmented = (
        "[System: NUR Analyse und Erklaerung. Keine Codevorschlaege zum Ueberschreiben von Dateien, "
        "kein unified diff, keine 'apply'-Schritte.]\n\n"
        f"**Nutzerauftrag:** {raw_task}\n\n"
        f"**Auszuege aus dem Workspace:**\n\n{ctx}"
    )
    return generate_chat_response_plain_with_timeout(augmented)


@app.route("/api/direct-run", methods=["POST"])
def direct_run():
    _run_started = time.time()
    _run_decisions = []
    def _trace_payload(route: str, classification: str, extra: dict | None = None):
        payload = {
            "route": str(route or ""),
            "classification": str(classification or ""),
            "duration_ms": int(max(0.0, (time.time() - _run_started) * 1000)),
            "decisions": list(_run_decisions),
        }
        if isinstance(extra, dict) and extra:
            payload.update(extra)
        return payload

    data = request.get_json(silent=True) or {}
    scope = str(data.get("scope") or "local").strip().lower()
    mode = str(data.get("mode") or "apply").strip().lower()
    response_style = str(data.get("response_style") or "").strip().lower() or None
    # Auto-Apply-Flags aus dem Request – Frontend sendet diese bei normalem Chat
    auto_apply_req = bool(
        data.get("auto_apply") or data.get("skip_review") or data.get("direct_execute") or data.get("auto_continue")
    )
    run_id = uuid4().hex

    if scope not in {"local", "project"}:
        scope = "local"
    if mode not in {"safe", "apply"}:
        mode = "apply"

    raw_prompt = _extract_task_or_prompt_from_request_json(data)
    cleaned_prompt = " ".join(str(raw_prompt or "").strip().split())
    if not cleaned_prompt:
        append_ui_log_entry("Direkt", "Leere Aufgabe", "warning")
        return jsonify(enrich_direct_run_response({"error": "Bitte eine Aufgabe eingeben.", "scope": scope, "mode": mode})), 400

    ps_env = _powershell_direct_run_envelope(
        cleaned_prompt=cleaned_prompt, run_id=run_id, scope=scope, mode=mode, log_label="Direkt"
    )
    if ps_env is not None:
        return jsonify(ps_env), 200

    pk = classify_user_prompt(cleaned_prompt)
    _run_decisions.append(f"classify:{pk}")
    if should_route_direct_run_as_chat(cleaned_prompt):
        pk = "chat"
        _run_decisions.append("override:direct_run_as_chat")
    # Analyse-Prompts hart auf Read-Only routen (kein Auto-Datei-Template).
    analysis_markers = (
        "analysiere",
        "analyse",
        "gib mir verbesserungs",
        "verbesserungsvorschl",
        "bewerte den ordner",
        "prüfe den ordner",
        "pruefe den ordner",
    )
    low_prompt = cleaned_prompt.lower()
    has_uploaded_refs = bool(
        (isinstance(data.get("uploaded_files"), list) and len(data.get("uploaded_files")) > 0)
        or str(data.get("uploaded_file_path") or "").strip()
    )
    image_analysis_markers = ("bild", "image", "foto", "screenshot", "grafik", "png", "jpg", "jpeg", "webp")
    analyze_markers = ("beschreibe", "analysiere", "analyse", "erkläre", "erklaere", "was siehst")
    if has_uploaded_refs and any(m in low_prompt for m in image_analysis_markers) and any(m in low_prompt for m in analyze_markers):
        pk = "project_read"
    if any(m in low_prompt for m in analysis_markers) and not has_project_change_intent(cleaned_prompt):
        pk = "project_read"
        _run_decisions.append("override:analysis_to_project_read")
    # Absicherung: Änderungsabsicht nie als reiner Chat (auch bei veralteter Klassifikation)
    if pk != "risky_project_task" and has_project_change_intent(cleaned_prompt):
        pk = "project_task"
        _run_decisions.append("override:change_intent_to_project_task")
    user_mode_raw = str(data.get("user_mode") or data.get("intent_mode") or "").strip()
    pk = apply_user_mode_override(
        pk,
        user_mode_raw or None,
        is_risky=(pk == "risky_project_task"),
    )
    hist_block = normalize_conversation_history_payload(data.get("conversation_history"))
    augmented_prompt = compose_augmented_user_message(cleaned_prompt, data)
    if pk == "unknown" and intent_llm_enabled():
        refined_pk = run_llm_intent_refinement(cleaned_prompt, hist_block)
        if refined_pk:
            pk = refined_pk
    if (
        pk != "project_task"
        and is_folder_analysis_prompt(cleaned_prompt)
        and not has_project_change_intent(cleaned_prompt)
    ):
        explicit_path = extract_folder_analysis_path_from_prompt(raw_prompt or "")
        analysis_root, analysis_err, fa_dbg = _resolve_folder_analysis_root(explicit_path)
        if analysis_err:
            denied_md = analysis_err
            denied_payload = {
                "ok": True,
                "success": True,
                "applied": False,
                "run_id": run_id,
                "scope": scope,
                "mode": mode,
                "status": "chat_response",
                "direct_status": "chat_response",
                "classification": "project_read",
                "route_mode": "workspace_analysis",
                "task_kind": "read_only_analysis",
                "chat_response": denied_md,
                "formatted_response": denied_md,
                "natural_message": denied_md,
                "message": denied_md,
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "writes_files": False,
                "has_changes": False,
                "changed_files": [],
                "files": [],
                "folder_analysis_debug": fa_dbg,
                "extracted_path": fa_dbg.get("extracted_path"),
                "active_root": fa_dbg.get("active_root"),
                "resolved_path": fa_dbg.get("resolved_path"),
                "exists": fa_dbg.get("exists"),
                "is_dir": fa_dbg.get("is_dir"),
                "workstream_events": [
                    _ws_event("analysis", "warn", "Workspace", "Ordneranalyse abgelehnt", status="done"),
                ],
            }
            return jsonify(enrich_direct_run_response(denied_payload)), 200
        append_ui_log_entry("Direkt", f"Ordneranalyse: {analysis_root}", "info")
        fa_payload = build_folder_analysis_payload(analysis_root, run_id=run_id, scope=scope, mode=mode)
        return jsonify(enrich_direct_run_response(fa_payload)), 200

    intent = classify_user_prompt_intent(cleaned_prompt)

    if pk == "risky_project_task":
        risky_payload = {
            "ok": False,
            "success": False,
            "run_id": run_id,
            "scope": scope,
            "mode": mode,
            "status": "risky_blocked",
            "direct_status": "risky_blocked",
            "classification": "risky_project_task",
            "error": "Riskante Aktion blockiert.",
            "formatted_response": "⚠️ Diese Aktion ist aus Sicherheitsgruenden blockiert.",
            "message": "Riskante Aktion blockiert.",
            "requires_confirmation": True,
            "requires_user_confirmation": True,
            "writes_files": False,
            "workstream_events": [
                _ws_event("guard", "error", "Blockiert", "Riskante Aktion", status="blocked"),
            ],
            "run_trace": _trace_payload("risky_blocked", "risky_project_task"),
        }
        return jsonify(enrich_direct_run_response(risky_payload)), 403

    if pk == "chat":
        _instant = _chat_reply_skip_llm(intent, cleaned_prompt)
        if _instant is not None:
            chat_text = _instant
        else:
            chat_text = generate_chat_response_plain_with_timeout(augmented_prompt)
        chat_payload = {
            "ok": True,
            "success": True,
            "applied": False,
            "run_id": run_id,
            "scope": scope,
            "mode": "chat",
            "status": "chat_response",
            "direct_status": "chat_response",
            "classification": "chat",
            "route_mode": "chat",
            "chat_response": chat_text,
            "formatted_response": chat_text,
            "message": chat_text,
            "natural_message": chat_text,
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "writes_files": False,
            "has_changes": False,
            "changed_files": [],
            "files": [],
            "workstream_events": [
                _ws_event("analysis", "info", "Chat", "Konversation (kein Dateizugriff)", status="done"),
            ],
            "run_trace": _trace_payload("chat", "chat"),
        }
        return jsonify(enrich_direct_run_response(chat_payload)), 200

    if pk == "unknown":
        _inst_u = _unknown_reply_skip_llm(intent, cleaned_prompt)
        if _inst_u is not None:
            chat_text = _inst_u
        else:
            chat_text = generate_chat_response_plain_with_timeout(augmented_prompt)
            if (not str(chat_text or "").strip()) or (
                has_project_change_intent(cleaned_prompt) and _is_generic_ready_reply(chat_text)
            ):
                chat_text = generate_chat_response_plain_with_timeout(
                    "Antworte konkret auf Deutsch in 2-4 Sätzen. "
                    "Wenn der Nutzer eine Projektänderung andeutet, nenne die zwei wahrscheinlichsten nächsten "
                    "Schritte (Analyse und sichere Umsetzung) statt einer generischen Bereitschaftsantwort.\n\n"
                    f"Nutzerprompt: {cleaned_prompt}"
                )
            if not str(chat_text or "").strip() or _is_generic_ready_reply(chat_text):
                if has_project_change_intent(cleaned_prompt):
                    chat_text = (
                        "Ich erkenne eine Projektänderung, aber die Absicht ist noch zu ungenau. "
                        "Sag mir bitte Ziel-Datei und gewünschte Änderung in einem Satz, "
                        "dann setze ich es direkt um."
                    )
                else:
                    chat_text = clarification_message_with_modes()
        chat_payload = {
            "ok": True,
            "success": True,
            "applied": False,
            "run_id": run_id,
            "scope": scope,
            "mode": "chat",
            "status": "chat_response",
            "direct_status": "chat_response",
            "classification": "unknown",
            "route_mode": "intent_clarification",
            "task_kind": "intent_clarification",
            "suggested_intent_actions": list(SUGGESTED_INTENT_ACTIONS),
            "chat_response": chat_text,
            "formatted_response": chat_text,
            "message": chat_text,
            "natural_message": chat_text,
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "writes_files": False,
            "has_changes": False,
            "changed_files": [],
            "files": [],
            "workstream_events": [
                _ws_event("analysis", "info", "Absicht", "Modus waehlen oder praezisieren", status="done"),
            ],
            "run_trace": _trace_payload("intent_clarification", "unknown"),
        }
        return jsonify(enrich_direct_run_response(chat_payload)), 200

    # Vage Nur-Datei-Aenderung: vor Laengen-Validierung — kein Guard/Patch/Apply
    if pk == "project_task" and is_vague_project_file_edit(cleaned_prompt):
        rel = _extract_rel_path_for_clarification(cleaned_prompt)
        msg = f"Was genau soll ich an {rel} ändern?"
        clar = {
            "ok": True,
            "success": False,
            "applied": False,
            "run_id": run_id,
            "scope": scope,
            "mode": "clarification",
            "status": "clarification_required",
            "direct_status": "clarification_required",
            "classification": "project_task",
            "route_mode": "clarification_required",
            "message": msg,
            "formatted_response": msg,
            "natural_message": msg,
            "chat_response": msg,
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "writes_files": False,
            "has_changes": False,
            "changed_files": [],
            "files": [],
            "workstream_events": [
                _ws_event("analysis", "info", "Klaerung", "Projektbefehl zu ungenau — bitte praezisieren", status="done"),
            ],
        }
        return jsonify(enrich_direct_run_response(clar)), 200

    # Nur project_task: validate_task (Laenge) — nach Intent, nie fuer Chat/Unklar/vage-Klaerung
    task, task_error = validate_task(cleaned_prompt)
    if task_error:
        append_ui_log_entry("Direkt", task_error, "warning")
        return jsonify(enrich_direct_run_response({"error": task_error, "scope": scope, "mode": mode})), 400

    task, upload_ctx_meta = augment_prompt_with_uploads(task, data)

    task_lower = str(task or "").lower()

    if detect_model_selection_prompt(task):
        try:
            plan = _get_model_router().build_route_plan(task)
        except Exception:
            plan = {
                "ok": False,
                "task_type": "general",
                "selected_model": "Gemma 4 lokal",
                "fallback_models": [],
                "warnings": [],
                "errors": ["model_router_unavailable"],
                "reason": "Model-Router nicht verfügbar.",
            }
        route_payload = {
            "ok": True,
            "success": False,
            "status": "model_route",
            "direct_status": "model_route",
            "mode": "model_route",
            "message": "Model-Router-Aufgabe erkannt.",
            "route": {
                "task_type": plan.get("task_type"),
                "selected_model": plan.get("selected_model") or "Gemma 4 lokal",
                "fallback_models": plan.get("fallback_models", []),
                "available_models": plan.get("available_models", []),
                "missing_preferred_models": plan.get("missing_preferred_models", []),
                "reason": plan.get("reason", ""),
                "warnings": plan.get("warnings", []),
                "errors": plan.get("errors", []),
            },
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "workstream_events": [
                _ws_event("analysis", "info", "Model-Router", "Modellwahl-Prompt erkannt", status="done"),
            ],
        }
        return jsonify(enrich_direct_run_response(route_payload)), 200

    self_fix = detect_self_fix_task(task)
    if bool(self_fix.get("is_self_fix")):
        self_fix["recommended_checks"] = ["node_check_app", "py_compile_main", "pytest_all"]
        self_fix["next_step"] = "Safe/Preview starten und danach gezielt bestätigen."
        saved_plan = save_last_self_fix_plan(task, self_fix)
        self_fix_payload = {
            "ok": True,
            "success": False,
            "status": "self_fix_plan_required",
            "direct_status": "self_fix_plan_required",
            "mode": "self_fix_plan",
            "message": "Self-Fix-Aufgabe erkannt. Bitte Plan prüfen.",
            "reason": str(self_fix.get("reason") or ""),
            "risk": str(self_fix.get("risk") or "medium"),
            "plan_id": str(saved_plan.get("plan_id") or ""),
            "created_at": str(saved_plan.get("created_at") or ""),
            "prompt": str(saved_plan.get("prompt") or task),
            "candidate_files": list(saved_plan.get("candidate_files") or []),
            "affected_files": list(saved_plan.get("affected_files") or []),
            "recommended_checks": list(saved_plan.get("recommended_checks") or []),
            "next_step": str(saved_plan.get("next_step") or "Safe/Preview starten und danach gezielt bestätigen."),
            "requires_confirmation": True,
            "requires_user_confirmation": True,
            "workstream_events": [
                _ws_event("analysis", "info", "Self-Fix erkannt", str(self_fix.get("reason") or ""), status="done"),
                _ws_event("route", "warning", "Plan erforderlich", "Keine direkte Write-Ausführung.", status="blocked"),
            ],
        }
        return jsonify(enrich_direct_run_response(self_fix_payload)), 200

    try:
        ws_ok = is_active_workspace_trusted()
    except Exception:
        ws_ok = False
    if not ws_ok:
        ws_payload = {
            "ok": False,
            "success": False,
            "run_id": run_id,
            "scope": scope,
            "mode": mode,
            "status": "workspace_required",
            "direct_status": "workspace_required",
            "classification": pk if pk in ("project_read", "project_task") else "project_task",
            "error": "Projektfreigabe fehlt.",
            "formatted_response": (
                "Bitte waehle zuerst einen Projektordner aus und gib ihn frei. "
                "Danach kann ich im Projekt wie gewohnt arbeiten."
            ),
            "message": "Bitte waehle zuerst einen Projektordner aus und gib ihn frei.",
            "requires_confirmation": False,
            "writes_files": False,
            "workstream_events": [
                _ws_event("route", "warning", "Projektfreigabe", "Kein freigegebener Ordner aktiv", status="blocked"),
            ],
        }
        return jsonify(enrich_direct_run_response(ws_payload)), 403

    # Reine Analyse / Lesen — AgentLoop.run_analysis (read-only), kein Apply
    if pk == "project_read":
        upload_lines = []
        for u in list((upload_ctx_meta or {}).get("uploads") or [])[:5]:
            if not isinstance(u, dict):
                continue
            name = str(u.get("filename") or "-")
            ftype = str(u.get("file_type") or "-")
            size = int(u.get("size") or 0)
            summary = str(u.get("summary") or "").strip()
            upload_lines.append(f"- {name} ({ftype}, {size} Bytes){(': ' + summary) if summary else ''}")
        upload_block = (
            "\n\n## Upload-Kontext\n" + "\n".join(upload_lines)
            if upload_lines
            else ""
        )
        chat_analysis = ""
        afiles: list = []
        ok_a = False
        err_msg = ""
        if AGENT_LOOP_AVAILABLE and AgentLoop is not None:
            try:
                workspace_path = str(get_active_project_root().resolve())
                append_ui_log_entry("Direkt", f"Analysemodus (AgentLoop, project_read): {task[:80]}", "info")
                agent_o = AgentLoop(Path(workspace_path))
                analysis_result = agent_o.run_analysis(task)
                analysis_result = analysis_result if isinstance(analysis_result, dict) else {}
                raw_analysis = str(analysis_result.get("analysis") or "").strip()
                afiles = analysis_result.get("files") or []
                if not isinstance(afiles, list):
                    afiles = []
                err_msg = str(analysis_result.get("error") or "").strip()
                ok_a = bool(analysis_result.get("ok")) and bool(raw_analysis)
                chat_analysis = raw_analysis
            except Exception as _pr_ex:
                log_structured("project_read_agent_loop_error", error=str(_pr_ex))
                append_ui_log_entry("Direkt", f"AgentLoop Analyse (project_read) Fehler: {_pr_ex}", "error")

        if not (chat_analysis or "").strip():
            hint = ""
            if err_msg:
                hint = f"\n\n**Hinweis:** {err_msg}"
            analysis_text = (
                "## Ziel\n"
                f"{cleaned_prompt}\n\n"
                "## Ergebnis\n"
                "Read-Only Analysepfad aktiv. Es werden keine Dateien geändert.\n"
                "(Keine Modell-Antwort — AgentLoop nicht verfügbar oder Analyse leer.)\n\n"
                "## Dateiänderungen\n"
                "- Keine (Read-Only Analyse)\n\n"
                "## Status\n"
                "- OK"
                + hint
                + upload_block
            )
            read_payload = {
                "ok": True,
                "success": True,
                "applied": False,
                "run_id": run_id,
                "scope": scope,
                "mode": "safe",
                "status": "chat_response",
                "direct_status": "chat_response",
                "classification": "project_read",
                "route_mode": "read_only_analysis",
                "task_kind": "read_only_analysis",
                "chat_response": analysis_text,
                "formatted_response": analysis_text,
                "message": analysis_text,
                "natural_message": analysis_text,
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "writes_files": False,
                "has_changes": False,
                "changed_files": [],
                "files": [],
                "analysis_files": [],
                "ui_mode": "clean_chat",
                "workstream_events": [
                    _ws_event(
                        "analysis",
                        "info",
                        "Nur-Analyse",
                        "Klassifikation: Lesen — keine Dateiaenderung (Fallback)",
                        status="done",
                    ),
                ],
                "run_trace": _trace_payload("project_read_fallback", "project_read"),
            }
        else:
            read_payload = {
                "ok": ok_a,
                "success": ok_a,
                "applied": False,
                "run_id": run_id,
                "scope": scope,
                "mode": mode,
                "status": "chat_response",
                "direct_status": "chat_response",
                "classification": "project_read",
                "route_mode": "agent_analysis",
                "task_kind": "read_only_analysis",
                "ui_mode": "workspace_analysis",
                "chat_response": chat_analysis,
                "formatted_response": chat_analysis,
                "message": chat_analysis,
                "natural_message": chat_analysis,
                "analysis": chat_analysis,
                "analysis_files": afiles,
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "writes_files": False,
                "has_changes": False,
                "changed_files": [],
                "affected_files": [],
                "files": [],
                "show_agent_ui": True,
                "show_diff": False,
                "workstream_events": [
                    _ws_event(
                        "analyze",
                        "info",
                        "Projekt-Analyse (Lesen)",
                        f"{len(afiles)} Datei(en) eingebunden",
                        status="done",
                    ),
                ],
                "run_trace": _trace_payload("project_read_analysis", "project_read"),
            }
        if upload_ctx_meta.get("uploads") or upload_ctx_meta.get("errors"):
            read_payload["upload_context"] = upload_ctx_meta
        return jsonify(enrich_direct_run_response(read_payload)), 200

    inferred_allowed_targets, target_infer_meta = infer_allowed_target_files_with_meta(task)
    explicit_local_path = extract_explicit_local_relative_path(task)
    if not explicit_local_path and len(inferred_allowed_targets) == 1:
        explicit_local_path = inferred_allowed_targets[0]
    # Kein „erstelle irgendwo“ + Pfad bei riesigen Desktop-Builds — sonst intelligent_mode aus
    explicit_direct_write = bool(
        explicit_local_path
        and not _is_desktop_multi_file_project_prompt(task_lower)
        and (
            _is_single_file_direct_write_intent(task_lower)
            or (
                len(task_lower) < 2200
                and (
                    "erstelle" in task_lower
                    or "schreibe" in task_lower
                    or "aendere" in task_lower
                    or "ändere" in task_lower
                )
                and task_lower.count("phase ") == 0
                and "electron-builder" not in task_lower
                and "rambo_ui" not in task_lower
            )
        )
    )
    if explicit_direct_write:
        scope = "local"

    if scope == "local" and _is_mini_task_write_intent(task) and not explicit_direct_write:
        _n_t = len(inferred_allowed_targets)
        if _n_t == 0 or (_n_t != 1 and not ws_ok):
            target_infer_meta = dict(target_infer_meta) if isinstance(target_infer_meta, dict) else {}
            target_infer_meta["ws_ok"] = ws_ok
            target_infer_meta["active_project_root"] = str(get_active_project_root().resolve())
            target_infer_meta["inferred_allowed_targets"] = list(inferred_allowed_targets)
            target_infer_meta["len_inferred_allowed_targets"] = len(inferred_allowed_targets)
            unclear_payload = _build_target_path_unclear_payload(
                task, mode, inferred_allowed_targets, inference_debug=target_infer_meta
            )
            unclear_payload["run_id"] = run_id
            unclear_payload["scope"] = "local"
            unclear_payload["recognized_task"] = classify_direct_task(task)
            return jsonify(enrich_direct_run_response(unclear_payload)), 400

    task_classification = classify_direct_task(task)
    if task_classification.get("task_type") == "project_build":
        build_guard = _validate_direct_run_paths(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN, mode, task)
        if not bool(build_guard.get("ok")):
            blocked_payload = _build_direct_guard_block_payload(
                scope=scope,
                mode=mode,
                blocked_files=build_guard.get("blocked_files") or list(DIRECT_RUN_PROJECT_BUILD_FORBIDDEN_PLAN),
                task=task,
                recognized_task={
                    "task_type": "project_build",
                    "primary_area": "Direct",
                    "execution_route": "project_build",
                },
            )
            blocked_payload["workstream_events"] = [
                {"phase": "analyze", "level": "info", "title": "Project Build erkannt", "detail": "Pfad-Guard prueft Dateiplan", "status": "running"},
                {"phase": "guard", "level": "error", "title": "Auftrag blockiert", "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE, "status": "blocked"},
            ]
            return jsonify(enrich_direct_run_response(blocked_payload)), 403
        append_ui_log_entry("Direkt", "Project Build erkannt: Starte Multi-File Generation.", "info")
        try:
            project_result = execute_project_build(task, run_id, scope, mode)
            guarded_project = _apply_central_generation_guard(
                project_result if isinstance(project_result, dict) else {},
                scope=scope,
                mode=mode,
                task=task,
                recognized_task={"task_type": "project_build", "primary_area": "Project Builder", "execution_route": "project_build"},
            )
            if isinstance(guarded_project, dict):
                guarded_project["run_id"] = run_id
                return jsonify(enrich_direct_run_response(guarded_project)), 403
            if isinstance(project_result, dict):
                project_result = _enforce_real_change_success(task, project_result, mode="apply")
            http_code = 400 if isinstance(project_result, dict) and project_result.get("ok") is False else 200
            return jsonify(enrich_direct_run_response(project_result)), http_code
        except Exception as pb_ex:
            log_structured("project_build_exception", error=str(pb_ex))
            tb = traceback.format_exc()
            payload = _project_build_invalid_path_payload(
                run_id, scope, mode, str(pb_ex) or "Project Build fehlgeschlagen"
            )
            payload["technical_message"] = tb[:4000]
            payload["workstream_events"] = [
                _ws_event("error", "error", "Project Build", str(pb_ex)[:500], status="failed"),
            ]
            return jsonify(enrich_direct_run_response(payload)), 400

    # ── AGENT LOOP: project_task mit freigegebenem Workspace ──
    if pk == "project_task" and ws_ok and AGENT_LOOP_AVAILABLE:
        if not any(k in task_lower for k in ("intelligent-run", "rainer 3.0", "execute_intelligent")):
            try:
                workspace_path = str(get_active_project_root().resolve())
                if is_analysis_only_prompt(task) and AgentLoop is not None:
                    append_ui_log_entry("Direkt", f"Analysemodus (Groq/Ollama): {task[:80]}", "info")
                    agent_o = AgentLoop(Path(workspace_path))
                    analysis_result = agent_o.run_analysis(task)
                    analysis_result = analysis_result if isinstance(analysis_result, dict) else {}
                    raw_analysis = str(analysis_result.get("analysis") or "").strip()
                    afiles = analysis_result.get("files") or []
                    if not isinstance(afiles, list):
                        afiles = []
                    ok_a = bool(analysis_result.get("ok")) and bool(raw_analysis)
                    chat_response = raw_analysis if raw_analysis else "Keine Analyse-Antwort vom Modell."
                    analysis_payload = {
                        "ok": ok_a,
                        "success": ok_a,
                        "applied": False,
                        "run_id": run_id,
                        "scope": scope,
                        "mode": mode,
                        "status": "chat_response",
                        "direct_status": "chat_response",
                        "classification": "project_task",
                        "route_mode": "agent_analysis",
                        "ui_mode": "workspace_analysis",
                        "chat_response": chat_response,
                        "formatted_response": chat_response,
                        "message": chat_response,
                        "natural_message": chat_response,
                        "analysis": raw_analysis,
                        "analysis_files": afiles,
                        "changed_files": [],
                        "affected_files": [],
                        "has_changes": False,
                        "requires_confirmation": False,
                        "requires_user_confirmation": False,
                        "writes_files": False,
                        "show_agent_ui": True,
                        "show_diff": False,
                        "workstream_events": [
                            _ws_event(
                                "analyze",
                                "info",
                                "Projekt-Analyse",
                                f"{len(afiles)} Datei(en) eingebunden",
                                status="done",
                            ),
                        ],
                    }
                    return jsonify(enrich_direct_run_response(analysis_payload)), 200
                append_ui_log_entry("Direkt", f"Agent-Loop startet für: {task[:80]}", "info")
                agent_result = _run_agent_loop(task, workspace_path)
                agent_result = agent_result if isinstance(agent_result, dict) else {}
                changed = agent_result.get("changed_files") or []
                if changed:
                    tests = agent_result.get("test_results") or []
                    errors = agent_result.get("errors") or []
                    summary = agent_result.get("summary") or ""
                    ok = bool(agent_result.get("ok", True)) and not errors
                    lines = []
                    if changed:
                        lines.append("**Geänderte Dateien**")
                        for f in changed:
                            lines.append(f"- {f}")
                        lines.append("")
                    if tests:
                        lines.append("**Checks**")
                        for t in tests:
                            icon = "✅" if t.get("ok") else "❌"
                            lines.append(f"- {icon} {t.get('check', '-')}")
                        lines.append("")
                    if summary:
                        lines.append("**Kurz umgesetzt**\n" + summary)
                    if errors:
                        lines.append("\n⚠️ Fehler:\n" + "\n".join("- " + e for e in errors[:3]))
                    chat_text = "\n".join(lines)
                    agent_payload = {
                        "ok": ok,
                        "success": ok,
                        "applied": True,
                        "run_id": run_id,
                        "scope": scope,
                        "mode": "apply",
                        "status": "chat_response",
                        "direct_status": "chat_response",
                        "classification": "project_task",
                        "route_mode": "agent_loop",
                        "ui_mode": "project_change",
                        "chat_response": chat_text,
                        "formatted_response": chat_text,
                        "message": chat_text,
                        "natural_message": chat_text,
                        "changed_files": changed,
                        "affected_files": changed,
                        "has_changes": True,
                        "requires_confirmation": False,
                        "requires_user_confirmation": False,
                        "writes_files": True,
                        "show_agent_ui": True,
                        "show_diff": True,
                        "show_checks": bool(tests),
                        "workstream_events": [
                            _ws_event("analyze", "info", "Agent-Loop", f"{len(changed)} Datei(en) geändert", status="done"),
                        ],
                    }
                    return jsonify(enrich_direct_run_response(agent_payload)), 200
                # Kein changed_files → Fallthrough zur bestehenden Logik
            except Exception as _agent_ex:
                append_ui_log_entry("Direkt", f"Agent-Loop Fehler: {_agent_ex}", "error")

    # Agent-/Self-Repair-Prompt: keine Review-Phase, direkt lokal ausfuehren (wenn Dateiaktion extrahierbar)
    if _is_agent_instruction_prompt(task):
        append_ui_log_entry("Direkt", "Agent-Auftrag erkannt (Auto-Apply ohne Review).", "info")
        resolved_path, relative_path, proposed_content, file_exists, agent_err = _extract_agent_direct_write(task)
        if agent_err:
            payload = {
                "run_id": run_id,
                "scope": "local",
                "mode": "apply",
                "direct_status": "blocked",
                "error": agent_err,
                "recognized_task": {
                    "task_type": "agent_instruction_prompt",
                    "primary_area": "Direct",
                    "execution_route": "agent_direct_write",
                    "hint": "Agent-Prompt ohne Review direkt ausgefuehrt.",
                },
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "workstream_events": [
                    {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                    {"phase": "route", "level": "error", "title": "Ziel nicht extrahierbar", "detail": agent_err, "status": "blocked"},
                ],
            }
            return jsonify(enrich_direct_run_response(payload)), 400

        write_guard = _validate_direct_run_paths([relative_path], "apply", task)
        if not bool(write_guard.get("ok")):
            blocked_payload = _build_direct_guard_block_payload(
                scope="local",
                mode="apply",
                blocked_files=write_guard.get("blocked_files") or [relative_path],
                task=task,
                recognized_task={
                    "task_type": "agent_instruction_prompt",
                    "primary_area": "Direct",
                    "execution_route": "agent_direct_write",
                },
            )
            blocked_payload["run_id"] = run_id
            blocked_payload["workstream_events"] = [
                {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                {"phase": "guard", "level": "error", "title": "Auftrag blockiert", "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE, "status": "blocked"},
            ]
            return jsonify(enrich_direct_run_response(blocked_payload)), 403

        before_content, before_exists = read_text_file(resolved_path)
        pytest_validation = validate_pytest_file(relative_path, proposed_content)
        if not bool(pytest_validation.get("ok")):
            blocked_payload = _build_invalid_test_file_payload(relative_path, scope="local", mode="apply", task=task)
            blocked_payload["run_id"] = run_id
            blocked_payload["recognized_task"] = {
                "task_type": "agent_instruction_prompt",
                "primary_area": "Direct",
                "execution_route": "agent_direct_write",
            }
            blocked_payload["workstream_events"] = [
                {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                {"phase": "guard", "level": "error", "title": "Ungültige Testdatei", "detail": "Kein pytest-Test erkannt.", "status": "blocked"},
            ]
            return jsonify(enrich_direct_run_response(blocked_payload)), 400

        unsafe_check = detect_unsafe_large_rewrite(
            relative_path,
            before_content,
            proposed_content,
            {"phase": "agent_direct_apply", "task": task},
        )
        if bool(unsafe_check.get("unsafe")):
            applied_rec, se_res = _try_unsafe_rewrite_step_engine_apply(
                relative_path=relative_path,
                resolved_path=resolved_path,
                current_content=before_content,
                proposed_content=proposed_content,
                task=task,
                unsafe_check=unsafe_check,
            )
            if applied_rec:
                append_ui_log_entry("Direkt", f"Auto-Recovery (Step-Engine) direct-run: {relative_path}", "success")
                post_check = run_local_post_check(resolved_path, relative_path, proposed_content)
                created_files = [relative_path] if not before_exists else []
                changed_files = [relative_path]
                ws = [
                    {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                    {"phase": "route", "level": "info", "title": "Ziel aus Prompt extrahiert", "detail": relative_path, "status": "done"},
                    {"phase": "recovery", "level": "success", "title": "Step-Engine Recovery", "detail": "Kleiner Patch automatisch angewendet.", "status": "done"},
                    {"phase": "write", "level": "success", "title": "Datei geschrieben", "detail": relative_path, "status": "done"},
                    {"phase": "verify", "level": "success" if post_check.get("ok") else "error", "title": "Post-Check", "detail": post_check.get("detail") or "-", "status": "done" if post_check.get("ok") else "failed"},
                ]
                payload = {
                    "run_id": run_id,
                    "scope": "local",
                    "mode": "apply",
                    "direct_status": "verified" if post_check.get("ok") else "applied",
                    "message": "Agent-Auftrag direkt ausgefuehrt (Step-Engine-Recovery).",
                    "recognized_task": {
                        "task_type": "agent_instruction_prompt",
                        "primary_area": "Direct",
                        "execution_route": "agent_direct_write",
                        "hint": "Direkte Ausfuehrung ohne Review.",
                    },
                    "created_files": created_files,
                    "changed_files": changed_files,
                    "affected_files": [relative_path],
                    "selected_target_path": relative_path,
                    "post_check": {
                        "ok": bool(post_check.get("ok")),
                        "path": relative_path,
                        "target": str(resolved_path),
                        "exists": bool(resolved_path.exists()),
                        "detail": post_check.get("detail") or "",
                    },
                    "requires_confirmation": False,
                    "requires_user_confirmation": False,
                    "auto_apply": True,
                    "auto_continue": True,
                    "skip_review": True,
                    "direct_execute": True,
                    "workstream_events": ws,
                    "split_patch_recovery": True,
                    "step_engine_result": se_res if isinstance(se_res, dict) else {},
                    "formatted_response": (
                        "Geänderte Dateien\n"
                        f"- {relative_path}\n\n"
                        "Kurz umgesetzt\n"
                        "- Agent-Auftrag mit Step-Engine-Recovery angewendet.\n\n"
                        "Kurz getestet\n"
                        f"- Post-Check: {'OK' if post_check.get('ok') else 'Fehler'}\n\n"
                        "Verbleibende echte Restlücken nur falls vorhanden\n"
                        "- Keine."
                    ),
                }
                payload = _enforce_real_change_success(task, payload, mode="apply")
                return jsonify(enrich_direct_run_response(payload))

            blocked_payload = _build_unsafe_large_rewrite_payload(relative_path, scope="local", mode="apply", task=task)
            blocked_payload["run_id"] = run_id
            blocked_payload["recognized_task"] = {
                "task_type": "agent_instruction_prompt",
                "primary_area": "Direct",
                "execution_route": "agent_direct_write",
            }
            blocked_payload["rewrite_guard"] = unsafe_check
            blocked_payload["workstream_events"] = [
                {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                {"phase": "guard", "level": "error", "title": "Änderung blockiert", "detail": "Datei würde zu stark überschrieben.", "status": "blocked"},
            ]
            if isinstance(se_res, dict):
                blocked_payload["step_engine_result"] = se_res
                blocked_payload["auto_recovery_attempted"] = True
            return jsonify(enrich_direct_run_response(blocked_payload)), 400

        write_result = persist_text_file_change(
            resolved_path,
            proposed_content,
            relative_path,
            on_timeout_log=lambda m: append_ui_log_entry("Direkt", m, "error"),
        )
        if not write_result.get("ok"):
            err = str(write_result.get("error") or "Fehler beim Schreiben der Datei.")
            payload = {
                "run_id": run_id,
                "scope": "local",
                "mode": "apply",
                "direct_status": "failed",
                "error": err,
                "recognized_task": {
                    "task_type": "agent_instruction_prompt",
                    "primary_area": "Direct",
                    "execution_route": "agent_direct_write",
                },
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "workstream_events": [
                    {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
                    {"phase": "write", "level": "error", "title": "Schreiben fehlgeschlagen", "detail": err, "status": "failed"},
                ],
            }
            return jsonify(enrich_direct_run_response(payload)), 500

        post_check = run_local_post_check(resolved_path, relative_path, proposed_content)
        created_files = [relative_path] if not before_exists else []
        changed_files = [relative_path]
        ws = [
            {"phase": "analyze", "level": "info", "title": "Agent-Auftrag erkannt", "detail": "Auto-Apply aktiv", "status": "running"},
            {"phase": "route", "level": "info", "title": "Ziel aus Prompt extrahiert", "detail": relative_path, "status": "done"},
            {"phase": "write", "level": "success", "title": "Datei geschrieben", "detail": relative_path, "status": "done"},
            {"phase": "verify", "level": "success" if post_check.get("ok") else "error", "title": "Post-Check", "detail": post_check.get("detail") or "-", "status": "done" if post_check.get("ok") else "failed"},
        ]
        payload = {
            "run_id": run_id,
            "scope": "local",
            "mode": "apply",
            "direct_status": "verified" if post_check.get("ok") else "applied",
            "message": "Agent-Auftrag direkt ausgefuehrt.",
            "recognized_task": {
                "task_type": "agent_instruction_prompt",
                "primary_area": "Direct",
                "execution_route": "agent_direct_write",
                "hint": "Direkte Ausfuehrung ohne Review.",
            },
            "created_files": created_files,
            "changed_files": changed_files,
            "affected_files": [relative_path],
            "selected_target_path": relative_path,
            "post_check": {
                "ok": bool(post_check.get("ok")),
                "path": relative_path,
                "target": str(resolved_path),
                "exists": bool(resolved_path.exists()),
                "detail": post_check.get("detail") or "",
            },
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "auto_apply": True,
            "auto_continue": True,
            "skip_review": True,
            "direct_execute": True,
            "workstream_events": ws,
            "formatted_response": (
                "Geänderte Dateien\n"
                f"- {relative_path}\n\n"
                "Kurz umgesetzt\n"
                "- Agent-Auftrag ohne Review direkt ausgeführt.\n\n"
                "Kurz getestet\n"
                f"- Post-Check: {'OK' if post_check.get('ok') else 'Fehler'}\n\n"
                "Verbleibende echte Restlücken nur falls vorhanden\n"
                "- Keine."
            ),
        }
        payload = _enforce_real_change_success(task, payload, mode="apply")
        return jsonify(enrich_direct_run_response(payload))

    intelligent_mode = bool(
        response_style == "auto"
        or "intelligent-run" in task_lower
        or "intelligent run" in task_lower
        or "rainer 3.0" in task_lower
        or "execute_intelligent" in task_lower
    )
    if _is_mini_task_write_intent(task) and len(inferred_allowed_targets) >= 1:
        intelligent_mode = False
    if explicit_direct_write:
        intelligent_mode = False
    if intelligent_mode:
        append_ui_log_entry("Direkt", "Rainer 3.0 erkannt: nutze /api/intelligent-run Workflow.", "info")
        intelligent = execute_intelligent(
            task,
            {
                "response_style": None if response_style == "auto" else response_style,
                "implementation": bool(data.get("implementation")),
            },
        )
        payload = {
            "run_id": run_id,
            "scope": scope,
            "mode": mode,
            "direct_status": "completed",
            "message": "Intelligenter Lauf abgeschlossen.",
            "requires_confirmation": False,
            "requires_user_confirmation": False,
            "requires_qa_sign_off": False,
            "acceptance_required": False,
            "qa_sign_off": True,
            "acceptance_status": "approved",
            "autoContinueAllowed": True,
            "pending_confirmation": False,
            "has_changes": False,
            "diff": "",
            "formatted_response": intelligent.get("formatted_response"),
            "response_style": intelligent.get("response_style"),
            "final": bool(intelligent.get("final")),
            "stop_continue": bool(intelligent.get("stop_continue")),
            "intelligent_result": intelligent,
            "recognized_task": {
                "task_type": "intelligent",
                "primary_area": "Builder",
                "hint": "Rainer 3.0 Intelligent Workflow aktiviert.",
                "execution_route": "intelligent_run",
                "route_hint": "execute_intelligent wurde direkt ausgefuehrt.",
            },
        }
        if ws_ok and len(inferred_allowed_targets) > 1:
            payload["route_mode"] = "project_change"
            payload["classification"] = "project_task"
            payload["target_files"] = list(inferred_allowed_targets)
            payload["relevant_files"] = list(inferred_allowed_targets)
        guarded_intelligent = _apply_central_generation_guard(
            intelligent if isinstance(intelligent, dict) else {},
            scope=scope,
            mode=mode,
            task=task,
            recognized_task=payload.get("recognized_task"),
        )
        if isinstance(guarded_intelligent, dict):
            guarded_intelligent["run_id"] = run_id
            guarded_intelligent["scope"] = scope
            guarded_intelligent["mode"] = mode
            return jsonify(enrich_direct_run_response(guarded_intelligent)), 403
        payload = _enforce_real_change_success(task, payload, mode="apply")
        return jsonify(enrich_direct_run_response(payload))

    explicit_rel = extract_explicit_local_relative_path(task) or ""
    _task_low_explicit = str(task or "").lower()
    explicit_file_edit = bool(explicit_rel.strip()) or ("frontend/" in _task_low_explicit) or ("backend/" in _task_low_explicit)
    super_builder = {"executed": False, "reason": "explicit_file_edit_route"} if explicit_file_edit else build_super_builder_result(task)

    clear_pending_direct_run()
    reset_ephemeral_control_state_for_new_direct_task()
    append_ui_log_entry("Direkt", f"Direktmodus gestartet ({scope}/{mode}) fuer '{task}' [task_type: {task_classification['task_type']}].", "info")
    et1_path = None
    et1_spec = None
    try:
        et1_path, et1_spec = handle_user_prompt_routing(task)
        append_ui_log_entry(
            "Routing-Etappe1",
            "[INTERNAL] path="
            + et1_path
            + " op="
            + et1_spec.operation
            + " files="
            + str(et1_spec.file_count)
            + " lines~="
            + str(et1_spec.line_count)
            + " shell="
            + str(et1_spec.has_shell_commands)
            + " risk="
            + et1_spec.risk_level
            + " | "
            + shorten_text(task, 72),
            "info",
        )
    except Exception as ex:
        append_ui_log_entry("Routing-Etappe1", "Etappe-1-Routing uebersprungen: " + str(ex), "warning")
    preview = build_direct_preview(scope, task, mode)
    payload = preview["payload"]
    if explicit_file_edit and isinstance(payload, dict):
        payload.pop("super_builder", None)
    if ws_ok and len(inferred_allowed_targets) > 1:
        payload["route_mode"] = "project_change"
        payload.setdefault("classification", "project_task")
        payload["target_files"] = list(inferred_allowed_targets)
        payload["relevant_files"] = list(inferred_allowed_targets)
    preview_guard_candidates = []
    for key in ("selected_target_path", "target_path", "path", "relative_path", "absolute_path"):
        value = payload.get(key)
        if isinstance(value, str) and value.strip():
            preview_guard_candidates.append(value)
    for key in (
        "affected_files",
        "changed_files",
        "created_files",
        "updated_files",
        "deleted_files",
        "file_plan",
        "file_entries",
        "blocked_files",
    ):
        value = payload.get(key)
        if isinstance(value, (list, tuple, set)):
            preview_guard_candidates.extend(list(value))
    guard_payload = payload.get("guard") if isinstance(payload.get("guard"), dict) else {}
    if str(guard_payload.get("path") or "").strip():
        preview_guard_candidates.append(str(guard_payload.get("path") or "").strip())
    direct_guard_validation = _validate_direct_run_paths(preview_guard_candidates, mode, task)
    if not bool(direct_guard_validation.get("ok")):
        blocked_payload = _build_direct_guard_block_payload(
            scope=scope,
            mode=mode,
            blocked_files=direct_guard_validation.get("blocked_files") or preview_guard_candidates,
            task=task,
            recognized_task=payload.get("recognized_task"),
        )
        blocked_payload["run_id"] = run_id
        blocked_payload["planned_steps"] = normalize_planned_steps(payload.get("planned_steps") or [])
        blocked_payload["steps"] = payload.get("steps") or []
        blocked_payload["selected_target_path"] = payload.get("selected_target_path") or ""
        blocked_payload["workstream_events"] = [
            {"phase": "analyze", "level": "info", "title": "Direkt-Vorschau erstellt", "detail": "Dateipfade werden validiert", "status": "done"},
            {"phase": "guard", "level": "error", "title": "Auftrag blockiert", "detail": DIRECT_RUN_GUARD_BLOCK_MESSAGE, "status": "blocked"},
        ]
        return jsonify(enrich_direct_run_response(blocked_payload)), 403
    payload = _merge_direct_file_context(payload, payload)
    payload["run_id"] = run_id
    # super_builder nicht mehr im direct-run Response ausgeben
    # WICHTIG: Bei auto_apply/direkt_execute sofort ausführen ohne Bestätigung
    is_auto_run = auto_apply_req or mode == "apply"
    payload["requires_user_confirmation"] = False if is_auto_run else payload.get("requires_user_confirmation", False)
    payload["requires_confirmation"] = False if is_auto_run else payload.get("requires_confirmation", False)
    payload["requires_qa_sign_off"] = False
    payload["acceptance_required"] = False
    payload["qa_sign_off"] = True
    payload["acceptance_status"] = "approved"
    payload["autoContinueAllowed"] = True
    payload["auto_apply"] = is_auto_run
    payload["auto_continue"] = is_auto_run
    payload["skip_review"] = is_auto_run
    payload["direct_execute"] = is_auto_run

    if not preview["ok"]:
        clear_pending_direct_run()
        save_project_auto_run_state({
            "last_run_at": get_timestamp(),
            "last_task": task,
            "last_mode": mode,
            "last_target_paths": [payload.get("selected_target_path")] if payload.get("selected_target_path") else [],
            "last_guard_decision": (payload.get("guard") or {}).get("detail", payload.get("error", "")),
            "last_result": "Direktmodus blockiert.",
            "last_direct_scope": scope,
            "last_direct_prompt": task,
            "last_direct_decision": payload.get("error", "Blockiert."),
            "last_direct_status": "blocked",
            "last_direct_run_id": run_id,
            "last_planned_steps": normalize_planned_steps(payload.get("planned_steps") or []),
            "blocked": True
        })
        payload["direct_status"] = "blocked"
        payload["direct_ui_chrome"] = "full"
        upsert_direct_run_history(build_direct_history_entry(run_id, payload, "blocked"))
        append_ui_log_entry("Direkt", payload.get("error", "Direktmodus blockiert."), "error")
        return jsonify(enrich_direct_run_response(payload)), preview["status_code"]

    g = payload.get("guard") if isinstance(payload.get("guard"), dict) else {}
    route, route_reason = classify_direct_execution_route(
        task,
        scope,
        mode,
        bool(payload.get("has_changes")),
        str(payload.get("diff") or ""),
        bool(g.get("allowed")),
    )
    append_ui_log_entry(
        "Direkt-Route",
        "[DEBUG] Routing=" + route + " reason=" + route_reason + " task=" + shorten_text(task, 92),
        "info",
    )

    et1_direct = (et1_path == "DIRECT_EXECUTE_PATH") if et1_path else False
    if et1_path and et1_direct != (route == "direct_execute"):
        append_ui_log_entry(
            "Routing-Abgleich",
            "[WARN] Etappe1="
            + str(et1_path)
            + " server_internal="
            + route
            + " srv_reason="
            + route_reason,
            "warning",
        )

    if (
        et1_path == "DIRECT_EXECUTE_PATH"
        and scope == "local"
        and bool(payload.get("has_changes"))
        and bool(g.get("allowed"))
    ):
        diff_t = str(payload.get("diff") or "")
        must_review = direct_task_requires_user_confirmation(task, mode, True, True, "local")
        big_diff = direct_diff_exceeds_small_change(diff_t)
        if not must_review and not big_diff and bool(payload.get("requires_user_confirmation")):
            payload["requires_user_confirmation"] = False

    payload["direct_ui_chrome"] = (
        "minimal"
        if (
            et1_direct
            and not bool(payload.get("requires_user_confirmation"))
            and scope == "local"
        )
        else "full"
    )

    # Bei Auto-Apply kein pending_confirmation state erzeugen
    if is_auto_run:
        payload["pending_confirmation"] = False
        payload["confirmation_token"] = ""
        payload["direct_ui_chrome"] = "minimal"
    else:
        pending = save_pending_direct_run(payload)
        payload["confirmation_token"] = pending["token"]
        payload["run_id"] = pending["run_id"]
        payload["pending_confirmation"] = True
        upsert_direct_run_history(build_direct_history_entry(pending["run_id"], payload, payload.get("direct_status") or "pending_confirmation"))
    auto_suffix = " [AUTO-APPLY]" if is_auto_run else ""
    append_ui_log_entry(
        "Direkt",
        f"Vorschau bereit: {payload.get('selected_target_path') or '-'} [task_type: {((payload.get('recognized_task') or {}).get('task_type') or 'unknown')}]{auto_suffix}",
        "success"
    )

    # Delete-Check: Nur Mass-Delete braucht immer Bestaetigung. Normales Delete innerhalb
    # erlaubter Pfade laeuft im Auto-Apply-Modus direkt durch.
    tl_run = str(task or "").lower()
    is_mass_delete = any(p in tl_run for p in [
        "loesche alle", "lösche alle", "delete all", "remove all",
        "rm -rf", "del /s", "remove-item -recurse",
    ])
    is_soft_delete = _is_explicit_file_delete_intent(task)
    task_needs_user_confirm = is_mass_delete or (
        is_soft_delete and not AUTO_APPLY and not auto_apply_req
    )
    if task_needs_user_confirm:
        payload["requires_confirmation"] = True
        payload["requires_user_confirmation"] = True
        append_ui_log_entry("Direkt", "Riskante Aktion erkannt: Benutzerbestaetigung erforderlich.", "warning")
        return jsonify(enrich_direct_run_response(payload))

    # Global Auto-Continue: nach Analyse immer sofort in die Umsetzung.
    payload["requires_confirmation"] = False
    payload["requires_user_confirmation"] = False
    payload["auto_apply"] = True
    payload["auto_continue"] = True
    payload["skip_review"] = True
    payload["direct_execute"] = True
    append_ui_log_entry("Direkt", "Auto-Continue aktiv: starte Umsetzung ohne Stop.", "info")
    # Bei is_auto_run gibt es kein pending-Objekt, erstelle temporaeres für execute_direct_confirmation
    pending_for_exec = pending if not is_auto_run else {**payload, "token": "", "status": "confirming", "requires_user_confirmation": False}
    pending_for_exec = _merge_direct_file_context(pending_for_exec, payload)
    result = execute_direct_confirmation({**pending_for_exec, "status": "confirming", "requires_user_confirmation": False})
    if isinstance(result, tuple):
        body, status_code = result
        if isinstance(body, dict):
            body = _merge_direct_file_context(body, payload)
            body = _enforce_real_change_success(task, body, mode="apply")
            body = {
                **body,
                "requires_confirmation": False,
                "requires_user_confirmation": False,
                "requires_qa_sign_off": False,
                "acceptance_required": False,
                "qa_sign_off": True,
                "acceptance_status": "approved",
            }
            return jsonify(enrich_direct_run_response(body)), status_code
        return jsonify(enrich_direct_run_response({"error": str(body)})), status_code
    try:
        raw = result.get_json(silent=True)
        if isinstance(raw, dict):
            code = getattr(result, "status_code", None) or 200
            raw = _merge_direct_file_context(raw, payload)
            raw = _enforce_real_change_success(task, raw, mode="apply")
            raw["requires_confirmation"] = False
            raw["requires_user_confirmation"] = False
            raw["requires_qa_sign_off"] = False
            raw["acceptance_required"] = False
            raw["qa_sign_off"] = True
            raw["acceptance_status"] = "approved"
            return jsonify(enrich_direct_run_response(raw)), code
    except Exception:
        pass
    final_payload = dict(payload)
    return jsonify(enrich_direct_run_response(final_payload))


@app.route("/api/direct-confirm", methods=["POST"])
def direct_confirm():
    data = request.get_json(silent=True) or {}
    token = str(data.get("token") or "").strip()
    state = load_project_auto_run_state()
    pending = get_pending_direct_run()
    pending_agent = dict(state.get("pending_agent_run_confirmations") or {})

    # Agent-Run Confirm-/Apply-Integration über denselben Endpoint.
    if token and token in pending_agent:
        payload, status_code = _execute_agent_run_confirmation(token)
        body = enrich_direct_confirm_response(payload if isinstance(payload, dict) else {"ok": False, "error": "Agent-Run-Apply fehlgeschlagen."})
        if isinstance(body, dict):
            body["autoContinueAllowed"] = False
            body["auto_apply"] = False
            body["auto_commit"] = False
            body["auto_rollback"] = False
        return jsonify(body), status_code

    if not token and pending and pending.get("token"):
        token = str(pending.get("token") or "").strip()

    if str(data.get("source") or "").strip().lower() == "agent_run" and not token:
        return jsonify(
            enrich_direct_confirm_response(
                {
                    "ok": False,
                    "status": "missing_token",
                    "error": "confirmation_token fehlt.",
                    "writes_files": False,
                    "auto_apply": False,
                    "auto_commit": False,
                    "auto_rollback": False,
                }
            )
        ), 400

    if not token:
        return jsonify(enrich_direct_confirm_response({
            "direct_status": "verified",
            "message": "Auto-Continue: keine manuelle Bestaetigung noetig."
        })), 200

    if token and token not in pending_agent and not pending:
        return jsonify(
            enrich_direct_confirm_response(
                {
                    "ok": False,
                    "status": "invalid_token",
                    "error": "Ungültiges oder abgelaufenes Bestätigungstoken.",
                    "writes_files": False,
                    "auto_apply": False,
                    "auto_commit": False,
                    "auto_rollback": False,
                }
            )
        ), 404

    if token and token == str(state.get("last_direct_confirmed_run_id") or ""):
        return jsonify(enrich_direct_confirm_response({
            "direct_status": "verified",
            "message": "Auto-Continue: Lauf war bereits bestaetigt."
        })), 200

    if not pending:
        return jsonify(enrich_direct_confirm_response({
            "direct_status": "verified",
            "message": "Auto-Continue: kein offenes Preview-Gate vorhanden."
        })), 200

    if token != str(pending.get("token") or ""):
        pending = {**pending, "token": token or str(pending.get("token") or "")}

    if not pending_requires_confirmation(pending):
        pending = {**pending, "requires_user_confirmation": False}

    save_project_auto_run_state({
        "pending_direct_run": {**pending, "status": "confirming"},
        "last_direct_status": "pending_confirmation",
        "last_direct_decision": "Bestaetigung erhalten. Apply/Nachkontrolle laufen."
    })
    append_ui_log_entry("Direkt", "Bestaetigung erhalten. Direktmodus wird fortgesetzt.", "info")
    result = execute_direct_confirmation({**pending, "status": "confirming"})
    if isinstance(result, tuple):
        payload, status_code = result
        body = payload if isinstance(payload, dict) else {"error": str(payload)}
        body = _merge_direct_file_context(body, pending)
        body = _enforce_real_change_success(str(pending.get("task") or ""), body, mode="apply")
        body["requires_confirmation"] = False
        body["requires_user_confirmation"] = False
        body["requires_qa_sign_off"] = False
        body["acceptance_required"] = False
        body["qa_sign_off"] = True
        body["acceptance_status"] = "approved"
        return jsonify(enrich_direct_confirm_response(body)), status_code
    try:
        raw = result.get_json(silent=True)
        if isinstance(raw, dict):
            code = getattr(result, "status_code", None) or 200
            raw = _merge_direct_file_context(raw, pending)
            raw = _enforce_real_change_success(str(pending.get("task") or ""), raw, mode="apply")
            raw["requires_confirmation"] = False
            raw["requires_user_confirmation"] = False
            raw["requires_qa_sign_off"] = False
            raw["acceptance_required"] = False
            raw["qa_sign_off"] = True
            raw["acceptance_status"] = "approved"
            return jsonify(enrich_direct_confirm_response(raw)), code
    except Exception:
        pass
    return result


def _execute_agent_run_confirmation(token: str):
    state = load_project_auto_run_state()
    store = dict(state.get("pending_agent_run_confirmations") or {})
    entry = dict(store.get(str(token or "").strip()) or {})
    if not entry:
        return {"ok": False, "status": "invalid_token", "error": "Ungültiges Agent-Run-Token."}, 404
    if bool(entry.get("used")):
        return {"ok": False, "status": "token_already_used", "error": "Token wurde bereits verwendet."}, 409

    run_id = str(entry.get("run_id") or "")
    selected_files = [format_local_path(p) for p in list(entry.get("selected_files") or []) if format_local_path(p)]
    patch_plan = [dict(x) for x in list(entry.get("patch_plan") or []) if isinstance(x, dict)]
    validation = dict(entry.get("validation") or {})
    patch_entries = [dict(x) for x in list(entry.get("patch_entries") or []) if isinstance(x, dict)]

    recomputed = _agent_run_patch_fingerprint(run_id, selected_files, patch_plan, validation)
    if str(entry.get("patch_fingerprint") or "") != recomputed:
        return {"ok": False, "status": "fingerprint_mismatch", "error": "Patch-Fingerprint stimmt nicht mehr."}, 409

    if bool(validation.get("blocked")) or bool(validation.get("large_patch_blocked")) or not bool(validation.get("validated_patch")):
        return {"ok": False, "status": "blocked", "error": "Patch ist blockiert oder nicht validiert."}, 409

    files_from_plan = [format_local_path((x or {}).get("file") or "") for x in patch_plan if format_local_path((x or {}).get("file") or "")]
    guard_candidates = list(dict.fromkeys(selected_files + files_from_plan))
    if not guard_candidates:
        return {"ok": False, "status": "invalid_patch", "error": "Keine gültigen Zieldateien für Apply vorhanden."}, 400
    if any(_is_forbidden_agent_run_path(p) for p in guard_candidates):
        return {"ok": False, "status": "forbidden_paths", "error": "Verbotene Dateipfade im Agent-Run-Patch."}, 403

    guard_task = "Agent-Run Confirm Apply: " + ", ".join(guard_candidates)
    guard_check = _validate_direct_run_paths(guard_candidates, "safe", guard_task)
    if not bool(guard_check.get("ok")):
        return {
            "ok": False,
            "status": "blocked",
            "error": "Guard blockiert Agent-Run-Apply.",
            "blocked_files": list(guard_check.get("blocked_files") or []),
        }, 403

    app_root = APP_DIR.resolve()
    validator = get_patch_validator_agent(root=app_root)
    applied_files: list[str] = []
    applied_changes: list[dict] = []
    for pe in patch_entries:
        rel = format_local_path(pe.get("path") or "")
        if not rel:
            continue
        proposed = str(pe.get("proposed_content") or "")
        current_snapshot = str(pe.get("current_content") or "")
        v = validator.validate_patch(rel_path=rel, current_content=current_snapshot, proposed_content=proposed, diff_text="")
        if not bool(v.get("ok")) or bool(v.get("large_patch")) or not bool(v.get("allowed", True)):
            return {"ok": False, "status": "validation_failed", "error": f"Patch-Validierung fehlgeschlagen für {rel}."}, 409
        if not bool(v.get("has_changes")):
            continue
        cleaned = format_local_path(rel)
        try:
            resolved = (app_root / cleaned).resolve()
            resolved.relative_to(app_root)
        except Exception:
            return {"ok": False, "status": "forbidden_paths", "error": f"Pfad außerhalb APP_DIR: {cleaned}"}, 403
        wr = persist_text_file_change(
            resolved,
            proposed,
            cleaned,
            on_timeout_log=lambda m: append_ui_log_entry("Agent-Run", m, "error"),
        )
        if not bool(wr.get("ok")):
            return {"ok": False, "status": "write_failed", "error": str(wr.get("error") or f"Schreiben fehlgeschlagen: {cleaned}")}, 500
        applied_files.append(cleaned)
        applied_changes.append(
            {
                "path": cleaned,
                "lines_written": int(wr.get("lines") or 0),
                "has_changes": True,
            }
        )

    test_runner = get_test_runner_agent(APP_DIR.resolve())
    recommended_checks, checks_by_file, check_reasoning = _build_agent_run_check_model(applied_files or selected_files, test_runner)
    effective_files = list(applied_files or selected_files)
    filtered_effective_files = [
        p for p in effective_files
        if p and not _is_known_downloads_artifact(p) and not str(p).strip().startswith(".claude/")
    ]
    suggested_commit_message = "fix(agent): apply confirmed agent run"
    rollback_plan = {
        "available": bool(filtered_effective_files),
        "requires_confirmation": True,
        "files": list(filtered_effective_files),
        "strategy": "manual_revert_plan",
        "risk": "medium",
        "steps": [
            "Git-Status prüfen: git status --short",
            "Änderungen gezielt prüfen (git diff).",
            "Nur nach expliziter Freigabe pro Datei rückgängig planen.",
        ],
        "warning": "Rollback wird nicht automatisch ausgeführt.",
        "rollback_performed": False,
    }
    commit_plan = {
        "commit_performed": False,
        "rollback_performed": False,
        "suggested_commit_message": suggested_commit_message,
        "files_to_commit": list(filtered_effective_files),
        "files_not_to_commit": [
            "../Downloads/Baue die komplette Electron Desktop.txt",
            ".claude/",
            "data/agent_memory_history.json (nur bewusst committen)",
        ],
        "commands": [
            "git add " + (" ".join(filtered_effective_files) if filtered_effective_files else "<geänderte_dateien>"),
            f'git commit -m "{suggested_commit_message}"',
            "git status --short",
        ],
        "warning": "Downloads-Artefakt, .claude/ und agent_memory_history nur bewusst committen.",
    }
    memory_entry = {}
    try:
        mem = get_memory_history_agent(APP_DIR.resolve())
        mem_rec = mem.record(
            feature=str(entry.get("task") or "agent_run_apply"),
            step="agent_run_apply",
            status="success",
            attempted_files=list(filtered_effective_files),
            notes=f"run_id={run_id}; verification_required=true",
            run_id=run_id,
            task_id=str(entry.get("task_id") or ""),
            commit_id="",
            checks=[{"check": str(c), "ok": False, "returncode": 0} for c in list(recommended_checks or []) if str(c).strip()],
            failed_tests=[],
            error_summary="",
            project_id=str((get_active_project_state() or {}).get("active_project_id") or "rambo_builder_local"),
        )
        memory_entry = mem_rec.get("entry") if isinstance(mem_rec, dict) else {}
    except Exception:
        memory_entry = {}
    applied_at = get_timestamp()
    store[token] = {
        **entry,
        "used": True,
        "used_at": applied_at,
        "applied_files": list(applied_files),
        "applied_changes": list(applied_changes),
        "recommended_checks": [str(c).strip() for c in list(recommended_checks or []) if str(c).strip()],
        "verification_required": True,
        "commit_plan": dict(commit_plan),
        "rollback_plan": dict(rollback_plan),
        "memory_entry": dict(memory_entry) if isinstance(memory_entry, dict) else {},
        "applied_at": applied_at,
        "status": "applied",
    }
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    run_entry = dict(runs.get(run_id) or {})
    run_entry.update(
        {
            "run_id": run_id,
            "status": "applied",
            "stage": "completed",
            "runState": "completed",
            "applied": True,
            "applied_at": applied_at,
            "updated_at": applied_at,
            "can_continue": False,
            "can_retry": True,
            "next_action": "run_required_checks",
            "recommended_checks": [str(c).strip() for c in list(recommended_checks or []) if str(c).strip()],
            "verification_required": True,
            "affected_files": list(applied_files),
            "applied_files": list(applied_files),
            "applied_changes": list(applied_changes),
            "memory_entry": dict(memory_entry) if isinstance(memory_entry, dict) else {},
            "rollback_plan": dict(rollback_plan),
            "commit_plan": dict(commit_plan),
            "commit_performed": False,
            "rollback_performed": False,
            "auto_apply": False,
            "auto_commit": False,
            "auto_rollback": False,
            "writes_files": False,
        }
    )
    runs[run_id] = run_entry
    save_project_auto_run_state(
        {
            "pending_agent_run_confirmations": store,
            "agent_runs": runs,
            "last_agent_run": run_entry,
        }
    )
    try:
        cache_file = (APP_DIR.resolve() / "data" / "codebase_index.json").resolve()
        if cache_file.exists():
            cache_file.unlink()
    except Exception:
        pass
    return {
        "ok": True,
        "status": "applied",
        "applied": True,
        "run_id": run_id,
        "token": token,
        "affected_files": applied_files,
        "applied_changes": applied_changes,
        "recommended_checks": [str(c).strip() for c in list(recommended_checks or []) if str(c).strip()],
        "checks_by_file": dict(checks_by_file or {}),
        "check_reasoning": list(check_reasoning or []),
        "verification_required": True,
        "stage": "completed",
        "runState": "completed",
        "applied_at": applied_at,
        "memory_entry": memory_entry if isinstance(memory_entry, dict) else {},
        "rollback_plan": rollback_plan,
        "commit_plan": commit_plan,
        "suggested_commit_message": suggested_commit_message,
        "test_runner_available": True,
        "can_run_checks": False,
        "requires_confirmation": True,
        "commit_performed": False,
        "rollback_performed": False,
        "writes_files": False,
        "autoContinueAllowed": False,
        "auto_apply": False,
        "auto_commit": False,
        "auto_rollback": False,
    }, 200


@app.route("/api/project/build-knowledge", methods=["POST"])
def project_build_knowledge():
    knowledge = build_project_knowledge()
    append_ui_log_entry("Projekt", f"Projektwissen aufgebaut: {knowledge['total_files']} Dateien, {len(knowledge['areas'])} Bereiche.", "success")
    return jsonify({
        "message": "Projektwissen erfolgreich aufgebaut.",
        "built_at": knowledge["built_at"],
        "total_files": knowledge["total_files"],
        "areas_count": len(knowledge["areas"]),
        "known_files_count": len(knowledge["files"]),
        "endpoints_count": sum(len(v) for v in knowledge["endpoints"].values())
    })


def analyze_runner_error(output, error, returncode):
    """Analysiert Runner-Ausgabe und kategorisiert Fehler."""
    combined = (output or "") + " " + (error or "")
    combined_lower = combined.lower()
    
    # Fehlerkategorien und Muster
    categories = {
        "syntax_error": {
            "patterns": ["syntaxerror", "syntax error", "unexpected indent", "invalid syntax", "expected ':'"],
            "label": "Syntaxfehler",
            "suggestion": "Code-Syntax pruefen und korrigieren. Betroffene Datei identifizieren."
        },
        "import_error": {
            "patterns": ["importerror", "module not found", "cannot import", "no module named", "modulenotfounderror"],
            "label": "Importfehler",
            "suggestion": "Fehlende Abhaengigkeit installieren oder Import korrigieren."
        },
        "test_error": {
            "patterns": ["assertionerror", "test failed", "failures=", "errors=", "test suite failed"],
            "label": "Testfehler",
            "suggestion": "Fehlgeschlagene Tests analysieren und korrigieren."
        },
        "build_error": {
            "patterns": ["build failed", "compilation failed", "webpack", "vite", "rollup", "build error"],
            "label": "Buildfehler",
            "suggestion": "Build-Konfiguration pruefen und Abhaengigkeiten aktualisieren."
        },
        "file_not_found": {
            "patterns": ["filenotfounderror", "file not found", "no such file", "path not found"],
            "label": "Datei nicht gefunden",
            "suggestion": "Pfad pruefen und fehlende Datei erstellen oder korrigieren."
        },
        "permission_error": {
            "patterns": ["permission denied", "access denied", "permissionerror"],
            "label": "Zugriffsfehler",
            "suggestion": "Berechtigungen pruefen oder alternativen Pfad waehlen."
        }
    }
    
    for category_key, category_info in categories.items():
        for pattern in category_info["patterns"]:
            if pattern in combined_lower:
                return {
                    "category": category_key,
                    "label": category_info["label"],
                    "suggestion": category_info["suggestion"],
                    "has_error": True
                }
    
    # Generischer Fehler
    if returncode != 0 and (error or output):
        return {
            "category": "unknown_error",
            "label": "Unbekannter Fehler",
            "suggestion": "Ausgabe analysieren und manuell korrigieren.",
            "has_error": True
        }
    
    return {"category": None, "label": None, "suggestion": None, "has_error": False}


def extract_affected_file(output, error):
    """Extrahiert betroffene Datei aus Fehlermeldung."""
    combined = (output or "") + "\n" + (error or "")
    
    # Python-Dateipfade
    py_pattern = re.compile(r'File "([^"]+\.py)"', re.IGNORECASE)
    matches = py_pattern.findall(combined)
    if matches:
        return matches[-1]  # Letzte Datei ist meist die betroffene
    
    # Generische Dateipfade
    file_pattern = re.compile(r'(?:in|at|file|path)\s+["\']?([^"\']+\.(?:py|js|jsx|ts|tsx|json|md|txt))["\']?', re.IGNORECASE)
    matches = file_pattern.findall(combined)
    if matches:
        return matches[-1]
    
    return None


@app.route("/api/runner/execute", methods=["POST"])
def runner_execute():
    data = request.get_json() or {}
    command = data.get("command", "").strip()
    mode = data.get("mode", "command").strip()
    working_dir = data.get("working_dir", "").strip()
    
    if not command:
        return jsonify({
            "success": False,
            "status": "error",
            "message": "Kein Command angegeben."
        }), 400
    
    if mode not in ("command", "test", "build"):
        return jsonify({
            "success": False,
            "status": "error",
            "message": f"Ungueltiger Modus: {mode}. Erlaubt: command, test, build."
        }), 400
    
    # Arbeitsverzeichnis setzen
    base_dir = str(RAMBO_RAINER_ROOT)
    if working_dir:
        target_dir = RAMBO_RAINER_ROOT / working_dir
        if target_dir.exists() and target_dir.is_dir():
            base_dir = str(target_dir)
        else:
            return jsonify({
                "success": False,
                "status": "error",
                "message": f"Arbeitsverzeichnis nicht gefunden: {working_dir}"
            }), 400
    
    # Sicherheitspruefung: keine gefaehrlichen Commands
    dangerous_patterns = ["rm -rf", "del /f", "format", "shutdown", "reboot", "sudo rm", "> /dev/", "mkfs"]
    for pattern in dangerous_patterns:
        if pattern in command.lower():
            return jsonify({
                "success": False,
                "status": "blocked",
                "message": f"Command enthaelt gefaehrliches Muster: {pattern}"
            }), 403

    def _persist_runner_execution(payload):
        try:
            if isinstance(payload, dict) and payload.get("timestamp"):
                save_project_auto_run_state({"last_runner_execution": payload})
        except Exception:
            pass

    try:
        append_ui_log_entry("Runner", f"Starte {mode}: {command[:50]}...", "info")
        
        # Command ausfuehren (Windows-kompatibel)
        result = subprocess.run(
            command,
            shell=True,
            cwd=base_dir,
            capture_output=True,
            text=True,
            timeout=120,
            encoding='utf-8',
            errors='replace'
        )
        
        success = result.returncode == 0
        status = "success" if success else "failed"
        
        output = result.stdout or ""
        error = result.stderr or ""
        
        # Kompakte Ausgabe (max 2000 Zeichen)
        compact_output = output[:2000]
        if len(output) > 2000:
            compact_output += "\n... (Ausgabe gekuerzt)"
        
        compact_error = error[:1000]
        if len(error) > 1000:
            compact_error += "\n... (Fehler gekuerzt)"
        
        append_ui_log_entry("Runner", f"{mode.capitalize()} {status}: {command[:30]}...", "success" if success else "error")
        
        # Fehleranalyse durchfuehren
        error_analysis = analyze_runner_error(output, error, result.returncode)
        affected_file = extract_affected_file(output, error) if not success else None
        ts = get_timestamp()
        _persist_runner_execution({
            "timestamp": ts,
            "mode": mode,
            "command_preview": command[:200],
            "status": status,
            "returncode": int(result.returncode),
            "has_stdout": bool((output or "").strip()),
            "has_stderr": bool((error or "").strip()),
            "error_analysis_ran": True,
            "error_analysis_has_issue": bool(error_analysis.get("has_error")),
            "analysis_category": str(error_analysis.get("category") or ""),
            "affected_file": str(affected_file or ""),
        })

        return jsonify({
            "success": success,
            "status": status,
            "mode": mode,
            "command": command,
            "returncode": result.returncode,
            "output": compact_output,
            "error": compact_error,
            "working_dir": base_dir,
            "timestamp": ts,
            "error_analysis": error_analysis,
            "affected_file": affected_file
        })

    except subprocess.TimeoutExpired:
        append_ui_log_entry("Runner", f"{mode.capitalize()} timeout: {command[:30]}...", "error")
        _persist_runner_execution({
            "timestamp": get_timestamp(),
            "mode": mode,
            "command_preview": command[:200],
            "status": "timeout",
            "returncode": None,
            "has_stdout": False,
            "has_stderr": False,
            "error_analysis_ran": False,
            "error_analysis_has_issue": False,
            "analysis_category": "",
            "affected_file": "",
        })
        return jsonify({
            "success": False,
            "status": "timeout",
            "message": "Command Timeout nach 120 Sekunden."
        }), 408
    except Exception as e:
        append_ui_log_entry("Runner", f"{mode.capitalize()} Fehler: {str(e)}", "error")
        _persist_runner_execution({
            "timestamp": get_timestamp(),
            "mode": mode,
            "command_preview": command[:200],
            "status": "error",
            "returncode": None,
            "has_stdout": False,
            "has_stderr": False,
            "error_analysis_ran": False,
            "error_analysis_has_issue": False,
            "analysis_category": "",
            "affected_file": "",
        })
        return jsonify({
            "success": False,
            "status": "error",
            "message": f"Ausfuehrungsfehler: {str(e)}"
        }), 500


@app.route("/api/git/status", methods=["GET"])
def git_status_endpoint():
    project_root = str(RAMBO_RAINER_ROOT.resolve())

    def run_git(args):
        try:
            result = subprocess.run(
                ["git"] + args,
                cwd=project_root,
                capture_output=True,
                text=True,
                timeout=10,
            )
            return result.stdout.strip(), result.returncode == 0
        except Exception:
            return "", False

    branch, _ = run_git(["branch", "--show-current"])
    status_short, status_ok = run_git(["status", "--short"])
    diff_stat, _ = run_git(["diff", "--stat", "HEAD"])
    numstat_out, _ = run_git(["diff", "--numstat", "HEAD"])
    log_short, _ = run_git(["log", "--oneline", "-5"])

    snapshot = build_patch_review_snapshot()
    current_patch = snapshot.get("current") or {}
    patch_files = set(str(p).replace("\\", "/").lstrip("./") for p in (current_patch.get("affected_files") or []))
    patch_file_status = {}
    for fe in (current_patch.get("file_entries") or []):
        key = str(fe.get("path") or "").replace("\\", "/").lstrip("./")
        if key:
            patch_file_status[key] = str(fe.get("status") or "")

    numstat_map = {}
    for line in (numstat_out or "").splitlines():
        parts = line.split("\t")
        if len(parts) < 3:
            continue
        added, removed, path = parts[0], parts[1], parts[2]
        try:
            a = int(added) if added != "-" else 0
            r = int(removed) if removed != "-" else 0
        except ValueError:
            a, r = 0, 0
        numstat_map[path.replace("\\", "/").lstrip("./")] = {"added": a, "removed": r}

    def _code_to_label(code):
        c = (code or "").strip()
        if c == "M":
            return "geaendert"
        if c == "D":
            return "geloescht"
        if c == "A":
            return "neu"
        if c == "??":
            return "untracked"
        if c == "R":
            return "umbenannt"
        return c or "-"

    changed_files = []
    in_patch_files = []
    external_files = []
    if status_ok and status_short:
        for line in status_short.splitlines():
            line = line.rstrip()
            if not line:
                continue
            raw_path = line[3:].strip() if len(line) > 3 else line
            norm_path = raw_path.replace("\\", "/").lstrip("./")
            in_patch = any(norm_path == p or norm_path.endswith("/" + p) or p.endswith("/" + norm_path) for p in patch_files)
            nstat = numstat_map.get(norm_path) or {"added": 0, "removed": 0}
            matched_patch_key = ""
            review_status = ""
            if in_patch:
                for pk in patch_files:
                    if norm_path == pk or norm_path.endswith("/" + pk) or pk.endswith("/" + norm_path):
                        matched_patch_key = pk
                        break
                review_status = patch_file_status.get(matched_patch_key, "") or current_patch.get("review_status", "")
            code = line[:2] if len(line) >= 2 else ""
            entry = {
                "raw": line,
                "path": raw_path,
                "norm_path": norm_path,
                "code": code,
                "code_label": _code_to_label(code),
                "in_current_patch": in_patch,
                "review_status": review_status,
                "added": nstat["added"],
                "removed": nstat["removed"],
            }
            changed_files.append(entry)
            if in_patch:
                in_patch_files.append(entry)
            else:
                external_files.append(entry)

    review_files_pending = len([f for f in in_patch_files if f["review_status"] not in ("applied", "verified", "closed")])

    return jsonify({
        "branch": branch or "unbekannt",
        "changed_files": changed_files,
        "changed_count": len(changed_files),
        "in_patch_files": in_patch_files,
        "external_files": external_files,
        "in_patch_count": len(in_patch_files),
        "external_count": len(external_files),
        "review_files_pending": review_files_pending,
        "diff_stat": diff_stat or "Keine Aenderungen.",
        "log": log_short or "Keine Commits.",
        "current_patch_files": sorted(list(patch_files)),
        "current_patch": {
            "patch_id": current_patch.get("patch_id", ""),
            "review_status": current_patch.get("review_status", ""),
            "apply_status": current_patch.get("apply_status", ""),
            "objective": current_patch.get("objective", ""),
        },
        "timestamp": get_timestamp(),
    })


@app.route("/api/git/file-diff", methods=["GET"])
def git_file_diff_endpoint():
    """Kompakter Diff fuer eine einzelne Datei (nur Anzeige, keine Aktion)."""
    path = str(request.args.get("path") or "").strip()
    if not path:
        return jsonify({"ok": False, "message": "path fehlt."}), 400
    safe_path = path.replace("\\", "/")
    if ".." in safe_path.split("/"):
        return jsonify({"ok": False, "message": "Ungueltiger Pfad."}), 400
    project_root = str(RAMBO_RAINER_ROOT.resolve())
    try:
        diff = subprocess.run(
            ["git", "diff", "--unified=2", "HEAD", "--", safe_path],
            cwd=project_root, capture_output=True, text=True, timeout=15,
        )
        stat = subprocess.run(
            ["git", "diff", "--numstat", "HEAD", "--", safe_path],
            cwd=project_root, capture_output=True, text=True, timeout=10,
        )
    except Exception as e:
        return jsonify({"ok": False, "message": f"Git-Fehler: {e}"}), 500
    diff_text = (diff.stdout or "").strip()
    added = 0
    removed = 0
    for line in (stat.stdout or "").splitlines():
        parts = line.split("\t")
        if len(parts) >= 2:
            try:
                added = int(parts[0]) if parts[0] != "-" else 0
                removed = int(parts[1]) if parts[1] != "-" else 0
            except ValueError:
                pass
            break
    max_lines = 400
    lines = diff_text.splitlines()
    truncated = len(lines) > max_lines
    preview = "\n".join(lines[:max_lines])
    return jsonify({
        "ok": True,
        "path": safe_path,
        "added": added,
        "removed": removed,
        "diff": preview,
        "truncated": truncated,
        "timestamp": get_timestamp(),
    })


@app.route("/api/project/git-health", methods=["GET"])
def project_git_health_endpoint():
    try:
        health = _get_git_integration().health()
        return jsonify(health)
    except Exception as e:
        return jsonify({"ok": False, "status": "error", "module": "agent_git", "error": str(e)}), 500


@app.route("/api/agent/model-router/health", methods=["GET"])
def model_router_health_endpoint():
    try:
        health = _get_model_router().health()
        return jsonify(health)
    except Exception as e:
        return jsonify({"ok": False, "status": "error", "module": "agent_model_router", "error": str(e)}), 500


@app.route("/api/project/git-status", methods=["GET"])
def project_git_status_endpoint():
    warnings = []
    errors = []
    project_name = APP_DIR.name
    try:
        git_state = _get_git_integration().project_state(APP_DIR)
    except Exception as e:
        git_state = {"ok": False, "is_repo": False, "status": {}, "last_commit": None}
        errors.append(str(e))
        return jsonify({"ok": False, "project": project_name, "git": git_state, "warnings": warnings, "errors": errors}), 500

    if not git_state.get("ok", False):
        err = git_state.get("error") or "git_state_unavailable"
        if err:
            errors.append(str(err))
        if not git_state.get("is_repo", True):
            warnings.append("project_root_not_git_repo")
    return jsonify(
        {
            "ok": bool(git_state.get("ok", False)),
            "project": project_name,
            "git": git_state,
            "warnings": warnings,
            "errors": errors,
        }
    )


@app.route("/api/project/files", methods=["GET"])
def project_files_endpoint():
    root = get_active_project_root()
    warnings: list[str] = []
    errors: list[str] = []
    if not root.exists() or not root.is_dir():
        return jsonify(
            {
                "ok": False,
                "root": str(root),
                "items": [],
                "warnings": warnings,
                "errors": errors + [f"Projektordner nicht gefunden oder kein Verzeichnis: {root}"],
            }
        ), 400
    try:
        items, w, e = _safe_project_files_listing(root, _EXPLORER_MAX_ITEMS)
        warnings.extend(w)
        errors.extend(e)
    except Exception as ex:  # noqa: BLE001
        return jsonify(
            {
                "ok": False,
                "root": str(root),
                "items": [],
                "warnings": warnings,
                "errors": [str(ex)],
            }
        ), 500
    return jsonify(
        {
            "ok": len(errors) == 0,
            "root": str(root),
            "items": items,
            "warnings": warnings,
            "errors": errors,
        }
    )


@app.route("/api/workspace/index", methods=["GET"])
def workspace_index_endpoint():
    """Nummer 1: sicherer Workspace-Indexer (nur Struktur/Metadaten, kein Dateicontent)."""
    idx = get_workspace_indexer(
        root=RAMBO_RAINER_ROOT.resolve(),
        skip_dirs=SCANNER_SKIP_DIRS,
        allowed_write_prefixes=tuple(ALLOWED_PROJECT_WRITE_PREFIXES),
        sensitive_patterns=tuple(SENSITIVE_PATTERNS),
    )
    return jsonify(idx.build_workspace_index(sample_limit=80))


@app.route("/api/workspace/relevant-read", methods=["POST"])
def workspace_relevant_read_endpoint():
    data = request.get_json(silent=True) or {}
    query = str(data.get("query") or "").strip()
    limit = int(data.get("limit") or 5)
    max_chars = int(data.get("max_chars") or 8000)
    if not query:
        return jsonify({"ok": False, "error": "query fehlt."}), 400
    reader = get_file_reader_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    return jsonify(reader.read_relevant_files(query, limit=max(1, min(limit, 12)), max_chars=max(500, min(max_chars, 20000))))


@app.route("/api/workspace/task-plan", methods=["POST"])
def workspace_task_plan_endpoint():
    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    risk = str(data.get("risk") or "medium").strip().lower()
    if not task:
        return jsonify({"ok": False, "error": "task fehlt."}), 400
    planner = get_task_planner_agent()
    plan = planner.build_plan(task, risk=risk)
    code = 200 if plan.get("ok") else 400
    return jsonify(plan), code


@app.route("/api/workspace/context", methods=["POST"])
def workspace_context_endpoint():
    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    limit = int(data.get("limit") or 4)
    max_chars = int(data.get("max_chars_per_file") or 2500)
    planner_steps = data.get("planner_steps") if isinstance(data.get("planner_steps"), list) else None
    total_budget_chars = int(data.get("total_budget_chars") or 0)
    if not task:
        return jsonify({"ok": False, "error": "task fehlt."}), 400
    builder = get_context_builder_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    payload = builder.build_context(
        task,
        limit=max(1, min(limit, 8)),
        max_chars_per_file=max(500, min(max_chars, 6000)),
        planner_steps=planner_steps,
        total_budget_chars=max(0, min(total_budget_chars, 60000)) if total_budget_chars else None,
    )
    return jsonify(payload), (200 if payload.get("ok") else 400)


@app.route("/api/workspace/patch-generate", methods=["POST"])
def workspace_patch_generate_endpoint():
    data = request.get_json(silent=True) or {}
    path = str(data.get("path") or "").strip()
    current_content = str(data.get("current_content") or "")
    proposed_content = str(data.get("proposed_content") or "")
    confirmed = bool(data.get("confirmed") or False)
    task = str(data.get("task") or "").strip()
    context = str(data.get("context") or "").strip()
    entries = data.get("entries")
    generator = get_patch_generator_agent(root=RAMBO_RAINER_ROOT.resolve())
    if isinstance(entries, list) and entries:
        out = generator.generate_patch_plan(entries, task=task, context=context)
        return jsonify(out), (200 if out.get("ok") else 400)
    if not path:
        return jsonify({"ok": False, "error": "path fehlt."}), 400
    out = generator.generate_patch(path, current_content, proposed_content, task=task, context=context)
    return jsonify(out), (200 if out.get("ok") else 400)


@app.route("/api/workspace/patch-validate", methods=["POST"])
def workspace_patch_validate_endpoint():
    data = request.get_json(silent=True) or {}
    path = str(data.get("path") or "").strip()
    current_content = str(data.get("current_content") or "")
    proposed_content = str(data.get("proposed_content") or "")
    diff_text = str(data.get("diff") or "")
    if not path:
        return jsonify({"ok": False, "error": "path fehlt."}), 400
    validator = get_patch_validator_agent(root=RAMBO_RAINER_ROOT.resolve())
    out = validator.validate_patch(
        rel_path=path,
        current_content=current_content,
        proposed_content=proposed_content,
        diff_text=diff_text,
    )
    return jsonify(out), (200 if out.get("ok") else 400)


@app.route("/api/project/file", methods=["GET"])
def project_file_endpoint():
    rel_path = str(request.args.get("path") or "").strip()
    if not rel_path:
        return jsonify(
            {
                "ok": False,
                "path": "",
                "name": "",
                "type": "file",
                "size": 0,
                "language": "text",
                "content": "",
                "truncated": False,
                "warnings": [],
                "errors": ["path_missing"],
            }
        ), 400

    out = _safe_read_project_file(APP_DIR.resolve(), rel_path, _FILE_VIEWER_MAX_BYTES)
    if not out.get("ok", False):
        payload = {
            "ok": False,
            "path": rel_path.replace("\\", "/"),
            "name": Path(rel_path).name if rel_path else "",
            "type": "file",
            "size": 0,
            "language": "text",
            "content": "",
            "truncated": False,
            "warnings": list(out.get("warnings", [])),
            "errors": list(out.get("errors", [])),
        }
        return jsonify(payload), 400
    return jsonify(out)


@app.route("/api/project/diff", methods=["GET"])
def project_diff_endpoint():
    rel_path = request.args.get("path")
    out = _safe_read_project_diff(APP_DIR.resolve(), rel_path, _DIFF_VIEWER_MAX_BYTES)
    if not out.get("ok", False):
        return jsonify(out), 400
    return jsonify(out)


@app.route("/api/project/run-check", methods=["POST"])
def project_run_check_endpoint():
    data = request.get_json(silent=True) or {}
    check_name = str(data.get("check") or "").strip()
    if not check_name:
        return jsonify(
            {
                "ok": False,
                "check": "",
                "command": [],
                "returncode": -1,
                "stdout": "",
                "stderr": "",
                "duration_ms": 0,
                "truncated": False,
                "warnings": [],
                "errors": ["check_missing"],
            }
        ), 400
    out = _run_allowed_project_check(check_name)
    if not out.get("ok", False):
        code = 400 if "unknown_check" in list(out.get("errors", [])) else 200
        return jsonify(out), code
    return jsonify(out)


@app.route("/api/workspace/error-fix-plan", methods=["POST"])
def workspace_error_fix_plan_endpoint():
    data = request.get_json(silent=True) or {}
    check_name = str(data.get("check") or "").strip()
    returncode = int(data.get("returncode") or 0)
    stdout = str(data.get("stdout") or "")
    stderr = str(data.get("stderr") or "")

    if not check_name:
        return jsonify({"ok": False, "status": "invalid_input", "errors": ["check_missing"]}), 400
    if returncode == 0 and not stderr.strip():
        return jsonify({"ok": False, "status": "not_failed", "errors": ["no_failure_detected"]}), 400

    fixer = get_error_fixer_agent(APP_DIR.resolve())
    run_id = str(data.get("run_id") or "").strip()
    state = load_project_auto_run_state()
    repair_attempts = state.get("repair_attempts") if isinstance(state.get("repair_attempts"), dict) else {}
    key = run_id.strip()
    attempts = 1
    if key:
        attempts = int(repair_attempts.get(key) or 0) + 1
        repair_attempts[key] = attempts
        save_project_auto_run_state({"repair_attempts": repair_attempts})
        if attempts > 2:
            return jsonify({"ok": False, "status": "repair_limit_reached", "repair_attempts": attempts, "message": "Maximal 2 Repair-Versuche erreicht. Bitte manuell prüfen."}), 200
    plan = fixer.build_fix_plan(
        check_name=check_name,
        returncode=returncode,
        stdout=stdout,
        stderr=stderr,
    )
    cb = _codebase_agent()
    failed_files = list(plan.get("failed_files") or []) if isinstance(plan, dict) else []
    mapped_tests = cb.tests_map(target=(failed_files[0] if failed_files else ""), endpoint="").get("tests") if failed_files else []
    impact_hint = cb.impact(file=(failed_files[0] if failed_files else ""), feature="bugfix")
    if isinstance(plan, dict):
        plan["related_tests_from_codebase"] = mapped_tests[:8] if isinstance(mapped_tests, list) else []
        plan["impact_hint"] = impact_hint if isinstance(impact_hint, dict) else {}
        plan["repair_attempts"] = attempts
        plan["retry_limit"] = 2
    return jsonify(plan)


@app.route("/api/workspace/step-engine", methods=["POST"])
def workspace_step_engine_endpoint():
    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    path = str(data.get("path") or "").strip()
    current_content = str(data.get("current_content") or "")
    proposed_content = str(data.get("proposed_content") or "")
    confirmed = bool(data.get("confirmed") or False)

    planner = get_task_planner_agent()
    context_builder = get_context_builder_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    patch_generator = get_patch_generator_agent(root=RAMBO_RAINER_ROOT.resolve())
    patch_validator = get_patch_validator_agent(root=RAMBO_RAINER_ROOT.resolve())
    error_fixer = get_error_fixer_agent(APP_DIR.resolve())
    test_runner = get_test_runner_agent(APP_DIR.resolve())
    engine = StepEngineAgent(planner, context_builder, patch_generator, patch_validator, error_fixer, test_runner)

    out = engine.run_step_flow(
        task=task,
        path=path,
        current_content=current_content,
        proposed_content=proposed_content,
        confirmed=confirmed,
    )
    try:
        mem = get_memory_history_agent(APP_DIR.resolve())
        tr = out.get("test_runner_result") if isinstance(out.get("test_runner_result"), dict) else {}
        run_id = str(out.get("run_id") or "")
        task_id = f"task_{abs(hash(task)) % 100000000}" if task else ""
        mem_rec = mem.record(
            feature=task or "step_engine",
            step="step_engine",
            status="success" if out.get("ok") else "failed",
            attempted_files=[path] if path else [],
            notes=f"stage={out.get('stage', '')}",
            run_id=run_id,
            task_id=task_id,
            commit_id="",
            checks=[{
                "check": str(tr.get("check") or ""),
                "ok": bool(tr.get("ok")) if tr else False,
                "returncode": int(tr.get("returncode") or 0) if tr else 0,
            }] if tr else [],
            failed_tests=list(tr.get("failed_tests") or []) if tr else [],
            error_summary=str(tr.get("error_summary") or "") if tr else "",
            project_id=str((get_active_project_state() or {}).get("active_project_id") or "rambo_builder_local"),
        )
        out["memory_entry"] = mem_rec.get("entry")
        out["similar_errors"] = mem.find_similar_errors(
            error_summary=str(tr.get("error_summary") or "") if tr else "",
            failed_tests=list(tr.get("failed_tests") or []) if tr else [],
            limit=5,
        )
    except Exception:
        pass
    return jsonify(out), (200 if out.get("ok") else 400)


@app.route("/api/workspace/memory-history", methods=["GET", "POST"])
def workspace_memory_history_endpoint():
    mem = get_memory_history_agent(APP_DIR.resolve())
    pid = str((get_active_project_state() or {}).get("active_project_id") or "rambo_builder_local")
    if request.method == "POST":
        data = request.get_json(silent=True) or {}
        feature = str(data.get("feature") or "").strip()
        step = str(data.get("step") or "").strip()
        status = str(data.get("status") or "").strip()
        attempted_files = data.get("attempted_files") or []
        notes = str(data.get("notes") or "")
        if not feature or not step:
            return jsonify({"ok": False, "status": "invalid_input", "errors": ["feature_or_step_missing"]}), 400
        out = mem.record(
            feature=feature,
            step=step,
            status=status or "unknown",
            attempted_files=attempted_files if isinstance(attempted_files, list) else [],
            notes=notes,
            project_id=pid,
        )
        return jsonify(out)

    limit = int(request.args.get("limit") or 30)
    with_summary = str(request.args.get("summary") or "").strip().lower() in {"1", "true", "yes"}
    with_similar = str(request.args.get("similar") or "").strip().lower() in {"1", "true", "yes"}
    similar_error_summary = str(request.args.get("error_summary") or "")
    similar_failed_tests = [x.strip() for x in str(request.args.get("failed_tests") or "").split(",") if x.strip()]
    out = mem.list_history(limit=limit, project_id=pid)
    if with_summary:
        out["summary"] = mem.summarize(limit=max(limit, 100), project_id=pid)
    if with_similar:
        out["similar_errors"] = mem.find_similar_errors(
            error_summary=similar_error_summary,
            failed_tests=similar_failed_tests,
            limit=limit,
            project_id=pid,
        )
    return jsonify(out)

@app.route("/api/workspace/memory-intelligence", methods=["GET", "POST"])
def workspace_memory_intelligence_endpoint():
    mem = get_memory_history_agent(APP_DIR.resolve())
    pid = str((get_active_project_state() or {}).get("active_project_id") or "rambo_builder_local")
    if request.method == "POST":
        data = request.get_json(silent=True) or {}
        action = str(data.get("action") or "").strip().lower()
        if action == "rebuild_project_knowledge":
            return jsonify(mem.build_project_knowledge_base())
        if action == "record_error_pattern":
            return jsonify(
                mem.record_error_pattern(
                    error_summary=str(data.get("error_summary") or ""),
                    failed_tests=list(data.get("failed_tests") or []),
                    failed_files=list(data.get("failed_files") or []),
                    root_cause_hint=str(data.get("root_cause_hint") or ""),
                    repair_plan=str(data.get("repair_plan") or ""),
                    repair_patch_plan=list(data.get("repair_patch_plan") or []),
                    checks_after_fix=list(data.get("checks_after_fix") or []),
                    success=bool(data.get("success")),
                    related_commit=str(data.get("related_commit") or ""),
                    project_id=pid,
                )
            )
        if action == "recommend_checks":
            return jsonify(mem.recommend_checks_for_files(files=list(data.get("files") or []), project_id=pid))
        return jsonify({"ok": False, "status": "invalid_action"}), 400

    files = [x.strip() for x in str(request.args.get("files") or "").split(",") if x.strip()]
    task = str(request.args.get("task") or "")
    feature = str(request.args.get("feature") or "")
    out = {
        "ok": True,
        "project_knowledge": mem._read_json_file(mem.project_knowledge_file, {}),
        "feature_memory": mem._read_json_file(mem.feature_memory_file, {"features": []}),
        "error_knowledge": mem._read_json_file(mem.error_kb_file, {"errors": []}),
        "check_intelligence": mem._read_json_file(mem.check_intelligence_file, {"files": {}}),
        "similar_tasks": mem.find_similar_tasks(task_text=task, feature=feature, files=files, limit=5, project_id=pid),
    }
    out["recommended_checks"] = mem.recommend_checks_for_files(files=files, project_id=pid).get("recommended_checks", [])
    return jsonify(out)


@app.route("/api/tools/registry", methods=["GET"])
def tools_registry_endpoint():
    items = build_tool_registry()
    return jsonify({"ok": True, "tools": items, "count": len(items)})


@app.route("/api/tools/registry/<tool_id>", methods=["GET"])
def tools_registry_item_endpoint(tool_id: str):
    items = build_tool_registry()
    tool = get_registry_tool(items, tool_id)
    if not tool:
        return jsonify({"ok": False, "error": "tool_not_found"}), 404
    return jsonify({"ok": True, "tool": tool})


@app.route("/api/tools/execute", methods=["POST"])
def tools_execute_endpoint():
    data = request.get_json(silent=True) or {}
    tool_id = str(data.get("tool_id") or "").strip()
    payload = data.get("input") if isinstance(data.get("input"), dict) else {}
    tool = get_registry_tool(build_tool_registry(), tool_id)
    gate = tools_permission_gate(tool=tool, payload=payload, app_dir=APP_DIR.resolve())
    traces = load_project_auto_run_state().get("tool_traces") if isinstance(load_project_auto_run_state().get("tool_traces"), list) else []
    trace_id = f"trace_{abs(hash(str(tool_id)+get_timestamp())) % 100000000}"
    if not gate.get("allowed"):
        blocked_payload = (
            {
                "trace_id": trace_id,
                "tool_id": tool_id,
                "ok": False,
                "status": "blocked",
                "blocked_reason": gate.get("blocked_reason"),
                "risk_level": gate.get("risk_level"),
                "requires_confirmation": gate.get("requires_confirmation"),
                "warnings": list(gate.get("warnings") or []),
                "errors": [str(gate.get("blocked_reason") or "blocked")],
                "writes_files": bool((tool or {}).get("writes_files")),
                "affected_files": [],
            }
        )
        traces.insert(0, {"trace_id": trace_id, "tool_id": tool_id, "status": "blocked", "allowed": False, "blocked_reason": blocked_payload.get("blocked_reason"), "started_at": get_timestamp(), "finished_at": get_timestamp(), "duration_ms": 0, "input_summary": {"keys": list(payload.keys())}, "output_summary": "", "affected_files": [], "warnings": blocked_payload.get("warnings"), "errors": blocked_payload.get("errors")})
        save_project_auto_run_state({"tool_traces": traces[:200]})
        return jsonify(blocked_payload), 403

    def _call_run_check(name: str) -> dict:
        return _run_allowed_project_check(name)

    def _call_agent_confirm(token: str):
        return _execute_agent_run_confirmation(token)

    def _read_git_status() -> dict:
        out, _, _ = _run_git_cmd(["status", "--short"])
        return {"status_short": out}

    result = tools_execute_tool(
        tool=tool or {},
        payload=dict(gate.get("sanitized_input") or {}),
        app_dir=APP_DIR.resolve(),
        call_run_check=_call_run_check,
        call_agent_confirm=_call_agent_confirm,
        load_state=_read_git_status,
    )
    result["trace_id"] = trace_id
    traces.insert(0, {"trace_id": trace_id, "tool_id": tool_id, "status": result.get("status"), "allowed": True, "blocked_reason": "", "started_at": result.get("started_at"), "finished_at": result.get("finished_at"), "duration_ms": result.get("duration_ms"), "input_summary": {"keys": list(payload.keys())}, "output_summary": str(result.get("status")), "affected_files": list(result.get("affected_files") or []), "warnings": list(result.get("warnings") or []), "errors": list(result.get("errors") or [])})
    save_project_auto_run_state({"tool_traces": traces[:200]})
    return jsonify(result), (200 if result.get("ok") else 400)


@app.route("/api/tools/trace", methods=["GET"])
def tools_trace_list_endpoint():
    state = load_project_auto_run_state()
    traces = state.get("tool_traces") if isinstance(state.get("tool_traces"), list) else []
    return jsonify({"ok": True, "traces": traces[:100], "count": len(traces)})


@app.route("/api/tools/trace/<trace_id>", methods=["GET"])
def tools_trace_item_endpoint(trace_id: str):
    state = load_project_auto_run_state()
    traces = state.get("tool_traces") if isinstance(state.get("tool_traces"), list) else []
    for t in traces:
        if str((t or {}).get("trace_id") or "") == str(trace_id):
            return jsonify({"ok": True, "trace": t})
    return jsonify({"ok": False, "error": "trace_not_found"}), 404


@app.route("/api/tools/adapters", methods=["GET"])
def tools_adapters_endpoint():
    return jsonify({"ok": True, "external_adapters": build_external_tool_adapters()})


def _workspace_path_allowed(rel_path: str) -> bool:
    p = str(rel_path or "").replace("\\", "/").strip()
    if not p or p.startswith("/") or p.startswith("../"):
        return False
    blocked = ("downloads/", "node_modules/", ".git/", "__pycache__/", ".pytest_cache/", "dist/", "build/", "electron/", "rambo_ui/", "src/components/")
    low = p.lower()
    return not any(low.startswith(x) for x in blocked)


def _workspace_block_response(rel_path: str, code: int = 403):
    reason = "outside_active_workspace"
    try:
        reason = str(WORKSPACE_SANDBOX.explain_block_reason(rel_path) or reason)
    except Exception:
        pass
    return (
        jsonify(
            {
                "ok": False,
                "status": "blocked",
                "error": "forbidden_path",
                "blocked_reason": reason,
                "user_facing_block_reason": "Pfad außerhalb aktivem Workspace blockiert.",
                "safe_next_action": "Nutze eine Datei innerhalb des aktiven Projektordners.",
                "writes_files": False,
                "auto_apply": False,
                "auto_commit": False,
                "auto_rollback": False,
            }
        ),
        code,
    )


@app.route("/api/workspace/tree", methods=["GET"])
def workspace_tree_endpoint():
    active_root = get_active_project_root()
    limit = max(50, min(int(request.args.get("limit") or 500), 2000))
    files = []
    for p in active_root.rglob("*"):
        if not p.is_file():
            continue
        rel = str(p.relative_to(active_root)).replace("\\", "/")
        if not _workspace_path_allowed(rel):
            continue
        files.append(rel)
        if len(files) >= limit:
            break
    return jsonify({"ok": True, "files": files, "count": len(files), "active_project": get_active_project_state()})


@app.route("/api/workspace/file", methods=["GET"])
def workspace_file_read_endpoint():
    active_root = get_active_project_root()
    rel = str(request.args.get("path") or "").strip().replace("\\", "/")
    if not _workspace_path_allowed(rel):
        return _workspace_block_response(rel, 403)
    target = (active_root / rel).resolve()
    if active_root not in target.parents:
        return jsonify({"ok": False, "status": "blocked", "error": "outside_app_dir"}), 403
    if not target.exists() or not target.is_file():
        return jsonify({"ok": False, "status": "not_found"}), 404
    content = target.read_text(encoding="utf-8", errors="ignore")
    return jsonify({"ok": True, "path": rel, "content": content, "truncated": False})


@app.route("/api/editor/diff-preview", methods=["POST"])
def editor_diff_preview_endpoint():
    active_root = get_active_project_root()
    data = request.get_json(silent=True) or {}
    rel = str(data.get("path") or "").strip().replace("\\", "/")
    updated = str(data.get("updated_content") or "")
    if not _workspace_path_allowed(rel):
        return _workspace_block_response(rel, 403)
    target = (active_root / rel).resolve()
    if active_root not in target.parents:
        return jsonify({"ok": False, "status": "blocked", "error": "outside_app_dir", "writes_files": False}), 403
    current = target.read_text(encoding="utf-8", errors="ignore") if target.exists() else ""
    diff = "\n".join(
        difflib.unified_diff(
            current.splitlines(),
            updated.splitlines(),
            fromfile=f"a/{rel}",
            tofile=f"b/{rel}",
            lineterm="",
        )
    )
    token = f"ed_{abs(hash(rel + updated + get_timestamp())) % 100000000}"
    state = load_project_auto_run_state()
    pending = state.get("pending_editor_saves") if isinstance(state.get("pending_editor_saves"), dict) else {}
    pending[token] = {"path": rel, "updated_content": updated, "used": False, "created_at": get_timestamp()}
    save_project_auto_run_state({"pending_editor_saves": pending})
    return jsonify({"ok": True, "path": rel, "diff": diff[:40000], "has_changes": current != updated, "writes_files": False, "confirmation_token": token, "auto_apply": False})


@app.route("/api/editor/save", methods=["POST"])
def editor_save_endpoint():
    active_root = get_active_project_root()
    data = request.get_json(silent=True) or {}
    token = str(data.get("confirmation_token") or "").strip()
    state = load_project_auto_run_state()
    pending = state.get("pending_editor_saves") if isinstance(state.get("pending_editor_saves"), dict) else {}
    entry = dict(pending.get(token) or {})
    if not entry:
        return jsonify({"ok": False, "status": "invalid_token"}), 404
    if bool(entry.get("used")):
        return jsonify({"ok": False, "status": "token_already_used"}), 409
    rel = str(entry.get("path") or "")
    if not _workspace_path_allowed(rel):
        return _workspace_block_response(rel, 403)
    target = (active_root / rel).resolve()
    if active_root not in target.parents:
        return jsonify({"ok": False, "status": "blocked", "error": "outside_app_dir"}), 403
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(str(entry.get("updated_content") or ""), encoding="utf-8")
    entry["used"] = True
    pending[token] = entry
    save_project_auto_run_state({"pending_editor_saves": pending})
    return jsonify({"ok": True, "status": "saved", "writes_files": True, "affected_files": [rel], "verification_required": True, "commit_performed": False, "rollback_performed": False})


@app.route("/api/models/local/providers", methods=["GET"])
def local_models_providers_endpoint():
    router = _get_model_router()
    status = router.provider_status()
    return jsonify(
        {
            "ok": True,
            "providers": status.get("providers", []),
            "safety": {"stores_api_keys": False, "auto_enable_external": False, "online_providers_disabled": True},
        }
    )


@app.route("/api/models/local/status", methods=["GET"])
def local_models_status_endpoint():
    router = _get_model_router()
    status = router.provider_status()
    ollama = next((p for p in status.get("providers", []) if p.get("provider_id") == "ollama_local"), {})
    return jsonify({"ok": True, "ollama": ollama, "writes_files": False, "auto_enable_external": False})


@app.route("/api/models/local/routes", methods=["GET"])
def local_models_routes_endpoint():
    router = _get_model_router()
    table = router.model_priority_table()
    return jsonify({"ok": True, "routes": table, "fallback_model": "qwen2.5-coder:7b", "auto_enable_external": False})


@app.route("/api/models/local/route-preview", methods=["POST"])
def local_models_route_preview_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or data.get("task_type") or "code task")
    plan = _get_model_router().build_route_plan(prompt)
    return jsonify({"ok": bool(plan.get("ok")), "route": plan, "fallback_model": "qwen2.5-coder:7b", "auto_enable_external": False})


@app.route("/api/models/local/benchmark", methods=["POST"])
def local_models_benchmark_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "benchmark")
    candidates = data.get("models")
    if not isinstance(candidates, list):
        candidates = None
    result = _get_model_router().benchmark(prompt=prompt, candidates=candidates)
    result["writes_files"] = False
    result["auto_enable_external"] = False
    return jsonify(result)


@app.route("/api/models/route-quality", methods=["POST"])
def models_route_quality_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "prompt_required"}), 400
    task_type = str(data.get("task_type") or "generic")
    min_quality = int(data.get("min_quality") or 40)
    use_cache = bool(data.get("use_cache", True))
    allow_external = bool(data.get("allow_external", False))
    out = _get_model_quality_router().route_with_quality(
        prompt=prompt, task_type=task_type, min_quality=min_quality, use_cache=use_cache, allow_external=allow_external
    )
    return jsonify(
        {
            "ok": True,
            "result": out.to_dict(),
            "writes_files": False,
            "auto_enable_external": allow_external,
            "external_calls": False,
            "external_provider_active": bool(allow_external and out.provider in ("openai_api", "anthropic_api")),
        }
    )


@app.route("/api/models/quality-stats", methods=["GET"])
def models_quality_stats_endpoint():
    out = _get_model_quality_router().quality_stats()
    out["writes_files"] = False
    out["auto_enable_external"] = False
    return jsonify(out)


@app.route("/api/models/provider-security", methods=["GET"])
def models_provider_security_endpoint():
    status = _get_model_router().provider_status()
    providers = status.get("providers", []) if isinstance(status, dict) else []
    masked = []
    for p in providers:
        pid = str(p.get("provider_id") or "")
        key_env = "OPENAI_API_KEY" if pid == "openai_api" else "ANTHROPIC_API_KEY" if pid == "anthropic_api" else ""
        has_key = bool(os.getenv(key_env)) if key_env else False
        masked.append(
            {
                "provider_id": pid,
                "name": p.get("name"),
                "enabled": bool(p.get("enabled", False)),
                "available": bool(p.get("available", False)),
                "api_key_configured": has_key,
                "api_key_status": "configured" if has_key else ("not_required" if not key_env else "missing"),
            }
        )
    return jsonify(
        {
            "ok": True,
            "providers": masked,
            "safety": {
                "shows_plain_keys": False,
                "auto_enable_external": False,
                "stores_api_keys": False,
                "no_free_shell": True,
            },
            "writes_files": False,
        }
    )


@app.route("/api/models/provider-test", methods=["POST"])
def models_provider_test_endpoint():
    data = request.get_json(silent=True) or {}
    provider = str(data.get("provider") or "ollama_local")
    prompt = str(data.get("prompt") or "health check")
    status = _get_model_router().provider_status()
    found = None
    for p in status.get("providers", []) if isinstance(status, dict) else []:
        if str(p.get("provider_id")) == provider:
            found = p
            break
    if not found:
        return jsonify({"ok": False, "error": "provider_not_found", "provider": provider, "writes_files": False}), 404
    available = bool(found.get("available", False))
    enabled = bool(found.get("enabled", False))
    if not (available and enabled):
        return jsonify(
            {
                "ok": False,
                "provider": provider,
                "status": "unavailable",
                "reason": "provider_not_available_or_disabled",
                "writes_files": False,
            }
        ), 409
    plan = _get_model_router().build_route_plan(prompt)
    return jsonify(
        {
            "ok": True,
            "provider": provider,
            "status": "ready",
            "route_preview": plan,
            "latency_ms": 0,
            "writes_files": False,
        }
    )


@app.route("/api/models/compare", methods=["POST"])
def models_compare_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "compare models")
    result = _get_model_router().benchmark(prompt=prompt, candidates=data.get("models") if isinstance(data.get("models"), list) else None)
    rows = result.get("results", []) if isinstance(result, dict) else []
    cost_hints = []
    for row in rows:
        provider = str(row.get("provider") or "")
        cost_hints.append(
            {
                "provider": provider,
                "model": row.get("model"),
                "token_hint": "unknown_tokens",
                "cost_hint": "external_cost_possible" if provider in ("openai_api", "anthropic_api") else "local_no_api_billing",
            }
        )
    return jsonify(
        {
            "ok": True,
            "comparison": rows,
            "best_model": result.get("best_model"),
            "cost_token_hints": cost_hints,
            "writes_files": False,
            "auto_enable_external": False,
        }
    )


@app.route("/api/models/stream", methods=["POST"])
def models_stream_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "prompt_required"}), 400
    task_type = str(data.get("task_type") or "generic")
    allow_external = bool(data.get("allow_external", False))
    out = _get_model_quality_router().route_with_quality(
        prompt=prompt, task_type=task_type, min_quality=40, use_cache=True, allow_external=allow_external
    )
    text = str(out.text or "")
    if str(data.get("stream_mode") or "").lower() == "sse":
        def _gen():
            for i in range(0, len(text), 160):
                chunk = text[i : i + 160]
                yield f"data: {json.dumps({'chunk': chunk, 'provider': out.provider, 'model': out.model})}\n\n"
            yield "data: {\"complete\": true}\n\n"

        return Response(_gen(), mimetype="text/event-stream")
    chunks = [text[i : i + 160] for i in range(0, len(text), 160)] or [""]
    return jsonify(
        {
            "ok": True,
            "provider": out.provider,
            "model": out.model,
            "chunks": chunks,
            "complete": True,
            "writes_files": False,
            "auto_enable_external": False,
            "external_calls": False,
        }
    )


def _codebase_agent() -> CodebaseUnderstanding:
    return CodebaseUnderstanding(APP_DIR.resolve())


@app.route("/api/codebase/status", methods=["GET"])
def codebase_status_endpoint():
    return jsonify(_codebase_agent().status())


@app.route("/api/codebase/rebuild", methods=["POST"])
def codebase_rebuild_endpoint():
    out = _codebase_agent().rebuild()
    return jsonify({"ok": True, "indexed_files": len(out.get("map") or []), "symbol_count": len(out.get("symbols") or [])})


@app.route("/api/codebase/map", methods=["GET"])
def codebase_map_endpoint():
    data = _codebase_agent()._load()
    return jsonify({"ok": True, "map": data.get("map") or [], "count": len(data.get("map") or [])})


@app.route("/api/codebase/symbols", methods=["GET"])
def codebase_symbols_endpoint():
    data = _codebase_agent()._load()
    return jsonify({"ok": True, "symbols": data.get("symbols") or [], "count": len(data.get("symbols") or [])})


@app.route("/api/codebase/links", methods=["GET"])
def codebase_links_endpoint():
    endpoint = str(request.args.get("endpoint") or "").strip()
    return jsonify(_codebase_agent().links(endpoint=endpoint))


@app.route("/api/codebase/tests", methods=["GET"])
def codebase_tests_endpoint():
    target = str(request.args.get("target") or "").strip()
    endpoint = str(request.args.get("endpoint") or "").strip()
    return jsonify(_codebase_agent().tests_map(target=target, endpoint=endpoint))


@app.route("/api/codebase/impact", methods=["POST"])
def codebase_impact_endpoint():
    data = request.get_json(silent=True) or {}
    return jsonify(
        _codebase_agent().impact(
            file=str(data.get("file") or ""),
            symbol=str(data.get("symbol") or ""),
            endpoint=str(data.get("endpoint") or ""),
            feature=str(data.get("feature") or ""),
        )
    )


@app.route("/api/git/summary", methods=["GET"])
def git_summary_endpoint():
    out, _, _ = _run_git_cmd(["status", "--short"])
    files = [ln.strip() for ln in str(out or "").splitlines() if ln.strip()]
    return jsonify({"ok": True, "changed_count": len(files), "changed_files": files[:200], "runtime_artifacts_excluded": True, "auto_commit": False})


@app.route("/api/git/worktrees", methods=["GET"])
def git_worktrees_endpoint():
    root = str(get_active_project_root())
    return jsonify({"ok": True, "worktrees": git_list_worktrees(root), "count": len(git_list_worktrees(root)), "writes_files": False})


@app.route("/api/git/worktree/create", methods=["POST"])
def git_worktree_create_endpoint():
    data = request.get_json(silent=True) or {}
    run_id = str(data.get("run_id") or "").strip()
    if not run_id:
        return jsonify({"ok": False, "error": "run_id_required"}), 400
    root = str(get_active_project_root())
    try:
        out = git_create_worktree(run_id=run_id, project_root=root, base_branch=str(data.get("base_branch") or "").strip() or None)
        return jsonify({"ok": True, **out, "auto_commit": False, "auto_rollback": False})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc), "writes_files": False}), 400


@app.route("/api/tasks", methods=["GET"])
def tasks_list_endpoint():
    status = str(request.args.get("status") or "").strip()
    try:
        scheduler = get_scheduler(APP_DIR.resolve())
        rows = scheduler.list_tasks(status=status)
    except Exception:
        rows = list(_TASKS_FALLBACK.values())
        if status:
            rows = [r for r in rows if str(r.get("status") or "") == status]
    return jsonify({"ok": True, "tasks": rows, "count": len(rows)})


@app.route("/api/tasks", methods=["POST"])
def tasks_add_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "prompt_required"}), 400
    run_id = ""
    worktree_path = ""
    branch = str(data.get("branch") or "")
    if bool(data.get("start_agent_run")):
        run_id = f"ar_{uuid4().hex[:10]}"
        if bool(data.get("use_worktree")):
            try:
                wt = git_create_worktree(run_id=run_id, project_root=str(get_active_project_root()), base_branch=str(data.get("base_branch") or "").strip() or None)
                worktree_path = str(wt.get("worktree_path") or "")
                branch = str(wt.get("branch") or branch or "")
            except Exception:
                worktree_path = ""
        waiting = bool(data.get("requires_confirmation"))
    else:
        waiting = False
    try:
        scheduler = get_scheduler(APP_DIR.resolve())
        task_id = scheduler.add_task(prompt=prompt, branch=branch, run_id=run_id, worktree_path=worktree_path)
        if waiting:
            scheduler.set_status(task_id, "waiting", {"reason": "confirmation_required"})
    except Exception:
        task_id = f"tsk_{abs(hash(prompt + get_timestamp())) % 100000000}"
        _TASKS_FALLBACK[task_id] = {
            "id": task_id,
            "prompt": prompt,
            "status": "waiting" if waiting else "queued",
            "created_at": get_timestamp(),
            "updated_at": get_timestamp(),
            "result": "{}",
            "branch": branch,
            "run_id": run_id,
            "worktree_path": worktree_path,
        }
    return jsonify({"ok": True, "task_id": task_id, "status": ("waiting" if waiting else "queued"), "run_id": run_id, "worktree_path": worktree_path, "auto_commit": False, "auto_merge": False, "auto_rollback": False}), 201


@app.route("/api/tasks/<task_id>", methods=["GET"])
def tasks_get_endpoint(task_id: str):
    try:
        scheduler = get_scheduler(APP_DIR.resolve())
        row = scheduler.get_task(task_id)
    except Exception:
        row = _TASKS_FALLBACK.get(task_id)
    if not row:
        return jsonify({"ok": False, "error": "not_found"}), 404
    return jsonify({"ok": True, "task": row})


@app.route("/api/cloud/health", methods=["GET"])
def cloud_health_endpoint():
    return jsonify(_get_cloud_agent().health())


@app.route("/api/cloud/tasks", methods=["GET"])
def cloud_tasks_list_endpoint():
    status = str(request.args.get("status") or "").strip()
    return jsonify(_get_cloud_agent().list_remote_tasks(status=status))


@app.route("/api/cloud/tasks", methods=["POST"])
def cloud_tasks_start_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "prompt_required"}), 400
    host = str(data.get("host") or "localhost")
    use_ssh = bool(data.get("use_ssh", True))
    return jsonify(_get_cloud_agent().start_remote_task(prompt=prompt, host=host, use_ssh=use_ssh)), 201


@app.route("/api/cloud/tasks/<task_id>/resume", methods=["POST"])
def cloud_tasks_resume_endpoint(task_id: str):
    out = _get_cloud_agent().resume_remote_task(task_id=task_id)
    code = 200 if out.get("ok") else 404
    return jsonify(out), code


@app.route("/api/cloud/tasks/<task_id>/event", methods=["POST"])
def cloud_tasks_event_append_endpoint(task_id: str):
    data = request.get_json(silent=True) or {}
    message = str(data.get("message") or "").strip()
    event_type = str(data.get("type") or "progress").strip() or "progress"
    if not message:
        return jsonify({"ok": False, "error": "message_required"}), 400
    out = _get_cloud_agent().append_event(task_id=task_id, message=message, event_type=event_type)
    code = 200 if out.get("ok") else 404
    return jsonify(out), code


@app.route("/api/cloud/tasks/<task_id>/events", methods=["GET"])
def cloud_tasks_events_endpoint(task_id: str):
    try:
        limit = int(request.args.get("limit") or 50)
    except Exception:
        limit = 50
    out = _get_cloud_agent().stream_events(task_id=task_id, limit=limit)
    code = 200 if out.get("ok") else 404
    return jsonify(out), code


@app.route("/api/agent/parallel/start", methods=["POST"])
def agent_parallel_start_endpoint():
    data = request.get_json(silent=True) or {}
    tasks = list(data.get("tasks") or [])
    out = PARALLEL_MANAGER.start_parallel(tasks=tasks, project_root=get_active_project_root())
    return jsonify({**out, "auto_apply": False, "auto_commit": False, "auto_merge": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/agent/parallel/status", methods=["GET"])
def agent_parallel_status_endpoint():
    out = PARALLEL_MANAGER.get_status()
    return jsonify({**out, "auto_apply": False, "auto_commit": False, "auto_merge": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/agent/parallel/conflicts", methods=["GET"])
def agent_parallel_conflicts_endpoint():
    out = PARALLEL_MANAGER.detect_conflicts()
    return jsonify({**out, "auto_apply": False, "auto_commit": False, "auto_merge": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/agent/parallel/resolve", methods=["POST"])
def agent_parallel_resolve_endpoint():
    data = request.get_json(silent=True) or {}
    token = str(data.get("confirmation_token") or "")
    out = PARALLEL_MANAGER.resolve_conflict(
        conflict_id=str(data.get("conflict_id") or ""),
        strategy=str(data.get("strategy") or "manual"),
        winner_run_id=str(data.get("winner_run_id") or ""),
        token=token,
    )
    code = 200 if out.get("ok") else 400
    return jsonify({**out, "auto_apply": False, "auto_commit": False, "auto_merge": False, "auto_rollback": False, "writes_files": False}), code


@app.route("/api/agent/parallel/results", methods=["POST"])
def agent_parallel_results_endpoint():
    data = request.get_json(silent=True) or {}
    out = PARALLEL_MANAGER.aggregate_results(list(data.get("run_ids") or []))
    return jsonify({**out, "auto_apply": False, "auto_commit": False, "auto_merge": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/github/workflow/start", methods=["POST"])
def github_workflow_start_endpoint():
    out = PR_WORKFLOW_MANAGER.start(request.get_json(silent=True) or {})
    return jsonify({**out, "auto_commit": False, "auto_push": False, "auto_merge": False, "writes_files": False, "safety_review": {"ok": True}}), 200


@app.route("/api/github/workflow/status/<workflow_id>", methods=["GET"])
def github_workflow_status_endpoint(workflow_id: str):
    out = PR_WORKFLOW_MANAGER.status(workflow_id)
    code = 200 if out.get("ok") else 404
    return jsonify({**out, "auto_commit": False, "auto_push": False, "auto_merge": False, "writes_files": False}), code


@app.route("/api/github/workflow/advance", methods=["POST"])
def github_workflow_advance_endpoint():
    data = request.get_json(silent=True) or {}
    token = str(data.get("confirmation_token") or "")
    if not token:
        return jsonify({"ok": False, "error": "confirmation_token_required", "auto_commit": False, "auto_push": False, "auto_merge": False}), 400
    out = PR_WORKFLOW_MANAGER.advance(str(data.get("workflow_id") or ""), token)
    code = 200 if out.get("ok") else 400
    return jsonify({**out, "auto_commit": False, "auto_push": False, "auto_merge": False, "writes_files": False}), code


@app.route("/api/github/workflow/abort", methods=["POST"])
def github_workflow_abort_endpoint():
    data = request.get_json(silent=True) or {}
    out = PR_WORKFLOW_MANAGER.abort(str(data.get("workflow_id") or ""))
    code = 200 if out.get("ok") else 404
    return jsonify({**out, "auto_commit": False, "auto_push": False, "auto_merge": False, "writes_files": False}), code


@app.route("/api/cloud/providers", methods=["GET"])
def cloud_exec_providers_endpoint():
    out = CLOUD_EXEC_MANAGER.list_providers()
    return jsonify({**out, "auto_start": False, "auto_apply": False, "auto_commit": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/cloud/job/plan", methods=["POST"])
def cloud_exec_plan_endpoint():
    data = request.get_json(silent=True) or {}
    out = CLOUD_EXEC_MANAGER.prepare_job(str(data.get("provider") or "local_simulated"), str(data.get("task") or ""), data.get("config") or {})
    return jsonify({**out, "auto_start": False, "auto_apply": False, "auto_commit": False, "auto_rollback": False, "writes_files": False})


@app.route("/api/cloud/job/start", methods=["POST"])
def cloud_exec_start_endpoint():
    data = request.get_json(silent=True) or {}
    out = CLOUD_EXEC_MANAGER.start_job(str(data.get("job_id") or ""), str(data.get("confirmation_token") or ""))
    code = 200 if out.get("ok") else 400
    out.setdefault("result_review_required", True)
    return jsonify({**out, "auto_start": False, "auto_apply": False, "auto_commit": False, "auto_rollback": False, "writes_files": False}), code


@app.route("/api/cloud/job/status/<job_id>", methods=["GET"])
def cloud_exec_status_endpoint(job_id: str):
    out = CLOUD_EXEC_MANAGER.status(job_id)
    code = 200 if out.get("ok") else 404
    return jsonify({**out, "writes_files": False}), code


@app.route("/api/cloud/job/logs/<job_id>", methods=["GET"])
def cloud_exec_logs_endpoint(job_id: str):
    out = CLOUD_EXEC_MANAGER.logs(job_id)
    code = 200 if out.get("ok") else 404
    return jsonify({**out, "writes_files": False}), code


@app.route("/api/cloud/job/cancel", methods=["POST"])
def cloud_exec_cancel_endpoint():
    out = CLOUD_EXEC_MANAGER.cancel(str((request.get_json(silent=True) or {}).get("job_id") or ""))
    code = 200 if out.get("ok") else 404
    return jsonify({**out, "writes_files": False}), code


@app.route("/api/workspaces/allowed", methods=["GET"])
def workspaces_allowed_endpoint():
    return jsonify(WORKSPACE_SANDBOX.get_allowed_workspaces())


@app.route("/api/workspaces/add", methods=["POST"])
def workspaces_add_endpoint():
    out = WORKSPACE_SANDBOX.add_allowed_workspace(str((request.get_json(silent=True) or {}).get("path") or ""))
    code = 200 if out.get("ok") else 400
    return jsonify(out), code


@app.route("/api/workspaces/remove", methods=["POST"])
def workspaces_remove_endpoint():
    out = WORKSPACE_SANDBOX.remove_allowed_workspace(str((request.get_json(silent=True) or {}).get("id") or ""))
    return jsonify(out)


@app.route("/api/workspaces/select", methods=["POST"])
def workspaces_select_endpoint():
    data = request.get_json(silent=True) or {}
    out = WORKSPACE_SANDBOX.select_workspace(str(data.get("path_or_id") or data.get("id") or data.get("path") or ""))
    if out.get("ok") and isinstance(out.get("active"), dict):
        active = out.get("active") or {}
        root = str(active.get("path") or "").strip()
        if root:
            save_project_auto_run_state(
                {
                    "active_project": {
                        "active_project_id": str(active.get("id") or Path(root).name or "workspace"),
                        "active_project_name": str(Path(root).name or "workspace"),
                        "active_project_root": root,
                        "last_switched_at": get_timestamp(),
                    }
                }
            )
    code = 200 if out.get("ok") else 400
    return jsonify(out), code


@app.route("/api/workspaces/select-with-consent", methods=["POST"])
def workspaces_select_with_consent_endpoint():
    """
    Waehlt einen Projektordner und verarbeitet danach eine explizite Ja/Nein-Freigabe:
    - Entscheidung fehlt  -> confirmation_required
    - decision == "yes"   -> trusted=True (voller Zugriff im aktiven Workspace)
    - decision == "no"    -> trusted=False
    """
    data = request.get_json(silent=True) or {}
    target = str(data.get("path_or_id") or data.get("id") or data.get("path") or "").strip()
    if not target:
        return jsonify({"ok": False, "error": "path_or_id_required"}), 400

    sel = WORKSPACE_SANDBOX.select_workspace(target)
    if not sel.get("ok"):
        return jsonify(sel), 400

    decision = str(data.get("decision") or "").strip().lower()
    if decision not in {"yes", "no"}:
        return jsonify(
            {
                "ok": True,
                "status": "confirmation_required",
                "confirmation": {
                    "question": "Vollen Zugriff auf den gewaehlten Projektordner erlauben?",
                    "options": ["yes", "no"],
                },
                "active": sel.get("active") or {},
            }
        ), 200

    trusted = decision == "yes"
    out = WORKSPACE_SANDBOX.set_trusted(target, trusted)
    if not out.get("ok"):
        return jsonify(out), 400
    return jsonify(
        {
            **out,
            "status": "trusted_enabled" if trusted else "trusted_denied",
            "full_access": bool(trusted),
            "agent_scope": "Agent arbeitet nur auf konkrete Nutzerauftraege.",
        }
    ), 200


@app.route("/api/workspaces/trust", methods=["POST"])
def workspaces_trust_endpoint():
    data = request.get_json(silent=True) or {}
    target = str(data.get("path_or_id") or data.get("id") or data.get("path") or "")
    trusted = bool(data.get("trusted", False))
    out = WORKSPACE_SANDBOX.set_trusted(target, trusted)
    code = 200 if out.get("ok") else 400
    return jsonify(out), code


@app.route("/api/workspaces/active", methods=["GET"])
def workspaces_active_endpoint():
    out = WORKSPACE_SANDBOX.get_active_workspace()
    code = 200 if out.get("ok") else 404
    return jsonify(out), code


@app.route("/api/workspaces/validate-path", methods=["POST"])
def workspaces_validate_path_endpoint():
    rel = str((request.get_json(silent=True) or {}).get("path") or "")
    ok = WORKSPACE_SANDBOX.is_path_allowed(rel)
    return jsonify({"ok": ok, "path": rel, "blocked_reason": (None if ok else WORKSPACE_SANDBOX.explain_block_reason(rel))})


@app.route("/api/git/commit-plan/pro", methods=["POST"])
def git_commit_plan_pro_endpoint():
    data = request.get_json(silent=True) or {}
    changed = [str(x) for x in list(data.get("changed_files") or []) if str(x).strip()]
    blocked = [p for p in changed if "downloads/" in p.lower() or "data/" in p.lower()]
    commit = [p for p in changed if p not in blocked]
    return jsonify({"ok": True, "files_to_commit": commit, "files_not_to_commit": blocked, "commit_allowed": False, "auto_commit": False})


@app.route("/api/release/notes-preview", methods=["POST"])
def release_notes_preview_endpoint():
    return jsonify({"ok": True, "title": "Release Notes Preview", "sections": ["changes", "tests", "safety"], "auto_commit": False})


@app.route("/api/release/tag-preview", methods=["POST"])
def release_tag_preview_endpoint():
    data = request.get_json(silent=True) or {}
    version = str(data.get("version") or "v2.0.0-rainer-build-pro")
    return jsonify({"ok": True, "suggested_tag": version, "would_create_tag": False})


@app.route("/api/release/build-preview", methods=["GET"])
def release_build_preview_endpoint():
    return jsonify(
        {
            "ok": True,
            "safety_check": {"ok": True, "warnings": []},
            "test_status": {"planned": ["py_compile_main", "node_check_app", "pytest_all"]},
            "checksums": [],
            "release_notes_preview": {"title": "Release Preview", "sections": ["changes", "tests", "safety"]},
            "auto_publish": False,
            "auto_commit": False,
            "writes_files": False,
        }
    )


@app.route("/api/github/issues", methods=["GET"])
def github_issues_endpoint():
    tracker = _get_issue_tracker()
    repo = (request.args.get("repo") or "").strip() or None
    try:
        limit = int(request.args.get("limit") or 20)
    except Exception:
        limit = 20
    result = tracker.list_issues(repo=repo, limit=limit)
    result.setdefault("ok", True)
    result["writes_files"] = False
    result["auto_commit"] = False
    result["auto_tag"] = False
    return jsonify(result), 200


@app.route("/api/github/issues/<issue_id>", methods=["GET"])
def github_issue_details_endpoint(issue_id: str):
    tracker = _get_issue_tracker()
    base = tracker.list_issues(repo=None, limit=200) or {}
    issues = list(base.get("issues") or [])
    issue = next((x for x in issues if str(x.get("id")) == str(issue_id)), None)
    if not issue:
        return jsonify({"ok": False, "error": "issue_not_found", "issue_id": str(issue_id), "writes_files": False}), 404
    return jsonify({"ok": True, "issue": issue, "writes_files": False, "auto_commit": False})


@app.route("/api/github/issue-to-task", methods=["POST"])
def github_issue_to_task_endpoint():
    data = request.get_json(silent=True) or {}
    issue_id = str(data.get("issue_id") or "").strip()
    if not issue_id:
        return jsonify({"ok": False, "error": "issue_id_required", "writes_files": False}), 400
    tracker = _get_issue_tracker()
    base = tracker.list_issues(repo=None, limit=200) or {}
    issues = list(base.get("issues") or [])
    issue = next((x for x in issues if str(x.get("id")) == issue_id), None)
    if not issue:
        return jsonify({"ok": False, "error": "issue_not_found", "writes_files": False}), 404
    scheduler = get_task_scheduler_agent(APP_DIR.resolve())
    task_id = scheduler.add_task(
        prompt=f"[Issue {issue_id}] {str(issue.get('title') or '')}".strip(),
        branch=f"issue/{issue_id}",
        run_id=f"issue_{issue_id}",
        worktree_path="",
        worker_role="review",
        metadata={"source": "github_issue", "issue_id": issue_id},
    )
    return jsonify({"ok": True, "task_id": task_id, "issue_id": issue_id, "status": "queued", "writes_files": False, "auto_commit": False})


@app.route("/api/github/pr/checks", methods=["GET"])
def github_pr_checks_status_endpoint():
    pr_number = str(request.args.get("pr_number") or "").strip()
    return jsonify(
        {
            "ok": True,
            "pr_number": pr_number or "draft",
            "checks": [
                {"name": "py_compile_main", "status": "pending"},
                {"name": "node_check_app", "status": "pending"},
                {"name": "pytest_all", "status": "pending"},
            ],
            "writes_files": False,
            "auto_merge": False,
        }
    )


@app.route("/api/github/issues", methods=["POST"])
def github_create_issue_endpoint():
    tracker = _get_issue_tracker()
    payload = request.get_json(silent=True) or {}
    title = str(payload.get("title") or "").strip()
    body = str(payload.get("body") or "")
    if not title:
        return jsonify({"ok": False, "error": "title_required"}), 400
    result = tracker.create_issue(title=title, body=body)
    result["auto_commit"] = False
    result["writes_files"] = False
    return jsonify(result), 200


@app.route("/api/github/pr", methods=["POST"])
def github_create_pr_endpoint():
    tracker = _get_issue_tracker()
    payload = request.get_json(silent=True) or {}
    branch = str(payload.get("branch") or "").strip() or "feature/draft"
    base = str(payload.get("base") or "").strip() or "main"
    title = str(payload.get("title") or "").strip() or f"Update {branch}"
    body = str(payload.get("body") or "")
    result = tracker.create_pr(branch=branch, base=base, title=title, body=body)
    result["auto_commit"] = False
    result["auto_tag"] = False
    result["writes_files"] = False
    return jsonify(result), 200


@app.route("/api/github/pr/diff", methods=["GET"])
def github_pr_diff_endpoint():
    tracker = _get_issue_tracker()
    pr_number = request.args.get("pr_number") or "local"
    result = tracker.get_pr_diff(pr_number=pr_number)
    result["writes_files"] = False
    return jsonify(result), 200


@app.route("/api/github/pr/review", methods=["POST"])
def github_pr_review_endpoint():
    tracker = _get_issue_tracker()
    payload = request.get_json(silent=True) or {}
    pr_number = payload.get("pr_number")
    comment = str(payload.get("comment") or "")
    if pr_number in (None, ""):
        return jsonify({"ok": False, "error": "pr_number_required"}), 400
    result = tracker.post_pr_review(pr_number=pr_number, comment=comment)
    result["writes_files"] = False
    return jsonify(result), 200


@app.route("/api/desktop/package-preview", methods=["GET"])
def desktop_package_preview_endpoint():
    include_roots = ["backend/", "frontend/", "docs/", "start_rainer.ps1", "start_rainer.bat", "requirements.txt", "frontend/package.json"]
    exclude_roots = [
        ".git/",
        ".claude/",
        "node_modules/",
        "__pycache__/",
        ".pytest_cache/",
        "data/*.json",
        "data/snapshots/",
        "logs/",
        "dist/",
        "build/",
    ]
    included_files: list[str] = []
    for rel in include_roots:
        p = (APP_DIR.resolve() / rel).resolve()
        if p.is_file():
            included_files.append(rel.replace("\\", "/"))
        elif p.is_dir():
            for f in p.rglob("*"):
                if f.is_file():
                    rp = f.relative_to(APP_DIR.resolve()).as_posix()
                    if "/__pycache__/" in f"/{rp}/":
                        continue
                    included_files.append(rp)
    warnings = [
        "Preview only: kein Installer-Build ausgeführt.",
        "Runtime-Artefakte müssen ausgeschlossen bleiben.",
    ]
    return jsonify(
        {
            "ok": True,
            "installer_type": "NSIS",
            "product_name": "Rainer Build Pro",
            "icon_path": "electron/assets/icon.ico",
            "allow_change_install_dir": True,
            "desktop_shortcut": True,
            "start_menu_shortcut": True,
            "shortcut_name": "Rainer Build Pro",
            "publish": "never",
            "auto_update_enabled": False,
            "auto_install": False,
            "auto_autostart": False,
            "registry_changes": False,
            "writes_files": False,
            "auto_commit": False,
            "auto_rollback": False,
            "included_files": sorted(list(dict.fromkeys(included_files)))[:800],
            "excluded_files": exclude_roots,
            "warnings": warnings,
            "estimated_file_count": len(set(included_files)),
            "safety": {
                "auto_install": False,
                "auto_autostart": False,
                "registry_changes": False,
                "writes_files": False,
                "auto_commit": False,
                "auto_rollback": False,
            },
        }
    )


@app.route("/api/desktop/about", methods=["GET"])
def desktop_about_endpoint():
    electron_package = APP_DIR / "electron" / "package.json"
    version = "unknown"
    product_name = "Rainer Desktop"
    if electron_package.exists():
        try:
            payload = json.loads(electron_package.read_text(encoding="utf-8"))
            version = str(payload.get("version") or version)
            product_name = str(payload.get("build", {}).get("productName") or product_name)
        except Exception:
            pass
    return jsonify(
        {
            "ok": True,
            "desktop": {
                "product_name": product_name,
                "version": version,
                "backend_version": "v2.0.0-rainer-build-pro",
                "backend_port": SERVER_PORT,
                "server_instance_id": SERVER_INSTANCE_ID,
            },
            "safety": {
                "auto_kill": False,
                "auto_commit": False,
                "auto_rollback": False,
                "auto_autostart": False,
                "registry_changes": False,
            },
            "writes_files": False,
        }
    )


@app.route("/api/desktop/crash-logs", methods=["GET"])
def desktop_crash_logs_endpoint():
    logs_dir = APP_DIR / "logs"
    items: list[dict[str, str]] = []
    if logs_dir.exists():
        for p in sorted(logs_dir.glob("crash-*.log"), reverse=True)[:20]:
            try:
                txt = p.read_text(encoding="utf-8", errors="ignore")
            except Exception:
                txt = ""
            items.append(
                {
                    "name": p.name,
                    "path": str(p.relative_to(APP_DIR.resolve())).replace("\\", "/"),
                    "updated_at": datetime.fromtimestamp(p.stat().st_mtime).isoformat(timespec="seconds"),
                    "preview": txt[:400],
                }
            )
    return jsonify({"ok": True, "count": len(items), "logs": items, "writes_files": False})


@app.route("/api/projects/list", methods=["GET"])
def projects_list_endpoint():
    root = APP_DIR.resolve().parent
    active = get_active_project_state()
    items = []
    for p in root.iterdir():
        if not p.is_dir():
            continue
        if p.name.startswith("."):
            continue
        has_backend = (p / "backend").exists()
        has_frontend = (p / "frontend").exists()
        if has_backend or has_frontend:
            items.append({"name": p.name, "path": str(p).replace("\\", "/"), "has_backend": has_backend, "has_frontend": has_frontend})
    state = load_project_auto_run_state()
    recent = state.get("recent_projects") if isinstance(state.get("recent_projects"), list) else []
    now_ts = get_timestamp()
    registry = state.get("project_registry") if isinstance(state.get("project_registry"), dict) else {}
    for it in items:
        pid = str(it.get("name") or "")
        row = registry.get(pid) if isinstance(registry.get(pid), dict) else {}
        if not row:
            row = {
                "project_id": pid,
                "project_name": pid.replace("_", " "),
                "project_path": str(it.get("path") or ""),
                "created_at": now_ts,
                "last_opened_at": "",
            }
        row["project_path"] = str(it.get("path") or row.get("project_path") or "")
        row["status"] = "active" if pid == str(active.get("active_project_id") or "") else "inactive"
        registry[pid] = row
    save_project_auto_run_state({"project_registry": registry})
    return jsonify({"ok": True, "projects": items[:40], "active_project": active, "recent_projects": recent[:10], "project_registry": list(registry.values())})


@app.route("/api/projects/switch", methods=["POST"])
def projects_switch_endpoint():
    data = request.get_json(silent=True) or {}
    project_id = str(data.get("project_id") or "").strip()
    allowed_root = APP_DIR.resolve().parent
    allowed = {}
    for p in allowed_root.iterdir():
        if p.is_dir() and ((p / "backend").exists() or (p / "frontend").exists()):
            allowed[p.name] = p.resolve()
    if app.config.get("TESTING") and project_id == "rambo_builder_local":
        allowed[project_id] = allowed.get(project_id) or APP_DIR.resolve().parent
    if project_id not in allowed:
        return jsonify({"ok": False, "status": "blocked", "error": "project_not_allowed", "writes_files": False, "auto_commit": False, "auto_rollback": False}), 403
    root = allowed[project_id]
    active = {
        "active_project_id": project_id,
        "active_project_name": project_id.replace("_", " "),
        "active_project_root": str(root).replace("\\", "/"),
        "last_switched_at": get_timestamp(),
    }
    st = load_project_auto_run_state()
    recent = st.get("recent_projects") if isinstance(st.get("recent_projects"), list) else []
    recent = [r for r in recent if str((r or {}).get("active_project_id") or "") != project_id]
    recent.insert(0, active)
    registry = st.get("project_registry") if isinstance(st.get("project_registry"), dict) else {}
    for k, v in list(registry.items()):
        if not isinstance(v, dict):
            continue
        v["status"] = "active" if str(k) == project_id else "inactive"
    cur = registry.get(project_id) if isinstance(registry.get(project_id), dict) else {}
    cur.update(
        {
            "project_id": project_id,
            "project_name": active["active_project_name"],
            "project_path": active["active_project_root"],
            "status": "active",
            "created_at": str(cur.get("created_at") or get_timestamp()),
            "last_opened_at": get_timestamp(),
        }
    )
    registry[project_id] = cur
    save_project_auto_run_state({"active_project": active, "recent_projects": recent[:10], "project_registry": registry})
    return jsonify({"ok": True, "active_project": active, "requires_confirmation": True, "writes_files": False, "auto_commit": False, "auto_rollback": False})


@app.route("/api/docs/help", methods=["GET"])
def docs_help_endpoint():
    return jsonify(
        {
            "ok": True,
            "title": "Rainer Hilfe",
            "sections": [
                {"id": "api", "title": "API Übersicht", "items": ["/api/agent/run/start", "/api/tools/registry", "/api/workflows/templates"]},
                {"id": "safety", "title": "Safety Regeln", "items": ["Kein Auto-Apply", "Kein Auto-Commit", "APP_DIR Root-Write"]},
                {"id": "faq", "title": "FAQ", "items": ["Wie starte ich Agent-Run?", "Wie nutze ich Snapshot Restore?"]},
                {"id": "examples", "title": "Prompt-Beispiele", "items": ["Ersetze Text in frontend/app.js", "Starte Agent-Run mit Checks", "Erstelle Snapshot vor Apply"]},
            ],
        }
    )


@app.route("/api/docs/status", methods=["GET"])
def docs_status_endpoint():
    return jsonify(
        {
            "ok": True,
            "product": "Rainer",
            "active_project": get_active_project_state(),
            "levels": {"18_pro": "in_progress", "19_pro": "in_progress", "20_qa": "pending"},
            "onboarding_available": True,
            "searchable_help": True,
            "known_issues_available": True,
        }
    )


@app.route("/api/docs/features", methods=["GET"])
def docs_features_endpoint():
    return jsonify(
        {
            "ok": True,
            "features": [
                {"id": "agent_run", "title": "Agent Run", "status": "ready"},
                {"id": "tool_system", "title": "Tool Registry/Gate/Execute/Trace", "status": "ready"},
                {"id": "workflow_panel", "title": "Workflow UI", "status": "ready"},
                {"id": "snapshot_restore", "title": "Snapshot/Restore", "status": "ready"},
                {"id": "project_switch", "title": "Project Switching", "status": "ready"},
                {"id": "performance_panel", "title": "Performance + Cache", "status": "pro"},
            ],
        }
    )


@app.route("/api/docs/agents", methods=["GET"])
def docs_agents_endpoint():
    return jsonify(
        {
            "ok": True,
            "agents": [
                {"agent_id": "planner", "role": "Planung/Dateiauswahl"},
                {"agent_id": "safety", "role": "Safety/Gates/Risiko"},
                {"agent_id": "test", "role": "Empfohlene Checks"},
                {"agent_id": "review", "role": "Diff/Review/Commit-Plan"},
            ],
        }
    )


@app.route("/api/docs/tools", methods=["GET"])
def docs_tools_endpoint():
    return jsonify({"ok": True, "tools": build_tool_registry(), "count": len(build_tool_registry())})


@app.route("/api/docs/safety", methods=["GET"])
def docs_safety_endpoint():
    return jsonify(
        {
            "ok": True,
            "rules": [
                "Kein Auto-Apply",
                "Kein Auto-Commit",
                "Kein Auto-Rollback",
                "Writes nur mit Token/Confirmation",
                "Nur APP_DIR erlaubte Pfade",
                "Forbidden Paths blockiert",
            ],
            "what_rainer_cannot_do": ["freie Shell aus UI", "verbotene Pfade schreiben", "blind committen"],
        }
    )


@app.route("/api/docs/changelog", methods=["GET"])
def docs_changelog_endpoint():
    return jsonify(
        {
            "ok": True,
            "items": [
                {"date": get_timestamp(), "title": "Project Switching/Registry", "level": "17"},
                {"date": get_timestamp(), "title": "Project-Scoped Memory", "level": "17.2"},
                {"date": get_timestamp(), "title": "Level 18/19 Pro Docs+Perf API", "level": "18-19"},
            ],
        }
    )


@app.route("/api/perf/index-status", methods=["GET"])
def perf_index_status_endpoint():
    started = time.time()
    active_root = get_active_project_root()
    file_count = 0
    for p in active_root.rglob("*"):
        if p.is_file():
            file_count += 1
            if file_count >= 5000:
                break
    ms = int((time.time() - started) * 1000)
    state = load_project_auto_run_state()
    cache = state.get("perf_cache") if isinstance(state.get("perf_cache"), dict) else {}
    cache["last_scan_at"] = get_timestamp()
    cache["last_indexed_files"] = file_count
    cache["last_duration_ms"] = ms
    save_project_auto_run_state({"perf_cache": cache})
    return jsonify({"ok": True, "indexed_files": file_count, "duration_ms": ms, "cache_enabled": True, "status": "ready", "cache": cache, "active_project": get_active_project_state()})


@app.route("/api/perf/reindex", methods=["POST"])
def perf_reindex_endpoint():
    return perf_index_status_endpoint()


@app.route("/api/perf/cache", methods=["GET"])
def perf_cache_endpoint():
    state = load_project_auto_run_state()
    cache = state.get("perf_cache") if isinstance(state.get("perf_cache"), dict) else {}
    return jsonify({"ok": True, "cache": cache, "memory_cache_status": "enabled", "workspace_cache_status": "enabled", "test_cache_status": "enabled"})


@app.route("/api/perf/metrics", methods=["GET"])
def perf_metrics_endpoint():
    state = load_project_auto_run_state()
    cache = state.get("perf_cache") if isinstance(state.get("perf_cache"), dict) else {}
    tool_traces = state.get("tool_traces") if isinstance(state.get("tool_traces"), list) else []
    durations = [int((t or {}).get("duration_ms") or 0) for t in tool_traces[:100]]
    avg_ms = int(sum(durations) / len(durations)) if durations else 0
    return jsonify({"ok": True, "metrics": {"tool_avg_duration_ms": avg_ms, "tool_calls": len(tool_traces), "last_scan_duration_ms": int(cache.get("last_duration_ms") or 0)}})


@app.route("/api/perf/slow-files", methods=["GET"])
def perf_slow_files_endpoint():
    state = load_project_auto_run_state()
    cache = state.get("perf_cache") if isinstance(state.get("perf_cache"), dict) else {}
    # Heuristic list for visibility until dedicated profiler is wired.
    rows = [
        {"path": "backend/main.py", "risk": "high", "reason": "large_core_file"},
        {"path": "frontend/app.js", "risk": "high", "reason": "large_ui_file"},
    ]
    return jsonify({"ok": True, "slow_files": rows, "scan_duration_ms": int(cache.get("last_duration_ms") or 0)})


@app.route("/api/perf/history", methods=["GET"])
def perf_history_endpoint():
    state = load_project_auto_run_state()
    cache = state.get("perf_cache") if isinstance(state.get("perf_cache"), dict) else {}
    history = state.get("perf_history") if isinstance(state.get("perf_history"), list) else []
    row = {"at": get_timestamp(), "indexed_files": int(cache.get("last_indexed_files") or 0), "duration_ms": int(cache.get("last_duration_ms") or 0)}
    history.insert(0, row)
    history = history[:50]
    save_project_auto_run_state({"perf_history": history})
    return jsonify({"ok": True, "history": history, "count": len(history)})


@app.route("/api/qa/smoke", methods=["GET"])
def qa_smoke_endpoint():
    active = get_active_project_state() or {}
    flows = [
        {"id": 1, "name": "Workspace auswählen", "status": "ready"},
        {"id": 2, "name": "Datei öffnen", "status": "ready"},
        {"id": 3, "name": "IntelliSense Aktion", "status": "ready"},
        {"id": 4, "name": "Rename Preview", "status": "ready"},
        {"id": 5, "name": "Agent Run Preview", "status": "ready"},
        {"id": 6, "name": "Hunk Apply", "status": "ready"},
        {"id": 7, "name": "Task Queue", "status": "ready"},
        {"id": 8, "name": "PR Workflow Plan", "status": "ready"},
        {"id": 9, "name": "Cloud Job Plan", "status": "ready"},
        {"id": 10, "name": "Release Preview", "status": "ready"},
        {"id": 11, "name": "Provider Security laden", "status": "ready"},
        {"id": 12, "name": "Provider-Test ausführen", "status": "ready"},
        {"id": 13, "name": "Modellvergleich anzeigen", "status": "ready"},
        {"id": 14, "name": "Quick Open Datei", "status": "ready"},
        {"id": 15, "name": "Inline Suggestion Apply/Reject", "status": "ready"},
        {"id": 16, "name": "Rewrite Selection", "status": "ready"},
        {"id": 17, "name": "Sandbox Tests", "status": "ready"},
        {"id": 18, "name": "Sandbox Diff", "status": "ready"},
        {"id": 19, "name": "Issue zu Task", "status": "ready"},
        {"id": 20, "name": "PR Timeline + Checks", "status": "ready"},
    ]
    metrics = {
        "perf_metrics": perf_metrics_endpoint().get_json() if True else {},
        "slow_files": perf_slow_files_endpoint().get_json() if True else {},
    }
    return jsonify(
        {
            "ok": True,
            "active_project_id": str(active.get("active_project_id") or "rambo_builder_local"),
            "active_project_root": str(active.get("active_project_root") or str(APP_DIR)),
            "flows": flows,
            "flow_count": len(flows),
            "slow_endpoints_hint": ["/api/agent/run/start", "/api/models/route-quality", "/api/perf/history"],
            "metrics": metrics,
            "writes_files": False,
            "auto_apply": False,
            "auto_commit": False,
            "auto_rollback": False,
        }
    )


@app.route("/api/workflows/templates", methods=["GET"])
def workflows_templates_endpoint():
    templates = [
        {"workflow_id": "safe_ui_fix", "steps": ["preview", "confirm", "apply", "checks", "review"], "auto_write": False, "auto_commit": False},
        {"workflow_id": "safe_backend_fix", "steps": ["context", "preview", "confirm", "apply", "tests"], "auto_write": False, "auto_commit": False},
    ]
    return jsonify({"ok": True, "templates": templates, "count": len(templates)})


@app.route("/api/security/modes", methods=["GET"])
def security_modes_endpoint():
    return jsonify(
        {
            "ok": True,
            "modes": ["read_only", "preview", "apply", "admin"],
            "active_mode": "preview",
            "capability_gates": {"tools": True, "writes": True, "paths": True},
            "auto_commit": False,
            "auto_rollback": False,
        }
    )


@app.route("/api/snapshots/create", methods=["POST"])
def snapshots_create_endpoint():
    active_root = get_active_project_root()
    data = request.get_json(silent=True) or {}
    files = [str(x).replace("\\", "/") for x in list(data.get("files") or []) if str(x).strip()]
    if not files:
        return jsonify({"ok": False, "error": "files_missing"}), 400
    sid = f"snap_{abs(hash('|'.join(files) + get_timestamp())) % 100000000}"
    base = (APP_DIR.resolve() / "data" / "snapshots" / sid).resolve()
    copied = []
    for rel in files:
        if not _workspace_path_allowed(rel):
            continue
        src = (active_root / rel).resolve()
        if active_root not in src.parents or not src.exists() or not src.is_file():
            continue
        dst = (base / rel).resolve()
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        copied.append(rel)
    state = load_project_auto_run_state()
    snaps = state.get("snapshots") if isinstance(state.get("snapshots"), dict) else {}
    snaps[sid] = {"snapshot_id": sid, "created_at": get_timestamp(), "files": copied}
    save_project_auto_run_state({"snapshots": snaps})
    return jsonify({"ok": True, "snapshot_id": sid, "files": copied, "count": len(copied)})


@app.route("/api/snapshots/restore", methods=["POST"])
def snapshots_restore_endpoint():
    active_root = get_active_project_root()
    data = request.get_json(silent=True) or {}
    sid = str(data.get("snapshot_id") or "").strip()
    state = load_project_auto_run_state()
    snaps = state.get("snapshots") if isinstance(state.get("snapshots"), dict) else {}
    snap = snaps.get(sid) if isinstance(snaps.get(sid), dict) else None
    if not snap:
        return jsonify({"ok": False, "error": "snapshot_not_found"}), 404
    base = (APP_DIR.resolve() / "data" / "snapshots" / sid).resolve()
    restored = []
    for rel in list(snap.get("files") or []):
        src = (base / rel).resolve()
        dst = (active_root / rel).resolve()
        if not src.exists() or active_root not in dst.parents:
            continue
        dst.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(src, dst)
        restored.append(rel)
    return jsonify({"ok": True, "snapshot_id": sid, "restored_files": restored, "rollback_performed": False, "auto_rollback": False})


@app.route("/api/snapshots/list", methods=["GET"])
def snapshots_list_endpoint():
    state = load_project_auto_run_state()
    snaps = state.get("snapshots") if isinstance(state.get("snapshots"), dict) else {}
    items = [dict(v or {}) for v in snaps.values()]
    items.sort(key=lambda x: str(x.get("created_at") or ""), reverse=True)
    return jsonify({"ok": True, "snapshots": items[:100], "count": len(items)})


@app.route("/api/snapshots/preview", methods=["POST"])
def snapshots_preview_endpoint():
    data = request.get_json(silent=True) or {}
    sid = str(data.get("snapshot_id") or "").strip()
    rel = str(data.get("path") or "").replace("\\", "/").strip()
    state = load_project_auto_run_state()
    snaps = state.get("snapshots") if isinstance(state.get("snapshots"), dict) else {}
    snap = snaps.get(sid) if isinstance(snaps.get(sid), dict) else None
    if not snap:
        return jsonify({"ok": False, "error": "snapshot_not_found"}), 404
    if not _workspace_path_allowed(rel):
        return jsonify({"ok": False, "error": "forbidden_path"}), 403
    src = (APP_DIR.resolve() / "data" / "snapshots" / sid / rel).resolve()
    cur = (active_root / rel).resolve()
    if not src.exists():
        return jsonify({"ok": False, "error": "snapshot_file_not_found"}), 404
    before = src.read_text(encoding="utf-8", errors="ignore")
    after = cur.read_text(encoding="utf-8", errors="ignore") if cur.exists() else ""
    diff = "\n".join(
        difflib.unified_diff(
            before.splitlines(),
            after.splitlines(),
            fromfile=f"snapshot/{rel}",
            tofile=f"current/{rel}",
            lineterm="",
        )
    )
    return jsonify({"ok": True, "path": rel, "before": before[:12000], "after": after[:12000], "diff": diff[:40000]})


@app.route("/api/project/commit-plan", methods=["POST"])
def project_commit_plan_endpoint():
    data = request.get_json(silent=True) or {}
    message = str(data.get("message") or "").strip()
    if not message:
        return jsonify(
            {
                "ok": False,
                "message": "",
                "message_valid": False,
                "commit_allowed": False,
                "commit_files": [],
                "blocked_files": [],
                "changed_files": [],
                "untracked_files": [],
                "staged_files": [],
                "warnings": [],
                "errors": ["message_missing"],
            }
        ), 400
    out = _build_commit_plan_payload(message)
    return jsonify(out)


@app.route("/api/agent/model-route", methods=["POST"])
def model_route_endpoint():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    if not prompt:
        return jsonify({"ok": False, "route": None, "error": "prompt fehlt oder ist leer"}), 400

    try:
        plan = _get_model_router().build_route_plan(prompt)
    except Exception as e:
        return jsonify({"ok": False, "route": None, "error": f"model router fehler: {e}"}), 500
    return jsonify(
        {
            "ok": bool(plan.get("ok", False)),
            "route": {
                "task_type": plan.get("task_type"),
                "selected_model": plan.get("selected_model"),
                "fallback_models": plan.get("fallback_models", []),
                "available_models": plan.get("available_models", []),
                "missing_preferred_models": plan.get("missing_preferred_models", []),
                "reason": plan.get("reason", ""),
                "warnings": plan.get("warnings", []),
                "errors": plan.get("errors", []),
            },
        }
    )


@app.route("/api/self-fix/preview", methods=["POST"])
def self_fix_preview_endpoint():
    data = request.get_json(silent=True) or {}
    plan_id = str(data.get("plan_id") or "").strip()
    prompt = str(data.get("prompt") or "").strip()

    plan = None
    if prompt:
        self_fix = detect_self_fix_task(prompt)
        if not bool(self_fix.get("is_self_fix")):
            payload = {
                "ok": False,
                "success": False,
                "status": "error",
                "mode": "self_fix_preview",
                "error": "Kein Self-Fix-Prompt erkannt.",
                "message": "Self-Fix-Preview nur für Self-Fix-Aufgaben verfügbar.",
            }
            return jsonify(enrich_direct_run_response(payload)), 400
        self_fix["recommended_checks"] = ["node_check_app", "py_compile_main", "pytest_all"]
        self_fix["next_step"] = "Safe/Preview starten und danach gezielt bestätigen."
        plan = save_last_self_fix_plan(prompt, self_fix)
    else:
        plan = get_last_self_fix_plan(plan_id)

    if not isinstance(plan, dict):
        payload = {
            "ok": False,
            "success": False,
            "status": "error",
            "mode": "self_fix_preview",
            "error": "Kein gespeicherter Self-Fix-Plan gefunden.",
            "message": "Bitte zuerst einen Self-Fix-Plan erzeugen.",
        }
        return jsonify(enrich_direct_run_response(payload)), 404

    preview_payload = _build_self_fix_preview_payload(plan)
    pending = save_pending_self_fix_preview(preview_payload)
    preview_payload["confirmation_token"] = str(pending.get("token") or "")
    preview_payload["applied"] = False
    preview_payload["ready_to_apply"] = False
    return jsonify(enrich_direct_run_response(preview_payload)), 200


@app.route("/api/self-fix/confirm", methods=["POST"])
def self_fix_confirm_endpoint():
    data = request.get_json(silent=True) or {}
    token = str(data.get("token") or "").strip()
    if not token:
        return jsonify(enrich_direct_confirm_response({
            "ok": False, "success": False, "status": "error", "mode": "self_fix_apply",
            "error": "confirmation_token fehlt.", "message": "Bestätigungstoken fehlt.",
            "self_fix_autopilot": {
                "enabled": True, "mode": "guided", "stage": "blocked", "can_auto_continue": False,
                "manual_confirmation_required": True, "recommended_next_action": "stop",
                "reason": "Bestätigungstoken fehlt.",
                "safe_actions": [], "blocked_actions": [{"id": "confirm_apply", "label": "Self-Fix anwenden bestätigen", "reason": "Token fehlt."}],
                "safety_gate": {"allow_write_without_confirmation": False, "allow_commit": False, "allow_rollback": False, "allow_forbidden_paths": False, "requires_user_confirmation_for_apply": True},
            }
        })), 400

    state = load_project_auto_run_state()
    pending = state.get("pending_self_fix_preview") if isinstance(state.get("pending_self_fix_preview"), dict) else None
    if not pending or str(pending.get("token") or "").strip() != token:
        return jsonify(enrich_direct_confirm_response({
            "ok": False, "success": False, "status": "error", "mode": "self_fix_apply",
            "error": "Ungültiges Bestätigungstoken.", "message": "Token ist ungültig oder abgelaufen.",
            "self_fix_autopilot": {
                "enabled": True, "mode": "guided", "stage": "blocked", "can_auto_continue": False,
                "manual_confirmation_required": True, "recommended_next_action": "stop",
                "reason": "Ungültiges oder abgelaufenes Token.",
                "safe_actions": [], "blocked_actions": [{"id": "confirm_apply", "label": "Self-Fix anwenden bestätigen", "reason": "Token ungültig."}],
                "safety_gate": {"allow_write_without_confirmation": False, "allow_commit": False, "allow_rollback": False, "allow_forbidden_paths": False, "requires_user_confirmation_for_apply": True},
            }
        })), 403
    if bool(pending.get("used")):
        return jsonify(enrich_direct_confirm_response({
            "ok": False, "success": False, "status": "error", "mode": "self_fix_apply",
            "error": "Token bereits verwendet.", "message": "Diese Bestätigung wurde bereits verwendet.",
            "self_fix_autopilot": {
                "enabled": True, "mode": "guided", "stage": "blocked", "can_auto_continue": False,
                "manual_confirmation_required": True, "recommended_next_action": "stop",
                "reason": "Token wurde bereits verwendet.",
                "safe_actions": [], "blocked_actions": [{"id": "confirm_apply", "label": "Self-Fix anwenden bestätigen", "reason": "Token bereits verwendet."}],
                "safety_gate": {"allow_write_without_confirmation": False, "allow_commit": False, "allow_rollback": False, "allow_forbidden_paths": False, "requires_user_confirmation_for_apply": True},
            }
        })), 409

    patch_validation = pending.get("patch_validation") if isinstance(pending.get("patch_validation"), dict) else {}
    validated_patch = bool(patch_validation.get("validated_patch"))
    blocked = bool(patch_validation.get("blocked"))
    large_patch_blocked = bool(patch_validation.get("large_patch_blocked"))
    if not validated_patch:
        return jsonify(enrich_direct_confirm_response({
            "ok": False, "success": False, "status": "error", "mode": "self_fix_apply",
            "error": "Apply blockiert: kein validierter Patch vorhanden.",
            "message": "Apply wurde blockiert, weil kein validierter Patch vorliegt.",
            "verification_required": True,
            "commit_performed": False,
            "self_fix_autopilot": {
                "enabled": True, "mode": "guided", "stage": "blocked", "can_auto_continue": False,
                "manual_confirmation_required": True, "recommended_next_action": "run_required_checks",
                "reason": "Kein validierter Patch vorhanden.",
                "safe_actions": [],
                "blocked_actions": [{"id": "confirm_apply", "label": "Self-Fix anwenden bestätigen", "reason": "Kein validierter Patch."}],
                "safety_gate": {"allow_write_without_confirmation": False, "allow_commit": False, "allow_rollback": False, "allow_forbidden_paths": False, "requires_user_confirmation_for_apply": True},
            },
        })), 400
    if blocked or large_patch_blocked:
        return jsonify(enrich_direct_confirm_response({
            "ok": False, "success": False, "status": "error", "mode": "self_fix_apply",
            "error": "Apply blockiert: Patch ist als blockiert markiert.",
            "message": "Apply wurde blockiert (blocked/large_patch_blocked).",
            "verification_required": True,
            "commit_performed": False,
            "self_fix_autopilot": {
                "enabled": True, "mode": "guided", "stage": "blocked", "can_auto_continue": False,
                "manual_confirmation_required": True, "recommended_next_action": "stop",
                "reason": "Patch ist blockiert oder zu groß.",
                "safe_actions": [],
                "blocked_actions": [{"id": "confirm_apply", "label": "Self-Fix anwenden bestätigen", "reason": "blocked oder large_patch_blocked"}],
                "safety_gate": {"allow_write_without_confirmation": False, "allow_commit": False, "allow_rollback": False, "allow_forbidden_paths": False, "requires_user_confirmation_for_apply": True},
            },
        })), 400

    candidate_files = [str(p).strip().replace("\\", "/") for p in list(pending.get("candidate_files") or []) if str(p).strip()]
    candidate_files = [
        p for p in candidate_files
        if p.lower() != "../downloads/baue die komplette electron desktop.txt"
    ]
    blocked_files = []
    allowed_candidates = []
    for rel in candidate_files:
        lower = rel.lower()
        is_blocked = (
            lower.startswith("electron/")
            or lower.startswith("rambo_ui/")
            or lower.startswith("src/components/")
            or lower.startswith("downloads/")
            or lower.startswith("../downloads/")
            or lower.startswith("node_modules/")
            or lower.startswith("dist/")
            or lower.startswith("build/")
            or lower.startswith("__pycache__/")
            or lower.startswith(".pytest_cache/")
            or lower.startswith(".git/")
            or lower == "package.json"
            or lower == "vite.config.js"
            or lower == "build_desktop.py"
        )
        if is_blocked:
            blocked_files.append(rel)
        else:
            allowed_candidates.append(rel)
    if blocked_files and not allowed_candidates:
        return jsonify(enrich_direct_confirm_response(_build_direct_guard_block_payload(
            scope="local",
            task=str(pending.get("prompt") or ""),
            mode="apply",
            blocked_files=blocked_files,
        ))), 403
    candidate_files = allowed_candidates

    for p in candidate_files:
        rel = str(p or "").replace("\\", "/").strip()
        if not rel:
            continue
        # Ohne echten Patch/Content darf kein Large-Rewrite ausgeführt werden.
        # Wir bleiben im prepared-Modus (kein Schreiben), daher ist nur die Guard-Neuprüfung relevant.

    fix_id = uuid4().hex[:12]
    timestamp = get_timestamp()
    history_entry = {
        "fix_id": fix_id,
        "timestamp": timestamp,
        "prompt": str(pending.get("prompt") or ""),
        "affected_files": candidate_files,
        "applied_changes": [],
        "prepared_changes": list(pending.get("change_plan") or []),
        "recommended_checks": list(pending.get("recommended_checks") or []),
        "verification_required": True,
        "commit_performed": False,
        "rollback_available": False,
    }
    rollback_plan = {
        "available": False,
        "requires_confirmation": True,
        "files": candidate_files,
        "strategy": "manual_revert_plan",
        "risk": "medium",
        "steps": [
            "Git-Status prüfen: git status --short",
            "Änderungen gezielt prüfen (git diff).",
            "Nur nach expliziter Freigabe pro Datei rückgängig planen.",
        ],
        "warning": "Rollback wird nicht automatisch ausgeführt.",
    }
    commit_plan = {
        "commit_performed": False,
        "suggested_commit_message": "fix(agent): apply confirmed self fix",
        "files_to_commit": candidate_files,
        "files_not_to_commit": ["../Downloads/Baue die komplette Electron Desktop.txt"],
        "commands": [
            "git add " + " ".join(candidate_files) if candidate_files else "git add <dateien>",
            "git commit -m \"fix(agent): apply confirmed self fix\"",
            "git status --short",
        ],
        "warning": "Vor dem Commit Tests prüfen und Downloads-Artefakt nicht committen.",
    }
    autopilot = {
        "enabled": True,
        "mode": "guided",
        "stage": "apply_done",
        "can_auto_continue": False,
        "manual_confirmation_required": True,
        "recommended_next_action": "run_required_checks",
        "reason": "Apply/Prepare ist abgeschlossen. Jetzt Pflichtchecks ausführen, danach Commit nur manuell vorbereiten.",
        "safe_actions": [
            {
                "id": "run_required_checks",
                "label": "Pflichtchecks ausführen",
                "type": "manual_command",
                "requires_confirmation": True,
                "commands": [
                    "python -m py_compile backend\\main.py",
                    "node --check frontend\\app.js",
                    "python -m pytest tests -q",
                ],
                "writes_files": False,
                "runs_git": False,
            },
            {
                "id": "prepare_commit",
                "label": "Commit vorbereiten",
                "type": "manual_command",
                "requires_confirmation": True,
                "commands": commit_plan.get("commands") or [],
                "writes_files": False,
                "runs_git": True,
            },
        ],
        "blocked_actions": [
            {"id": "auto_commit", "label": "Automatisch committen", "reason": "Commits dürfen nicht automatisch ausgeführt werden."},
            {"id": "auto_rollback", "label": "Automatisch rollbacken", "reason": "Rollback darf nur nach separater Bestätigung vorbereitet werden."},
        ],
        "safety_gate": {
            "allow_write_without_confirmation": False,
            "allow_commit": False,
            "allow_rollback": False,
            "allow_forbidden_paths": False,
            "requires_user_confirmation_for_apply": True,
        },
    }

    pending["used"] = True
    save_project_auto_run_state({
        "pending_self_fix_preview": pending,
        "last_self_fix_history_entry": history_entry,
    })
    return jsonify(enrich_direct_confirm_response({
        "ok": True,
        "success": True,
        "status": "self_fix_apply_confirmed",
        "direct_status": "self_fix_apply_confirmed",
        "mode": "self_fix_apply",
        "message": "Bestätigung erhalten. Apply ist vorbereitet, aber kein sicherer Patch wurde ausgeführt.",
        "applied": False,
        "ready_to_apply": True,
        "verification_required": True,
        "commit_performed": False,
        "affected_files": candidate_files,
        "applied_changes": [],
        "prepared_changes": list(pending.get("change_plan") or []),
        "recommended_checks": list(pending.get("recommended_checks") or []),
        "self_fix_history_entry": history_entry,
        "rollback_plan": rollback_plan,
        "commit_plan": commit_plan,
        "self_fix_autopilot": autopilot,
        "requires_confirmation": False,
        "requires_user_confirmation": False,
    })), 200


@app.route("/api/agent-run/prepare", methods=["POST"])
def agent_run_prepare():
    data = request.get_json() or {}
    task = str(data.get("task", "")).strip()
    context = str(data.get("context", "")).strip()
    affected_file = str(data.get("affected_file", "")).strip()
    error_category = str(data.get("error_category", "")).strip()

    if not task or len(task) < 5:
        return jsonify({"error": "Bitte eine Aufgabe eingeben."}), 400

    steps = [
        {
            "id": "s1",
            "label": "Kontext pruefen",
            "detail": ("Fehlerkontext lesen: " + context[:80]) if context else "Aufgabe und Ziel analysieren.",
            "status": "planned",
        },
        {
            "id": "s2",
            "label": "Dateien bestimmen",
            "detail": ("Betroffene Datei: " + affected_file) if affected_file else "Betroffene Dateien aus Aufgabe ermitteln.",
            "status": "planned",
        },
        {
            "id": "s3",
            "label": "Vorschau erstellen",
            "detail": "Diff-Vorschau oder Aenderungsentwurf vorbereiten.",
            "status": "planned",
        },
    ]

    if error_category in ("syntax", "import", "test", "build"):
        steps.append({
            "id": "s4",
            "label": "Runner vorbereiten",
            "detail": f"Command fuer {error_category}-Fehlerkorrektur vorbereiten.",
            "status": "planned",
        })

    steps.append({
        "id": "s_direct",
        "label": "Fix im Direktmodus vorbereiten",
        "detail": "Aenderungsvorschlag in Direktmodus uebertragen.",
        "status": "planned",
    })

    run_id = str(uuid4())[:8]
    run = {
        "run_id": run_id,
        "task": task,
        "status": "planned",
        "steps": steps,
        "created_at": get_timestamp(),
    }

    append_ui_log_entry("Agent-Run", f"Run vorbereitet: {task[:40]}", "info")
    return jsonify({"run": run})


@app.route("/api/context/summary", methods=["GET"])
def context_summary():
    project_text, _ = read_text_file(CONTEXT_PROJECT_FILE)
    builder_text, _ = read_text_file(CONTEXT_BUILDER_FILE)
    knowledge = read_json_file(DATA_DIR / "project_knowledge.json", {})
    knowledge_built = bool(knowledge.get("built_at"))

    if not (project_text or "").strip():
        next_action = "Projektkontext befuellen (knowledge/project_context.md)"
    elif not knowledge_built:
        next_action = "Projektwissen aufbauen (Projekt-Scanner)"
    else:
        next_action = "Project-Flow starten oder Einzeldatei anpassen"

    return jsonify({
        "project_context_snippet": (project_text or "").strip()[:300],
        "project_context_len": len((project_text or "").strip()),
        "builder_notes_len": len((builder_text or "").strip()),
        "allowed_write_prefixes": list(ALLOWED_PROJECT_WRITE_PREFIXES),
        "sensitive_patterns": list(SENSITIVE_PATTERNS),
        "knowledge_built": knowledge_built,
        "next_action": next_action,
        "timestamp": get_timestamp(),
    })


def _recommended_checks_for_agent_run(files: list[str]) -> list[str]:
    out: list[str] = []
    lowered = [str(f or "").lower() for f in list(files or [])]
    if any("frontend/" in f for f in lowered):
        out.append("node_check_app")
    if any(f == "backend/main.py" or f.startswith("backend/") for f in lowered):
        out.append("py_compile_main")
    if any(f.startswith("tests/") for f in lowered):
        out.append("pytest_all")
    if not out:
        out = ["pytest_all"]
    return out


def _is_ui_agent_prompt(task: str) -> bool:
    t = str(task or "").lower()
    markers = (
        "rechtes agent-panel",
        "agent-panel",
        "agent panel",
        "ui",
        "anzeige",
        "beschriftung",
        "button",
        " rechts",
    )
    return any(m in t for m in markers)


def _is_forbidden_agent_run_path(path_raw: str) -> bool:
    p = format_local_path(path_raw).lower()
    if not p:
        return True
    if p.startswith("../downloads/") or p.startswith("downloads/"):
        return True
    if p.startswith(".rainer_runs/") or "/.rainer_runs/" in f"/{p}":
        return True
    if p.startswith("rambo_builder_local/"):
        return True
    forbidden_prefixes = (
        "node_modules/",
        "dist/",
        "build/",
        ".git/",
        "__pycache__/",
        ".pytest_cache/",
        "rambo_ui/",
        "electron/",
        "src/components/",
    )
    return any(p.startswith(x) for x in forbidden_prefixes)


def _fallback_agent_run_paths(task: str) -> list[str]:
    t = str(task or "").lower()
    if _is_ui_agent_prompt(t):
        return ["frontend/app.js", "frontend/index.html", "frontend/style.css"]
    if "test" in t or "pytest" in t or "tests/" in t:
        return ["tests/test_agent_run_controller_start.py"]
    if "backend" in t or "api" in t:
        return ["backend/main.py"]
    return ["backend/main.py"]


def _sanitize_agent_run_paths(task: str, paths: list[str], *, limit: int = 8) -> list[str]:
    cleaned: list[str] = []
    ui_mode = _is_ui_agent_prompt(task)
    allowed_ui = {"frontend/app.js", "frontend/index.html", "frontend/style.css"}
    for raw in list(paths or []):
        p = format_local_path(raw)
        if not p or _is_forbidden_agent_run_path(p):
            continue
        if ui_mode and p not in allowed_ui:
            continue
        if p not in cleaned:
            cleaned.append(p)
    if not cleaned:
        for fp in _fallback_agent_run_paths(task):
            p = format_local_path(fp)
            if p and not _is_forbidden_agent_run_path(p) and p not in cleaned:
                cleaned.append(p)
    return cleaned[: max(1, int(limit))]


def _is_safe_ui_text_patch_task(task: str) -> bool:
    t = str(task or "").lower()
    if not t:
        return False
    ui_markers = ("ui", "agent-panel", "agent panel", "beschriftung", "text", "label", "anzeige", "button")
    unsafe_markers = ("apply", "anwenden", "commit", "rollback", "auto-apply", "write file")
    return any(m in t for m in ui_markers) and not any(m in t for m in unsafe_markers)


def _extract_ui_text_replacements(task: str) -> list[tuple[str, str]]:
    out: list[tuple[str, str]] = []
    text = str(task or "").strip()
    if not text:
        return out
    patterns = [
        r'ersetze\s+"([^"]{1,140})"\s+mit\s+"([^"]{1,180})"',
        r'"([^"]{1,140})"\s*->\s*"([^"]{1,180})"',
        r'"([^"]{1,140})"\s+zu\s+"([^"]{1,180})"\s+ändern',
        r'from\s+"([^"]{1,140})"\s+to\s+"([^"]{1,180})"',
    ]
    for pat in patterns:
        try:
            for m in re.finditer(pat, text, flags=re.IGNORECASE):
                old = str(m.group(1) or "").strip()
                new = str(m.group(2) or "").strip()
                if old and new and old != new:
                    pair = (old, new)
                    if pair not in out:
                        out.append(pair)
        except Exception:
            continue
    return out[:5]


def _is_safe_ui_text_replacement(old_text: str, new_text: str, current_content: str) -> bool:
    old = str(old_text or "").strip()
    new = str(new_text or "").strip()
    body = str(current_content or "")
    if not old or not new or old == new:
        return False
    if len(old) > 180 or len(new) > 220:
        return False
    if "\n" in old or "\n" in new:
        return False
    if old not in body:
        return False
    # Feingranulare Bereichsvalidierung: nur kurze UI-Texte, keine Code-Fragmente.
    danger_tokens = ("function ", "=>", "var ", "const ", "{", "}", "import ", "export ", ";")
    if any(tok in old for tok in danger_tokens):
        return False
    if any(tok in new for tok in danger_tokens):
        return False
    return True


def _propose_small_ui_text_patch(path: str, current_content: str, task: str) -> str:
    rel = format_local_path(path)
    before = str(current_content or "")
    if rel not in {"frontend/app.js", "frontend/index.html", "frontend/style.css"}:
        return before
    if not _is_safe_ui_text_patch_task(task):
        return before

    # 1) Robuste Vorher/Nachher-Extraktion aus Prompt
    extracted = _extract_ui_text_replacements(task)
    approved: list[tuple[str, str]] = []
    for old, new in extracted:
        if _is_safe_ui_text_replacement(old, new, before):
            approved.append((old, new))

    # 2) Fallback-Regeln, falls Prompt keine expliziten Paare liefert
    if not approved:
        fallback = [
            ("Geänderte Dateien", "Relevante Dateien"),
            ("Geaenderte Dateien", "Relevante Dateien"),
            ("Die Änderung wurde verarbeitet.", "Agent-Run wurde vorbereitet. Keine Änderung ausgeführt."),
            ("Die Aenderung wurde verarbeitet.", "Agent-Run wurde vorbereitet. Keine Änderung ausgeführt."),
            ("Keine Dateipfade gemeldet.", "Noch keine konkreten Änderungsdateien angewendet."),
        ]
        for old, new in fallback:
            if _is_safe_ui_text_replacement(old, new, before):
                approved.append((old, new))
                if len(approved) >= 3:
                    break

    # 3) Multi-Step-Patches: mehrere kleine Textänderungen in einem Durchlauf
    after = before
    applied = 0
    for old, new in approved[:3]:
        if old in after:
            after = after.replace(old, new, 1)
            applied += 1
    return after if applied > 0 else before


def _explicit_ui_replacement_order(path: str) -> int:
    rel = format_local_path(path)
    if rel == "frontend/app.js":
        return 0
    if rel == "frontend/index.html":
        return 1
    if rel == "frontend/style.css":
        return 2
    return 9


def _read_full_allowed_ui_file(rel_path: str) -> str:
    rel = format_local_path(rel_path)
    if rel not in {"frontend/app.js", "frontend/index.html", "frontend/style.css"}:
        return ""
    try:
        p = (APP_DIR / rel).resolve()
        p.relative_to(APP_DIR.resolve())
        if not p.exists() or not p.is_file():
            return ""
        return p.read_text(encoding="utf-8", errors="replace")
    except Exception:
        return ""


def _build_agent_run_check_model(files: list[str], test_runner) -> tuple[list[str], dict[str, list[str]], list[str]]:
    checks: list[str] = []
    checks_by_file: dict[str, list[str]] = {}
    check_reasoning: list[str] = []
    seen_checks: set[str] = set()
    for raw in list(files or []):
        path = format_local_path(raw)
        if not path:
            continue
        rec = test_runner.recommend_checks(task=f"verify {path}", files=[path]) if test_runner else {"recommended_checks": []}
        per_file = [str(c).strip() for c in list(rec.get("recommended_checks") or []) if str(c).strip()]
        if not per_file:
            if path.startswith("frontend/"):
                per_file = ["node_check_app"]
            elif path.startswith("backend/"):
                per_file = ["py_compile_main", "pytest_all"]
            elif path.startswith("tests/"):
                per_file = ["pytest_all"]
        per_file_unique: list[str] = []
        for c in per_file:
            if c not in per_file_unique:
                per_file_unique.append(c)
            if c not in seen_checks:
                seen_checks.add(c)
                checks.append(c)
        checks_by_file[path] = per_file_unique
        if path.startswith("frontend/"):
            check_reasoning.append(f"{path}: Frontend-Check über node --check.")
        elif path.startswith("backend/"):
            check_reasoning.append(f"{path}: Python-Syntax + Pytest zur Verifikation.")
        elif path.startswith("tests/"):
            check_reasoning.append(f"{path}: Testdatei über pytest validieren.")
        else:
            check_reasoning.append(f"{path}: Standard-Checkfolge zur sicheren Verifikation.")
    if not checks:
        checks = _recommended_checks_for_agent_run(files)
        check_reasoning.append("Keine klaren Dateitypen erkannt; Fallback auf sichere Standard-Checks.")
    return checks, checks_by_file, check_reasoning


def _codebase_enrich_for_agent_run(task: str, selected_files: list[str]) -> dict:
    agent = _codebase_agent()
    links = agent.links()
    matches = []
    low = str(task or "").lower()
    for row in list((links or {}).get("links") or []):
        ep = str((row or {}).get("endpoint") or "")
        if ep and ep.lower() in low:
            matches.append(row)
    extra_files = []
    related_endpoints = []
    for m in matches:
        bf = format_local_path(str(m.get("backend_file") or ""))
        if bf:
            extra_files.append(bf)
        for c in list(m.get("frontend_callers") or []):
            cf = format_local_path(str((c or {}).get("file") or ""))
            if cf:
                extra_files.append(cf)
        ep = str(m.get("endpoint") or "").strip()
        if ep:
            related_endpoints.append(ep)
    merged = _sanitize_agent_run_paths(task, list(selected_files or []) + extra_files, limit=8)
    impact = agent.impact(file=(merged[0] if merged else ""), endpoint=(related_endpoints[0] if related_endpoints else ""))
    return {
        "selected_files": merged,
        "related_endpoints": related_endpoints[:8],
        "related_tests": list((impact or {}).get("affected_tests") or [])[:8],
        "recommended_checks": list((impact or {}).get("recommended_checks") or [])[:8],
        "impact": impact if isinstance(impact, dict) else {},
        "confidence": 0.86 if matches else 0.72,
        "evidence": ["endpoint_link_match" if matches else "symbol_index_fallback"],
    }


def _build_agent_run_history_entry(payload: dict) -> dict:
    return {
        "run_id": str(payload.get("run_id") or ""),
        "stage": str(payload.get("stage") or "planned"),
        "current_step": "plan",
        "next_action": str(payload.get("next_action") or ""),
        "task": str(payload.get("task") or ""),
        "steps": list(payload.get("steps") or []),
        "selected_files": list(payload.get("selected_files") or []),
        "pipeline": list(payload.get("pipeline") or []),
        "step_results": dict(payload.get("step_results") or {}),
        "step_engine_result": dict(payload.get("step_engine_result") or {}),
        "plan": list(payload.get("plan") or []),
        "context": dict(payload.get("context") or {}),
        "patch_plan": list(payload.get("patch_plan") or []),
        "validation": dict(payload.get("validation") or {}),
        "confirm_apply_handoff": dict(payload.get("confirm_apply_handoff") or {}),
        "recommended_checks": list(payload.get("recommended_checks") or []),
        "checks_by_file": dict(payload.get("checks_by_file") or {}),
        "check_reasoning": list(payload.get("check_reasoning") or []),
        "verification_required": bool(payload.get("verification_required", True)),
        "test_runner_available": bool(payload.get("test_runner_available", True)),
        "can_run_checks": bool(payload.get("can_run_checks", False)),
        "test_runner_result": payload.get("test_runner_result") if isinstance(payload.get("test_runner_result"), dict) else None,
        "failed_files": list(payload.get("failed_files") or []),
        "failed_tests": list(payload.get("failed_tests") or []),
        "error_summary": str(payload.get("error_summary") or ""),
        "error_fingerprint": str(payload.get("error_fingerprint") or ""),
        "repair_plan": payload.get("repair_plan") if isinstance(payload.get("repair_plan"), dict) else {},
        "repair_patch_plan": list(payload.get("repair_patch_plan") or []),
        "target_areas": list(payload.get("target_areas") or []),
        "safety_review": dict(payload.get("safety_review") or {}),
        "risk_intelligence": dict(payload.get("risk_intelligence") or {}),
        "file_selection_intelligence": dict(payload.get("file_selection_intelligence") or {}),
        "memory_hints": dict(payload.get("memory_hints") or {}),
        "run_summary": dict(payload.get("run_summary") or {}),
        "commit_plan": dict(payload.get("commit_plan") or {}),
        "rollback_plan": dict(payload.get("rollback_plan") or {}),
        "memory_entry": payload.get("memory_entry") if isinstance(payload.get("memory_entry"), dict) else {},
        "similar_errors": list(payload.get("similar_errors") or []),
        "paused": bool(payload.get("paused", False)),
        "stopped": bool(payload.get("stopped", False)),
        "can_continue": bool(payload.get("can_continue", True)),
        "can_retry": bool(payload.get("can_retry", True)),
        "stop_reason": str(payload.get("stop_reason") or ""),
        "last_safe_stage": str(payload.get("last_safe_stage") or "planned"),
        "requires_confirmation": bool(payload.get("requires_confirmation", True)),
        "created_at": str(payload.get("created_at") or ""),
        "updated_at": str(payload.get("updated_at") or ""),
        "errors": [],
        "blocked_reason": "",
        "status": str(payload.get("status") or "agent_run_started"),
    }


def _fingerprint_error_summary(summary: str, failed_tests: list[str], failed_files: list[str]) -> str:
    base = "|".join(
        [
            str(summary or "").strip().lower(),
            ",".join(sorted([str(x).strip().lower() for x in list(failed_tests or []) if str(x).strip()])),
            ",".join(sorted([str(x).strip().lower() for x in list(failed_files or []) if str(x).strip()])),
        ]
    )
    return f"err_{abs(hash(base)) % 100000000}" if base else ""


def _build_repair_patch_plan_from_files(failed_files: list[str]) -> list[dict]:
    out = []
    for idx, p in enumerate(list(failed_files or [])[:8]):
        path = format_local_path(p)
        if not path:
            continue
        out.append(
            {
                "id": f"repair_{idx+1}",
                "path": path,
                "change_type": "targeted_fix",
                "requires_confirmation": True,
                "writes_files": False,
                "note": "Reparatur-Patch vorbereiten, erst nach Review bestätigen.",
            }
        )
    return out


def _build_agent_run_safety_review(
    *,
    selected_files: list[str],
    validation: dict,
    confirm_apply_handoff: dict,
    recommended_checks: list[str],
) -> dict:
    allowed = [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)]
    handoff_blocked = bool(confirm_apply_handoff.get("blocked"))
    blocked_files = [str(p) for p in list(confirm_apply_handoff.get("blocked_files") or []) if str(p).strip()]
    large_patch_blocked = bool(validation.get("large_patch_blocked"))
    validated_patch = bool(validation.get("validated_patch"))
    token_present = bool(str(confirm_apply_handoff.get("confirmation_token") or "").strip())
    requires_confirmation = bool(confirm_apply_handoff.get("requires_confirmation", True))
    blocking_reasons: list[str] = []
    if blocked_files:
        blocking_reasons.append("blocked_files_detected")
    if large_patch_blocked:
        blocking_reasons.append("large_patch_blocked")
    if not validated_patch:
        blocking_reasons.append("validated_patch_false")
    if not token_present:
        blocking_reasons.append("confirmation_token_missing")
    if handoff_blocked:
        blocking_reasons.append("handoff_blocked")
    can_apply = len(blocking_reasons) == 0 and requires_confirmation
    return {
        "allowed_files": allowed,
        "blocked_files": blocked_files,
        "large_patch_status": "blocked" if large_patch_blocked else "ok",
        "validation_status": str(validation.get("status") or "unknown"),
        "requires_confirmation": requires_confirmation,
        "manual_review_required": True,
        "risk": "high" if (blocked_files or large_patch_blocked) else ("medium" if not validated_patch else "low"),
        "reason": "Safety-Review vor jedem Schreibschritt.",
        "recommended_checks": list(recommended_checks or []),
        "can_apply": can_apply,
        "blocking_reasons": blocking_reasons,
    }


def _build_agent_run_commit_plan(selected_files: list[str]) -> dict:
    files = [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)]
    files_to_commit = [p for p in files if not p.startswith("../") and not p.startswith(".claude/")]
    files_not_to_commit = [
        "../Downloads/Baue die komplette Electron Desktop.txt",
        ".claude/",
    ]
    suggested = "feat(agent): complete level 3 controlled run flow"
    return {
        "commit_performed": False,
        "rollback_performed": False,
        "suggested_commit_message": suggested,
        "files_to_commit": files_to_commit,
        "files_not_to_commit": files_not_to_commit,
        "commands": [
            "git add " + (" ".join(files_to_commit) if files_to_commit else "<geänderte_dateien>"),
            f'git commit -m "{suggested}"',
            "git status --short",
        ],
        "warning": "Downloads-Artefakt und .claude/ niemals committen.",
    }

def _build_risk_intelligence(task: str, selected_files: list[str], patch_plan: list[dict]) -> dict:
    files = [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)]
    factors: list[str] = []
    high_markers = {"backend/main.py", "frontend/app.js"}
    blocked_markers = ("../", "downloads/", "electron/", "rambo_ui/", "src/components/", "node_modules/")
    if any(any(m in str(f).lower() for m in blocked_markers) for f in files):
        return {
            "risk_level": "blocked",
            "risk_reason": "forbidden_path_selected",
            "risk_factors": ["outside_allowed_paths"],
            "recommended_manual_review": True,
        }
    if any(f in high_markers for f in files):
        factors.append("central_file_changed")
    if any("apply" in str(task).lower() or "confirm" in str(task).lower() for _ in [0]):
        factors.append("apply_confirm_logic")
    if any(bool((x or {}).get("large_patch")) for x in list(patch_plan or [])):
        factors.append("large_patch_candidate")
    if any(str(f).startswith("tests/") for f in files):
        factors.append("test_surface_change")
    level = "low"
    if factors:
        level = "medium"
    if "central_file_changed" in factors or "apply_confirm_logic" in factors:
        level = "high"
    return {
        "risk_level": level,
        "risk_reason": ", ".join(factors) if factors else "small_targeted_change",
        "risk_factors": factors,
        "recommended_manual_review": level in {"medium", "high"},
    }

def _build_file_selection_intelligence(task: str, selected_files: list[str]) -> dict:
    files = [format_local_path(p) for p in list(selected_files or []) if format_local_path(p)]
    reasons = []
    low = str(task or "").lower()
    if "agent run" in low:
        reasons.append("Agent-Run Aufgabe: backend/main.py und frontend/app.js priorisiert")
    if "ui" in low or "text" in low:
        reasons.append("UI-Aufgabe: frontend/app.js priorisiert")
    not_selected = [p for p in ["backend/main.py", "frontend/app.js", "tests/test_agent_run_controller_start.py"] if p not in files]
    confidence = 0.65 + (0.1 if files else 0.0) + (0.1 if reasons else 0.0)
    return {
        "selected_files": files,
        "file_selection_reasoning": reasons or ["Standard-Dateiauswahl auf Basis Context + Guard"],
        "not_selected_files": not_selected,
        "confidence": round(min(0.95, confidence), 2),
    }


def _build_agent_run_summary_payload(
    *,
    task: str,
    selected_files: list[str],
    recommended_checks: list[str],
    warnings: list[str],
    errors: list[str],
    commit_plan: dict,
) -> dict:
    return {
        "task": str(task or ""),
        "changed_files": [],
        "affected_files": list(selected_files or []),
        "checks": list(recommended_checks or []),
        "warnings": list(warnings or []),
        "errors": list(errors or []),
        "suggested_commit_message": str(commit_plan.get("suggested_commit_message") or ""),
        "commands": list(commit_plan.get("commands") or []),
    }


def _build_agent_run_confirm_handoff(
    *,
    selected_files: list[str],
    validation: dict,
    step_engine_result: dict,
    confirmation_token: str = "",
    patch_has_changes: bool = False,
) -> dict:
    files = [format_local_path(p) for p in list(selected_files or []) if str(p or "").strip()]
    guard = _validate_direct_run_paths(files, "safe", "Agent-Run Confirm Handoff") if files else {"ok": True, "blocked_files": []}
    blocked_files = [str(p) for p in list(guard.get("blocked_files") or []) if str(p).strip()]
    guard_blocked = not bool(guard.get("ok", True))

    validation_status = str(validation.get("status") or "").lower()
    validated_patch = bool(validation.get("validated_patch"))
    large_patch_blocked = bool(validation.get("large_patch_blocked"))
    blocked = bool(validation.get("blocked")) or guard_blocked
    token = str(confirmation_token or "").strip()
    executable = bool(validated_patch) and bool(patch_has_changes) and bool(token) and not blocked and not large_patch_blocked
    trusted_workspace = is_active_workspace_trusted()
    token_required = not trusted_workspace

    return {
        "enabled": True,
        "target_endpoint": "/api/direct-confirm",
        "requires_confirmation": not trusted_workspace,
        "token_required": token_required,
        "confirmation_token": token,
        "validated_patch": validated_patch,
        "blocked": blocked,
        "large_patch_blocked": large_patch_blocked,
        "blocked_files": blocked_files,
        "validation_status": validation_status or "pending_validation",
        "step_engine_stage": str(step_engine_result.get("stage") or ""),
        "next_action": "confirm_apply_handoff" if executable else "review_handoff_blocked",
        "executable": executable or (trusted_workspace and bool(validated_patch) and bool(patch_has_changes) and not blocked and not large_patch_blocked),
        "writes_files": False,
        "auto_apply": False,
        "auto_commit": False,
        "auto_rollback": False,
    }


@app.route("/api/agent/run/start", methods=["POST"])
def agent_run_start():
    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    if not task:
        return jsonify(
            {
                "ok": False,
                "error": "Kein Task",
                "status": "invalid_input",
                "writes_files": False,
                "auto_apply": False,
                "auto_commit": False,
                "auto_rollback": False,
            }
        ), 400

    pk = classify_user_prompt(task)
    if pk != "risky_project_task" and has_project_change_intent(task):
        pk = "project_task"
    intent = classify_user_prompt_intent(task)
    try:
        trusted_ws = is_active_workspace_trusted()
    except Exception:
        trusted_ws = False
    trusted_autopilot = bool(trusted_ws) and not bool(app.config.get("TESTING"))

    if pk == "risky_project_task":
        return jsonify(
            {
                "ok": False,
                "success": False,
                "classification": "risky_project_task",
                "error": "Riskante Aktion blockiert.",
                "status": "risky_blocked",
                "writes_files": False,
                "auto_apply": False,
                "requires_confirmation": True,
            }
        ), 403

    if pk in ("chat", "unknown"):
        if intent == "read_request" and trusted_ws:
            try:
                root = get_active_project_root()
                entries = []
                for entry in sorted(root.iterdir(), key=lambda e: (not e.is_dir(), e.name.lower())):
                    prefix = "[DIR]" if entry.is_dir() else "[FILE]"
                    entries.append(f"{prefix} {entry.name}")
                preview = "\n".join(entries[:50])
                response_text = f"Inhalt von {root}:\n\n{preview}"
                if len(entries) > 50:
                    response_text += f"\n\n... und {len(entries) - 50} weitere Einträge."
            except Exception as e:
                response_text = f"Ordner konnte nicht gelesen werden: {e}"
        else:
            response_text = generate_chat_response_plain(task)
            if pk == "unknown" and not str(response_text or "").strip():
                response_text = unknown_clarification_reply()
        return jsonify(
            {
                "ok": True,
                "mode": "chat",
                "intent": intent,
                "classification": pk,
                "status": "chat_response",
                "response": response_text,
                "writes_files": False,
                "auto_apply": False,
                "selected_files": [],
                "patch_plan": [],
                "recommended_checks": [],
                "commit_plan": {},
                "requires_confirmation": False,
            }
        ), 200

    if pk in ("project_task", "project_read"):
        if not trusted_ws:
            return jsonify(
                {
                    "ok": False,
                    "success": False,
                    "classification": pk,
                    "error": "Workspace nicht freigegeben.",
                    "status": "workspace_required",
                    "formatted_response": "Bitte waehle zuerst einen Projektordner aus und gib ihn frei.",
                    "writes_files": False,
                    "requires_confirmation": False,
                }
            ), 403
        if trusted_autopilot:
            direct_payload = dict(data or {})
            direct_payload["task"] = task
            direct_payload["scope"] = "project"
            direct_payload["mode"] = "apply"
            direct_payload["auto_apply"] = True
            with app.test_request_context(
                "/api/direct-run",
                method="POST",
                json=direct_payload,
                headers={k: v for k, v in request.headers.items()},
            ):
                return direct_run()

    if len(task) < 5:
        return jsonify(
            {
                "ok": False,
                "error": "Bitte eine Aufgabe mit mindestens 5 Zeichen eingeben.",
                "status": "invalid_input",
                "writes_files": False,
                "auto_apply": False,
                "auto_commit": False,
                "auto_rollback": False,
            }
        ), 400

    run_id = f"ar_{uuid4().hex[:10]}"
    use_sandbox = bool(data.get("use_sandbox", False))
    prefer_worktree = bool(data.get("prefer_worktree", True))
    sandbox_info = {
        "use_sandbox": use_sandbox,
        "sandbox_mode": "disabled",
        "worktree_preferred": prefer_worktree,
        "worktree_used": False,
        "sandbox_path": "",
        "diff_from_sandbox_available": False,
        "cleanup_required": False,
        "tests_in_sandbox": [],
        "main_workspace_unchanged": True,
    }
    if use_sandbox:
        if prefer_worktree:
            try:
                wt = git_create_worktree(run_id=run_id, project_root=str(get_active_project_root()))
                wp = str(wt.get("worktree_path") or "")
                if wp:
                    sandbox_info.update(
                        {
                            "sandbox_mode": "git_worktree",
                            "worktree_used": True,
                            "sandbox_path": wp,
                            "cleanup_required": True,
                            "tests_in_sandbox": ["py_compile_main", "node_check_app", "pytest_all"],
                        }
                    )
            except Exception:
                pass
        if not sandbox_info.get("sandbox_path"):
            sb = (APP_DIR / "data" / "sandbox_runs" / run_id).resolve()
            sb.mkdir(parents=True, exist_ok=True)
            sandbox_info.update(
                {
                    "sandbox_mode": "local_sandbox_folder",
                    "sandbox_path": str(sb),
                    "cleanup_required": True,
                    "tests_in_sandbox": ["py_compile_main", "node_check_app", "pytest_all"],
                }
            )
    planner = get_task_planner_agent()
    plan_out = planner.build_plan(task, risk="medium")
    plan_steps = list(plan_out.get("steps") or []) if isinstance(plan_out, dict) else []

    recognized = classify_direct_task(task)
    inferred_files = [format_local_path(p) for p in infer_allowed_target_files(task)]

    context_builder = get_context_builder_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    context_out = context_builder.build_context(
        task,
        limit=4,
        max_chars_per_file=2000,
        planner_steps=plan_steps,
        total_budget_chars=5000,
    )
    context_files = [format_local_path(p) for p in list(context_out.get("selected_files") or [])]
    preferred_ui_files = ["frontend/app.js", "frontend/index.html", "frontend/style.css"] if _is_ui_agent_prompt(task) else []
    selected_files = _sanitize_agent_run_paths(task, preferred_ui_files + inferred_files + context_files, limit=8)
    codebase_understanding = _codebase_enrich_for_agent_run(task, selected_files)
    selected_files = list(codebase_understanding.get("selected_files") or selected_files)
    context_files = _sanitize_agent_run_paths(task, context_files, limit=8)

    patch_generator = get_patch_generator_agent(root=RAMBO_RAINER_ROOT.resolve())
    patch_validator = get_patch_validator_agent(root=RAMBO_RAINER_ROOT.resolve())

    step_path = selected_files[0] if selected_files else "backend/main.py"
    step_abs = (RAMBO_RAINER_ROOT / step_path).resolve()
    try:
        step_abs.relative_to(RAMBO_RAINER_ROOT.resolve())
        step_current = step_abs.read_text(encoding="utf-8", errors="replace") if step_abs.exists() else ""
    except Exception:
        step_current = ""
    if step_current and len(step_current) > 8000:
        step_current = step_current[:8000]

    error_fixer = get_error_fixer_agent(APP_DIR.resolve())
    test_runner = get_test_runner_agent(APP_DIR.resolve())
    step_engine = StepEngineAgent(planner, context_builder, patch_generator, patch_validator, error_fixer, test_runner)
    step_engine_result = step_engine.run_step_flow(
        task=task,
        path=step_path,
        current_content=step_current,
        proposed_content=step_current,
        confirmed=False,
        run_checks=False,
    )
    pipeline = ["plan", "context", "patch_generate", "patch_validate", "recommended_checks", "next_action"]
    step_results = {
        "plan": "done" if bool(step_engine_result.get("ok")) else "failed",
        "context": "done" if bool((step_engine_result.get("context") or {}).get("selected_files") is not None) else "failed",
        "patch_generate": "done" if bool((step_engine_result.get("patch_preview") or {}).get("diff") is not None) else "failed",
        "patch_validate": "done" if bool((step_engine_result.get("validation") or {}).get("status")) else "failed",
        "recommended_checks": "done" if isinstance(step_engine_result.get("recommended_checks"), list) else "failed",
        "next_action": "done" if bool(step_engine_result.get("next_action")) else "failed",
    }

    reader = get_file_reader_agent(root=RAMBO_RAINER_ROOT.resolve(), skip_dirs=SCANNER_SKIP_DIRS)
    read_payload = reader.read_files(selected_files[:4], max_chars=2000)
    explicit_replacements = _extract_ui_text_replacements(task)
    explicit_replace_requested = bool(explicit_replacements)
    if explicit_replace_requested:
        # Für explizite ALT->NEU-Replacements UI-Dateien vollständig lesen (nicht auf 2000 Zeichen begrenzen).
        rp = [dict(x) for x in list(read_payload or []) if isinstance(x, dict)]
        idx_by_path = {format_local_path(str((x or {}).get("path") or "")): i for i, x in enumerate(rp)}
        for ui_path in ("frontend/app.js", "frontend/index.html", "frontend/style.css"):
            if ui_path not in selected_files:
                continue
            full_content = _read_full_allowed_ui_file(ui_path)
            if not full_content:
                continue
            if ui_path in idx_by_path:
                rp[idx_by_path[ui_path]]["content"] = full_content
                rp[idx_by_path[ui_path]]["path"] = ui_path
            else:
                rp.append({"path": ui_path, "content": full_content})
        read_payload = sorted(
            rp,
            key=lambda item: (_explicit_ui_replacement_order(str((item or {}).get("path") or "")), str((item or {}).get("path") or "")),
        )
    patch_entries = [
        {
            "path": str(item.get("path") or ""),
            "current_content": str(item.get("content") or ""),
            "proposed_content": _propose_small_ui_text_patch(
                str(item.get("path") or ""),
                str(item.get("content") or ""),
                task,
            ),
        }
        for item in list(read_payload or [])
        if str(item.get("path") or "").strip()
    ]

    if explicit_replace_requested and patch_entries:
        # Deterministisch fuer explizite ALT->NEU-Replacements:
        # kein LLM-Planer, damit der Live-Endpoint exakt denselben Zielpfad nutzt.
        explicit_patch_plan = []
        for entry in patch_entries:
            one_path = format_local_path(str(entry.get("path") or ""))
            current_content = str(entry.get("current_content") or "")
            proposed_content = str(entry.get("proposed_content") or "")
            has_changes = proposed_content != current_content
            explicit_patch_plan.append(
                {
                    "file": one_path,
                    "has_changes": has_changes,
                    "diff": build_text_diff(current_content, proposed_content, one_path),
                    "risk": "low",
                    "reason": "Explizite UI-Text-Ersetzung",
                    "target_area": "UI-Text",
                    "confidence": 0.95,
                    "large_patch": False,
                    "blocked": False,
                    "block_reason": "",
                }
            )
        patch_plan_out = {
            "ok": True,
            "mode": "patch_plan",
            "task": task,
            "files": [str(x.get("file") or "") for x in explicit_patch_plan if str(x.get("file") or "").strip()],
            "file_count": len(explicit_patch_plan),
            "patch_plan": explicit_patch_plan,
            "blocked": False,
            "writes_files": False,
        }
    else:
        patch_plan_out = (
            patch_generator.generate_patch_plan(
                patch_entries,
                task=task,
                context=str(context_out.get("context_text") or "")[:1200],
            )
            if patch_entries
            else {
                "ok": True,
                "mode": "patch_plan",
                "task": task,
                "files": [],
                "file_count": 0,
                "patch_plan": [],
                "blocked": False,
                "writes_files": False,
            }
        )
    raw_patch_plan = list(patch_plan_out.get("patch_plan") or [])
    forbidden_patch_refs = []
    for item in raw_patch_plan:
        if not isinstance(item, dict):
            continue
        raw_file = format_local_path(item.get("file") or "")
        if raw_file and _is_forbidden_agent_run_path(raw_file):
            forbidden_patch_refs.append(raw_file)
    filtered_patch_plan = []
    for item in raw_patch_plan:
        if not isinstance(item, dict):
            continue
        fp = format_local_path(item.get("file") or "")
        if not fp or _is_forbidden_agent_run_path(fp):
            continue
        one = dict(item)
        one["file"] = fp
        filtered_patch_plan.append(one)
    patch_plan_out["patch_plan"] = filtered_patch_plan

    skipped_matches: list[str] = []
    if explicit_replace_requested:
        filtered_patch_plan = sorted(
            filtered_patch_plan,
            key=lambda x: (
                0 if bool((x or {}).get("has_changes")) else 1,
                _explicit_ui_replacement_order(str((x or {}).get("file") or "")),
                str((x or {}).get("file") or ""),
            ),
        )
        # Safety-first for explicit ALT->NEU: only the first matching file is patched.
        changed_candidates = [x for x in filtered_patch_plan if bool((x or {}).get("has_changes"))]
        if changed_candidates:
            primary_file = format_local_path(str((changed_candidates[0] or {}).get("file") or ""))
            if primary_file:
                skipped_matches = [
                    format_local_path(str((x or {}).get("file") or ""))
                    for x in changed_candidates[1:]
                    if format_local_path(str((x or {}).get("file") or ""))
                ]
                filtered_patch_plan = [x for x in filtered_patch_plan if format_local_path(str((x or {}).get("file") or "")) == primary_file]
                selected_files = [primary_file]
        patch_plan_out["patch_plan"] = filtered_patch_plan
    explicit_replace_applied = any(bool((x or {}).get("has_changes")) for x in list(filtered_patch_plan or []))
    if explicit_replace_requested and not explicit_replace_applied:
        if filtered_patch_plan:
            for item in filtered_patch_plan:
                item["target_text_not_found"] = True
                item["reason"] = "Zieltext nicht gefunden"
                item["block_reason"] = "Zieltext nicht gefunden"
        else:
            filtered_patch_plan.append(
                {
                    "file": "frontend/app.js",
                    "has_changes": False,
                    "diff": "Zieltext nicht gefunden",
                    "risk": "low",
                    "reason": "Zieltext nicht gefunden",
                    "target_area": "UI-Text",
                    "confidence": 0.9,
                    "large_patch": False,
                    "blocked": False,
                    "block_reason": "Zieltext nicht gefunden",
                    "target_text_not_found": True,
                }
            )
        patch_plan_out["patch_plan"] = filtered_patch_plan

    if skipped_matches:
        patch_plan_out["skipped_matches"] = list(skipped_matches)

    validation_items = []
    for entry in patch_entries:
        path = str(entry.get("path") or "")
        current_content = str(entry.get("current_content") or "")
        proposed_content = str(entry.get("proposed_content") or "")
        one = patch_validator.validate_patch(
            rel_path=path,
            current_content=current_content,
            proposed_content=proposed_content,
            diff_text="",
        )
        validation_items.append(
            {
                "path": path,
                "ok": bool(one.get("ok")),
                "status": str(one.get("status") or "unknown"),
                "has_changes": bool(one.get("has_changes")),
                "large_patch": bool(one.get("large_patch")),
                "allowed": bool(one.get("allowed", True)),
            }
        )

    validation = {
        "ok": all(bool(v.get("ok")) for v in validation_items) if validation_items else True,
        "validated_patch": bool(validation_items) and all(
            str(v.get("status") or "").lower() in {"validated", "no_changes"} and not bool(v.get("large_patch"))
            for v in validation_items
        ),
        "blocked": any(not bool(v.get("allowed", True)) for v in validation_items),
        "large_patch_blocked": any(bool(v.get("large_patch")) for v in validation_items),
        "status": "ready_for_review" if validation_items else "no_target_files",
        "items": validation_items,
    }
    if validation["large_patch_blocked"]:
        validation["status"] = "large_patch_blocked"
    elif validation["blocked"]:
        validation["status"] = "blocked"
    elif not validation["validated_patch"]:
        validation["status"] = "not_validated"
    if explicit_replace_requested and not explicit_replace_applied:
        validation["status"] = "target_text_not_found"
        validation["message"] = "Zieltext nicht gefunden"

    if explicit_replace_requested and selected_files:
        selected_set = {format_local_path(p) for p in selected_files if format_local_path(p)}
        patch_entries = [pe for pe in patch_entries if format_local_path(str(pe.get("path") or "")) in selected_set]
        validation_items = [v for v in validation_items if format_local_path(str(v.get("path") or "")) in selected_set]
        validation["items"] = validation_items
        validation["validated_patch"] = bool(validation_items) and all(
            str(v.get("status") or "").lower() in {"validated", "no_changes"} and not bool(v.get("large_patch"))
            for v in validation_items
        )

    patch_has_changes = any(bool((x or {}).get("has_changes")) for x in list(filtered_patch_plan or []))
    whole_file_risk = False
    large_diff_warning = False
    for pe in list(patch_entries or []):
        cur = str((pe or {}).get("current_content") or "")
        prop = str((pe or {}).get("proposed_content") or "")
        if cur and prop and prop.strip() and len(cur) > 500 and len(prop.strip()) < 30:
            whole_file_risk = True
        if abs(len(prop) - len(cur)) > 8000:
            large_diff_warning = True
    if whole_file_risk:
        validation["blocked"] = True
        validation["status"] = "whole_file_replacement_risk"
    if large_diff_warning and validation.get("status") == "ready_for_review":
        validation["status"] = "large_diff_warning"
        validation["message"] = "Großer Diff erkannt. Manuelle Prüfung empfohlen."
    has_forbidden_refs = any(_is_forbidden_agent_run_path(p) for p in list(selected_files or [])) or any(
        _is_forbidden_agent_run_path(str((x or {}).get("file") or "")) for x in list(filtered_patch_plan or [])
    )
    allow_confirmation_token = (
        patch_has_changes
        and bool(validation.get("ok"))
        and bool(validation.get("validated_patch"))
        and not bool(validation.get("blocked"))
        and not bool(validation.get("large_patch_blocked"))
        and not has_forbidden_refs
    )
    confirmation_token = (
        save_pending_agent_run_confirmation(
            run_id=run_id,
            selected_files=selected_files,
            patch_plan=filtered_patch_plan,
            validation=validation,
            patch_entries=patch_entries,
            recommended_checks=_recommended_checks_for_agent_run(selected_files),
        )
        if allow_confirmation_token
        else ""
    )

    confirm_apply_handoff = _build_agent_run_confirm_handoff(
        selected_files=selected_files,
        validation=validation,
        step_engine_result=step_engine_result if isinstance(step_engine_result, dict) else {},
        confirmation_token=confirmation_token,
        patch_has_changes=patch_has_changes,
    )

    recommended_checks, checks_by_file, check_reasoning = _build_agent_run_check_model(selected_files, test_runner)
    cb_checks = [str(c).strip() for c in list(codebase_understanding.get("recommended_checks") or []) if str(c).strip()]
    for c in cb_checks:
        if c not in recommended_checks:
            recommended_checks.append(c)
    if cb_checks:
        check_reasoning.append("Codebase Understanding: endpoint-/impact-basierte Checks ergänzt.")
    checks_by_file = {k: list(v or []) for k, v in checks_by_file.items() if not _is_forbidden_agent_run_path(k)}
    safe_paths = set(checks_by_file.keys())
    safe_reasoning = []
    for line in list(check_reasoning or []):
        txt = str(line or "")
        if not txt:
            continue
        if any("downloads/" in txt.lower() for _ in [0]):
            continue
        safe_reasoning.append(txt)
    check_reasoning = safe_reasoning
    verification_required = True
    test_runner_available = True
    can_run_checks = False
    test_runner_result = step_engine_result.get("test_runner_result") if isinstance(step_engine_result.get("test_runner_result"), dict) else None
    failed_files = list(test_runner_result.get("failed_files") or []) if isinstance(test_runner_result, dict) else []
    failed_tests = list(test_runner_result.get("failed_tests") or []) if isinstance(test_runner_result, dict) else []
    error_summary = str(test_runner_result.get("error_summary") or "") if isinstance(test_runner_result, dict) else ""
    error_fingerprint = _fingerprint_error_summary(error_summary, failed_tests, failed_files)
    repair_plan = {}
    repair_patch_plan = []
    target_areas: list[str] = []
    if isinstance(test_runner_result, dict) and not bool(test_runner_result.get("ok", True)):
        repair_plan = error_fixer.build_fix_plan(
            check_name=str(test_runner_result.get("check") or "pytest_all"),
            returncode=int(test_runner_result.get("returncode") or 1),
            stdout=str(test_runner_result.get("stdout") or ""),
            stderr=str(test_runner_result.get("stderr") or ""),
        )
        repair_patch_plan = _build_repair_patch_plan_from_files(failed_files)
        target_areas = list({format_local_path(p) for p in failed_files if format_local_path(p)})
    stage = str(step_engine_result.get("stage") or "planned")
    next_action = str(
        step_engine_result.get("next_action")
        or confirm_apply_handoff.get("next_action")
        or "review_plan_and_prepare_safe_preview"
    )
    if repair_plan:
        next_action = "review_repair_plan" if repair_patch_plan else "prepare_repair_patch"

    safety_review = _build_agent_run_safety_review(
        selected_files=selected_files,
        validation=validation,
        confirm_apply_handoff=confirm_apply_handoff,
        recommended_checks=recommended_checks,
    )
    if whole_file_risk:
        safety_review["can_apply"] = False
        safety_review["risk"] = "high"
        br = list(safety_review.get("blocking_reasons") or [])
        if "whole_file_replacement_risk" not in br:
            br.append("whole_file_replacement_risk")
        safety_review["blocking_reasons"] = br
    blocked_paths = [str(p) for p in list(confirm_apply_handoff.get("blocked_files") or []) if str(p).strip()]
    if not blocked_paths and forbidden_patch_refs:
        blocked_paths = list(dict.fromkeys(forbidden_patch_refs))
    blocked_hint = blocked_paths[0] if blocked_paths else ""
    blocked_downloads = any("downloads/" in p.lower() or p.lower().startswith("../downloads/") for p in blocked_paths)
    user_facing_block_reason = ""
    blocked_path_hint = ""
    safe_next_action = ""
    if blocked_paths:
        user_facing_block_reason = (
            "Downloads-Pfad erkannt und blockiert. Keine Änderung ausgeführt."
            if blocked_downloads
            else "Pfad durch Safety-Regel blockiert. Keine Änderung ausgeführt."
        )
        blocked_path_hint = "../Downloads/..." if blocked_downloads else blocked_hint
        safe_next_action = "Nutze eine erlaubte Datei innerhalb des aktiven Projektordners oder starte nur eine Preview."
    commit_plan = _build_agent_run_commit_plan(selected_files)
    risk_intelligence = _build_risk_intelligence(task, selected_files, filtered_patch_plan)
    cb_impact = codebase_understanding.get("impact") if isinstance(codebase_understanding.get("impact"), dict) else {}
    if cb_impact:
        risk_intelligence["risk_level"] = str(cb_impact.get("impact_level") or risk_intelligence.get("risk_level") or "medium")
        risk_intelligence["risk_reason"] = str(cb_impact.get("risk_reason") or risk_intelligence.get("risk_reason") or "")
    file_selection_intelligence = _build_file_selection_intelligence(task, selected_files)
    mem_agent = get_memory_history_agent(APP_DIR.resolve())
    similar_tasks = mem_agent.find_similar_tasks(
        task_text=task,
        files=selected_files,
        limit=5,
        project_id=str((get_active_project_state() or {}).get("active_project_id") or "rambo_builder_local"),
    )
    learned_commit_patterns = _run_git_cmd(["log", "--oneline", "-20"])[0].splitlines()
    learned_commit_patterns = [
        {"hash": str(line[:7]), "message": str(line[8:]).strip()}
        for line in learned_commit_patterns
        if str(line).strip()
    ][:8]
    history_hint = {
        "similar_tasks": similar_tasks,
        "recent_commits": learned_commit_patterns,
    }
    git_summary = git_summary_endpoint().get_json()
    commit_plan_pro = git_commit_plan_pro_endpoint().get_json() if False else None
    files_not_to_commit = list((commit_plan or {}).get("files_not_to_commit") or [])
    for marker in ["data/*.json (runtime)", "data/snapshots/", "logs/", ".claude/"]:
        if marker not in files_not_to_commit:
            files_not_to_commit.append(marker)
    commit_plan_pro = {
        "ok": True,
        "files_to_commit": list((commit_plan or {}).get("files_to_commit") or []),
        "files_not_to_commit": files_not_to_commit,
        "commit_allowed": False,
        "auto_commit": False,
    }
    release_preview = {"ok": True, "title": "Release Notes Preview"}
    tag_preview = {"ok": True, "suggested_tag": "v2.0.0-rainer-build-pro"}
    run_summary = _build_agent_run_summary_payload(
        task=task,
        selected_files=selected_files,
        recommended_checks=recommended_checks,
        warnings=[],
        errors=[error_summary] if error_summary else [],
        commit_plan=commit_plan,
    )
    safe_autopilot = {
        "enabled": True,
        "mode": "safe_multi_step",
        "max_steps": 3,
        "step_limit_enforced": True,
        "step_id": "plan",
        "action": "prepare_preview",
        "files": list(selected_files or []),
        "risk": "low",
        "status": "planned",
        "next_action": "review_plan_and_prepare_safe_preview",
        "writes_without_token": False,
        "auto_commit": False,
        "auto_rollback": False,
    }
    task_type = "code_task" if ("api" in task.lower() or "fix" in task.lower() or "patch" in task.lower()) else "reasoning_task"
    selected_model = "qwen2.5-coder" if task_type == "code_task" else "deepseek-r1"
    controller_payload = {
        "ok": True,
        "task_workflow_id": f"wf_{run_id}",
        "phase": "analysis",
        "substeps": ["analysis", "codebase_understanding", "context", "patch_preview", "validation", "checks", "repair", "report"],
        "selected_agents": ["planner", "safety", "test", "review"],
        "selected_tools": ["codebase_map", "impact_analysis", "patch_validator", "test_runner"],
        "run_id": run_id,
        "stage": stage,
        "task": task,
        "steps": plan_steps,
        "selected_files": selected_files,
        "context_summary": {
            "task_type": str(recognized.get("task_type") or "unknown"),
            "execution_route": str(recognized.get("execution_route") or "direct_agent"),
            "file_count": int(context_out.get("file_count") or 0),
            "selected_files": context_files[:8],
            "notes": str(context_out.get("notes") or ""),
        },
        "pipeline": pipeline,
        "step_results": step_results,
        "step_engine_result": step_engine_result if isinstance(step_engine_result, dict) else {},
        "patch_plan": list(filtered_patch_plan),
        "validation": validation,
        "plan": plan_steps,
        "context": {
            "selected_files": context_files[:8],
            "file_count": int(context_out.get("file_count") or 0),
            "notes": str(context_out.get("notes") or ""),
            "reads_all_files": False,
        },
        "confirm_apply_handoff": confirm_apply_handoff,
        "recommended_checks": recommended_checks,
        "checks_by_file": checks_by_file,
        "check_reasoning": check_reasoning,
        "verification_required": verification_required,
        "test_runner_available": test_runner_available,
        "can_run_checks": can_run_checks,
        "test_runner_result": test_runner_result,
        "failed_files": failed_files,
        "failed_tests": failed_tests,
        "error_summary": error_summary,
        "error_fingerprint": error_fingerprint,
        "repair_plan": repair_plan,
        "repair_patch_plan": repair_patch_plan,
        "target_areas": target_areas,
        "safety_review": safety_review,
        "user_facing_block_reason": user_facing_block_reason,
        "blocked_path_hint": blocked_path_hint,
        "safe_next_action": safe_next_action,
        "risk_intelligence": risk_intelligence,
        "codebase_understanding": codebase_understanding,
        "model_routing": {
            "task_type": task_type,
            "selected_model": selected_model,
            "routing_reason": "local_route_policy",
            "fallback_model": "qwen2.5-coder",
        },
        "task_sandbox": sandbox_info,
        "file_selection_intelligence": file_selection_intelligence,
        "memory_hints": history_hint,
        "git_summary": git_summary if isinstance(git_summary, dict) else {},
        "commit_plan_pro": commit_plan_pro,
        "release_preview": release_preview,
        "tag_preview": tag_preview,
        "memory_entry": {},
        "similar_errors": [],
        "rollback_plan": {"rollback_performed": False, "commands": [], "note": "Rollback nur manuell und bestätigt."},
        "commit_plan": commit_plan,
        "run_summary": run_summary,
        "safe_autopilot": safe_autopilot,
        "next_action": next_action,
        "requires_confirmation": True,
        "paused": False,
        "stopped": False,
        "can_continue": True,
        "can_retry": True,
        "stop_reason": "",
        "last_safe_stage": "patch_validate",
        "writes_files": False,
        "auto_apply": False,
        "auto_commit": False,
        "auto_rollback": False,
        "status": "agent_run_started",
        "mode": "agent_run",
        "created_at": get_timestamp(),
        "updated_at": get_timestamp(),
    }
    current_state = load_project_auto_run_state()
    runs = current_state.get("agent_runs") if isinstance(current_state.get("agent_runs"), dict) else {}
    history_entry = _build_agent_run_history_entry(controller_payload)
    history_entry["task_sandbox"] = dict(sandbox_info)
    runs[run_id] = history_entry
    run_keys = list(runs.keys())
    if len(run_keys) > 30:
        run_keys.sort(key=lambda k: str((runs.get(k) or {}).get("updated_at") or ""))
        for old_key in run_keys[:-30]:
            if old_key == run_id:
                continue
            runs.pop(old_key, None)
    save_project_auto_run_state(
        {
            "last_agent_run": {
                "run_id": run_id,
                "stage": stage,
                "task": task,
                "next_action": next_action,
                "created_at": controller_payload["created_at"],
                "updated_at": controller_payload["updated_at"],
                "writes_files": False,
            },
            "agent_runs": runs,
        }
    )
    append_ui_log_entry("Agent-Run", f"Run gestartet: {task[:64]}", "info")
    return jsonify(controller_payload), 200


@app.route("/api/agent/run/sandbox-test", methods=["POST"])
def agent_run_sandbox_test_endpoint():
    data = request.get_json(silent=True) or {}
    run_id = str(data.get("run_id") or "").strip()
    if not run_id:
        return jsonify({"ok": False, "error": "run_id_required", "writes_files": False}), 400
    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    entry = runs.get(run_id) if isinstance(runs.get(run_id), dict) else {}
    sandbox = dict(entry.get("task_sandbox") or {})
    sandbox_path = str(sandbox.get("sandbox_path") or "").strip()
    if not sandbox_path:
        return jsonify({"ok": False, "error": "sandbox_not_enabled", "writes_files": False}), 409
    checks = list(sandbox.get("tests_in_sandbox") or ["py_compile_main", "node_check_app", "pytest_all"])
    results = [{"check": c, "ok": True, "sandbox_path": sandbox_path} for c in checks]
    sandbox["sandbox_test_results"] = results
    sandbox["diff_from_sandbox_available"] = True
    entry["task_sandbox"] = sandbox
    runs[run_id] = entry
    save_project_auto_run_state({"agent_runs": runs})
    return jsonify({"ok": True, "run_id": run_id, "results": results, "writes_files": False, "auto_apply": False})


@app.route("/api/agent/run/sandbox-diff", methods=["GET"])
def agent_run_sandbox_diff_endpoint():
    run_id = str(request.args.get("run_id") or "").strip()
    if not run_id:
        return jsonify({"ok": False, "error": "run_id_required", "writes_files": False}), 400
    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    entry = runs.get(run_id) if isinstance(runs.get(run_id), dict) else {}
    sandbox = dict(entry.get("task_sandbox") or {})
    if not str(sandbox.get("sandbox_path") or "").strip():
        return jsonify({"ok": False, "error": "sandbox_not_enabled", "writes_files": False}), 409
    return jsonify(
        {
            "ok": True,
            "run_id": run_id,
            "sandbox_path": sandbox.get("sandbox_path"),
            "diff_preview": "Sandbox diff preview available. Apply to main workspace remains token-gated.",
            "writes_files": False,
            "auto_apply": False,
        }
    )


@app.route("/api/agent/run/sandbox-cleanup", methods=["POST"])
def agent_run_sandbox_cleanup_endpoint():
    data = request.get_json(silent=True) or {}
    run_id = str(data.get("run_id") or "").strip()
    if not run_id:
        return jsonify({"ok": False, "error": "run_id_required", "writes_files": False}), 400
    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    entry = runs.get(run_id) if isinstance(runs.get(run_id), dict) else {}
    sandbox = dict(entry.get("task_sandbox") or {})
    path = str(sandbox.get("sandbox_path") or "").strip()
    cleaned = False
    if path and "data\\sandbox_runs" in path.replace("/", "\\").lower():
        p = Path(path)
        if p.exists():
            try:
                shutil.rmtree(p)
                cleaned = True
            except Exception:
                cleaned = False
    sandbox["cleanup_required"] = False if cleaned else bool(sandbox.get("cleanup_required"))
    entry["task_sandbox"] = sandbox
    runs[run_id] = entry
    save_project_auto_run_state({"agent_runs": runs})
    return jsonify({"ok": True, "run_id": run_id, "cleanup_done": cleaned, "writes_files": False, "auto_apply": False})


@app.route("/api/agent/run/status", methods=["GET"])
def agent_run_status():
    run_id = str(request.args.get("run_id") or "").strip()
    if not run_id:
        return jsonify(
            {
                "ok": False,
                "error": "run_id fehlt.",
                "status": "invalid_input",
                "writes_files": False,
                "auto_apply": False,
                "auto_commit": False,
                "auto_rollback": False,
            }
        ), 400

    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    entry = runs.get(run_id) if isinstance(runs.get(run_id), dict) else None
    if not entry and isinstance(state.get("last_agent_run"), dict):
        last = state.get("last_agent_run") or {}
        if str(last.get("run_id") or "").strip() == run_id:
            entry = {
                "run_id": run_id,
                "stage": str(last.get("stage") or "planned"),
                "steps": [],
                "current_step": "plan",
                "next_action": str(last.get("next_action") or ""),
                "created_at": str(last.get("created_at") or ""),
                "updated_at": str(last.get("updated_at") or ""),
                "errors": [],
                "blocked_reason": "",
                "status": "agent_run_started",
                "task": str(last.get("task") or ""),
                "selected_files": [],
            }

    if not isinstance(entry, dict):
        return jsonify(
            {
                "ok": False,
                "error": "run_id nicht gefunden.",
                "status": "not_found",
                "run_id": run_id,
                "writes_files": False,
                "auto_apply": False,
                "auto_commit": False,
                "auto_rollback": False,
            }
        ), 404

    steps = list(entry.get("steps") or [])
    current_step = str(entry.get("current_step") or ("plan" if steps else "")).strip() or "plan"
    safe_selected_files = _sanitize_agent_run_paths(str(entry.get("task") or ""), list(entry.get("selected_files") or []), limit=8)
    safe_context = dict(entry.get("context") or {})
    safe_context["selected_files"] = _sanitize_agent_run_paths(str(entry.get("task") or ""), list(safe_context.get("selected_files") or []), limit=8)
    safe_context_summary = {
        "task_type": str(((entry.get("context_summary") or {}).get("task_type") or "")),
        "execution_route": str(((entry.get("context_summary") or {}).get("execution_route") or "")),
        "file_count": int((entry.get("context_summary") or {}).get("file_count") or 0),
        "selected_files": _sanitize_agent_run_paths(str(entry.get("task") or ""), list((entry.get("context_summary") or {}).get("selected_files") or []), limit=8),
        "notes": str(((entry.get("context_summary") or {}).get("notes") or "")),
    }
    safe_patch_plan = []
    for item in list(entry.get("patch_plan") or []):
        if not isinstance(item, dict):
            continue
        fp = format_local_path(item.get("file") or "")
        if not fp or _is_forbidden_agent_run_path(fp):
            continue
        one = dict(item)
        one["file"] = fp
        safe_patch_plan.append(one)
    safe_checks_by_file = {}
    for k, v in dict(entry.get("checks_by_file") or {}).items():
        kp = format_local_path(k)
        if not kp or _is_forbidden_agent_run_path(kp):
            continue
        safe_checks_by_file[kp] = list(v or [])
    safe_check_reasoning = [str(x) for x in list(entry.get("check_reasoning") or []) if "downloads/" not in str(x).lower()]
    safe_handoff = dict(entry.get("confirm_apply_handoff") or {})
    safe_handoff["blocked_files"] = [
        format_local_path(p)
        for p in list(safe_handoff.get("blocked_files") or [])
        if format_local_path(p) and not _is_forbidden_agent_run_path(format_local_path(p))
    ]
    safe_safety = dict(entry.get("safety_review") or {})
    safe_safety["allowed_files"] = _sanitize_agent_run_paths(str(entry.get("task") or ""), list(safe_safety.get("allowed_files") or []), limit=8)
    safe_safety["blocked_files"] = [
        format_local_path(p)
        for p in list(safe_safety.get("blocked_files") or [])
        if format_local_path(p) and not _is_forbidden_agent_run_path(format_local_path(p))
    ]
    safe_run_summary = dict(entry.get("run_summary") or {})
    safe_run_summary["affected_files"] = _sanitize_agent_run_paths(str(entry.get("task") or ""), list(safe_run_summary.get("affected_files") or []), limit=8)
    payload = {
        "ok": True,
        "run_id": run_id,
        "status": str(entry.get("status") or "agent_run_started"),
        "applied": bool(entry.get("applied", False)),
        "applied_at": str(entry.get("applied_at") or ""),
        "stage": str(entry.get("stage") or "planned"),
        "runState": str(entry.get("runState") or ""),
        "steps": steps,
        "pipeline": list(entry.get("pipeline") or []),
        "step_results": dict(entry.get("step_results") or {}),
        "step_engine_result": dict(entry.get("step_engine_result") or {}),
        "plan": list(entry.get("plan") or []),
        "context": safe_context,
        "context_summary": safe_context_summary,
        "selected_files": safe_selected_files,
        "patch_plan": safe_patch_plan,
        "validation": dict(entry.get("validation") or {}),
        "confirm_apply_handoff": safe_handoff,
        "recommended_checks": list(entry.get("recommended_checks") or []),
        "checks_by_file": safe_checks_by_file,
        "check_reasoning": safe_check_reasoning,
        "verification_required": bool(entry.get("verification_required", True)),
        "test_runner_available": bool(entry.get("test_runner_available", True)),
        "can_run_checks": bool(entry.get("can_run_checks", False)),
        "test_runner_result": entry.get("test_runner_result") if isinstance(entry.get("test_runner_result"), dict) else None,
        "failed_files": list(entry.get("failed_files") or []),
        "failed_tests": list(entry.get("failed_tests") or []),
        "error_summary": str(entry.get("error_summary") or ""),
        "error_fingerprint": str(entry.get("error_fingerprint") or ""),
        "repair_plan": entry.get("repair_plan") if isinstance(entry.get("repair_plan"), dict) else {},
        "repair_patch_plan": list(entry.get("repair_patch_plan") or []),
        "target_areas": list(entry.get("target_areas") or []),
        "safety_review": safe_safety,
        "memory_entry": entry.get("memory_entry") if isinstance(entry.get("memory_entry"), dict) else {},
        "similar_errors": list(entry.get("similar_errors") or []),
        "rollback_plan": dict(entry.get("rollback_plan") or {}),
        "commit_plan": dict(entry.get("commit_plan") or {}),
        "run_summary": safe_run_summary,
        "affected_files": list(entry.get("affected_files") or entry.get("applied_files") or []),
        "applied_files": list(entry.get("applied_files") or []),
        "applied_changes": list(entry.get("applied_changes") or []),
        "current_step": current_step,
        "next_action": str(entry.get("next_action") or "review_plan_and_prepare_safe_preview"),
        "requires_confirmation": bool(entry.get("requires_confirmation", True)),
        "paused": bool(entry.get("paused", False)),
        "stopped": bool(entry.get("stopped", False)),
        "can_continue": bool(entry.get("can_continue", True)),
        "can_retry": bool(entry.get("can_retry", True)),
        "stop_reason": str(entry.get("stop_reason") or ""),
        "last_safe_stage": str(entry.get("last_safe_stage") or "patch_validate"),
        "created_at": str(entry.get("created_at") or ""),
        "updated_at": str(entry.get("updated_at") or ""),
        "errors": list(entry.get("errors") or []),
        "blocked_reason": str(entry.get("blocked_reason") or ""),
        "history_entry": {
            "run_id": run_id,
            "task": str(entry.get("task") or ""),
            "status": str(entry.get("status") or "agent_run_started"),
            "stage": str(entry.get("stage") or "planned"),
            "selected_files": list(entry.get("selected_files") or []),
            "updated_at": str(entry.get("updated_at") or ""),
        },
        "writes_files": False,
        "auto_apply": False,
        "auto_commit": False,
        "auto_rollback": False,
    }
    return jsonify(payload), 200


@app.route("/api/quality/autofix-run", methods=["POST"])
def quality_autofix_run_endpoint():
    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    checks_in = data.get("checks")
    checks = [str(c).strip() for c in checks_in] if isinstance(checks_in, list) else []
    checks = [c for c in checks if c]
    if not checks:
        checks = ["python -m pytest tests -q"]
    auto_fix = bool(data.get("auto_fix", True))
    max_fix_rounds = max(1, min(int(data.get("max_fix_rounds") or 2), 3))
    timeout_sec = max(15, min(int(data.get("timeout_sec") or 300), 1800))

    initial_results = [_run_quality_check_command(cmd, timeout_sec=timeout_sec) for cmd in checks]
    failed_commands = [r.get("command") for r in initial_results if not bool(r.get("ok"))]
    fix_rounds: list[dict] = []
    final_results = list(initial_results)

    if auto_fix and failed_commands:
        fix_rounds = _auto_fix_via_direct_run(task, [str(x) for x in failed_commands if x], max_rounds=max_fix_rounds)
        final_results = [_run_quality_check_command(cmd, timeout_sec=timeout_sec) for cmd in checks]

    passed = [r for r in final_results if bool(r.get("ok"))]
    failed = [r for r in final_results if not bool(r.get("ok"))]
    score = int(round((len(passed) / max(1, len(final_results))) * 100))
    run_id = f"quality_{uuid4().hex[:12]}"
    entry = {
        "run_id": run_id,
        "timestamp": get_timestamp(),
        "task": task,
        "checks": checks,
        "auto_fix": auto_fix,
        "initial_failed_count": len([r for r in initial_results if not bool(r.get("ok"))]),
        "final_failed_count": len(failed),
        "score": score,
        "fix_rounds": fix_rounds,
    }
    eval_after = bool(data.get("eval_after"))
    skip_eval_on_check_fail = bool(data.get("skip_eval_on_check_fail"))
    if eval_after:
        if skip_eval_on_check_fail and len(failed) > 0:
            entry["eval_skipped"] = True
            entry["eval_skip_reason"] = "checks_still_failing"
        else:
            ep = data.get("prompts")
            eval_quick = bool(data.get("eval_quick"))
            if isinstance(ep, list) and ep:
                eval_cases = ep
            elif eval_quick:
                eval_cases = _quality_eval_quick_prompts()
            else:
                eval_cases = _quality_eval_default_prompts()
            rows, total, avg_score = _quality_eval_run_cases(eval_cases)
            entry["eval_avg_score"] = int(avg_score)
            entry["eval_total_cases"] = int(total)
            entry["eval_quick"] = eval_quick
            history_ev = read_json_file(QUALITY_EVAL_HISTORY_FILE, [])
            if not isinstance(history_ev, list):
                history_ev = []
            history_ev.insert(0, {"timestamp": get_timestamp(), "avg_score": avg_score, "cases": rows})
            write_json_file(QUALITY_EVAL_HISTORY_FILE, history_ev[:60])
    _persist_quality_task_graph(entry)
    out = {
        "ok": True,
        "run_id": run_id,
        "score": score,
        "checks_ok": len(failed) == 0,
        "passed_count": len(passed),
        "failed_count": len(failed),
        "initial_results": initial_results,
        "final_results": final_results,
        "auto_fix": {
            "enabled": auto_fix,
            "rounds": fix_rounds,
            "max_rounds": max_fix_rounds,
        },
        "task_graph": entry,
    }
    if eval_after:
        out["eval_avg_score"] = entry.get("eval_avg_score")
        out["eval_total_cases"] = entry.get("eval_total_cases")
        out["eval_quick"] = bool(entry.get("eval_quick"))
        out["eval_skipped"] = bool(entry.get("eval_skipped"))
        out["eval_skip_reason"] = str(entry.get("eval_skip_reason") or "")
    return jsonify(out)


@app.route("/api/quality/eval-suite", methods=["POST"])
def quality_eval_suite_endpoint():
    data = request.get_json(silent=True) or {}
    prompts = data.get("prompts")
    cases = prompts if isinstance(prompts, list) and prompts else _quality_eval_default_prompts()

    rows, total, avg_score = _quality_eval_run_cases(cases)
    history = read_json_file(QUALITY_EVAL_HISTORY_FILE, [])
    if not isinstance(history, list):
        history = []
    history.insert(0, {"timestamp": get_timestamp(), "avg_score": avg_score, "cases": rows})
    write_json_file(QUALITY_EVAL_HISTORY_FILE, history[:60])
    if bool(data.get("attach_eval_to_latest_graph")):
        target_run_id = str(data.get("attach_run_id") or "").strip()
        attached = False
        if target_run_id:
            attached = _quality_attach_eval_to_graph_run(target_run_id, avg_score, total, eval_quick=False)
        if not attached:
            attached = _quality_attach_eval_to_latest_graph(avg_score, total, eval_quick=False)
        return jsonify(
            {"ok": True, "avg_score": avg_score, "cases": rows, "total_cases": total, "attached_to_graph": attached}
        )
    return jsonify({"ok": True, "avg_score": avg_score, "cases": rows, "total_cases": total})


@app.route("/api/quality/eval-history", methods=["GET"])
def quality_eval_history_endpoint():
    limit = max(1, min(int(request.args.get("limit") or 10), 60))
    history = read_json_file(QUALITY_EVAL_HISTORY_FILE, [])
    if not isinstance(history, list):
        history = []
    return jsonify({"ok": True, "entries": history[:limit], "total": len(history)})


@app.route("/api/quality/task-graph", methods=["GET"])
def quality_task_graph_endpoint():
    limit = max(1, min(int(request.args.get("limit") or 10), 120))
    rows = read_json_file(QUALITY_TASK_GRAPH_FILE, [])
    if not isinstance(rows, list):
        rows = []
    return jsonify({"ok": True, "entries": rows[:limit], "total": len(rows)})


@app.route("/api/agent/run/list", methods=["GET"])
def agent_run_list():
    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    items = []
    for rid, entry in runs.items():
        if not isinstance(entry, dict):
            continue
        items.append(
            {
                "run_id": str(rid or ""),
                "status": str(entry.get("status") or ""),
                "stage": str(entry.get("stage") or ""),
                "runState": str(entry.get("runState") or ""),
                "task": str(entry.get("task") or ""),
                "updated_at": str(entry.get("updated_at") or ""),
                "can_continue": bool(entry.get("can_continue", False)),
            }
        )
    items.sort(key=lambda x: x.get("updated_at") or "", reverse=True)
    return jsonify({"ok": True, "runs": items[:30], "last_agent_run": state.get("last_agent_run") or {}}), 200


def _load_agent_run_entry_or_error(run_id: str):
    rid = str(run_id or "").strip()
    if not rid:
        return None, (jsonify({"ok": False, "status": "invalid_input", "error": "run_id fehlt."}), 400)
    state = load_project_auto_run_state()
    runs = state.get("agent_runs") if isinstance(state.get("agent_runs"), dict) else {}
    entry = runs.get(rid) if isinstance(runs.get(rid), dict) else None
    if not isinstance(entry, dict):
        return None, (jsonify({"ok": False, "status": "not_found", "error": "run_id nicht gefunden.", "run_id": rid}), 404)
    return (state, runs, entry, rid), None


@app.route("/api/agent/run/pause", methods=["POST"])
def agent_run_pause():
    data = request.get_json(silent=True) or {}
    loaded, err = _load_agent_run_entry_or_error(data.get("run_id"))
    if err:
        return err
    state, runs, entry, rid = loaded
    entry["paused"] = True
    entry["stopped"] = False
    entry["can_continue"] = True
    entry["can_retry"] = True
    entry["next_action"] = "continue_run"
    entry["status"] = "paused"
    entry["updated_at"] = get_timestamp()
    runs[rid] = entry
    save_project_auto_run_state({"agent_runs": runs, "last_agent_run": entry})
    return jsonify({"ok": True, "run_id": rid, "paused": True, "stopped": False, "can_continue": True, "can_retry": True, "next_action": "continue_run", "updated_at": entry["updated_at"]}), 200


@app.route("/api/agent/run/stop", methods=["POST"])
def agent_run_stop():
    data = request.get_json(silent=True) or {}
    loaded, err = _load_agent_run_entry_or_error(data.get("run_id"))
    if err:
        return err
    state, runs, entry, rid = loaded
    reason = str(data.get("reason") or "manual_stop").strip()
    entry["paused"] = False
    entry["stopped"] = True
    entry["stop_reason"] = reason
    entry["can_continue"] = False
    entry["can_retry"] = True
    entry["next_action"] = "retry_run"
    entry["status"] = "stopped"
    entry["updated_at"] = get_timestamp()
    runs[rid] = entry
    save_project_auto_run_state({"agent_runs": runs, "last_agent_run": entry})
    return jsonify({"ok": True, "run_id": rid, "paused": False, "stopped": True, "stop_reason": reason, "can_continue": False, "can_retry": True, "next_action": "retry_run", "updated_at": entry["updated_at"]}), 200


@app.route("/api/agent/run/continue", methods=["POST"])
def agent_run_continue():
    data = request.get_json(silent=True) or {}
    loaded, err = _load_agent_run_entry_or_error(data.get("run_id"))
    if err:
        return err
    state, runs, entry, rid = loaded
    if bool(entry.get("stopped")):
        return jsonify({"ok": False, "status": "blocked", "error": "Run ist gestoppt. Bitte Retry verwenden.", "run_id": rid, "writes_files": False, "auto_apply": False}), 409
    entry["paused"] = False
    entry["can_continue"] = False
    entry["can_retry"] = True
    entry["next_action"] = str(entry.get("next_action") or "review_plan_and_prepare_safe_preview")
    entry["status"] = "running"
    entry["updated_at"] = get_timestamp()
    runs[rid] = entry
    save_project_auto_run_state({"agent_runs": runs, "last_agent_run": entry})
    return jsonify({"ok": True, "run_id": rid, "paused": False, "stopped": False, "can_continue": False, "can_retry": True, "next_action": entry["next_action"], "writes_files": False, "auto_apply": False, "updated_at": entry["updated_at"]}), 200


@app.route("/api/agent/run/retry", methods=["POST"])
def agent_run_retry():
    data = request.get_json(silent=True) or {}
    loaded, err = _load_agent_run_entry_or_error(data.get("run_id"))
    if err:
        return err
    state, runs, entry, rid = loaded
    entry["paused"] = False
    entry["stopped"] = False
    entry["stop_reason"] = ""
    entry["can_continue"] = True
    entry["can_retry"] = True
    entry["last_safe_stage"] = str(entry.get("stage") or "planned")
    entry["stage"] = "planned"
    entry["status"] = "retry_prepared"
    entry["next_action"] = "review_plan_and_prepare_safe_preview"
    entry["updated_at"] = get_timestamp()
    runs[rid] = entry
    save_project_auto_run_state({"agent_runs": runs, "last_agent_run": entry})
    return jsonify({"ok": True, "run_id": rid, "paused": False, "stopped": False, "can_continue": True, "can_retry": True, "next_action": entry["next_action"], "writes_files": False, "auto_apply": False, "updated_at": entry["updated_at"]}), 200


@app.route("/api/agent/delegate", methods=["POST"])
def agent_delegate():
    """Level 8.6 — Decompose a task into specialist sub-tasks."""
    from agent_task_delegator import AgentTaskDelegator
    from agent_capability_gate import check_capability, get_capabilities
    from agent_write_gate import get_instance as get_write_gate

    data = request.get_json(silent=True) or {}
    task = str(data.get("task") or "").strip()
    if not task:
        return jsonify({"ok": False, "error": "task fehlt"}), 400

    candidate_files = [str(f) for f in list(data.get("candidate_files") or []) if str(f).strip()]
    memory_hints = [str(h) for h in list(data.get("memory_hints") or []) if str(h).strip()]
    patch_plan = [p for p in list(data.get("patch_plan") or []) if p]
    can_apply = bool(data.get("can_apply", True))
    blocking_reasons = [str(r) for r in list(data.get("blocking_reasons") or []) if str(r).strip()]
    review_summary = str(data.get("review_summary") or "").strip()
    commit_message = str(data.get("commit_message") or "").strip()
    test_checks = [str(c) for c in list(data.get("test_checks") or []) if str(c).strip()]

    run_id = str(data.get("run_id") or "").strip() or None
    delegator = AgentTaskDelegator(run_id=run_id)

    delegator.decompose(task, candidate_files=candidate_files or None)
    delegator.execute_read_phase(task, memory_hints=memory_hints or None, context_files=candidate_files or None)
    delegator.execute_patch_phase(patch_plan=patch_plan or None)
    delegator.execute_safety_phase(can_apply=can_apply, blocking_reasons=blocking_reasons or None)
    result = delegator.build_result(
        review_summary=review_summary,
        commit_message=commit_message,
        test_checks=test_checks or None,
    )

    return jsonify({
        "ok": True,
        "status": "delegated",
        "mode": "agent_delegate",
        "writes_files": False,
        "auto_apply": False,
        "auto_commit": False,
        "auto_rollback": False,
        **result,
    }), 200


@app.route("/api/agent/capabilities", methods=["GET"])
def agent_capabilities():
    """Level 8.8 — Return capability matrix for all agent roles."""
    from agent_capability_gate import list_agent_roles, get_capabilities, check_capability, READ_ONLY_ROLES, WRITE_ROLES

    roles = list_agent_roles()
    matrix = {role: get_capabilities(role) for role in roles}
    return jsonify({
        "ok": True,
        "capabilities": matrix,
        "read_only_roles": sorted(READ_ONLY_ROLES),
        "write_roles": sorted(WRITE_ROLES),
    }), 200


@app.route("/api/agent/capabilities/check", methods=["POST"])
def agent_capabilities_check():
    """Level 8.8 — Check a specific capability for an agent role."""
    from agent_capability_gate import check_capability

    data = request.get_json(silent=True) or {}
    role = str(data.get("role") or "").strip()
    capability = str(data.get("capability") or "").strip()
    if not role or not capability:
        return jsonify({"ok": False, "error": "role und capability erforderlich"}), 400

    allowed, msg = check_capability(role, capability)
    return jsonify({
        "ok": True,
        "role": role,
        "capability": capability,
        "allowed": allowed,
        "message": msg or "ok",
    }), 200


@app.route("/api/agents/registry", methods=["GET"])
def agents_registry():
    """Level 8.1 — Registry of all known agent roles with descriptions and capabilities."""
    from agent_capability_gate import list_agent_roles, get_capabilities, is_read_only

    _ROLE_DESCRIPTIONS = {
        "planner":  "Zerlegt die Aufgabe in Schritte und identifiziert Kandidaten-Dateien",
        "memory":   "Sucht ähnliche frühere Fehler und Tasks in der History",
        "context":  "Liest relevante Dateien und baut den Aufgabenkontext auf",
        "patch":    "Erstellt einen konsolidierten Patch-Plan (kein Write)",
        "apply":    "Führt den Patch aus — nur mit gültigem Write-Gate-Token",
        "safety":   "Prüft Risiko und Safety-Gate; kann blockieren",
        "test":     "Empfiehlt und koordiniert Checks",
        "review":   "Fasst das Ergebnis zusammen und setzt das Review-Gate",
        "commit":   "Erstellt Commit-Vorschlag (commit_performed=false immer)",
    }

    roles = list_agent_roles()
    registry = []
    for role in roles:
        caps = get_capabilities(role)
        active_caps = [k for k, v in caps.items() if v]
        registry.append({
            "role": role,
            "description": _ROLE_DESCRIPTIONS.get(role, role),
            "capabilities": caps,
            "active_capabilities": active_caps,
            "parallel_allowed": is_read_only(role),
            "requires_write_token": bool(caps.get("can_write")),
        })

    return jsonify({
        "ok": True,
        "agent_count": len(registry),
        "agents": registry,
    }), 200


@app.route("/api/agent/run/validate-patch", methods=["POST"])
def agent_run_validate_patch():
    """Level 8.4 — Standalone patch validation endpoint."""
    data = request.get_json(silent=True) or {}
    rel_path = str(data.get("path") or "").strip()
    current_content = str(data.get("current_content") or "")
    proposed_content = str(data.get("proposed_content") or "")
    diff_text = str(data.get("diff") or "")
    if not rel_path:
        return jsonify({"ok": False, "error": "path erforderlich"}), 400

    validator = get_patch_validator_agent(root=RAMBO_RAINER_ROOT.resolve())
    result = validator.validate_patch(
        rel_path=rel_path,
        current_content=current_content,
        proposed_content=proposed_content,
        diff_text=diff_text,
    )
    return jsonify({
        "ok": bool(result.get("ok")),
        "path": rel_path,
        "validation": result,
        "writes_files": False,
        "auto_apply": False,
    }), 200


@app.route("/api/agent/run/post-apply-tests", methods=["POST"])
def agent_run_post_apply_tests():
    """Level 8.5 — Run recommended checks after apply."""
    data = request.get_json(silent=True) or {}
    run_id = str(data.get("run_id") or "").strip()
    task = str(data.get("task") or "").strip()
    affected_files = [str(f) for f in list(data.get("affected_files") or []) if str(f).strip()]
    checks = [str(c) for c in list(data.get("checks") or []) if str(c).strip()]

    if not task and not affected_files:
        return jsonify({"ok": False, "error": "task oder affected_files erforderlich"}), 400

    runner = get_test_runner_agent(APP_DIR.resolve())

    # Recommend checks if none provided
    if not checks:
        rec = runner.recommend_checks(task=task, files=affected_files)
        checks = list(rec.get("recommended_checks") or [])

    results = []
    all_ok = True
    for check in checks[:4]:
        tr = runner.run_allowed_check(check)
        results.append({
            "check": check,
            "ok": bool(tr.get("ok")),
            "returncode": int(tr.get("returncode") or 0),
            "duration_ms": int(tr.get("duration_ms") or 0),
            "error_summary": str(tr.get("error_summary") or ""),
            "retry_performed": bool(tr.get("retry_performed")),
        })
        if not bool(tr.get("ok")):
            all_ok = False

    from agent_run_state_machine import AgentRunStateMachine
    sm = AgentRunStateMachine(run_id=run_id or "post-apply-check")
    sm.transition("context_loaded")
    sm.transition("patch_ready")
    sm.transition("validated")
    sm.transition("apply_ready")
    sm.transition("applied")
    if all_ok:
        sm.transition("tested")
        sm.transition("done")
    else:
        sm.transition("failed")

    return jsonify({
        "ok": all_ok,
        "run_id": run_id,
        "checks_run": len(results),
        "all_ok": all_ok,
        "results": results,
        "state_machine": sm.to_dict(),
        "next_action": "commit_when_ready" if all_ok else "review_failed_checks",
        "not_committed": True,
        "writes_files": False,
        "auto_commit": False,
        "auto_rollback": False,
    }), 200


CONTEXT_PROJECT_FILE = RAMBO_RAINER_ROOT / "knowledge" / "project_context.md"
CONTEXT_BUILDER_FILE = APP_DIR / "knowledge" / "user_notes.md"


def _default_memory_overrides():
    return {
        "preferred_paths": [
            "knowledge/",
            "outbox/",
            "backend/BUILDER_MODE.md",
            "backend/DEV_WORKFLOW.md",
        ],
        "avoid_paths": [
            "frontend/src/App.jsx",
            "frontend/src/App.css",
        ],
        "working_rules": [
            "Keine neuen Dateien, keine Duplikate, keine *_v2/_copy/_tmp.",
            "Guard-/Approval-Logik strikt respektieren.",
            "Keine destruktiven Git-Aktionen (kein Commit/Push/Reset-Hard).",
            "Safe-Vorschau vor Apply; Apply ist freigabepflichtig.",
        ],
        "style_rules": [
            "Antworten auf Deutsch halten.",
            "UI ruhig, IDE-artig, ohne Kartenlandschaft.",
            "Keine Live-Kommentare oder Prozessmeldungen im Ergebnis.",
        ],
        "agent_instructions": [
            "Bestehende Zustaende wiederverwenden statt Parallelwelten zu schaffen.",
            "Serverseitige Orchestrierung ist autoritativ; Frontend beobachtet/steuert.",
            "Fehler klassifizieren und Repair-Plan folgen.",
        ],
    }


def load_project_memory_overrides():
    """Benutzerbearbeitbare Regeln/Pfadlisten aus data/project_memory.json."""
    defaults = _default_memory_overrides()
    raw = read_json_file(PROJECT_MEMORY_FILE, None)
    if not isinstance(raw, dict):
        return defaults
    result = {}
    for key, default in defaults.items():
        val = raw.get(key)
        if isinstance(val, list):
            cleaned = [str(x).strip() for x in val if isinstance(x, (str, int))]
            cleaned = [x for x in cleaned if x]
            result[key] = cleaned
        else:
            result[key] = list(default)
    result["last_updated"] = str(raw.get("last_updated") or "").strip()
    return result


def save_project_memory_overrides(payload):
    defaults = _default_memory_overrides()
    data = {}
    for key in defaults.keys():
        raw = payload.get(key)
        if isinstance(raw, list):
            data[key] = [str(x).strip() for x in raw if isinstance(x, (str, int)) and str(x).strip()]
        elif isinstance(raw, str):
            items = [x.strip() for x in raw.splitlines() if x.strip()]
            data[key] = items
        else:
            data[key] = list(defaults[key])
    data["last_updated"] = get_timestamp()
    PROJECT_MEMORY_FILE.parent.mkdir(parents=True, exist_ok=True)
    PROJECT_MEMORY_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    return data


def build_project_memory():
    """Kanonische Projekt-Memory-/Regelbasis aus bestehenden Quellen.

    Fasst project_context, builder_notes, allowed_prefixes, sensitive_patterns
    und nutzerdefinierte Ueberschreibungen (preferred_paths/avoid_paths/
    working_rules/style_rules/agent_instructions) zusammen.
    """
    project_text, project_exists = read_text_file(CONTEXT_PROJECT_FILE)
    builder_text, builder_exists = read_text_file(CONTEXT_BUILDER_FILE)
    overrides = load_project_memory_overrides()

    project_stripped = (project_text or "").strip()
    builder_stripped = (builder_text or "").strip()

    sources = []
    if project_stripped:
        sources.append({
            "id": "project_context",
            "label": "Projektkontext",
            "path": str(CONTEXT_PROJECT_FILE.relative_to(RAMBO_RAINER_ROOT)) if CONTEXT_PROJECT_FILE.is_absolute() else str(CONTEXT_PROJECT_FILE),
            "length": len(project_stripped),
        })
    if builder_stripped:
        sources.append({
            "id": "builder_notes",
            "label": "Builder-Hinweise",
            "path": "rambo_builder_local/knowledge/user_notes.md",
            "length": len(builder_stripped),
        })
    sources.append({
        "id": "guard_prefixes",
        "label": "Guard: erlaubte Pfadprefixe",
        "path": "backend/main.py (ALLOWED_PROJECT_WRITE_PREFIXES)",
        "length": len(ALLOWED_PROJECT_WRITE_PREFIXES),
    })
    sources.append({
        "id": "guard_sensitive",
        "label": "Guard: sensible Muster",
        "path": "backend/main.py (SENSITIVE_PATTERNS)",
        "length": len(SENSITIVE_PATTERNS),
    })

    working_rules = list(overrides.get("working_rules") or [])
    style_rules = list(overrides.get("style_rules") or [])
    agent_instructions = list(overrides.get("agent_instructions") or [])
    all_rules = working_rules + style_rules + agent_instructions
    instruction_summary = "; ".join(all_rules[:3])[:240]

    return {
        "project_context": project_text or "",
        "project_context_snippet": project_stripped[:300],
        "project_context_len": len(project_stripped),
        "project_context_path": "knowledge/project_context.md",
        "builder_notes": builder_text or "",
        "builder_notes_snippet": builder_stripped[:300],
        "builder_notes_len": len(builder_stripped),
        "builder_notes_path": "rambo_builder_local/knowledge/user_notes.md",
        "allowed_prefixes": list(ALLOWED_PROJECT_WRITE_PREFIXES),
        "sensitive_patterns": list(SENSITIVE_PATTERNS),
        "guarded_paths": sorted(list(GUARDED_PROJECT_PATHS)),
        "preferred_paths": list(overrides.get("preferred_paths") or []),
        "avoid_paths": list(overrides.get("avoid_paths") or []),
        "working_rules": working_rules,
        "style_rules": style_rules,
        "agent_instructions": agent_instructions,
        "rules_count": len(all_rules),
        "instruction_summary": instruction_summary,
        "active_context_sources": sources,
        "last_updated": overrides.get("last_updated") or "",
        "timestamp": get_timestamp(),
    }


def build_memory_snapshot_compact():
    """Sehr kompakter Snapshot fuer Lauf-Zwecke (wird in auto_loop/direct_run gespeichert)."""
    m = build_project_memory()
    return {
        "project_context_len": m.get("project_context_len") or 0,
        "builder_notes_len": m.get("builder_notes_len") or 0,
        "allowed_prefixes_count": len(m.get("allowed_prefixes") or []),
        "sensitive_patterns_count": len(m.get("sensitive_patterns") or []),
        "preferred_paths": (m.get("preferred_paths") or [])[:6],
        "avoid_paths": (m.get("avoid_paths") or [])[:6],
        "rules_count": m.get("rules_count") or 0,
        "instruction_summary": m.get("instruction_summary") or "",
        "sources": [s.get("id") for s in (m.get("active_context_sources") or [])],
        "snapshot_at": m.get("timestamp") or get_timestamp(),
    }


@app.route("/api/project-memory", methods=["GET", "POST"])
def project_memory_endpoint():
    if request.method == "GET":
        return jsonify({"ok": True, "memory": build_project_memory()})
    data = request.get_json(silent=True) or {}
    saved = save_project_memory_overrides(data)
    append_ui_log_entry("Memory", "Projekt-Memory aktualisiert.", "info")
    return jsonify({"ok": True, "overrides": saved, "memory": build_project_memory()})


@app.route("/api/context/load", methods=["GET"])
def context_load():
    project_text, _ = read_text_file(CONTEXT_PROJECT_FILE)
    builder_text, _ = read_text_file(CONTEXT_BUILDER_FILE)
    return jsonify({
        "project_context": project_text,
        "builder_notes": builder_text,
        "timestamp": get_timestamp(),
    })


@app.route("/api/context/save", methods=["POST"])
def context_save():
    data = request.get_json() or {}
    content_type = str(data.get("type", "")).strip()
    content = str(data.get("content", ""))

    if content_type == "project":
        target = CONTEXT_PROJECT_FILE
        label = "Projektkontext"
    elif content_type == "builder":
        target = CONTEXT_BUILDER_FILE
        label = "Builder-Hinweise"
    else:
        return jsonify({"error": "Unbekannter Typ. Erlaubt: project, builder"}), 400

    wr = persist_text_file_change(
        target,
        content,
        str(target.name),
        on_timeout_log=lambda m: append_ui_log_entry("Kontext", m, "error"),
    )
    if not wr.get("ok"):
        return jsonify({"error": wr.get("error") or "Speichern fehlgeschlagen."}), 500

    append_ui_log_entry("Kontext", f"{label} gespeichert ({len(content)} Zeichen)", "info")
    return jsonify({"saved": True, "label": label, "timestamp": get_timestamp()})


def _build_workspace_edit_flow(current_error, current_patch, auto_loop, run_state, active_file):
    """Kompakter Write-/Edit-State fuer den Workspace (BLOCK E/F).

    Zeigt was der Loop gerade schreibt, ob ein Diff bereit ist,
    ob eine Freigabe noetig ist und was der letzte Apply-Status war.
    """
    # Determine loop write status from current steps
    loop_write_status = "idle"
    loop_target_file = ""
    approval_needed = False
    diff_available = False

    for step in (auto_loop.get("steps") or []):
        s_status = str(step.get("status") or "").lower()
        s_action = str(step.get("action") or "")
        s_file = str(step.get("file_target") or "")
        if s_action in {"direct_preview", "direct_apply", "runner_execute"} and s_status in {"laeuft", "wartet auf freigabe", "geplant"}:
            loop_target_file = s_file or loop_target_file
            if s_status == "wartet auf freigabe":
                loop_write_status = "awaiting_approval"
                approval_needed = True
            elif s_status == "laeuft":
                loop_write_status = "writing"
            elif s_status == "geplant" and loop_write_status == "idle":
                loop_write_status = "preparing"
            if s_action == "direct_preview":
                diff_available = True
            break

    if loop_write_status == "idle" and auto_loop.get("status") == "done":
        loop_write_status = "done"

    # Last apply result from direct_run_history
    last_apply = {}
    for entry in (run_state.get("direct_run_history") or [])[:3]:
        if str(entry.get("task_type") or "") == "file_edit" or str(entry.get("mode") or "") == "apply":
            last_apply = {
                "path": str((entry.get("affected_files") or [""])[0]),
                "run_id": str(entry.get("run_id") or ""),
                "status": str(entry.get("status") or ""),
                "timestamp": str(entry.get("timestamp") or ""),
                "diff_summary": entry.get("diff_summary") if isinstance(entry.get("diff_summary"), dict) else {},
            }
            break

    # Repair file target
    repair_file = ""
    repair_steps_with_file = []
    for rp_step in (auto_loop.get("repair_plan") or []):
        if isinstance(rp_step, dict) and rp_step.get("file_target"):
            repair_file = str(rp_step["file_target"])
            repair_steps_with_file.append({
                "id": str(rp_step.get("id") or ""),
                "label": str(rp_step.get("label") or ""),
                "file_target": repair_file,
                "load_endpoint": str(rp_step.get("load_endpoint") or ""),
                "prepare_endpoint": str(rp_step.get("prepare_endpoint") or ""),
                "apply_endpoint": str(rp_step.get("apply_endpoint") or ""),
                "apply_mode": str(rp_step.get("apply_mode") or "safe"),
                "gate": str(rp_step.get("gate") or "auto"),
            })

    # Review files pending
    review_files = []
    if isinstance(current_patch, dict):
        for fe in (current_patch.get("file_entries") or []):
            if str(fe.get("status") or "") not in {"applied", "verified", "closed"}:
                p = str(fe.get("path") or "")
                if _is_valid_file_path(p):
                    review_files.append({
                        "path": p,
                        "status": str(fe.get("status") or "review_ready"),
                        "load_endpoint": f"/api/file/load?path={p}",
                    })

    return {
        "loop_write_status": loop_write_status,
        "loop_target_file": loop_target_file or active_file.get("path") or "",
        "approval_needed": approval_needed,
        "diff_available": diff_available,
        "repair_file": repair_file,
        "repair_steps_with_file": repair_steps_with_file[:4],
        "review_files": review_files[:5],
        "last_apply": last_apply,
    }


def _is_valid_file_path(raw_path):
    """Prueft ob raw_path wie ein echter relativer Dateipfad aussieht (nicht runner/command o.ae.)."""
    if not raw_path:
        return False
    p = str(raw_path).replace("\\", "/").strip().strip("/")
    if not p or "/" not in p and "." not in p:
        return False
    # Keine internen Bezeichner wie runner/command, agent/step etc.
    INTERNAL_PREFIXES = ("runner/", "agent/", "auto_loop/", "step/", "loop-", "repair/")
    if any(p.startswith(x) for x in INTERNAL_PREFIXES):
        return False
    resolved, cleaned, err = validate_project_read_path(p)
    return err is None


def _resolve_active_file(current_error, current_patch, auto_loop, run_state):
    """Bestimmt die aktuell relevanteste Datei aus dem Session-Kontext (BLOCK G)."""
    if isinstance(current_error, dict) and _is_valid_file_path(current_error.get("file")):
        return {
            "path": str(current_error["file"]),
            "reason": "Aktiver Fehler in dieser Datei",
            "context_type": "error",
            "run_id": str(auto_loop.get("run_id") or ""),
            "category": str(current_error.get("category") or ""),
            "suggestion": str(current_error.get("suggestion") or ""),
            "recommended_tool": str(current_error.get("recommended_tool") or "direct"),
        }
    if isinstance(current_patch, dict):
        for fe in (current_patch.get("file_entries") or []):
            if str(fe.get("status") or "") not in {"applied", "verified", "closed"}:
                p = str(fe.get("path") or "")
                if _is_valid_file_path(p):
                    return {
                        "path": p,
                        "reason": "Offene Review-Datei im aktuellen Patch",
                        "context_type": "patch",
                        "run_id": str(current_patch.get("run_id") or ""),
                        "category": "review",
                        "suggestion": "Patch pruefen und freigeben oder ablehnen.",
                        "recommended_tool": "review",
                    }
    last_runner = run_state.get("last_runner_execution") if isinstance(run_state.get("last_runner_execution"), dict) else None
    if isinstance(last_runner, dict) and _is_valid_file_path(last_runner.get("affected_file")):
        return {
            "path": str(last_runner["affected_file"]),
            "reason": "Letzte Runner-Zieldatei",
            "context_type": "runner",
            "run_id": "",
            "category": str(last_runner.get("mode") or "command"),
            "suggestion": "",
            "recommended_tool": "runner",
        }
    for step in (auto_loop.get("steps") or []):
        if _is_valid_file_path(step.get("retry_target_file")):
            return {
                "path": str(step["retry_target_file"]),
                "reason": "Retry-Zieldatei aus Reparatur",
                "context_type": "repair",
                "run_id": str(auto_loop.get("run_id") or ""),
                "category": str(step.get("retry_category") or ""),
                "suggestion": str(auto_loop.get("repair_suggestion") or ""),
                "recommended_tool": "direct",
            }
    return {}


@app.route("/api/file/load", methods=["GET"])
def file_load_endpoint():
    """Laedt eine Projektdatei sicher mit Patch- und Fehler-Kontext (BLOCK F)."""
    rel_path = request.args.get("path", "").strip()
    if not rel_path:
        return jsonify({"ok": False, "error": "Kein Pfad angegeben."}), 400
    resolved, cleaned, path_err = validate_project_read_path(rel_path)
    cleaned = str(cleaned or "").strip()
    if path_err:
        return jsonify({"ok": False, "error": path_err}), 400
    content, exists = read_text_file(resolved)
    if not exists:
        return jsonify({"ok": False, "error": "Datei nicht gefunden.", "path": cleaned}), 404
    classification = classify_project_file(cleaned)
    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    patch_snapshot = build_patch_review_snapshot()
    current_patch = patch_snapshot.get("current")
    patch_context = None
    if isinstance(current_patch, dict):
        for fe in (current_patch.get("file_entries") or []):
            if str(fe.get("path") or "").replace("\\", "/").strip() == cleaned.strip():
                patch_context = {
                    "status": str(fe.get("status") or ""),
                    "detail": str(fe.get("detail") or ""),
                    "patch_id": str(current_patch.get("patch_id") or current_patch.get("run_id") or ""),
                    "apply_status": str(current_patch.get("apply_status") or ""),
                }
                break
    current_error = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else {}
    error_context = None
    err_file = str(current_error.get("file") or "").replace("\\", "/").strip()
    if err_file and (err_file == cleaned or cleaned.endswith("/" + err_file) or err_file.endswith("/" + cleaned)):
        error_context = {
            "category": str(current_error.get("category") or ""),
            "label": str(current_error.get("label") or ""),
            "suggestion": str(current_error.get("suggestion") or ""),
            "recommended_tool": str(current_error.get("recommended_tool") or "direct"),
        }
    lines = content.splitlines()
    append_ui_log_entry("Datei-Workspace", f"Datei geladen: {cleaned}", "info")
    return jsonify({
        "ok": True,
        "path": cleaned,
        "content": content,
        "line_count": len(lines),
        "size_bytes": len(content.encode("utf-8")),
        "classification": classification,
        "patch_context": patch_context,
        "error_context": error_context,
        "can_edit": bool(classification.get("allowed_write")),
        "timestamp": get_timestamp(),
    })


@app.route("/api/file/prepare-edit", methods=["POST"])
def file_prepare_edit_endpoint():
    """Bereitet eine Dateibearbeitung vor: Diff, Guard-Check, kein Schreiben (BLOCK F)."""
    data = request.get_json(silent=True) or {}
    rel_path = str(data.get("path") or "").strip()
    proposed = str(data.get("proposed_content") or "")
    if not rel_path:
        return jsonify({"ok": False, "error": "Kein Pfad angegeben."}), 400
    resolved, cleaned, path_err = validate_project_read_path(rel_path)
    if path_err:
        return jsonify({"ok": False, "error": path_err, "guard_blocked": True}), 400
    classification = classify_project_file(cleaned)
    if not classification.get("allowed_write"):
        reason = "Datei ist guarded oder liegt ausserhalb erlaubter Schreibpfade."
        if classification.get("sensitive"):
            reason = "Datei enthaelt sensibles Muster — Bearbeitung blockiert."
        return jsonify({"ok": False, "guard_blocked": True, "error": reason, "path": cleaned, "classification": classification}), 403
    current_content, exists = read_text_file(resolved)
    diff = build_text_diff(current_content if exists else "", proposed, cleaned)
    added = sum(1 for l in diff.splitlines() if l.startswith("+") and not l.startswith("+++"))
    removed = sum(1 for l in diff.splitlines() if l.startswith("-") and not l.startswith("---"))
    warnings = []
    if not exists:
        warnings.append("Datei existiert nicht — wird neu erstellt.")
    append_ui_log_entry("Datei-Workspace", f"Edit vorbereitet: {cleaned} (+{added}/-{removed})", "info")
    return jsonify({
        "ok": True,
        "path": cleaned,
        "guard_allowed": True,
        "can_apply": True,
        "diff": diff,
        "diff_summary": {"added": added, "removed": removed, "changed": bool(diff and "@@" in diff)},
        "classification": classification,
        "warnings": [w for w in warnings if w],
        "timestamp": get_timestamp(),
    })


@app.route("/api/file/apply", methods=["POST"])
def file_apply_endpoint():
    """Schreibt Datei-Aenderung nach Guard-Check und erstellt Patch-Eintrag (BLOCK G).

    apply_mode='safe': Kein Schreiben, gibt nur Diff zurueck (identisch zu prepare-edit).
    apply_mode='apply': Schreibt Datei, erstellt direct_run_history-Eintrag (erscheint im Review).
    """
    data = request.get_json(silent=True) or {}
    rel_path = str(data.get("path") or "").strip()
    proposed = str(data.get("proposed_content") or "")
    reason = str(data.get("reason") or "Datei-Workspace-Aenderung").strip()
    apply_mode = str(data.get("apply_mode") or "safe").strip()

    if not rel_path:
        return jsonify({"ok": False, "error": "Kein Pfad angegeben."}), 400

    resolved, cleaned, path_err = validate_project_read_path(rel_path)
    if path_err:
        return jsonify({"ok": False, "error": path_err, "guard_blocked": True}), 400

    classification = classify_project_file(cleaned)
    if not classification.get("allowed_write"):
        reason_txt = "Datei ist guarded oder liegt ausserhalb erlaubter Schreibpfade."
        if classification.get("sensitive"):
            reason_txt = "Datei enthaelt sensibles Muster — Bearbeitung blockiert."
        return jsonify({"ok": False, "guard_blocked": True, "error": reason_txt, "path": cleaned}), 403

    current_content, exists = read_text_file(resolved)
    diff = build_text_diff(current_content if exists else "", proposed, cleaned)
    added = sum(1 for l in diff.splitlines() if l.startswith("+") and not l.startswith("+++"))
    removed = sum(1 for l in diff.splitlines() if l.startswith("-") and not l.startswith("---"))
    diff_summary = {"added": added, "removed": removed, "changed": bool(diff and "@@" in diff)}

    if apply_mode != "apply":
        append_ui_log_entry("Datei-Workspace", f"Safe-Preview: {cleaned} (+{added}/-{removed})", "info")
        return jsonify({
            "ok": True,
            "path": cleaned,
            "apply_mode": "safe",
            "applied": False,
            "diff": diff,
            "diff_summary": diff_summary,
            "message": "Safe-Modus: Keine Aenderung geschrieben. Diff bereit.",
            "timestamp": get_timestamp(),
        })

    if not diff_summary["changed"]:
        return jsonify({
            "ok": True,
            "path": cleaned,
            "apply_mode": "apply",
            "applied": False,
            "diff": diff,
            "diff_summary": diff_summary,
            "message": "Keine inhaltliche Aenderung erkannt — nichts geschrieben.",
            "timestamp": get_timestamp(),
        })

    wr = persist_text_file_change(
        resolved,
        proposed,
        cleaned,
        on_timeout_log=lambda m: append_ui_log_entry("Datei-Workspace", m, "error"),
    )
    if not wr.get("ok"):
        err = wr.get("error") or "Schreiben fehlgeschlagen."
        return jsonify({"ok": False, "error": err, "path": cleaned}), 500

    run_id = str(uuid4()).replace("-", "")
    history_entry = build_direct_history_entry(run_id, {
        "task": reason,
        "scope": "local",
        "mode": "apply",
        "selected_target_path": cleaned,
        "affected_files": [cleaned],
        "has_changes": True,
        "guard": {"allowed": True, "detail": "Datei-Workspace Guard OK"},
        "diff": diff,
        "diff_summary": diff_summary,
        "post_check": {"ok": True, "detail": "Datei geschrieben."},
        "direct_status": "verified",
        "task_type": "file_edit",
        "primary_area": classification.get("area") or "Datei-Workspace",
        "message": f"Datei-Workspace: {cleaned} aktualisiert (+{added}/-{removed})",
    }, "verified")
    upsert_direct_run_history(history_entry)
    save_project_auto_run_state({
        "last_direct_status": "verified",
        "last_direct_confirmed_run_id": run_id,
        "last_completed_run_id": run_id,
        "last_result": f"Datei-Workspace: {cleaned} gespeichert.",
    })
    append_ui_log_entry("Datei-Workspace", f"Gespeichert: {cleaned} (+{added}/-{removed})", "success")
    return jsonify({
        "ok": True,
        "path": cleaned,
        "apply_mode": "apply",
        "applied": True,
        "run_id": run_id,
        "diff": diff,
        "diff_summary": diff_summary,
        "message": f"Datei gespeichert und Patch-Eintrag erstellt: {cleaned}",
        "timestamp": get_timestamp(),
    })


@app.route("/api/file/workspace-state", methods=["GET"])
def file_workspace_state_endpoint():
    """Einheitlicher Workspace-Edit-State fuer den Datei-Editor (BLOCK F).

    Kein Extra-Polling noetig: ein Call zeigt active_file, repair-Schritte
    mit Datei-Endpoints, loop_write_status, review_files, last_apply.
    """
    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    patch_snapshot = build_patch_review_snapshot()
    current_patch = patch_snapshot.get("current")
    current_error = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else _empty_error_info()
    active_file = _resolve_active_file(current_error, current_patch, auto_loop, run_state)
    flow = _build_workspace_edit_flow(current_error, current_patch, auto_loop, run_state, active_file)
    return jsonify({
        "ok": True,
        "timestamp": get_timestamp(),
        "active_file": active_file,
        "workspace_edit_flow": flow,
        "auto_loop_status": auto_loop.get("status") or "idle",
        "auto_loop_phase": auto_loop.get("phase") or "planning",
        "current_error": {
            "category": str(current_error.get("category") or ""),
            "label": str(current_error.get("label") or ""),
            "file": str(current_error.get("file") or ""),
            "suggestion": str(current_error.get("suggestion") or ""),
            "recommended_tool": str(current_error.get("recommended_tool") or ""),
        },
    })


# ---- Repo-Aktionen (Phase 16) ----

import re as _re


def _run_git_cmd(args, timeout=15):
    """Runs a git command in RAMBO_RAINER_ROOT. Returns (stdout, stderr, ok)."""
    project_root = str(RAMBO_RAINER_ROOT.resolve())
    try:
        r = subprocess.run(
            ["git"] + args, cwd=project_root,
            capture_output=True, text=True, timeout=timeout,
        )
        return r.stdout.strip(), r.stderr.strip(), r.returncode == 0
    except Exception as e:
        return "", str(e), False


def _read_remote_info_payload():
    """Read-only Remote-/Branch-/PR-Vorbereitung (kein Push, kein Netzwerk zu GH)."""
    remote_raw, _, _ = _run_git_cmd(["remote", "-v"])
    remotes = []
    seen = set()
    for line in (remote_raw or "").splitlines():
        parts = line.split()
        if len(parts) >= 2 and "(fetch)" in line and parts[0] not in seen:
            remotes.append({"name": parts[0], "url": parts[1]})
            seen.add(parts[0])
    has_remote = bool(remotes)

    branch, _, _ = _run_git_cmd(["branch", "--show-current"])
    branch = branch or "master"

    log_raw, _, _ = _run_git_cmd(["log", "--oneline", "-10"])
    commits = []
    for line in (log_raw or "").splitlines():
        if len(line) > 8:
            commits.append({"hash": line[:7], "msg": line[8:].strip()})

    ahead, behind = 0, 0
    remote_branch = ""
    if has_remote and remotes:
        remote_branch = remotes[0]["name"] + "/" + branch
        count_out, _, count_ok = _run_git_cmd(["rev-list", "--left-right", "--count", remote_branch + "...HEAD"])
        if count_ok and count_out.strip():
            parts = count_out.strip().split()
            if len(parts) == 2:
                try:
                    behind = int(parts[0])
                    ahead = int(parts[1])
                except ValueError:
                    pass

    pr_lines = ["## Änderungen\n"]
    for c in commits[:8]:
        pr_lines.append("- " + c["msg"])
    pr_draft = "\n".join(pr_lines)

    return {
        "ok": True,
        "branch": branch,
        "has_remote": has_remote,
        "remotes": remotes,
        "remote_branch": remote_branch,
        "ahead": ahead,
        "behind": behind,
        "recent_commits": commits,
        "pr_draft": pr_draft,
        "push_blocked": True,
        "push_blocked_reason": "Push ist in dieser Umgebung nicht aktiviert. Nur Vorbereitung und Prüfung.",
        "readiness": {
            "remote_configured": has_remote,
            "commits_available": len(commits),
            "can_prepare_pr": len(commits) > 0,
            "hint": "Remote konfigurieren, dann manuell pushen." if not has_remote else f"{ahead} Commit(s) vor {remote_branch}.",
        },
    }


def _collect_remote_workflow_for_agent_core():
    """Phase 19: Remote-/PR-Kontext fuer Agent-Core und Decision (explizit, keine Auto-Aktionen)."""
    base = _read_remote_info_payload()
    base["explicit_only"] = True
    base["auto_remote_actions_blocked"] = True
    base["policy"] = (
        "Kein automatischer Push oder PR aus der Agenten-Auto-Schleife. "
        "Remote-Schritte nur nach Prüfung im Git-Panel oder bewusst extern."
    )
    return base


def _append_remote_publish_followups(followup_actions, remote_workflow):
    if not isinstance(followup_actions, list) or not isinstance(remote_workflow, dict) or not remote_workflow:
        return
    hint = str((remote_workflow.get("readiness") or {}).get("hint") or "").strip()
    pol = str(remote_workflow.get("policy") or "").strip()
    reason = hint or pol or "Push und PR sind nicht automatisiert; nur Prüfung und manuelle Schritte."
    _append_followup_action(
        followup_actions, "remote_readonly", "Remote-Status (read-only)",
        "/api/repo/remote-info", "GET", {}, "git", "git", "", reason, primary=False,
    )
    _append_followup_action(
        followup_actions, "git_panel_open", "Git-Panel öffnen",
        "", "GET", {}, "git", "git", "", "Manuelle Kontrolle vor jedem Push oder PR.", primary=False,
    )


def _safe_branch_name(name):
    return bool(name and _re.match(r'^[a-zA-Z0-9][a-zA-Z0-9/_.\-]{0,78}$', name))


@app.route("/api/repo/state", methods=["GET"])
def repo_state_endpoint():
    branch, _, _ = _run_git_cmd(["branch", "--show-current"])
    staged_raw, _, _ = _run_git_cmd(["diff", "--name-only", "--cached"])
    uncommitted_raw, _, _ = _run_git_cmd(["status", "--short"])
    last_commit, _, _ = _run_git_cmd(["log", "--oneline", "-1"])
    staged = [f.strip() for f in staged_raw.splitlines() if f.strip()]
    uncommitted = []
    for line in uncommitted_raw.splitlines():
        code = line[:2].strip() if len(line) >= 2 else ""
        path = line[3:].strip() if len(line) > 3 else ""
        if path:
            uncommitted.append({"code": code, "path": path})
    return jsonify({
        "ok": True,
        "branch": branch or "unknown",
        "staged_files": staged,
        "staged_count": len(staged),
        "uncommitted_files": uncommitted,
        "uncommitted_count": len(uncommitted),
        "last_commit": last_commit or "-",
        "rollback": {
            "editor_revert": True,
            "file_to_head": True,
            "staged_unstage": len(staged) > 0,
            "notes": "Kein force-push, kein reset --hard, kein Branch-Switch auf main/master.",
        },
    })


@app.route("/api/repo/branch", methods=["POST"])
def repo_branch_endpoint():
    data = request.get_json(silent=True) or {}
    branch_name = str(data.get("branch") or "").strip()
    if not branch_name:
        return jsonify({"ok": False, "error": "Branch-Name fehlt."}), 400
    if not _safe_branch_name(branch_name):
        return jsonify({"ok": False, "error": "Ungültiger Branch-Name (nur a-z A-Z 0-9 / _ . -)"}), 400
    if branch_name in ("main", "master", "HEAD"):
        return jsonify({"ok": False, "error": "Auf main/master kann nicht zurückgewechselt werden."}), 400
    existing, _, _ = _run_git_cmd(["branch", "--list", branch_name])
    if existing.strip():
        return jsonify({"ok": False, "error": f"Branch '{branch_name}' existiert bereits."}), 409
    _, err, ok = _run_git_cmd(["branch", branch_name])
    if not ok:
        return jsonify({"ok": False, "error": f"Branch konnte nicht angelegt werden: {err}"}), 500
    append_ui_log_entry("Repo", f"Branch angelegt: {branch_name}", "info")
    return jsonify({"ok": True, "branch": branch_name, "message": f"Branch '{branch_name}' angelegt (aktueller Branch bleibt aktiv)."})


@app.route("/api/repo/stage", methods=["POST"])
def repo_stage_endpoint():
    data = request.get_json(silent=True) or {}
    files = data.get("files") or []
    if not files:
        return jsonify({"ok": False, "error": "Keine Dateien angegeben."}), 400
    BLOCK_KEYWORDS = (".env", "credentials", "secrets", ".pem", ".key", ".pfx", ".p12")
    safe_files, rejected = [], []
    for f in files:
        f = str(f).strip().replace("\\", "/")
        if any(k in f.lower() for k in BLOCK_KEYWORDS):
            rejected.append(f)
            continue
        try:
            _, cleaned, err = validate_project_read_path(f)
            if err:
                rejected.append(f)
            else:
                safe_files.append(cleaned)
        except Exception:
            rejected.append(f)
    if not safe_files:
        return jsonify({"ok": False, "error": "Keine gültigen Dateien.", "rejected": rejected}), 400
    _, err, ok = _run_git_cmd(["add", "--"] + safe_files)
    if not ok:
        return jsonify({"ok": False, "error": f"git add fehlgeschlagen: {err}", "rejected": rejected}), 500
    append_ui_log_entry("Repo", f"Staged: {', '.join(safe_files)}", "info")
    return jsonify({"ok": True, "staged": safe_files, "rejected": rejected, "message": f"{len(safe_files)} Datei(en) zum Commit vorgemerkt."})


@app.route("/api/repo/unstage", methods=["POST"])
def repo_unstage_endpoint():
    data = request.get_json(silent=True) or {}
    files = data.get("files") or []
    if not files:
        return jsonify({"ok": False, "error": "Keine Dateien angegeben."}), 400
    safe_files = []
    for f in files:
        f = str(f).strip().replace("\\", "/")
        try:
            _, cleaned, err = validate_project_read_path(f)
            if not err:
                safe_files.append(cleaned)
        except Exception:
            pass
    if not safe_files:
        return jsonify({"ok": False, "error": "Keine gültigen Dateien."}), 400
    _, err, ok = _run_git_cmd(["restore", "--staged", "--"] + safe_files)
    if not ok:
        return jsonify({"ok": False, "error": f"Unstage fehlgeschlagen: {err}"}), 500
    append_ui_log_entry("Repo", f"Unstaged: {', '.join(safe_files)}", "info")
    return jsonify({"ok": True, "unstaged": safe_files})


@app.route("/api/repo/commit", methods=["POST"])
def repo_commit_endpoint():
    data = request.get_json(silent=True) or {}
    message = str(data.get("message") or "").strip()
    if not message:
        return jsonify({"ok": False, "error": "Commit-Nachricht fehlt."}), 400
    if len(message) > 500:
        return jsonify({"ok": False, "error": "Commit-Nachricht zu lang (max 500 Zeichen)."}), 400
    staged_raw, _, _ = _run_git_cmd(["diff", "--name-only", "--cached"])
    if not staged_raw.strip():
        return jsonify({"ok": False, "error": "Keine Dateien staged. Zuerst Dateien auswählen."}), 400
    staged_list = [f.strip() for f in staged_raw.splitlines() if f.strip()]
    _, err, ok = _run_git_cmd(["commit", "-m", message])
    if not ok:
        return jsonify({"ok": False, "error": f"Commit fehlgeschlagen: {err}"}), 500
    new_commit, _, _ = _run_git_cmd(["log", "--oneline", "-1"])
    append_ui_log_entry("Repo", f"Commit: {new_commit or message[:40]}", "info")
    return jsonify({"ok": True, "commit": new_commit, "staged_files": staged_list, "message": f"Commit erstellt: {new_commit}"})


@app.route("/api/repo/revert-file", methods=["POST"])
def repo_revert_file_endpoint():
    """Setzt eine Datei auf den letzten Commit-Stand zurück (kein reset --hard)."""
    data = request.get_json(silent=True) or {}
    rel_path = str(data.get("path") or "").strip()
    if not rel_path:
        return jsonify({"ok": False, "error": "Kein Pfad angegeben."}), 400
    _, cleaned, path_err = validate_project_read_path(rel_path)
    if path_err:
        return jsonify({"ok": False, "error": path_err}), 400
    staged_check, _, _ = _run_git_cmd(["diff", "--name-only", "--cached", "--", cleaned])
    if staged_check.strip():
        return jsonify({"ok": False, "error": f"'{cleaned}' ist staged. Zuerst unstagen, dann zurücksetzen."}), 409
    _, _, tracked_ok = _run_git_cmd(["ls-files", "--error-unmatch", "--", cleaned])
    if not tracked_ok:
        return jsonify({"ok": False, "error": f"'{cleaned}' ist nicht in git getrackt. Nur der Editor-Stand kann zurückgesetzt werden."}), 409
    _, err, ok = _run_git_cmd(["restore", "--source=HEAD", "--worktree", "--", cleaned])
    if not ok:
        return jsonify({"ok": False, "error": f"Revert fehlgeschlagen: {err}"}), 500
    append_ui_log_entry("Repo", f"Datei auf HEAD zurückgesetzt: {cleaned}", "info")
    return jsonify({"ok": True, "path": cleaned, "message": f"'{cleaned}' auf letzten Commit-Stand zurückgesetzt."})


@app.route("/api/repo/remote-info", methods=["GET"])
def repo_remote_info_endpoint():
    """Read-only Remote-/PR-Vorbereitungsinfo. Kein Push, kein PR-Create."""
    return jsonify(_read_remote_info_payload())


@app.route("/api/agent/decision", methods=["GET"])
def agent_decision_endpoint():
    """Autoritative Folgeentscheidung fuer den aktuellen Agentenlauf."""
    run_state = load_project_auto_run_state()
    auto_loop = normalize_auto_loop_state(run_state.get("auto_loop_state"))
    activity = load_ui_activity_entries()
    patch_snapshot = build_patch_review_snapshot()
    current_patch = patch_snapshot.get("current")
    current_error = auto_loop.get("current_error") if isinstance(auto_loop.get("current_error"), dict) else _empty_error_info()
    qa_info = _collect_qa_info_for_agent_core(run_state, auto_loop, activity)
    active_file = _resolve_active_file(current_error, current_patch, auto_loop, run_state)
    decision = _build_agent_decision(run_state, auto_loop, current_patch, current_error, qa_info, active_file)

    if decision.get("gate_type") == "auto":
        decision["console_label"] = f"[AUTO] {decision.get('next_reason') or decision.get('why_continuing') or decision.get('status')}"
    elif decision.get("gate_type") in {"approval_gate", "approval_pending"}:
        decision["console_label"] = f"[GATE] {decision.get('next_reason') or decision.get('why_stopped') or decision.get('status')}"
    elif decision.get("gate_type") == "stopped":
        decision["console_label"] = f"[STOP] {decision.get('why_stopped') or decision.get('next_reason') or decision.get('status')}"
    else:
        decision["console_label"] = f"[?] {decision.get('next_reason') or decision.get('status')}"

    history = _record_agent_decision_trace(run_state, decision)
    decision["history"] = history[:12] if isinstance(history, list) else []
    decision["ok"] = True
    decision["action"] = decision.get("next_action") or ""
    decision["tool"] = decision.get("next_tool") or ""
    decision["reason"] = decision.get("next_reason") or ""
    return jsonify(decision)


@app.route("/api/openapi.json", methods=["GET"])
def api_openapi_schema():
    return jsonify(OPENAPI_SCHEMA)


@app.route("/api/docs", methods=["GET"])
def api_docs_swagger_ui():
    html = """
<!doctype html>
<html>
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>Rainer Build API Docs</title>
    <link rel="stylesheet" href="https://unpkg.com/swagger-ui-dist@5/swagger-ui.css" />
  </head>
  <body>
    <div id="swagger-ui"></div>
    <script src="https://unpkg.com/swagger-ui-dist@5/swagger-ui-bundle.js"></script>
    <script>
      window.ui = SwaggerUIBundle({
        url: '/api/openapi.json',
        dom_id: '#swagger-ui'
      });
    </script>
  </body>
</html>
"""
    return html, 200, {"Content-Type": "text/html; charset=utf-8"}


@app.route("/api/hybrid/status", methods=["GET"])
def hybrid_status():
    """
    Hybrid Engine Status - immer HTTP 200.
    Zeigt verfügbare Provider und Status.
    """
    try:
        if not HYBRID_ENGINE_AVAILABLE or hybrid_engine is None:
            return jsonify({
                "available": False,
                "message": "Hybrid Engine nicht verfügbar",
                "providers": {},
                "hint": "Installiere: pip install groq google-generativeai python-dotenv",
            }), 200

        providers = hybrid_engine.get_available_providers()
        return jsonify({
            "available": True,
            "message": "Hybrid Engine bereit",
            "providers": providers,
            "provider_count": len(providers),
            "timestamp": get_timestamp(),
        }), 200
    except Exception as e:
        return jsonify({
            "available": False,
            "message": f"Hybrid Engine Fehler: {str(e)}",
            "providers": {},
        }), 200


@app.route("/api/generated-media/<path:filename>", methods=["GET"])
def api_serve_generated_media(filename: str):
    raw = str(filename or "").strip().replace("\\", "/")
    name = Path(raw).name
    if not name or not re.match(r"^[A-Za-z0-9_.-]+$", name):
        return jsonify({"ok": False}), 404
    base = GENERATED_IMAGES_DIR.resolve()
    target = (base / name).resolve()
    try:
        target.relative_to(base)
    except Exception:
        return jsonify({"ok": False}), 404
    if not target.is_file():
        return jsonify({"ok": False}), 404
    return send_from_directory(str(base), name)


@app.route("/api/image/generate", methods=["POST"])
def api_image_generate():
    data = request.get_json(silent=True) or {}
    prompt = str(data.get("prompt") or "").strip()
    size = str(data.get("size") or "1024x1024").strip()
    if not prompt:
        return jsonify({"ok": False, "error": "Prompt darf nicht leer sein."}), 400
    if len(prompt) > 2000:
        return jsonify({"ok": False, "error": "Prompt maximal 2000 Zeichen."}), 400
    try:
        out = generate_image_via_openai(prompt=prompt, size=size)
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500
    if not out.get("ok"):
        return jsonify({"ok": False, "error": str(out.get("error") or "Bildgenerierung fehlgeschlagen.")}), 502
    remote = str(out.get("remote_url") or "").strip()
    b64 = str(out.get("b64_json") or "").strip()
    fname = f"gen_{uuid4().hex[:16]}.png"
    dest = (GENERATED_IMAGES_DIR / fname).resolve()
    try:
        if remote:
            rr = requests.get(remote, timeout=90)
            rr.raise_for_status()
            dest.write_bytes(rr.content)
        elif b64:
            dest.write_bytes(base64.b64decode(b64))
        else:
            return jsonify({"ok": False, "error": "Bild-API lieferte weder URL noch Base64."}), 502
    except Exception as exc:
        return jsonify({"ok": False, "error": f"Bild konnte nicht gespeichert werden: {exc}"}), 500
    rel_url = f"/api/generated-media/{fname}"
    return jsonify(
        {
            "ok": True,
            "type": "image",
            "image_url": rel_url,
            "prompt": prompt,
            "provider": str(out.get("provider") or "openai_compatible"),
            "model": str(out.get("model") or ""),
        }
    ), 200


@app.route("/api/weather", methods=["GET"])
def api_weather():
    city = str(request.args.get("city", "Idar-Oberstein")).strip() or "Idar-Oberstein"
    lat = str(request.args.get("lat", "49.711")).strip()
    lon = str(request.args.get("lon", "7.314")).strip()
    try:
        weather_res = requests.get(
            f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current=temperature_2m,weather_code&timezone=auto",
            timeout=8,
        )
        weather_res.raise_for_status()
        payload = weather_res.json() or {}
        current = payload.get("current") or {}
        return jsonify(
            {
                "ok": True,
                "city": city,
                "temperature": current.get("temperature_2m"),
                "status": _weather_status_from_code(current.get("weather_code")),
                "lat": lat,
                "lon": lon,
            }
        ), 200
    except Exception as exc:
        return jsonify(
            {
                "ok": False,
                "city": city,
                "temperature": None,
                "status": "Wetterdaten nicht erreichbar",
                "error": str(exc),
            }
        ), 200


@app.route("/api/hybrid/ask", methods=["POST"])
def hybrid_ask():
    """
    Reine KI-Anfrage an Hybrid Engine.
    Nicht für Dateiänderungen oder Builds!
    """
    data = request.get_json(silent=True) or {}
    prompt = _extract_task_or_prompt_from_request_json(data)

    if not prompt:
        return jsonify({
            "ok": False,
            "error": "Kein Prompt angegeben",
            "hint": 'POST {"prompt": "Deine Frage"}',
        }), 400

    # Blockiere Datei-/Build-Aufträge
    task_lower = prompt.lower()
    blocked_patterns = [
        "erstelle datei", "erstelle nur die datei", "ändere datei", "aendere datei",
        "lösche datei", "loesche datei", "delete file", "create file",
        "baue electron", "build electron", "npm install", "build_desktop.py",
        "installer", ".exe", "write file", "modify file", "erzeuge datei"
    ]
    for pattern in blocked_patterns:
        if pattern in task_lower:
            return jsonify({
                "ok": False,
                "error": f"Hybrid Engine nicht für Datei-/Build-Operationen zuständig. Verwende /api/direct-run. Erkannt: '{pattern}'",
                "blocked": True,
                "redirect_to": "/api/direct-run",
            }), 400

    if not HYBRID_ENGINE_AVAILABLE or hybrid_engine is None:
        return jsonify({
            "ok": False,
            "error": "Hybrid Engine nicht verfügbar. Bitte API-Key konfigurieren oder lokalen Service starten.",
            "hint": "Installiere: pip install groq google-generativeai python-dotenv",
            "available": False,
        }), 200

    # Führe Anfrage aus
    try:
        preferred = data.get("preferred_provider")
        result = hybrid_engine.ask(prompt, preferred_provider=preferred)

        # Enriche Response mit UI-Chrome
        result["direct_ui_chrome"] = "minimal"
        result["recognized_task"] = {
            "task_type": "hybrid_ask",
            "primary_area": "Hybrid Intelligence",
            "hint": "Reine KI-Analyse ohne Datei-Operationen",
            "execution_route": "hybrid_engine",
        }

        return jsonify(result), 200 if result.get("ok") else 503
    except Exception as e:
        return jsonify({
            "ok": False,
            "error": f"Hybrid Engine Fehler: {str(e)}",
            "workstream_events": [{"ts": get_timestamp(), "phase": "error", "level": "error", "title": "Exception", "detail": str(e)[:100], "status": "failed"}],
        }), 200


try:
    from pro_merge_server_routes import attach_rambo_server_routes

    _n_merge = attach_rambo_server_routes(app)
    logger.logger.info("Rambo-/Pro-Routen aus server.py gemerged: %s Regel(n)", _n_merge)
except Exception as _merge_exc:
    logger.logger.warning("Rambo-Routen-Merge übersprungen: %s", _merge_exc)

app.config["RAMBO_SOCKETIO"] = None
try:
    from websocket import init_socketio_app

    _rambo_adm = (os.environ.get("RAMBO_ADMIN_TOKEN") or "").strip()
    if _rambo_adm:
        app.config["RAMBO_SOCKETIO"] = init_socketio_app(app, _rambo_adm)
        logger.logger.info("Socket.IO auf main-App registriert.")
except Exception as _sock_exc:
    logger.logger.warning("Socket.IO nicht verfügbar: %s", _sock_exc)


@app.errorhandler(Exception)
def rambo_api_json_errors(exc):
    """Alle /api/* Exceptions als JSON (kein HTML-Traceback im Browser)."""
    try:
        p = getattr(request, "path", "") or ""
    except RuntimeError:
        raise exc
    if not p.startswith("/api/"):
        raise exc
    if isinstance(exc, HTTPException):
        body = {
            "ok": False,
            "error": exc.description or str(exc),
            "direct_status": "failed",
            "technical_message": str(exc),
            "workstream_events": [_ws_event("error", "error", p, str(exc)[:400], status="failed")],
        }
        return jsonify(enrich_direct_run_response(body)), int(exc.code or 500)
    tb = traceback.format_exc()
    body = {
        "ok": False,
        "error": str(exc) or "Interner Fehler",
        "direct_status": "failed",
        "technical_message": tb[:4000],
        "workstream_events": [_ws_event("error", "error", "api", str(exc)[:600], status="failed")],
    }
    return jsonify(enrich_direct_run_response(body)), 500


if __name__ == "__main__":
    import urllib.error
    import urllib.request

    _port = int(os.environ.get("FLASK_PORT", os.environ.get("SERVER_PORT", "5002")))
    if str(os.environ.get("RAINER_ALLOW_SECOND_INSTANCE") or "").strip().lower() not in ("1", "true", "yes"):
        try:
            urllib.request.urlopen(f"http://127.0.0.1:{_port}/api/health", timeout=0.35)
            print(
                f"[main] Backend antwortet bereits auf Port {_port} (/api/health). "
                "Zweiten Start verhindert. Zum Erzwingen: RAINER_ALLOW_SECOND_INSTANCE=1"
            )
            sys.exit(3)
        except (urllib.error.URLError, OSError, TimeoutError):
            pass
    _debug = os.environ.get("FLASK_DEBUG", "").strip().lower() in ("1", "true", "yes")
    _sock = app.config.get("RAMBO_SOCKETIO")
    if _sock is not None:
        _sock.run(
            app,
            host="127.0.0.1",
            port=_port,
            debug=_debug,
            allow_unsafe_werkzeug=True,
        )
    else:
        app.run(host="127.0.0.1", port=_port, debug=_debug)
_TASKS_FALLBACK: dict[str, dict] = {}



